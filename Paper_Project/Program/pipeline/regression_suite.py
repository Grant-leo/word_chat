"""
regression_suite.py - synthetic regression checks for the Word pipeline.

The suite creates temporary, non-private DOCX/MD fixtures and verifies the
engine behavior that is hard to cover by manual inspection alone:

- inline math stays in the current paragraph
- display math remains native OMML display math
- mixed text/image/math paragraphs do not drop tokens
- markdown tables/code/math keep structure
- build_manifest.json drives body element QA counts
- template profiles avoid private source filenames
- visual QA fails closed when required render tools are unavailable
"""
from __future__ import annotations

import argparse
import base64
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from types import SimpleNamespace
from typing import Any, Callable, Dict, List

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from lxml import etree

PIPELINE_DIR = Path(__file__).resolve().parent
if str(PIPELINE_DIR) not in sys.path:
    sys.path.insert(0, str(PIPELINE_DIR))

from content_parser import extract as extract_docx_content
from content_parser_modules.caption_flow import (
    is_figure_caption,
    normalize_caption_spacing,
    pair_figure_blocks,
)
from content_parser_modules.body_dispatcher import append_text_or_code, parse_body_sections
from content_parser_modules.formula_extractor import _strip_trailing_formula_labels_from_xml
from content_parser_modules.image_extractor import ImageRegistry
from content_parser_modules.paragraph_stream import append_stream_run_group
from content_parser_modules.reference_collector import ReferenceCollector, is_reference_heading
from content_parser_modules.section_builder import (
    filter_content_sections,
    make_body_section,
    mark_first_body_page_break,
    postprocess_section_paragraphs,
)
from content_parser_modules.text_cleaner import (
    clean_code_text,
    clean_text_artifacts,
)
from formula_semantics import (
    CATEGORY_CONTAMINATED,
    CATEGORY_DISPLAY_MATH,
    CATEGORY_QUANTITY_TEXT,
    classify_formula_text,
    is_formula_problem_text,
    looks_like_formula_text as semantic_looks_like_formula_text,
    split_inline_math_spans,
)
from format_extractor import extract as extract_docx_format
from latex_omath import latex_to_omath
from md_parser import extract_content as extract_md_content
from pipeline_runner.artifacts import build_content_markdown, write_content_artifacts, write_format_artifacts
from pipeline_runner.build_phase import generate_and_build_docx_phase
from pipeline_runner.cli import build_arg_parser, dispatch_cli
from pipeline_runner.contracts import (
    has_contract_errors,
    validate_build_manifest,
    validate_content_data,
    validate_format_data,
    validate_qa_report,
)
from pipeline_runner.context import (
    create_unique_output_dir,
    normalize_qa_level,
    resolve_inputs,
    write_workflow_mode,
)
from pipeline_runner.dependencies import load_optional_dependencies
from pipeline_runner.execution import ScriptExecutionResult, run_generated_script
from pipeline_runner.io import normalize_mode, scan_inputs
from pipeline_runner.qa import QADependencies, run_qa_phases
from pipeline_runner.summary import build_completion_summary
from pipeline_runner.template_phase import write_template_profile_phase, write_template_requirements_phase
from pipeline_runner.verification import VerificationError, _merge_content_results, double_verify
from qa_conformance import check_conformance
from qa_checker import check_output, write_reports
from script_generator import generate
from script_generator import (
    RUNTIME_TEMPLATE,
    _extract_page_and_header,
    _front_matter_sections,
    _infer_style_profiles,
    _infer_template_rules,
    _normalize_numbered_section_order,
)
from template_profiler import profile_format
from privacy import sanitize_value


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAFgwJ/lc0ndwAAAABJRU5ErkJggg=="
)

CASES: List[Callable[[], None]] = []
TEMP_DIRS: List[Path] = []
KEEP_ARTIFACTS = False


def case(fn: Callable[[], None]) -> Callable[[], None]:
    CASES.append(fn)
    return fn


def fail(message: str) -> None:
    raise AssertionError(message)


def assert_true(condition: bool, message: str) -> None:
    if not condition:
        fail(message)


def new_workdir(name: str) -> Path:
    work = Path(tempfile.mkdtemp(prefix=f"wordchat_{name}_"))
    TEMP_DIRS.append(work)
    return work


def write_json(path: Path, value: Dict[str, Any]) -> None:
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")


def write_sample_png(path: Path, width: int = 480, height: int = 270) -> None:
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((8, 8, width - 8, height - 8), outline=(30, 80, 160), width=3)
    draw.line((24, height - 40, width - 24, 40), fill=(180, 50, 50), width=4)
    image.save(path)


def base_format(source: str = "synthetic.md") -> Dict[str, Any]:
    return {
        "_meta": {
            "source": source,
            "sha256": "synthetic",
            "paragraphs": 1,
            "tables": 0,
            "sections": 1,
        },
        "paragraphs": [
            {
                "text": "Synthetic body paragraph used for style inference.",
                "runs": [{"font": "Times New Roman", "size_pt": 12}],
                "align": "JUSTIFY",
            }
        ],
        "tables": [],
        "sections": [
            {
                "page_width_cm": 21.0,
                "page_height_cm": 29.7,
                "margin_top_cm": 2.54,
                "margin_bottom_cm": 2.54,
                "margin_left_cm": 3.17,
                "margin_right_cm": 3.17,
            }
        ],
        "cover": [],
        "style_profiles": {},
    }


def base_content(paragraphs: List[Any], meta_tables: int = 0) -> Dict[str, Any]:
    return {
        "_meta": {
            "source": "synthetic.docx",
            "sha256": "synthetic",
            "paragraphs": max(1, len(paragraphs)),
            "tables_count": meta_tables,
            "images_extracted": 0,
        },
        "title_info": {"title_cn": "Synthetic Thesis"},
        "sections": [
            {
                "heading": "1 Introduction",
                "level": 1,
                "role": "body",
                "paragraphs": paragraphs,
                "images": [],
            }
        ],
        "references": ["[1] Synthetic reference."],
    }


@case
def pipeline_contracts_accept_current_handoffs() -> None:
    manifest = {
        "schema_version": 1,
        "counts": {
            "content_images_rendered": 0,
            "content_tables_rendered": 1,
            "content_formulas_rendered": 2,
        },
    }
    qa_report = {
        "schema_version": 1,
        "mode": "developer",
        "passed": True,
        "counts": {},
        "issues": [],
        "next_action": "ok",
    }
    all_issues = (
        validate_format_data(base_format())
        + validate_content_data(base_content(["Body text"]))
        + validate_build_manifest(manifest)
        + validate_qa_report(qa_report)
    )
    assert_true(not has_contract_errors(all_issues), "valid current handoff structures should satisfy contracts")


@case
def pipeline_contracts_report_structural_errors() -> None:
    issues = (
        validate_format_data({"paragraphs": {}})
        + validate_content_data({"sections": {}})
        + validate_build_manifest({"counts": {"content_images_rendered": -1}})
        + validate_qa_report({"passed": "yes", "counts": [], "issues": [{"code": ""}]})
    )
    codes = {issue.code for issue in issues}
    assert_true("FORMAT_SECTIONS_MISSING" in codes, "missing format sections was not reported")
    assert_true("FORMAT_PARAGRAPHS_NOT_LIST" in codes, "invalid format paragraphs was not reported")
    assert_true("CONTENT_SECTIONS_NOT_LIST" in codes, "invalid content sections was not reported")
    assert_true("MANIFEST_COUNT_INVALID" in codes, "invalid manifest count was not reported")
    assert_true("QA_REPORT_PASSED_NOT_BOOL" in codes, "invalid QA passed flag was not reported")
    assert_true("QA_REPORT_ISSUE_CODE_MISSING" in codes, "invalid QA issue code was not reported")
    assert_true(has_contract_errors(issues), "structural contract errors should be marked as errors")


@case
def pipeline_io_helpers_scan_inputs_and_normalize_modes() -> None:
    work = new_workdir("pipeline_io")
    (work / "paper.docx").write_text("x", encoding="utf-8")
    (work / "notes.md").write_text("x", encoding="utf-8")
    (work / "~$paper.docx").write_text("x", encoding="utf-8")
    (work / "ignore.txt").write_text("x", encoding="utf-8")
    assert_true(scan_inputs(str(work)) == ["notes.md", "paper.docx"], "scan_inputs should skip temp and unsupported files")
    assert_true(scan_inputs(str(work), exts=(".docx",)) == ["paper.docx"], "scan_inputs should respect extension filters")
    assert_true(normalize_mode("developer") == "developer", "developer mode should be preserved")
    assert_true(normalize_mode("bad-mode") == "user", "unknown mode should fall back to user")


@case
def pipeline_cli_dispatches_md_parameter_and_single_file_modes() -> None:
    parser = build_arg_parser()
    calls: List[Dict[str, Any]] = []

    def fake_run(template_file, content_file, **kwargs):
        calls.append({"template": template_file, "content": content_file, **kwargs})
        return "ok"

    def expect_exit_zero(args, template_dir="", inputs_dir=""):
        try:
            dispatch_cli(args, run_pipeline=fake_run, template_dir=template_dir, inputs_dir=inputs_dir)
        except SystemExit as exc:
            assert_true(exc.code == 0, f"dispatch returned nonzero exit: {exc.code}")
            return
        fail("dispatch should exit through exit_from_result")

    expect_exit_zero(parser.parse_args(["--md", "paper.md", "--mode", "developer", "--qa-level", "basic", "--no-qa"]))
    assert_true(calls[-1]["template"] is None and calls[-1]["content"] is None, "MD mode should not require template/content")
    assert_true(calls[-1]["md_file"] == "paper.md", "MD mode did not pass md_file")
    assert_true(calls[-1]["mode"] == "developer" and calls[-1]["run_qa"] is False, "MD mode options were not preserved")

    expect_exit_zero(parser.parse_args(["--template", "t.docx", "--content", "c.docx"]))
    assert_true(calls[-1]["template"] == "t.docx" and calls[-1]["content"] == "c.docx", "parameter mode files changed")
    assert_true(calls[-1]["mode"] == "user", "non-interactive auto mode should default to user")

    work = new_workdir("pipeline_cli")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    (template_dir / "only.docx").write_bytes(b"")
    (inputs_dir / "only.md").write_text("# Paper", encoding="utf-8")
    expect_exit_zero(parser.parse_args(["--mode", "user"]), template_dir=str(template_dir), inputs_dir=str(inputs_dir))
    assert_true(calls[-1]["template"] == "only.docx" and calls[-1]["content"] == "only.md", "single-file interactive dispatch changed")


@case
def pipeline_context_resolves_inputs_outputs_and_workflow() -> None:
    work = new_workdir("pipeline_context")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    outputs_dir = work / "Outputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    outputs_dir.mkdir()
    (template_dir / "template.docx").write_text("x", encoding="utf-8")
    (inputs_dir / "paper.md").write_text("# Paper", encoding="utf-8")

    resolution = resolve_inputs("template.docx", "paper.md", None, str(template_dir), str(inputs_dir))
    assert_true(resolution.ok, f"expected inputs to resolve: {resolution.error}")
    assert_true(resolution.inputs.use_md_content is True, "markdown content flag was not set")
    assert_true(resolution.inputs.content_name == "paper", "content stem was not preserved")
    assert_true(normalize_qa_level("bad-level") == "strict", "invalid QA level should fall back to strict")

    first_dir, first_name = create_unique_output_dir(str(outputs_dir), "paper", today="2026-05-26")
    second_dir, second_name = create_unique_output_dir(str(outputs_dir), "paper", today="2026-05-26")
    assert_true(first_name == "2026-05-26_paper", "first output folder name changed")
    assert_true(second_name == "2026-05-26_paper_2", "duplicate output folder suffix changed")
    workflow_path = write_workflow_mode(
        first_dir,
        mode="developer",
        template_path=resolution.inputs.template_path,
        content_path=resolution.inputs.content_path,
        run_qa=True,
        qa_level="visual",
        golden_dir="Golden",
        update_golden=False,
        require_wps=True,
    )
    workflow = json.loads(Path(workflow_path).read_text(encoding="utf-8"))
    assert_true(workflow["mode"] == "developer", "workflow mode was not written")
    assert_true(workflow["qa_level"] == "visual", "workflow QA level was not written")
    assert_true(workflow["require_wps"] is True, "workflow require_wps flag was not written")


@case
def pipeline_dependencies_loads_optional_modules_and_reports_missing() -> None:
    def marker(name):
        return lambda *args, **kwargs: name

    modules = {
        "qa_checker": SimpleNamespace(check_and_write=marker("qa")),
        "qa_conformance": SimpleNamespace(check_and_write=marker("conformance"), write_requirements=marker("requirements")),
        "template_profiler": SimpleNamespace(write_profile=marker("profile")),
        "md_parser": SimpleNamespace(extract_format=marker("md_format"), extract_content=marker("md_content")),
    }

    def fake_import_module(name):
        if name == "qa_visual":
            raise ImportError("visual unavailable")
        return modules[name]

    deps = load_optional_dependencies(import_module=fake_import_module)
    assert_true(deps.qa_check_and_write() == "qa", "qa checker dependency was not loaded")
    assert_true(deps.conformance_check_and_write() == "conformance", "conformance dependency was not loaded")
    assert_true(deps.write_template_requirements() == "requirements", "template requirements dependency was not loaded")
    assert_true(deps.write_template_profile() == "profile", "template profiler dependency was not loaded")
    assert_true(deps.extract_md_format() == "md_format" and deps.extract_md_content() == "md_content", "md parser dependencies were not loaded")
    assert_true(deps.visual_check_and_write is None, "missing visual dependency should be None")
    assert_true("visual unavailable" in deps.optional_import_detail("qa_visual"), "missing dependency detail was not preserved")
    assert_true(deps.optional_import_detail("qa_checker") == "", "available dependency should not report an error detail")


@case
def pipeline_artifacts_write_format_and_content_handoffs() -> None:
    work = new_workdir("pipeline_artifacts")
    fmt_json_path, fmt_md_path = write_format_artifacts(base_format(), "# Format", str(work))
    assert_true(Path(fmt_json_path).exists(), "format.json was not written")
    assert_true(Path(fmt_md_path).read_text(encoding="utf-8") == "# Format", "format markdown changed")

    content = base_content(
        [
            {"text": "A paragraph with math", "math": [{"text": "E=mc^2"}]},
            "Plain paragraph",
        ]
    )
    content["sections"][0]["images"] = ["fig1.png"]
    cnt_json_path, cnt_md_path = write_content_artifacts(content, str(work), str(work / "paper.docx"))
    summary = Path(cnt_md_path).read_text(encoding="utf-8")
    assert_true(Path(cnt_json_path).exists(), "content.json was not written")
    assert_true("# 内容提取" in summary, "content report title missing")
    assert_true("- [图片] fig1.png" in summary, "content report image line missing")
    assert_true("(+1公式)" in summary, "content report math count missing")
    assert_true("## 参考文献" in build_content_markdown(content, str(work / "paper.docx")), "references section missing")


@case
def pipeline_execution_runs_generated_script_with_utf8_output() -> None:
    work = new_workdir("pipeline_execution")
    script = work / "build_generated.py"
    script.write_text("print('生成完成')\n", encoding="utf-8")
    result = run_generated_script(str(script), str(work), python_executable=sys.executable)
    assert_true(result.returncode == 0, f"generated script returned {result.returncode}: {result.stderr}")
    assert_true("生成完成" in result.stdout, "UTF-8 stdout was not decoded")


@case
def pipeline_build_phase_generates_and_blocks_on_failure() -> None:
    work = new_workdir("pipeline_build_phase")
    steps: List[str] = []
    calls: List[str] = []

    def fake_step(label):
        steps.append(label)

    def fake_generate(fmt_json_path, cnt_json_path, out_dir, output_docx_name):
        calls.append(f"generate:{Path(out_dir).name}:{output_docx_name}")
        Path(out_dir, "build_generated.py").write_text("print('ok')\n", encoding="utf-8")
        return 11

    def fake_run(gen_py_path, out_dir, python_executable):
        calls.append(f"run:{Path(gen_py_path).name}:{python_executable}")
        return ScriptExecutionResult(returncode=0, stdout="built\n", stderr="")

    ok = generate_and_build_docx_phase(
        "format.json",
        "content.json",
        str(work),
        "folder",
        output_docx_name="out.docx",
        generate_script=fake_generate,
        run_generated_script=fake_run,
        python_executable="python",
        step=fake_step,
    )
    assert_true(ok is True, "build phase should pass when generated script succeeds")
    assert_true(
        steps[:2]
        == [
            "Phase 4/6: 生成构建脚本",
            "Phase 5/6: 构建最终 docx（生成静态目录；可用 Word COM 时写入页码）",
        ],
        f"unexpected build steps: {steps}",
    )
    assert_true(calls == [f"generate:{work.name}:out.docx", "run:build_generated.py:python"], f"unexpected build calls: {calls}")

    def failing_run(gen_py_path, out_dir, python_executable):
        return ScriptExecutionResult(returncode=1, stdout="", stderr="boom")

    failed = generate_and_build_docx_phase(
        "format.json",
        "content.json",
        str(work),
        "folder",
        output_docx_name="out.docx",
        generate_script=fake_generate,
        run_generated_script=failing_run,
        python_executable="python",
        step=fake_step,
    )
    assert_true(failed is False, "build phase should block when generated script fails")


@case
def pipeline_qa_runs_strict_and_blocks_failures() -> None:
    work = new_workdir("pipeline_qa")
    calls: List[str] = []

    def passing_qa(out_dir, mode, output_docx_name):
        calls.append("qa")
        return {"passed": True, "issues": [], "counts": {}, "mode": mode}

    def passing_conformance(out_dir, mode, output_docx_name, project_root):
        calls.append("conformance")
        return {"passed": True, "issues": []}

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=passing_conformance,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    assert_true(
        run_qa_phases(str(work), mode="developer", output_docx_name="最终论文.docx", qa_level="strict", project_root=str(work), deps=deps),
        "strict QA should pass when structural and conformance QA pass",
    )
    assert_true(calls == ["qa", "conformance"], f"unexpected QA call order: {calls}")

    def failing_qa(out_dir, mode, output_docx_name):
        return {
            "passed": False,
            "counts": {},
            "issues": [{"severity": "error", "code": "TEST_ERROR", "message": "Synthetic failure", "active_owner": "developer"}],
            "repair_plan": {"steps": []},
        }

    failing_deps = QADependencies(
        qa_check_and_write=failing_qa,
        conformance_check_and_write=passing_conformance,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    assert_true(
        not run_qa_phases(str(work), mode="developer", output_docx_name="最终论文.docx", qa_level="basic", project_root=str(work), deps=failing_deps),
        "QA should block the pipeline when the required QA report fails",
    )


@case
def pipeline_summary_mentions_outputs_and_mode() -> None:
    summary = build_completion_summary("2026-05-27_demo", "最终论文.docx", "developer")
    assert_true("Outputs/2026-05-27_demo/" in summary, "output directory missing from completion summary")
    assert_true("最终论文.docx" in summary, "output docx missing from completion summary")
    assert_true("当前模式: 开发者" in summary, "developer mode missing from completion summary")
    assert_true("build_generated.py" in summary, "user fine-tuning target missing from completion summary")


@case
def pipeline_verification_double_verify_arbitrates_format_mismatch() -> None:
    calls: List[int] = []

    def flaky_format_extractor(path):
        calls.append(len(calls) + 1)
        paragraph_count = 1 if len(calls) != 2 else 2
        fmt = {
            "_meta": {"paragraphs": paragraph_count},
            "paragraphs": [{"runs": []} for _ in range(paragraph_count)],
            "tables": [],
            "sections": [],
        }
        return fmt, "# Format"

    result, report = double_verify(flaky_format_extractor, "synthetic.docx", "Format")
    assert_true(len(calls) == 3, "format mismatch should trigger a third arbitration run")
    assert_true(len(result["paragraphs"]) == 1, "third run should arbitrate back to the majority shape")
    assert_true(report == "# Format", "format markdown payload should be preserved")


@case
def pipeline_verification_arbitrates_content_mismatch() -> None:
    work = new_workdir("pipeline_verify_content_majority")
    calls: List[int] = []

    def flaky_content_extractor(path, output_dir=None):
        calls.append(len(calls) + 1)
        references = [{"text": "Synthetic reference"}] if len(calls) == 2 else []
        return {
            "_meta": {"images_extracted": 0, "image_extract_failures": [], "non_body_images": []},
            "sections": [{"heading": "Body", "paragraphs": ["Text"]}],
            "references": references,
        }

    result = double_verify(flaky_content_extractor, "synthetic.docx", "Content", output_dir=str(work))
    assert_true(len(calls) == 3, "content mismatch should trigger a third arbitration run")
    assert_true(result["references"] == [], "third run should arbitrate content back to the majority shape")


@case
def pipeline_verification_converges_incremental_content_and_materializes_images() -> None:
    work = new_workdir("pipeline_verify_content_converge")
    calls: List[int] = []

    def incremental_content_extractor(path, output_dir=None):
        calls.append(len(calls) + 1)
        run_no = len(calls)
        base = Path(path).stem
        fig_dir = Path(output_dir) / base / "figures"
        fig_dir.mkdir(parents=True, exist_ok=True)

        image_names = [f"img{i}.png" for i in range(1, run_no + 1)]
        for image_name in image_names:
            (fig_dir / image_name).write_bytes(PNG_1X1)

        references = [{"text": f"[{i}] Ref {i}"} for i in range(1, run_no + 1)]
        return {
            "_meta": {
                "images_extracted": len(image_names),
                "images_dir": str(fig_dir),
                "image_extract_failures": [{"target": "transient", "error": "miss"}] if run_no == 1 else [],
                "non_body_images": [],
            },
            "sections": [
                {
                    "heading": "Body",
                    "level": 1,
                    "paragraphs": [{"role": "image", "image": image_names[-1]}],
                    "images": [image_names[-1]],
                }
            ],
            "references": references,
        }

    result = double_verify(incremental_content_extractor, "synthetic.docx", "Content", output_dir=str(work))
    assert_true(len(calls) == 5, "content convergence should keep collecting unstable recoverable content up to the cap")
    assert_true(result["_meta"].get("converged_extraction"), "convergence metadata was not recorded")
    assert_true(len(result["references"]) == 5, f"references were not merged: {result['references']}")
    assert_true(result["sections"][0]["images"] == ["img1.png", "img2.png", "img3.png", "img4.png", "img5.png"], "section images were not unioned")

    paragraph_images = sorted(
        item.get("image")
        for item in result["sections"][0]["paragraphs"]
        if isinstance(item, dict) and item.get("image")
    )
    assert_true(paragraph_images == ["img1.png", "img2.png", "img3.png", "img4.png", "img5.png"], f"paragraph images were not merged: {paragraph_images}")
    final_fig_dir = work / "synthetic" / "figures"
    assert_true((final_fig_dir / "img1.png").exists() and (final_fig_dir / "img5.png").exists(), "converged images were not materialized")
    assert_true(not (work / "_extract_verify_runs").exists(), "isolated verification directories were not cleaned")
    assert_true(result["_meta"].get("image_extract_failures") == [], "transient image failures should not survive convergence")


@case
def pipeline_verification_inserts_recovered_items_near_anchors() -> None:
    def content_with(paragraphs, images=None):
        return {
            "_meta": {
                "images_extracted": len(images or []),
                "image_extract_failures": [],
                "non_body_images": [],
            },
            "sections": [
                {
                    "heading": "Body",
                    "level": 1,
                    "role": "body",
                    "paragraphs": paragraphs,
                    "images": images or [],
                }
            ],
            "references": [],
        }

    merged, _reason = _merge_content_results(
        [
            content_with(["before", {"role": "image", "image": "img.png"}, "middle", "after"], ["img.png"]),
            content_with(["before", {"role": "formula", "source": "latex", "latex": "E=mc^2", "text": "E=mc^2"}, "middle", "after"]),
            content_with(["before", {"role": "table", "table_rows": [["A"], ["1"]]}, "middle", "after"]),
        ]
    )
    paragraphs = merged["sections"][0]["paragraphs"]
    middle_index = paragraphs.index("middle")
    formula_index = next(i for i, item in enumerate(paragraphs) if isinstance(item, dict) and item.get("role") == "formula")
    table_index = next(i for i, item in enumerate(paragraphs) if isinstance(item, dict) and item.get("table_rows"))
    assert_true(formula_index < middle_index, f"recovered formula drifted to section tail: {paragraphs}")
    assert_true(table_index < middle_index, f"recovered table drifted to section tail: {paragraphs}")


@case
def pipeline_verification_fails_when_no_majority_signature() -> None:
    calls: List[int] = []

    def unstable_content_extractor(path):
        calls.append(len(calls) + 1)
        return {
            "_meta": {"images_extracted": len(calls), "image_extract_failures": [], "non_body_images": []},
            "sections": [{"heading": f"Body {idx}", "paragraphs": ["Text"]} for idx in range(len(calls))],
            "references": [],
        }

    try:
        double_verify(unstable_content_extractor, "synthetic.docx", "Content")
    except VerificationError as exc:
        assert_true("verification failed" in str(exc), "failure should explain that verification failed")
        assert_true(len(calls) == 3, "unresolved mismatch should stop after the third run")
        return
    fail("unresolved extraction mismatch should raise VerificationError")


@case
def run_pipeline_completion_step_is_named() -> None:
    root = PIPELINE_DIR.parents[2]
    text = (root / "run_pipeline.py").read_text(encoding="utf-8")
    assert_true("step('完成')" in text or 'step("完成")' in text, "completion step should have a user-visible label")
    assert_true("step('??')" not in text and 'step("??")' not in text, "placeholder completion step leaked into run_pipeline.py")


@case
def pipeline_template_phase_writes_profile_and_requirements() -> None:
    work = new_workdir("pipeline_template_phase")
    calls: List[str] = []

    def fake_profile(fmt, out_dir, project_root):
        calls.append(f"profile:{project_root}")
        return {
            "capabilities": {
                "has_cover": True,
                "has_heading_styles": False,
                "has_caption_styles": True,
            },
            "risk_flags": {"mixed_headers": True, "low_sample": False},
        }

    profile = write_template_profile_phase(
        base_format(),
        str(work),
        project_root="ROOT",
        write_template_profile=fake_profile,
    )
    assert_true(profile["capabilities"]["has_cover"] is True, "profile return value was not preserved")

    def fake_requirements(fmt, content, out_dir):
        calls.append("requirements")
        Path(out_dir, "template_requirements.json").write_text("{}", encoding="utf-8")

    ok = write_template_requirements_phase(
        base_format(),
        base_content(["Body"]),
        str(work),
        write_template_requirements=fake_requirements,
        optional_import_detail=lambda name: " detail",
    )
    assert_true(ok is True, "requirements phase should report success")
    assert_true(calls == ["profile:ROOT", "requirements"], f"unexpected template phase calls: {calls}")
    assert_true((work / "template_requirements.json").exists(), "requirements writer was not called")

    skipped = write_template_requirements_phase(
        base_format(),
        base_content(["Body"]),
        str(work),
        write_template_requirements=None,
        optional_import_detail=lambda name: " detail",
    )
    assert_true(skipped is False, "missing requirements writer should report skipped")


def run_generated_case(name: str, content: Dict[str, Any], fmt: Dict[str, Any] | None = None) -> Dict[str, Any]:
    work = new_workdir(name)
    fmt_path = work / "format.json"
    cnt_path = work / "content.json"
    out_docx = work / "out.docx"
    write_json(fmt_path, fmt or base_format())
    write_json(cnt_path, content)
    write_json(work / "workflow_mode.json", {"mode": "developer"})

    generate(str(fmt_path), str(cnt_path), str(work), "out.docx")
    build_py = work / "build_generated.py"
    subprocess.run([sys.executable, "-m", "py_compile", str(build_py)], check=True, timeout=120)
    result = subprocess.run(
        [sys.executable, str(build_py)],
        cwd=str(work),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=180,
    )
    if result.returncode != 0:
        fail(f"{name}: generated build failed: {result.stderr[:1000] or result.stdout[:1000]}")
    assert_true(out_docx.exists(), f"{name}: output docx was not created")

    with zipfile.ZipFile(out_docx) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    manifest = json.loads((work / "build_manifest.json").read_text(encoding="utf-8"))
    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    return {"work": work, "xml": xml, "manifest": manifest, "report": report}


def omath_count(xml: str) -> int:
    return len(re.findall(r"<(?:[A-Za-z_][\w.-]*:)?oMath\b", xml))


def omath_para_count(xml: str) -> int:
    return len(re.findall(r"<(?:[A-Za-z_][\w.-]*:)?oMathPara\b", xml))


def make_vml_picture_docx(src_docx: Path, dst_docx: Path) -> None:
    tmp = Path(tempfile.mkdtemp(prefix="wordchat_vml_zip_"))
    try:
        with zipfile.ZipFile(src_docx) as zf:
            zf.extractall(tmp)
        document_xml = tmp / "word" / "document.xml"
        xml = document_xml.read_text(encoding="utf-8")
        m = re.search(r'r:embed="([^"]+)"', xml)
        if not m:
            fail("source docx did not contain a drawing relationship")
        rid = m.group(1)
        xml = re.sub(
            r"<w:drawing>.*?</w:drawing>",
            f'<w:pict><v:shape><v:imagedata r:id="{rid}"/></v:shape></w:pict>',
            xml,
            count=1,
            flags=re.S,
        )
        document_xml.write_text(xml, encoding="utf-8")
        with zipfile.ZipFile(dst_docx, "w", zipfile.ZIP_DEFLATED) as zf:
            for path in tmp.rglob("*"):
                if path.is_file():
                    zf.write(path, path.relative_to(tmp).as_posix())
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


@case
def inline_rich_text_stays_in_paragraph() -> None:
    content = base_content(
        [
            {
                "role": "rich_text",
                "text": "before x2 after",
                "runs": [
                    {"type": "text", "text": "before "},
                    {
                        "type": "math",
                        "text": "x2",
                        "math": [{"type": "inline", "latex": "x^2", "text": "x2"}],
                    },
                    {"type": "text", "text": " after"},
                ],
                "math": [{"type": "inline", "latex": "x^2", "text": "x2"}],
            }
        ]
    )
    result = run_generated_case("inline_rich", content)
    counts = result["manifest"]["counts"]
    xml = result["xml"]
    assert_true(counts["inline_formulas_rendered"] == 1, "inline formula was not counted")
    assert_true(counts["display_formulas_rendered"] == 0, "inline formula became display math")
    assert_true(omath_count(xml) == 1, "expected one native oMath")
    assert_true(omath_para_count(xml) == 0, "inline case should not create oMathPara")
    before_idx = xml.find("before ")
    math_idx = xml.find("<m:oMath", before_idx)
    after_idx = xml.find(" after", math_idx)
    assert_true(before_idx >= 0 and before_idx < math_idx < after_idx, "inline math order drifted")
    assert_true(not result["report"]["issues"], f"unexpected QA issues: {result['report']['issues']}")


@case
def legacy_text_math_item_is_inline() -> None:
    content = base_content(
        [
            {
                "text": "legacy inline math",
                "math": [{"type": "inline", "latex": "y^2", "text": "y2"}],
            }
        ]
    )
    result = run_generated_case("legacy_inline", content)
    counts = result["manifest"]["counts"]
    assert_true(counts["inline_formulas_rendered"] == 1, "legacy inline math was not rendered inline")
    assert_true(counts["display_formulas_rendered"] == 0, "legacy inline math became display math")
    assert_true(not result["report"]["issues"], f"unexpected QA issues: {result['report']['issues']}")


@case
def display_formula_remains_display() -> None:
    content = base_content([
        {"role": "formula", "latex": "a=b+c", "text": "a=b+c", "numbered": False}
    ])
    result = run_generated_case("display_formula", content)
    counts = result["manifest"]["counts"]
    assert_true(counts["display_formulas_rendered"] == 1, "display formula was not counted")
    assert_true(omath_para_count(result["xml"]) == 1, "display formula did not create oMathPara")
    assert_true(not result["report"]["issues"], f"unexpected QA issues: {result['report']['issues']}")


@case
def latex_delimited_text_formula_with_number_renders_native() -> None:
    content = base_content([
        {"role": "formula", "source": "text", "text": r"$$E_{total}=\sum_{t=1}^{24}P(t)\Delta t$$ (1.1)"}
    ])
    result = run_generated_case("latex_delimited_numbered_text_formula", content)
    assert_true(result["manifest"]["counts"]["display_formulas_rendered"] == 1, "numbered LaTeX text formula was not rendered")
    assert_true(omath_para_count(result["xml"]) == 1, "numbered LaTeX text formula did not create display OMML")
    assert_true("[LaTeX error" not in result["xml"] and "$$" not in result["xml"], "LaTeX text leaked into final XML")
    assert_true(not result["report"]["issues"], f"unexpected QA issues: {result['report']['issues']}")


@case
def latex_delimited_appendix_formula_with_letter_number_renders_native() -> None:
    content = base_content([
        {
            "role": "formula",
            "source": "text",
            "text": r"$$R_{green}=\frac{E_{renew}}{E_{total}}\times100\%$$" + " \uff08A.1\uff09",
        }
    ])
    result = run_generated_case("latex_delimited_appendix_formula", content)
    assert_true(result["manifest"]["counts"]["display_formulas_rendered"] == 1, "appendix-numbered LaTeX text formula was not rendered")
    assert_true(omath_para_count(result["xml"]) == 1, "appendix-numbered formula did not create display OMML")
    assert_true("[LaTeX error" not in result["xml"] and "$$" not in result["xml"], "appendix formula leaked LaTeX/error text")
    assert_true(not result["report"]["issues"], f"unexpected QA issues: {result['report']['issues']}")


@case
def source_formula_label_cleanup_preserves_arguments() -> None:
    for expr in ["f(1)", "x^{(1)}", "P=f(1)"]:
        _xml, text, had_label = _strip_trailing_formula_labels_from_xml(latex_to_omath(expr, display=True))
        assert_true(not had_label, f"formula argument was mistaken for an equation label: {expr} -> {text}")
    _xml, text, had_label = _strip_trailing_formula_labels_from_xml(latex_to_omath("E=mc^2(1.1)", display=True))
    assert_true(had_label and "(1.1)" not in text, f"equation label was not stripped: {text}")


@case
def multiple_display_omml_entries_render_all() -> None:
    content = base_content([
        {
            "role": "formula",
            "source": "omml",
            "text": "",
            "math": [
                {"type": "display", "xml": latex_to_omath("x=1", display=True), "text": "x=1"},
                {"type": "display", "xml": latex_to_omath("y=2", display=True), "text": "y=2"},
            ],
            "numbered": False,
        }
    ])
    result = run_generated_case("multi_display_omml", content)
    counts = result["manifest"]["counts"]
    assert_true(counts["display_formulas_rendered"] == 2, f"multiple OMML formulas were not all rendered: {counts}")
    assert_true(omath_para_count(result["xml"]) == 2, "expected two display OMML paragraphs")
    assert_true(not result["report"]["issues"], f"unexpected QA issues: {result['report']['issues']}")


@case
def table_manifest_matches_structured_body_tables() -> None:
    content = base_content(
        [
            {"role": "table", "table_rows": [["A", "B"], ["1", "2"]]},
        ],
        meta_tables=99,
    )
    result = run_generated_case("table_manifest", content)
    assert_true(result["manifest"]["counts"]["content_tables_rendered"] == 1, "body table was not counted")
    assert_true(result["report"]["counts"]["content_tables"] == 1, "QA used raw doc.tables instead of structured tables")
    assert_true(not result["report"]["issues"], f"unexpected QA issues: {result['report']['issues']}")


@case
def qa_manifest_detects_missing_table_render() -> None:
    content = base_content([{"role": "table", "table_rows": [["A"], ["B"]]}])
    result = run_generated_case("qa_missing_table", content)
    manifest_path = result["work"] / "build_manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["counts"]["content_tables_rendered"] = 0
    write_json(manifest_path, manifest)
    report = check_output(str(result["work"]), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("TABLE_COUNT_MISMATCH" in codes, "QA did not trust manifest for table mismatch")


@case
def code_table_is_not_body_table() -> None:
    content = base_content(
        [
            {
                "role": "code",
                "code": "interface GigabitEthernet0/0/1\nip address 10.0.0.1 255.255.255.0",
                "table_rows": [["interface GigabitEthernet0/0/1\nip address 10.0.0.1 255.255.255.0"]],
            }
        ],
        meta_tables=1,
    )
    result = run_generated_case("code_table", content)
    assert_true(result["manifest"]["counts"]["content_tables_rendered"] == 0, "code box was counted as body table")
    assert_true(result["report"]["counts"]["content_tables"] == 0, "QA counted code table_rows as body table")
    assert_true(not result["report"]["issues"], f"unexpected QA issues: {result['report']['issues']}")


@case
def image_manifest_matches_rendered_body_images() -> None:
    img_src = new_workdir("image_src")
    write_sample_png(img_src / "dot.png")
    content = base_content([
        {"role": "figure", "image": "dot.png", "caption": "Figure 1 sample"}
    ])
    content["_meta"]["images_dir"] = str(img_src)
    content["_meta"]["images_extracted"] = 1
    content["sections"][0]["images"] = ["dot.png"]
    result = run_generated_case("image_manifest", content)
    assert_true(result["manifest"]["counts"]["content_images_rendered"] == 1, "body image was not counted")
    assert_true(result["report"]["counts"]["content_images"] == 1, "QA did not count body image occurrence")
    assert_true(not result["report"]["issues"], f"unexpected QA issues: {result['report']['issues']}")


@case
def caption_flow_pairs_images_with_nearby_captions() -> None:
    paragraphs = [
        {"role": "image", "image": "a.png"},
        "intervening prose",
        {"role": "figure_caption", "text": "Fig. 1 Demo"},
        {"role": "image", "image": "b.png"},
        {"role": "image", "image": "c.png"},
        {"role": "figure_caption", "text": "Figure 2 First"},
        {"role": "figure_caption", "text": "Figure 3 Second"},
    ]
    paired = pair_figure_blocks(paragraphs)
    assert_true(paired[0] == {"role": "figure", "image": "a.png", "caption": "Fig. 1 Demo"}, "image/body/caption was not paired")
    assert_true(paired[1] == "intervening prose", "intervening prose was not preserved after pairing")
    assert_true(paired[2] == {"role": "figure", "image": "b.png", "caption": "Figure 2 First"}, "first stacked image was not paired")
    assert_true(paired[3] == {"role": "figure", "image": "c.png", "caption": "Figure 3 Second"}, "second stacked image was not paired")
    assert_true(is_figure_caption("\u56fe1\u7ed3\u679c\u5bf9\u6bd4"), "Chinese figure caption was not detected")
    assert_true(normalize_caption_spacing("\u56fe1\u7ed3\u679c") == "\u56fe 1 \u7ed3\u679c", "Chinese figure caption spacing was not normalized")


@case
def text_cleaner_removes_editor_noise_and_preserves_code_lines() -> None:
    assert_true(clean_text_artifacts("[label](https://example.test)") == "label", "markdown link label was not preserved")
    assert_true(clean_text_artifacts("Plain Text") == "", "editor noise was not removed")
    assert_true(clean_text_artifacts("\u590d\u5236") == "", "Chinese editor noise was not removed")
    code = clean_code_text(" interface  Gi0/0/1 \nPlain Text\n ip   address  10.0.0.1 ")
    assert_true(code == "interface Gi0/0/1\nip address 10.0.0.1", f"code cleanup changed unexpectedly: {code!r}")


@case
def paragraph_stream_run_group_preserves_text_and_math_semantics() -> None:
    section = {"paragraphs": []}

    def append_text(section_obj: Dict[str, Any], text: str, in_appendix: bool = False) -> None:
        section_obj["paragraphs"].append({"role": "text", "text": text, "appendix": in_appendix})

    append_stream_run_group(
        section,
        [{"type": "text", "text": "plain text"}],
        append_text_or_code_func=append_text,
        in_appendix=True,
    )
    assert_true(section["paragraphs"][0] == {"role": "text", "text": "plain text", "appendix": True}, "plain run group bypassed text callback")

    append_stream_run_group(
        section,
        [{"type": "math", "text": "a=b", "math": [{"text": "a=b", "had_number_label": True}]}],
        append_text_or_code_func=append_text,
    )
    formula = section["paragraphs"][1]
    assert_true(formula["role"] == "formula", "math-only run group did not become a formula")
    assert_true(formula["numbered"] is True, "math had_number_label was not preserved")


@case
def reference_collector_exits_before_backmatter_and_keeps_tables() -> None:
    collector = ReferenceCollector(
        clean_text_func=clean_text_artifacts,
        is_backmatter_heading_func=lambda text: text == "Appendix A",
        normalize_heading_spacing_func=lambda text: text,
        classify_section_role_func=lambda heading, level: "appendix",
        table_rows_look_like_code_func=lambda rows: rows and rows[0] and rows[0][0].startswith("interface"),
        code_text_from_table_rows_func=lambda rows, clean_code_func=None: clean_code_func(rows[0][0]) if clean_code_func else rows[0][0],
        clean_code_func=clean_code_text,
    )
    assert_true(is_reference_heading("\u53c2\u8003\u6587\u732e"), "Chinese reference heading was not detected")
    assert_true(collector.start_if_heading("References"), "English reference heading was not accepted")
    assert_true(collector.consume_text("[1]  Example  Paper"), "active collector did not consume reference text")
    assert_true(collector.consume_table_rows([["interface  Gi0/0/1\n ip  address  10.0.0.1"]]), "active collector did not consume reference table")
    backmatter = collector.exit_to_backmatter_section("Appendix A", 1)
    assert_true(backmatter and backmatter["role"] == "appendix", "collector did not exit on backmatter")
    refs = collector.finish()
    assert_true(refs[0] == "[1] Example Paper", f"reference text was not cleaned: {refs[0]!r}")
    assert_true(refs[1]["role"] == "code", "reference code table was not preserved as code")
    assert_true(refs[1]["code"] == "interface Gi0/0/1\nip address 10.0.0.1", "reference code table was not cleaned")


@case
def section_builder_filters_placeholder_and_finalizes_blocks() -> None:
    sections = [
        make_body_section(),
        {"heading": "Abstract", "level": 1, "role": "en_abstract", "paragraphs": ["Summary"], "images": []},
        {"heading": "1 Introduction", "level": 1, "role": "body", "paragraphs": [], "images": []},
        {
            "heading": "2 Results",
            "level": 1,
            "role": "body",
            "paragraphs": [
                {"role": "image", "image": "fig1.png"},
                {"role": "figure_caption", "text": "Fig. 1. Result overview"},
            ],
            "images": ["fig1.png"],
        },
    ]
    filtered = filter_content_sections(sections)
    assert_true(filtered[0]["heading"] == "Abstract", "initial body placeholder was not filtered")
    assert_true(filtered[1]["heading"] == "1 Introduction", "empty structural heading was dropped")
    mark_first_body_page_break(filtered)
    assert_true(not filtered[0].get("page_break_before"), "front matter was marked as first body page")
    assert_true(filtered[1].get("page_break_before") is True, "first body section was not marked for page break")
    postprocess_section_paragraphs(filtered)
    assert_true(filtered[2]["paragraphs"][0]["role"] == "figure", "image and figure caption were not paired")
    assert_true(filtered[2]["paragraphs"][0]["caption"] == "Fig. 1. Result overview", "figure caption text changed")


@case
def body_dispatcher_routes_sections_references_tables_and_appendix_code() -> None:
    work = new_workdir("body_dispatcher")
    fig_dir = work / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("报名序号: [报名序号]")
    doc.add_paragraph("图 1 系统结构示意")
    doc.add_paragraph("参考文献")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "[1] Zhang S. Synthetic reference."
    doc.add_paragraph("附录 A")
    doc.add_paragraph("interface Gi0/0/1\n ip address 10.0.0.1")

    scratch = {"paragraphs": []}
    append_text_or_code(scratch, r"$$a=b$$")
    assert_true(scratch["paragraphs"][0]["role"] == "formula", "dispatcher text append did not preserve display math")

    result = parse_body_sections(doc, 0, ImageRegistry(str(fig_dir), "dispatch_img"))
    sections = filter_content_sections(result.sections)
    assert_true(sections[0]["heading"] == "1 Introduction", "body heading was not routed into a section")
    assert_true(result.placeholders_removed == 1, "unfilled placeholder paragraph was not filtered")
    assert_true(sections[0]["paragraphs"][0]["role"] == "figure_caption", "figure caption was not classified")
    assert_true(result.references and result.references[0]["role"] == "table", "reference table was not captured before appendix")
    assert_true(sections[1]["role"] == "appendix", "back matter did not exit reference collection")
    assert_true(sections[1]["paragraphs"][0]["role"] == "code", "appendix command text was not routed as code")


@case
def content_parser_preserves_table_at_start_of_body() -> None:
    work = new_workdir("parser_table_first")

    docx = work / "table_first_then_text.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Accuracy"
    table.cell(1, 1).text = "98%"
    doc.add_paragraph("Body paragraph after the opening table.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    items = [item for sec in content.get("sections") or [] for item in sec.get("paragraphs") or []]
    table_items = [item for item in items if isinstance(item, dict) and item.get("table_rows")]
    assert_true(len(table_items) == 1, f"opening table was not preserved in content stream: {items}")
    assert_true(table_items[0]["table_rows"][1] == ["Accuracy", "98%"], "opening table rows changed")

    table_only_docx = work / "table_only.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Only"
    table.cell(0, 1).text = "Table"
    doc.save(table_only_docx)

    table_only = extract_docx_content(str(table_only_docx), output_dir=str(work / "only_out"))
    only_items = [item for sec in table_only.get("sections") or [] for item in sec.get("paragraphs") or []]
    assert_true(
        any(isinstance(item, dict) and item.get("table_rows") for item in only_items),
        f"table-only document became empty content: {table_only}",
    )


@case
def content_parser_does_not_count_cover_table_as_body_table() -> None:
    work = new_workdir("parser_cover_table_count")
    docx = work / "cover_table_count.docx"
    doc = Document()
    cover = doc.add_table(rows=1, cols=2)
    cover.cell(0, 0).text = "\u5b66\u6821\u7f16\u7801"
    cover.cell(0, 1).text = "10001"
    title = "\u9762\u5411\u7eff\u7535\u6d88\u7eb3\u7684\u591a\u6e90\u80fd\u6e90\u8c03\u5ea6\u7814\u7a76"
    p = doc.add_paragraph(title)
    p.runs[0].font.size = Pt(16)
    doc.add_paragraph("\u6458\u8981\uff1a\u672c\u6587\u6458\u8981\u7528\u4e8e\u9a8c\u8bc1\u5c01\u9762\u8868\u683c\u4e0d\u5e94\u8ba1\u5165\u6b63\u6587\u8868\u683c\u3002")
    doc.add_paragraph("\u5173\u952e\u8bcd\uff1a\u7eff\u7535\uff1b\u8c03\u5ea6")
    doc.add_heading("\u7b2c1\u7ae0 \u7eea\u8bba", 1)
    doc.add_paragraph("\u6b63\u6587\u6bb5\u843d\u3002")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    meta = content.get("_meta") or {}
    assert_true(meta.get("source_tables_count") == 1, f"source table count was not recorded: {meta}")
    assert_true(meta.get("tables_count") == 0, f"cover table was counted as a body table: {meta}")


@case
def low_res_image_fragment_is_reported_and_not_upscaled() -> None:
    img_src = new_workdir("tiny_image_src")
    write_sample_png(img_src / "tiny.png", width=76, height=18)
    content = base_content([
        {"role": "figure", "image": "tiny.png", "caption": "Figure 1 broken label shard"}
    ])
    content["_meta"]["images_dir"] = str(img_src)
    content["_meta"]["images_extracted"] = 1
    content["sections"][0]["images"] = ["tiny.png"]
    result = run_generated_case("tiny_image_fragment", content)
    codes = [item["code"] for item in result["report"]["issues"]]
    assert_true("LOW_RES_IMAGE_FRAGMENT" in codes, "QA did not report a low-resolution image fragment")
    issue = next(item for item in result["report"]["issues"] if item["code"] == "LOW_RES_IMAGE_FRAGMENT")
    assert_true(issue["severity"] == "warning", f"contained image fragment should be a warning: {issue}")
    assert_true(result["manifest"]["counts"].get("content_image_fragments_contained") == 1, "contained image fragment was not counted")
    m = re.search(r"<wp:extent[^>]+cx=\"(\d+)\"[^>]+cy=\"(\d+)\"", result["xml"])
    assert_true(bool(m), "generated DOCX did not contain an image extent")
    width_inches = int(m.group(1)) / 914400
    assert_true(width_inches < 2.0, f"tiny image was upscaled to page width: {width_inches:.2f}in")


@case
def small_square_icon_is_not_reported_as_low_res_fragment() -> None:
    img_src = new_workdir("small_icon_src")
    write_sample_png(img_src / "icon.png", width=64, height=64)
    content = base_content([
        {"role": "figure", "image": "icon.png", "caption": "QR code icon"}
    ])
    content["_meta"]["images_dir"] = str(img_src)
    content["_meta"]["images_extracted"] = 1
    content["sections"][0]["images"] = ["icon.png"]
    result = run_generated_case("small_square_icon", content)
    codes = [item["code"] for item in result["report"]["issues"]]
    assert_true("LOW_RES_IMAGE_FRAGMENT" not in codes, f"legitimate small square image was reported as fragment: {result['report']['issues']}")


@case
def qa_manifest_detects_missing_image_render() -> None:
    img_src = new_workdir("image_missing_src")
    write_sample_png(img_src / "dot.png")
    content = base_content([
        {"role": "figure", "image": "dot.png", "caption": "Figure 1 sample"}
    ])
    content["_meta"]["images_dir"] = str(img_src)
    content["_meta"]["images_extracted"] = 1
    content["sections"][0]["images"] = ["dot.png"]
    result = run_generated_case("qa_missing_image", content)
    manifest_path = result["work"] / "build_manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["counts"]["content_images_rendered"] = 0
    write_json(manifest_path, manifest)
    report = check_output(str(result["work"]), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("IMAGE_COUNT_MISMATCH" in codes, "QA did not trust manifest for image mismatch")
    repair = report.get("repair_plan") or {}
    assert_true(repair.get("blocking_errors", 0) >= 1, "repair plan did not count blocking errors")
    steps = repair.get("steps") or []
    assert_true(steps and steps[0].get("severity") == "error", "repair plan did not prioritize errors")
    assert_true(any(step.get("code") == "IMAGE_COUNT_MISMATCH" for step in steps), "repair plan omitted image repair guidance")
    write_reports(report, str(result["work"]))
    assert_true((result["work"] / "qa_repair_plan.md").exists(), "repair plan markdown was not written")
    assert_true((result["work"] / "qa_fix_prompt.txt").exists(), "AI fix prompt was not written")


@case
def qa_manifest_detects_missing_formula_render() -> None:
    content = base_content([
        {"role": "formula", "latex": "a=b+c", "text": "a=b+c", "numbered": False}
    ])
    result = run_generated_case("qa_missing_formula", content)
    manifest_path = result["work"] / "build_manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["counts"]["content_formulas_rendered"] = 0
    write_json(manifest_path, manifest)
    report = check_output(str(result["work"]), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("FORMULA_COUNT_MISMATCH" in codes, "QA did not trust manifest for formula mismatch")


@case
def omml_display_item_builds_display_formula() -> None:
    xml = latex_to_omath("a=b+c", display=True)
    content = base_content([
        {
            "role": "formula",
            "source": "omml",
            "text": "a=b+c",
            "math": [{"type": "display", "xml": xml, "text": "a=b+c"}],
            "numbered": False,
        }
    ])
    result = run_generated_case("omml_display", content)
    assert_true(result["manifest"]["counts"]["display_formulas_rendered"] == 1, "OMML display formula was not counted")
    assert_true(omath_para_count(result["xml"]) == 1, "OMML display formula did not remain display")


@case
def content_parser_preserves_mixed_docx_tokens() -> None:
    work = new_workdir("parser_mixed")
    img = work / "dot.png"
    img.write_bytes(PNG_1X1)
    docx = work / "mixed.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("before ")
    p._element.append(etree.fromstring(latex_to_omath("x^2", display=False).encode("utf-8")))
    p.add_run(" after")
    p_img = doc.add_paragraph()
    p_img.add_run("image before ")
    p_img.add_run().add_picture(str(img))
    p_img._element.append(etree.fromstring(latex_to_omath("z^2", display=False).encode("utf-8")))
    p_img.add_run(" image after")
    p_disp = doc.add_paragraph()
    p_disp._element.append(etree.fromstring(latex_to_omath("a=b+c", display=True).encode("utf-8")))
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    rich_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "rich_text"]
    formula_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "formula"]
    image_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "image"]
    assert_true(any([r.get("type") for r in item.get("runs", [])] == ["text", "math", "text"] for item in rich_items), "mixed text/math run order was not preserved")
    assert_true(any((item.get("runs") or [])[0].get("text") == "before " and (item.get("runs") or [])[-1].get("text") == " after" for item in rich_items), "inline formula surrounding spaces were not preserved")
    assert_true(image_items, "image token disappeared from mixed paragraph")
    assert_true(any((item.get("math") or [{}])[0].get("type") == "display" for item in formula_items), "display OMML was not extracted as formula")


@case
def content_parser_extracts_vml_pictures() -> None:
    work = new_workdir("parser_vml")
    img = work / "dot.png"
    img.write_bytes(PNG_1X1)
    src_docx = work / "drawing.docx"
    vml_docx = work / "vml.docx"
    doc = Document()
    doc.add_paragraph().add_run().add_picture(str(img))
    doc.save(src_docx)
    make_vml_picture_docx(src_docx, vml_docx)
    content = extract_docx_content(str(vml_docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    assert_true(any(isinstance(p, dict) and p.get("role") == "image" for p in paragraphs), "VML picture was not extracted")


@case
def content_parser_detects_latex_delimited_formula_paragraphs() -> None:
    work = new_workdir("parser_latex_delimited")
    docx = work / "latex_delimited.docx"
    doc = Document()
    doc.add_paragraph("1 Formula Cases")
    doc.add_paragraph(r"$$L=\lim_{n\to\infty}\frac{1}{n}\sum_{i=1}^{n}x_i$$")
    doc.add_paragraph(r"$$M=\begin{matrix}a&b\\c&d\end{matrix}$$")
    doc.add_paragraph(r"$$I=\int_0^T f(t)\,dt$$")
    doc.add_paragraph(r"$$E_{total}=\sum_{t=1}^{24}P(t)\Delta t$$ (1.1)")
    doc.add_paragraph(r"$$R_{green}=\frac{E_{renew}}{E_{total}}\times100\%$$" + " \uff08A.1\uff09")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    formulas = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "formula"]
    assert_true(len(formulas) == 5, f"LaTeX-delimited paragraphs were not all formulas: {paragraphs}")
    assert_true(all(f.get("source") == "latex" and f.get("latex") for f in formulas), "LaTeX delimiters were not stripped into latex fields")
    assert_true(all("$$" not in (f.get("latex") or "") for f in formulas), "LaTeX delimiters leaked into formula latex")
    assert_true(any(f.get("text", "").endswith("\uff08A.1\uff09") and f.get("latex", "").startswith("R_{green}") for f in formulas), "appendix formula label was not handled")


@case
def formula_semantics_classifies_quantities_and_contamination() -> None:
    quantity = "全年发电量为 172.04 MWh"
    equation = "P_total(t)=P_AK+P_PM+P_NH3+P_L(t)=20.75+6*p_L(t)"
    contaminated = "当PRE(t)-PL(t)<0.1*41.5=4.15MW时,该时段无法开机;产量约束为="

    quantity_result = classify_formula_text(quantity)
    equation_result = classify_formula_text(equation)
    contaminated_result = classify_formula_text(contaminated)

    assert_true(quantity_result.category == CATEGORY_QUANTITY_TEXT, f"quantity became {quantity_result}")
    assert_true(not semantic_looks_like_formula_text(quantity), "quantity/unit text was mistaken for a formula")
    assert_true(equation_result.category == CATEGORY_DISPLAY_MATH, f"equation became {equation_result}")
    assert_true(semantic_looks_like_formula_text(equation), "standalone equation was not recognized")
    assert_true(contaminated_result.category == CATEGORY_CONTAMINATED, f"contaminated formula became {contaminated_result}")
    assert_true(not semantic_looks_like_formula_text(contaminated), "contaminated narrative was accepted as a formula")


@case
def content_parser_marks_formula_semantic_problems() -> None:
    work = new_workdir("parser_formula_semantics")
    docx = work / "formula_semantics.docx"
    quantity = "全年发电量为 172.04 MWh"
    equation = "P_total(t)=P_AK+P_PM+P_NH3+P_L(t)=20.75+6*p_L(t)"
    contaminated = "当PRE(t)-PL(t)<0.1*41.5=4.15MW时,该时段无法开机;产量约束为="
    doc = Document()
    doc.add_paragraph("1 Formula Semantics")
    doc.add_paragraph(quantity)
    doc.add_paragraph(equation)
    doc.add_paragraph(contaminated)
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    formulas = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "formula"]
    problems = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "formula_problem"]

    assert_true(quantity in paragraphs, "quantity text should remain ordinary text")
    assert_true(any(p.get("text") == equation for p in formulas), "standalone equation was not extracted as formula")
    assert_true(any(p.get("text") == contaminated for p in problems), "contaminated formula text was not marked")


@case
def formula_semantics_splits_inline_math_without_units() -> None:
    text = "模型中 P_total(t)=20.75+6*p_L(t)，全年电量为 172.04 MWh。"
    spans = split_inline_math_spans(text)
    assert_true(len(spans) == 1, f"expected one inline formula span, got {spans}")
    assert_true(spans[0]["text"] == "P_total(t)=20.75+6*p_L(t)", f"wrong inline span: {spans}")
    assert_true("172.04" not in spans[0]["text"], "quantity text was absorbed into inline formula")
    assert_true(not split_inline_math_spans("$100$ cost"), "plain dollar amount was mistaken for inline math")
    assert_true(not split_inline_math_spans("$$x=1$$"), "display math delimiter was mistaken for inline math")
    dollar_spans = split_inline_math_spans("变量 $x_i$ 表示第 i 个样本。")
    assert_true(dollar_spans and dollar_spans[0]["text"] == "x_i", f"valid dollar inline math was lost: {dollar_spans}")
    mixed = "由于约束仅为 u(t) = n 且 u(t) ∈ {0, 1}，这是标准选择问题。"
    mixed_spans = split_inline_math_spans(mixed)
    assert_true(not any("且" in str(s.get("text")) or "{" in str(s.get("text")) for s in mixed_spans), f"unsafe inline span accepted: {mixed_spans}")


@case
def prose_with_inline_formula_is_not_formula_problem() -> None:
    text = "假设三：每小时为一个调度时段。题目明确园区运行分析时段为 1 小时，即 ∆t = 1。"
    spans = split_inline_math_spans(text)
    assert_true(any(s.get("text") == "t = 1" for s in spans), f"inline formula was not found: {spans}")
    assert_true(not is_formula_problem_text(text), "ordinary prose with inline math should not block as formula_problem")


@case
def content_parser_builds_rich_text_for_plain_inline_formula() -> None:
    work = new_workdir("parser_plain_inline_formula")
    docx = work / "plain_inline_formula.docx"
    paragraph = "模型中 P_total(t)=20.75+6*p_L(t)，全年电量为 172.04 MWh。"
    doc = Document()
    doc.add_paragraph("1 Inline Formula")
    doc.add_paragraph(paragraph)
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    rich_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "rich_text"]
    assert_true(rich_items, f"inline formula paragraph was not converted to rich_text: {paragraphs}")
    runs = rich_items[0].get("runs") or []
    assert_true([r.get("type") for r in runs] == ["text", "math", "text"], f"inline formula runs are wrong: {runs}")
    result = run_generated_case("plain_inline_formula_render", content)
    assert_true(result["manifest"]["counts"]["inline_formulas_rendered"] == 1, "plain inline formula did not render as native inline math")
    assert_true(omath_count(result["xml"]) >= 1 and omath_para_count(result["xml"]) == 0, "plain inline formula rendered with wrong OMML shape")


@case
def content_parser_keeps_prose_inline_formula_as_rich_text() -> None:
    work = new_workdir("parser_prose_inline_formula")
    docx = work / "prose_inline_formula.docx"
    doc = Document()
    doc.add_paragraph("1 Formula Prose")
    doc.add_paragraph("假设三：每小时为一个调度时段。题目明确园区运行分析时段为 1 小时，即 ∆t = 1。")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    assert_true(not any(isinstance(p, dict) and p.get("role") == "formula_problem" for p in paragraphs), f"prose inline formula became formula_problem: {paragraphs}")
    rich = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "rich_text"]
    assert_true(rich and any(run.get("type") == "math" for run in rich[0].get("runs") or []), f"prose inline formula was not rich_text math: {paragraphs}")


@case
def numeric_formula_denominator_is_not_heading() -> None:
    work = new_workdir("parser_numeric_fragment_heading")
    docx = work / "numeric_fragment_heading.docx"
    doc = Document()
    doc.add_paragraph("论文题目: Numeric Fragment")
    doc.add_paragraph("8 离网运行分析")
    p = doc.add_paragraph("41 .5")
    run = p.runs[0]
    run.font.size = Pt(16)
    run.bold = True
    doc.add_paragraph("This text should stay under the previous heading.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    assert_true("41 .5" not in headings, f"numeric formula denominator became heading: {headings}")
    assert_true("41 .5" in body_text and "previous heading" in body_text, f"numeric fragment/body text was lost: {body_text}")


@case
def content_parser_repairs_split_ratio_formula_layout() -> None:
    work = new_workdir("parser_split_ratio_layout")
    docx = work / "split_ratio_layout.docx"
    doc = Document()
    doc.add_paragraph("1 Split Formula Layout")
    doc.add_paragraph("日总用电量和新能源发电量分别为：")
    doc.add_paragraph("24 24")
    doc.add_paragraph("Etotal = ∑ Ptotal(t) · ∆t, ERE = ∑ PRE(t) · ∆t")
    doc.add_paragraph("t=1")
    doc.add_paragraph("三项绿电直连指标定义为：")
    doc.add_paragraph("t=1")
    doc.add_paragraph("rself")
    p1 = doc.add_paragraph()
    p1._element.append(etree.fromstring(latex_to_omath(r"=Etotal-Esell-Ebuy\times100\%(6)(3)", display=True).encode("utf-8")))
    doc.add_paragraph("E")
    doc.add_paragraph("rgreen")
    doc.add_paragraph("RE")
    p2 = doc.add_paragraph()
    p2._element.append(etree.fromstring(latex_to_omath(r"=ERE-Esell\times100\%(7)(4)", display=True).encode("utf-8")))
    doc.add_paragraph("E")
    doc.add_paragraph("rup")
    doc.add_paragraph("total")
    p3 = doc.add_paragraph()
    p3._element.append(etree.fromstring(latex_to_omath(r"=Esell\times100\%(8)(5)", display=True).encode("utf-8")))
    doc.add_paragraph("E")
    doc.add_paragraph("RE")
    doc.add_paragraph("吨氨成本模型")
    doc.add_paragraph("吨氨成本由购电成本、风光度电成本和运维成本三部分构成，扣除售电收入：")
    doc.add_paragraph("∑24")
    doc.add_paragraph("[Pbuy(t) · λbuy(t) − Psell(t) · λsell] · 1000 + CRE + COM")
    doc.add_paragraph("其中风光度电成本 CRE = ∑24")
    doc.add_paragraph("[Pw(t) · 0.15 + Ppv(t) · 0.12] · 1000（元），运维成本")
    doc.add_paragraph("COM = ∑24")
    doc.add_paragraph("u(t)·(10×0.1+10×0.15+0.75×0.002)·1000(元),日产氨量QNH3=1.5×24=36")
    doc.add_paragraph("24")
    doc.add_paragraph("u(t)=n=Qtarget/3")
    doc.add_paragraph("t=1")
    doc.add_paragraph("Cton =")
    doc.add_paragraph("1")
    doc.add_paragraph("Qtarget")
    doc.add_paragraph("24")
    doc.add_paragraph("t=1")
    doc.add_paragraph("∆c(t) · u(t) + Cfixed]")
    doc.add_paragraph("(13)")
    doc.add_paragraph("max")
    doc.add_paragraph("α")
    doc.add_paragraph("24")
    doc.add_paragraph("QNH3 = 3 α(t) (20)")
    doc.add_paragraph("t=1")
    doc.add_paragraph("CRE=∑_{t=1}^{24} [Pw(t) · 0.15 + Ppv(t) · 0.12] · 1000")
    doc.add_paragraph("COM=∑24 (10×0.1+10×0.15+0.75×0.002)·1000")
    doc.add_paragraph("S = ∑10")
    doc.add_paragraph("x_i")
    doc.add_paragraph("K = ∑10")
    doc.add_paragraph("42")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    repaired = [p for p in paragraphs if isinstance(p, dict) and str(p.get("source") or "").startswith("repaired_")]
    repaired_text = "\n".join(p.get("text") or "" for p in repaired)
    residue = [p for p in paragraphs if p in ("rself", "rgreen", "rup", "RE") or (isinstance(p, str) and p.strip() == "E")]
    split_sum_problems = [p for p in paragraphs if isinstance(p, dict) and p.get("problem") == "split_sum_index_unknown"]

    assert_true(any(p.get("source") == "repaired_sum_bounds" for p in repaired), f"split sum bounds were not repaired: {paragraphs}")
    assert_true(sum(1 for p in repaired if p.get("source") == "repaired_ratio_cluster") == 3, f"ratio formulas were not repaired: {repaired}")
    assert_true(any(p.get("source") == "repaired_sum_prefix" for p in repaired), f"split leading sum prefix was not repaired: {paragraphs}")
    assert_true(sum(1 for p in repaired if p.get("source") == "repaired_labeled_sum_continuation") >= 2, f"labeled sum continuations were not repaired: {repaired}")
    assert_true(any(p.get("source") == "repaired_missing_sum_symbol" for p in repaired), f"missing sum symbol layout was not repaired: {repaired}")
    assert_true(any(p.get("source") == "repaired_fraction_sum_layout" for p in repaired), f"fraction plus sum layout was not repaired: {repaired}")
    assert_true(any(p.get("source") == "repaired_max_sum_layout" for p in repaired), f"max/min sum layout was not repaired: {repaired}")
    assert_true(any(p.get("source") == "repaired_inline_sum_missing_lower" for p in repaired), f"inline sum missing lower was not repaired from context: {repaired}")
    assert_true("rself=" in repaired_text and "rgreen=" in repaired_text and "rup=" in repaired_text, f"ratio variables missing: {repaired_text}")
    assert_true(all(r"\mathrm" in (p.get("latex") or "") for p in repaired if p.get("source") == "repaired_ratio_cluster"), "multi-letter repaired variables should use roman subscripts")
    latex_blob = "\n".join(p.get("latex") or "" for p in repaired)
    assert_true(r"\Deltat" not in latex_blob and r"\lambdabuy" not in latex_blob, f"greek-letter suffixes collapsed into invalid commands: {latex_blob}")
    assert_true(r"\Delta t" in latex_blob and r"\lambda_{\mathrm{buy}}" in latex_blob, f"greek-letter suffixes were not preserved: {latex_blob}")
    assert_true(r"\sum_{i=1}^{10}" in latex_blob, f"split sum index was not inferred from subscript variable: {latex_blob}")
    assert_true(r"\sum_{t=1}^{24}" in latex_blob and r"\frac{1}{Q_{\mathrm{target}}}" in latex_blob, f"time-indexed sum/fraction layout was not rendered: {latex_blob}")
    assert_true(r"\max_{\alpha}" in latex_blob, f"max sum layout did not preserve optimization variable: {latex_blob}")
    assert_true(not any((p.get("text") or "").startswith("K=") for p in repaired), f"index-less split sum was repaired by guessing: {repaired}")
    assert_true(split_sum_problems, "index-less split sum should be flagged instead of repaired by guessing")
    assert_true("(5)" not in latex_blob, f"stale source equation label leaked into repaired latex: {latex_blob}")
    assert_true(not residue, f"split ratio fragments leaked as body text: {residue}")
    result = run_generated_case("split_ratio_layout_render", content)
    assert_true(result["manifest"]["counts"]["display_formulas_rendered"] >= 7, "repaired split formulas did not render as display math")


@case
def content_parser_extracts_table_cell_images_and_flags_header_images() -> None:
    work = new_workdir("parser_table_header_images")
    img = work / "dot.png"
    img.write_bytes(PNG_1X1)
    docx = work / "table_header_images.docx"
    doc = Document()
    doc.sections[0].header.paragraphs[0].add_run().add_picture(str(img))
    doc.add_paragraph("1 Images")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).paragraphs[0].add_run().add_picture(str(img))
    table.cell(0, 1).text = "image in table cell"
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_images = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "image" and p.get("location") == "table_cell"]
    assert_true(table_images, "table-cell image was not promoted into the content image stream")
    assert_true(content["_meta"]["images_extracted"] == 1, "header image should not be counted as a body image")
    assert_true(content["_meta"].get("non_body_images"), "header image was not recorded as a non-body image")


@case
def content_parser_keeps_references_before_english_appendix() -> None:
    work = new_workdir("parser_refs_appendix")
    docx = work / "refs_appendix.docx"
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Body paragraph before references.")
    doc.add_paragraph("References")
    doc.add_paragraph("[1] Synthetic reference one.")
    doc.add_paragraph("[2] Synthetic reference two.")
    doc.add_paragraph("Appendix A Reproducible Commands")
    doc.add_paragraph("python run_pipeline.py --mode developer")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    assert_true(len(content.get("references") or []) == 2, f"reference entries were not preserved: {content.get('references')}")
    assert_true(all("python run_pipeline" not in str(ref) for ref in content.get("references") or []), "appendix command leaked into references")
    appendix = [sec for sec in content.get("sections") or [] if sec.get("role") == "appendix"]
    assert_true(appendix and any("python run_pipeline" in str(p) for p in appendix[0].get("paragraphs") or []), "English appendix section was not preserved")


@case
def content_parser_extracts_english_title_before_abstract() -> None:
    work = new_workdir("parser_english_title")
    docx = work / "english_title.docx"
    doc = Document()
    p = doc.add_paragraph("中文论文标题示例用于测试")
    p.runs[0].font.size = Pt(16)
    doc.add_paragraph("Research on Template Driven Document Quality Assessment")
    doc.add_paragraph("Abstract: This is the abstract body.")
    doc.add_paragraph("Key words: document automation")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Body text.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    assert_true(content.get("title_info", {}).get("title_en") == "Research on Template Driven Document Quality Assessment", "English title before abstract was not extracted")
    assert_true(not any((sec.get("heading") or "").startswith("Research on") for sec in content.get("sections") or []), "English title leaked into body sections")


@case
def content_parser_extracts_docx_title_style_without_explicit_font_size() -> None:
    work = new_workdir("parser_title_style")
    docx = work / "title_style.docx"
    title = "\u9762\u5411\u4e2d\u6587\u6bd5\u4e1a\u8bba\u6587\u7684\u673a\u5668\u5b66\u4e60\u6392\u7248\u6d41\u6c34\u7ebf\u7814\u7a76"
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_heading("\u6458\u8981", 1)
    doc.add_paragraph("\u672c\u6587\u6458\u8981\u7528\u4e8e\u9a8c\u8bc1\u9898\u540d\u6837\u5f0f\u63d0\u53d6\u3002")
    doc.add_heading("1 \u7eea\u8bba", 1)
    doc.add_paragraph("Body text.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    assert_true(content.get("title_info", {}).get("title_cn") == title, f"title style was not extracted: {content.get('title_info')}")
    assert_true(title not in headings, f"title style leaked into body sections: {headings}")
    assert_true("\u6458\u8981" in headings, f"abstract heading was lost after title extraction: {headings}")


@case
def content_parser_extracts_plain_title_before_abstract() -> None:
    work = new_workdir("parser_plain_front_title")
    docx = work / "plain_front_title.docx"
    title = "\u9762\u5411\u7eff\u7535\u6d88\u7eb3\u7684\u591a\u6e90\u80fd\u6e90\u8c03\u5ea6\u4e0e\u7ecf\u6d4e\u6027\u8bc4\u4f30\u7814\u7a76"
    doc = Document()
    doc.add_paragraph(title)
    doc.add_paragraph("\u6458\u8981\uff1a\u672c\u6587\u6458\u8981\u7528\u4e8e\u9a8c\u8bc1\u65e0\u663e\u5f0f\u5b57\u53f7\u7684\u524d\u7f6e\u9898\u540d\u3002")
    doc.add_paragraph("\u5173\u952e\u8bcd\uff1a\u7eff\u7535\uff1b\u8c03\u5ea6")
    doc.add_heading("\u7b2c1\u7ae0 \u7eea\u8bba", 1)
    doc.add_paragraph("\u6b63\u6587\u6bb5\u843d\u3002")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    assert_true(content.get("title_info", {}).get("title_cn") == title, f"plain front title was not extracted: {content.get('title_info')}")
    assert_true(title not in headings, f"plain title leaked into body sections: {headings}")


@case
def content_parser_does_not_treat_chapter_heading_as_title() -> None:
    work = new_workdir("parser_chapter_not_title")
    docx = work / "chapter_not_title.docx"
    heading = "\u7b2c\u4e00\u7ae0 \u7eea\u8bba\u4e0e\u7814\u7a76\u80cc\u666f"
    doc = Document()
    doc.add_heading(heading, 1)
    doc.add_paragraph("\u8fd9\u662f\u6b63\u6587\u7b2c\u4e00\u6bb5\uff0c\u7528\u4e8e\u9a8c\u8bc1\u76f4\u63a5\u4ece\u7ae0\u8282\u5f00\u59cb\u7684\u6587\u6863\u3002")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    assert_true(content.get("title_info", {}).get("title_cn") != heading, f"chapter heading was misdetected as title: {content.get('title_info')}")
    assert_true(heading in headings, f"chapter heading was not preserved as a body section: {headings}")


@case
def content_parser_splits_chinese_enumerated_headings_after_keywords() -> None:
    work = new_workdir("parser_cn_enum")
    img = work / "dot.png"
    img.write_bytes(PNG_1X1)
    docx = work / "cn_enum.docx"
    doc = Document()
    doc.add_paragraph("摘要")
    doc.add_paragraph("本文用于测试。")
    doc.add_paragraph("关键词：")
    doc.add_paragraph("绿电直连；优化运行")
    doc.add_paragraph("一、问题重述")
    doc.add_paragraph("正文段落。")
    doc.add_paragraph().add_run().add_picture(str(img))
    doc.add_paragraph("图 1-1 示例图片")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    assert_true("一、问题重述" in headings, "Chinese enumerated heading was not split into a body section")
    body_sec = next(sec for sec in content.get("sections") or [] if sec.get("heading") == "一、问题重述")
    assert_true(body_sec.get("images"), "Image after Chinese enumerated heading was not assigned to body section")
    kw_sec = next((sec for sec in content.get("sections") or [] if sec.get("role") == "cn_keywords"), {})
    assert_true(not kw_sec.get("images"), "Image after body heading leaked into keyword front matter")


@case
def md_parser_keeps_table_code_and_rich_math() -> None:
    work = new_workdir("md_parser")
    md = work / "sample.md"
    md.write_text(
        "\n".join(
            [
                "# Sample",
                "",
                "Alpha $x^2$ beta.",
                "",
                "$$a=b+c$$",
                "",
                "| A | B |",
                "| --- | --- |",
                "| 1 | 2 |",
                "",
                "```text",
                "interface G0/0/1",
                "ip address 10.0.0.1 255.255.255.0",
                "```",
            ]
        ),
        encoding="utf-8",
    )
    content = extract_md_content(str(md), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    assert_true(any(isinstance(p, dict) and p.get("role") == "rich_text" and p.get("runs") for p in paragraphs), "MD inline math did not become rich_text runs")
    rich = next(p for p in paragraphs if isinstance(p, dict) and p.get("role") == "rich_text")
    assert_true(rich["runs"][0]["text"] == "Alpha " and rich["runs"][-1]["text"] == " beta.", "MD inline math surrounding spaces were not preserved")
    assert_true(any(isinstance(p, dict) and p.get("role") == "formula" for p in paragraphs), "MD display formula missing")
    assert_true(any(isinstance(p, dict) and p.get("table_rows") and p.get("role") == "table" for p in paragraphs), "MD table missing")
    assert_true(any(isinstance(p, dict) and p.get("role") == "code" for p in paragraphs), "MD code block missing")
    assert_true(content["_meta"]["tables_count"] == 1, "MD table count should be one")


@case
def md_title_only_document_keeps_body_section() -> None:
    work = new_workdir("md_title_only")
    md = work / "title_only.md"
    md.write_text("# Only Title\n\nBody text under the title without another heading.", encoding="utf-8")
    content = extract_md_content(str(md), output_dir=str(work))
    assert_true(content["sections"], "MD body after title-only H1 was dropped")
    assert_true(content["sections"][0]["paragraphs"], "MD title-only body paragraph missing")


@case
def md_image_keeps_paragraph_position() -> None:
    work = new_workdir("md_image_order")
    img = work / "dot.png"
    img.write_bytes(PNG_1X1)
    md = work / "image_order.md"
    md.write_text("# Image Order\n\nBefore image ![dot](dot.png) after image.", encoding="utf-8")
    content = extract_md_content(str(md), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    roles = [p.get("role") if isinstance(p, dict) else "text" for p in paragraphs]
    assert_true(roles == ["text", "image", "text"], f"MD image order drifted: {roles}")
    assert_true(content["_meta"]["images_extracted"] == 1, "MD image was not copied")


@case
def md_missing_images_are_reported_to_qa() -> None:
    work = new_workdir("md_missing_images")
    md = work / "missing_images.md"
    md.write_text(
        "# Missing Images\n\nLocal ![local](missing.png) and remote ![remote](https://example.com/a.png).",
        encoding="utf-8",
    )
    content = extract_md_content(str(md), output_dir=str(work))
    missing = content["_meta"].get("missing_images") or []
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    assert_true(len(missing) == 2, f"missing image references were not recorded: {missing}")
    assert_true(any(isinstance(p, dict) and p.get("role") == "missing_image" for p in paragraphs), "missing image marker not preserved in content stream")

    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("CONTENT_IMAGE_MISSING" in codes, "QA did not report missing Markdown images")
    assert_true(report["passed"] is False, "missing images should fail QA")


@case
def md_parser_keeps_references_before_english_appendix() -> None:
    work = new_workdir("md_refs_appendix")
    md = work / "refs_appendix.md"
    md.write_text(
        "\n".join([
            "# Demo",
            "",
            "Body text.",
            "",
            "## References",
            "[1] Synthetic Markdown reference one.",
            "[2] Synthetic Markdown reference two.",
            "",
            "## Appendix A Commands",
            "python run_pipeline.py --mode developer",
        ]),
        encoding="utf-8",
    )
    content = extract_md_content(str(md), output_dir=str(work))
    assert_true(len(content.get("references") or []) == 2, f"Markdown references were not preserved: {content.get('references')}")
    assert_true(all("python run_pipeline" not in str(ref) for ref in content.get("references") or []), "Markdown appendix command leaked into references")
    assert_true(any((sec.get("heading") or "").startswith("Appendix") for sec in content.get("sections") or []), "Markdown appendix section was not preserved")


@case
def qa_counts_mixed_inline_and_section_images() -> None:
    work = new_workdir("mixed_image_count")
    content = base_content([
        {"role": "image", "image": "inline.png"},
    ])
    content["sections"][0]["images"] = ["inline.png", "section_only.png"]
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    assert_true(report["counts"]["content_images"] == 2, f"mixed image count lost section-only images: {report['counts']}")


@case
def qa_reports_non_body_images_and_raw_latex_text() -> None:
    work = new_workdir("qa_non_body_latex")
    docx = work / "out.docx"
    doc = Document()
    doc.add_paragraph(r"$$x^2+y^2=z^2$$")
    doc.save(docx)
    content = base_content(["Body text"])
    content["_meta"]["non_body_images"] = [{"location": "section_1_header", "target": "media/image1.png"}]
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("NON_BODY_IMAGE_UNSUPPORTED" in codes, "QA did not flag unsupported header/footer images")
    assert_true("LATEX_DELIMITER_TEXT" in codes, "QA did not flag raw LaTeX delimiters left in final DOCX")
    assert_true(report["passed"] is False, "non-body images and raw LaTeX text should fail QA")


@case
def qa_reports_duplicate_front_matter_headings() -> None:
    work = new_workdir("qa_duplicate_front_heading")
    docx = work / "out.docx"
    doc = Document()
    doc.add_paragraph("摘  要")
    doc.add_paragraph("Synthetic title")
    doc.add_paragraph("摘要")
    doc.add_paragraph("Abstract body.")
    doc.save(docx)
    content = {
        "_meta": {"source": "synthetic.docx", "paragraphs": 1, "tables_count": 0, "images_extracted": 0},
        "title_info": {"title_cn": "Synthetic title"},
        "sections": [
            {"heading": "摘要", "level": 1, "role": "cn_abstract", "paragraphs": ["Abstract body."], "images": []}
        ],
        "references": ["[1] Synthetic reference."],
    }
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    (work / "build_generated.py").write_text("# synthetic\n", encoding="utf-8")
    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("DUPLICATE_FRONT_MATTER_HEADING" in codes, "QA did not report duplicate front matter heading")
    assert_true(report["passed"] is False, "duplicate front matter heading should fail QA")


@case
def content_parser_skips_source_toc_and_records_placeholders() -> None:
    work = new_workdir("source_toc_placeholder")
    docx = work / "source_toc.docx"
    doc = Document()
    doc.add_paragraph("报名序号: [报名序号]")
    doc.add_paragraph("论文题目: Synthetic Energy Paper")
    doc.add_paragraph("目 录")
    doc.add_paragraph("一、 问题重述")
    doc.add_paragraph("172.04 MWh should not become a heading")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("一、 问题重述")
    doc.add_paragraph("This is the real body paragraph.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    assert_true(content.get("title_info", {}).get("title_cn") == "Synthetic Energy Paper", "labeled title was not extracted")
    assert_true(content.get("_meta", {}).get("source_toc_skipped_paragraphs", 0) >= 3, "source TOC block was not skipped")
    assert_true(content.get("_meta", {}).get("source_placeholders"), "source placeholders were not recorded")
    assert_true("172.04 MWh should not become a heading" not in headings, "TOC/body sentence became a heading")
    assert_true("一、 问题重述" in headings, "real body heading was lost after TOC skip")


@case
def content_parser_skips_source_toc_with_roman_page_numbers() -> None:
    work = new_workdir("source_toc_roman_pages")
    docx = work / "source_toc_roman_pages.docx"
    doc = Document()
    doc.add_paragraph("\u8bba\u6587\u9898\u76ee: Roman Page TOC")
    doc.add_paragraph("\u76ee\u5f55")
    doc.add_paragraph("\u6458\u8981........................................ I")
    doc.add_paragraph("Abstract.................................... II")
    doc.add_paragraph("1 \u7eea\u8bba..................................... 1")
    doc.add_paragraph("2 \u7cfb\u7edf\u8bbe\u8ba1................................. 5")
    doc.add_paragraph("\u53c2\u8003\u6587\u732e................................... 32")
    doc.add_paragraph("1 \u7eea\u8bba")
    doc.add_paragraph("This real body paragraph must not be collected as a reference.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    assert_true(content.get("_meta", {}).get("source_toc_skipped_paragraphs", 0) >= 5, "roman-page source TOC entries were not skipped")
    assert_true(not content.get("references"), f"source TOC reference line activated reference collection: {content.get('references')}")
    assert_true("1 \u7eea\u8bba" in headings, f"real body heading after roman TOC was lost: {headings}")
    assert_true("real body paragraph" in body_text, f"real body after roman TOC was lost: {body_text}")


@case
def content_parser_exits_source_toc_without_page_break() -> None:
    work = new_workdir("source_toc_no_break")
    docx = work / "source_toc_no_break.docx"
    doc = Document()
    doc.add_paragraph("论文题目: No Break TOC")
    doc.add_paragraph("目录")
    doc.add_paragraph("1 绪论 1")
    doc.add_paragraph("1 正文第一章")
    doc.add_paragraph("This paragraph must survive even though no page break follows the source TOC.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    assert_true(content.get("_meta", {}).get("source_toc_skipped_paragraphs", 0) >= 2, "source TOC entries were not skipped")
    assert_true("1 正文第一章" in headings, f"real heading after source TOC was lost: {headings}")
    assert_true("must survive" in body_text, f"body text after source TOC was lost: {body_text}")


@case
def content_parser_skips_unpaged_source_toc_without_boundary() -> None:
    work = new_workdir("source_toc_unpaged_no_boundary")
    docx = work / "source_toc_unpaged_no_boundary.docx"
    doc = Document()
    doc.add_paragraph("论文题目: Unpaged TOC")
    doc.add_paragraph("目录")
    doc.add_paragraph("一、 问题重述")
    doc.add_paragraph("二、 问题分析")
    doc.add_paragraph("三、 模型假设")
    doc.add_paragraph("一、 问题重述")
    doc.add_paragraph("This is the real body paragraph after the repeated first heading.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    assert_true(content.get("_meta", {}).get("source_toc_skipped_paragraphs", 0) >= 4, "unpaged source TOC block was not skipped")
    assert_true("二、 问题分析" not in headings and "三、 模型假设" not in headings, f"unpaged TOC entries leaked as sections: {headings}")
    assert_true("一、 问题重述" in headings, f"real repeated heading was lost: {headings}")
    assert_true("real body paragraph" in body_text, f"body text after repeated heading was lost: {body_text}")


@case
def content_parser_skips_source_toc_with_duplicate_formula_fragments_before_boundary() -> None:
    work = new_workdir("source_toc_duplicate_fragments")
    docx = work / "source_toc_duplicate_fragments.docx"
    doc = Document()
    doc.add_paragraph("论文题目: Duplicate Fragment TOC")
    doc.add_paragraph("目 录")
    doc.add_paragraph("一、 问题重述")
    doc.add_paragraph("二、 问题分析")
    doc.add_paragraph("41 .5")
    doc.add_paragraph("41 .5")
    doc.add_paragraph("九、 问题五")
    doc.add_paragraph("十、 灵敏度分析")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("一、 问题重述")
    doc.add_paragraph("This is the real body paragraph.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    assert_true(content.get("_meta", {}).get("source_toc_skipped_paragraphs", 0) >= 8, "source TOC with duplicate fragments was not skipped to boundary")
    assert_true("九、 问题五" not in headings and "十、 灵敏度分析" not in headings, f"late TOC entries leaked: {headings}")
    assert_true("一、 问题重述" in headings and "real body paragraph" in body_text, f"real body after source TOC was lost: {headings} / {body_text}")


@case
def content_parser_preserves_real_directory_section_content() -> None:
    work = new_workdir("real_directory_section")
    docx = work / "real_directory_section.docx"
    doc = Document()
    doc.add_paragraph("论文题目: Directory Section")
    directory_heading = doc.add_paragraph("目录")
    directory_heading.style = doc.styles["Heading 1"]
    doc.add_paragraph("本节介绍产品目录与数据字典，不是源文档自动目录。")
    doc.add_paragraph("1 正文")
    doc.add_paragraph("This paragraph must remain in the document.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    assert_true(not content.get("_meta", {}).get("source_toc_skipped_paragraphs"), "real directory section was treated as a source TOC")
    assert_true("产品目录与数据字典" in body_text, f"real directory section content was lost: {body_text}")
    assert_true("must remain" in body_text, f"body text after real directory section was lost: {body_text}")


@case
def content_parser_preserves_real_directory_section_before_page_break() -> None:
    work = new_workdir("real_directory_section_page_break")
    docx = work / "real_directory_section_page_break.docx"
    doc = Document()
    doc.add_paragraph("论文题目: Directory Section With Page Break")
    directory_heading = doc.add_paragraph("目录")
    directory_heading.style = doc.styles["Heading 1"]
    subheading = doc.add_paragraph("一、 产品目录")
    subheading.style = doc.styles["Heading 2"]
    doc.add_paragraph("本节介绍产品目录与数据字典，不是源文档自动目录。")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("1 正文")
    doc.add_paragraph("This paragraph must remain after the page break.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    assert_true(not content.get("_meta", {}).get("source_toc_skipped_paragraphs"), "real directory section before a page break was treated as source TOC")
    assert_true("目录" in headings and "一、 产品目录" in headings, f"real directory headings were lost: {headings}")
    assert_true("产品目录与数据字典" in body_text and "must remain" in body_text, f"real directory/page-break content was lost: {body_text}")


@case
def qa_reports_toc_pollution_placeholders_and_formula_fragments() -> None:
    work = new_workdir("qa_semantic_guards")
    docx = work / "out.docx"
    doc = Document()
    doc.add_paragraph("报名序号: [报名序号]")
    doc.add_paragraph("172.04 MWh should not be a heading")
    doc.save(docx)
    content = base_content(["E", "rgreen", "RE"])
    content["_meta"]["source_placeholders"] = [{"paragraph": 1, "text": "报名序号: [报名序号]"}]
    content["sections"][0]["heading"] = "172.04 MWh should not be a heading"
    content["sections"][0]["paragraphs"] = [
        {"role": "formula", "text": "x=1(1)(1)", "numbered": True},
        {"role": "formula", "text": "吨/日对应的开机时段数n=Q/qrate∈24,21,18,15,12.产量约束为=", "numbered": True},
        "E",
        "rgreen",
        "RE",
    ]
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    (work / "build_generated.py").write_text("# synthetic\n", encoding="utf-8")
    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    for code in ["CONTENT_TOC_POLLUTION", "UNFILLED_PLACEHOLDER_TEXT", "FORMULA_NUMBER_CONFLICT", "FORMULA_TEXT_FRAGMENTED", "PLACEHOLDER_TEXT_LEFT"]:
        assert_true(code in codes, f"QA did not report {code}")
    formula_issue = next(item for item in report["issues"] if item["code"] == "FORMULA_TEXT_FRAGMENTED")
    assert_true("contaminated formula text" in formula_issue.get("detail", ""), "QA did not explain narrative text inside a formula")
    assert_true(formula_issue["severity"] == "warning", "fragmented formula text should be downgraded to warning")
    assert_true(report["passed"] is False, "semantic guard issues should fail QA")


@case
def source_omml_is_made_wps_compatible() -> None:
    xml = latex_to_omath("x=1", display=True)
    root = etree.fromstring(xml.encode("utf-8"))
    M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    for mr in root.iter(f"{{{M}}}r"):
        for rpr in list(mr.findall(f"{{{M}}}rPr")):
            mr.remove(rpr)
    raw_without_rpr = etree.tostring(root, encoding="unicode")
    content = base_content([
        {
            "role": "formula",
            "source": "omml",
            "text": "x=1",
            "math": [{"type": "display", "xml": raw_without_rpr, "text": "x=1"}],
            "numbered": True,
        }
    ])
    result = run_generated_case("source_omml_wps", content)
    conf = check_conformance(str(result["work"]), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in conf["issues"]]
    assert_true("OMML_WPS_COMPAT" not in codes, f"source OMML was not normalized for WPS: {conf['issues']}")


@case
def conformance_style_check_ignores_static_toc_lines() -> None:
    content = base_content(
        [
            "Body paragraph after heading.",
        ]
    )
    content["sections"] = [
        {
            "heading": "2 Custom Chapter",
            "level": 1,
            "role": "body",
            "paragraphs": ["Body paragraph after heading."],
            "images": [],
        }
    ]
    result = run_generated_case("conformance_toc", content)
    conf = check_conformance(str(result["work"]), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in conf["issues"]]
    assert_true("STYLE_MISMATCH" not in codes, f"conformance matched TOC lines instead of body headings: {conf['issues']}")
    assert_true("docx_sections" in conf["counts"], "conformance report should name Word section count as docx_sections")
    assert_true("sections" not in conf["counts"], "ambiguous conformance count key 'sections' should not be emitted")


@case
def script_generator_honors_body_page_break_before() -> None:
    content = base_content([])
    content["sections"] = [
        {
            "heading": "1 First",
            "level": 1,
            "role": "body",
            "paragraphs": ["First body paragraph."],
            "images": [],
            "page_break_before": True,
        },
        {
            "heading": "2 Second",
            "level": 1,
            "role": "body",
            "paragraphs": ["Second body paragraph."],
            "images": [],
            "page_break_before": True,
        },
    ]
    result = run_generated_case("body_page_break_before", content)
    page_breaks = len(re.findall(r'<w:br[^>]+w:type="page"', result["xml"]))
    assert_true(page_breaks >= 1, "body page_break_before was not rendered as a page break")


@case
def script_generator_renders_chinese_backmatter_headings() -> None:
    content = base_content([])
    content["sections"] = [
        {
            "heading": "1 Body",
            "level": 1,
            "role": "body",
            "paragraphs": ["Body paragraph."],
            "images": [],
        },
        {
            "heading": "\u81f4\u8c22",
            "level": 1,
            "role": "acknowledgement",
            "paragraphs": ["Thanks paragraph."],
            "images": [],
        },
        {
            "heading": "\u9644\u5f55",
            "level": 1,
            "role": "appendix",
            "paragraphs": ["Appendix paragraph."],
            "images": [],
        },
    ]
    result = run_generated_case("backmatter_unicode_headings", content)
    xml_compact = re.sub(r"\s+", "", result["xml"])
    assert_true("\u81f4\u8c22" in xml_compact, "acknowledgement heading was not rendered as Chinese text")
    assert_true("\u9644\u5f55" in xml_compact, "appendix heading was not rendered as Chinese text")
    codes = [item["code"] for item in result["report"]["issues"]]
    assert_true("CONTENT_HEADING_MISSING" not in codes, f"backmatter headings were still reported missing: {codes}")


@case
def script_generator_renders_backmatter_rich_items() -> None:
    work = new_workdir("backmatter_item_assets")
    img_dir = work / "figures"
    img_dir.mkdir()
    (img_dir / "dot.png").write_bytes(PNG_1X1)

    content = base_content([])
    content["_meta"]["images_dir"] = str(img_dir)
    content["_meta"]["images_extracted"] = 1
    content["_meta"]["tables_count"] = 1
    content["sections"] = [
        {
            "heading": "1 Body",
            "level": 1,
            "role": "body",
            "paragraphs": ["Body paragraph."],
            "images": [],
        },
        {
            "heading": "\u9644\u5f55",
            "level": 1,
            "role": "appendix",
            "paragraphs": [
                {"role": "table", "table_rows": [["A", "B"], ["1", "2"]]},
                {"role": "formula", "source": "latex", "latex": "E=mc^2", "text": "E=mc^2", "numbered": False},
                {"role": "image", "image": "dot.png"},
            ],
            "images": ["dot.png"],
        },
    ]
    result = run_generated_case("backmatter_rich_items", content)
    counts = result["manifest"]["counts"]
    assert_true(counts["content_tables_rendered"] == 1, f"appendix table was not rendered: {counts}")
    assert_true(counts["content_formulas_rendered"] == 1, f"appendix formula was not rendered: {counts}")
    assert_true(counts["content_images_rendered"] == 1, f"appendix image was not rendered: {counts}")


@case
def md_rich_math_builds_inline_omml() -> None:
    content = base_content(
        [
            {
                "role": "rich_text",
                "text": "Alpha x^2 beta.",
                "runs": [
                    {"type": "text", "text": "Alpha "},
                    {
                        "type": "math",
                        "text": "x^2",
                        "math": [{"type": "inline", "latex": "x^2", "text": "x^2"}],
                    },
                    {"type": "text", "text": " beta."},
                ],
                "math": [{"type": "inline", "latex": "x^2", "text": "x^2"}],
            }
        ]
    )
    result = run_generated_case("md_rich_build", content)
    assert_true(result["manifest"]["counts"]["inline_formulas_rendered"] == 1, "MD rich inline math not rendered inline")
    assert_true(omath_para_count(result["xml"]) == 0, "MD rich inline math created display formula")


@case
def latex_omath_display_flag_is_honored() -> None:
    inline_xml = latex_to_omath("x^2", display=False)
    display_xml = latex_to_omath("x^2", display=True)
    assert_true("oMathPara" not in inline_xml, "inline latex_to_omath wrapped in oMathPara")
    assert_true("oMathPara" in display_xml, "display latex_to_omath did not wrap in oMathPara")


@case
def latex_omath_limit_accepts_multitoken_subscript() -> None:
    xml = latex_to_omath(r"L=\lim_{n\to\infty}\frac{1}{n}\sum_{i=1}^{n}x_i", display=True)
    assert_true("[LaTeX error" not in xml, "limit with n\\to\\infty subscript produced a LaTeX error")
    assert_true("oMathPara" in xml and "lim" in xml, "limit formula did not render as display OMML")


@case
def latex_omath_invisible_delimiters_hide_separators() -> None:
    xml = latex_to_omath(r"\frac{E_{\mathrm{total}}-E_{\mathrm{sell}}-E_{\mathrm{buy}}}{E_{\mathrm{RE}}}+\sum_{t=1}^{24}x_t", display=True)
    root = etree.fromstring(xml.encode("utf-8"))
    ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
    delimiters = root.findall(".//m:d", ns)
    assert_true(delimiters, "complex formula did not create grouped delimiter elements")
    for delim in delimiters:
        entries = delim.findall("./m:e", ns)
        if len(entries) <= 1:
            continue
        dpr = delim.find("./m:dPr", ns)
        sep = dpr.find("./m:sepChr", ns) if dpr is not None else None
        assert_true(sep is not None and sep.get(f"{{{ns['m']}}}val") == "", "invisible delimiter can render visible vertical separators")
    texts = "".join(t.text or "" for t in root.findall(".//m:t", ns))
    assert_true("t=1" in texts and "24" in texts, f"plain multi-character scripts were split incorrectly: {texts}")
    styled_xml = latex_to_omath(r"\frac{\mathrm{abc}\mathrm{def}}{x}+\frac{\mathrm{abc}+\mathrm{def}}{x}", display=True)
    styled_root = etree.fromstring(styled_xml.encode("utf-8"))
    styled_runs = []
    for run in styled_root.findall(".//m:r", ns):
        text = "".join(t.text or "" for t in run.findall("./m:t", ns))
        sty = run.find("./m:rPr/m:sty", ns)
        styled_runs.append((text, sty.get(f"{{{ns['m']}}}val") if sty is not None else None))
    assert_true(("abcdef", "p") in styled_runs, f"merged \\mathrm runs lost upright style: {styled_runs}")
    assert_true(("abc+def", "p") in styled_runs or (("abc", "p") in styled_runs and ("def", "p") in styled_runs), f"mixed-style grouped runs lost \\mathrm style: {styled_runs}")


@case
def latex_omath_keeps_literals_operators_and_brackets_upright() -> None:
    xml = latex_to_omath(r"P(t)=\max(0,PRE(t)-P_{total}(t))+\frac{x_1}{2}+\{z\}", display=True)
    root = etree.fromstring(xml.encode("utf-8"))
    ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
    upright_chars = set("0123456789()[]{}=+-*/×÷<>≤≥≈≠,:;.%")
    bad_runs = []
    variable_runs = []
    for run in root.findall(".//m:r", ns):
        text = "".join(t.text or "" for t in run.findall("./m:t", ns))
        sty = run.find("./m:rPr/m:sty", ns)
        style = sty.get(f"{{{ns['m']}}}val") if sty is not None else None
        if text in {"P", "t", "x"} and style is None:
            variable_runs.append(text)
        if any(ch in upright_chars for ch in text) and style != "p":
            bad_runs.append((text, style))
    assert_true(not bad_runs, f"formula literals/operators/brackets should be upright, got {bad_runs}")
    assert_true({"P", "t", "x"}.issubset(set(variable_runs)), f"variables should keep default math style: {variable_runs}")


@case
def template_profile_sanitizes_private_source() -> None:
    fmt = base_format(source="private_school_template.docx")
    fmt["_meta"]["assets_dir"] = r"E:\private\Templates\private_school_template_assets"
    profile = profile_format(fmt, project_root=r"E:\private")
    text = json.dumps(profile, ensure_ascii=False)
    assert_true("private_school_template" not in text, "template profile leaked source filename")
    assert_true(profile["source"]["source_ext"] == ".docx", "template profile lost source extension")
    assert_true(profile["source"]["has_assets_dir"] is True, "template profile lost assets flag")


@case
def format_extractor_stops_cover_before_spaced_abstract_heading() -> None:
    work = new_workdir("format_cover_stop_abstract")
    docx = work / "cover_stop.docx"
    doc = Document()
    doc.add_paragraph("Cover title")
    doc.add_paragraph("摘  要")
    doc.add_paragraph("Template abstract sample paragraph should not be replayed as cover.")
    doc.save(docx)

    fmt, _ = extract_docx_format(str(docx))
    cover_text = "\n".join(
        "".join(run.get("t", "") for run in el.get("r", []))
        for el in fmt.get("cover") or []
        if isinstance(el, dict)
    )
    assert_true("Cover title" in cover_text, "cover text before abstract was not extracted")
    assert_true("摘" not in cover_text and "Template abstract sample" not in cover_text, "abstract page leaked into cover extraction")


@case
def numbered_english_heading_is_not_front_title() -> None:
    cnt = {
        "sections": [
            {
                "heading": "1 Introduction",
                "level": 1,
                "role": "body",
                "paragraphs": ["Body"],
                "images": [],
            }
        ]
    }
    front = _front_matter_sections(cnt)
    assert_true(front.get("en_title") in ("", None), "numbered English body heading became English title")
    assert_true(0 not in front.get("front_indices", set()), "numbered English body heading was marked as front matter")


@case
def english_appendix_heading_is_not_front_title() -> None:
    cnt = {
        "sections": [
            {"heading": "1 Introduction", "level": 1, "role": "body", "paragraphs": ["Body"], "images": []},
            {"heading": "Appendix A Commands", "level": 1, "role": "appendix", "paragraphs": ["python run_pipeline.py"], "images": []},
        ]
    }
    front = _front_matter_sections(cnt)
    assert_true(front.get("en_title") in ("", None), "English appendix heading became English title")
    assert_true(1 not in front.get("front_indices", set()), "English appendix was marked as front matter")


@case
def script_generator_section_order_preserves_content_while_sorting_subsections() -> None:
    sections = [
        {"heading": "第1章 Intro", "level": 1, "paragraphs": ["chapter"], "images": []},
        {"heading": "1.2 Later", "level": 2, "paragraphs": ["later"], "images": []},
        {"heading": "1.2.2 Detail B", "level": 3, "paragraphs": ["b"], "images": []},
        {"heading": "1.2.1 Detail A", "level": 3, "paragraphs": ["a"], "images": []},
        {"heading": "1.1 Earlier", "level": 2, "paragraphs": ["earlier"], "images": []},
        {"heading": "第2章 Methods", "level": 1, "paragraphs": ["next"], "images": []},
    ]
    ordered = _normalize_numbered_section_order(sections)
    headings = [sec["heading"] for sec in ordered]
    assert_true(headings[:5] == ["第1章 Intro", "1.1 Earlier", "1.2 Later", "1.2.1 Detail A", "1.2.2 Detail B"], "numbered sections were not sorted safely")
    assert_true(ordered[1]["paragraphs"] == ["earlier"], "section content moved away from its heading")
    assert_true(headings[-1] == "第2章 Methods", "next numbered chapter was not preserved")


@case
def script_generator_template_rules_extract_page_header_and_flags() -> None:
    fmt = {
        "sections": [
            {
                "page_width_cm": 21.0,
                "page_height_cm": 29.7,
                "margin_top_cm": 2.5,
                "margin_bottom_cm": 2.4,
                "margin_left_cm": 2.8,
                "margin_right_cm": 2.2,
                "header": [
                    {
                        "text": "Thesis Header",
                        "alignment": "DEFAULT",
                        "runs": [{"text": "Thesis Header", "font": "Times New Roman", "size_pt": 9, "italic": True}],
                    }
                ],
            }
        ],
        "paragraphs": [
            {"text": "图1 系统结构"},
            {"text": "中文摘要不分自然段，英文题目要求大写字母，公式居中并编号，英文参考文献左对齐，参考文献悬挂缩进2字符。"},
        ],
    }
    page = _extract_page_and_header(fmt)
    rules = _infer_template_rules(fmt)
    assert_true(page["mt"] == 2.5 and page["mr"] == 2.2, "page margin extraction changed")
    assert_true(page["header"]["align"] == "CENTER", "DEFAULT header alignment was not normalized")
    assert_true(rules["cn_abstract_single_paragraph"] is True, "Chinese abstract rule was not detected")
    assert_true(rules["formula_center"] is True and rules["formula_numbered"] is True, "formula rules were not detected")
    assert_true(rules["reference_hanging_chars"] == 2.0, "reference hanging indent rule changed")


@case
def script_generator_style_profiles_apply_template_text_rules() -> None:
    body_text = "这是一个足够长的中文正文样本，用于模拟模板正文段落的自然排版特征，确保样式推断不会只依赖说明文字。"
    fmt = {
        "style_profiles": {},
        "paragraphs": [
            {
                "text": "论文正文使用宋体小四号，固定值28磅，首行缩进2字符，两端对齐。",
                "runs": [{"text": "论文正文使用宋体小四号，固定值28磅，首行缩进2字符，两端对齐。", "font": "宋体", "size_pt": 12}],
                "align": "LEFT",
            },
            {
                "text": "第1章 一级标题黑体三号居中加粗",
                "runs": [{"text": "第1章 一级标题黑体三号居中加粗", "font": "黑体", "size_pt": 16, "bold": True}],
                "align": "CENTER",
            },
            {
                "text": "图标题宋体五号居中",
                "runs": [{"text": "图标题宋体五号居中", "font": "宋体", "size_pt": 10.5}],
                "align": "CENTER",
            },
            {
                "text": "参考文献中文使用宋体小四号，悬挂缩进2字符。",
                "runs": [{"text": "参考文献中文使用宋体小四号，悬挂缩进2字符。", "font": "宋体", "size_pt": 12}],
                "align": "LEFT",
            },
            {
                "text": body_text,
                "runs": [{"text": body_text, "font": "宋体", "size_pt": 12}],
                "align": "JUSTIFY",
                "ls": 1.5,
                "indent": 0.74,
            },
        ],
    }
    profiles = _infer_style_profiles(fmt)
    assert_true(profiles["body"]["font"] == "宋体", "body font rule was not applied")
    assert_true(profiles["body"]["line_spacing_fixed_pt"] == 28.0, "fixed body line spacing rule was not applied")
    assert_true(profiles["body"]["align"] == "JUSTIFY", "body alignment rule was not normalized")
    assert_true(profiles["body"]["first_indent_cm"] > 0, "body first-line indent rule was not applied")
    assert_true(profiles["h1"]["font"] == "黑体" and profiles["h1"]["align"] == "CENTER", "h1 style inference changed")
    assert_true(profiles["figure_caption"]["size"] == 10.5 and profiles["figure_caption"]["align"] == "CENTER", "figure caption style rule changed")
    assert_true(profiles["reference"]["font"] == "宋体", "reference style rule changed")


@case
def script_generator_runtime_base_fragment_is_injected() -> None:
    assert_true("__BASE_RUNTIME__" not in RUNTIME_TEMPLATE, "base runtime placeholder leaked into generated template")
    assert_true(
        RUNTIME_TEMPLATE.lstrip().startswith("# -*- coding: utf-8 -*-"),
        "generated script header should come from base runtime",
    )
    assert_true("DATA = json.loads(__DATA_BLOB__)" in RUNTIME_TEMPLATE, "base runtime data bootstrap was not injected")
    assert_true("def apply_run_profile" in RUNTIME_TEMPLATE, "base run styling helper was not injected")
    assert_true("def add_text" in RUNTIME_TEMPLATE, "base text helper was not injected")
    assert_true("def setup_section" in RUNTIME_TEMPLATE, "base section setup helper was not injected")
    assert_true("def force_cover_headerless" in RUNTIME_TEMPLATE, "base cover cleanup helper was not injected")
    assert_true(
        RUNTIME_TEMPLATE.index("def add_text") < RUNTIME_TEMPLATE.index("def render_cover_and_declarations"),
        "base text helpers should be defined before cover rendering",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def setup_section") < RUNTIME_TEMPLATE.index("def build_document"),
        "base page helpers should be defined before build orchestration",
    )


@case
def script_generator_runtime_content_helpers_fragment_is_injected() -> None:
    assert_true(
        "__CONTENT_HELPERS_RUNTIME__" not in RUNTIME_TEMPLATE,
        "content helper runtime placeholder leaked into generated template",
    )
    assert_true("def is_front_section_index" in RUNTIME_TEMPLATE, "front-section predicate was not injected")
    assert_true("def normalize_caption" in RUNTIME_TEMPLATE, "caption normalizer was not injected")
    assert_true("def clean_text_artifacts" in RUNTIME_TEMPLATE, "text cleanup helper was not injected")
    assert_true("def clean_formula_text" in RUNTIME_TEMPLATE, "formula cleanup helper was not injected")
    assert_true("def add_caption" in RUNTIME_TEMPLATE, "caption renderer helper was not injected")
    assert_true(RUNTIME_TEMPLATE.count("def clean_text_artifacts") == 1, "text cleanup helper should be injected exactly once")
    assert_true(
        RUNTIME_TEMPLATE.index("def clean_formula_text") < RUNTIME_TEMPLATE.index("def text_formula_to_latex"),
        "formula cleanup should be defined before plain-text formula conversion",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def add_caption") < RUNTIME_TEMPLATE.index("def render_table"),
        "caption helper should be defined before table/image rendering",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def is_front_section_index") < RUNTIME_TEMPLATE.index("def render_body"),
        "front-section predicate should be defined before body rendering",
    )


@case
def script_generator_runtime_formula_fragment_is_injected() -> None:
    assert_true("__FORMULA_RUNTIME__" not in RUNTIME_TEMPLATE, "formula runtime placeholder leaked into generated template")
    assert_true("def append_inline_formula" in RUNTIME_TEMPLATE, "inline formula runtime fragment was not injected")
    assert_true("def add_rich_text_runs" in RUNTIME_TEMPLATE, "rich_text runtime fragment was not injected")


@case
def script_generator_runtime_formula_text_fragment_is_injected() -> None:
    assert_true("__FORMULA_TEXT_RUNTIME__" not in RUNTIME_TEMPLATE, "formula text runtime placeholder leaked into generated template")
    assert_true("def chapter_number_from_heading" in RUNTIME_TEMPLATE, "chapter-number helper was not injected")
    assert_true("def text_formula_to_latex" in RUNTIME_TEMPLATE, "plain-text formula conversion helper was not injected")
    assert_true("def split_formula_number" in RUNTIME_TEMPLATE, "formula-number parsing helper was not injected")
    assert_true("def formula_has_number" in RUNTIME_TEMPLATE, "formula-number predicate was not injected")
    assert_true(RUNTIME_TEMPLATE.count("def text_formula_to_latex") == 1, "plain-text formula conversion helper should be injected exactly once")
    assert_true(
        RUNTIME_TEMPLATE.index("def clean_formula_text") < RUNTIME_TEMPLATE.index("def text_formula_to_latex"),
        "plain-text formula conversion should be defined after formula text cleanup",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def text_formula_to_latex") < RUNTIME_TEMPLATE.index("def render_formula"),
        "plain-text formula conversion should be defined before formula rendering",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def chapter_number_from_heading") < RUNTIME_TEMPLATE.index("def render_body"),
        "chapter-number helper should be defined before body rendering",
    )


@case
def script_generator_runtime_formula_render_fragment_is_injected() -> None:
    assert_true("__FORMULA_RENDER_RUNTIME__" not in RUNTIME_TEMPLATE, "formula render runtime placeholder leaked into generated template")
    assert_true("def next_formula_label" in RUNTIME_TEMPLATE, "formula label counter helper was not injected")
    assert_true("def render_plain_formula" in RUNTIME_TEMPLATE, "plain formula renderer was not injected")
    assert_true("def render_formula" in RUNTIME_TEMPLATE, "formula renderer was not injected")
    assert_true(RUNTIME_TEMPLATE.count("def render_formula") == 1, "formula renderer should be injected exactly once")
    assert_true(
        RUNTIME_TEMPLATE.index("def text_formula_to_latex") < RUNTIME_TEMPLATE.index("def render_formula"),
        "formula renderer should be defined after plain-text formula conversion",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def render_formula") < RUNTIME_TEMPLATE.index("def render_paragraph_item"),
        "formula renderer should be defined before body paragraph dispatch",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("FORMULA_COUNTERS = {}") < RUNTIME_TEMPLATE.index("if __name__ == '__main__'"),
        "formula counters should be initialized before the generated-script entrypoint runs",
    )


@case
def script_generator_runtime_media_table_fragment_is_injected() -> None:
    assert_true("__MEDIA_TABLE_RUNTIME__" not in RUNTIME_TEMPLATE, "media/table runtime placeholder leaked into generated template")
    assert_true("def render_table" in RUNTIME_TEMPLATE, "table runtime fragment was not injected")
    assert_true("def render_image" in RUNTIME_TEMPLATE, "image runtime fragment was not injected")
    assert_true("def add_code_block" in RUNTIME_TEMPLATE, "code-block runtime fragment was not injected")


@case
def script_generator_runtime_references_fragment_is_injected() -> None:
    assert_true("__REFERENCES_RUNTIME__" not in RUNTIME_TEMPLATE, "references runtime placeholder leaked into generated template")
    assert_true("def render_reference_entries" in RUNTIME_TEMPLATE, "reference runtime fragment was not injected")
    assert_true("def render_backmatter_section" in RUNTIME_TEMPLATE, "backmatter runtime fragment was not injected")
    assert_true("def collect_structural_backmatter" in RUNTIME_TEMPLATE, "structural backmatter runtime fragment was not injected")


@case
def script_generator_runtime_toc_fragment_is_injected() -> None:
    assert_true("__TOC_RUNTIME__" not in RUNTIME_TEMPLATE, "TOC runtime placeholder leaked into generated template")
    assert_true("def collect_toc_entries" in RUNTIME_TEMPLATE, "TOC entry collection runtime fragment was not injected")
    assert_true("def add_toc" in RUNTIME_TEMPLATE, "TOC rendering runtime fragment was not injected")
    assert_true("def _infer_heading_pages_from_word_com" in RUNTIME_TEMPLATE, "Word COM TOC page-resolution fragment was not injected")


@case
def script_generator_runtime_build_fragment_is_injected() -> None:
    assert_true("__BUILD_RUNTIME__" not in RUNTIME_TEMPLATE, "build runtime placeholder leaked into generated template")
    assert_true("def reset_build_stats" in RUNTIME_TEMPLATE, "build stats runtime fragment was not injected")
    assert_true("def write_build_manifest" in RUNTIME_TEMPLATE, "manifest writer runtime fragment was not injected")
    assert_true("def build_document" in RUNTIME_TEMPLATE, "document build runtime fragment was not injected")
    assert_true("def main" in RUNTIME_TEMPLATE, "generated-script main runtime fragment was not injected")


@case
def script_generator_runtime_cover_fragment_is_injected() -> None:
    assert_true("__COVER_RUNTIME__" not in RUNTIME_TEMPLATE, "cover runtime placeholder leaked into generated template")
    assert_true("def render_cover_and_declarations" in RUNTIME_TEMPLATE, "cover renderer runtime fragment was not injected")
    assert_true("def render_cover_table" in RUNTIME_TEMPLATE, "cover table runtime fragment was not injected")
    assert_true("def compute_cover_skip_indices" in RUNTIME_TEMPLATE, "cover spacer runtime fragment was not injected")
    assert_true("def set_cover_cell_margins" in RUNTIME_TEMPLATE, "cover cell-margin helper should not collide with body table helper")
    assert_true(
        RUNTIME_TEMPLATE.index("def render_cover_and_declarations") < RUNTIME_TEMPLATE.index("def render_front_matter"),
        "cover runtime should be defined before front matter rendering",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def set_cell_borders") < RUNTIME_TEMPLATE.index("def add_code_block"),
        "shared border helper should be defined before media/table runtime uses it",
    )


@case
def script_generator_runtime_front_matter_fragment_is_injected() -> None:
    assert_true("__FRONT_MATTER_RUNTIME__" not in RUNTIME_TEMPLATE, "front matter runtime placeholder leaked into generated template")
    assert_true("def section_text" in RUNTIME_TEMPLATE, "front matter section_text helper was not injected")
    assert_true("def add_keywords" in RUNTIME_TEMPLATE, "front matter keyword helper was not injected")
    assert_true("def render_front_matter" in RUNTIME_TEMPLATE, "front matter renderer runtime fragment was not injected")
    assert_true(RUNTIME_TEMPLATE.count("def render_front_matter") == 1, "front matter renderer should be injected exactly once")
    assert_true(
        RUNTIME_TEMPLATE.index("def add_rich_text_item") < RUNTIME_TEMPLATE.index("def render_front_matter"),
        "front matter renderer should be defined after rich-text formula helpers",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def add_toc") < RUNTIME_TEMPLATE.index("def render_front_matter"),
        "front matter renderer should be defined after TOC helpers",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def render_front_matter") < RUNTIME_TEMPLATE.index("def is_front_section_index"),
        "front matter renderer should stay before body/front-index helpers",
    )


@case
def script_generator_runtime_body_fragment_is_injected() -> None:
    assert_true("__BODY_RUNTIME__" not in RUNTIME_TEMPLATE, "body runtime placeholder leaked into generated template")
    assert_true("def render_paragraph_item" in RUNTIME_TEMPLATE, "body paragraph dispatcher runtime fragment was not injected")
    assert_true("def render_body" in RUNTIME_TEMPLATE, "body renderer runtime fragment was not injected")
    assert_true(RUNTIME_TEMPLATE.count("def render_body") == 1, "body renderer should be injected exactly once")
    assert_true(
        RUNTIME_TEMPLATE.index("def render_formula") < RUNTIME_TEMPLATE.index("def render_paragraph_item"),
        "body renderer should be defined after formula rendering helpers",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def render_table") < RUNTIME_TEMPLATE.index("def render_paragraph_item"),
        "body renderer should be defined after table/image helpers",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def render_reference_entries") < RUNTIME_TEMPLATE.index("def render_body"),
        "body renderer should be defined after reference/backmatter helpers",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def render_body") < RUNTIME_TEMPLATE.index("def build_document"),
        "body renderer should be defined before document build orchestration",
    )


@case
def privacy_sanitizes_absolute_paths() -> None:
    data = {
        "path": r"X:\workspace\project\Outputs\private\file.docx",
        "tmp": str(Path(tempfile.gettempdir()) / "abc" / "file.pdf"),
    }
    sanitized = sanitize_value(data, project_root=r"X:\workspace\project")
    text = json.dumps(sanitized, ensure_ascii=False)
    assert_true(r"X:\workspace" not in text and "project" not in text, "project path leaked")
    assert_true("<PROJECT>" in text and "<TEMP>" in text, "path labels missing")


@case
def visual_sample_pages_pick_useful_pages() -> None:
    import qa_visual

    pages = ["cover", "contents", "blank", "1. Introduction", "middle"] + ["body"] * 7
    samples = qa_visual._sample_pages(12, pages)
    assert_true(1 in samples and 2 in samples and 4 in samples and 6 in samples, "sample page selection missed key pages")


@case
def visual_qa_fails_closed_without_pdf_tools() -> None:
    import qa_visual

    work = new_workdir("visual_closed")
    (work / "final.docx").write_bytes(b"not a real docx; export is monkeypatched")
    fake_pdf = work / "fake.pdf"
    fake_pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")
    original_export = qa_visual._export_pdf
    original_pdfinfo = qa_visual._pdfinfo
    original_pages_text = qa_visual._pdf_pages_text
    original_render = qa_visual._render_samples
    try:
        qa_visual._export_pdf = lambda _docx, _visual_dir: str(fake_pdf)
        qa_visual._pdfinfo = lambda _pdf: {"available": False}
        qa_visual._pdf_pages_text = lambda _pdf, _visual_dir: []
        qa_visual._render_samples = lambda _pdf, _visual_dir, _pages: []
        report = qa_visual.check_visual(str(work), output_docx_name="final.docx")
    finally:
        qa_visual._export_pdf = original_export
        qa_visual._pdfinfo = original_pdfinfo
        qa_visual._pdf_pages_text = original_pages_text
        qa_visual._render_samples = original_render
    codes = [item["code"] for item in report["issues"]]
    assert_true(report["passed"] is False, "visual QA passed without pdfinfo/text validation")
    assert_true("PDFINFO_UNAVAILABLE" in codes, "missing pdfinfo was not reported")


@case
def sample_pages_empty_when_page_count_unknown() -> None:
    import qa_visual

    assert_true(qa_visual._sample_pages(0, []) == [], "sample pages should be empty for unknown page count")


@case
def run_pipeline_missing_inputs_returns_nonzero() -> None:
    root = PIPELINE_DIR.parents[2]
    result = subprocess.run(
        [
            sys.executable,
            str(root / "run_pipeline.py"),
            "--template",
            "__missing_template__.docx",
            "--content",
            "__missing_content__.docx",
        ],
        cwd=str(root),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=60,
    )
    assert_true(result.returncode != 0, "run_pipeline returned success for missing inputs")


@case
def qa_checker_cli_failure_returns_nonzero() -> None:
    work = new_workdir("qa_cli_nonzero")
    content = base_content([])
    content["_meta"]["missing_images"] = [{"source": "missing.png", "reason": "not_found"}]
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    result = subprocess.run(
        [
            sys.executable,
            str(PIPELINE_DIR / "qa_checker.py"),
            str(work),
            "--mode",
            "developer",
            "--docx",
            "out.docx",
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=60,
    )
    assert_true(result.returncode != 0, "qa_checker CLI returned success for a failed report")


@case
def qa_missing_image_detail_is_sanitized() -> None:
    work = new_workdir("qa_missing_image_privacy")
    private_path = str(work / "private" / "missing.png")
    content = base_content([])
    content["_meta"]["missing_images"] = [{"source": private_path, "reason": "not_found"}]
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    detail = "\n".join(str(item.get("detail") or "") for item in report["issues"])
    assert_true(private_path not in detail, "QA leaked an absolute missing-image path")
    assert_true("<TEMP>" in detail or "<ABS_PATH>" in detail or "<PROJECT>" in detail, "QA missing-image detail was not sanitized")


def run_cases(selected: str | None = None) -> int:
    passed = 0
    failed = 0
    matched = 0
    for fn in CASES:
        if selected and selected not in fn.__name__:
            continue
        matched += 1
        try:
            fn()
            print(f"PASS {fn.__name__}")
            passed += 1
        except Exception as exc:
            failed += 1
            print(f"FAIL {fn.__name__}: {exc}")
    if selected and matched == 0:
        print(f"RESULT passed=0 failed=1")
        print(f"FAIL no tests matched filter: {selected}")
        return 1
    print(f"RESULT passed={passed} failed={failed}")
    return 1 if failed else 0


def main() -> None:
    global KEEP_ARTIFACTS
    parser = argparse.ArgumentParser(description="Run synthetic regression checks for the Word pipeline.")
    parser.add_argument("--keep", action="store_true", help="Keep temporary test artifacts.")
    parser.add_argument("--filter", default=None, help="Run cases whose function name contains this text.")
    args = parser.parse_args()
    KEEP_ARTIFACTS = bool(args.keep)
    try:
        code = run_cases(args.filter)
    finally:
        if not KEEP_ARTIFACTS and "code" in locals() and code == 0:
            for path in TEMP_DIRS:
                shutil.rmtree(path, ignore_errors=True)
        elif TEMP_DIRS:
            print("ARTIFACTS")
            for path in TEMP_DIRS:
                print(path)
    raise SystemExit(code)


if __name__ == "__main__":
    main()
