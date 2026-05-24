"""
script_generator.py — role-driven DOCX build script generator.

输入: format.json + content.json
输出: build_generated.py；运行该脚本生成最终论文 docx。

修复点：
- 封面按模板元素重建，支持校徽/图片 assets，学位/学校编码行左对齐。
- 中文摘要、英文摘要、目录、正文使用独立 role；英文摘要另起一页，正文在目录后另起一节。
- 标题设置 outline level，编号和标题之间自动补一个空格。
- 正文 [N] / [1,2] / [1-3] 引用自动上标。
- 表格渲染为三线表，保留单元格换行。
- 图/表题走 caption role，居中并规范“图 1 标题”的空格。
- 代码块保留换行和空格，使用等宽字体。
- 参考文献单条一个段落，悬挂缩进，中文宋体/英文 Times New Roman 混排。
"""
from __future__ import annotations

import json
import os
import re
import shutil
from typing import Any, Dict, List, Optional


def _is_cjk(text: str) -> bool:
    return any('\u4e00' <= c <= '\u9fff' for c in str(text or ''))


def _ascii_ratio(text: str) -> float:
    text = str(text or '')
    if not text:
        return 0.0
    return sum(1 for c in text if c.isascii() and c.isalpha()) / max(len(text), 1)


_CN_SIZE_PATTERNS = [
    ('小二', 18.0), ('二号', 22.0),
    ('小三', 15.0), ('三号', 16.0),
    ('小四', 12.0), ('四号', 14.0),
    ('小五', 9.0), ('五号', 10.5),
]


def _text_blob(fmt: Dict[str, Any]) -> str:
    return '\n'.join(str(p.get('text') or '') for p in fmt.get('paragraphs') or [])


def _find_instruction(texts: str, *needles: str) -> str:
    """Return a compact instruction line/paragraph that contains all needles."""
    chunks = re.split(r'[\r\n]+', texts)
    chunks += re.split(r'[。；;]\s*', texts)
    for chunk in chunks:
        if all(n in chunk for n in needles):
            return chunk.strip()
    return ''


def _find_regex_instruction(texts: str, pattern: str) -> str:
    m = re.search(pattern, texts, re.S)
    return m.group(1).strip() if m else ''


def _size_from_text(text: str, default: Optional[float] = None) -> Optional[float]:
    text = str(text or '')
    m = re.search(r'(\d+(?:\.\d+)?)\s*pt', text, re.I)
    if m:
        return float(m.group(1))
    for name, size in _CN_SIZE_PATTERNS:
        if name in text:
            return size
    return default


def _font_from_text(text: str, default: Optional[str] = None) -> Optional[str]:
    text = str(text or '')
    if 'timesnewroman' in re.sub(r'\s+', '', text).lower():
        return 'Times New Roman'
    for font in ('Times New Roman', '黑体', '宋体', '楷体_GB2312', '楷体', '仿宋', '微软雅黑', '华文中宋', '方正小标宋简体'):
        if font.lower() in text.lower():
            return font
    return default


def _align_from_text(text: str, default: Optional[str] = None) -> Optional[str]:
    if '居中' in text:
        return 'CENTER'
    if '右对齐' in text or '靠右' in text:
        return 'RIGHT'
    if '左对齐' in text or '靠左' in text or '左侧' in text:
        return 'LEFT'
    if '两端对齐' in text or '右侧也要对齐' in text:
        return 'JUSTIFY'
    return default


def _line_spacing_from_text(text: str) -> Dict[str, Any]:
    text = str(text or '')
    m = re.search(r'固定值\s*(\d+(?:\.\d+)?)\s*磅', text)
    if m:
        v = float(m.group(1))
        return {'line_spacing_val': v, 'line_spacing_rule': 'exact', 'line_spacing_fixed_pt': v}
    if re.search(r'1\.5\s*倍|1\.5倍', text):
        return {'line_spacing_val': 1.5, 'line_spacing_rule': 'auto', 'line_spacing_fixed_pt': None}
    if '单倍' in text:
        return {'line_spacing_val': 1.0, 'line_spacing_rule': 'auto', 'line_spacing_fixed_pt': None}
    return {}


def _spacing_before_after_from_text(text: str, line_pt: Optional[float] = None) -> Dict[str, Any]:
    text = str(text or '')
    out: Dict[str, Any] = {}
    m = re.search(r'段前段后各?\s*(\d+(?:\.\d+)?)\s*磅', text)
    if m:
        v = float(m.group(1))
        out['space_before_pt'] = v; out['space_after_pt'] = v
        return out
    m = re.search(r'段前段后\s*(\d+(?:\.\d+)?)\s*磅', text)
    if m:
        v = float(m.group(1))
        out['space_before_pt'] = v; out['space_after_pt'] = v
        return out
    if re.search(r'段前段后各?\s*1\s*行', text):
        v = float(line_pt or 28.0)
        out['space_before_pt'] = v; out['space_after_pt'] = v
    elif re.search(r'段前段后\s*0(?:\.5|点5|半)\s*行', text):
        v = float(line_pt or 28.0) * 0.5
        out['space_before_pt'] = v; out['space_after_pt'] = v
    elif re.search(r'段前段后\s*0\s*行|段前段后0', text):
        out['space_before_pt'] = 0.0; out['space_after_pt'] = 0.0
    return out


def _indent_from_text(text: str, size_pt: Optional[float] = None, default: Optional[float] = None) -> Optional[float]:
    text = str(text or '')
    m = re.search(r'缩进\s*(\d+(?:\.\d+)?)\s*(?:个)?(?:汉)?字(?:符)?', text)
    if m:
        chars = float(m.group(1))
        return round(chars * float(size_pt or 12.0) * 0.0352778, 2)
    return default


def _profile_from_instruction(text: str, base: Dict[str, Any], **defaults: Any) -> Dict[str, Any]:
    prof = dict(base)
    prof.update(defaults)
    font = _font_from_text(text)
    size = _size_from_text(text)
    align = _align_from_text(text)
    if font: prof['font'] = font
    if size: prof['size'] = size
    if align: prof['align'] = align
    if '加粗' in text:
        prof['bold'] = True
    if '不加粗' in text:
        prof['bold'] = False
    prof.update(_line_spacing_from_text(text))
    line_pt = prof.get('line_spacing_fixed_pt') or (float(prof.get('size') or 12) * float(prof.get('line_spacing_val') or 1.5))
    prof.update(_spacing_before_after_from_text(text, line_pt))
    ind = _indent_from_text(text, prof.get('size'))
    if ind is not None:
        prof['first_indent_cm'] = ind
    return _normalize_profile(prof, base)


def _first_run(p: Dict[str, Any]) -> Dict[str, Any]:
    runs = [r for r in (p.get('runs') or []) if r.get('text', '').strip() or r.get('size_pt')]
    if not runs:
        return (p.get('runs') or [{}])[0] if p.get('runs') else {}

    def score(r: Dict[str, Any]) -> float:
        txt = r.get('text', '') or ''
        font = r.get('font', '') or ''
        size = float(r.get('size_pt') or 0)
        return size + (5 if _is_cjk(txt) else 0) + (3 if font and font not in ('Arial', 'Times New Roman', 'Calibri') else 0)

    return max(runs, key=score)


def _first_text_run(p: Dict[str, Any]) -> Dict[str, Any]:
    for r in p.get('runs') or []:
        if str(r.get('text') or '').strip():
            return r
    return _first_run(p)


def _profile_from_para_first_text(p: Dict[str, Any], fallback: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    r = _first_text_run(p)
    prof = {
        'font': r.get('font') or (fallback or {}).get('font') or '宋体',
        'size': r.get('size_pt') or (fallback or {}).get('size') or 12,
        'bold': bool(r.get('bold', (fallback or {}).get('bold', False))),
        'italic': bool(r.get('italic', (fallback or {}).get('italic', False))),
        'align': p.get('alignment') or p.get('align') or (fallback or {}).get('align') or 'LEFT',
        'line_spacing_val': p.get('line_spacing_val') if p.get('line_spacing_val') is not None else p.get('ls', (fallback or {}).get('line_spacing_val')),
        'line_spacing_rule': p.get('line_spacing_rule') or (fallback or {}).get('line_spacing_rule'),
        'line_spacing_fixed_pt': p.get('line_spacing_fixed_pt') or (fallback or {}).get('line_spacing_fixed_pt'),
        'space_before_pt': p.get('space_before_pt') if p.get('space_before_pt') is not None else (fallback or {}).get('space_before_pt', 0),
        'space_after_pt': p.get('space_after_pt') if p.get('space_after_pt') is not None else (fallback or {}).get('space_after_pt', 0),
        'first_indent_cm': p.get('first_indent_cm') if p.get('first_indent_cm') is not None else p.get('indent', (fallback or {}).get('first_indent_cm', 0)),
    }
    return _normalize_profile(prof, fallback)


def _normalize_profile(prof: Optional[Dict[str, Any]], fallback: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    fallback = fallback or {}
    p = dict(fallback)
    if prof:
        for k, v in prof.items():
            if v is not None or k in ('line_spacing_fixed_pt',):
                p[k] = v
    p.setdefault('font', '宋体')
    p.setdefault('size', 12)
    p.setdefault('bold', False)
    p.setdefault('italic', False)
    p.setdefault('align', 'LEFT')
    if p.get('align') == 'DEFAULT':
        p['align'] = fallback.get('align', 'LEFT')
    p.setdefault('line_spacing_val', 1.5)
    p.setdefault('line_spacing_fixed_pt', None)
    p.setdefault('space_before_pt', 0)
    p.setdefault('space_after_pt', 0)
    p.setdefault('first_indent_cm', 0)
    try:
        p['size'] = float(p.get('size') or fallback.get('size') or 12)
    except Exception:
        p['size'] = 12.0
    for k in ('space_before_pt', 'space_after_pt', 'first_indent_cm'):
        try:
            p[k] = float(p.get(k) or 0)
        except Exception:
            p[k] = 0.0
    try:
        if p.get('line_spacing_fixed_pt') is not None:
            p['line_spacing_fixed_pt'] = float(p.get('line_spacing_fixed_pt'))
    except Exception:
        p['line_spacing_fixed_pt'] = None
    try:
        if p.get('line_spacing_val') is not None:
            p['line_spacing_val'] = float(p.get('line_spacing_val'))
    except Exception:
        p['line_spacing_val'] = fallback.get('line_spacing_val', 1.5)
    if p.get('line_spacing_val') and p.get('line_spacing_val') > 20 and not p.get('line_spacing_fixed_pt'):
        p['line_spacing_val'] = fallback.get('line_spacing_val', 1.5)
    return p


def _profile_from_para(p: Dict[str, Any], fallback: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    r = _first_run(p)
    prof = {
        'font': r.get('font') or (fallback or {}).get('font') or '宋体',
        'size': r.get('size_pt') or (fallback or {}).get('size') or 12,
        'bold': bool(r.get('bold', (fallback or {}).get('bold', False))),
        'italic': bool(r.get('italic', (fallback or {}).get('italic', False))),
        'align': p.get('alignment') or p.get('align') or (fallback or {}).get('align') or 'LEFT',
        'line_spacing_val': p.get('line_spacing_val') if p.get('line_spacing_val') is not None else p.get('ls', (fallback or {}).get('line_spacing_val')),
        'line_spacing_rule': p.get('line_spacing_rule') or (fallback or {}).get('line_spacing_rule'),
        'line_spacing_fixed_pt': p.get('line_spacing_fixed_pt') or (fallback or {}).get('line_spacing_fixed_pt'),
        'space_before_pt': p.get('space_before_pt') if p.get('space_before_pt') is not None else (fallback or {}).get('space_before_pt', 0),
        'space_after_pt': p.get('space_after_pt') if p.get('space_after_pt') is not None else (fallback or {}).get('space_after_pt', 0),
        'first_indent_cm': p.get('first_indent_cm') if p.get('first_indent_cm') is not None else p.get('indent', (fallback or {}).get('first_indent_cm', 0)),
    }
    return _normalize_profile(prof, fallback)


def _infer_style_profiles(fmt: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    existing = {k: _normalize_profile(v) for k, v in (fmt.get('style_profiles') or {}).items()}
    paras = fmt.get('paragraphs') or []
    profiles: Dict[str, Dict[str, Any]] = {}

    def put(role: str, p: Dict[str, Any]) -> None:
        if p and role not in profiles:
            profiles[role] = _profile_from_para(p)

    for p in paras:
        txt = (p.get('text') or '').strip()
        if not txt:
            continue
        compact = re.sub(r'\s+', '', txt)
        up = txt.upper()
        if ('论文' in txt and '题目' in txt and len(txt) < 100 and ('居中' in txt or p.get('style') == '论文题目')):
            put('cn_title', p)
        if re.match(r'^摘\s*要(?:[（(]|$)', txt) and len(txt) < 40:
            put('cn_abstract_heading', p)
        if txt.startswith('摘要是') or ('中文摘要' in txt and len(txt) > 60):
            put('cn_abstract_body', p)
        if txt.startswith('关键词'):
            put('cn_keywords', p)
        if '英文题目' in txt and ('Times' in txt or 'Roman' in txt):
            profiles.setdefault('en_title', _profile_from_para_first_text(p))
        if up.startswith('ABSTRACT') and len(txt) < 40:
            profiles.setdefault('en_abstract_heading', _profile_from_para_first_text(p))
        if len(txt) > 80 and _ascii_ratio(txt[:160]) > 0.55:
            put('en_abstract_body', p)
        if up.startswith('KEY WORD') or up.startswith('KEYWORDS'):
            profiles.setdefault('en_keywords', _profile_from_para_first_text(p))
        if compact in ('目录', '目目录') or compact.startswith('目录'):
            put('toc_title', p)
        if ('一级标题' in txt or re.match(r'^第[一二三四五六七八九十\d]+章\s+', txt)) and len(txt) < 100:
            put('h1', p)
        if ('二级标题' in txt or re.match(r'^\d+\.\d+\s+', txt)) and len(txt) < 100:
            put('h2', p)
        if ('三级标题' in txt or re.match(r'^\d+\.\d+\.\d+\s+', txt)) and len(txt) < 100:
            put('h3', p)
        if re.match(r'^(图|表)\s*\d+', txt) and len(txt) < 80:
            put('figure_caption' if txt.startswith('图') else 'table_caption', p)
        if '参考文献' in txt and len(txt) < 40:
            put('reference_heading', p)

    def heading_score(p: Dict[str, Any], level: int) -> float:
        txt = (p.get('text') or '').strip()
        r = _first_run(p)
        font = r.get('font') or ''
        size = float(r.get('size_pt') or 0)
        score = size
        if p.get('style', '').lower().startswith('heading'):
            score += 20
        if '\t' in txt or re.search(r'\s\d+$', txt):
            score -= 18
        if font and font not in ('Arial', 'Times New Roman', 'Calibri'):
            score += 8
        if level == 1 and re.match(r'^第[一二三四五六七八九十\d]+章\s+', txt):
            score += 10
        if level == 2 and re.match(r'^\d+\.\d+\s+', txt):
            score += 10
        if level == 3 and re.match(r'^\d+\.\d+\.\d+\s+', txt):
            score += 10
        return score

    for role, level, pat in [
        ('h1', 1, r'^第[一二三四五六七八九十\d]+章\s+'),
        ('h2', 2, r'^\d+\.\d+\s+'),
        ('h3', 3, r'^\d+\.\d+\.\d+\s+'),
    ]:
        cands = [p for p in paras if re.match(pat, (p.get('text') or '').strip()) and len((p.get('text') or '').strip()) < 100]
        if cands:
            profiles[role] = _profile_from_para(max(cands, key=lambda x: heading_score(x, level)))
            profiles[role]['bold'] = True
            profiles[role]['first_indent_cm'] = 0
            if level == 1:
                profiles[role]['align'] = 'CENTER'

    body_cands = []
    for p in paras:
        txt = (p.get('text') or '').strip()
        if len(txt) < 80:
            continue
        if any(k in txt[:80] for k in ['本人郑重声明', '本人在导师', '原创性声明', '版权使用', '格式要求', '字体要求', '行距：', '字号', '页眉页脚']):
            continue
        if _is_cjk(txt):
            body_cands.append(p)
    if body_cands:
        put('body', body_cands[0])

    for role, prof in existing.items():
        profiles.setdefault(role, prof)

    body = _normalize_profile(profiles.get('body') or {'font': '宋体', 'size': 12, 'align': 'JUSTIFY', 'line_spacing_fixed_pt': 28, 'first_indent_cm': 0.74})
    profiles['body'] = body
    profiles.setdefault('h1', _normalize_profile({'font': '黑体', 'size': 16, 'bold': True, 'align': 'CENTER', 'line_spacing_fixed_pt': body.get('line_spacing_fixed_pt'), 'first_indent_cm': 0}, body))
    profiles.setdefault('h2', _normalize_profile({'font': '黑体', 'size': 14, 'bold': True, 'align': 'LEFT', 'line_spacing_fixed_pt': body.get('line_spacing_fixed_pt'), 'first_indent_cm': 0}, body))
    profiles.setdefault('h3', _normalize_profile({'font': '黑体', 'size': 12, 'bold': True, 'align': 'LEFT', 'line_spacing_fixed_pt': body.get('line_spacing_fixed_pt'), 'first_indent_cm': 0}, body))
    profiles.setdefault('cn_title', profiles['h1'])
    profiles.setdefault('cn_abstract_heading', profiles['h1'])
    profiles.setdefault('cn_abstract_body', body)
    profiles.setdefault('cn_keywords', _normalize_profile({'first_indent_cm': 0, 'bold': False}, body))
    profiles.setdefault('en_title', _normalize_profile({'font': 'Times New Roman', 'bold': True, 'align': 'CENTER'}, profiles['h1']))
    profiles.setdefault('en_abstract_heading', _normalize_profile({'font': 'Times New Roman', 'size': 16, 'bold': True, 'align': 'CENTER', 'first_indent_cm': 0}, profiles['h1']))
    profiles.setdefault('en_abstract_body', _normalize_profile({'font': 'Times New Roman', 'line_spacing_fixed_pt': None, 'line_spacing_val': 1.5, 'first_indent_cm': 0.9, 'align': 'JUSTIFY'}, body))
    profiles.setdefault('en_keywords', _normalize_profile({'font': 'Times New Roman', 'bold': False, 'first_indent_cm': 0, 'align': 'LEFT'}, profiles['en_abstract_body']))
    profiles.setdefault('toc_title', profiles['h1'])
    profiles.setdefault('figure_caption', _normalize_profile({'font': '宋体', 'size': 10.5, 'align': 'CENTER', 'first_indent_cm': 0, 'space_before_pt': 6, 'space_after_pt': 6, 'line_spacing_fixed_pt': 28}, body))
    profiles.setdefault('table_caption', dict(profiles['figure_caption']))
    profiles.setdefault('table_body', _normalize_profile({'font': '宋体', 'size': 10.5, 'align': 'CENTER', 'first_indent_cm': 0, 'line_spacing_fixed_pt': None, 'line_spacing_val': 1.0}, body))
    profiles.setdefault('table_header', _normalize_profile({'bold': True}, profiles['table_body']))
    profiles.setdefault('formula', _normalize_profile({'align': 'CENTER', 'first_indent_cm': 0}, body))
    profiles.setdefault('code', _normalize_profile({
        # No template-specific code style found: inherit the thesis body profile
        # so configuration examples obey the same font size and line spacing as
        # the paper instead of falling back to IDE-style monospace formatting.
        'font': body.get('font') or '宋体',
        'size': body.get('size') or 12,
        'align': 'LEFT',
        'first_indent_cm': 0,
        'line_spacing_fixed_pt': body.get('line_spacing_fixed_pt'),
        'line_spacing_val': body.get('line_spacing_val'),
        'line_spacing_rule': body.get('line_spacing_rule'),
        'space_before_pt': 0,
        'space_after_pt': 0,
    }, body))
    profiles.setdefault('reference', _normalize_profile({'font': '宋体', 'size': 12, 'align': 'JUSTIFY', 'first_indent_cm': 0, 'space_before_pt': 6, 'space_after_pt': 6, 'line_spacing_fixed_pt': 28}, body))
    profiles.setdefault('reference_heading', profiles['h1'])
    return _apply_template_text_rules(fmt, profiles)


def _apply_template_text_rules(fmt: Dict[str, Any], profiles: Dict[str, Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    """Apply prose rules found in the template, without school-specific constants."""
    texts = _text_blob(fmt)
    body = _normalize_profile(profiles.get('body') or {})
    body_rule = (
        _find_regex_instruction(texts, r'摘要内容为([^。；;\n]*)') or
        _find_regex_instruction(texts, r'论文正文[^。；;\n]*?(宋体[^。；;\n]*?(?:固定值|行距)[^。；;\n]*)') or
        _find_instruction(texts, '正文')
    )
    if body_rule:
        body = _profile_from_instruction(body_rule, body, align='JUSTIFY')
    abstract_body_rule = _find_regex_instruction(texts, r'摘要内容为([^。；;\n]*)') or _find_instruction(texts, '摘要内容')
    if abstract_body_rule and ('宋体' in abstract_body_rule or '小四' in abstract_body_rule):
        body = _profile_from_instruction(abstract_body_rule, body, align='JUSTIFY')
    body.setdefault('font', '宋体')
    body.setdefault('size', 12.0)
    body.setdefault('align', 'JUSTIFY')
    body.setdefault('line_spacing_fixed_pt', 28.0)
    body.setdefault('line_spacing_val', 28.0)
    body.setdefault('line_spacing_rule', 'exact')
    profiles['body'] = body

    def role(name: str, **kw: Any) -> None:
        base = dict(body)
        base.update(kw)
        profiles[name] = _normalize_profile(base, body)

    cn_title_rule = (
        _find_regex_instruction(texts, r'(?:毕业)?论文(?:（设计）)?题目为([^。；;\n]*黑体[^。；;\n]*)') or
        _find_instruction(texts, '毕业论文', '题目', '黑体') or
        _find_instruction(texts, '论文题目', '黑体', '居中')
    )
    if cn_title_rule:
        profiles['cn_title'] = _profile_from_instruction(cn_title_rule, body, first_indent_cm=0.0)
    cn_abs_head_rule = _find_regex_instruction(texts, r'[“"]?摘要[”"]?为([^。；;\n]*)') or _find_instruction(texts, '摘要', '居中')
    if cn_abs_head_rule:
        profiles['cn_abstract_heading'] = _profile_from_instruction(cn_abs_head_rule, body, first_indent_cm=0.0)
    if abstract_body_rule:
        profiles['cn_abstract_body'] = _profile_from_instruction(abstract_body_rule, body, first_indent_cm=_indent_from_text(abstract_body_rule, 12.0, body.get('first_indent_cm')))
    kw_rule = _find_instruction(texts, '关键词')
    if kw_rule:
        profiles['cn_keywords'] = _profile_from_instruction(kw_rule, body, align='LEFT', first_indent_cm=0.0)

    en_rule = _find_regex_instruction(texts, r'英文标题和摘要([^。；;\n]*)') or _find_instruction(texts, '英文标题', '摘要') or _find_instruction(texts, '英文题目', '摘要')
    en_title_rule = _find_regex_instruction(texts, r'论文题目为([^。；;\n]*Times\s*New\s*Roman[^。；;\n]*)') or _find_instruction(texts, '英文题目') or en_rule
    if en_title_rule:
        profiles['en_title'] = _profile_from_instruction(en_title_rule, body, font='Times New Roman', align='CENTER', first_indent_cm=0.0)
    if en_rule:
        en_spacing = _line_spacing_from_text(en_rule)
        en_spacing.update(_spacing_before_after_from_text(en_rule, float(body.get('size') or 12) * 1.5))
        if profiles.get('en_title'):
            profiles['en_title'] = _normalize_profile({**profiles['en_title'], **en_spacing}, profiles['en_title'])
        en_body = _normalize_profile({'font': 'Times New Roman', 'size': body.get('size', 12), 'bold': False, 'align': 'JUSTIFY', 'first_indent_cm': body.get('first_indent_cm'), **en_spacing}, body)
        profiles['en_abstract_body'] = en_body
        profiles['en_abstract_heading'] = _normalize_profile({'font': 'Times New Roman', 'size': profiles.get('en_title', en_body).get('size', 16), 'bold': False, 'align': 'CENTER', 'first_indent_cm': 0.0, **en_spacing}, en_body)
        profiles['en_keywords'] = _normalize_profile({'font': 'Times New Roman', 'size': body.get('size', 12), 'bold': False, 'align': 'LEFT', 'first_indent_cm': 0.0, **en_spacing}, en_body)

    toc_title_rule = _find_instruction(texts, '目录', '黑体') or _find_instruction(texts, '【目录】')
    if toc_title_rule:
        profiles['toc_title'] = _profile_from_instruction(toc_title_rule, body, first_indent_cm=0.0)
    toc_rule = _find_regex_instruction(texts, r'中文：([^。；;\n]*一级标题[^。；;\n]*)') or _find_instruction(texts, '一级标题', '二级', '三级', '目录') or _find_instruction(texts, '一级标题', '宋体', '四号')
    if toc_rule:
        role('toc1', font=_font_from_text(toc_rule, '宋体'), size=14.0, bold=bool(re.search(r'一级标题[^。；;\n]*加粗', toc_rule)), align='LEFT', first_indent_cm=0.0)
        role('toc2', font=_font_from_text(toc_rule, '宋体'), size=12.0, bold=False, align='LEFT', first_indent_cm=0.0)
        role('toc3', font=_font_from_text(toc_rule, '宋体'), size=12.0, bold=False, align='LEFT', first_indent_cm=0.0)

    h_rules = [
        ('h1', _find_regex_instruction(texts, r'(第1章[^。；;\n]*一级标题[^。；;\n]*)') or _find_instruction(texts, '第1章', '标题') or _find_instruction(texts, '一级标题')),
        ('h2', _find_regex_instruction(texts, r'(1\.1[^。；;\n]*二级标题[^。；;\n]*)') or _find_instruction(texts, '二级标题')),
        ('h3', _find_regex_instruction(texts, r'(1\.1\.1[^。；;\n]*三级标题[^。；;\n]*)') or _find_instruction(texts, '三级标题')),
    ]
    for h_role, h_rule in h_rules:
        if h_rule:
            defaults = {
                'align': 'CENTER' if h_role == 'h1' else 'LEFT',
                'first_indent_cm': 0.0 if h_role == 'h1' else _indent_from_text(h_rule, _size_from_text(h_rule, body.get('size')), body.get('first_indent_cm')),
            }
            profiles[h_role] = _profile_from_instruction(h_rule, body, **defaults)

    fig_rule = _find_instruction(texts, '图标题') or _find_instruction(texts, '图题')
    tab_rule = _find_instruction(texts, '表标题') or _find_instruction(texts, '表题')
    table_detail_rule = _find_instruction(texts, '表内容') or _find_instruction(texts, '表格', '五号')
    if fig_rule:
        profiles['figure_caption'] = _profile_from_instruction(fig_rule, body, font='宋体', size=10.5, bold=False, align='CENTER', first_indent_cm=0.0)
    if tab_rule:
        profiles['table_caption'] = _profile_from_instruction(tab_rule, body, font='宋体', size=10.5, bold=False, align='CENTER', first_indent_cm=0.0)
    if table_detail_rule:
        table_body = _profile_from_instruction(table_detail_rule, body, font='宋体', size=10.5, align='CENTER', first_indent_cm=0.0)
        table_body['font'] = '宋体'
        if '单倍' in table_detail_rule:
            table_body.update({'line_spacing_fixed_pt': None, 'line_spacing_val': 1.0, 'line_spacing_rule': 'auto'})
        profiles['table_body'] = _normalize_profile(table_body, body)
        profiles['table_header'] = _normalize_profile({'bold': True}, profiles['table_body'])

    ref_rule = _find_regex_instruction(texts, r'参考文献中中文使用([^。；;\n]*)') or _find_instruction(texts, '参考文献', '小四') or _find_instruction(texts, '参考文献格式要求')
    if ref_rule:
        reference = _profile_from_instruction(ref_rule, body, font='宋体', size=12.0, bold=False, align='JUSTIFY', first_indent_cm=0.0, space_before_pt=0.0, space_after_pt=0.0)
        reference['font'] = '宋体'
        profiles['reference'] = reference
    if _find_instruction(texts, '英文参考文献', '左对齐'):
        profiles['reference_english'] = _normalize_profile({'align': 'LEFT'}, profiles.get('reference', body))
    profiles['reference_heading'] = profiles.get('reference_heading') or profiles['h1']
    formula_rule = _find_instruction(texts, '公式应') or _find_instruction(texts, '公式', '居中')
    if formula_rule:
        profiles['formula'] = _profile_from_instruction(formula_rule, body, align='CENTER', first_indent_cm=0.0)
    return profiles


def _extract_page_and_header(fmt: Dict[str, Any]) -> Dict[str, Any]:
    sections = fmt.get('sections') or []
    s0 = sections[0] if sections else {}
    page = {
        'page_w': s0.get('page_width_cm', 21.0),
        'page_h': s0.get('page_height_cm', 29.7),
        'mt': s0.get('margin_top_cm', 2.54),
        'mb': s0.get('margin_bottom_cm', 2.54),
        'ml': s0.get('margin_left_cm', 2.54),
        'mr': s0.get('margin_right_cm', 2.54),
        'header': None,
    }
    for sec in sections:
        for h in sec.get('header', []) or []:
            text = (h.get('text') or '').strip()
            if not text:
                continue
            run = next((r for r in h.get('runs', []) if str(r.get('text', '')).strip()), next((r for r in h.get('runs', []) if r.get('size_pt')), {}))
            page['header'] = {
                'text': text,
                'align': h.get('alignment') if h.get('alignment') != 'DEFAULT' else 'CENTER',
                'font': run.get('font') or '宋体',
                'size': run.get('size_pt') or 9,
                'bold': bool(run.get('bold', False)),
                'italic': bool(run.get('italic', False)),
            }
            return page
    return page


def _section_role(sec: Dict[str, Any]) -> str:
    role = (sec.get('role') or '').strip()
    if role:
        return role
    h = (sec.get('heading') or '').strip()
    compact = re.sub(r'[\s：:]+', '', h).lower()
    if compact in ('摘要', '中文摘要'):
        return 'cn_abstract'
    if h.startswith('关键词') or compact in ('关键词', '关键字'):
        return 'cn_keywords'
    if compact == 'abstract':
        return 'en_abstract'
    if h.upper().replace(' ', '').startswith('KEYWORDS') or re.match(r'(?i)^key\s*words?', h):
        return 'en_keywords'
    if h.startswith('参考文献') or re.match(r'(?i)^references?$', h):
        return 'references'
    if re.search(r'致\s*谢', h):
        return 'acknowledgement'
    if re.search(r'附\s*录', h):
        return 'appendix'
    return 'body'


def _front_matter_sections(cnt: Dict[str, Any]) -> Dict[str, Any]:
    result: Dict[str, Any] = {'cn_abs': None, 'cn_kw': None, 'en_title': '', 'en_abs': None, 'en_kw': None, 'front_indices': set()}
    sections = cnt.get('sections') or []
    for idx, sec in enumerate(sections):
        role = _section_role(sec)
        h = (sec.get('heading') or '').strip()
        if role == 'cn_abstract':
            result['cn_abs'] = sec; result['front_indices'].add(idx)
        elif role == 'cn_keywords':
            result['cn_kw'] = sec; result['front_indices'].add(idx)
        elif role == 'en_abstract':
            result['en_abs'] = sec; result['front_indices'].add(idx)
        elif role == 'en_keywords':
            result['en_kw'] = sec; result['front_indices'].add(idx)
        elif sec.get('level') == 1 and _ascii_ratio(h) > 0.55 and not re.match(r'(?i)^chapter\s+\d+', h):
            if not result['en_title']:
                result['en_title'] = h; result['front_indices'].add(idx)
    return result


def _infer_template_rules(fmt: Dict[str, Any]) -> Dict[str, Any]:
    """Infer layout rules from template instruction text, not from fixed school names."""
    texts = _text_blob(fmt)
    caption_samples = [str(p.get('text') or '').strip() for p in fmt.get('paragraphs') or []
                       if re.match(r'^(图|表)\s*\d+(?:[.-]\d+)?\s+', str(p.get('text') or '').strip())]
    caption_number_space = None
    for sample in caption_samples:
        if re.match(r'^(图|表)\s+\d', sample):
            caption_number_space = True; break
        if re.match(r'^(图|表)\d', sample):
            caption_number_space = False; break
    ref_indent_chars = 2.0 if re.search(r'参考文献[^\n。；;]{0,120}悬挂缩进[^\n。；;]{0,20}2\s*字符|悬挂缩进[^\n。；;]{0,20}2\s*字符', texts) else None
    return {
        'cn_abstract_single_paragraph': bool(re.search(r'中文摘要[^\n。；;]{0,80}不分自然段|不分自然段[^\n。；;]{0,80}中文摘要', texts)),
        'en_title_upper': bool(re.search(r'英文题目[^\n。；;]{0,120}(大写字母|大写)', texts)),
        'caption_number_space': caption_number_space,
        'formula_center': bool(re.search(r'公式[^\n。；;]{0,60}居中', texts)),
        'formula_numbered': bool(re.search(r'公式[^\n。；;]{0,80}(编号|括弧|括号)', texts)),
        'reference_hanging_chars': ref_indent_chars,
        'reference_english_left': bool(re.search(r'英文参考文献[^\n。；;]{0,80}左对齐', texts)),
        'toc_indents_cm': [0.0, 0.74, 1.48],
    }


RUNTIME_TEMPLATE = r'''
# -*- coding: utf-8 -*-
"""
build_generated.py — generated by role-driven script_generator.py.
运行: python build_generated.py
"""
import json
import math
import os
import re
import shutil
import subprocess
import tempfile

from lxml import etree
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.image.image import Image as DocxImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from latex_omath import latex_to_omath

DATA = json.loads(__DATA_BLOB__)
BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(BASE, __OUT_DOCX__)

doc = Document()
TOC_PAGE_MAP = {}
USE_NATIVE_TOC = False

ALIGN = {
    'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
    'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
    'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
    'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
    'DISTRIBUTE': WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    'DEFAULT': WD_ALIGN_PARAGRAPH.LEFT,
}

CJK_FONTS = {'宋体', '黑体', '楷体', '微软雅黑', '仿宋', '华文宋体', '华文中宋'}


def _style_by_name(name, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        try:
            return doc.styles.add_style(name, style_type)
        except Exception:
            return None


def _set_style_font(style, east_asia='宋体', latin=None, size=12, bold=False, italic=False):
    if style is None:
        return
    latin = latin or ('Times New Roman' if east_asia in CJK_FONTS else east_asia)
    style.font.name = latin
    style.font.size = Pt(float(size))
    style.font.bold = bool(bold)
    style.font.italic = bool(italic)
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)
    rFonts.set(qn('w:eastAsia'), east_asia)
    rFonts.set(qn('w:hint'), 'eastAsia')


def _configure_para_style(style, prof, left_indent_cm=0.0):
    if style is None:
        return
    _set_style_font(style, prof.get('font') or '宋体',
                    'Times New Roman' if (prof.get('font') in CJK_FONTS) else (prof.get('font') or 'Times New Roman'),
                    prof.get('size') or 12, prof.get('bold', False), prof.get('italic', False))
    pf = style.paragraph_format
    fixed = prof.get('line_spacing_fixed_pt')
    if fixed:
        pf.line_spacing = Pt(float(fixed))
    elif prof.get('line_spacing_val'):
        pf.line_spacing = float(prof.get('line_spacing_val'))
    pf.space_before = Pt(float(prof.get('space_before_pt') or 0))
    pf.space_after = Pt(float(prof.get('space_after_pt') or 0))
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(float(left_indent_cm or 0))
    pf.alignment = ALIGN.get(prof.get('align') or 'LEFT', WD_ALIGN_PARAGRAPH.LEFT)


def configure_global_styles():
    """Lock styles regenerated by Word/WPS, especially TOC 1/2/3."""
    indents = (DATA.get('rules') or {}).get('toc_indents_cm') or [0.0, 0.74, 1.48]
    for style_name, role_name, indent in [
        ('TOC 1', 'toc1', indents[0] if len(indents) > 0 else 0.0),
        ('TOC 2', 'toc2', indents[1] if len(indents) > 1 else 0.74),
        ('TOC 3', 'toc3', indents[2] if len(indents) > 2 else 1.48),
        ('toc 1', 'toc1', indents[0] if len(indents) > 0 else 0.0),
        ('toc 2', 'toc2', indents[1] if len(indents) > 1 else 0.74),
        ('toc 3', 'toc3', indents[2] if len(indents) > 2 else 1.48),
    ]:
        _configure_para_style(_style_by_name(style_name), profile(role_name), indent)
    for style_name, role_name in [('Heading 1', 'h1'), ('Heading 2', 'h2'), ('Heading 3', 'h3')]:
        _configure_para_style(_style_by_name(style_name), profile(role_name), 0.0)
    _configure_para_style(_style_by_name('Normal'), profile('body'), 0.0)


def has_cjk(text):
    return any('\u4e00' <= c <= '\u9fff' for c in str(text or ''))


def ascii_ratio(text):
    text = str(text or '')
    return sum(1 for c in text if c.isascii() and c.isalpha()) / max(len(text), 1)


def set_run_fonts(run, east_asia='宋体', latin=None):
    latin = latin or ('Times New Roman' if east_asia in CJK_FONTS else east_asia)
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)
    rFonts.set(qn('w:eastAsia'), east_asia)
    rFonts.set(qn('w:hint'), 'eastAsia')


def profile(name):
    return DATA['profiles'].get(name) or DATA['profiles']['body']


def apply_paragraph_profile(p, prof, outline_level=None, first_indent_override=None):
    p.alignment = ALIGN.get(prof.get('align') or 'LEFT', WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    fixed = prof.get('line_spacing_fixed_pt')
    if fixed:
        pf.line_spacing = Pt(float(fixed))
    elif prof.get('line_spacing_val'):
        pf.line_spacing = float(prof.get('line_spacing_val'))
    if prof.get('space_before_pt') is not None:
        pf.space_before = Pt(float(prof.get('space_before_pt') or 0))
    if prof.get('space_after_pt') is not None:
        pf.space_after = Pt(float(prof.get('space_after_pt') or 0))
    indent = first_indent_override if first_indent_override is not None else prof.get('first_indent_cm')
    if indent:
        pf.first_line_indent = Cm(float(indent))
    else:
        pf.first_line_indent = Cm(0)
    if outline_level is not None:
        pPr = p._element.get_or_add_pPr()
        # Give WPS/Word two independent signals for TOC generation:
        # 1) built-in HeadingN style id; 2) explicit outline level.  Direct
        # formatting is applied after style assignment, so the template look is
        # preserved while WPS can still generate a normal directory.
        level_int = max(0, min(int(outline_level), 8))
        pstyle = pPr.find(qn('w:pStyle'))
        if pstyle is None:
            pstyle = OxmlElement('w:pStyle')
            pPr.insert(0, pstyle)
        pstyle.set(qn('w:val'), 'Heading' + str(level_int + 1))
        old = pPr.find(qn('w:outlineLvl'))
        if old is not None:
            pPr.remove(old)
        ol = OxmlElement('w:outlineLvl')
        ol.set(qn('w:val'), str(level_int))
        pPr.append(ol)


def apply_run_profile(run, prof, text='', superscript=False, force_latin=None):
    font = prof.get('font') or '宋体'
    latin = force_latin
    if not latin:
        if font in CJK_FONTS:
            latin = 'Times New Roman'
        else:
            latin = font
    east_asia = font if font in CJK_FONTS else font
    if ascii_ratio(text) > 0.75 and font in CJK_FONTS:
        east_asia = font
        latin = 'Times New Roman'
    set_run_fonts(run, east_asia, latin)
    run.font.size = Pt(float(prof.get('size') or 12))
    run.bold = bool(prof.get('bold', False))
    run.italic = bool(prof.get('italic', False))
    run.font.superscript = bool(superscript)
    run.font.color.rgb = RGBColor(0, 0, 0)


def citation_pattern():
    return re.compile(r'(\[\d+(?:\s*[-,，、]\s*\d+)*\])')


def add_text_runs(p, text, prof, superscript_citations=False):
    text = str(text or '')
    pos = 0
    pat = citation_pattern() if superscript_citations else None
    matches = list(pat.finditer(text)) if pat else []
    if not matches:
        r = p.add_run(text)
        apply_run_profile(r, prof, text)
        return
    for m in matches:
        if m.start() > pos:
            seg = text[pos:m.start()]
            r = p.add_run(seg)
            apply_run_profile(r, prof, seg)
        r = p.add_run(m.group(1))
        apply_run_profile(r, prof, m.group(1), superscript=True)
        pos = m.end()
    if pos < len(text):
        seg = text[pos:]
        r = p.add_run(seg)
        apply_run_profile(r, prof, seg)


def add_text(text, role='body', first_indent=True, outline_level=None, bold_prefix=None):
    prof = profile(role)
    text = clean_text_artifacts(text)
    if not text:
        return None
    p = doc.add_paragraph()
    apply_paragraph_profile(p, prof, outline_level=outline_level, first_indent_override=(prof.get('first_indent_cm') if first_indent else 0))
    superscript = role == 'body'
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        p1 = dict(prof); p1['bold'] = True
        apply_run_profile(r1, p1, bold_prefix)
        add_text_runs(p, text[len(bold_prefix):], prof, superscript)
    else:
        add_text_runs(p, text, prof, superscript)
    return p


def normalize_heading_spacing(text):
    t = str(text or '').strip()
    t = re.sub(r'^(第[一二三四五六七八九十百千万\d]+章)\s*(\S)', r'\1 \2', t)
    t = re.sub(r'^(\d+(?:\.\d+)*)(?![.\d])\s*(\S)', r'\1 \2', t)
    return t


def add_heading(text, level):
    level = max(1, min(int(level or 1), 3))
    # 一级标题居中不缩进；二/三级标题按模板首行缩进2字符。
    p = add_text(normalize_heading_spacing(text), role='h' + str(level), first_indent=(level > 1), outline_level=level - 1)
    if p is not None:
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
    return p


def text_width_cm():
    page = DATA.get('page') or {}
    try:
        return max(4.0, float(page.get('page_w') or 21.0) - float(page.get('ml') or 2.54) - float(page.get('mr') or 2.54))
    except Exception:
        return 15.0


def text_width_inches(scale=1.0):
    return max(1.0, text_width_cm() * float(scale) / 2.54)


def text_height_cm():
    page = DATA.get('page') or {}
    try:
        return max(4.0, float(page.get('page_h') or 29.7) - float(page.get('mt') or 2.54) - float(page.get('mb') or 2.54))
    except Exception:
        return 24.0


def text_height_inches(scale=1.0):
    return max(1.0, text_height_cm() * float(scale) / 2.54)


def role_line_height_pt(role_name):
    prof = profile(role_name)
    if prof.get('line_spacing_fixed_pt'):
        return float(prof.get('line_spacing_fixed_pt') or 0)
    size = float(prof.get('size') or 12)
    return size * float(prof.get('line_spacing_val') or 1.0)


def caption_block_inches(role_name='figure_caption'):
    prof = profile(role_name)
    total_pt = role_line_height_pt(role_name) + float(prof.get('space_before_pt') or 0) + float(prof.get('space_after_pt') or 0)
    return max(0.0, total_pt / 72.0)


def configure_picture_paragraph(p, keep_with_next=True):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.keep_with_next = bool(keep_with_next)
    pf.keep_together = True


def fit_picture_dimensions(path, has_caption=True):
    max_width = Inches(text_width_inches(1.0))
    caption_inches = caption_block_inches('figure_caption') if has_caption else 0.0
    max_height = Inches(max(1.0, text_height_inches() - caption_inches))
    try:
        image = DocxImage.from_file(path)
        width, height = image.scaled_dimensions(width=max_width)
        if height > max_height:
            width, height = image.scaled_dimensions(height=max_height)
        return width, height
    except Exception:
        return max_width, None


def setup_section(sec):
    page = DATA['page']
    sec.page_width = Cm(float(page.get('page_w') or 21.0))
    sec.page_height = Cm(float(page.get('page_h') or 29.7))
    sec.top_margin = Cm(float(page.get('mt') or 2.54))
    sec.bottom_margin = Cm(float(page.get('mb') or 2.54))
    sec.left_margin = Cm(float(page.get('ml') or 2.54))
    sec.right_margin = Cm(float(page.get('mr') or 2.54))


def add_page_field(paragraph):
    r = paragraph.add_run()
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    r._element.append(begin); r._element.append(instr); r._element.append(end)
    return r


def set_page_numbering(sec, fmt='decimal', start=1):
    sectPr = sec._sectPr
    pg = sectPr.find(qn('w:pgNumType'))
    if pg is None:
        pg = OxmlElement('w:pgNumType')
        sectPr.append(pg)
    # New sections created by python-docx copy the previous section's
    # w:start.  If this continuation section should not restart numbering,
    # remove the copied attribute instead of leaving another start=1 behind.
    if start is None:
        start_attr = qn('w:start')
        if start_attr in pg.attrib:
            del pg.attrib[start_attr]
    else:
        pg.set(qn('w:start'), str(start))
    pg.set(qn('w:fmt'), fmt)


def apply_header_footer(sec, page_fmt='decimal', start=1):
    hdr = DATA['page'].get('header') or {}
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    hp = sec.header.paragraphs[0]
    hp.text = ''
    hp.alignment = ALIGN.get(hdr.get('align', 'CENTER'), WD_ALIGN_PARAGRAPH.CENTER)
    if hdr.get('text'):
        r = hp.add_run(hdr['text'])
        apply_run_profile(r, {'font': hdr.get('font') or '宋体', 'size': hdr.get('size') or 9, 'bold': hdr.get('bold', False), 'italic': hdr.get('italic', False)}, hdr['text'])
        pPr = hp._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4'); bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom); pPr.append(pBdr)
    fp = sec.footer.paragraphs[0]
    fp.text = ''
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = add_page_field(fp)
    apply_run_profile(r, {'font': hdr.get('font') or '宋体', 'size': hdr.get('size') or 9}, '')
    set_page_numbering(sec, page_fmt, start)


def clear_header_footer(sec):
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    for part in (sec.header, sec.footer):
        for p in part.paragraphs:
            p.text = ''


def _paragraph_is_empty_for_cleanup(p):
    if p.text.strip():
        return False
    xml = p._element.xml
    if '<w:drawing' in xml or '<w:pict' in xml or 'oMath' in xml:
        return False
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
        return False
    return True


def remove_trailing_empty_body_paragraphs(limit=12):
    removed = 0
    while len(doc.paragraphs) > 1 and removed < limit and _paragraph_is_empty_for_cleanup(doc.paragraphs[-1]):
        el = doc.paragraphs[-1]._element
        el.getparent().remove(el)
        removed += 1


def add_section_with_header(page_fmt='decimal', start=1):
    remove_trailing_empty_body_paragraphs()
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(sec)
    apply_header_footer(sec, page_fmt, start)
    return sec


def enable_update_fields_on_open():
    settings = doc.settings._element
    upd = settings.find(qn('w:updateFields'))
    if upd is None:
        upd = OxmlElement('w:updateFields')
        settings.append(upd)
    upd.set(qn('w:val'), 'true')


def remove_initial_empty_paragraph():
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text.strip():
        el = doc.paragraphs[0]._element
        el.getparent().remove(el)


def force_cover_headerless():
    if not doc.sections:
        return
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    for part in (sec.header, sec.first_page_header, sec.footer, sec.first_page_footer):
        part.is_linked_to_previous = False
        for p in part.paragraphs:
            p.text = ''


def normalize_label(text):
    return re.sub(r'[\s：:]+', '', str(text or ''))


def para_text_from_cover_el(el):
    if el.get('type') in ('para', 'empty', 'image'):
        return ''.join(r.get('t', '') for r in el.get('r', []))
    if el.get('type') == 'table':
        parts = []
        for row in el.get('rows', []):
            for cell in row:
                for pp in cell.get('p', []):
                    parts.append(''.join(r.get('t', '') for r in pp.get('r', [])))
        return ''.join(parts)
    return ''


def apply_cover_run(run, rd):
    font = rd.get('fn') or rd.get('fe') or '宋体'
    east = rd.get('fe') or font
    latin = rd.get('fn') or ('Times New Roman' if east in CJK_FONTS else east)
    set_run_fonts(run, east, latin)
    if rd.get('sz'):
        run.font.size = Pt(float(rd.get('sz')))
    run.bold = bool(rd.get('b', False))


def apply_cover_paragraph_format(p, el):
    am = {'left': 'LEFT', 'center': 'CENTER', 'right': 'RIGHT', 'both': 'JUSTIFY', 'distribute': 'DISTRIBUTE'}
    if el.get('al'):
        p.alignment = ALIGN.get(am.get(el.get('al'), 'LEFT'), WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    line = el.get('ls_val')
    rule = el.get('ls_rule')
    if line:
        try:
            n = int(line)
            pf.line_spacing = Pt(n / 20.0) if rule in ('exact', 'atLeast') else n / 240.0
        except Exception:
            pass
    if el.get('sp_before'):
        try: pf.space_before = Pt(int(el.get('sp_before')) / 20.0)
        except Exception: pass
    if el.get('sp_after'):
        try: pf.space_after = Pt(int(el.get('sp_after')) / 20.0)
        except Exception: pass
    if el.get('fl_indent'):
        try: pf.first_line_indent = Pt(int(el.get('fl_indent')) / 20.0)
        except Exception: pass


def asset_path(name):
    if not name:
        return None
    bases = []
    if DATA.get('assets_dir'):
        bases.append(DATA.get('assets_dir'))
    bases.extend([os.path.join(BASE, 'assets'), BASE, os.getcwd()])
    for b in bases:
        p = os.path.join(b, name) if not os.path.isabs(name) else name
        if os.path.exists(p):
            return p
    return None


def image_width_from_extent(extent, default_inches=1.2):
    if not extent:
        return Inches(default_inches)
    try:
        cx = int(extent.get('cx') or 0)
        if cx > 0:
            return Inches(cx / 914400.0)
    except Exception:
        pass
    return Inches(default_inches)


def image_dimensions_from_extent(extent, default_inches=1.2):
    if not extent:
        return Inches(default_inches), None
    try:
        cx = int(extent.get('cx') or 0)
        cy = int(extent.get('cy') or 0)
        width = Inches(cx / 914400.0) if cx > 0 else Inches(default_inches)
        height = Inches(cy / 914400.0) if cy > 0 else None
        return width, height
    except Exception:
        return Inches(default_inches), None


def add_asset_picture(run, rd, default_inches=1.2):
    path = asset_path(rd.get('asset') or rd.get('image'))
    if not path:
        return False
    try:
        width, height = image_dimensions_from_extent(rd.get('extent'), default_inches)
        if height is not None:
            run.add_picture(path, width=width, height=height)
        else:
            run.add_picture(path, width=width)
        return True
    except Exception:
        return False


def render_cover_para(el):
    p = doc.add_paragraph()
    apply_cover_paragraph_format(p, el)
    for rd in el.get('r', []):
        rr = p.add_run(rd.get('t', ''))
        apply_cover_run(rr, rd)
    return p


def render_cover_image(el):
    p = doc.add_paragraph()
    apply_cover_paragraph_format(p, el)
    extent = el.get('extent') or {}
    try:
        cy = int(extent.get('cy') or 0)
        if cy > 0:
            p.paragraph_format.line_spacing = Pt((cy / 12700.0) + 4.0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
    except Exception:
        pass
    if not el.get('al'):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    add_asset_picture(r, el, default_inches=1.35)
    return p


def set_cell_borders(cell, **sides):
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    tcB = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        val = sides.get(side, 'nil')
        el = OxmlElement('w:' + side)
        if isinstance(val, dict):
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', '8')))
            el.set(qn('w:color'), val.get('color', '000000'))
        else:
            el.set(qn('w:val'), val)
            el.set(qn('w:sz'), '0' if val in ('nil', 'none') else '8')
            el.set(qn('w:color'), '000000')
        el.set(qn('w:space'), '0')
        tcB.append(el)
    tcPr.append(tcB)


def set_table_indent(table, twips=0):
    tblPr = table._tbl.tblPr
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), str(int(twips)))
    tblInd.set(qn('w:type'), 'dxa')




def _set_or_remove_attr(el, name, value):
    attr = qn('w:' + name)
    if value is None:
        if attr in el.attrib:
            del el.attrib[attr]
    else:
        el.set(attr, str(value))


def set_table_width(table, spec):
    if not spec:
        return
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.insert(0, tblW)
    for k, v in spec.items():
        _set_or_remove_attr(tblW, k, v)


def set_table_alignment_from_jc(table, jc):
    if jc in ('left', 'start'):
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
    elif jc == 'right':
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    elif jc == 'center':
        table.alignment = WD_TABLE_ALIGNMENT.CENTER


def set_table_layout(table, layout_type):
    if not layout_type:
        return
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), layout_type)


def set_table_cell_margins(table, margins):
    if not margins:
        return
    tblPr = table._tbl.tblPr
    cellMar = tblPr.find(qn('w:tblCellMar'))
    if cellMar is not None:
        tblPr.remove(cellMar)
    cellMar = OxmlElement('w:tblCellMar')
    for side, attrs in margins.items():
        el = OxmlElement('w:' + side)
        for k, v in (attrs or {}).items():
            _set_or_remove_attr(el, k, v)
        cellMar.append(el)
    tblPr.append(cellMar)


def set_table_grid(table, grid_cols):
    if not grid_cols:
        return
    tblGrid = table._tbl.tblGrid
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        table._tbl.insert(0, tblGrid)
    for i, w in enumerate(grid_cols):
        if i < len(tblGrid.gridCol_lst):
            gc = tblGrid.gridCol_lst[i]
        else:
            gc = OxmlElement('w:gridCol')
            tblGrid.append(gc)
        gc.set(qn('w:w'), str(w))


def apply_cover_table_props(table, el):
    props = el.get('tblPr') or {}
    if props:
        set_table_width(table, props.get('tblW'))
        if props.get('tblInd'):
            try:
                set_table_indent(table, int(props['tblInd'].get('w') or 0))
            except Exception:
                pass
        set_table_alignment_from_jc(table, props.get('jc'))
        layout = props.get('tblLayout') or {}
        set_table_layout(table, layout.get('type'))
        set_table_cell_margins(table, props.get('cellMar'))
        set_table_grid(table, props.get('grid_cols'))


def set_cell_margins(cell, margins):
    if not margins:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcMar'))
    if old is not None:
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, attrs in margins.items():
        el = OxmlElement('w:' + side)
        for k, v in (attrs or {}).items():
            _set_or_remove_attr(el, k, v)
        mar.append(el)
    tcPr.append(mar)


def apply_cell_props(cell, cell_data):
    tcPr_data = cell_data.get('tcPr') or {}
    tcW = tcPr_data.get('tcW') or {}
    if tcW:
        tcPr = cell._tc.get_or_add_tcPr()
        old = tcPr.find(qn('w:tcW'))
        if old is None:
            old = OxmlElement('w:tcW')
            tcPr.insert(0, old)
        for k, v in tcW.items():
            _set_or_remove_attr(old, k, v)
    if tcPr_data.get('tcMar'):
        set_cell_margins(cell, tcPr_data.get('tcMar'))
    valign = tcPr_data.get('vAlign')
    if valign == 'top':
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    elif valign == 'bottom':
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    else:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_no_wrap(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    if tcPr.find(qn('w:noWrap')) is None:
        tcPr.append(OxmlElement('w:noWrap'))


def apply_row_props(row, row_props):
    if not row_props:
        return
    trPr = row._tr.get_or_add_trPr()
    height = row_props.get('height') or {}
    if height:
        old = trPr.find(qn('w:trHeight'))
        if old is None:
            old = OxmlElement('w:trHeight')
            trPr.append(old)
        for k, v in height.items():
            _set_or_remove_attr(old, k, v)
    if row_props.get('cantSplit'):
        if trPr.find(qn('w:cantSplit')) is None:
            trPr.append(OxmlElement('w:cantSplit'))


def cover_table_sample_value(el, label_suffix='题目'):
    for row in el.get('rows') or []:
        if len(row) < 2:
            continue
        label = normalize_label(first_cell_label(row))
        if label.endswith(label_suffix):
            return ''.join(r.get('t', '') for pp in row[1].get('p', []) for r in pp.get('r', []))
    return ''


def cover_table_value_cell(row):
    if not row or len(row) < 2:
        return None
    return row[1]


def cover_text_units(text):
    units = 0.0
    for ch in str(text or ''):
        if ch.isspace():
            continue
        units += 0.5 if ch.isascii() else 1.0
    return units


def cover_title_capacity_chars(el, label_suffix='题目'):
    for row in el.get('rows') or []:
        label = normalize_label(first_cell_label(row))
        if not label.endswith(label_suffix):
            continue
        cell = cover_table_value_cell(row)
        if not cell:
            continue
        width_dxa = 0
        try:
            width_dxa = int(cell.get('w') or 0)
        except Exception:
            width_dxa = 0
        try:
            tcw = (cell.get('tcPr') or {}).get('tcW') or {}
            width_dxa = max(width_dxa, int(tcw.get('w') or 0))
        except Exception:
            pass
        font_size = 16.0
        sizes = []
        for pp in cell.get('p') or []:
            for rd in pp.get('r') or []:
                try:
                    if rd.get('sz'):
                        sizes.append(float(rd.get('sz')))
                except Exception:
                    pass
        if sizes:
            font_size = max(sizes)
        if width_dxa <= 0 or font_size <= 0:
            return 20.0
        width_pt = width_dxa / 20.0
        # CJK cover titles are close to one em per character.  A small
        # reserve avoids treating text that barely fits as one-line content.
        return max(1.0, width_pt / (font_size * 1.05))
    return 20.0


def estimate_cover_title_lines(text, el):
    capacity = cover_title_capacity_chars(el)
    return max(1, int(math.ceil(cover_text_units(text) / max(capacity, 1.0))))


def is_cover_empty_paragraph(el):
    """True for structurally empty cover paragraphs, including section markers."""
    if el.get('type') != 'empty':
        return False
    return not ''.join(r.get('t', '') for r in el.get('r', [])).strip()


def is_elastic_cover_empty(el):
    """Empty spacer that can be removed without deleting a section break."""
    return is_cover_empty_paragraph(el) and not el.get('section_break_after')


def compute_cover_skip_indices(cover):
    """Remove template spacer paragraphs only when replacement content is longer.

    This implements the template instruction 'delete one return above/below the table
    if the title uses two lines' without checking any school-specific text.  It
    compares the sample title in the template with the actual title and removes
    only structurally empty spacer paragraphs adjacent to the cover info table.
    """
    info_idx = next((i for i, el in enumerate(cover) if el.get('role') == 'cover_info_table'), None)
    if info_idx is None:
        return set()
    info_el = cover[info_idx]
    actual = str((DATA.get('cover_info') or {}).get('paper_title') or DATA.get('title_cn') or '').strip()
    sample = cover_table_sample_value(info_el, '题目')
    if not actual or not sample:
        return set()
    extra_lines = max(0, estimate_cover_title_lines(actual, info_el) - estimate_cover_title_lines(sample, info_el))
    skip = set()
    j = info_idx - 1
    before = []
    while j >= 0 and is_elastic_cover_empty(cover[j]):
        before.append(j); j -= 1
    # Pre-table blank paragraphs are elastic vertical budget.  Generated
    # content must keep the full cover on page one, so these can be removed.
    for idx in before:
        skip.add(idx)
    j = info_idx + 1
    after = []
    while j < len(cover) and is_elastic_cover_empty(cover[j]):
        after.append(j); j += 1
    # Keep a bounded gap before the next visible paragraph (normally a
    # committee/signature line).  Keeping every template spacer can push that
    # paragraph to page two, while removing all of them visually attaches it
    # to the final table row.
    keep_after = max(1, min(2, len(after) - extra_lines))
    after_remove = max(0, len(after) - keep_after)
    for idx in after[:after_remove]:
        skip.add(idx)
    # Also delete trailing blank spacer paragraphs immediately before an
    # empty section-break carrier. Keep the marker itself so the render loop
    # can still create the intended next-page section.
    for marker_idx, marker_el in enumerate(cover):
        if cover_element_is_empty_section_marker(marker_el):
            k = marker_idx - 1
            while k >= 0 and is_elastic_cover_empty(cover[k]):
                skip.add(k)
                k -= 1
    return skip


def cover_element_is_empty_section_marker(el):
    # A section break can be stored on a visually empty paragraph.  Render the
    # break, but never render that paragraph itself; otherwise it becomes a
    # blank page between declaration/front-matter sections.
    return bool(el.get('section_break_after')) and is_cover_empty_paragraph(el)

def first_cell_label(row):
    if not row:
        return ''
    cell = row[0]
    return ''.join(r.get('t', '') for pp in cell.get('p', []) for r in pp.get('r', []))


def is_code_like_cover_table(el):
    if el.get('role') == 'cover_code_table':
        return True
    rows = el.get('rows') or []
    labels = [normalize_label(first_cell_label(row)) for row in rows if row]
    return bool(labels) and len(rows) <= 2 and all(x.endswith('编码') for x in labels if x)


def render_cover_table(el):
    rows = el.get('rows', [])
    if not rows:
        return None
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    apply_cover_table_props(table, el)
    code_like_table = is_code_like_cover_table(el)
    if code_like_table and ncols >= 2:
        table.autofit = False
        try:
            table.columns[0].width = Cm(2.45)
            table.columns[1].width = Cm(1.75)
        except Exception:
            pass
    # Fallback only when the extractor did not provide table-level properties.
    if not (el.get('tblPr') or {}).get('jc'):
        table.alignment = WD_TABLE_ALIGNMENT.LEFT if code_like_table else WD_TABLE_ALIGNMENT.CENTER
    if code_like_table and not (el.get('tblPr') or {}).get('tblInd'):
        set_table_indent(table, 0)
    cover_info = DATA.get('cover_info') or {}
    label_map = {
        '学校编码': cover_info.get('school_code', '') or cover_info.get('degree_code', ''),
        '学位编码': cover_info.get('degree_code', '') or cover_info.get('school_code', ''),
        '论文题目': cover_info.get('paper_title', ''),
        '学生姓名': cover_info.get('student_name', ''),
        '学号': cover_info.get('student_id', ''),
        '所属学院': cover_info.get('college', ''),
        '专业班级': cover_info.get('class_name', ''),
        '指导老师': cover_info.get('advisor', ''),
        '指导教师': cover_info.get('advisor', ''),
    }
    norm_map = {normalize_label(k): v for k, v in label_map.items() if v}
    row_props = (el.get('tblPr') or {}).get('rows') or []
    for ri, row in enumerate(rows):
        if ri < len(table.rows):
            apply_row_props(table.rows[ri], row_props[ri] if ri < len(row_props) else {})
        row_label = ''
        if row and row[0].get('p'):
            row_label = ''.join(r.get('t', '') for r in row[0]['p'][0].get('r', []))
        row_key = normalize_label(row_label)
        row_value = norm_map.get(row_key, '')
        if not row_value:
            for k, v in norm_map.items():
                if k and (k in row_key or row_key in k):
                    row_value = v; break
        force_left = code_like_table or row_key.endswith('编码')
        for ci in range(ncols):
            cell = table.rows[ri].cells[ci]
            cell.text = ''
            cell_data = row[ci] if ci < len(row) else {'p': []}
            apply_cell_props(cell, cell_data)
            if code_like_table:
                set_cell_no_wrap(cell)
            if cell_data.get('w') and not (cell_data.get('tcPr') or {}).get('tcW'):
                try: cell.width = Cm(float(cell_data.get('w')) / 567.0)
                except Exception: pass
            paras = cell_data.get('p') or [{'r': []}]
            for pi, pp in enumerate(paras):
                p = cell.paragraphs[0] if pi == 0 else cell.add_paragraph()
                apply_cover_paragraph_format(p, pp)
                if force_left:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                use_value = (ci == 1 and row_value)
                if use_value:
                    rd = (pp.get('r') or [{}])[0]
                    rr = p.add_run(row_value)
                    apply_cover_run(rr, rd)
                    if code_like_table and not rd.get('sz'):
                        rr.font.size = Pt(10.5)
                    continue
                for rd in pp.get('r', []) or [{}]:
                    rr = p.add_run(rd.get('t', ''))
                    apply_cover_run(rr, rd)
                    if code_like_table and not rd.get('sz'):
                        rr.font.size = Pt(10.5)
                    if rd.get('asset') or rd.get('image'):
                        add_asset_picture(rr, rd)
            borders = cell_data.get('borders') or {}
            if borders:
                set_cell_borders(cell, **borders)
    return table


def render_cover_and_declarations():
    setup_section(doc.sections[0])
    clear_header_footer(doc.sections[0])
    cover = DATA.get('cover') or []
    if not cover:
        return add_section_with_header('upperRoman', 1)
    front_started = False
    skip_indices = compute_cover_skip_indices(cover)
    for idx, el in enumerate(cover):
        # A paragraph whose only purpose is to carry a section break must not
        # be rendered as a blank page. Treat it as a structural marker.
        if idx in skip_indices:
            continue
        should_render = not cover_element_is_empty_section_marker(el)
        if should_render:
            if el.get('type') in ('para', 'empty'):
                render_cover_para(el)
            elif el.get('type') == 'table':
                render_cover_table(el)
            elif el.get('type') == 'image':
                render_cover_image(el)
        if el.get('section_break_after'):
            if not front_started:
                front_started = True
                add_section_with_header('upperRoman', 1)
            else:
                add_section_with_header('upperRoman', None)
    if not front_started:
        add_section_with_header('upperRoman', 1)
    return doc.sections[-1]


def section_text(sec):
    out = []
    for para in sec.get('paragraphs', []) or []:
        if isinstance(para, str):
            out.append(para.strip())
        elif isinstance(para, dict) and para.get('text'):
            out.append(str(para.get('text')).strip())
    return '\n'.join(x for x in out if x)


def add_rich_text_item(item, role='body', first_indent=True, chapter=None):
    """Render a content item that may contain plain text plus extracted math.

    Markdown inline math arrives as {"text": "...", "math": [...]}.  The
    generator cannot recover exact inline positions after extraction, so it
    keeps the readable text paragraph and renders each math object as native
    OMML immediately after it.  This preserves editability instead of silently
    dropping formulas or leaving them as plain text.
    """
    if isinstance(item, str):
        return add_text(item, role=role, first_indent=first_indent)
    if not isinstance(item, dict):
        return None
    text = str(item.get('text') or '').strip()
    if text:
        add_text(text, role=role, first_indent=first_indent)
    if item.get('math') and not item.get('latex') and not item.get('xml'):
        for m in item.get('math') or []:
            render_formula({
                'latex': m.get('latex'),
                'xml': m.get('xml'),
                'text': m.get('text') or '',
                'numbered': False,
            }, chapter)
    elif item.get('role') == 'formula' or item.get('latex') or item.get('xml'):
        render_formula(item, chapter)
    return None


def add_keywords(label, value, role):
    prof = profile(role)
    p = doc.add_paragraph()
    apply_paragraph_profile(p, prof, first_indent_override=0)
    r1 = p.add_run(label)
    p_bold = dict(prof); p_bold['bold'] = True
    apply_run_profile(r1, p_bold, label)
    r2 = p.add_run(value)
    p_norm = dict(prof); p_norm['bold'] = False
    apply_run_profile(r2, p_norm, value)
    return p


def add_blank_line(role='body'):
    return add_text('', role=role, first_indent=False)


def collect_toc_entries():
    entries = []
    for i, sec in enumerate(DATA.get('sections') or []):
        if is_front_section_index(i):
            continue
        h = (sec.get('heading') or '').strip()
        role = sec.get('role') or ''
        if not h or h == '正文':
            continue
        if is_reference_heading(h) or is_backmatter_heading(h) or is_caption_heading(h) or role in ('references', 'acknowledgement', 'appendix'):
            continue
        level = max(1, min(int(sec.get('level') or 1), 3))
        entries.append({'level': level, 'text': normalize_heading_spacing(h)})
    if DATA.get('references'):
        entries.append({'level': 1, 'text': '参考文献'})
    ack_sections, app_sections = collect_structural_backmatter()
    pure_refs, ack_from_refs, app_from_refs = split_refs_backmatter(DATA.get('references') or [])
    if ack_sections or ack_from_refs:
        entries.append({'level': 1, 'text': '致  谢'})
    if app_sections or app_from_refs:
        entries.append({'level': 1, 'text': '附  录'})
    return entries


def add_toc_line(text, level, page_text=''):
    level = max(1, min(int(level or 1), 3))
    prof = profile('toc' + str(level))
    p = doc.add_paragraph()
    apply_paragraph_profile(p, prof, first_indent_override=0)
    indents = (DATA.get('rules') or {}).get('toc_indents_cm') or [0.0, 0.74, 1.48]
    p.paragraph_format.left_indent = Cm(float(indents[level - 1] if len(indents) >= level else 0.0))
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(max(1.0, text_width_cm() - 0.15)), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(text)
    apply_run_profile(r, prof, text)
    p.add_run('\t')
    r2 = p.add_run(str(page_text))
    page_prof = dict(profile('toc2'))
    apply_run_profile(r2, page_prof, str(page_text))
    return p


def add_wps_toc_field():
    """Insert a real Word/WPS TOC field instead of a fake static directory.

    The generated DOCX only needs correct heading styles/outline levels.  WPS
    can then populate page numbers by Update Field / Generate Directory.  This
    avoids any dependency on LibreOffice or pdftotext and does not guess pages.
    """
    p = doc.add_paragraph()
    apply_paragraph_profile(p, profile('body'), first_indent_override=0)

    def append_run_with(el):
        r = OxmlElement('w:r')
        r.append(el)
        p._element.append(r)
        return r

    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    append_run_with(begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r' TOC \o "1-3" \h \z \u '
    append_run_with(instr)

    sep = OxmlElement('w:fldChar')
    sep.set(qn('w:fldCharType'), 'separate')
    append_run_with(sep)

    hint = p.add_run('请在 WPS 中右键“更新域”或“生成目录”，目录将按正文标题自动生成。')
    hint_prof = dict(profile('body'))
    hint_prof['italic'] = True
    apply_run_profile(hint, hint_prof, hint.text)

    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    append_run_with(end)


def add_toc():
    enable_update_fields_on_open()
    configure_global_styles()
    add_text('目  录', role='toc_title', first_indent=False)
    if USE_NATIVE_TOC:
        add_wps_toc_field()
        return
    entries = collect_toc_entries()
    for ent in entries:
        key = _norm_for_pdf_match(ent.get('text') or '')
        add_toc_line(ent.get('text') or '', ent.get('level') or 1, TOC_PAGE_MAP.get(key, ''))


def render_front_matter():
    front = DATA.get('front') or {}
    title_cn = DATA.get('title_cn') or ''
    if title_cn:
        add_text(title_cn, role='cn_title', first_indent=False)
    cn_abs = front.get('cn_abs')
    if cn_abs:
        add_text('摘 要', role='cn_abstract_heading', first_indent=False)
        cn_items = []
        for para in cn_abs.get('paragraphs', []) or []:
            if isinstance(para, dict) and (para.get('math') or para.get('role') == 'formula' or para.get('latex') or para.get('xml')):
                cn_items.append(para)
            else:
                text = para if isinstance(para, str) else para.get('text', '')
                if str(text).strip():
                    cn_items.append(str(text).strip())
        if DATA.get('rules', {}).get('cn_abstract_single_paragraph') and cn_items:
            plain_items = [x for x in cn_items if isinstance(x, str)]
            rich_items = [x for x in cn_items if not isinstance(x, str)]
            if plain_items:
                add_text(''.join(plain_items), role='cn_abstract_body', first_indent=True)
            for item in rich_items:
                add_rich_text_item(item, role='cn_abstract_body', first_indent=True)
        else:
            for item in cn_items:
                add_rich_text_item(item, role='cn_abstract_body', first_indent=True)
    cn_kw = front.get('cn_kw')
    if cn_kw and cn_abs:
        add_blank_line('cn_abstract_body')
    if cn_kw:
        val = section_text(cn_kw)
        if val:
            add_keywords('关键词：', val, 'cn_keywords')
    has_en = bool(front.get('en_title') or front.get('en_abs') or front.get('en_kw'))
    if has_en:
        doc.add_page_break()
    en_title = front.get('en_title') or ''
    if en_title and DATA.get('rules', {}).get('en_title_upper'):
        en_title = en_title.upper()
    if en_title:
        add_text(en_title, role='en_title', first_indent=False)
    en_abs = front.get('en_abs')
    if en_abs:
        add_text('ABSTRACT', role='en_abstract_heading', first_indent=False)
        for para in en_abs.get('paragraphs', []) or []:
            add_rich_text_item(para, role='en_abstract_body', first_indent=True)
    en_kw = front.get('en_kw')
    if en_kw:
        val = section_text(en_kw).replace('；', ';')
        if val:
            add_keywords('KEY WORDS: ', val, 'en_keywords')
    add_section_with_header('upperRoman', None)
    add_toc()


def is_front_section_index(i):
    return i in set(DATA.get('front_indices') or [])


def is_reference_heading(h):
    h = str(h or '').strip()
    return h.startswith('参考文献') or bool(re.match(r'(?i)^references?$', h))


def is_ack_heading(h):
    return bool(re.search(r'致\s*谢', str(h or '')))


def is_appendix_heading(h):
    return bool(re.search(r'附\s*录', str(h or '')))


def is_backmatter_heading(h):
    return is_ack_heading(h) or is_appendix_heading(h)


def is_caption_heading(h):
    return bool(re.match(r'^(图|表)\s*\d+(?:[.-]\d+)?\s*', str(h or '').strip()))


def normalize_caption(text):
    t = str(text or '').strip()
    space = (DATA.get('rules') or {}).get('caption_number_space')
    if space is True:
        t = re.sub(r'^(图|表)\s*(\d+(?:[.-]\d+)?)\s*', r'\1 \2 ', t)
    elif space is False:
        t = re.sub(r'^(图|表)\s*(\d+(?:[.-]\d+)?)\s*', r'\1\2 ', t)
    else:
        t = re.sub(r'^(图|表)\s*(\d+(?:[.-]\d+)?)\s*', r'\1 \2 ', t)
    return t.strip()


def clean_markdown_links(text):
    def repl(m):
        label = (m.group(1) or '').strip()
        target = (m.group(2) or '').strip()
        return label or target
    return re.sub(r'\[([^\]]+)\]\(([^)]+)\)', repl, str(text or ''))


def is_noise_text(text):
    return str(text or '').strip() in {'复制', 'Copy', 'Plain Text', '纯文本'}


def clean_text_artifacts(text, preserve_newlines=False):
    t = clean_markdown_links(text).replace('\u00a0', ' ')
    if preserve_newlines:
        lines = []
        for line in t.replace('\r\n', '\n').replace('\r', '\n').split('\n'):
            s = re.sub(r'[ \t]+', ' ', line).strip()
            if not is_noise_text(s):
                lines.append(s)
        return '\n'.join(lines).strip()
    t = re.sub(r'\s+', ' ', t).strip()
    return '' if is_noise_text(t) else t


def clean_code_text(text):
    return clean_text_artifacts(text, preserve_newlines=True)


def clean_formula_text(text):
    t = clean_text_artifacts(text)
    if t.count('|') >= 3:
        t = t.replace('|', '')
    t = re.sub(r'\s+', ' ', t).strip()
    return t


def add_caption(text, role='figure_caption'):
    caption_text = normalize_caption(text)
    if role == 'table_caption':
        m = re.match(r'^表\s*(\d+)(?:[-.](\d+))?', caption_text)
        if m:
            ch = int(m.group(1) or 0)
            no = int(m.group(2) or 0)
            if no:
                TABLE_COUNTERS[ch] = max(TABLE_COUNTERS.get(ch, 0), no)
    p = add_text(caption_text, role=role, first_indent=False)
    if p is None:
        return None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    if role == 'table_caption':
        p.paragraph_format.keep_with_next = True
    return p


FORMULA_COUNTERS = {}
TABLE_COUNTERS = {}


def chapter_number_from_heading(text):
    t = str(text or '').strip()
    m = re.match(r'^第(\d+)章', t)
    if m:
        return int(m.group(1))
    cn = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}
    m = re.match(r'^第([一二三四五六七八九十])章', t)
    if m:
        return cn.get(m.group(1))
    m = re.match(r'^(\d+)(?:\.|\s)', t)
    return int(m.group(1)) if m else None


def strip_heading_number(text):
    t = str(text or '').strip()
    t = re.sub(r'^第[一二三四五六七八九十百千万\d]+章\s*', '', t)
    t = re.sub(r'^\d+(?:\.\d+)*\s*', '', t)
    return t.strip()


def is_table_item(item):
    return isinstance(item, dict) and item.get('table_rows') and item.get('role') != 'code'


def is_code_table_item(item):
    return isinstance(item, dict) and item.get('table_rows') and (item.get('role') == 'code' or rows_look_like_code(item.get('table_rows') or []))


def looks_like_table_title(text):
    t = strip_heading_number(clean_text_artifacts(text))
    if not t or len(t) > 50:
        return False
    if re.match(r'^(图|表)\s*\d+', t) or re.match(r'^代码\s*\d+', t):
        return False
    if re.search(r'[。！？；;=<>]|如下|所示', t):
        return False
    return True


def next_table_caption(title, chapter=None):
    title = strip_heading_number(title)
    ch = chapter or 0
    TABLE_COUNTERS[ch] = TABLE_COUNTERS.get(ch, 0) + 1
    label = f'{ch}-{TABLE_COUNTERS[ch]}' if ch else str(TABLE_COUNTERS[ch])
    return f'表 {label} {title}'.strip()


def latex_escape_text(text):
    return str(text or '').replace('\\', r'\backslash ').replace('{', r'\{').replace('}', r'\}')


def latex_text_arg(text):
    return str(text or '').replace('\\', r'\backslash ').replace('{', r'\{').replace('}', r'\}')


def split_formula_number(text):
    t = str(text or '').strip()
    m = re.search(r'[（(]\s*(\d+\s*[-.]\s*\d+)\s*[）)]\s*$', t)
    if not m:
        return t, ''
    label = re.sub(r'\s+', '', m.group(1)).replace('.', '-')
    return t[:m.start()].strip(), label


def formula_token_to_latex(token):
    if not token:
        return ''
    if re.fullmatch(r'[\u4e00-\u9fff]+', token):
        return r'\text{' + latex_text_arg(token) + '}'
    if re.fullmatch(r'[A-Za-z]+', token):
        return r'\mathrm{' + token + '}'
    mapping = {
        '×': r'\times',
        '÷': r'\div',
        '≤': r'\leq',
        '≥': r'\geq',
        '≈': r'\approx',
        '≒': r'\approx',
        '％': r'\%',
        '%': r'\%',
        '²': '^{2}',
        '³': '^{3}',
        '（': '(',
        '）': ')',
        '，': ',',
        '。': '.',
        '：': '=',
        '＝': '=',
        '＋': '+',
        '－': '-',
    }
    return mapping.get(token, token)


def expression_to_latex(text):
    s = str(text or '').strip()
    if not s:
        return ''
    tokens = re.findall(r'[\u4e00-\u9fff]+|[A-Za-z]+|\d+(?:,\d{3})*(?:\.\d+)?|[²³]|.', s)
    out = []
    for tok in tokens:
        if tok.isspace():
            continue
        out.append(formula_token_to_latex(tok))
    return ''.join(out)


def formula_colon_split(text):
    t = str(text or '').strip()
    for sep in ('：', ':'):
        if sep in t:
            left, right = t.split(sep, 1)
            if left.strip() and right.strip() and re.search(r'\d|[=＝+\-*/×÷]', right):
                return left.strip(), right.strip()
    return '', t


def text_formula_to_latex(text):
    body, existing_label = split_formula_number(clean_formula_text(text))
    if not body:
        return '', existing_label
    body = body.replace('＝', '=').replace('＋', '+').replace('－', '-')
    return expression_to_latex(body), existing_label


def formula_latex_from_text(text):
    t = clean_formula_text(text)
    if not t:
        return ''
    if t.startswith('$') and t.endswith('$'):
        return t.strip('$').strip()
    latex, _existing_label = text_formula_to_latex(t)
    return latex


def formula_has_number(text):
    return bool(re.search(r'[（(]\s*\d+\s*[-.]\s*\d+\s*[）)]\s*$', str(text or '').strip()))


def next_formula_label(chapter):
    ch = chapter or 0
    FORMULA_COUNTERS[ch] = FORMULA_COUNTERS.get(ch, 0) + 1
    return f'{ch}-{FORMULA_COUNTERS[ch]}' if ch else str(FORMULA_COUNTERS[ch])


def render_plain_formula(text, chapter=None):
    text = clean_formula_text(text)
    if not text:
        return None
    rules = DATA.get('rules') or {}
    if rules.get('formula_numbered') and not formula_has_number(text):
        text = text + '(' + next_formula_label(chapter) + ')'
    p = doc.add_paragraph()
    apply_paragraph_profile(p, profile('formula'), first_indent_override=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    apply_run_profile(r, profile('formula'), text)
    return p


def render_formula(item, chapter=None):
    if isinstance(item, str):
        item = {'text': item}
    latex = str(item.get('latex') or '').strip()
    text = clean_formula_text(item.get('text') or '')
    xml = item.get('xml')
    existing_label = ''
    if not latex and not xml:
        latex, existing_label = text_formula_to_latex(text)
    elif text:
        _body, existing_label = split_formula_number(text)
    if not xml and not latex:
        return render_plain_formula(text, chapter)
    rules = DATA.get('rules') or {}
    numbered = item.get('numbered')
    should_number = bool(existing_label) or (bool(rules.get('formula_numbered')) if numbered is None else bool(numbered))
    if latex and should_number and r'\tag' not in latex and r'\begin{equation}' not in latex and r'\begin{align}' not in latex:
        label = existing_label or next_formula_label(chapter)
        latex = latex + r'\tag{' + label + '}'
    p = doc.add_paragraph()
    apply_paragraph_profile(p, profile('formula'), first_indent_override=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        xml_str = xml or latex_to_omath(latex, display=True)
        p._element.append(etree.fromstring(xml_str.encode('utf-8') if isinstance(xml_str, str) else xml_str))
    except Exception:
        r = p.add_run(text or latex)
        apply_run_profile(r, profile('formula'), text or latex)
    return p


def looks_like_code_line(text):
    t = str(text or '').strip()
    if not t or len(t) > 240:
        return False
    if re.match(r'^[A-Za-z0-9_.-]+[>#]', t):
        return True
    if re.match(r'^(interface|vlan|ip route|ip address|router|switchport|acl|rule|nat|dhcp|dns|ospf|bgp|display|show|ping|tracert|undo|quit|return|sysname|description|gateway|firewall|security-policy)\b', t, re.I):
        return True
    if re.match(r'^[a-z][a-z0-9_-]+\s+[-A-Za-z0-9_/.:]+', t) and any(ch in t for ch in ['/', '.', '-', '_']):
        return True
    return False




def rows_look_like_code(rows):
    flat = []
    for row in rows or []:
        for cell in row or []:
            for line in str(cell or '').splitlines():
                if line.strip():
                    flat.append(line.strip())
    if not flat:
        return False
    if len(rows or []) >= 4 and max((len(r) for r in rows or []), default=0) <= 2:
        code_hits = sum(1 for x in flat if looks_like_code_line(x))
        return code_hits >= max(2, len(flat) // 3)
    if max((len(r) for r in rows or []), default=0) == 1 and len(flat) >= 2:
        return sum(1 for x in flat if looks_like_code_line(x)) >= 2
    return False


def code_text_from_rows(rows):
    lines = []
    for row in rows or []:
        cells = [str(c or '').rstrip() for c in row]
        if len(cells) == 1:
            lines.append(cells[0])
        else:
            lines.append('    '.join(cells).rstrip())
    return clean_code_text('\n'.join(lines).rstrip())

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is not None:
        tcPr.remove(tcMar)
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        el = OxmlElement('w:' + side)
        el.set(qn('w:w'), str(int(val)))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def add_code_block(text):
    """Render code/configuration as a bordered block.

    A one-cell table is used instead of a normal paragraph so the output has a
    real solid frame, which is the conventional way to present command/config
    blocks in thesis appendices and network-design papers.  The detection of
    code remains semantic; no vendor, school, or fixed heading text is used.
    """
    text = clean_code_text(text)
    if not text:
        return None
    prof = profile('code')
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    try:
        _w = text_width_cm()
        table.columns[0].width = Cm(_w)
        table.rows[0].cells[0].width = Cm(_w)
    except Exception:
        pass
    cell = table.rows[0].cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_borders(cell,
                     top={'val': 'single', 'sz': '8', 'color': '000000'},
                     left={'val': 'single', 'sz': '8', 'color': '000000'},
                     bottom={'val': 'single', 'sz': '8', 'color': '000000'},
                     right={'val': 'single', 'sz': '8', 'color': '000000'})
    set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
    cell.text = ''
    lines = text.splitlines() or ['']
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        apply_paragraph_profile(p, prof, first_indent_override=0)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        apply_run_profile(r, prof, line)
    # Add a tiny spacing paragraph after the code box so following text does not
    # visually touch the frame; it has no text and therefore cannot enter TOC.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    spacer.paragraph_format.line_spacing = 1
    return table


def apply_three_line_borders(table):
    rows = len(table.rows)
    if rows == 0:
        return
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            sides = {'top': 'nil', 'left': 'nil', 'bottom': 'nil', 'right': 'nil', 'insideH': 'nil', 'insideV': 'nil'}
            if ri == 0:
                sides['top'] = {'val': 'single', 'sz': '12', 'color': '000000'}
                sides['bottom'] = {'val': 'single', 'sz': '8', 'color': '000000'}
            if ri == rows - 1:
                sides['bottom'] = {'val': 'single', 'sz': '12', 'color': '000000'}
            set_cell_borders(cell, **sides)


def repeat_table_header(row):
    try:
        trPr = row._tr.get_or_add_trPr()
        tbl_header = trPr.find(qn('w:tblHeader'))
        if tbl_header is None:
            tbl_header = OxmlElement('w:tblHeader')
            trPr.append(tbl_header)
        tbl_header.set(qn('w:val'), 'true')
    except Exception:
        pass


def prevent_row_split(row):
    try:
        trPr = row._tr.get_or_add_trPr()
        cant = trPr.find(qn('w:cantSplit'))
        if cant is None:
            cant = OxmlElement('w:cantSplit')
            trPr.append(cant)
    except Exception:
        pass


def should_keep_table_together(rows):
    if not rows or len(rows) > 10:
        return False
    text_cells = [str(cell or '') for row in rows for cell in row]
    if any(len(cell) > 90 for cell in text_cells):
        return False
    estimated_lines = 0
    for row in rows:
        row_lines = 1
        for cell in row:
            parts = str(cell or '').split('\n') or ['']
            cell_lines = sum(max(1, math.ceil(len(part) / 14)) for part in parts)
            row_lines = max(row_lines, cell_lines)
        estimated_lines += row_lines
    return estimated_lines <= 18


def keep_table_together(table):
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True
                if ri < len(table.rows) - 1:
                    p.paragraph_format.keep_with_next = True


def render_table(rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        if ri == 0:
            repeat_table_header(table.rows[ri])
        prevent_row_split(table.rows[ri])
        prof = profile('table_header' if ri == 0 else 'table_body')
        for ci in range(ncols):
            text = row[ci] if ci < len(row) else ''
            cell = table.rows[ri].cells[ci]
            cell.text = ''
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            parts = str(text or '').split('\n') or ['']
            for pi, part in enumerate(parts):
                p = cell.paragraphs[0] if pi == 0 else cell.add_paragraph()
                apply_paragraph_profile(p, prof, first_indent_override=0)
                r = p.add_run(part)
                apply_run_profile(r, prof, part)
    apply_three_line_borders(table)
    if should_keep_table_together(rows):
        keep_table_together(table)
    return table


def keep_paragraph_with_previous(p):
    # Word has no high-level python-docx flag for keep-with-previous.  We use
    # keepNext on the preceding image paragraph and keepLines on the caption so
    # a figure title is less likely to drift to the next page alone.
    try:
        pPr = p._element.get_or_add_pPr()
        keep = pPr.find(qn('w:keepLines'))
        if keep is None:
            keep = OxmlElement('w:keepLines')
            pPr.append(keep)
    except Exception:
        pass


def render_image(filename, caption=''):
    img_dir = DATA.get('images_dir') or ''
    candidates = []
    if os.path.isabs(img_dir):
        candidates.append(os.path.join(img_dir, filename))
    candidates += [
        os.path.join(BASE, img_dir, filename),
        os.path.abspath(os.path.join(os.getcwd(), img_dir, filename)),
        os.path.abspath(os.path.join(BASE, '..', img_dir, filename)),
        os.path.join(BASE, 'figures', filename),
    ]
    path = next((p for p in candidates if p and os.path.exists(p)), None)
    if not path:
        return None
    p = doc.add_paragraph()
    configure_picture_paragraph(p, keep_with_next=bool(caption))
    r = p.add_run()
    try:
        width, height = fit_picture_dimensions(path, has_caption=bool(caption))
        r.add_picture(path, width=width, height=height)
    except Exception:
        return None
    if caption:
        cap = add_caption(caption, 'figure_caption')
        cap.paragraph_format.keep_together = True
        keep_paragraph_with_previous(cap)
    return p


def clean_ref_text(ref):
    text = re.sub(r'\s+', ' ', str(ref or '')).strip()
    if text.startswith('[') and ']' in text:
        prefix, rest = text.split(']', 1)
        if prefix[1:].isdigit():
            return rest.strip()
    parts = text.split(None, 1)
    if len(parts) == 2 and parts[0].strip('.、[]').isdigit():
        return parts[1].strip()
    return text


def split_refs_backmatter(refs):
    pure, ack, app = [], [], []
    mode = 'refs'
    for raw in refs or []:
        if isinstance(raw, dict):
            if mode == 'ack':
                ack.append(raw)
            elif mode == 'app':
                app.append(raw)
            continue
        text = str(raw or '').strip()
        if not text:
            continue
        if is_ack_heading(text):
            mode = 'ack'; continue
        if is_appendix_heading(text):
            mode = 'app'
            if normalize_label(text) != '附录':
                app.append(text)
            continue
        (pure if mode == 'refs' else ack if mode == 'ack' else app).append(text)
    return pure, ack, app


def add_reference_mixed_runs(p, text, prof):
    # Chinese parts use the role's CJK font; Latin/numeric punctuation uses Times New Roman.
    for seg in re.findall(r'[\u4e00-\u9fff]+|[^\u4e00-\u9fff]+', text):
        r = p.add_run(seg)
        if has_cjk(seg):
            apply_run_profile(r, prof, seg, force_latin='Times New Roman')
        else:
            p_latin = dict(prof); p_latin['font'] = 'Times New Roman'
            apply_run_profile(r, p_latin, seg, force_latin='Times New Roman')


def render_reference_entries(refs):
    if not refs:
        return
    doc.add_page_break()
    add_text('参考文献', role='reference_heading', first_indent=False, outline_level=0)
    base_prof = profile('reference')
    hang_chars = (DATA.get('rules') or {}).get('reference_hanging_chars')
    try:
        hang_cm = float(hang_chars) * float(base_prof.get('size') or 12) * 0.0352778 if hang_chars else 0.74
    except Exception:
        hang_cm = 0.74
    for idx, raw in enumerate(refs, 1):
        if isinstance(raw, dict):
            continue
        text = clean_ref_text(raw)
        if not text:
            continue
        prof = profile('reference_english') if (DATA.get('rules') or {}).get('reference_english_left') and ascii_ratio(text[:120]) > 0.55 else base_prof
        p = doc.add_paragraph()
        apply_paragraph_profile(p, prof, first_indent_override=0)
        p.paragraph_format.left_indent = Cm(hang_cm)
        p.paragraph_format.first_line_indent = Cm(-hang_cm)
        p.paragraph_format.keep_together = True
        add_reference_mixed_runs(p, '[' + str(idx) + '] ' + text, prof)


def render_backmatter_section(title, paragraphs, code_sensitive=False):
    if not paragraphs:
        return
    doc.add_page_break()
    add_heading(title, 1)
    for item in paragraphs:
        if isinstance(item, dict) and (item.get('code') or item.get('role') == 'code'):
            add_code_block(item.get('code') or item.get('text') or '')
            continue
        text = str(item.get('text') if isinstance(item, dict) else item or '').strip()
        if not text:
            continue
        if code_sensitive and looks_like_code_line(text):
            add_code_block(text)
        else:
            add_text(text, role='body', first_indent=not code_sensitive)


def collect_structural_backmatter():
    ack, app = [], []
    mode = None
    for i, sec in enumerate(DATA.get('sections') or []):
        if is_front_section_index(i):
            continue
        h = (sec.get('heading') or '').strip()
        role = sec.get('role') or ''
        if role == 'acknowledgement' or is_ack_heading(h):
            mode = 'ack'; continue
        if role == 'appendix' or is_appendix_heading(h):
            mode = 'app'
            if normalize_label(h) != '附录':
                app.append(h)
            continue
        if mode in ('ack', 'app'):
            if sec.get('level') and h:
                (ack if mode == 'ack' else app).append(h)
            for para in sec.get('paragraphs', []) or []:
                (ack if mode == 'ack' else app).append(para)
    return ack, app


def render_paragraph_item(item, code_sensitive=False, chapter=None):
    if isinstance(item, dict) and (item.get('role') == 'formula' or item.get('math')):
        if item.get('math') and not item.get('latex') and not item.get('xml'):
            text = str(item.get('text') or '').strip()
            if text:
                add_text(text, role='body', first_indent=True)
            for m in item.get('math') or []:
                render_formula({
                    'latex': m.get('latex'),
                    'xml': m.get('xml'),
                    'text': m.get('text') or '',
                    'numbered': False,
                }, chapter)
        else:
            render_formula(item, chapter)
        return
    if isinstance(item, dict) and item.get('role') == 'figure':
        render_image(item.get('image') or item.get('filename') or item.get('asset') or '', item.get('caption') or '')
        return
    if isinstance(item, dict) and (item.get('role') == 'image' or item.get('image')):
        render_image(item.get('image') or item.get('filename') or item.get('asset') or '')
        return
    if isinstance(item, dict) and item.get('table_rows'):
        rows = item.get('table_rows') or []
        if item.get('role') == 'code' or rows_look_like_code(rows):
            add_code_block(item.get('code') or code_text_from_rows(rows))
        else:
            render_table(rows)
        return
    if isinstance(item, dict) and (item.get('role') == 'figure_caption'):
        add_caption(item.get('text') or '', 'figure_caption')
        return
    if isinstance(item, dict) and (item.get('role') == 'table_caption'):
        add_caption(item.get('text') or '', 'table_caption')
        return
    if isinstance(item, dict) and (item.get('code') or item.get('role') == 'code'):
        add_code_block(item.get('code') or item.get('text') or '')
        return
    text = str(item.get('text') if isinstance(item, dict) else item or '').strip()
    if not text:
        return
    if len(text) > 20 and any(k in text[:80] for k in ['完成后删除', '格式要求', '字体要求', '页眉页脚']):
        return
    if code_sensitive and looks_like_code_line(text):
        add_code_block(text)
    elif re.match(r'^图\s*\d+', text):
        add_caption(text, 'figure_caption')
    elif re.match(r'^表\s*\d+', text):
        add_caption(text, 'table_caption')
    else:
        add_text(text, role='body', first_indent=True)


def render_body():
    add_section_with_header('decimal', 1)
    fig_no = 0
    current_chapter = None
    for i, sec in enumerate(DATA.get('sections') or []):
        if is_front_section_index(i):
            continue
        h = (sec.get('heading') or '').strip()
        role = sec.get('role') or ''
        if is_reference_heading(h) or is_backmatter_heading(h) or role in ('references', 'acknowledgement', 'appendix'):
            continue
        if is_caption_heading(h):
            add_caption(h, 'figure_caption' if str(h).strip().startswith('图') else 'table_caption')
        elif h and h != '正文':
            add_heading(h, sec.get('level') or 1)
            if int(sec.get('level') or 1) == 1:
                current_chapter = chapter_number_from_heading(h) or current_chapter
        paragraphs = sec.get('paragraphs', []) or []
        has_inline_images = any(isinstance(x, dict) and (x.get('role') in ('image', 'figure') or x.get('image')) for x in paragraphs)
        # New content_parser keeps images in the paragraph stream.  For old
        # content.json files, fall back to section-level images, but do not
        # invent a caption from the heading because that caused figure-title
        # mismatch.
        if not has_inline_images:
            for img in sec.get('images', []) or []:
                render_image(img, '')
        idx = 0
        while idx < len(paragraphs):
            para = paragraphs[idx]
            nxt = paragraphs[idx + 1] if idx + 1 < len(paragraphs) else None
            if isinstance(para, str) and is_table_item(nxt) and looks_like_table_title(para):
                add_caption(next_table_caption(para, current_chapter), 'table_caption')
                render_table(nxt.get('table_rows') or [])
                idx += 2
                continue
            if is_table_item(para):
                prev = paragraphs[idx - 1] if idx > 0 else None
                has_caption = isinstance(prev, dict) and prev.get('role') == 'table_caption'
                if not has_caption and idx == 0 and h and looks_like_table_title(h):
                    add_caption(next_table_caption(h, current_chapter), 'table_caption')
            render_paragraph_item(para, code_sensitive=False, chapter=current_chapter)
            idx += 1
    pure_refs, ack_from_refs, app_from_refs = split_refs_backmatter(DATA.get('references') or [])
    ack_sections, app_sections = collect_structural_backmatter()
    render_reference_entries(pure_refs)
    render_backmatter_section('致  谢', ack_sections or ack_from_refs, code_sensitive=False)
    render_backmatter_section('附  录', app_sections or app_from_refs, code_sensitive=True)




def _norm_for_pdf_match(text):
    return re.sub(r'\s+', '', str(text or '')).lower()


def _extract_pdf_pages(pdf_path):
    exe = shutil.which('pdftotext')
    if not exe:
        return []
    with tempfile.TemporaryDirectory() as td:
        txt = os.path.join(td, 'out.txt')
        cmd = [exe, '-layout', pdf_path, txt]
        subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=60)
        if not os.path.exists(txt):
            return []
        data = open(txt, 'r', encoding='utf-8', errors='ignore').read()
    return data.split('\f')


def _make_pdf_for_pagination(docx_path):
    soffice = shutil.which('libreoffice') or shutil.which('soffice')
    if not soffice:
        for p in [r'C:\Program Files\LibreOffice\program\soffice.exe',
                  r'C:\Program Files (x86)\LibreOffice\program\soffice.exe']:
            if os.path.exists(p):
                soffice = p
                break
    if not soffice:
        return None
    td = tempfile.mkdtemp(prefix='toc_pages_')
    profile = os.path.join(td, 'lo_profile')
    home = os.path.join(td, 'home')
    os.makedirs(profile, exist_ok=True)
    os.makedirs(home, exist_ok=True)
    try:
        from pathlib import Path
        profile_uri = Path(profile).as_uri()
    except Exception:
        profile_uri = 'file://' + profile.replace(' ', '%20')
    cmd = [
        soffice, '--headless', '--norestore', '--nofirststartwizard',
        f'-env:UserInstallation={profile_uri}',
        '--convert-to', 'pdf', '--outdir', td, docx_path,
    ]
    env = dict(os.environ)
    env['HOME'] = home
    subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=120, env=env)
    base = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
    pdf = os.path.join(td, base)
    return pdf if os.path.exists(pdf) and os.path.getsize(pdf) > 0 else None


def _infer_heading_pages_from_pdf(docx_path=None):
    """Render the current DOCX to PDF and infer static TOC page numbers.

    This is the no-hardcode TOC pass: page numbers are not guessed from chapter
    count, school name, or fixed page offsets.  The first rendered body heading
    defines Arabic page 1, then every TOC entry is located in the rendered PDF
    text after the TOC pages.
    """
    try:
        pdf = _make_pdf_for_pagination(docx_path or OUT)
        if not pdf:
            return {}
        pages = _extract_pdf_pages(pdf)
        if not pages:
            return {}
        norm_pages = [_norm_for_pdf_match(p) for p in pages]
        entries = collect_toc_entries()
        if not entries:
            return {}

        first = _norm_for_pdf_match(entries[0]['text'])
        toc_last = 0
        for i, text in enumerate(norm_pages):
            if _norm_for_pdf_match('目录') in text:
                toc_last = i

        body_start = None
        for i in range(toc_last + 1, len(norm_pages)):
            if first and first in norm_pages[i]:
                body_start = i
                break
        if body_start is None:
            return {}

        page_map = {}
        search_from = body_start
        for ent in entries:
            key = _norm_for_pdf_match(ent['text'])
            if not key:
                continue
            found = None
            for i in range(search_from, len(norm_pages)):
                if key in norm_pages[i]:
                    found = i
                    break
            if found is None:
                # Some PDF extractors insert or drop punctuation/spaces.  Fall
                # back to a looser key built from the first substantial token.
                loose = re.sub(r'[^0-9a-zA-Z\u4e00-\u9fff]+', '', key)[:16]
                if loose:
                    for i in range(search_from, len(norm_pages)):
                        if loose in norm_pages[i]:
                            found = i
                            break
            if found is not None:
                page_map[key] = found - body_start + 1
                search_from = min(found, len(norm_pages) - 1)
        return page_map
    except Exception:
        return {}


def _rewrite_static_toc_pages(page_map):
    if not page_map:
        return False
    try:
        d = Document(OUT)
        in_toc = False
        changed = False
        for p in d.paragraphs:
            txt = p.text.strip()
            if txt == '目录':
                in_toc = True
                continue
            if not in_toc:
                continue
            pPr = p._element.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                break
            if '\t' not in p.text:
                continue
            label = p.text.split('\t', 1)[0].strip()
            key = _norm_for_pdf_match(label)
            if key not in page_map:
                continue
            # Preserve paragraph formatting; rebuild simple runs only.
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            prof = profile('body')
            r1 = p.add_run(label)
            apply_run_profile(r1, prof, label)
            p.add_run('\t')
            r2 = p.add_run(str(page_map[key]))
            apply_run_profile(r2, prof, str(page_map[key]))
            changed = True
        if changed:
            d.save(OUT)
        return changed
    except Exception:
        return False


def update_static_toc_pages():
    page_map = _infer_heading_pages_from_pdf()
    _rewrite_static_toc_pages(page_map)


def _infer_heading_pages_from_word_com(docx_path=None):
    """Use Word pagination to compute TOC page numbers without updating fields."""
    try:
        import win32com.client  # type: ignore
    except Exception:
        return {}
    word = None
    doc_obj = None
    try:
        entries = collect_toc_entries()
        if not entries:
            return {}
        entry_keys = [_norm_for_pdf_match(e.get('text') or '') for e in entries]
        wanted = set(k for k in entry_keys if k)
        found = {}
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False
        doc_obj = word.Documents.Open(os.path.abspath(docx_path or OUT), ReadOnly=True)
        try:
            doc_obj.Repaginate()
        except Exception:
            pass
        for para in doc_obj.Paragraphs:
            try:
                if int(para.OutlineLevel) not in (1, 2, 3):
                    continue
            except Exception:
                continue
            text = str(para.Range.Text or '').replace('\r', '').replace('\x07', '').strip()
            key = _norm_for_pdf_match(text)
            if key in wanted and key not in found:
                try:
                    found[key] = int(para.Range.Information(3))
                except Exception:
                    pass
            if len(found) >= len(wanted):
                break
        if not found:
            return {}
        first_key = next((k for k in entry_keys if k in found), None)
        if not first_key:
            return {}
        first_page = found[first_key]
        return {k: max(1, v - first_page + 1) for k, v in found.items()}
    except Exception:
        return {}
    finally:
        try:
            if doc_obj is not None:
                doc_obj.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass


def update_fields_with_word_com():
    """Update TOC/fields through Microsoft Word when available on Windows."""
    try:
        import win32com.client  # type: ignore
    except Exception:
        return False
    word = None
    doc_obj = None
    try:
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False
        doc_obj = word.Documents.Open(os.path.abspath(OUT), ReadOnly=False)
        try:
            for toc in doc_obj.TablesOfContents:
                toc.Update()
        except Exception:
            pass
        try:
            doc_obj.Fields.Update()
        except Exception:
            pass
        doc_obj.Save()
        return True
    except Exception:
        return False
    finally:
        try:
            if doc_obj is not None:
                doc_obj.Close(SaveChanges=True)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass


def build_document(toc_page_map=None, native_toc=True):
    """Build the whole DOCX once, optionally with resolved static TOC pages."""
    global doc, TOC_PAGE_MAP, USE_NATIVE_TOC, FORMULA_COUNTERS, TABLE_COUNTERS
    TOC_PAGE_MAP = dict(toc_page_map or {})
    USE_NATIVE_TOC = bool(native_toc)
    FORMULA_COUNTERS = {}
    TABLE_COUNTERS = {}
    doc = Document()
    configure_global_styles()
    setup_section(doc.sections[0])
    clear_header_footer(doc.sections[0])
    remove_initial_empty_paragraph()
    render_cover_and_declarations()
    render_front_matter()
    render_body()
    force_cover_headerless()
    doc.save(OUT)


def main():
    build_document({}, native_toc=False)
    page_map = _infer_heading_pages_from_word_com()
    if page_map:
        build_document(page_map, native_toc=False)
    suffix = 'static TOC pages resolved by Word COM' if page_map else 'static TOC generated without resolved page numbers'
    print(f'Saved: {OUT}  ({suffix})')


if __name__ == '__main__':
    main()
'''



def _heading_num_tuple(text: str) -> Optional[tuple]:
    m = re.match(r'^(\d+(?:\.\d+)*)\b', str(text or '').strip())
    if not m:
        return None
    try:
        return tuple(int(x) for x in m.group(1).split('.'))
    except Exception:
        return None


def _chapter_num(text: str) -> Optional[int]:
    s = str(text or '').strip()
    m = re.match(r'^第(\d+)章', s)
    if m:
        return int(m.group(1))
    cn = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}
    m = re.match(r'^第([一二三四五六七八九十]+)章', s)
    if m:
        t = m.group(1)
        if t == '十':
            return 10
        if t.startswith('十'):
            return 10 + cn.get(t[1:], 0)
        if t.endswith('十'):
            return cn.get(t[0], 1) * 10
        if '十' in t:
            a, b = t.split('十', 1)
            return cn.get(a, 1) * 10 + cn.get(b, 0)
        return cn.get(t)
    return None


def _sort_h3_inside_h2(block: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    if len(block) <= 2:
        return block
    head = block[0]
    chunks: List[List[Dict[str, Any]]] = []
    cur: List[Dict[str, Any]] = []
    for sec in block[1:]:
        num = _heading_num_tuple(sec.get('heading', ''))
        if sec.get('level') == 3 and num and len(num) == 3:
            if cur:
                chunks.append(cur)
            cur = [sec]
        else:
            if cur:
                cur.append(sec)
            else:
                chunks.append([sec])
    if cur:
        chunks.append(cur)
    nums = [_heading_num_tuple(c[0].get('heading', '')) for c in chunks]
    if len(chunks) >= 2 and all(nums) and nums != sorted(nums):
        chunks = sorted(chunks, key=lambda c: _heading_num_tuple(c[0].get('heading', '')))
    return [head] + [sec for c in chunks for sec in c]


def _normalize_numbered_section_order(sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Safely reorder numeric subsection blocks within each chapter.

    It fixes order drift such as 2.3 appearing after 2.9 while preserving every
    paragraph/table/image.  It does not rename duplicate numbers, because that
    is a content judgment rather than a formatting operation.
    """
    out: List[Dict[str, Any]] = []
    i = 0
    n = len(sections or [])
    while i < n:
        sec = sections[i]
        ch = _chapter_num(sec.get('heading', '')) if sec.get('level') == 1 else None
        if ch is None:
            out.append(sec)
            i += 1
            continue
        out.append(sec)
        i += 1
        chapter_items: List[Dict[str, Any]] = []
        while i < n:
            nxt = sections[i]
            if nxt.get('level') == 1 and _chapter_num(nxt.get('heading', '')) is not None:
                break
            chapter_items.append(nxt)
            i += 1
        blocks: List[List[Dict[str, Any]]] = []
        cur: List[Dict[str, Any]] = []
        for item in chapter_items:
            num = _heading_num_tuple(item.get('heading', ''))
            if item.get('level') == 2 and num and len(num) == 2 and num[0] == ch:
                if cur:
                    blocks.append(cur)
                cur = [item]
            else:
                if cur:
                    cur.append(item)
                else:
                    blocks.append([item])
        if cur:
            blocks.append(cur)
        nums = [_heading_num_tuple(b[0].get('heading', '')) for b in blocks]
        if blocks and all(num and len(num) == 2 and num[0] == ch for num in nums) and nums != sorted(nums):
            blocks = sorted(blocks, key=lambda b: _heading_num_tuple(b[0].get('heading', '')))
        for b in blocks:
            out.extend(_sort_h3_inside_h2(b))
    return out

def generate(format_json_path: str, content_json_path: str, output_dir: str, output_docx_name: str = '最终论文.docx') -> int:
    os.makedirs(output_dir, exist_ok=True)
    with open(format_json_path, 'r', encoding='utf-8') as f:
        fmt = json.load(f)
    with open(content_json_path, 'r', encoding='utf-8') as f:
        cnt = json.load(f)

    profiles = _infer_style_profiles(fmt)
    page = _extract_page_and_header(fmt)
    front = _front_matter_sections(cnt)
    cover_info = cnt.get('cover_info', {}) or {}
    title_cn = cover_info.get('paper_title') or cnt.get('title_info', {}).get('title_cn') or ''
    rules = _infer_template_rules(fmt)

    src_assets = fmt.get('_meta', {}).get('assets_dir') or ''
    runtime_assets = src_assets
    if src_assets and os.path.isdir(src_assets):
        dst_assets = os.path.join(output_dir, 'assets')
        if os.path.abspath(src_assets) != os.path.abspath(dst_assets):
            shutil.rmtree(dst_assets, ignore_errors=True)
            shutil.copytree(src_assets, dst_assets)
        runtime_assets = 'assets'

    src_images = cnt.get('_meta', {}).get('images_dir') or ''
    runtime_images = src_images
    if src_images and os.path.isdir(src_images):
        dst_images = os.path.join(output_dir, 'figures')
        if os.path.abspath(src_images) != os.path.abspath(dst_images):
            shutil.rmtree(dst_images, ignore_errors=True)
            shutil.copytree(src_images, dst_images)
        runtime_images = 'figures'

    latex_src = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'latex_omath.py')
    if os.path.exists(latex_src):
        shutil.copy2(latex_src, os.path.join(output_dir, 'latex_omath.py'))

    normalized_sections = _normalize_numbered_section_order(cnt.get('sections', []))

    data_blob = {
        'fmt_meta': fmt.get('_meta', {}),
        'content_meta': cnt.get('_meta', {}),
        'page': page,
        'profiles': profiles,
        'cover': fmt.get('cover', []),
        'cover_info': cover_info,
        'title_cn': title_cn,
        'sections': normalized_sections,
        'references': cnt.get('references', []),
        'front_indices': sorted(front['front_indices']),
        'front': {k: v for k, v in front.items() if k != 'front_indices'},
        'images_dir': runtime_images,
        'assets_dir': runtime_assets,
        'rules': rules,
    }
    blob = json.dumps(data_blob, ensure_ascii=False, indent=2)
    code = RUNTIME_TEMPLATE.replace('__DATA_BLOB__', repr(blob)).replace('__OUT_DOCX__', repr(os.path.basename(output_docx_name)))
    out_py = os.path.join(output_dir, 'build_generated.py')
    with open(out_py, 'w', encoding='utf-8') as f:
        f.write(code)
    return len(code)


if __name__ == '__main__':
    import sys
    if len(sys.argv) < 4:
        print('Usage: python script_generator.py format.json content.json output_dir [output.docx]')
        raise SystemExit(2)
    out_name = sys.argv[4] if len(sys.argv) > 4 else '最终论文.docx'
    n = generate(sys.argv[1], sys.argv[2], sys.argv[3], out_name)
    print(f'Generated build script, {n} chars')
