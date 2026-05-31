"""Template format extraction orchestration."""
from __future__ import annotations

import hashlib
import os

from docx import Document

try:
    from path_safety import ensure_safe_output_dir, safe_rmtree_generated_child
    from format_extractor_modules.cover import extract_cover as _extract_cover
    from format_extractor_modules.ooxml import ALIGN_MAP, paragraph_metrics as _paragraph_metrics
    from format_extractor_modules.pdf_template import extract_pdf_template as _extract_pdf_template
    from format_extractor_modules.style_profiles import build_style_profiles as _build_style_profiles
    from format_extractor_modules.style_resolver import StyleResolver
except ImportError:  # pragma: no cover - package-style imports
    from ..path_safety import ensure_safe_output_dir, safe_rmtree_generated_child
    from .cover import extract_cover as _extract_cover
    from .ooxml import ALIGN_MAP, paragraph_metrics as _paragraph_metrics
    from .pdf_template import extract_pdf_template as _extract_pdf_template
    from .style_profiles import build_style_profiles as _build_style_profiles
    from .style_resolver import StyleResolver

def _default_output_dir():
    return os.path.abspath(os.path.join(os.getcwd(), "Outputs", "_format_extractor_extract"))


def extract(docx_path, output_dir=None):
    if str(docx_path).lower().endswith(".pdf"):
        return _extract_pdf_template(docx_path)
    return extract_docx_template(docx_path, output_dir=output_dir)


def extract_docx_template(docx_path, output_dir=None):
    """Extract all template formatting. Returns (format_dict, markdown_report)."""
    doc = Document(docx_path)
    resolver = StyleResolver(doc)

    fmt = {
        "_meta": {
            "source": os.path.basename(docx_path),
            "sha256": hashlib.sha256(open(docx_path, "rb").read()).hexdigest()[:16],
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "sections": len(doc.sections),
        },
        "sections": [],
        "paragraphs": [],
        "tables": [],
    }

    md_lines = []
    md_lines.append(f"# 模版格式提取 — {os.path.basename(docx_path)}\n")
    md_lines.append(f"**段落**: {len(doc.paragraphs)} | **表格**: {len(doc.tables)} | **节**: {len(doc.sections)}\n")

    md_lines.append("## 页面设置\n")
    for index, section in enumerate(doc.sections):
        section_info = {
            "index": index,
            "page_width_cm": round(section.page_width.cm, 1),
            "page_height_cm": round(section.page_height.cm, 1),
            "margin_top_cm": round(section.top_margin.cm, 1),
            "margin_bottom_cm": round(section.bottom_margin.cm, 1),
            "margin_left_cm": round(section.left_margin.cm, 1),
            "margin_right_cm": round(section.right_margin.cm, 1),
            "diff_first_page": section.different_first_page_header_footer,
            "header": [],
            "footer": [],
        }
        if section.header:
            for paragraph in section.header.paragraphs:
                runs = []
                for run in paragraph.runs:
                    info = resolver.resolve(paragraph._element, run._element)
                    runs.append(
                        {
                            "text": run.text,
                            "font": info["font"],
                            "size_pt": info["size"],
                            "bold": info["bold"],
                            "italic": info["italic"],
                        }
                    )
                section_info["header"].append({"text": paragraph.text, "alignment": ALIGN_MAP.get(paragraph.alignment, "CENTER"), "runs": runs})
        if section.footer:
            for paragraph in section.footer.paragraphs:
                section_info["footer"].append({"text": paragraph.text, "alignment": ALIGN_MAP.get(paragraph.alignment, "CENTER")})
        fmt["sections"].append(section_info)

        md_lines.append(f'**节{index}**: {section_info["page_width_cm"]}x{section_info["page_height_cm"]}cm')
        md_lines.append(
            f'  边距: T{section_info["margin_top_cm"]} B{section_info["margin_bottom_cm"]} '
            f'L{section_info["margin_left_cm"]} R{section_info["margin_right_cm"]}'
        )
        for header in section_info["header"]:
            md_lines.append(f'  页眉: {header["alignment"]} | {header["text"][:100]}')
        for footer in section_info["footer"]:
            md_lines.append(f'  页脚: {footer["alignment"]} | {footer["text"][:100]}')
        md_lines.append("")

    md_lines.append("## 正文格式\n")
    for index, paragraph in enumerate(doc.paragraphs):
        pinfo = {
            "index": index,
            "style": paragraph.style.name if paragraph.style else "Normal",
            "text": paragraph.text,
            "runs": [],
            "has_page_break": False,
        }
        for run in paragraph.runs:
            info = resolver.resolve(paragraph._element, run._element)
            has_page_break = 'w:br w:type="page"' in run._element.xml or "w:br w:type='page'" in run._element.xml
            if has_page_break:
                pinfo["has_page_break"] = True
            pinfo["runs"].append(
                {
                    "text": run.text,
                    "font": info["font"],
                    "size_pt": info["size"],
                    "bold": info["bold"],
                    "italic": info["italic"],
                }
            )
            if "align" not in pinfo:
                pinfo["align"] = info["align"]
                pinfo["ls"] = info["ls"]
                pinfo["indent"] = info["indent"]
                pinfo["left_indent_cm"] = info.get("left_indent_cm")
                pinfo["right_indent_cm"] = info.get("right_indent_cm")
                pinfo["hanging_indent_cm"] = info.get("hanging_indent_cm")

        metrics = _paragraph_metrics(paragraph._element)
        if "align" not in pinfo:
            pinfo["align"] = metrics["alignment"]
            pinfo["ls"] = metrics["line_spacing_val"] or 1.15
            pinfo["indent"] = metrics["first_indent_cm"] or 0
        pinfo["alignment"] = metrics["alignment"] if metrics["alignment"] != "DEFAULT" else pinfo.get("align", "DEFAULT")
        pinfo["line_spacing_val"] = metrics["line_spacing_val"] if metrics["line_spacing_val"] is not None else pinfo.get("ls")
        pinfo["line_spacing_rule"] = metrics["line_spacing_rule"]
        pinfo["line_spacing_fixed_pt"] = metrics["line_spacing_fixed_pt"]
        pinfo["space_before_pt"] = metrics["space_before_pt"]
        pinfo["space_after_pt"] = metrics["space_after_pt"]
        pinfo["first_indent_cm"] = metrics["first_indent_cm"] if metrics["first_indent_cm"] is not None else pinfo.get("indent", 0)
        pinfo["left_indent_cm"] = metrics["left_indent_cm"] if metrics["left_indent_cm"] is not None else pinfo.get("left_indent_cm")
        pinfo["right_indent_cm"] = metrics["right_indent_cm"] if metrics["right_indent_cm"] is not None else pinfo.get("right_indent_cm")
        pinfo["hanging_indent_cm"] = metrics["hanging_indent_cm"] if metrics["hanging_indent_cm"] is not None else pinfo.get("hanging_indent_cm")

        fmt["paragraphs"].append(pinfo)

        if paragraph.text.strip() or pinfo["has_page_break"]:
            flags = f'align={pinfo.get("align", "?")} ls={pinfo.get("ls", "?")}'
            if pinfo.get("indent"):
                flags += f' indent={pinfo["indent"]}cm'
            if pinfo["has_page_break"]:
                flags += " [PAGE BREAK]"
            text = paragraph.text[:100].replace("\n", "\\n")
            run_formats = []
            for run in pinfo["runs"][:4]:
                parts = [f'{run.get("font", "?")}', f'{run.get("size_pt", "?")}pt']
                if run.get("bold"):
                    parts.append("B")
                if run.get("italic"):
                    parts.append("I")
                run_formats.append("|".join(parts))
            md_lines.append(f'**P{index}** [{pinfo["style"]}] {flags}')
            if run_formats:
                md_lines.append(f'  runs: {" / ".join(run_formats)}')
            if text:
                md_lines.append(f"  > {text}")
            md_lines.append("")

    md_lines.append("## 表格\n")
    for table_index, table in enumerate(doc.tables):
        table_info = {"index": table_index, "rows": len(table.rows), "cols": len(table.columns), "cells": []}
        for row_index, row in enumerate(table.rows):
            row_cells = []
            for cell_index, cell in enumerate(row.cells):
                cell_runs = []
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        info = resolver.resolve(paragraph._element, run._element)
                        cell_runs.append({"text": run.text, "font": info["font"], "size_pt": info["size"], "bold": info["bold"]})
                row_cells.append({"row": row_index, "col": cell_index, "text": cell.text, "runs": cell_runs})
            table_info["cells"].append(row_cells)
            if row_index < 5:
                md_lines.append(f'  Row{row_index}: {[cell["text"][:50] for cell in row_cells]}')
        fmt["tables"].append(table_info)
        md_lines.append("")

    md_lines.append("\n## 验证\n")
    md_lines.append(f'- 段落: JSON={len(fmt["paragraphs"])} docx={len(doc.paragraphs)} ✓')
    md_lines.append(f'- 表格: JSON={len(fmt["tables"])} docx={len(doc.tables)} ✓')
    md_lines.append(f'- 节:   JSON={len(fmt["sections"])} docx={len(doc.sections)} ✓')

    base = os.path.splitext(os.path.basename(docx_path))[0]
    asset_root = ensure_safe_output_dir(output_dir or _default_output_dir())
    asset_dir = os.path.join(asset_root, f"{base}_assets")
    safe_rmtree_generated_child(asset_dir, asset_root, allowed_suffixes=("_assets",))
    fmt["_meta"]["assets_dir"] = os.path.abspath(asset_dir)
    fmt["cover"] = _extract_cover(doc, asset_dir)

    normal_style = doc.styles["Normal"]
    fmt["normal_style"] = {
        "font_name": normal_style.font.name,
        "font_size_pt": round(normal_style.font.size / 12700, 1) if normal_style.font.size else None,
        "line_spacing": normal_style.paragraph_format.line_spacing,
        "line_spacing_rule": str(normal_style.paragraph_format.line_spacing_rule) if normal_style.paragraph_format.line_spacing_rule else None,
    }

    fmt["style_profiles"] = _build_style_profiles(fmt)
    return fmt, "\n".join(md_lines)


