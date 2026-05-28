"""DOCX content extraction orchestration."""
from __future__ import annotations

import hashlib
import os
import shutil

from docx import Document

try:
    from content_parser_modules.placeholders import (
        is_unfilled_placeholder_text as _is_unfilled_placeholder_text,
        placeholder_samples as _placeholder_samples,
    )
    from content_parser_modules.front_matter import extract_front_matter
    from content_parser_modules.text_cleaner import clean_text_artifacts as _clean_text_artifacts
    from content_parser_modules.body_dispatcher import parse_body_sections
    from content_parser_modules.image_extractor import ImageRegistry, non_body_image_entries as _non_body_image_entries
    from content_parser_modules.section_builder import (
        filter_content_sections,
        mark_first_body_page_break,
        postprocess_section_paragraphs,
    )
except ImportError:  # pragma: no cover - package-style imports
    from .placeholders import (
        is_unfilled_placeholder_text as _is_unfilled_placeholder_text,
        placeholder_samples as _placeholder_samples,
    )
    from .front_matter import extract_front_matter
    from .text_cleaner import clean_text_artifacts as _clean_text_artifacts
    from .body_dispatcher import parse_body_sections
    from .image_extractor import ImageRegistry, non_body_image_entries as _non_body_image_entries
    from .section_builder import (
        filter_content_sections,
        mark_first_body_page_break,
        postprocess_section_paragraphs,
    )

def _content_placeholder_samples(content, limit=8):
    samples = []

    def add(value):
        text = str(value or "").strip()
        if text and _is_unfilled_placeholder_text(text):
            samples.append(text[:120])

    for value in (content.get('title_info') or {}).values():
        add(value)
    for value in (content.get('cover_info') or {}).values():
        add(value)
    for section in content.get('sections') or []:
        add(section.get('heading'))
        for item in section.get('paragraphs') or []:
            if isinstance(item, str):
                add(item)
            elif isinstance(item, dict):
                add(item.get('text') or item.get('code'))
                for row in item.get('table_rows') or []:
                    for cell in row:
                        add(cell)
            if len(samples) >= limit:
                return samples[:limit]
    return samples[:limit]


def _count_structured_body_tables(sections):
    total = 0
    for section in sections or []:
        for item in section.get('paragraphs') or []:
            if isinstance(item, dict) and item.get('table_rows') and item.get('role') != 'code':
                total += 1
    return total


def extract(docx_path, output_dir='Inputs'):
    """Extract content from a content docx into structured JSON + copy images."""
    doc = Document(docx_path)
    base = os.path.splitext(os.path.basename(docx_path))[0]

    # Setup output dirs.  Recreate figures for each extraction so repeated
    # verification passes do not accumulate stale/duplicated files.
    content_dir = os.path.join(output_dir, base)
    fig_dir = os.path.join(content_dir, 'figures')
    shutil.rmtree(fig_dir, ignore_errors=True)
    os.makedirs(fig_dir, exist_ok=True)
    image_registry = ImageRegistry(fig_dir, f'{base}_img')

    content = {
        '_meta': {
            'source': os.path.basename(docx_path),
            'sha256': hashlib.sha256(open(docx_path, 'rb').read()).hexdigest()[:16],
            'paragraphs': len(doc.paragraphs),
            'source_tables_count': len(doc.tables),
            'tables_count': 0,
        },
        'title_info': {},
        'sections': [],
        'references': [],
    }
    source_placeholders = _placeholder_samples(doc.paragraphs)
    content['_meta']['source_placeholders'] = source_placeholders

    front_matter = extract_front_matter(doc, clean_text_func=_clean_text_artifacts)
    text_start = int(front_matter.get('text_start') or 0)
    content['title_info'].update(front_matter.get('title_info') or {})
    if front_matter.get('cover_info'):
        content['cover_info'] = front_matter['cover_info']

    body_result = parse_body_sections(doc, text_start, image_registry)
    content['sections'] = filter_content_sections(body_result.sections)
    if body_result.references:
        content['references'] = body_result.references

    mark_first_body_page_break(content['sections'])
    postprocess_section_paragraphs(content['sections'])
    content['_meta']['tables_count'] = _count_structured_body_tables(content['sections'])
    remaining_placeholders = _content_placeholder_samples(content)
    if remaining_placeholders:
        content['_meta']['remaining_placeholders'] = remaining_placeholders
    if body_result.placeholders_removed:
        content['_meta']['body_placeholders_removed'] = body_result.placeholders_removed
    if source_placeholders and not remaining_placeholders:
        content['_meta']['source_placeholders_auto_removed'] = len(source_placeholders)

    # Count saved images without running a second extraction pass.
    # Re-extracting here used to create duplicate filenames and made figure
    # captions drift away from their intended images.
    content['_meta']['images_extracted'] = len([
        f for f in os.listdir(fig_dir)
        if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tif', '.tiff'))
    ])
    content['_meta']['images_dir'] = os.path.abspath(fig_dir)
    content['_meta']['image_extract_failures'] = image_registry.failures
    content['_meta']['non_body_images'] = _non_body_image_entries(doc)
    if body_result.source_toc_skipped:
        content['_meta']['source_toc_skipped_paragraphs'] = body_result.source_toc_skipped

    return content


