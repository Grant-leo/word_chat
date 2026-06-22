"""Body element dispatch for content_parser.py.

This module owns the central DOCX body traversal.  The public parser keeps
front-matter, output directory, and metadata orchestration; body_dispatcher
decides how paragraph/table OOXML becomes sections, references, images, tables,
code blocks, captions, and formulas.
"""
from __future__ import annotations

from dataclasses import dataclass
import re
from typing import Any, Dict, List

from docx.text.paragraph import Paragraph

try:
    from formula_semantics import is_formula_problem_text
except ImportError:  # pragma: no cover - package-style imports
    from ..formula_semantics import is_formula_problem_text

try:
    from content_parser_modules.caption_flow import (
        is_figure_caption,
        is_table_caption,
        normalize_caption_spacing,
    )
    from content_parser_modules.formula_extractor import (
        _formula_item_from_text,
        _formula_problem_item_from_text,
        _looks_like_formula_text,
        _omml_text_looks_like_body,
        _rich_text_item_from_inline_formula_spans,
        set_clean_text_artifacts_func,
    )
    from content_parser_modules.heading_detector import (
        classify_section_role,
        detect_heading_level,
        is_backmatter_heading,
        normalize_heading_spacing,
    )
    from content_parser_modules.image_extractor import image_items_from_ooxml, images_from_run_ooxml
    from content_parser_modules.paragraph_stream import (
        append_stream_run_group,
        paragraph_visible_text,
        paragraph_stream_items,
    )
    from content_parser_modules.placeholders import is_unfilled_placeholder_text
    from content_parser_modules.placeholders import is_template_instruction_text
    from content_parser_modules.reference_collector import ReferenceCollector
    from content_parser_modules.section_builder import make_body_section
    from content_parser_modules.source_toc import (
        is_source_toc_title,
        source_toc_skip_count_after_title,
    )
    from content_parser_modules.table_extractor import (
        code_text_from_table_rows,
        extract_table_from_ooxml,
        looks_like_code_line,
        table_rows_look_like_code,
    )
    from content_parser_modules.text_cleaner import clean_code_text, clean_text_artifacts
except ImportError:  # pragma: no cover - package-style imports
    from .caption_flow import (
        is_figure_caption,
        is_table_caption,
        normalize_caption_spacing,
    )
    from .formula_extractor import (
        _formula_item_from_text,
        _formula_problem_item_from_text,
        _looks_like_formula_text,
        _omml_text_looks_like_body,
        _rich_text_item_from_inline_formula_spans,
        set_clean_text_artifacts_func,
    )
    from .heading_detector import (
        classify_section_role,
        detect_heading_level,
        is_backmatter_heading,
        normalize_heading_spacing,
    )
    from .image_extractor import image_items_from_ooxml, images_from_run_ooxml
    from .paragraph_stream import append_stream_run_group, paragraph_visible_text, paragraph_stream_items
    from .placeholders import is_unfilled_placeholder_text
    from .placeholders import is_template_instruction_text
    from .reference_collector import ReferenceCollector
    from .section_builder import make_body_section
    from .source_toc import is_source_toc_title, source_toc_skip_count_after_title
    from .table_extractor import (
        code_text_from_table_rows,
        extract_table_from_ooxml,
        looks_like_code_line,
        table_rows_look_like_code,
    )
    from .text_cleaner import clean_code_text, clean_text_artifacts


set_clean_text_artifacts_func(clean_text_artifacts)


_BODY_TRANSPARENT_CONTAINERS = {"customXml", "smartTag"}
_BODY_ACCEPTED_REVISION_CONTAINERS = {"ins", "moveTo"}
_BODY_DELETED_REVISION_CONTAINERS = {"del", "moveFrom"}


@dataclass
class BodyDispatchResult:
    sections: List[Dict[str, Any]]
    references: List[Any]
    source_toc_skipped: int = 0
    placeholders_removed: int = 0
    content_control_paragraphs_in_place: int = 0


def append_text_or_code(section: Dict[str, Any], text: str, in_appendix: bool = False) -> None:
    """Append semantic blocks while preserving captions, code and inline math."""
    if not text:
        return
    text = clean_text_artifacts(text)
    if not text:
        return
    if is_unfilled_placeholder_text(text) or is_template_instruction_text(text):
        return
    if is_figure_caption(text):
        section["paragraphs"].append({"role": "figure_caption", "text": normalize_caption_spacing(text)})
    elif is_table_caption(text):
        section["paragraphs"].append({"role": "table_caption", "text": normalize_caption_spacing(text)})
    elif re.match(r"^\s*\$\$.+\$\$\s*(?:[\(\uff08]\s*\d+(?:\s*[-.]\s*\d+)?\s*[\)\uff09])?\s*$", text, re.S):
        section["paragraphs"].append(_formula_item_from_text(text))
    else:
        rich_item = _rich_text_item_from_inline_formula_spans(text)
        rich_has_text = bool(
            rich_item
            and any(
                r.get("type") == "text"
                and str(r.get("text") or "").strip(" \t\r\n，,。.;；:：()（）")
                for r in rich_item.get("runs") or []
            )
        )
        if rich_item and rich_has_text:
            section["paragraphs"].append(rich_item)
            return
        if is_formula_problem_text(text):
            section["paragraphs"].append(_formula_problem_item_from_text(text))
        elif _looks_like_formula_text(text):
            section["paragraphs"].append(_formula_item_from_text(text))
        elif in_appendix and (looks_like_code_line(text) or "\n" in text):
            section["paragraphs"].append({"role": "code", "code": clean_code_text(text)})
        elif rich_item:
            section["paragraphs"].append(rich_item)
        elif _omml_text_looks_like_body(text):
            section["paragraphs"].append(text)
        else:
            section["paragraphs"].append(text)


def _section_from_heading(text: str, level: int) -> Dict[str, Any]:
    clean_heading = text.split("（")[0].strip()
    if is_unfilled_placeholder_text(clean_heading) or is_template_instruction_text(text):
        return {}
    if not clean_heading:
        return {}
    m = re.match(r"(?i)^(Abstract\s*:?|Key\s*words?\s*:?|摘要\s*[：:]|关键词\s*[：:])\s*", text)
    if m:
        heading_part = m.group(1).strip()
        body_part = text[m.end() :].strip()
        body_part = re.sub(r"^[（(][^）)]*[）)]\s*", "", body_part)
    else:
        heading_part = clean_heading
        body_part = ""
    heading_part = normalize_heading_spacing(heading_part)
    role = classify_section_role(heading_part, level)
    section: Dict[str, Any] = {
        "heading": heading_part,
        "level": level,
        "role": role,
        "paragraphs": [],
        "images": [],
    }
    if role == "en_abstract":
        section["page_break_before"] = True
    if body_part:
        append_text_or_code(section, body_part, in_appendix=False)
    return section


def _append_paragraph_stream(section: Dict[str, Any], paragraph: Any, image_registry: Any, notes: Dict[str, Dict[str, str]] | None = None) -> None:
    """Append text/image/math items from one Word paragraph in OOXML order."""
    text = paragraph_visible_text(paragraph).strip()
    stream_items = paragraph_stream_items(paragraph, image_registry, notes=notes)
    if not stream_items and text:
        stream_items = [{"role": "text", "text": text}]
    rich_runs: List[Dict[str, Any]] = []
    in_appendix = bool(re.search(r"(\u9644\s*\u5f55|\u914d\u7f6e|\u547d\u4ee4|\u4ee3\u7801)", section.get("heading", "")))

    def flush_rich_runs() -> None:
        nonlocal rich_runs
        append_stream_run_group(
            section,
            rich_runs,
            append_text_or_code_func=append_text_or_code,
            in_appendix=in_appendix,
        )
        rich_runs = []

    for item in stream_items:
        if item.get("role") == "image":
            flush_rich_runs()
            section["images"].append(item.get("image"))
            section["paragraphs"].append(item)
        elif item.get("role") == "math_inline":
            rich_runs.append({"type": "math", "text": item.get("text") or "", "math": item.get("math") or []})
        elif item.get("role") == "note_ref":
            rich_runs.append({
                "type": "note_ref",
                "note_type": item.get("note_type") or "footnote",
                "source_id": item.get("source_id") or "",
                "text": item.get("text") or "",
            })
        elif item.get("role") == "formula":
            flush_rich_runs()
            section["paragraphs"].append(item)
        elif item.get("role") == "text":
            txt = item.get("text") or ""
            if txt:
                rich_runs.append({"type": "text", "text": txt})
    flush_rich_runs()


def parse_body_sections(doc: Any, text_start: int, image_registry: Any, notes: Dict[str, Dict[str, str]] | None = None) -> BodyDispatchResult:
    """Parse DOCX body children after front matter into content sections."""
    current_section = make_body_section()
    sections = [current_section]
    ref_collector = ReferenceCollector(
        clean_text_func=clean_text_artifacts,
        is_backmatter_heading_func=is_backmatter_heading,
        normalize_heading_spacing_func=normalize_heading_spacing,
        classify_section_role_func=classify_section_role,
        table_rows_look_like_code_func=table_rows_look_like_code,
        code_text_from_table_rows_func=code_text_from_table_rows,
        clean_code_func=clean_code_text,
    )

    body_children = list(doc.element.body)
    p_idx = 0
    source_toc_skip_remaining = 0
    source_toc_skipped = 0
    placeholders_removed = 0
    content_control_paragraphs_in_place = 0

    def local_name(elem: Any) -> str:
        return elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

    def child_by_local_name(elem: Any, name: str) -> Any | None:
        for child in list(elem):
            if local_name(child) == name:
                return child
        return None

    def attr_value(elem: Any | None, name: str) -> str | None:
        if elem is None:
            return None
        value = elem.get(name)
        if value is not None:
            return value
        for key, val in elem.attrib.items():
            if str(key).split("}")[-1] == name:
                return val
        return None

    def int_attr(elem: Any | None, name: str) -> int | None:
        value = attr_value(elem, name)
        try:
            parsed = int(value or 0)
        except Exception:
            return None
        return parsed if parsed > 0 else None

    def paragraph_section_properties(elem: Any) -> Any | None:
        if local_name(elem) != "p":
            return None
        p_pr = child_by_local_name(elem, "pPr")
        if p_pr is None:
            return None
        return child_by_local_name(p_pr, "sectPr")

    def section_page_setup(sect_pr: Any | None) -> Dict[str, Any]:
        if sect_pr is None:
            return {}
        pg_sz = child_by_local_name(sect_pr, "pgSz")
        pg_mar = child_by_local_name(sect_pr, "pgMar")
        width = int_attr(pg_sz, "w")
        height = int_attr(pg_sz, "h")
        orientation = str(attr_value(pg_sz, "orient") or "").strip().lower()
        if orientation not in ("landscape", "portrait"):
            orientation = ""
        if not orientation and width and height:
            orientation = "landscape" if width > height else "portrait"
        margins: Dict[str, int] = {}
        for side in ("top", "bottom", "left", "right", "header", "footer", "gutter"):
            value = int_attr(pg_mar, side)
            if value:
                margins[side] = value
        setup: Dict[str, Any] = {}
        if orientation:
            setup["orientation"] = orientation
        if width:
            setup["page_width_twips"] = width
        if height:
            setup["page_height_twips"] = height
        if margins:
            setup["margins_twips"] = margins
        return setup

    section_page_by_elem: Dict[int, Dict[str, Any]] = {}

    def assign_section_page_setup(elem: Any, setup: Dict[str, Any]) -> None:
        tag = local_name(elem)
        if tag in ("p", "tbl"):
            section_page_by_elem[id(elem)] = setup
        elif tag in _BODY_TRANSPARENT_CONTAINERS or tag in _BODY_ACCEPTED_REVISION_CONTAINERS or tag == "sdt":
            for child in list(elem):
                assign_section_page_setup(child, setup)
        else:
            for child in list(elem):
                if local_name(child) in ("p", "tbl", "sdt") or local_name(child) in _BODY_TRANSPARENT_CONTAINERS | _BODY_ACCEPTED_REVISION_CONTAINERS:
                    assign_section_page_setup(child, setup)

    def build_section_page_map() -> None:
        pending: List[Any] = []

        def flush(sect_pr: Any | None) -> None:
            nonlocal pending
            setup = section_page_setup(sect_pr)
            if setup:
                for elem in pending:
                    assign_section_page_setup(elem, setup)
            pending = []

        for child in body_children:
            if local_name(child) == "sectPr":
                flush(child)
                continue
            pending.append(child)
            sect_pr = paragraph_section_properties(child)
            if sect_pr is not None:
                flush(sect_pr)
        if pending:
            fallback_sect_pr = None
            try:
                fallback_sect_pr = doc.sections[-1]._sectPr
            except Exception:
                fallback_sect_pr = None
            flush(fallback_sect_pr)

    build_section_page_map()

    def sdt_content_children(elem: Any) -> List[Any]:
        for part in list(elem):
            if local_name(part) == "sdtContent":
                return list(part)
        return []

    def countable_content_control_paragraph(paragraph: Any) -> bool:
        text = clean_text_artifacts(str(paragraph_visible_text(paragraph) or ""), preserve_newlines=True).strip()
        return bool(text and not is_unfilled_placeholder_text(text) and not is_template_instruction_text(text))

    def handle_paragraph(paragraph: Any, source_paragraph_index: int | None = None) -> None:
        nonlocal current_section, source_toc_skip_remaining, source_toc_skipped, placeholders_removed
        text = paragraph_visible_text(paragraph).strip()
        level = detect_heading_level(paragraph, text_override=text)

        if is_unfilled_placeholder_text(text) or is_template_instruction_text(text):
            placeholders_removed += 1
            return

        if source_toc_skip_remaining > 0:
            source_toc_skip_remaining -= 1
            source_toc_skipped += 1
            return

        if is_source_toc_title(text) and source_paragraph_index is not None:
            skip_after_title = source_toc_skip_count_after_title(doc.paragraphs, source_paragraph_index)
            if skip_after_title:
                source_toc_skip_remaining = skip_after_title
                source_toc_skipped += 1
                return

        if ref_collector.start_if_heading(text):
            return

        backmatter_section = ref_collector.exit_to_backmatter_section(text, level)
        if backmatter_section is not None:
            current_section = backmatter_section
            sections.append(current_section)
            return

        if ref_collector.consume_text(text):
            return
        if level > 0:
            section = _section_from_heading(text, level)
            if section:
                current_section = section
                sections.append(current_section)
        else:
            _append_paragraph_stream(current_section, paragraph, image_registry, notes=notes)

    def handle_table(tbl_elem: Any) -> None:
        table_data = extract_table_from_ooxml(
            tbl_elem,
            clean_text_func=clean_text_artifacts,
            image_rels=doc.part.rels,
            image_registry=image_registry,
            image_items_func=image_items_from_ooxml,
            image_run_items_func=images_from_run_ooxml,
            notes=notes,
        )
        rows = table_data.get("table_rows") or []
        if rows:
            if ref_collector.consume_table_rows(rows):
                pass
            elif table_rows_look_like_code(rows) and not table_data.get("table_cell_items"):
                current_section["paragraphs"].append(
                    {
                        "role": "code",
                        "code": code_text_from_table_rows(rows, clean_code_func=clean_code_text),
                        "table_rows": rows,
                    }
                )
            else:
                table_item = {"role": "table", "table_rows": rows}
                source_section_setup = section_page_by_elem.get(id(tbl_elem)) or {}
                if source_section_setup.get("orientation") == "landscape":
                    table_item["source_section_page_setup"] = source_section_setup
                if table_data.get("table_merges"):
                    table_item["table_merges"] = table_data.get("table_merges")
                if table_data.get("table_col_widths_twips"):
                    table_item["table_col_widths_twips"] = table_data.get("table_col_widths_twips")
                for layout_key in (
                    "table_row_grid_before",
                    "table_row_heights_twips",
                    "table_repeat_header_rows",
                    "table_cell_margins_twips",
                    "table_borders",
                    "table_cell_overrides",
                    "table_cell_items",
                ):
                    if layout_key == "table_repeat_header_rows":
                        if layout_key in table_data:
                            table_item[layout_key] = table_data.get(layout_key)
                    elif table_data.get(layout_key):
                        table_item[layout_key] = table_data.get(layout_key)
                current_section["paragraphs"].append(table_item)

    def dispatch_body_child(child: Any, wrapped: bool = False, in_content_control: bool = False) -> None:
        nonlocal p_idx, content_control_paragraphs_in_place
        tag = local_name(child)

        if tag == "p":
            if p_idx < text_start:
                if not wrapped:
                    p_idx += 1
                return
            if wrapped:
                paragraph = Paragraph(child, doc)
                if in_content_control and countable_content_control_paragraph(paragraph):
                    content_control_paragraphs_in_place += 1
                handle_paragraph(paragraph, source_paragraph_index=None)
            else:
                source_index = p_idx
                paragraph = doc.paragraphs[p_idx] if p_idx < len(doc.paragraphs) else Paragraph(child, doc)
                p_idx += 1
                handle_paragraph(paragraph, source_paragraph_index=source_index)

        elif tag == "tbl":
            if p_idx < text_start:
                return
            handle_table(child)

        elif tag in _BODY_TRANSPARENT_CONTAINERS:
            if p_idx < text_start:
                return
            for nested in list(child):
                dispatch_body_child(nested, wrapped=True, in_content_control=in_content_control)

        elif tag in _BODY_ACCEPTED_REVISION_CONTAINERS:
            if p_idx < text_start:
                return
            for nested in list(child):
                dispatch_body_child(nested, wrapped=True, in_content_control=in_content_control)

        elif tag in _BODY_DELETED_REVISION_CONTAINERS:
            return

        elif tag == "sdt":
            if p_idx < text_start:
                return
            for nested in sdt_content_children(child):
                dispatch_body_child(nested, wrapped=True, in_content_control=True)

    for child in body_children:
        dispatch_body_child(child, wrapped=False)

    return BodyDispatchResult(
        sections=sections,
        references=ref_collector.finish(),
        source_toc_skipped=source_toc_skipped,
        placeholders_removed=placeholders_removed,
        content_control_paragraphs_in_place=content_control_paragraphs_in_place,
    )
