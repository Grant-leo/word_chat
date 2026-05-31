"""Controlled user-level repair loop for generated build scripts.

The loop is intentionally conservative: it may rebuild the current output and
may patch only ``Outputs/<run>/build_generated.py`` with known-safe repairs.
Reusable engine changes remain developer work outside this module.
"""
from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
import json
import os
from pathlib import Path
import re
from typing import Any, Callable, Dict, List

try:
    from privacy import sanitize_value
except Exception:  # pragma: no cover - best-effort report hardening
    def sanitize_value(value, project_root=None):
        return value

try:
    from .qa import _write_structural_dependency_handoff
except Exception:  # pragma: no cover - repair loop can still write a minimal fallback
    _write_structural_dependency_handoff = None


NEEDS_USER_AUTO_LEVELS = {
    "needs_user_file",
    "needs_user_input",
    "needs_user_confirmation",
    "optional_user_input",
}
ENVIRONMENT_AUTO_LEVELS = {"needs_environment"}
NEEDS_USER_CODES = {
    "CONTENT_EMPTY",
    "CONTENT_IMAGE_MISSING",
    "IMAGE_EXTRACT_FAILED",
    "NON_BODY_IMAGE_UNSUPPORTED",
    "PDF_TEMPLATE_UNSUPPORTED",
    "PDF_TEMPLATE_DEPENDENCY_MISSING",
    "MISSING_CONTENT_JSON",
    "MISSING_FORMAT_JSON",
}
REBUILD_ONLY_CODES = {"MISSING_DOCX", "DOCX_XML_UNREADABLE"}
PLACEHOLDER_MARKER = "# AUTO_REPAIR_PLACEHOLDER_CLEANUP_V1"
REFERENCE_EAST_ASIA_MARKER = "# AUTO_REPAIR_REFERENCE_EAST_ASIA_FONT_V1"
REPORT_BLOCKER_GUIDES = {
    "STRUCTURAL_QA_UNAVAILABLE": {
        "auto_level": "needs_user_input",
        "user_action": "修复 qa_checker.py / qa_checker_modules 的导入或依赖后，重新运行完整流水线；先查看 qa_report.md 和 qa_repair_plan.md。",
    },
    "STRUCTURAL_QA_FAILED": {
        "auto_level": "needs_user_input",
        "user_action": "修复 qa_checker.py / qa_checker_modules 的运行异常后，重新运行完整流水线；先查看 qa_report.md 和 qa_repair_plan.md。",
    },
    "CONFORMANCE_QA_UNAVAILABLE": {
        "auto_level": "needs_user_input",
        "user_action": "修复 strict conformance QA 依赖后重跑；先查看 conformance_report.md。",
    },
    "CONFORMANCE_QA_FAILED": {
        "auto_level": "needs_user_input",
        "user_action": "修复 qa_conformance.py / qa_conformance_modules 的运行异常后重跑 strict QA；先查看 conformance_report.md。",
    },
    "VISUAL_QA_UNAVAILABLE": {
        "auto_level": "needs_user_input",
        "user_action": "修复 visual QA 依赖后重跑；先查看 visual_report.md。",
    },
    "VISUAL_QA_FAILED": {
        "auto_level": "needs_user_input",
        "user_action": "修复 qa_visual.py / qa_visual_modules、Word COM 或 Poppler 渲染异常后重跑 visual QA；先查看 visual_report.md。",
    },
    "PDF_EXPORT_FAILED": {
        "auto_level": "needs_user_input",
        "user_action": "安装或修复 Microsoft Word COM/PDF 导出环境后重跑 visual QA。",
    },
    "PDFINFO_UNAVAILABLE": {
        "auto_level": "needs_user_input",
        "user_action": "安装 Poppler 的 pdfinfo 后重跑 visual QA。",
    },
    "PDFINFO_FAILED": {
        "auto_level": "needs_user_input",
        "user_action": "检查 Poppler/pdfinfo 是否可运行，或打开 visual_report.md 查看 PDF 元信息读取错误。",
    },
    "PDF_PAGE_COUNT_INVALID": {
        "auto_level": "needs_user_confirmation",
        "user_action": "打开导出的 PDF/DOCX 确认是否为空白；若导出异常，先修复 Word/PDF 导出环境后重跑。",
    },
    "PDFTOTEXT_UNAVAILABLE": {
        "auto_level": "needs_user_input",
        "user_action": "安装 Poppler 的 pdftotext 后重跑 visual QA。",
    },
    "PDFTOTEXT_FAILED": {
        "auto_level": "needs_user_input",
        "user_action": "检查 Poppler/pdftotext 是否可运行，或打开 visual_report.md 查看文本提取错误。",
    },
    "SAMPLE_RENDER_FAILED": {
        "auto_level": "needs_user_input",
        "user_action": "安装 Poppler 的 pdftoppm 后重跑 visual QA。",
    },
    "ALL_PAGE_RENDER_FAILED": {
        "auto_level": "needs_user_input",
        "user_action": "安装或修复 Poppler 的 pdftoppm 后重跑 visual QA。",
    },
    "PAGE_IMAGE_UNREADABLE": {
        "auto_level": "needs_user_input",
        "user_action": "检查 visual_qa/all_pages 下的 PNG 是否损坏，并修复 PDF 渲染/图像读取环境后重跑。",
    },
    "WPS_EXPORT_UNAVAILABLE": {
        "auto_level": "needs_user_input",
        "user_action": "安装/配置 WPS COM，或取消 --require-wps 后重跑 visual QA。",
    },
    "WPS_PDFINFO_UNAVAILABLE": {
        "auto_level": "needs_user_input",
        "user_action": "WPS 已导出 PDF，但 pdfinfo 不可用；先确认 Poppler 可用，并检查 WPS 导出的 PDF 能正常打开，修复后重跑 visual QA。",
    },
    "WPS_PDFINFO_FAILED": {
        "auto_level": "needs_user_input",
        "user_action": "WPS 已导出 PDF，但页面信息读取失败；先确认 WPS 导出的 PDF 能正常打开，再修复 PDF/Poppler 环境并重跑 visual QA。",
    },
    "WPS_PAGE_COUNT_INVALID": {
        "auto_level": "needs_user_confirmation",
        "user_action": "WPS 导出的 PDF 没有有效页面；先用 WPS 打开最终 DOCX 和导出的 PDF 检查是否为空白，修复后重跑 visual QA。",
    },
    "WPS_PAGE_COUNT_MISMATCH": {
        "auto_level": "needs_user_confirmation",
        "user_action": "分别打开 Word 与 WPS 导出的 PDF 比对分页差异；确认是兼容性差异还是排版脚本问题。修复后重跑 visual QA。",
    },
    "WPS_PAGE_SIZE_MISMATCH": {
        "auto_level": "needs_user_confirmation",
        "user_action": "分别打开 Word 与 WPS 导出的 PDF 比对纸张大小、页面尺寸和横竖方向；确认是模板页面设置、WPS 兼容性差异还是排版脚本问题。修复后重跑 visual QA。",
    },
    "WPS_TEXT_PAGE_MISMATCH": {
        "auto_level": "needs_user_confirmation",
        "user_action": "分别打开 Word 与 WPS 导出的 PDF 比对正文、目录、公式和图片内容；确认 WPS 文本页缺失是导出、字体兼容还是排版生成问题。修复后重跑 visual QA。",
    },
    "WPS_SAMPLE_RENDER_FAILED": {
        "auto_level": "needs_user_confirmation",
        "user_action": "打开 WPS 导出的 PDF 和 visual_report.md，检查 WPS 样张 PNG 是否能完整生成；修复 WPS/Poppler 渲染问题后重跑 visual QA。",
    },
    "WPS_SAMPLE_IMAGE_MISMATCH": {
        "auto_level": "needs_user_confirmation",
        "user_action": "分别打开 Word 样张 visual_qa/samples/ 和 WPS 样张 visual_qa/wps/samples/，比对公式、图片、表格和正文画面差异；确认是 WPS 兼容还是排版生成问题。修复后重跑 visual QA。",
    },
    "GOLDEN_BASELINE_MISMATCH": {
        "auto_level": "needs_user_confirmation",
        "user_action": "打开 visual_report.md 和 visual_qa/samples/ 对比页面；确认变化正确后用 --update-golden 更新基线，或继续修复排版。",
    },
}


@dataclass(frozen=True)
class RepairLoopResult:
    ok: bool
    status: str
    report_path: str
    rounds: int
    final_errors: int


def run_repair_loop(
    out_dir: str,
    *,
    mode: str,
    output_docx_name: str,
    qa_level: str,
    project_root: str,
    max_rounds: int,
    stop_no_improve: int,
    deps: Any,
    run_generated_script: Callable[..., Any],
    python_executable: str,
    golden_dir: str | None = None,
    update_golden: bool = False,
    require_wps: bool = False,
) -> RepairLoopResult:
    """Run a bounded build-script repair loop and write audit reports."""

    max_rounds = max(1, int(max_rounds or 1))
    stop_no_improve = max(1, int(stop_no_improve or 1))
    out_path = Path(out_dir).resolve()
    history: List[Dict[str, Any]] = []

    state = _ensure_qa_state(
        out_path,
        mode=mode,
        output_docx_name=output_docx_name,
        qa_level=qa_level,
        project_root=project_root,
        deps=deps,
        golden_dir=golden_dir,
        update_golden=update_golden,
        require_wps=require_wps,
    )
    history.append(_round_snapshot(0, "initial", state, actions=[], previous_error_codes=[]))

    if state["total_errors"] == 0:
        return _finish(
            out_path,
            ok=True,
            status="converged",
            mode=mode,
            output_docx_name=output_docx_name,
            qa_level=qa_level,
            max_rounds=max_rounds,
            stop_no_improve=stop_no_improve,
            history=history,
        )

    best_errors = state["total_errors"]
    stagnant_rounds = 0

    for round_no in range(1, max_rounds + 1):
        blockers = _blocking_steps(state)
        if blockers:
            history.append(_round_snapshot(round_no, "blocked", state, actions=[], previous_error_codes=history[-1].get("error_codes") or []))
            status, stop_detail = _blocker_stop_reason(blockers)
            return _finish(
                out_path,
                ok=False,
                status=status,
                mode=mode,
                output_docx_name=output_docx_name,
                qa_level=qa_level,
                max_rounds=max_rounds,
                stop_no_improve=stop_no_improve,
                history=history,
                stop_detail=stop_detail,
                blockers=blockers,
            )

        actions = _apply_safe_repairs(out_path, state)
        if not actions:
            history.append(_round_snapshot(round_no, "no_supported_auto_repair", state, actions=[], previous_error_codes=history[-1].get("error_codes") or []))
            return _finish(
                out_path,
                ok=False,
                status="stopped_no_supported_auto_repair",
                mode=mode,
                output_docx_name=output_docx_name,
                qa_level=qa_level,
                max_rounds=max_rounds,
                stop_no_improve=stop_no_improve,
                history=history,
                stop_detail="No known-safe build_generated.py repair matched the current QA errors.",
            )

        build_result = _rebuild_current_docx(
            out_path,
            output_docx_name=output_docx_name,
            run_generated_script=run_generated_script,
            python_executable=python_executable,
        )
        if build_result["returncode"] != 0:
            history.append(_round_snapshot(round_no, "build_failed", state, actions=actions, build=build_result, previous_error_codes=history[-1].get("error_codes") or []))
            return _finish(
                out_path,
                ok=False,
                status="stopped_build_failed",
                mode=mode,
                output_docx_name=output_docx_name,
                qa_level=qa_level,
                max_rounds=max_rounds,
                stop_no_improve=stop_no_improve,
                history=history,
                stop_detail=build_result.get("stderr") or build_result.get("stdout") or "build_generated.py failed",
            )

        state = _run_qa_state(
            out_path,
            mode=mode,
            output_docx_name=output_docx_name,
            qa_level=qa_level,
            project_root=project_root,
            deps=deps,
            golden_dir=golden_dir,
            update_golden=update_golden,
            require_wps=require_wps,
        )
        history.append(_round_snapshot(round_no, "after_repair", state, actions=actions, build=build_result, previous_error_codes=history[-1].get("error_codes") or []))

        if state["total_errors"] == 0:
            return _finish(
                out_path,
                ok=True,
                status="converged",
                mode=mode,
                output_docx_name=output_docx_name,
                qa_level=qa_level,
                max_rounds=max_rounds,
                stop_no_improve=stop_no_improve,
                history=history,
            )

        if state["total_errors"] < best_errors:
            best_errors = state["total_errors"]
            stagnant_rounds = 0
        else:
            stagnant_rounds += 1
        if stagnant_rounds >= stop_no_improve:
            return _finish(
                out_path,
                ok=False,
                status="stopped_no_improvement",
                mode=mode,
                output_docx_name=output_docx_name,
                qa_level=qa_level,
                max_rounds=max_rounds,
                stop_no_improve=stop_no_improve,
                history=history,
                stop_detail=f"QA error count did not improve for {stagnant_rounds} consecutive repair rounds.",
            )

    return _finish(
        out_path,
        ok=False,
        status="stopped_max_rounds",
        mode=mode,
        output_docx_name=output_docx_name,
        qa_level=qa_level,
        max_rounds=max_rounds,
        stop_no_improve=stop_no_improve,
        history=history,
        stop_detail=f"Reached max repair rounds: {max_rounds}.",
    )


def _ensure_qa_state(out_path: Path, **kwargs: Any) -> Dict[str, Any]:
    if not (out_path / "qa_report.json").exists() or _required_report_missing(out_path, kwargs["qa_level"]):
        return _run_qa_state(out_path, **kwargs)
    return _read_state(out_path, kwargs["qa_level"])


def _required_report_missing(out_path: Path, qa_level: str) -> bool:
    if qa_level in ("strict", "visual") and not (out_path / "conformance_report.json").exists():
        return True
    if qa_level == "visual" and not (out_path / "visual_report.json").exists():
        return True
    return False


def _run_qa_state(
    out_path: Path,
    *,
    mode: str,
    output_docx_name: str,
    qa_level: str,
    project_root: str,
    deps: Any,
    golden_dir: str | None = None,
    update_golden: bool = False,
    require_wps: bool = False,
) -> Dict[str, Any]:
    if deps.qa_check_and_write is None:
        _write_structural_failure_report(
            out_path,
            mode=mode,
            code="STRUCTURAL_QA_UNAVAILABLE",
            detail=_optional_detail(deps, "qa_checker"),
            project_root=project_root,
        )
        return _read_state(out_path, qa_level)
    try:
        deps.qa_check_and_write(str(out_path), mode=mode, output_docx_name=output_docx_name)
    except Exception as exc:
        _write_structural_failure_report(
            out_path,
            mode=mode,
            code="STRUCTURAL_QA_FAILED",
            detail=_exception_detail(exc, project_root),
            project_root=project_root,
        )
        return _read_state(out_path, qa_level)
    if qa_level in ("strict", "visual") and deps.conformance_check_and_write is None:
        _write_dependency_report(
            out_path,
            report_name="conformance_report",
            mode=mode,
            code="CONFORMANCE_QA_UNAVAILABLE",
            message="strict conformance QA is required but qa_conformance.py is unavailable.",
            detail=_optional_detail(deps, "qa_conformance"),
            project_root=project_root,
        )
        return _read_state(out_path, qa_level)
    if qa_level in ("strict", "visual"):
        try:
            deps.conformance_check_and_write(
                str(out_path),
                mode=mode,
                output_docx_name=output_docx_name,
                project_root=project_root,
            )
        except Exception as exc:
            _write_dependency_report(
                out_path,
                report_name="conformance_report",
                mode=mode,
                code="CONFORMANCE_QA_FAILED",
                message="strict conformance QA crashed before it could finish.",
                detail=_exception_detail(exc, project_root),
                project_root=project_root,
            )
            return _read_state(out_path, qa_level)
    if qa_level == "visual" and deps.visual_check_and_write is None:
        _write_dependency_report(
            out_path,
            report_name="visual_report",
            mode=mode,
            code="VISUAL_QA_UNAVAILABLE",
            message="visual QA is required but qa_visual.py is unavailable.",
            detail=_optional_detail(deps, "qa_visual"),
            project_root=project_root,
        )
        return _read_state(out_path, qa_level)
    if qa_level == "visual":
        try:
            deps.visual_check_and_write(
                str(out_path),
                output_docx_name=output_docx_name,
                project_root=project_root,
                render_all_pages=True,
                require_wps=bool(require_wps),
                golden_dir=os.path.abspath(golden_dir) if golden_dir else None,
                update_golden=bool(update_golden),
            )
        except Exception as exc:
            _write_dependency_report(
                out_path,
                report_name="visual_report",
                mode=mode,
                code="VISUAL_QA_FAILED",
                message="visual QA crashed before it could finish.",
                detail=_exception_detail(exc, project_root),
                project_root=project_root,
            )
            return _read_state(out_path, qa_level)
    return _read_state(out_path, qa_level)


def _exception_detail(exc: Exception, project_root: str | None = None) -> str:
    return str(_safe_report_value(f"{exc.__class__.__name__}: {exc}", project_root) or "")


def _optional_detail(deps: Any, name: str) -> str:
    detail = getattr(deps, "optional_import_detail", None)
    if detail is None:
        return ""
    try:
        return str(detail(name) or "")
    except Exception:
        return ""


def _write_structural_failure_report(
    out_path: Path,
    *,
    mode: str,
    code: str,
    detail: str,
    project_root: str | None = None,
) -> None:
    if _write_structural_dependency_handoff is not None:
        if code == "STRUCTURAL_QA_FAILED":
            _write_structural_dependency_handoff(
                str(out_path),
                mode=mode,
                detail=detail,
                project_root=project_root,
                code=code,
                message="Structural QA crashed before it could finish, so the pipeline cannot prove this DOCX is safe to deliver.",
                title="结构 QA 执行失败",
                why="qa_checker.py 运行中抛出异常；不能把本轮输出标记为已通过。",
                user_action="让 Agent 先查看 qa_report.md 和 qa_repair_plan.md，修复 qa_checker.py / qa_checker_modules 后重跑完整流水线；不要把这次输出当作已通过。",
                developer_action="检查 qa_checker.py / qa_checker_modules 的异常堆栈、输入报告和依赖状态，修复后重跑 targeted regression、完整 regression 与真实流水线。",
                next_action="结构 QA 运行中断。先修复 qa_checker.py / qa_checker_modules 的异常，再重新运行完整流水线；先查看 qa_report.md 和 qa_repair_plan.md。",
            )
            return
        _write_structural_dependency_handoff(
            str(out_path),
            mode=mode,
            detail=detail,
            project_root=project_root,
        )
        return

    _write_dependency_report(
        out_path,
        report_name="qa_report",
        mode=mode,
        code=code,
        message="Structural QA could not run, so auto repair stopped before declaring the output safe.",
        detail=detail,
        project_root=project_root,
    )
    resume_command = _workflow_resume_command_fallback(out_path, mode)
    issue = {
        "code": code,
        "severity": "error",
        "title": "结构 QA 不可用",
        "why": "结构 QA 没有成功运行，无法证明当前输出可交付。",
        "detail": _safe_report_value(detail, project_root),
        "auto_level": "dependency_repair",
        "target": "Paper_Project/Program/pipeline/qa_checker.py / qa_checker_modules",
        "user_action": "让 Agent 修复结构 QA 后重跑完整流水线。",
        "developer_action": "修复 qa_checker.py / qa_checker_modules 后重跑回归测试和真实流水线。",
    }
    plan = {
        "schema_version": 1,
        "passed": False,
        "summary": "结构 QA 没有运行成功，auto-repair 已停止。",
        "mode": mode,
        "blocking_errors": 1,
        "warnings": 0,
        "next_action": f"修复 qa_checker.py / qa_checker_modules 后重新运行完整流水线。{(' 修复后运行：`' + resume_command + '`') if resume_command else ''}",
        "resume_scope": "full_pipeline",
        "resume_command": resume_command,
        "steps": [issue],
    }
    (out_path / "qa_repair_plan.json").write_text(json.dumps(plan, ensure_ascii=False, indent=2), encoding="utf-8")
    (out_path / "qa_repair_plan.md").write_text(
        "# QA 修复向导\n\n"
        "- 结果：需要修复\n"
        f"- 优先动作：{plan['next_action']}\n"
        "- 先打开 `qa_report.md`。\n",
        encoding="utf-8",
    )
    (out_path / "qa_fix_prompt.txt").write_text(str(plan["next_action"]) + "\n", encoding="utf-8")


def _write_dependency_report(
    out_path: Path,
    *,
    report_name: str,
    mode: str,
    code: str,
    message: str,
    detail: str,
    project_root: str | None = None,
) -> None:
    next_action = (REPORT_BLOCKER_GUIDES.get(code) or {}).get("user_action") or "安装或修复所需 QA 依赖后，重新运行完整流水线。"
    message = _safe_report_value(message, project_root)
    detail = _safe_report_value(detail, project_root)
    next_action = _safe_report_value(next_action, project_root)
    issue = {
        "code": code,
        "severity": "error",
        "message": message,
        "detail": detail,
    }
    report = {
        "schema_version": 1,
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "mode": mode,
        "output_dir_name": out_path.name,
        "passed": False,
        "counts": {},
        "issues": [issue],
        "next_action": next_action,
    }
    (out_path / f"{report_name}.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        f"# {report_name.replace('_', ' ')}",
        "",
        "- 结果：未通过",
        f"- 问题码：`{code}`",
        f"- 信息：{message}",
        f"- 下一步：{next_action}",
    ]
    if detail:
        lines.append(f"- 细节：`{detail}`")
    (out_path / f"{report_name}.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def _safe_report_value(value: Any, project_root: str | None = None) -> Any:
    try:
        return sanitize_value(value, project_root)
    except Exception:
        return value


def _read_state(out_path: Path, qa_level: str) -> Dict[str, Any]:
    qa_report = _load_json(out_path / "qa_report.json", {})
    repair_plan = _load_json(out_path / "qa_repair_plan.json", qa_report.get("repair_plan") or {})
    reports: Dict[str, Dict[str, Any]] = {"qa": qa_report}
    if qa_level in ("strict", "visual"):
        reports["conformance"] = _load_json(out_path / "conformance_report.json", {})
    if qa_level == "visual":
        reports["visual"] = _load_json(out_path / "visual_report.json", {})
    issues = []
    for label, report in reports.items():
        for item in report.get("issues") or []:
            issue = dict(item)
            issue["report"] = label
            issues.append(issue)
    errors = [item for item in issues if item.get("severity") == "error"]
    warnings = [item for item in issues if item.get("severity") == "warning"]
    passed = bool(qa_report.get("passed")) and all(
        bool(report.get("passed", True)) for label, report in reports.items() if label != "qa" and report
    )
    report_mode = str(qa_report.get("mode") or next((r.get("mode") for r in reports.values() if r.get("mode")), "user"))
    workflow_resume_command = _workflow_resume_command_fallback(out_path, report_mode)
    return {
        "passed": passed,
        "reports": reports,
        "qa_report": qa_report,
        "repair_plan": repair_plan,
        "issues": issues,
        "errors": errors,
        "warnings": warnings,
        "error_codes": [str(item.get("code") or "") for item in errors],
        "warning_codes": [str(item.get("code") or "") for item in warnings],
        "qa_errors": sum(1 for item in qa_report.get("issues") or [] if item.get("severity") == "error"),
        "qa_warnings": sum(1 for item in qa_report.get("issues") or [] if item.get("severity") == "warning"),
        "total_errors": len(errors),
        "total_warnings": len(warnings),
        "workflow_resume_command": workflow_resume_command,
        "workflow_resume_scope": "full_pipeline" if workflow_resume_command else "",
    }


def _blocking_steps(state: Dict[str, Any]) -> List[Dict[str, Any]]:
    steps = state.get("repair_plan", {}).get("steps") or []
    error_codes = set(state.get("error_codes") or [])
    blockers = []
    for step in steps:
        code = str(step.get("code") or "")
        auto_level = str(step.get("auto_level") or "")
        severity = str(step.get("severity") or "")
        if code in error_codes and severity == "error" and (auto_level in NEEDS_USER_AUTO_LEVELS or auto_level in ENVIRONMENT_AUTO_LEVELS or code in NEEDS_USER_CODES):
            blockers.append(
                {
                    "code": code,
                    "auto_level": auto_level,
                    "user_action": step.get("user_action") or "",
                    "detail": step.get("detail") or "",
                }
            )
    blocked_codes = {item.get("code") for item in blockers}
    for issue in state.get("errors") or []:
        code = str(issue.get("code") or "")
        if not code or code in blocked_codes:
            continue
        guide = REPORT_BLOCKER_GUIDES.get(code)
        if not guide:
            continue
        blockers.append(
            {
                "code": code,
                "auto_level": guide.get("auto_level") or "needs_user_input",
                "user_action": guide.get("user_action") or "查看对应 QA 报告后重跑。",
                "detail": issue.get("detail") or issue.get("message") or "",
                "report": issue.get("report") or "",
            }
        )
    return blockers


def _blocker_stop_reason(blockers: List[Dict[str, Any]]) -> tuple[str, str]:
    levels = {str(item.get("auto_level") or "") for item in blockers}
    codes = {str(item.get("code") or "") for item in blockers}
    if levels and levels <= ENVIRONMENT_AUTO_LEVELS:
        return "stopped_dependency_or_environment", "Dependency or environment repair is required before automatic repair can continue."
    if (levels and levels <= {"needs_user_file"}) or (codes and codes <= NEEDS_USER_CODES):
        return "stopped_needs_user_file", "User file/input is required before automatic repair can continue."
    return "stopped_needs_user_input", "User confirmation, dependency repair, or manual inspection is required before automatic repair can continue."


def _apply_safe_repairs(out_path: Path, state: Dict[str, Any]) -> List[Dict[str, Any]]:
    codes = set(state.get("error_codes") or [])
    actions: List[Dict[str, Any]] = []
    if "PLACEHOLDER_TEXT_LEFT" in codes:
        action = _patch_placeholder_cleanup(out_path)
        if action:
            actions.append(action)
    if "STYLE_MISMATCH" in codes:
        action = _patch_reference_east_asia_fonts(out_path, state)
        if action:
            actions.append(action)
    if codes & REBUILD_ONLY_CODES:
        actions.append({"kind": "rebuild", "codes": sorted(codes & REBUILD_ONLY_CODES)})
    return actions


def _patch_placeholder_cleanup(out_path: Path) -> Dict[str, Any] | None:
    build_path = _safe_build_path(out_path)
    text = build_path.read_text(encoding="utf-8")
    if PLACEHOLDER_MARKER in text:
        return {"kind": "script_patch", "code": "PLACEHOLDER_TEXT_LEFT", "status": "already_applied"}
    needle = "if __name__ == '__main__':\n    main()\n"
    if needle not in text:
        return None
    patched = text.replace(needle, _placeholder_cleanup_fragment() + "\n" + needle.replace("main()", "main()\n    _auto_repair_after_main()"))
    build_path.write_text(patched, encoding="utf-8")
    return {
        "kind": "script_patch",
        "code": "PLACEHOLDER_TEXT_LEFT",
        "target": "build_generated.py",
        "summary": "Injected a post-build cleanup that removes obvious placeholder paragraphs from the final DOCX.",
    }


def _placeholder_cleanup_fragment() -> str:
    return r'''
# AUTO_REPAIR_PLACEHOLDER_CLEANUP_V1
def _auto_repair_placeholder_hit(text):
    import re
    t = str(text or '').strip()
    if not t:
        return False
    return bool(re.search(r'(\[[^\]\n]*(?:XX|XXX|TODO|FIXME|待填|待填写|请输入|报名|序号|姓名|学号|学院|专业|班级|题目|指导|教师|日期|编码)[^\]\n]*\])|(\{\{[^}]+\}\}|TODO|FIXME|待填写|待补全|XXXX)', t, re.I))


def _auto_repair_delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _auto_repair_clean_placeholders(docx_path):
    import os
    from docx import Document as _AutoRepairDocument
    if not os.path.exists(docx_path):
        return 0
    docx = _AutoRepairDocument(docx_path)
    removed = 0
    for paragraph in list(docx.paragraphs):
        if _auto_repair_placeholder_hit(paragraph.text):
            _auto_repair_delete_paragraph(paragraph)
            removed += 1
    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in list(cell.paragraphs):
                    if _auto_repair_placeholder_hit(paragraph.text):
                        paragraph.clear()
                        removed += 1
    if removed:
        docx.save(docx_path)
    return removed


def _auto_repair_after_main():
    try:
        removed = _auto_repair_clean_placeholders(OUT)
        if removed:
            print(f'Auto repair: removed {removed} placeholder paragraph(s)')
    except Exception as exc:
        print(f'Auto repair warning: placeholder cleanup failed: {exc}')
'''


def _patch_reference_east_asia_fonts(out_path: Path, state: Dict[str, Any]) -> Dict[str, Any] | None:
    details = "\n".join(str(item.get("detail") or "") for item in state.get("errors") or [])
    if "reference" not in details or "eastAsia font" not in details:
        return None
    build_path = _safe_build_path(out_path)
    text = build_path.read_text(encoding="utf-8")
    if REFERENCE_EAST_ASIA_MARKER in text:
        return {"kind": "script_patch", "code": "STYLE_MISMATCH", "status": "already_applied"}
    old = """def add_reference_mixed_runs(p, text, prof):
    # Chinese parts use the role's CJK font; Latin/numeric punctuation uses Times New Roman.
    for seg in re.findall(r'[\\u4e00-\\u9fff]+|[^\\u4e00-\\u9fff]+', text):
        r = p.add_run(seg)
        if has_cjk(seg):
            apply_run_profile(r, prof, seg, force_latin='Times New Roman')
        else:
            p_latin = dict(prof); p_latin['font'] = 'Times New Roman'
            apply_run_profile(r, p_latin, seg, force_latin='Times New Roman')
"""
    new = """def add_reference_mixed_runs(p, text, prof):
    # AUTO_REPAIR_REFERENCE_EAST_ASIA_FONT_V1
    # Keep the template CJK font in w:eastAsia for numeric/Latin label runs.
    for seg in re.findall(r'[\\u4e00-\\u9fff]+|[^\\u4e00-\\u9fff]+', text):
        r = p.add_run(seg)
        apply_run_profile(r, prof, seg, force_latin='Times New Roman')
"""
    if old not in text:
        return None
    build_path.write_text(text.replace(old, new, 1), encoding="utf-8")
    return {
        "kind": "script_patch",
        "code": "STYLE_MISMATCH",
        "target": "build_generated.py",
        "summary": "Patched reference rendering so numeric/Latin runs preserve the role CJK eastAsia font.",
    }


def _safe_build_path(out_path: Path) -> Path:
    build_path = (out_path / "build_generated.py").resolve()
    if build_path.name != "build_generated.py" or build_path.parent != out_path.resolve():
        raise RuntimeError(f"Unsafe build script path: {build_path}")
    if not build_path.exists():
        raise RuntimeError(f"Missing build script: {build_path}")
    return build_path


def _rebuild_current_docx(
    out_path: Path,
    *,
    output_docx_name: str,
    run_generated_script: Callable[..., Any],
    python_executable: str,
) -> Dict[str, Any]:
    build_path = _safe_build_path(out_path)
    result = run_generated_script(str(build_path), str(out_path), python_executable=python_executable)
    return {
        "returncode": int(getattr(result, "returncode", 1)),
        "stdout": _sanitize_log(str(getattr(result, "stdout", "")), out_path)[-2000:],
        "stderr": _sanitize_log(str(getattr(result, "stderr", "")), out_path)[-2000:],
        "docx_exists": (out_path / output_docx_name).exists(),
    }


def _round_snapshot(
    round_no: int,
    phase: str,
    state: Dict[str, Any],
    *,
    actions: List[Dict[str, Any]],
    build: Dict[str, Any] | None = None,
    previous_error_codes: List[str] | None = None,
) -> Dict[str, Any]:
    previous_codes = set(previous_error_codes or [])
    current_codes = set(state.get("error_codes") or [])
    repair_plan = state.get("repair_plan") or {}
    return {
        "round": round_no,
        "phase": phase,
        "passed": state.get("passed"),
        "total_errors": state.get("total_errors", 0),
        "total_warnings": state.get("total_warnings", 0),
        "qa_errors": state.get("qa_errors", 0),
        "qa_warnings": state.get("qa_warnings", 0),
        "error_codes": list(state.get("error_codes") or []),
        "warning_codes": list(state.get("warning_codes") or []),
        "new_error_codes": sorted(current_codes - previous_codes),
        "resolved_error_codes": sorted(previous_codes - current_codes),
        "repair_next_action": repair_plan.get("next_action") or "",
        "repair_resume_scope": repair_plan.get("resume_scope") or "",
        "repair_resume_command": repair_plan.get("resume_command") or "",
        "workflow_resume_scope": state.get("workflow_resume_scope") or "",
        "workflow_resume_command": state.get("workflow_resume_command") or "",
        "actions": actions,
        "build": build or {},
    }


def _finish(
    out_path: Path,
    *,
    ok: bool,
    status: str,
    mode: str,
    output_docx_name: str,
    qa_level: str,
    max_rounds: int,
    stop_no_improve: int,
    history: List[Dict[str, Any]],
    stop_detail: str = "",
    blockers: List[Dict[str, Any]] | None = None,
) -> RepairLoopResult:
    final = history[-1] if history else {}
    display_out = _report_display_path(out_path)
    display_docx = _join_report_path(display_out, output_docx_name)
    final_warnings = int(final.get("total_warnings") or 0)
    handoff = _repair_loop_handoff(
        ok=ok,
        status=status,
        mode=mode,
        final=final,
        display_out=display_out,
        blockers=blockers or [],
        stop_detail=stop_detail,
    )
    manual_checks = ["用 Word/WPS 打开最终 DOCX，核对分页、图片、公式、表格和目录。"]
    if final_warnings:
        manual_checks.append("查看 qa_report.md 和 repair_loop_report.md 中的剩余 warning，确认不会影响交付。")
    report = {
        "schema_version": 1,
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "status": status,
        "ok": bool(ok),
        "mode": mode,
        "qa_level": qa_level,
        "output_dir": display_out,
        "final_docx": display_docx,
        "max_rounds": max_rounds,
        "stop_no_improve": stop_no_improve,
        "rounds_run": max(0, len([r for r in history if r.get("round", 0) > 0])),
        "final_errors": int(final.get("total_errors") or 0),
        "final_warnings": final_warnings,
        "final_error_codes": final.get("error_codes") or [],
        "final_warning_codes": final.get("warning_codes") or [],
        "next_action": handoff["next_action"],
        "resume_scope": handoff["resume_scope"],
        "resume_command": handoff["resume_command"],
        "warning_policy": (
            "剩余 warning 不会阻断自动 QA 收敛，但仍可能影响交付质量，交付前必须用 Word/WPS 人工确认。"
            if final_warnings
            else "当前启用的 QA 未报告剩余 warning。"
        ),
        "stop_detail": stop_detail,
        "blockers": blockers or [],
        "manual_check_required": manual_checks,
        "rounds": history,
    }
    report_path = out_path / "repair_loop_report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    (out_path / "repair_loop_report.md").write_text(_report_to_markdown(report), encoding="utf-8")
    return RepairLoopResult(
        ok=bool(ok),
        status=status,
        report_path=_join_report_path(display_out, "repair_loop_report.json"),
        rounds=int(report["rounds_run"]),
        final_errors=int(report["final_errors"]),
    )


def _repair_loop_handoff(
    *,
    ok: bool,
    status: str,
    mode: str,
    final: Dict[str, Any],
    display_out: str,
    blockers: List[Dict[str, Any]],
    stop_detail: str,
) -> Dict[str, str]:
    if ok:
        return {
            "next_action": "自动修复已收敛；请用 Word/WPS 打开最终 DOCX 做最终视觉核对。",
            "resume_scope": "final_review",
            "resume_command": "",
        }

    if status == "stopped_build_failed":
        command = _default_rebuild_command(display_out)
        return {
            "next_action": f"自动修复后重建失败；先打开 `repair_loop_report.md` 查看构建错误，再让 Agent 检查本次 `build_generated.py`，修复后运行：`{command}`。",
            "resume_scope": "current_docx",
            "resume_command": command,
        }

    if blockers:
        first = blockers[0]
        code = str(first.get("code") or "").strip()
        action = str(first.get("user_action") or "").strip() or str(first.get("detail") or "").strip()
        prefix = f"优先处理 `{code}`：{action}" if code else action
        if status == "stopped_needs_user_file":
            scope = "input_files"
        elif status == "stopped_dependency_or_environment":
            scope = "environment"
        else:
            scope = "manual_or_dependency"
        command = str(final.get("repair_resume_command") or final.get("workflow_resume_command") or "").strip()
        route = "补齐或更换输入文件后重新运行完整流水线。" if scope == "input_files" else "完成依赖修复、人工确认或环境修复后重新运行对应 QA。"
        if command:
            route += f" 可运行：`{command}`"
        return {
            "next_action": " ".join(part for part in (prefix, route) if part).strip(),
            "resume_scope": scope,
            "resume_command": command,
        }

    command = str(final.get("repair_resume_command") or "").strip()
    scope = str(final.get("repair_resume_scope") or "").strip()
    if not scope:
        scope = "current_docx" if mode == "user" else "full_pipeline"
    if not command and scope == "current_docx":
        command = _default_rebuild_command(display_out)
    repair_next = str(final.get("repair_next_action") or "").strip()
    if status == "stopped_no_improvement":
        reason = "自动修复已停止：QA 错误数没有继续改善。"
    elif status == "stopped_no_supported_auto_repair":
        reason = "自动修复已停止：当前问题没有匹配到安全的自动修复动作。"
    elif status == "stopped_max_rounds":
        reason = "自动修复已停止：已达到最大修复轮次。"
    else:
        reason = "自动修复已停止。"
    if stop_detail:
        reason += f" {stop_detail}"
    next_action = repair_next or "请打开 `qa_report.md`、`qa_repair_plan.md` 和 `build_generated.py`，让 Agent 继续处理剩余错误。"
    if command and command not in next_action:
        next_action += f" 修复后运行：`{command}`"
    return {
        "next_action": f"{reason} {next_action}".strip(),
        "resume_scope": scope,
        "resume_command": command,
    }


def _default_rebuild_command(display_out: str) -> str:
    return "python " + _quote_cmd_arg(_join_report_path(display_out, "build_generated.py"))


def _quote_cmd_arg(value: str) -> str:
    text = str(value)
    if not text:
        return '""'
    if re.search(r'[\s"&|<>^]', text):
        return '"' + text.replace('"', r'\"') + '"'
    return text


def _workflow_resume_command_fallback(out_path: Path, fallback_mode: str) -> str:
    try:
        workflow = json.loads((out_path / "workflow_mode.json").read_text(encoding="utf-8"))
    except Exception:
        return ""

    mode = str(workflow.get("mode") or fallback_mode or "user")
    template = workflow.get("template")
    content = workflow.get("content")
    md_file = workflow.get("md")
    if not md_file and template and content and str(template).lower().endswith(".md") and template == content:
        md_file = template

    args = ["python", "run_pipeline.py", "--mode", mode]
    if md_file:
        args.extend(["--md", str(md_file)])
    elif template and content:
        args.extend(["--template", str(template), "--content", str(content)])
    else:
        return ""

    qa_level = str(workflow.get("qa_level") or "").strip().lower()
    if qa_level in {"basic", "strict", "visual"}:
        args.extend(["--qa-level", qa_level])
    if workflow.get("auto_repair"):
        args.append("--auto-repair")
        if workflow.get("repair_max_rounds"):
            args.extend(["--repair-max-rounds", str(workflow.get("repair_max_rounds"))])
        if workflow.get("repair_stop_no_improve"):
            args.extend(["--repair-stop-no-improve", str(workflow.get("repair_stop_no_improve"))])
    if workflow.get("require_wps"):
        args.append("--require-wps")
    if workflow.get("update_golden"):
        args.append("--update-golden")
    if workflow.get("golden_dir"):
        args.extend(["--golden-dir", str(workflow.get("golden_dir"))])
    return " ".join(_quote_cmd_arg(arg) for arg in args)


def _report_to_markdown(report: Dict[str, Any]) -> str:
    lines = [
        "# 自动修复闭环报告",
        "",
        f"- 状态：`{report.get('status')}`",
        f"- 结果：`{'已收敛' if report.get('ok') else '已停止'}`",
        f"- 模式：`{report.get('mode')}`",
        f"- QA 等级：`{report.get('qa_level')}`",
        f"- 最终 DOCX：`{report.get('final_docx')}`",
        f"- 最终错误：`{report.get('final_errors')}`",
        f"- 最终警告：`{report.get('final_warnings')}`",
        f"- 下一步：{report.get('next_action')}",
        f"- 修复范围：`{report.get('resume_scope')}`",
    ]
    if report.get("resume_command"):
        lines.append(f"- 恢复命令：`{report.get('resume_command')}`")
    if report.get("stop_detail"):
        lines.append(f"- 停止原因：{report.get('stop_detail')}")
    lines.append(f"- warning 处理策略：{report.get('warning_policy')}")
    blockers = report.get("blockers") or []
    if blockers:
        lines.extend(["", "## 需要用户补充", ""])
        for item in blockers:
            lines.append(f"- `{item.get('code')}` ({item.get('auto_level')}): {item.get('user_action')}")
    lines.extend(["", "## 修复轮次", ""])
    for item in report.get("rounds") or []:
        lines.append(
            f"- 第 `{item.get('round')}` 轮 `{item.get('phase')}`："
            f"错误 `{item.get('total_errors')}`，警告 `{item.get('total_warnings')}`，"
            f"错误码 `{', '.join(item.get('error_codes') or []) or '-'}`"
        )
        for action in item.get("actions") or []:
            lines.append(f"  - 动作：`{action.get('kind')}` {action.get('summary') or action.get('code') or action.get('codes')}")
    lines.extend([
        "",
        "## 人工检查",
        "",
        "- 自动 QA 收敛不等于 100% 正确。",
        "- 用 Word/WPS 打开最终 DOCX，核对分页、图片、公式、表格和目录。",
    ])
    if int(report.get("final_warnings") or 0):
        lines.append("- 交付前查看 qa_report.md 和 repair_loop_report.md 中的剩余 warning。")
    else:
        lines.append("- 当前启用的 QA 未报告剩余 warning。")
    lines.append("")
    return "\n".join(lines)


def _report_display_path(path: Path) -> str:
    parts = list(path.parts)
    if "Outputs" in parts:
        return "/".join(parts[parts.index("Outputs"):])
    return path.name


def _join_report_path(parent: str, child: str) -> str:
    return (str(parent).rstrip("/\\") + "/" + str(child).lstrip("/\\")).replace("\\", "/")


def _sanitize_log(text: str, out_path: Path) -> str:
    safe = _report_display_path(out_path)
    raw = str(out_path)
    return str(text or "").replace(raw, safe).replace(raw.replace("\\", "/"), safe)


def _load_json(path: Path, default: Any) -> Any:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default
