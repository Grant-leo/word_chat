"""
content_parser.py — 从文本资料 docx 提取结构化内容
输出: content.json (章节结构 + 段落 + 图片路径 + 表格 + 参考文献)
"""
from docx import Document
from docx.shared import Pt, Inches
from lxml import etree
import json, os, re, shutil, hashlib

try:
    from formula_semantics import (
        CATEGORY_CONTAMINATED,
        classify_formula_text,
        formula_should_number as semantic_formula_should_number,
        is_formula_label,
        is_formula_problem_text,
        looks_like_formula_text as semantic_looks_like_formula_text,
        split_inline_math_spans,
    )
except ImportError:  # pragma: no cover - package-style imports
    from .formula_semantics import (
        CATEGORY_CONTAMINATED,
        classify_formula_text,
        formula_should_number as semantic_formula_should_number,
        is_formula_label,
        is_formula_problem_text,
        looks_like_formula_text as semantic_looks_like_formula_text,
        split_inline_math_spans,
    )

def emu_to_pt(emu):
    if emu is None: return None
    return round(emu / 12700, 1)


def _math_text(elem):
    """Extract plain text from a math OOXML element."""
    M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    parts = []
    for t in elem.iter(f'{{{M}}}t'):
        if t.text:
            parts.append(t.text)
    return ''.join(parts)


_FORMULA_TRAILING_LABEL_RE = re.compile(r'((?:\s*[\(\uff08]\s*\d+(?:\s*[-.]\s*\d+)?\s*[\)\uff09])+\s*)$')


def _split_trailing_formula_labels(text):
    t = str(text or '').strip()
    m = _FORMULA_TRAILING_LABEL_RE.search(t)
    if not m:
        return t, []
    labels = re.findall(r'[\(\uff08]\s*(\d+(?:\s*[-.]\s*\d+)?)\s*[\)\uff09]', m.group(1))
    return t[:m.start()].strip(), labels


def _should_strip_trailing_formula_labels(text):
    body, labels = _split_trailing_formula_labels(text)
    if not body or not labels:
        return False
    # A single trailing "(1)" can be a function argument or superscript marker
    # after OOXML is flattened to text, e.g. f(1) or x^{(1)}. Treat it as an
    # equation number only when the preceding body has equation-like structure.
    if len(labels) == 1 and re.search(r'[A-Za-z\u0370-\u03ff]$', body):
        return False
    if len(labels) >= 2:
        return True
    return bool(re.search(r'[=＝≈≤≥<>]', body) and re.search(r'\d|[+*/×÷%·]', body))


def _strip_trailing_formula_labels(text):
    """Remove stale equation numbers copied from source documents."""
    if not _should_strip_trailing_formula_labels(text):
        return str(text or '').strip()
    body, _labels = _split_trailing_formula_labels(text)
    return body


def _strip_trailing_formula_labels_from_xml(xml):
    """Strip trailing formula labels from m:t nodes while preserving OMML."""
    try:
        root = etree.fromstring(str(xml or '').encode('utf-8'))
    except Exception:
        return xml, '', False
    M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    nodes = [n for n in root.iter(f'{{{M}}}t') if n.text]
    original = ''.join(n.text or '' for n in nodes)
    stripped = _strip_trailing_formula_labels(original)
    if not original or stripped == original:
        return xml, original, False
    remove_chars = len(original) - len(stripped)
    for node in reversed(nodes):
        if remove_chars <= 0:
            break
        txt = node.text or ''
        if remove_chars >= len(txt):
            node.text = ''
            remove_chars -= len(txt)
        else:
            node.text = txt[:-remove_chars]
            remove_chars = 0
    return etree.tounicode(root, with_tail=False), stripped, True


def extract_math(para):
    """Extract OOXML math elements from a paragraph. Returns (text, math_list).
    math_list entries: {'type': 'inline'|'display', 'xml': escaped_xml_string, 'text': plain_text}
    Text is cleaned of formula garbling."""
    xml = para._element.xml
    if 'm:oMath' not in xml and 'oMathPara' not in xml:
        return para.text, []

    math_list = []
    root = para._element

    # Extract m:oMathPara (display formulas) — these ARE the paragraph
    for omp in root.findall('{http://schemas.openxmlformats.org/officeDocument/2006/math}oMathPara'):
        raw = etree.tounicode(omp, with_tail=False)
        math_list.append({'type': 'display', 'xml': raw, 'text': _math_text(omp)})

    # Extract m:oMath (inline formulas) — embedded in runs
    for om in root.iter('{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath'):
        # Skip if already inside an oMathPara
        parent_tag = om.getparent().tag.split('}')[-1] if '}' in om.getparent().tag else om.getparent().tag
        if parent_tag == 'oMathPara':
            continue
        raw = etree.tounicode(om, with_tail=False)
        math_list.append({'type': 'inline', 'xml': raw, 'text': _math_text(om)})

    # Reconstruct text without formula garbling: get text from runs, skip math-only runs
    text_parts = []
    for child in root:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            # Skip runs that contain only math (no w:t)
            has_text = child.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') is not None
            if has_text:
                text_parts.append(child.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t').text or '')
        elif tag == 'oMathPara' or tag == 'oMath':
            # Don't add math XML to text
            pass
        elif tag == 'pPr':
            pass
        else:
            # Other elements (like w:r with math only) — skip
            pass

    text = ''.join(text_parts).strip()
    return text, math_list


def detect_heading_level(para):
    """Detect heading level using OOXML-direct size + heuristics."""
    if not para.runs:
        return 0
    text = para.text.strip()
    if not text:
        return 0
    # Figure/table captions are captions, not outline headings or TOC entries.
    if re.match(r'^(图|表)\s*\d+(?:[.-]\d+)?\s*', text):
        return 0
    if re.fullmatch(r'[\d\s.,，．。]+', text) or re.fullmatch(r'\d+\s*[.．]\s*\d+', text):
        return 0

    # ── Label-style headings (Abstract:, Key words:, 摘要：, 关键词：) ──
    # These are detected by text pattern regardless of paragraph length or OOXML formatting,
    # so that mid-length paragraphs like "Key words: memetics; ..." (~100 chars) are caught.
    # Case-insensitive for English labels so "key words:", "KEY WORDS:", etc. are all detected.
    label_patterns = [
        (r'(?i)^(Abstract\s*:?)', 2),
        (r'(?i)^(Key\s*words?\s*:?)', 2),
        (r'^(摘要\s*[：:]?)', 2),
        (r'^(关键词\s*[：:]?)', 2),
    ]
    for pat, lvl in label_patterns:
        if re.match(pat, text):
            return lvl

    # ── Explicit heading patterns should win over font heuristics ──
    # Chapter: 第1章 / 第一章 / 1 绪论 / Chapter 1
    if re.match(r'^第[一二三四五六七八九十\d]+章', text):
        return 1
    m_count_heading = re.match(r'^\d+\s+([\u4e00-\u9fff])', text)
    if m_count_heading and m_count_heading.group(1) in set('种个天年月日吨项组类场'):
        return 2 if _looks_like_heading_style(para) else 0
    if re.match(r'^(?:Chapter\s*)?\d+\s+[\u4e00-\u9fffA-Za-z]', text) and not re.match(r'^\d+\.\d+', text):
        return 1
    if len(text) <= 80 and re.match(r'^[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e\u5343]+[\u3001\uff0e.]\s*\S+', text):
        return 1
    if len(text) <= 80 and re.match(r'^\d+[\u3001\uff0e]\s*\S+', text):
        return 1
    if len(text) <= 80 and re.match(r'^[\uff08(][\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+[\uff09)]\s*\S+', text):
        return 2
    # Numbered: 1.1 = level 2, 1.1.1 = level 3
    if re.match(r'^\d+\.\d+\s*(?:MWh|MW|kWh|kg|h|吨|元|%|[+\-−=,，）。])', text, re.I):
        return 0
    if re.match(r'^\d+\.\d+\.\d+\s*', text):
        return 3
    if re.match(r'^\d+\.\d+\s*', text):
        return 2
    # ── Numbered heading patterns for long paragraphs (>200 chars) ──
    if len(text) > 200:
        heading_patterns = [
            r'^(\d+\.\s+\w)', r'^(\d+\.\d+\s+\w)',
        ]
        for pat in heading_patterns:
            if re.match(pat, text):
                return 1 if re.match(r'^(\d+\.\s+\w)', text) else 2
        return 0

    # Skip formatting notes (Chinese parenthetical instructions or short bracketed text)
    if text.startswith('（') and (text.endswith('）') or len(text) < 60):
        return 0

    style_level = _heading_level_from_style(para)
    if style_level and len(text) <= 80 and not text.endswith('。'):
        if re.search(r'[=×÷ΣΠ∫√∞≈±]', text) and len(text) < 100:
            return 0
        return style_level

    # Get size from OOXML directly (bypass python-docx API None issue)
    max_size = 0
    is_bold = False
    for r in para.runs:
        rPr = r._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        if rPr is not None:
            for child in rPr:
                tag = child.tag.split('}')[-1]
                if tag == 'sz' or tag == 'szCs':
                    try:
                        max_size = max(max_size, int(child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')) / 2.0)
                    except:
                        pass
                if tag == 'b':
                    val = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if str(val).lower() not in ('0', 'false', 'off'):
                        is_bold = True

    if not is_bold or max_size < 12:
        return 0

    # Exclude non-heading patterns
    if re.search(r'\(\d{4}[-–]\d{4}\)', text):  # year range
        return 0
    if re.search(r'\d{6}', text):  # postal code
        return 0
    if 'University' in text and len(text) > 100:  # long affiliation
        return 0
    if len(text.split()) <= 2 and not any(c.isdigit() for c in text):
        return 0
    # Formulas: contain math operators and are centered/italic
    if re.search(r'[=×÷ΣΠ∫√∞≈±]', text) and len(text) < 100:
        return 0

    # A heading must be reasonably short and not a full sentence
    if len(text) > 80:
        return 0
    # Full Chinese sentences (ending with 。) are body text, not headings
    if text.endswith('。'):
        return 0

    # Level detection
    numbered = bool(re.match(r'^[\d]+\.[\d]*\s', text))  # "1. " or "2.1 "

    if max_size >= 15:
        return 1  # 小三号及以上 = h1
    if max_size >= 14:
        return 2 if (numbered or len(text) < 60) else 1
    if max_size >= 12:
        return 2 if numbered else 3
    return 0


class ImageRegistry:
    """Per-extraction image registry.

    The same DOCX image relationship can be encountered more than once during
    verification or when Word duplicates drawing markup.  Saving by a content
    hash avoids the historical bug where one logical figure produced hundreds
    or thousands of copied files.  The registry is local to one extraction, so
    there is no university-, title-, or path-specific hardcoding.
    """

    def __init__(self, fig_dir, prefix='img'):
        self.fig_dir = fig_dir
        self.prefix = prefix
        self.counter = 0
        self.by_hash = {}
        self.failures = []

    def save_relationship_image(self, rel):
        if 'image' not in getattr(rel, 'reltype', ''):
            return None
        try:
            blob = rel.target_part.blob
            digest = hashlib.sha256(blob).hexdigest()[:20]
            if digest in self.by_hash:
                return self.by_hash[digest]

            ext = rel.target_ref.rsplit('.', 1)[-1].lower()
            if ext not in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'tif', 'tiff', 'webp'):
                ext = 'png'
            self.counter += 1
            fname = f'{self.prefix}_{self.counter:03d}.{ext}'
            fpath = os.path.join(self.fig_dir, fname)
            with open(fpath, 'wb') as f:
                f.write(blob)
            self.by_hash[digest] = fname
            return fname
        except Exception as exc:
            self.failures.append({
                'target': getattr(rel, 'target_ref', ''),
                'error': str(exc)[:200],
            })
            return None


def extract_images_from_para(para, fig_dir, prefix='img', registry=None):
    """Extract inline images by rId. Returns filenames in paragraph order.

    Fixes two duplicate sources without hardcoding:
    1) `seen_rids` is paragraph-scoped, not run-scoped;
    2) image bytes are de-duplicated by SHA-256 within the extraction pass.
    """
    saved = []
    registry = registry or ImageRegistry(fig_dir, prefix)
    A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    seen_rids = set()
    for run in para.runs:
        xml = run._element.xml
        if 'w:drawing' not in xml and 'wp:inline' not in xml and 'wp:anchor' not in xml:
            continue
        for blip in run._element.iter(f'{{{A_NS}}}blip'):
            embed = blip.get(f'{{{R_NS}}}embed')
            if not embed or embed in seen_rids:
                continue
            seen_rids.add(embed)
            if embed in para.part.rels:
                fname = registry.save_relationship_image(para.part.rels[embed])
                if fname:
                    saved.append(fname)
            else:
                registry.failures.append({'target': embed, 'error': 'relationship id not found'})
    return saved




def _local_name(el):
    return el.tag.split('}')[-1] if '}' in el.tag else el.tag


_PLACEHOLDER_RE = re.compile(
    r'(\[[^\]\n]*(?:\u62a5\u540d|\u5e8f\u53f7|\u59d3\u540d|\u5b66\u53f7|\u5b66\u9662|\u4e13\u4e1a|\u73ed\u7ea7|\u9898\u76ee|\u6307\u5bfc|\u6559\u5e08|\u65e5\u671f|\u7f16\u7801|\u5f85\u586b|\u8bf7\u8f93\u5165|XX|XXX)[^\]\n]*\])'
    r'|(\{\{[^}]+\}\}|TODO|FIXME|\u5f85\u586b\u5199|\u5f85\u8865\u5168|XXXX)',
    re.I,
)


def _is_unfilled_placeholder_text(text):
    return bool(_PLACEHOLDER_RE.search(str(text or '')))


def _placeholder_samples(paragraphs, limit=8):
    out = []
    for idx, para in enumerate(paragraphs, 1):
        text = str(getattr(para, 'text', '') or '').strip()
        if text and _is_unfilled_placeholder_text(text):
            out.append({'paragraph': idx, 'text': text[:120]})
            if len(out) >= limit:
                break
    return out


def _extract_labeled_title(text):
    t = str(text or '').strip()
    m = re.match(r'^\s*(?:\u8bba\u6587\u9898\u76ee|\u9898\u76ee|\u6807\u9898)\s*[:\uff1a]\s*(.+?)\s*$', t)
    if not m:
        return ''
    value = m.group(1).strip()
    return '' if _is_unfilled_placeholder_text(value) else value


def _compact_text(text):
    return re.sub(r'[\s\u3000]+', '', str(text or '')).upper()


def _is_source_toc_title(text):
    compact = _compact_text(text)
    return compact in {'\u76ee\u5f55', '\u76ee\u6b21', 'CONTENTS', 'TABLEOFCONTENTS'}


def _is_source_toc_entry(text):
    t = str(text or '').strip()
    if not t:
        return True
    if len(t) > 160:
        return False
    if _is_source_toc_title(t):
        return True
    if re.search(r'(?:\.{2,}|…+|·{2,}|_{2,})\s*\d+\s*$', t):
        return True
    toc_prefix = r'(?:第[一二三四五六七八九十百千万\d]+章|[一二三四五六七八九十]+[、.．]|\d+(?:\.\d+)*|摘要|ABSTRACT|关键词|KEY\s*WORDS?|参考文献|致谢|附录|APPENDIX|ACKNOWLEDGEMENTS?)'
    return bool(re.match(r'^' + toc_prefix + r'\s+.+\s+\d+\s*$', t, re.I))


def _is_unpaged_source_toc_entry(text, para=None):
    t = str(text or '').strip()
    if not t or len(t) > 100:
        return False
    if _is_source_toc_entry(t):
        return True
    if re.search(r'[。！？!?；;]\s*$', t):
        return False
    if _is_unfilled_placeholder_text(t):
        return False
    prefix = (
        r'(?:第[一二三四五六七八九十百千万\d]+章\s*\S+'
        r'|[一二三四五六七八九十]+[、.．]\s*\S+'
        r'|\d+(?:\.\d+)*\s+\S+'
        r'|摘要|ABSTRACT|关键词|KEY\s*WORDS?'
        r'|参考文献|REFERENCES?|致谢|ACKNOWLEDGEMENTS?'
        r'|附录|APPENDIX(?:\s+\S+)?)'
    )
    if re.match(r'^' + prefix + r'\s*$', t, re.I):
        return True
    if para is not None and _looks_like_heading_style(para) and len(t) <= 80:
        return True
    return False


def _simple_cn_number(value):
    s = str(value or '').strip()
    if not s:
        return None
    if s.isdigit():
        return int(s)
    digits = {'零': 0, '一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9}
    if s == '十':
        return 10
    if '十' in s:
        left, right = s.split('十', 1)
        tens = digits.get(left, 1 if left == '' else 0)
        ones = digits.get(right, 0) if right else 0
        return tens * 10 + ones
    if len(s) == 1 and s in digits:
        return digits[s]
    return None


def _toc_entry_key(text):
    t = str(text or '').strip()
    t = re.sub(r'(?:\.{2,}|…+|·{2,}|_{2,})\s*(?:[ivxlcdm]+|\d+)\s*$', '', t, flags=re.I)
    t = re.sub(r'\s+(?:[ivxlcdm]+|\d+)\s*$', '', t, flags=re.I)
    return _compact_text(t)


def _toc_entry_order(text):
    t = str(text or '').strip()
    m = re.match(r'^第([一二三四五六七八九十百千万\d]+)章', t)
    if m:
        n = _simple_cn_number(m.group(1))
        return ('chapter', n) if n is not None else None
    m = re.match(r'^([一二三四五六七八九十]+)[、.．]', t)
    if m:
        n = _simple_cn_number(m.group(1))
        return ('cn', n) if n is not None else None
    m = re.match(r'^(\d+)(?:\.\d+)*\b', t)
    if m:
        return ('num', int(m.group(1)))
    return None


def _source_toc_skip_count_after_title(paragraphs, title_idx, max_scan=160):
    """Return how many paragraphs after a source TOC title should be skipped."""
    plist = list(paragraphs)
    n = len(plist)
    if title_idx < 0 or title_idx >= n:
        return 0
    first_visible = None
    scan_end = min(n, title_idx + 1 + max_scan)
    for idx in range(title_idx + 1, scan_end):
        text = str(getattr(plist[idx], 'text', '') or '').strip()
        if text:
            first_visible = idx
            break
    if first_visible is None:
        return 0
    first_text = str(getattr(plist[first_visible], 'text', '') or '').strip()
    if not _is_unpaged_source_toc_entry(first_text, plist[first_visible]):
        return 0

    first_boundary_idx = None
    for idx in range(title_idx + 1, scan_end):
        try:
            if _paragraph_has_page_or_section_break(plist[idx]._element):
                first_boundary_idx = idx
                break
        except Exception:
            continue

    visible_count = 0
    saw_paged_entry = False
    saw_boundary = False
    boundary_idx = None
    title_has_heading_style = _looks_like_heading_style(plist[title_idx])
    non_toc_before_boundary = False
    skip_until = title_idx
    seen_keys = set()
    last_order = None
    for idx in range(title_idx + 1, scan_end):
        para = plist[idx]
        text = str(getattr(para, 'text', '') or '').strip()
        try:
            has_boundary = _paragraph_has_page_or_section_break(para._element)
        except Exception:
            has_boundary = False
        if not text:
            skip_until = idx
            if has_boundary:
                saw_boundary = True
                boundary_idx = idx
                break
            continue
        paged_entry = _is_source_toc_entry(text) and not _is_source_toc_title(text)
        unpaged_entry = _is_unpaged_source_toc_entry(text, para)
        if not (paged_entry or unpaged_entry):
            if first_boundary_idx is not None and not title_has_heading_style:
                non_toc_before_boundary = True
                skip_until = idx
                if has_boundary:
                    saw_boundary = True
                    boundary_idx = idx
                    break
                continue
            break

        key = _toc_entry_key(text)
        order = _toc_entry_order(text)
        if visible_count > 0:
            if key and key in seen_keys:
                if first_boundary_idx is not None and not title_has_heading_style:
                    non_toc_before_boundary = True
                    skip_until = idx
                    if has_boundary:
                        saw_boundary = True
                        boundary_idx = idx
                        break
                    continue
                break
            if saw_paged_entry and not paged_entry:
                if first_boundary_idx is not None and not title_has_heading_style:
                    non_toc_before_boundary = True
                    skip_until = idx
                    if has_boundary:
                        saw_boundary = True
                        boundary_idx = idx
                        break
                    continue
                break
            if order and last_order and order[0] == last_order[0] and order[1] <= last_order[1]:
                if first_boundary_idx is not None and not title_has_heading_style:
                    non_toc_before_boundary = True
                    skip_until = idx
                    if has_boundary:
                        saw_boundary = True
                        boundary_idx = idx
                        break
                    continue
                break

        visible_count += 1
        saw_paged_entry = saw_paged_entry or paged_entry
        if key:
            seen_keys.add(key)
        if order:
            last_order = order
        skip_until = idx
        if has_boundary:
            saw_boundary = True
            boundary_idx = idx
            break

    if saw_boundary:
        if saw_paged_entry or visible_count >= 2 or (visible_count >= 1 and non_toc_before_boundary and not title_has_heading_style):
            return max(skip_until, boundary_idx or skip_until) - title_idx
        return 0
    if visible_count < 2 and not saw_paged_entry:
        return 0
    return max(0, skip_until - title_idx)


def _paragraph_has_page_or_section_break(p_elem):
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    if p_elem.find(f'.//{{{W}}}sectPr') is not None:
        return True
    for br in p_elem.iter(f'{{{W}}}br'):
        if br.get(f'{{{W}}}type') == 'page':
            return True
    return False


def _paragraph_style_id(para):
    try:
        pPr = para._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            st = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
            if st is not None:
                return st.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') or ''
    except Exception:
        pass
    return ''


def _heading_level_from_style(para):
    style_id = _paragraph_style_id(para)
    try:
        style_name = para.style.name if para.style else ''
    except Exception:
        style_name = ''
    compact = _compact_text(style_id + style_name)
    m = re.search(r'HEADING([1-6])', compact)
    if m:
        return int(m.group(1))
    m = re.search(r'标题([1-6])', compact)
    if m:
        return int(m.group(1))
    if re.search(r'HEADING|TITLE|CHAPTER', compact) or '标题' in compact or '章' in compact:
        return 1
    return 0


def _looks_like_heading_style(para):
    return bool(_heading_level_from_style(para))


def _run_text_preserve_breaks(r_elem):
    """Return visible text carried by a run, preserving explicit breaks."""
    parts = []
    for child in r_elem:
        name = _local_name(child)
        if name == 't':
            parts.append(child.text or '')
        elif name in ('tab',):
            parts.append('\t')
        elif name in ('br', 'cr'):
            parts.append('\n')
    return ''.join(parts)


def _math_entry_from_ooxml(math_elem, math_type='inline'):
    raw = etree.tounicode(math_elem, with_tail=False)
    clean_xml, clean_text, had_label = _strip_trailing_formula_labels_from_xml(raw)
    if not clean_text:
        clean_text = _strip_trailing_formula_labels(_math_text(math_elem))
    entry = {'type': math_type, 'xml': clean_xml, 'text': clean_text}
    if clean_text:
        entry['formula_semantics'] = classify_formula_text(clean_text).to_dict()
    if had_label:
        entry['had_number_label'] = True
    return entry


def _images_from_run_ooxml(run_elem, rels, registry, seen_rids, location='body'):
    """Extract images from one OOXML run in its exact paragraph position."""
    A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    V_NS = 'urn:schemas-microsoft-com:vml'
    out = []

    def add_rid(rid):
        if not rid or rid in seen_rids:
            return
        seen_rids.add(rid)
        if rid in rels:
            fname = registry.save_relationship_image(rels[rid])
            if fname:
                item = {'role': 'image', 'image': fname}
                if location and location != 'body':
                    item['location'] = location
                out.append(item)
        else:
            registry.failures.append({'target': rid, 'error': 'relationship id not found'})

    for blip in run_elem.iter(f'{{{A_NS}}}blip'):
        add_rid(blip.get(f'{{{R_NS}}}embed') or blip.get(f'{{{R_NS}}}link'))
    for imagedata in run_elem.iter(f'{{{V_NS}}}imagedata'):
        add_rid(imagedata.get(f'{{{R_NS}}}id') or imagedata.get(f'{{{R_NS}}}embed'))
    return out


def _image_items_from_ooxml(container_elem, rels, registry, location='body'):
    """Extract all image runs from an arbitrary OOXML container.

    Body paragraphs use `paragraph_stream_items()`, but images can also live
    inside table cells.  Keeping this helper generic lets table-cell drawings
    enter the same content image pipeline instead of disappearing silently.
    """
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    seen_rids = set()
    out = []
    for run_elem in container_elem.iter(f'{{{W_NS}}}r'):
        out.extend(_images_from_run_ooxml(run_elem, rels, registry, seen_rids, location=location))
    return out


def _non_body_image_entries(doc):
    """Return header/footer images that are outside the body content stream."""
    entries = []
    seen = set()
    for sec_idx, section in enumerate(doc.sections):
        for attr in (
            'header', 'first_page_header', 'even_page_header',
            'footer', 'first_page_footer', 'even_page_footer',
        ):
            try:
                part = getattr(section, attr).part
            except Exception:
                continue
            for rid, rel in getattr(part, 'rels', {}).items():
                if 'image' not in getattr(rel, 'reltype', ''):
                    continue
                try:
                    digest = hashlib.sha256(rel.target_part.blob).hexdigest()[:20]
                except Exception:
                    digest = f'{id(part)}:{rid}'
                key = (attr, digest)
                if key in seen:
                    continue
                seen.add(key)
                entries.append({
                    'location': f'section_{sec_idx + 1}_{attr}',
                    'target': getattr(rel, 'target_ref', ''),
                })
    return entries


def paragraph_stream_items(para, registry):
    """Yield paragraph text/image/math items in true OOXML run order.

    The previous extractor saved all images first and appended paragraph text
    afterwards, which could create: image -> body text -> caption.  This routine
    flushes text before every drawing or math run, then emits the token at the
    exact location where Word stores it.  It is structural and does not depend
    on a school name, figure number, or fixed paragraph index.
    """
    items = []
    buf = []
    seen_rids = set()

    def flush_text():
        text = ''.join(buf)
        buf.clear()
        if text.strip():
            items.append({'role': 'text', 'text': text})

    def append_math(math_elem, math_type='inline'):
        entry = _math_entry_from_ooxml(math_elem, math_type)
        flush_text()
        semantic = classify_formula_text(entry.get('text') or '')
        if is_formula_problem_text(entry.get('text') or ''):
            items.append({
                'role': 'formula_problem',
                'problem': 'contaminated_formula_text',
                'source': 'omml',
                'text': entry.get('text') or '',
                'math': [entry],
                'formula_semantics': semantic.to_dict(),
            })
            return
        if _omml_text_looks_like_body(entry.get('text') or ''):
            items.append({'role': 'text', 'text': entry.get('text') or ''})
            return
        if math_type == 'display':
            items.append({
                'role': 'formula',
                'source': 'omml',
                'text': entry.get('text') or '',
                'math': [entry],
                'numbered': bool(entry.get('had_number_label')) or _formula_should_number(entry.get('text') or ''),
            })
        else:
            items.append({
                'role': 'math_inline',
                'source': 'omml',
                'text': entry.get('text') or '',
                'math': [entry],
            })

    def consume_run(run_elem):
        for part in run_elem:
            name = _local_name(part)
            if name in ('drawing', 'pict'):
                flush_text()
                items.extend(_images_from_run_ooxml(run_elem, para.part.rels, registry, seen_rids))
            elif name == 'oMath':
                append_math(part, 'inline')
            elif name == 't':
                if part.text:
                    buf.append(part.text)
            elif name in ('tab',):
                buf.append('\t')
            elif name in ('br', 'cr'):
                buf.append('\n')

    for child in para._element:
        name = _local_name(child)
        if name == 'r':
            consume_run(child)
        elif name in ('hyperlink',):
            for r in child:
                if _local_name(r) != 'r':
                    continue
                consume_run(r)
        elif name in ('oMath', 'oMathPara'):
            append_math(child, 'display' if name == 'oMathPara' else 'inline')
    flush_text()
    return items


def _append_stream_run_group(section, runs, in_appendix=False):
    if not runs:
        return
    text = ''.join(str(r.get('text') or '') for r in runs).strip()
    math_items = []
    for run in runs:
        if run.get('type') == 'math':
            math_items.extend(run.get('math') or [])
    if not math_items:
        _append_text_or_code(section, text, in_appendix=in_appendix)
        return
    semantic = classify_formula_text(text)
    if is_formula_problem_text(text):
        section['paragraphs'].append({
            'role': 'formula_problem',
            'problem': 'contaminated_formula_text',
            'source': 'omml',
            'text': text,
            'math': math_items,
            'formula_semantics': semantic.to_dict(),
        })
        return
    non_math_text = ''.join(str(r.get('text') or '') for r in runs if r.get('type') != 'math').strip()
    if non_math_text:
        section['paragraphs'].append({
            'role': 'rich_text',
            'text': text,
            'runs': runs,
            'math': math_items,
        })
    else:
        section['paragraphs'].append({
            'role': 'formula',
            'source': 'omml',
            'text': text,
            'math': math_items,
            'numbered': any(m.get('had_number_label') for m in math_items) or _formula_should_number(text),
        })


def _caption_kind(item):
    if isinstance(item, dict):
        role = item.get('role')
        if role in ('figure_caption', 'table_caption'):
            return role
        text = item.get('text') or ''
    else:
        text = str(item or '')
    if _is_figure_caption(text):
        return 'figure_caption'
    if _is_table_caption(text):
        return 'table_caption'
    return None


def _is_image_item(item):
    return isinstance(item, dict) and item.get('role') == 'image' and item.get('image')


def _caption_text(item):
    if isinstance(item, dict):
        return item.get('text') or ''
    return str(item or '')


def pair_figure_blocks(paragraphs):
    """Pair images with captions while preserving all text.

    Besides the normal `image -> caption` layout, this fixes the two observed
    drift patterns without hardcoding figure numbers or page positions:
      * image, body text, caption  -> image+caption, body text
      * image, image, caption, caption -> image+caption, image+caption

    Look-ahead is deliberately small.  When no nearby figure caption exists the
    image token is left unchanged so content is never dropped or invented.
    """
    out = []
    i = 0
    n = len(paragraphs or [])
    while i < n:
        item = paragraphs[i]
        if not _is_image_item(item):
            out.append(item)
            i += 1
            continue

        images = []
        while i < n and _is_image_item(paragraphs[i]):
            images.append(paragraphs[i])
            i += 1

        # Look ahead through a small local window for captions.  Non-caption
        # text between image and caption is temporarily held, then emitted after
        # the paired figure so it can no longer split picture and title.
        j = i
        held_text = []
        captions = []
        max_probe = min(n, i + max(8, len(images) * 3 + 3))
        while j < max_probe and len(captions) < len(images):
            nxt = paragraphs[j]
            if _is_image_item(nxt):
                break
            kind = _caption_kind(nxt)
            if kind == 'figure_caption':
                captions.append(nxt)
                j += 1
                # Continue consuming immediately stacked figure captions for
                # the image,image,caption,caption case.
                continue
            if kind == 'table_caption':
                break
            held_text.append(nxt)
            j += 1

        if captions:
            for idx, img in enumerate(images):
                cap = captions[idx] if idx < len(captions) else None
                if cap is not None:
                    out.append({'role': 'figure', 'image': img.get('image'), 'caption': _caption_text(cap)})
                else:
                    out.append(img)
            if len(captions) > len(images):
                out.extend(captions[len(images):])
            out.extend(held_text)
            i = j
            continue

        # No nearby caption: keep all images in their original place and leave
        # subsequent text to be processed normally.
        out.extend(images)
    return out




def _normalize_role_heading(text):
    return re.sub(r'\s+', ' ', str(text or '').strip())


def _ascii_alpha_ratio(text):
    text = str(text or '')
    if not text:
        return 0.0
    return sum(1 for c in text if c.isascii() and c.isalpha()) / max(len(text), 1)


def _classify_section_role(heading, level=0):
    """Map a detected heading to a semantic role used by the renderer."""
    h = _normalize_role_heading(heading)
    h_compact = re.sub(r'[\s：:]+', '', h).lower()
    if h_compact in ('摘要', '中文摘要'):
        return 'cn_abstract'
    if h_compact in ('关键词', '关键字') or h.startswith('关键词'):
        return 'cn_keywords'
    if h_compact in ('abstract', 'englishabstract'):
        return 'en_abstract'
    if h.upper().replace(' ', '').startswith('KEYWORDS') or re.match(r'(?i)^key\s*words?', h):
        return 'en_keywords'
    if re.match(r'(?i)^references?$', h) or h.startswith('参考文献'):
        return 'references'
    if re.match(r'(?i)^acknowledg(?:e)?ments?\b|^acknowledgment\b', h):
        return 'acknowledgement'
    if re.search(r'致\s*谢', h):
        return 'acknowledgement'
    if re.match(r'(?i)^append(?:ix|ices)\b', h):
        return 'appendix'
    if re.search(r'附\s*录', h):
        return 'appendix'
    if level and level > 0:
        return 'heading'
    return 'body'


def _is_backmatter_heading(text):
    h = _normalize_role_heading(text)
    return bool(
        re.match(r'(?i)^acknowledg(?:e)?ments?\b|^acknowledgment\b', h)
        or re.match(r'(?i)^append(?:ix|ices)\b', h)
        or re.search(r'(致\s*谢|附\s*录)', h)
    )


def _split_heading_number(text):
    """Return (number, title) for headings such as '1.1标题' or '第1章绪论'."""
    t = str(text or '').strip()
    m = re.match(r'^(第[一二三四五六七八九十百千万\d]+章)\s*(.+)$', t)
    if m:
        return m.group(1), m.group(2).strip()
    m = re.match(r'^(\d+(?:\.\d+)*)\s*(.+)$', t)
    if m:
        return m.group(1), m.group(2).strip()
    return '', t


def _normalize_heading_spacing(text):
    num, title = _split_heading_number(text)
    return f'{num} {title}'.strip() if num and title else str(text or '').strip()


def _is_figure_caption(text):
    t = str(text or '').strip()
    return bool(
        re.match(r'^图\s*\d+(?:[.-]\d+)?\s*[^\d\s]', t)
        or re.match(r'(?i)^(?:fig\.?|figure)\s*\d+(?:[.-]\d+)?\s+[^\d\s]', t)
    )


def _is_table_caption(text):
    t = str(text or '').strip()
    return bool(
        re.match(r'^表\s*\d+(?:[.-]\d+)?\s*[^\d\s]', t)
        or re.match(r'(?i)^table\s*\d+(?:[.-]\d+)?\s+[^\d\s]', t)
    )


def _normalize_caption_spacing(text):
    t = str(text or '').strip()
    t = re.sub(r'(?i)^(fig\.?|figure)\s*(\d+(?:[.-]\d+)?)\s*', r'Fig. \2 ', t)
    t = re.sub(r'(?i)^table\s*(\d+(?:[.-]\d+)?)\s*', r'Table \1 ', t)
    return re.sub(r'^(图|表)\s*(\d+(?:[.-]\d+)?)\s*', r'\1 \2 ', t).strip()


def _clean_markdown_links(text):
    def repl(m):
        label = (m.group(1) or '').strip()
        target = (m.group(2) or '').strip()
        return label or target
    return re.sub(r'\[([^\]]+)\]\(([^)]+)\)', repl, str(text or ''))


def _clean_text_artifacts(text, preserve_newlines=False):
    """Remove generic editor/clipboard artifacts without changing content semantics."""
    t = _clean_markdown_links(text)
    t = t.replace('\u00a0', ' ')
    if preserve_newlines:
        lines = []
        for line in t.replace('\r\n', '\n').replace('\r', '\n').split('\n'):
            s = re.sub(r'[ \t]+', ' ', line).strip()
            if _is_noise_text(s):
                continue
            lines.append(s)
        return '\n'.join(lines).strip()
    t = re.sub(r'\s+', ' ', t).strip()
    return '' if _is_noise_text(t) else t


def _is_noise_text(text):
    t = str(text or '').strip()
    return t in {'复制', 'Copy', 'Plain Text', '纯文本'}


def _clean_code_text(text):
    return _clean_text_artifacts(text, preserve_newlines=True)


def _clean_formula_text(text):
    t = _clean_text_artifacts(text)
    if t.count('|') >= 3:
        t = t.replace('|', '')
    t = re.sub(r'\s+', ' ', t).strip()
    return t


def _paragraph_plain_text_from_ooxml(p_elem):
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    pieces = []
    for r in p_elem.findall(f'{{{W}}}r'):
        part = ''.join(t.text or '' for t in r.findall(f'{{{W}}}t'))
        if r.find(f'{{{W}}}br') is not None and not part:
            part = '\n'
        pieces.append(part)
    return ''.join(pieces)


def _extract_table_rows_from_ooxml(tbl_elem):
    """Preserve cell paragraph breaks so code/config tables do not collapse into one line."""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    rows = []
    for tr in tbl_elem.findall(f'{{{W}}}tr'):
        cells = []
        for tc in tr.findall(f'{{{W}}}tc'):
            paras = []
            for p in tc.findall(f'{{{W}}}p'):
                txt = _clean_text_artifacts(_paragraph_plain_text_from_ooxml(p), preserve_newlines=True).rstrip()
                if txt:
                    paras.append(txt)
            cells.append('\n'.join(paras).strip())
        rows.append(cells)
    return rows

def _looks_like_code_line(text):
    """Heuristic for network/device configuration or command-line code.

    Kept generic: it detects syntax-like command lines rather than any
    particular vendor, university, or template.
    """
    t = (text or '').strip()
    if not t or len(t) > 220:
        return False
    if re.match(r'^[A-Za-z0-9_.-]+[>#]', t):
        return True
    if re.match(r'^(interface|vlan|ip route|ip address|router|switchport|acl|rule|nat|dhcp|dns|ospf|bgp|display|show|ping|tracert|undo|quit|return|sysname|description|gateway|firewall|security-policy)\b', t, re.I):
        return True
    if re.match(r'^[a-z][a-z0-9_-]+\s+[-A-Za-z0-9_/.:]+', t) and any(ch in t for ch in ['/', '.', '-', '_']):
        return True
    return False



def _table_rows_look_like_code(rows):
    """Classify one-/two-column command tables as code, not academic tables."""
    flat = []
    for row in rows or []:
        for cell in row or []:
            for line in str(cell or '').splitlines():
                if line.strip():
                    flat.append(line.strip())
    if not flat:
        return False
    ncols = max((len(r) for r in rows or []), default=0)
    hits = sum(1 for x in flat if _looks_like_code_line(x))
    if ncols <= 1 and len(flat) >= 2 and hits >= 2:
        return True
    if ncols <= 2 and len(flat) >= 4 and hits >= max(2, len(flat) // 3):
        return True
    return False


def _code_text_from_table_rows(rows):
    lines = []
    for row in rows or []:
        cells = [str(c or '').rstrip() for c in row]
        if len(cells) == 1:
            lines.append(cells[0])
        else:
            lines.append('    '.join(cells).rstrip())
    return _clean_code_text('\n'.join(lines).rstrip())


def _looks_like_formula_text(text):
    """Detect standalone calculation/formula paragraphs.

    The rule is intentionally structural: formulas are short standalone lines
    with equality/calculation operators. Some thesis sources store formulas as
    plain text, including definition lines without numbers and continuation
    lines that start with "=".
    """
    return semantic_looks_like_formula_text(text)


def _latex_escape_text(text):
    return str(text or '').replace('\\', r'\backslash ').replace('{', r'\{').replace('}', r'\}')


def _latex_from_formula_text(text):
    t = str(text or '').strip()
    if t.startswith('$$') and t.endswith('$$'):
        return t[2:-2].strip()
    if t.startswith('$') and t.endswith('$'):
        return t[1:-1].strip()
    return ''


def _formula_should_number(text):
    if _omml_text_looks_like_body(text):
        return False
    return semantic_formula_should_number(text)


def _omml_text_looks_like_body(text):
    t = str(text or '').strip()
    if len(t) > 220:
        return True
    cjk = len(re.findall(r'[\u4e00-\u9fff]', t))
    if len(t) > 35 and cjk > 18 and re.search(r'(表明|显示|说明|分析|结果|选择|问题|模型|成本|指标)', t):
        return True
    if len(t) > 90 and cjk > 20 and re.search(r'[。；;，,\.]', t):
        return True
    return False


def _formula_item_from_text(text):
    clean = _clean_formula_text(text)
    semantic = classify_formula_text(clean)
    item = {
        'role': 'formula',
        'source': 'text',
        'text': clean,
        'numbered': _formula_should_number(clean),
        'formula_semantics': semantic.to_dict(),
    }
    latex = _latex_from_formula_text(clean)
    if latex:
        item['source'] = 'latex'
        item['latex'] = latex
    return item


def _formula_problem_item_from_text(text):
    clean = _clean_formula_text(text)
    semantic = classify_formula_text(clean)
    return {
        'role': 'formula_problem',
        'problem': 'contaminated_formula_text',
        'source': 'text',
        'text': clean,
        'formula_semantics': semantic.to_dict(),
    }


def _rich_text_item_from_inline_formula_spans(text):
    spans = split_inline_math_spans(text)
    if not spans:
        return None
    runs = []
    math_entries = []
    pos = 0
    for span in spans:
        start = int(span.get('start') or 0)
        end = int(span.get('end') or start)
        if start > pos:
            runs.append({'type': 'text', 'text': text[pos:start]})
        formula_text = str(span.get('text') or '').strip()
        if not formula_text:
            pos = max(pos, end)
            continue
        entry = {
            'type': 'inline',
            'text': formula_text,
            'formula_semantics': span,
        }
        if span.get('latex'):
            entry['latex'] = span.get('latex')
        math_entries.append(entry)
        runs.append({'type': 'math', 'text': formula_text, 'math': [entry]})
        pos = max(pos, end)
    if pos < len(text):
        runs.append({'type': 'text', 'text': text[pos:]})
    if not math_entries:
        return None
    return {
        'role': 'rich_text',
        'text': text,
        'runs': runs,
        'math': math_entries,
    }


def _item_text(item):
    if isinstance(item, str):
        return item
    if isinstance(item, dict):
        return item.get('text') or item.get('code') or ''
    return ''


def _item_role(item):
    return item.get('role') if isinstance(item, dict) else 'text'


def _is_formula_like_item(item):
    return isinstance(item, dict) and (item.get('role') == 'formula' or item.get('latex') or item.get('xml') or item.get('math'))


def _is_split_formula_fragment(item):
    text = str(_item_text(item) or '').strip()
    if not text or len(text) > 18:
        return False
    if re.search(r'[\u4e00-\u9fff]', text):
        return False
    return bool(re.fullmatch(r'[A-Za-z0-9_\s∆ΔλΛµμ%+\-*/·×÷=<>≤≥≈∈∑().,α-ωΑ-Ω]+', text))


def _is_ratio_variable_fragment(text):
    t = re.sub(r'\s+', '', str(text or ''))
    return bool(re.fullmatch(r'[a-zα-ω][A-Za-z0-9_α-ωΑ-Ω]*', t))


def _latex_identifier(token):
    t = re.sub(r'\s+', '', str(token or ''))
    if not t:
        return ''
    greek = {
        '∆': r'\Delta',
        'Δ': r'\Delta',
        'λ': r'\lambda',
        'Λ': r'\Lambda',
        'μ': r'\mu',
        'α': r'\alpha',
        'β': r'\beta',
        'γ': r'\gamma',
    }
    if t in greek:
        return greek[t]
    m = re.fullmatch(r'([∆ΔλΛμ])([A-Za-z0-9]+)', t)
    if m:
        base, suffix = m.groups()
        command = greek.get(base, base)
        if base in ('∆', 'Δ'):
            return command + ' ' + _latex_identifier(suffix)
        return command + r'_{\mathrm{' + suffix + '}}'
    if re.fullmatch(r'[A-Za-z][A-Za-z0-9]*', t):
        if len(t) == 1:
            return t
        return t[0] + r'_{\mathrm{' + t[1:] + '}}'
    return t


def _latex_math_expr(text, sum_lower=None, sum_upper=None):
    s = str(text or '').strip()
    s = s.replace('−', '-').replace('－', '-').replace('＝', '=')
    s = s.replace('×', r'\times ').replace('·', r'\cdot ').replace('÷', r'\div ')
    s = s.replace('%', r'\%')
    if sum_lower and sum_upper:
        s = s.replace('∑', r'\sum_{' + sum_lower + '}^{' + str(sum_upper) + '}')
    else:
        s = s.replace('∑', r'\sum')
    s = re.sub(r'[∆ΔλΛμ][A-Za-z0-9]*', lambda m: _latex_identifier(m.group(0)), s)
    s = re.sub(r'(?<![\\{])\b[A-Za-z][A-Za-z0-9]*\b', lambda m: _latex_identifier(m.group(0)), s)
    s = re.sub(r'\s+', ' ', s).strip()
    return s


def _repaired_formula_item(text, latex, numbered=False, repair='split_formula_layout'):
    semantic = classify_formula_text(text)
    return {
        'role': 'formula',
        'source': 'repaired_' + repair,
        'text': text,
        'latex': latex,
        'numbered': bool(numbered),
        'formula_semantics': semantic.to_dict(),
    }


def _split_formula_problem_item(text, problem='split_formula_layout'):
    clean = _clean_formula_text(text)
    semantic = classify_formula_text(clean)
    return {
        'role': 'formula_problem',
        'problem': problem,
        'source': 'repair',
        'text': clean,
        'formula_semantics': semantic.to_dict(),
    }


def _repair_split_sum_bounds(items, start):
    upper_text = str(_item_text(items[start]) or '').strip()
    uppers = re.findall(r'\d+', upper_text)
    if not uppers or start + 1 >= len(items):
        return None
    formula = items[start + 1]
    formula_text = str(_item_text(formula) or '').strip()
    if '∑' not in formula_text:
        return None
    clean_formula_text, stale_labels = _split_trailing_formula_labels(formula_text)
    if not stale_labels:
        clean_formula_text = _strip_trailing_formula_labels(formula_text)
    j = start + 2
    lowers = []
    while j < len(items) and len(lowers) < max(1, len(uppers)):
        txt = str(_item_text(items[j]) or '').strip()
        if re.fullmatch(r'[A-Za-z]\s*=\s*\d+', txt):
            lowers.append(re.sub(r'\s+', '', txt))
            j += 1
            continue
        break
    if not lowers:
        return None
    lower = lowers[0]
    upper = uppers[0]
    latex = _latex_math_expr(clean_formula_text, sum_lower=lower, sum_upper=upper)
    repaired = _repaired_formula_item(clean_formula_text, latex, numbered=bool(stale_labels or (formula.get('numbered') if isinstance(formula, dict) else False)), repair='sum_bounds')
    return [repaired], j


def _repair_missing_sum_symbol_bounds(items, start):
    upper_text = str(_item_text(items[start]) or '').strip()
    m_upper = re.fullmatch(r'\d+', upper_text)
    if not m_upper or start + 2 >= len(items):
        return None
    formula = items[start + 1]
    formula_text = str(_item_text(formula) or '').strip()
    lower_text = str(_item_text(items[start + 2]) or '').strip()
    if '∑' in formula_text or not re.fullmatch(r'[A-Za-z]\s*=\s*\d+', lower_text):
        return None
    if not re.search(r'[A-Za-z][A-Za-z0-9]*\s*\(\s*[A-Za-z]\s*\)', formula_text):
        return None
    lower = re.sub(r'\s+', '', lower_text)
    upper = m_upper.group(0)
    text = f'∑_{{{lower}}}^{{{upper}}} {formula_text}'
    latex = _latex_math_expr('∑' + formula_text, sum_lower=lower, sum_upper=upper)
    repaired = _repaired_formula_item(text, latex, numbered=bool(formula.get('numbered') if isinstance(formula, dict) else False), repair='missing_sum_symbol')
    return [repaired], start + 3


def _infer_sum_lower_from_context(out, upper):
    upper = str(upper or '').strip()
    if not upper:
        return None
    for prev in reversed(out[-12:]):
        text = str(_item_text(prev) or '')
        m = re.search(r'∑_\{\s*([A-Za-z]\s*=\s*1)\s*\}\^\{\s*' + re.escape(upper) + r'\s*\}', text)
        if m:
            return re.sub(r'\s+', '', m.group(1))
        if '∑' in text and upper in text:
            lower = _infer_sum_lower(text)
            if lower:
                return lower
    return None


def _repair_labeled_inline_sum_missing_lower(item, out):
    text = str(_item_text(item) or '').strip()
    m = re.match(r'^(?P<label>[A-Za-z][A-Za-z0-9_,]*\s*=\s*)∑\s*(?P<upper>\d+)\s+(?P<expr>.+)$', text)
    if not m:
        return None
    expr, tail = _split_formula_expression_tail(m.group('expr'))
    if not expr:
        return None
    upper = m.group('upper')
    lower = _infer_sum_lower(expr)
    if not lower:
        context_lower = _infer_sum_lower_from_context(out, upper)
        if upper == '24' and context_lower == 't=1':
            lower = context_lower
    if not lower:
        return None
    label = re.sub(r'\s+', '', m.group('label') or '')
    display_text = f'{label}∑_{{{lower}}}^{{{upper}}} {expr}'
    latex = _latex_math_expr(label + '∑' + expr, sum_lower=lower, sum_upper=upper)
    repaired = [_repaired_formula_item(display_text, latex, numbered=bool(item.get('numbered') if isinstance(item, dict) else False), repair='inline_sum_missing_lower')]
    _append_repair_tail(repaired, tail)
    return repaired, None


def _repair_fraction_sum_layout(items, start):
    label = str(_item_text(items[start]) or '').strip()
    if not re.fullmatch(r'[A-Za-z][A-Za-z0-9_]*\s*=', label) or start + 5 >= len(items):
        return None
    numerator = str(_item_text(items[start + 1]) or '').strip()
    denominator = str(_item_text(items[start + 2]) or '').strip()
    upper = str(_item_text(items[start + 3]) or '').strip()
    lower = str(_item_text(items[start + 4]) or '').strip()
    expr = str(_item_text(items[start + 5]) or '').strip()
    if not numerator or not denominator or not re.fullmatch(r'\d+', upper) or not re.fullmatch(r'[A-Za-z]\s*=\s*\d+', lower):
        return None
    if not expr or not re.search(r'[A-Za-z∆ΔλΛμ]', expr):
        return None
    clean_expr = expr.strip()
    if clean_expr.endswith(']') and '[' not in clean_expr:
        clean_expr = '[' + clean_expr
    lower = re.sub(r'\s+', '', lower)
    lhs = label.replace(' ', '')
    frac_latex = r'\frac{' + _latex_math_expr(numerator) + '}{' + _latex_math_expr(denominator) + '}'
    sum_latex = _latex_math_expr('∑' + clean_expr, sum_lower=lower, sum_upper=upper)
    latex = _latex_identifier(lhs.rstrip('=')) + '=' + frac_latex + sum_latex
    display_text = f'{lhs}{numerator}/{denominator} ∑_{{{lower}}}^{{{upper}}} {clean_expr}'
    numbered = False
    next_i = start + 6
    if next_i < len(items) and is_formula_label(str(_item_text(items[next_i]) or '').strip()):
        numbered = True
        next_i += 1
    return [_repaired_formula_item(display_text, latex, numbered=numbered, repair='fraction_sum_layout')], next_i


def _repair_max_sum_layout(items, start):
    op = str(_item_text(items[start]) or '').strip().lower()
    if op not in {'max', 'min'} or start + 4 >= len(items):
        return None
    opt_var = str(_item_text(items[start + 1]) or '').strip()
    upper = str(_item_text(items[start + 2]) or '').strip()
    formula = items[start + 3]
    formula_text = str(_item_text(formula) or '').strip()
    lower = str(_item_text(items[start + 4]) or '').strip()
    if not _is_split_formula_fragment(opt_var) or not re.fullmatch(r'\d+', upper) or not re.fullmatch(r'[A-Za-z]\s*=\s*\d+', lower):
        return None
    if '∑' in formula_text:
        return None
    clean_formula, labels = _split_trailing_formula_labels(formula_text)
    clean_formula = clean_formula or formula_text
    lower = re.sub(r'\s+', '', lower)
    lower_var = lower.split('=', 1)[0]
    summand_match = re.search(r'([A-Za-zα-ωΑ-Ω][A-Za-z0-9_α-ωΑ-Ω]*\s*\(\s*' + re.escape(lower_var) + r'\s*\))', clean_formula)
    if not summand_match:
        return None
    summand = summand_match.group(1)
    display_body = clean_formula[:summand_match.start()] + f'∑_{{{lower}}}^{{{upper}}} {summand}' + clean_formula[summand_match.end():]
    latex_body_src = clean_formula[:summand_match.start()] + '∑' + summand + clean_formula[summand_match.end():]
    latex_body = _latex_math_expr(latex_body_src, sum_lower=lower, sum_upper=upper)
    latex = ('\\max' if op == 'max' else '\\min') + '_{' + _latex_identifier(opt_var) + '} ' + latex_body
    return [_repaired_formula_item(f'{op} {opt_var} {display_body}', latex, numbered=bool(labels or (formula.get('numbered') if isinstance(formula, dict) else False)), repair='max_sum_layout')], start + 5


def _split_percentage_suffix(rhs):
    text = str(rhs or '').strip()
    m = re.match(r'^(.*?)(?:[×x*]\s*100\s*%|\\times\s*100\s*%)$', text)
    if m:
        return m.group(1).strip(), True
    return text, False


def _append_repair_tail(out, tail):
    text = str(tail or '').strip()
    if not text:
        return
    rich = _rich_text_item_from_inline_formula_spans(text)
    out.append(rich or text)


def _split_formula_expression_tail(text):
    s = str(text or '').strip()
    if not s:
        return '', ''
    m = re.search(r'[\u4e00-\u9fff]', s)
    if not m:
        return s, ''
    idx = m.start()
    if idx > 0 and s[idx - 1] in '（(':
        idx -= 1
    return s[:idx].strip().rstrip('，,。;；'), s[idx:].strip()


def _infer_sum_lower(expr):
    text = str(expr or '')
    candidates = []
    candidates.extend(re.findall(r'\b[A-Za-z][A-Za-z0-9]*\s*\(\s*([A-Za-z])\s*\)', text))
    candidates.extend(re.findall(r'\b[A-Za-z][A-Za-z0-9]*\s*_\s*([A-Za-z])\b', text))
    candidates = [c for c in candidates if re.fullmatch(r'[A-Za-z]', c or '')]
    if not candidates:
        return None
    unique = set(candidates)
    if len(unique) != 1:
        return None
    return candidates[0] + '=1'


def _repair_split_sum_prefix(items, start):
    prefix = str(_item_text(items[start]) or '').strip()
    m = re.fullmatch(r'∑\s*(\d+)', prefix)
    if not m or start + 1 >= len(items):
        return None
    formula_text = str(_item_text(items[start + 1]) or '').strip()
    if not formula_text or re.search(r'[\u4e00-\u9fff]', prefix):
        return None
    expr, tail = _split_formula_expression_tail(formula_text)
    if not expr:
        return None
    upper = m.group(1)
    lower = _infer_sum_lower(expr)
    if not lower:
        repaired = [_split_formula_problem_item(prefix + ' ' + expr, problem='split_sum_index_unknown')]
        _append_repair_tail(repaired, tail)
        return repaired, start + 2
    text = f'∑_{{{lower}}}^{{{upper}}} {expr}'
    latex = _latex_math_expr('∑' + expr, sum_lower=lower, sum_upper=upper)
    repaired = [_repaired_formula_item(text, latex, repair='sum_prefix')]
    _append_repair_tail(repaired, tail)
    return repaired, start + 2


def _repair_labeled_sum_continuation(items, start, out=None):
    current_text = str(_item_text(items[start]) or '').strip()
    m = re.match(r'^(?P<prefix>.*?)(?P<label>[A-Za-z][A-Za-z0-9_,]*\s*=\s*)∑\s*(?P<upper>\d+)\s*$', current_text)
    if not m or start + 1 >= len(items):
        return None
    next_text = str(_item_text(items[start + 1]) or '').strip()
    expr, tail = _split_formula_expression_tail(next_text)
    if not expr:
        return None
    prefix = (m.group('prefix') or '').strip()
    label = re.sub(r'\s+', '', m.group('label') or '')
    upper = m.group('upper')
    lower = _infer_sum_lower(expr)
    if not lower:
        context_lower = _infer_sum_lower_from_context(out or [], upper)
        if upper == '24' and context_lower == 't=1':
            lower = context_lower
    repaired = []
    if prefix:
        repaired.append(prefix)
    if not lower:
        repaired.append(_split_formula_problem_item(label + '∑' + upper + ' ' + expr, problem='split_sum_index_unknown'))
        _append_repair_tail(repaired, tail)
        return repaired, start + 2
    latex = _latex_math_expr(label + '∑' + expr, sum_lower=lower, sum_upper=upper)
    text = f'{label}∑_{{{lower}}}^{{{upper}}} {expr}'
    repaired.append(_repaired_formula_item(text, latex, repair='labeled_sum_continuation'))
    _append_repair_tail(repaired, tail)
    return repaired, start + 2


def _repair_split_ratio_cluster(items, start):
    first_var = str(_item_text(items[start]) or '').strip()
    if not _is_ratio_variable_fragment(first_var) or start + 1 >= len(items):
        return None
    first_formula = items[start + 1]
    if not _is_formula_like_item(first_formula):
        return None
    first_rhs = str(_item_text(first_formula) or '').strip()
    if not first_rhs.startswith('=') or '100' not in first_rhs:
        return None

    formulas = [first_formula]
    variables = [first_var]
    denominators = []
    current_formula_idx = start + 1
    j = current_formula_idx + 1
    while True:
        fragments = []
        while j < len(items) and _is_split_formula_fragment(items[j]) and not _is_formula_like_item(items[j]):
            fragments.append(str(_item_text(items[j]) or '').strip())
            j += 1
        if j < len(items) and _is_formula_like_item(items[j]) and str(_item_text(items[j]) or '').strip().startswith('=') and '100' in str(_item_text(items[j]) or ''):
            next_var_pos = None
            for pos, frag in enumerate(fragments):
                if _is_ratio_variable_fragment(frag):
                    next_var_pos = pos
                    break
            if next_var_pos is None:
                return None
            denom_parts = fragments[:next_var_pos] + fragments[next_var_pos + 1:]
            if not denom_parts:
                return None
            denominators.append(''.join(denom_parts))
            variables.append(fragments[next_var_pos])
            formulas.append(items[j])
            current_formula_idx = j
            j += 1
            continue
        if fragments:
            denominators.append(''.join(fragments))
        break

    if len(formulas) < 2 or len(denominators) != len(formulas):
        return None

    repaired = []
    for var, formula, denom in zip(variables, formulas, denominators):
        rhs = str(_item_text(formula) or '').strip().lstrip('=').strip()
        rhs = _strip_trailing_formula_labels(rhs)
        numerator, has_percent = _split_percentage_suffix(rhs)
        lhs_latex = _latex_identifier(var)
        denom_latex = _latex_identifier(denom)
        numerator_latex = _latex_math_expr(numerator)
        latex = lhs_latex + r'=\frac{' + numerator_latex + '}{' + denom_latex + '}'
        if has_percent:
            latex += r'\times100\%'
        text = f'{var}=({numerator})/({denom})' + ('×100%' if has_percent else '')
        repaired.append(_repaired_formula_item(text, latex, numbered=bool(formula.get('numbered') if isinstance(formula, dict) else False), repair='ratio_cluster'))
    return repaired, j


def repair_split_formula_layouts(paragraphs):
    """Repair formula layouts that were already fragmented in a source DOCX."""
    out = []
    i = 0
    while i < len(paragraphs):
        if _is_split_formula_fragment(paragraphs[i]):
            max_sum_repair = _repair_max_sum_layout(paragraphs, i)
            if max_sum_repair:
                repaired, next_i = max_sum_repair
                out.extend(repaired)
                i = next_i
                continue
            fraction_sum_repair = _repair_fraction_sum_layout(paragraphs, i)
            if fraction_sum_repair:
                repaired, next_i = fraction_sum_repair
                out.extend(repaired)
                i = next_i
                continue
            sum_repair = _repair_split_sum_bounds(paragraphs, i)
            if sum_repair:
                repaired, next_i = sum_repair
                out.extend(repaired)
                i = next_i
                continue
            missing_sum_repair = _repair_missing_sum_symbol_bounds(paragraphs, i)
            if missing_sum_repair:
                repaired, next_i = missing_sum_repair
                out.extend(repaired)
                i = next_i
                continue
            sum_prefix_repair = _repair_split_sum_prefix(paragraphs, i)
            if sum_prefix_repair:
                repaired, next_i = sum_prefix_repair
                out.extend(repaired)
                i = next_i
                continue
            ratio_repair = _repair_split_ratio_cluster(paragraphs, i)
            if ratio_repair:
                if out and re.fullmatch(r'[A-Za-z]\s*=\s*\d+', str(_item_text(out[-1]) or '').strip()):
                    out.pop()
                repaired, next_i = ratio_repair
                out.extend(repaired)
                i = next_i
                continue
        labeled_sum_repair = _repair_labeled_sum_continuation(paragraphs, i, out)
        if labeled_sum_repair:
            repaired, next_i = labeled_sum_repair
            out.extend(repaired)
            i = next_i
            continue
        inline_sum_repair = _repair_labeled_inline_sum_missing_lower(paragraphs[i], out)
        if inline_sum_repair:
            repaired, _ = inline_sum_repair
            out.extend(repaired)
            i += 1
            continue
        out.append(paragraphs[i])
        i += 1
    return out

def _append_text_or_code(section, text, in_appendix=False):
    """Append semantic blocks while preserving captions, code and inline citations."""
    if not text:
        return
    text = _clean_text_artifacts(text)
    if not text:
        return
    semantic = classify_formula_text(text)
    if _is_figure_caption(text):
        section['paragraphs'].append({'role': 'figure_caption', 'text': _normalize_caption_spacing(text)})
    elif _is_table_caption(text):
        section['paragraphs'].append({'role': 'table_caption', 'text': _normalize_caption_spacing(text)})
    elif re.match(r'^\s*\$\$.+\$\$\s*$', text, re.S):
        section['paragraphs'].append(_formula_item_from_text(text))
    else:
        rich_item = _rich_text_item_from_inline_formula_spans(text)
        rich_has_text = bool(rich_item and any(
            r.get('type') == 'text' and str(r.get('text') or '').strip(' \t\r\n，,。.;；:：()（）')
            for r in rich_item.get('runs') or []
        ))
        if rich_item and rich_has_text:
            section['paragraphs'].append(rich_item)
            return
        if is_formula_problem_text(text):
            section['paragraphs'].append(_formula_problem_item_from_text(text))
        elif _looks_like_formula_text(text):
            section['paragraphs'].append(_formula_item_from_text(text))
        elif in_appendix and (_looks_like_code_line(text) or '\n' in text):
            section['paragraphs'].append({'role': 'code', 'code': _clean_code_text(text)})
        elif rich_item:
            section['paragraphs'].append(rich_item)
        elif _omml_text_looks_like_body(text):
            section['paragraphs'].append(text)
        else:
            section['paragraphs'].append(text)

def extract(docx_path, output_dir='Inputs'):
    """Extract content from a content docx into structured JSON + copy images."""
    doc = Document(docx_path)
    base = os.path.splitext(os.path.basename(docx_path))[0]

    # Setup output dirs.  Recreate figures for each extraction so repeated
    # verification passes do not accumulate stale/duplicated files.
    content_dir = os.path.join(output_dir, base)
    fig_dir = os.path.join(content_dir, 'figures')
    shutil.rmtree(fig_dir, ignore_errors=True)
    os.makedirs(fig_dir, exist_ok=True)
    image_registry = ImageRegistry(fig_dir, f'{base}_img')

    content = {
        '_meta': {
            'source': os.path.basename(docx_path),
            'sha256': hashlib.sha256(open(docx_path, 'rb').read()).hexdigest()[:16],
            'paragraphs': len(doc.paragraphs),
            'tables_count': len(doc.tables),
        },
        'title_info': {},
        'sections': [],
        'references': [],
    }
    content['_meta']['source_placeholders'] = _placeholder_samples(doc.paragraphs)

    # ── Extract title info ──
    # Find the largest-text paragraph in the first 20 paragraphs
    text_start = 0
    best_title = ('', 0, 0)  # (text, size, index)
    for i, p in enumerate(doc.paragraphs[:30]):
        txt = p.text.strip()
        if not txt or len(txt) < 10:
            continue
        labeled_title = _extract_labeled_title(txt)
        if labeled_title:
            content['title_info']['title_cn'] = labeled_title
            content.setdefault('cover_info', {})['paper_title'] = labeled_title
            text_start = max(text_start, i + 1)
            continue
        if _is_unfilled_placeholder_text(txt):
            continue
        if txt.startswith('（') and txt.endswith('）'):
            continue
        # Get max font size from runs
        max_sz = 0
        for r in p.runs:
            if r.font.size:
                max_sz = max(max_sz, r.font.size.pt)
        if max_sz >= 14 and max_sz > best_title[1]:
            clean = txt.split('（')[0].strip().replace('\n', ' ').replace('\r', '')
            if clean and not clean.startswith('年') and not clean.startswith('本科'):
                best_title = (clean, max_sz, i)

    if best_title[0]:
        content['title_info']['title_cn'] = best_title[0]
        text_start = best_title[2] + 1
        for j, p in enumerate(doc.paragraphs[text_start:min(text_start + 6, len(doc.paragraphs))], start=text_start):
            txt = _clean_text_artifacts(p.text)
            if not txt:
                continue
            if re.match(r'(?i)^(abstract|key\s*words?)\b', txt) or txt.startswith(('摘要', '关键词')):
                break
            if _classify_section_role(txt, 1) in {'references', 'acknowledgement', 'appendix'}:
                continue
            if _ascii_alpha_ratio(txt) > 0.55 and not re.match(r'^\d+(?:\.\d+)*\s+', txt):
                content['title_info']['title_en'] = txt
                text_start = j + 1
                break

    # ── Extract cover info from content docx tables ──
    cover_info = {}
    _COVER_LABEL_MAP = {
        '学校编码': 'school_code', '学位编码': 'degree_code', '论文题目': 'paper_title',
        '学生姓名': 'student_name', '学号': 'student_id', '学    号': 'student_id',
        '所属学院': 'college', '专业班级': 'class_name',
        '指导老师': 'advisor', '指导教师': 'advisor',
    }
    for table in doc.tables[:5]:
        for row in table.rows:
            if len(row.cells) >= 2:
                label = row.cells[0].text.strip()
                value = row.cells[1].text.strip()
                if label and value:
                    for kw, key in _COVER_LABEL_MAP.items():
                        if kw in label:
                            content['cover_info'] = content.get('cover_info', {})
                            content['cover_info'][key] = value
                            cover_info[key] = value
                            break
    # Cover table is the most reliable source for the paper title.
    # The old heuristic sometimes picked declaration/footer text with large font.
    if cover_info.get('paper_title'):
        content['title_info']['title_cn'] = cover_info['paper_title']

    # ── Parse sections ──
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    current_section = {'heading': '正文', 'level': 1, 'role': 'body', 'paragraphs': [], 'images': []}
    sections = [current_section]
    ref_section = None
    collected_references = []

    # Skip body elements before text_start
    _body_children = list(doc.element.body)
    _p_idx = 0  # paragraph index in body children
    _started = False
    _source_toc_skip_remaining = 0
    _source_toc_skipped = 0
    for _child in _body_children:
        _tag = _child.tag.split('}')[-1]

        if _tag == 'p':
            if _p_idx < text_start:
                _p_idx += 1
                continue
            _started = True
            p = doc.paragraphs[_p_idx]
            _p_idx += 1
            text = p.text.strip()

            level = detect_heading_level(p)

            if _source_toc_skip_remaining > 0:
                _source_toc_skip_remaining -= 1
                _source_toc_skipped += 1
                continue

            if _is_source_toc_title(text):
                skip_after_title = _source_toc_skip_count_after_title(doc.paragraphs, _p_idx - 1)
                if skip_after_title:
                    _source_toc_skip_remaining = skip_after_title
                    _source_toc_skipped += 1
                    continue

            if re.match(r'(?i)^references?\b', text) or text.startswith('参考文献'):
                ref_section = {'heading': text, 'entries': []}
                continue

            # Back matter after references (致谢/附录) must not be swallowed
            # by the reference collector. Treat it as normal sections again.
            if ref_section is not None and _is_backmatter_heading(text):
                if ref_section.get('entries'):
                    collected_references.extend(ref_section['entries'])
                _h = _normalize_heading_spacing(re.split(r'[:：]', text, maxsplit=1)[0].strip())
                _level = level or 1
                current_section = {'heading': _h, 'level': _level, 'role': _classify_section_role(_h, _level), 'paragraphs': [], 'images': []}
                sections.append(current_section)
                ref_section = None
                continue

            if ref_section is not None:
                clean_ref_text = _clean_text_artifacts(text)
                if clean_ref_text:
                    ref_section['entries'].append(clean_ref_text)
            elif level > 0:
                clean_heading = text.split('（')[0].strip()
                if not clean_heading:
                    continue
                m = re.match(r'(?i)^(Abstract\s*:?|Key\s*words?\s*:?|摘要\s*[：:]|关键词\s*[：:])\s*', text)
                if m:
                    heading_part = m.group(1).strip()
                    body_part = text[m.end():].strip()
                    body_part = re.sub(r'^[（(][^）)]*[）)]\s*', '', body_part)
                else:
                    heading_part = clean_heading
                    body_part = ''
                heading_part = _normalize_heading_spacing(heading_part)
                role = _classify_section_role(heading_part, level)
                current_section = {'heading': heading_part, 'level': level, 'role': role, 'paragraphs': [], 'images': []}
                if role == 'en_abstract':
                    current_section['page_break_before'] = True
                sections.append(current_section)
                if body_part:
                    _append_text_or_code(current_section, body_part, in_appendix=False)
            else:
                # Preserve exact OOXML run order within the paragraph.  This is
                # crucial when Word stores explanatory text, a drawing, and a
                # caption in nearby runs/paragraphs.
                stream_items = paragraph_stream_items(p, image_registry)
                if not stream_items and text:
                    stream_items = [{'role': 'text', 'text': text}]
                rich_runs = []
                in_appendix = bool(re.search(r'(附\s*录|配置|命令|代码)', current_section.get('heading','')))

                def _flush_rich_runs():
                    nonlocal rich_runs
                    _append_stream_run_group(current_section, rich_runs, in_appendix=in_appendix)
                    rich_runs = []

                for _it in stream_items:
                    if _it.get('role') == 'image':
                        _flush_rich_runs()
                        current_section['images'].append(_it.get('image'))
                        current_section['paragraphs'].append(_it)
                    elif _it.get('role') == 'math_inline':
                        rich_runs.append({'type': 'math', 'text': _it.get('text') or '', 'math': _it.get('math') or []})
                    elif _it.get('role') == 'formula':
                        _flush_rich_runs()
                        current_section['paragraphs'].append(_it)
                    elif _it.get('role') == 'text':
                        txt = _it.get('text') or ''
                        if txt:
                            rich_runs.append({'type': 'text', 'text': txt})
                _flush_rich_runs()

        elif _tag == 'tbl' and _started:
            # Body table — preserve paragraph breaks inside each cell.
            _rows = _extract_table_rows_from_ooxml(_child)
            _table_images = _image_items_from_ooxml(_child, doc.part.rels, image_registry, location='table_cell')
            if _rows:
                if ref_section is not None:
                    if _table_rows_look_like_code(_rows):
                        ref_section['entries'].append({'role': 'code', 'code': _code_text_from_table_rows(_rows), 'table_rows': _rows})
                    else:
                        ref_section['entries'].append({'role': 'table', 'table_rows': _rows})
                elif _table_rows_look_like_code(_rows):
                    current_section['paragraphs'].append({'role': 'code', 'code': _code_text_from_table_rows(_rows), 'table_rows': _rows})
                else:
                    current_section['paragraphs'].append({'role': 'table', 'table_rows': _rows})
            if _table_images and ref_section is None:
                for _img in _table_images:
                    current_section['images'].append(_img.get('image'))
                    current_section['paragraphs'].append(_img)

    if ref_section and ref_section['entries']:
        collected_references.extend(ref_section['entries'])
        ref_section = None
    if collected_references:
        content['references'] = collected_references

    # Filter empty placeholder sections, but KEEP structural headings.
    # A level-1 chapter can legitimately have no direct paragraphs because it is
    # followed immediately by 1.1/1.2 subsections. Dropping it breaks TOC and body
    # structure, so keep every non-placeholder heading even when empty.
    content['sections'] = []
    for s in sections:
        # If real headings were detected, the initial placeholder section is
        # almost always cover/declaration/title residue. Keep it only when it
        # is the sole section in an unstructured document.
        if s['heading'] == '正文' and len(sections) > 1:
            continue
        if s['paragraphs'] or s['images'] or (s.get('heading') and s.get('heading') != '正文'):
            content['sections'].append(s)
    if collected_references:
        content['references'] = collected_references

    # Mark the first real body chapter so renderers can start it on a new page.
    _front_roles = {'cn_abstract', 'cn_keywords', 'en_abstract', 'en_keywords'}
    for _s in content['sections']:
        if _s.get('role') not in _front_roles and not _s.get('page_break_before'):
            _s.setdefault('page_break_before', True)
            break

    # Pair images with following figure captions after all sections are built.
    for _s in content['sections']:
        _s['paragraphs'] = repair_split_formula_layouts(_s.get('paragraphs') or [])
        _s['paragraphs'] = pair_figure_blocks(_s.get('paragraphs') or [])

    # Count saved images without running a second extraction pass.
    # Re-extracting here used to create duplicate filenames and made figure
    # captions drift away from their intended images.
    content['_meta']['images_extracted'] = len([
        f for f in os.listdir(fig_dir)
        if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tif', '.tiff'))
    ])
    content['_meta']['images_dir'] = os.path.abspath(fig_dir)
    content['_meta']['image_extract_failures'] = image_registry.failures
    content['_meta']['non_body_images'] = _non_body_image_entries(doc)
    if _source_toc_skipped:
        content['_meta']['source_toc_skipped_paragraphs'] = _source_toc_skipped

    return content


if __name__ == '__main__':
    import sys
    path = sys.argv[1] if len(sys.argv) > 1 else 'Templates/模版.docx'
    content = extract(path)
    json_path = os.path.splitext(path)[0] + '_content.json'
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(content, f, ensure_ascii=False, indent=2)
    print(f'Content JSON → {json_path}')
    print(f'Sections: {len(content["sections"])}  References: {len(content["references"])}  Images: {content["_meta"]["images_extracted"]}')
