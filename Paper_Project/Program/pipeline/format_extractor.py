"""
format_extractor.py — OOXML-direct format extraction with style resolution.
Solves python-docx API limitation: no more None values from style inheritance.
"""
import json, os, hashlib, re
from io import BytesIO

try:
    from PIL import Image
except Exception:
    Image = None

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ALIGN_MAP = {0: 'LEFT', 1: 'CENTER', 2: 'RIGHT', 3: 'JUSTIFY', None: 'DEFAULT'}

def _tag(el):
    return el.tag.split('}')[-1] if '}' in el.tag else el.tag

def _val(el, attr='w:val', default=None):
    return el.get(qn(attr), default)

def _pt(half_pts_str):
    """Convert half-points string to float. '28' -> 14.0"""
    try: return int(half_pts_str) / 2.0
    except: return None

def _emu_to_pt(emu):
    try: return round(int(emu) / 12700, 1)
    except: return None


def _twips_to_pt(v):
    try: return round(int(v) / 20.0, 2)
    except Exception: return None

def _twips_to_cm(v):
    try: return round(int(v) / 567.0, 2)
    except Exception: return None

def _paragraph_metrics(p_elem):
    """Extract paragraph metrics directly from OOXML.

    Returns both Word-style raw values and python-friendly normalized values.
    The important part is preserving fixed line spacing: w:line="560"
    with w:lineRule="exact" means exactly 28 pt, not a 2.33 multiple.
    """
    info = {
        'alignment': 'DEFAULT', 'line_spacing_val': None, 'line_spacing_rule': None,
        'line_spacing_fixed_pt': None, 'space_before_pt': None, 'space_after_pt': None,
        'first_indent_cm': 0,
    }
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        return info
    jc = pPr.find(qn('w:jc'))
    if jc is not None:
        v = _val(jc)
        info['alignment'] = {'left':'LEFT','center':'CENTER','right':'RIGHT','both':'JUSTIFY','distribute':'DISTRIBUTE'}.get(v, 'DEFAULT')
    spacing = pPr.find(qn('w:spacing'))
    if spacing is not None:
        line = spacing.get(qn('w:line'))
        rule = spacing.get(qn('w:lineRule'))
        info['line_spacing_rule'] = rule
        if line:
            try:
                n = int(line)
                if rule in ('exact', 'atLeast'):
                    info['line_spacing_fixed_pt'] = round(n / 20.0, 2)
                    info['line_spacing_val'] = info['line_spacing_fixed_pt']
                else:
                    info['line_spacing_val'] = round(n / 240.0, 4)
            except Exception:
                pass
        info['space_before_pt'] = _twips_to_pt(spacing.get(qn('w:before')))
        info['space_after_pt'] = _twips_to_pt(spacing.get(qn('w:after')))
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        fi = ind.get(qn('w:firstLine'))
        if fi:
            info['first_indent_cm'] = _twips_to_cm(fi) or 0
    return info


def _first_real_run(p):
    for r in p.get('runs', []):
        if r.get('text', '').strip() or r.get('size_pt'):
            return r
    return (p.get('runs') or [{}])[0] if p.get('runs') else {}


def _profile_from_paragraph(p):
    r = _first_real_run(p)
    return {
        'font': r.get('font') or '宋体',
        'size': r.get('size_pt') or 12,
        'bold': bool(r.get('bold', False)),
        'italic': bool(r.get('italic', False)),
        'align': p.get('alignment') or p.get('align') or 'LEFT',
        'line_spacing_val': p.get('line_spacing_val') or p.get('ls'),
        'line_spacing_rule': p.get('line_spacing_rule'),
        'line_spacing_fixed_pt': p.get('line_spacing_fixed_pt'),
        'space_before_pt': p.get('space_before_pt'),
        'space_after_pt': p.get('space_after_pt'),
        'first_indent_cm': p.get('first_indent_cm') if p.get('first_indent_cm') is not None else p.get('indent', 0),
    }


def _build_style_profiles(fmt):
    """Infer semantic style profiles from the template examples/instructions.

    This is intentionally role-based instead of school-specific.  The generator
    consumes these roles and no longer hardcodes CENTER/黑体/16pt etc.
    """
    profiles = {}
    paras = fmt.get('paragraphs', [])

    def put(role, p):
        if role not in profiles and p:
            profiles[role] = _profile_from_paragraph(p)

    for p in paras:
        txt = (p.get('text') or '').strip()
        if not txt:
            continue
        no_space = txt.replace(' ', '')
        # Front matter roles
        if '论文' in txt and '题目' in txt and ('居中' in txt or p.get('style') == '论文题目'):
            put('cn_title', p)
        if no_space.startswith('摘要') or no_space.startswith('摘要（') or no_space.startswith('摘要(') or no_space.startswith('摘要'):
            if len(txt) < 30:
                put('cn_abstract_heading', p)
        if txt.startswith('摘要是') or ('中文摘要300' in txt and len(txt) > 50):
            put('cn_abstract_body', p)
        if txt.startswith('关键词'):
            put('cn_keywords', p)
        if '英文题目' in txt and ('Times' in txt or 'Roman' in txt):
            put('en_title', p)
        if txt.upper().startswith('ABSTRACT') and len(txt) < 40:
            put('en_abstract_heading', p)
        if len(txt) > 80 and sum(1 for c in txt[:120] if c.isascii() and c.isalpha()) > 50:
            put('en_abstract_body', p)
        if txt.upper().startswith('KEY WORD') or txt.upper().startswith('KEYWORDS'):
            put('en_keywords', p)
        if no_space in ('目录', '目 录'.replace(' ', '')) or no_space.startswith('目录'):
            put('toc_title', p)
        # Body heading examples
        if ('一级标题' in txt or '第1章' in txt) and len(txt) < 80:
            put('h1', p)
        if ('二级标题' in txt or re.match(r'^1\.1\s+', txt)) and len(txt) < 80:
            put('h2', p)
        if ('三级标题' in txt or re.match(r'^1\.1\.1\s+', txt)) and len(txt) < 80:
            put('h3', p)

    # Prefer actual body heading examples over TOC entries or textual notes.
    def _score_heading_candidate(p, level):
        txt = (p.get('text') or '').strip()
        r = _first_real_run(p)
        font = r.get('font') or ''
        size = r.get('size_pt') or 0
        score = 0
        if font and font not in ('Arial', 'Times New Roman', 'Calibri'):
            score += 10
        if level == 1 and re.match(r'^第[一二三四五六七八九十\d]+章\s+', txt):
            score += 8
        if level == 2 and re.match(r'^\d+\.\d+\s+', txt):
            score += 8
        if level == 3 and re.match(r'^\d+\.\d+\.\d+\s+', txt):
            score += 8
        if size:
            score += min(float(size), 20) / 2
        if '标题' in txt and '（' in txt:
            score -= 4  # instruction line, not rendered sample
        if '目录' in txt:
            score -= 8
        return score

    for _role, _level, _pat in [
        ('h1', 1, r'^第[一二三四五六七八九十\d]+章\s+'),
        ('h2', 2, r'^\d+\.\d+\s+'),
        ('h3', 3, r'^\d+\.\d+\.\d+\s+'),
    ]:
        _cands = [p for p in paras if re.match(_pat, (p.get('text') or '').strip()) and len((p.get('text') or '').strip()) < 80]
        if _cands:
            _best = max(_cands, key=lambda p: _score_heading_candidate(p, _level))
            profiles[_role] = _profile_from_paragraph(_best)
            profiles[_role]['bold'] = bool(profiles[_role].get('bold')) or True
            profiles[_role]['first_indent_cm'] = 0 if _level == 1 else profiles[_role].get('first_indent_cm', 0)

    # Body = most common long CJK paragraph that is not declaration/format note
    candidates = []
    for p in paras:
        txt = (p.get('text') or '').strip()
        if len(txt) < 80:
            continue
        if any(k in txt[:60] for k in ['本人郑重声明', '本人在导师', '格式', '要求', '行距', '字号', '页眉']):
            continue
        if any('\u4e00' <= c <= '\u9fff' for c in txt[:120]):
            candidates.append(p)
    if candidates:
        put('body', candidates[0])

    # Fallbacks based on available roles
    body = profiles.get('body') or {'font':'宋体','size':12,'align':'JUSTIFY','line_spacing_fixed_pt':28,'first_indent_cm':0.74}
    profiles.setdefault('body', body)
    profiles.setdefault('h1', {**body, 'font':'黑体', 'size':16, 'bold':True, 'align':'CENTER', 'first_indent_cm':0})
    profiles.setdefault('h2', {**body, 'font':'黑体', 'size':14, 'bold':True, 'align':'LEFT', 'first_indent_cm':0})
    profiles.setdefault('h3', {**body, 'font':'黑体', 'size':12, 'bold':True, 'align':'LEFT', 'first_indent_cm':0})
    profiles.setdefault('cn_title', profiles['h1'])
    profiles.setdefault('cn_abstract_heading', profiles['h1'])
    profiles.setdefault('cn_abstract_body', body)
    profiles.setdefault('cn_keywords', {**body, 'first_indent_cm':0})
    profiles.setdefault('en_title', {**profiles['h1'], 'font':'Times New Roman'})
    profiles.setdefault('en_abstract_heading', {**profiles['h1'], 'font':'Times New Roman', 'size':16, 'bold':True, 'align':'CENTER', 'first_indent_cm':0})
    profiles.setdefault('en_abstract_body', {**body, 'font':'Times New Roman', 'line_spacing_fixed_pt':None, 'line_spacing_val':1.5, 'first_indent_cm':0.9})
    profiles.setdefault('figure_caption', {**body, 'font':'宋体', 'size':10.5, 'align':'CENTER', 'first_indent_cm':0, 'space_before_pt':6, 'space_after_pt':6, 'line_spacing_fixed_pt':28})
    profiles.setdefault('table_caption', {**profiles['figure_caption']})
    profiles.setdefault('code', {**body, 'font':'Consolas', 'size':10.5, 'align':'LEFT', 'first_indent_cm':0, 'line_spacing_fixed_pt':None, 'line_spacing_val':1.0})
    profiles.setdefault('reference', {**body, 'font':'宋体', 'size':12, 'align':'JUSTIFY', 'first_indent_cm':0, 'line_spacing_fixed_pt':28, 'space_before_pt':6, 'space_after_pt':6})
    profiles.setdefault('en_keywords', {**profiles['en_abstract_body'], 'bold':True, 'first_indent_cm':0})
    profiles.setdefault('toc_title', profiles['h1'])

    # Sanitize inherited python-docx Length objects that can appear as huge
    # numbers in style-derived line_spacing. Headings without explicit line
    # spacing should inherit the body line-spacing role.
    for _role in ('h1', 'h2', 'h3', 'cn_title', 'cn_abstract_heading', 'toc_title'):
        _p = profiles.get(_role, {})
        try:
            _ls = float(_p.get('line_spacing_val') or 0)
        except Exception:
            _ls = 0
        if _ls > 10 and not _p.get('line_spacing_fixed_pt'):
            _p['line_spacing_val'] = body.get('line_spacing_val')
            _p['line_spacing_fixed_pt'] = body.get('line_spacing_fixed_pt')
            _p['line_spacing_rule'] = body.get('line_spacing_rule')
    return profiles


class StyleResolver:
    """Resolve formatting by walking style inheritance tree in styles.xml."""

    def __init__(self, doc):
        self.styles = {}  # style_id -> {font, size, bold, ...}
        self._load_styles(doc)

    def _load_styles(self, doc):
        for style in doc.styles:
            sid = style.style_id
            entry = {}
            f = getattr(style, 'font', None)
            if f is not None:
                if f.name: entry['font'] = f.name
                if f.size:  entry['size'] = f.size.pt
                if f.bold is not None: entry['bold'] = f.bold
                if f.italic is not None: entry['italic'] = f.italic
            # Paragraph styles have paragraph_format; character styles don't
            try:
                pf = style.paragraph_format
                if pf.line_spacing: entry['ls'] = pf.line_spacing
                if pf.alignment is not None: entry['align'] = ALIGN_MAP.get(pf.alignment, 'DEFAULT')
                if pf.first_line_indent: entry['indent'] = pf.first_line_indent.cm
            except AttributeError:
                pass  # character style — no paragraph formatting
            try:
                entry['base'] = style.base_style.style_id if style.base_style else None
            except (AttributeError, ValueError):
                entry['base'] = None
            self.styles[sid] = entry

    def resolve(self, p_elem, r_elem):
        """Return fully resolved formatting for a run."""
        result = {'font': None, 'size': None, 'bold': False, 'italic': False,
                  'ls': None, 'align': 'DEFAULT', 'indent': 0}

        # 1. Direct formatting on the run (highest priority)
        rPr = r_elem.find(qn('w:rPr'))
        if rPr is not None:
            for child in rPr:
                ct = _tag(child)
                if ct == 'rFonts':
                    result['font'] = _val(child, 'w:ascii') or _val(child, 'w:hAnsi')
                elif ct == 'sz':
                    result['size'] = _pt(_val(child))
                elif ct == 'b':
                    result['bold'] = True
                elif ct == 'i':
                    result['italic'] = True
                elif ct == 'color':
                    result['color'] = _val(child)

        # 2. Paragraph properties
        pPr = p_elem.find(qn('w:pPr'))
        style_id = None
        if pPr is not None:
            for child in pPr:
                ct = _tag(child)
                if ct == 'pStyle':
                    style_id = _val(child)
                elif ct == 'jc':
                    v = _val(child)
                    align_map = {'left': 'LEFT', 'center': 'CENTER', 'right': 'RIGHT', 'both': 'JUSTIFY'}
                    result['align'] = align_map.get(v, 'DEFAULT')
                elif ct == 'spacing':
                    line = child.get(qn('w:line'))
                    if line: result['ls'] = int(line) / 240.0
                elif ct == 'ind':
                    fi = child.get(qn('w:firstLine'))
                    if fi: result['indent'] = round(int(fi) / 567.0, 1)  # twips -> cm

        # 3. Resolve from style tree
        self._resolve_style(style_id, result)

        return result

    def _resolve_style(self, sid, result):
        """Walk style inheritance chain to fill missing values."""
        visited = set()
        while sid and sid not in visited:
            visited.add(sid)
            s = self.styles.get(sid, {})
            if result['font'] is None: result['font'] = s.get('font')
            if result['size'] is None: result['size'] = s.get('size')
            if result['align'] == 'DEFAULT': result['align'] = s.get('align', 'DEFAULT')
            if result['ls'] is None: result['ls'] = s.get('ls')
            sid = s.get('base')
        # Final fallbacks
        if result['font'] is None: result['font'] = 'Times New Roman'
        if result['size'] is None: result['size'] = 12.0
        if result['ls'] is None: result['ls'] = 1.15
        if result['align'] == 'DEFAULT': result['align'] = 'LEFT'


def extract(docx_path):
    """Extract ALL formatting with style resolution. Returns (format_dict, md_text)."""
    doc = Document(docx_path)
    resolver = StyleResolver(doc)

    fmt = {
        '_meta': {
            'source': os.path.basename(docx_path),
            'sha256': hashlib.sha256(open(docx_path, 'rb').read()).hexdigest()[:16],
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'sections': len(doc.sections),
        },
        'sections': [],
        'paragraphs': [],
        'tables': [],
    }

    md_lines = []
    md_lines.append(f'# 模版格式提取 — {os.path.basename(docx_path)}\n')
    md_lines.append(f'**段落**: {len(doc.paragraphs)} | **表格**: {len(doc.tables)} | **节**: {len(doc.sections)}\n')

    # ── Sections ──
    md_lines.append('## 页面设置\n')
    for i, sec in enumerate(doc.sections):
        si = {
            'index': i,
            'page_width_cm': round(sec.page_width.cm, 1),
            'page_height_cm': round(sec.page_height.cm, 1),
            'margin_top_cm': round(sec.top_margin.cm, 1),
            'margin_bottom_cm': round(sec.bottom_margin.cm, 1),
            'margin_left_cm': round(sec.left_margin.cm, 1),
            'margin_right_cm': round(sec.right_margin.cm, 1),
            'diff_first_page': sec.different_first_page_header_footer,
            'header': [], 'footer': [],
        }
        if sec.header:
            for p in sec.header.paragraphs:
                runs = []
                for r in p.runs:
                    info = resolver.resolve(p._element, r._element)
                    runs.append({'text': r.text, 'font': info['font'], 'size_pt': info['size'],
                                 'bold': info['bold'], 'italic': info['italic']})
                si['header'].append({'text': p.text, 'alignment': ALIGN_MAP.get(p.alignment, 'CENTER'), 'runs': runs})
        if sec.footer:
            for p in sec.footer.paragraphs:
                si['footer'].append({'text': p.text, 'alignment': ALIGN_MAP.get(p.alignment, 'CENTER')})
        fmt['sections'].append(si)

        md_lines.append(f'**节{i}**: {si["page_width_cm"]}x{si["page_height_cm"]}cm')
        md_lines.append(f'  边距: T{si["margin_top_cm"]} B{si["margin_bottom_cm"]} L{si["margin_left_cm"]} R{si["margin_right_cm"]}')
        for h in si['header']:
            md_lines.append(f'  页眉: {h["alignment"]} | {h["text"][:100]}')
        for f in si['footer']:
            md_lines.append(f'  页脚: {f["alignment"]} | {f["text"][:100]}')
        md_lines.append('')

    # ── Paragraphs ──
    md_lines.append('## 正文格式\n')
    for idx, p in enumerate(doc.paragraphs):
        pf = p.paragraph_format
        pinfo = {
            'index': idx,
            'style': p.style.name if p.style else 'Normal',
            'text': p.text,
            'runs': [],
            'has_page_break': False,
        }
        for r in p.runs:
            info = resolver.resolve(p._element, r._element)
            has_pb = 'w:br w:type="page"' in r._element.xml or "w:br w:type='page'" in r._element.xml
            if has_pb: pinfo['has_page_break'] = True
            # Use resolved values (never None)
            pinfo['runs'].append({
                'text': r.text,
                'font': info['font'],
                'size_pt': info['size'],
                'bold': info['bold'],
                'italic': info['italic'],
            })
            # Carry paragraph-level formatting on first run
            if 'align' not in pinfo:
                pinfo['align'] = info['align']
                pinfo['ls'] = info['ls']
                pinfo['indent'] = info['indent']

        # Add paragraph-level aliases consumed by the generator.
        # Older output used only align/ls/indent; newer generator uses the
        # clearer alignment/line_spacing_val/first_indent_cm names.
        m = _paragraph_metrics(p._element)
        if 'align' not in pinfo:
            pinfo['align'] = m['alignment']
            pinfo['ls'] = m['line_spacing_val'] or 1.15
            pinfo['indent'] = m['first_indent_cm'] or 0
        pinfo['alignment'] = m['alignment'] if m['alignment'] != 'DEFAULT' else pinfo.get('align', 'DEFAULT')
        pinfo['line_spacing_val'] = m['line_spacing_val'] if m['line_spacing_val'] is not None else pinfo.get('ls')
        pinfo['line_spacing_rule'] = m['line_spacing_rule']
        pinfo['line_spacing_fixed_pt'] = m['line_spacing_fixed_pt']
        pinfo['space_before_pt'] = m['space_before_pt']
        pinfo['space_after_pt'] = m['space_after_pt']
        pinfo['first_indent_cm'] = m['first_indent_cm'] if m['first_indent_cm'] is not None else pinfo.get('indent', 0)

        fmt['paragraphs'].append(pinfo)

        # MD
        if p.text.strip() or pinfo['has_page_break']:
            flags = f"align={pinfo.get('align','?')} ls={pinfo.get('ls','?')}"
            if pinfo.get('indent'): flags += f' indent={pinfo["indent"]}cm'
            if pinfo['has_page_break']: flags += ' [PAGE BREAK]'
            txt = p.text[:100].replace('\n', '\\n')
            run_fmts = []
            for r in pinfo['runs'][:4]:
                parts = [f'{r.get("font","?")}', f'{r.get("size_pt","?")}pt']
                if r.get('bold'): parts.append('B')
                if r.get('italic'): parts.append('I')
                run_fmts.append('|'.join(parts))
            md_lines.append(f'**P{idx}** [{pinfo["style"]}] {flags}')
            if run_fmts: md_lines.append(f'  runs: {" / ".join(run_fmts)}')
            if txt: md_lines.append(f'  > {txt}')
            md_lines.append('')

    # ── Tables ──
    md_lines.append('## 表格\n')
    for ti, table in enumerate(doc.tables):
        tinfo = {'index': ti, 'rows': len(table.rows), 'cols': len(table.columns), 'cells': []}
        for ri, row in enumerate(table.rows):
            row_cells = []
            for ci, cell in enumerate(row.cells):
                cell_runs = []
                for p in cell.paragraphs:
                    for r in p.runs:
                        info = resolver.resolve(p._element, r._element)
                        cell_runs.append({
                            'text': r.text, 'font': info['font'], 'size_pt': info['size'], 'bold': info['bold'],
                        })
                row_cells.append({'row': ri, 'col': ci, 'text': cell.text, 'runs': cell_runs})
            tinfo['cells'].append(row_cells)
            if ri < 5:
                md_lines.append(f'  Row{ri}: {[c["text"][:50] for c in row_cells]}')
        fmt['tables'].append(tinfo)
        md_lines.append('')

    # ── Verification ──
    md_lines.append(f'\n## 验证\n')
    md_lines.append(f'- 段落: JSON={len(fmt["paragraphs"])} docx={len(doc.paragraphs)} ✓')
    md_lines.append(f'- 表格: JSON={len(fmt["tables"])} docx={len(doc.tables)} ✓')
    md_lines.append(f'- 节:   JSON={len(fmt["sections"])} docx={len(doc.sections)} ✓')

    # ── Cover extraction: walk body elements until abstract/body content ──
    asset_dir = os.path.splitext(docx_path)[0] + '_assets'
    fmt['_meta']['assets_dir'] = os.path.abspath(asset_dir)
    fmt['cover'] = _extract_cover(doc, asset_dir)

    # ── Normal style: extract for cover empty paragraph spacing ──
    ns = doc.styles['Normal']
    fmt['normal_style'] = {
        'font_name': ns.font.name,
        'font_size_pt': round(ns.font.size / 12700, 1) if ns.font.size else None,
        'line_spacing': ns.paragraph_format.line_spacing,
        'line_spacing_rule': str(ns.paragraph_format.line_spacing_rule) if ns.paragraph_format.line_spacing_rule else None,
    }

    # Semantic style profiles consumed by script_generator.py.
    fmt['style_profiles'] = _build_style_profiles(fmt)

    return fmt, '\n'.join(md_lines)



def _crop_blob_by_src_rect(blob, src_rect, ext):
    """Crop a DOCX image blob according to DrawingML a:srcRect.

    Word stores crop values as 1/1000 percent.  If a template uses one
    composite logo image and displays two cropped views, saving the raw blob
    would reproduce the composite image twice.  Cropping here turns each view
    into its own asset, while keeping the renderer role-driven.
    """
    if not src_rect or Image is None:
        return blob, ext
    try:
        vals = {k: int(src_rect.get(k, 0) or 0) for k in ('l', 't', 'r', 'b')}
    except Exception:
        return blob, ext
    if not any(vals.values()):
        return blob, ext
    try:
        img = Image.open(BytesIO(blob))
        w, h = img.size
        # OOXML srcRect is expressed in 1/100000 of the source dimension.
        left = max(0, int(w * vals['l'] / 100000.0))
        top = max(0, int(h * vals['t'] / 100000.0))
        right = min(w, int(w * (1 - vals['r'] / 100000.0)))
        bottom = min(h, int(h * (1 - vals['b'] / 100000.0)))
        if right <= left or bottom <= top:
            return blob, ext
        out = BytesIO()
        img.crop((left, top, right, bottom)).save(out, format='PNG')
        return out.getvalue(), 'png'
    except Exception:
        return blob, ext


def _cover_table_role(rows_data):
    """Infer cover table role from structure instead of school-specific text."""
    def cell_text(cell):
        return ''.join(r.get('t', '') for pp in cell.get('p', []) for r in pp.get('r', []))
    first_col = []
    for row in rows_data or []:
        if row:
            first_col.append(re.sub(r'[\s：:]+', '', cell_text(row[0])))
    if first_col and len(rows_data or []) <= 2 and all(x.endswith('编码') for x in first_col if x):
        return 'cover_code_table'
    if len(rows_data or []) >= 3 and sum(1 for x in first_col if x.endswith(('题目', '姓名', '学号', '学院', '班级', '老师', '教师'))) >= 2:
        return 'cover_info_table'
    return 'cover_table'


def _ooxml_attrs(el, ns):
    """Return OOXML attributes without namespace prefixes for JSON storage."""
    if el is None:
        return {}
    out = {}
    for k, v in el.attrib.items():
        key = k.split('}')[-1] if '}' in k else k
        out[key] = v
    return out


def _extract_margin_box(parent, ns, child_name):
    box = parent.find(f'{{{ns}}}{child_name}') if parent is not None else None
    if box is None:
        return {}
    res = {}
    for side in ('top', 'left', 'bottom', 'right', 'start', 'end'):
        el = box.find(f'{{{ns}}}{side}')
        if el is not None:
            res[side] = _ooxml_attrs(el, ns)
    return res


def _extract_tbl_props(tbl_elem):
    """Extract table-level layout properties so cover tables can be replayed.

    This is deliberately structural: alignment, indent, width, grid columns,
    layout mode, margins and row heights are copied from the template instead
    of being guessed from labels such as school code or committee text.
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    props = {'tblW': {}, 'tblInd': {}, 'jc': None, 'tblLayout': None,
             'cellMar': {}, 'grid_cols': [], 'rows': []}
    tblPr = tbl_elem.find(f'{{{W}}}tblPr')
    if tblPr is not None:
        for name in ('tblW', 'tblInd', 'tblLayout'):
            el = tblPr.find(f'{{{W}}}{name}')
            if el is not None:
                props[name] = _ooxml_attrs(el, W)
        jc = tblPr.find(f'{{{W}}}jc')
        if jc is not None:
            props['jc'] = jc.get(f'{{{W}}}val')
        props['cellMar'] = _extract_margin_box(tblPr, W, 'tblCellMar')
    grid = tbl_elem.find(f'{{{W}}}tblGrid')
    if grid is not None:
        for gc in grid.findall(f'{{{W}}}gridCol'):
            w = gc.get(f'{{{W}}}w')
            if w:
                props['grid_cols'].append(w)
    for tr in tbl_elem.findall(f'{{{W}}}tr'):
        trp = {'height': {}, 'cantSplit': False}
        trPr = tr.find(f'{{{W}}}trPr')
        if trPr is not None:
            h = trPr.find(f'{{{W}}}trHeight')
            if h is not None:
                trp['height'] = _ooxml_attrs(h, W)
            trp['cantSplit'] = trPr.find(f'{{{W}}}cantSplit') is not None
        props['rows'].append(trp)
    return props


def _extract_tc_props(tcPr):
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    info = {'tcW': {}, 'tcMar': {}, 'vAlign': None, 'gridSpan': None, 'vMerge': None}
    if tcPr is None:
        return info
    tcW = tcPr.find(f'{{{W}}}tcW')
    if tcW is not None:
        info['tcW'] = _ooxml_attrs(tcW, W)
    info['tcMar'] = _extract_margin_box(tcPr, W, 'tcMar')
    va = tcPr.find(f'{{{W}}}vAlign')
    if va is not None:
        info['vAlign'] = va.get(f'{{{W}}}val')
    gs = tcPr.find(f'{{{W}}}gridSpan')
    if gs is not None:
        info['gridSpan'] = gs.get(f'{{{W}}}val')
    vm = tcPr.find(f'{{{W}}}vMerge')
    if vm is not None:
        info['vMerge'] = vm.get(f'{{{W}}}val') or 'continue'
    return info

def _extract_cover(doc, assets_dir=None):
    """Walk template body from start, extract cover+declaration elements with full formatting.
    Stops at abstract/TOC/body content. Returns list of element dicts or [] if no cover."""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

    if assets_dir:
        os.makedirs(assets_dir, exist_ok=True)
    image_counter = 0

    def _save_image_by_rid(rid, src_rect=None):
        nonlocal image_counter
        if not rid or rid not in doc.part.rels:
            return None
        rel = doc.part.rels[rid]
        if 'image' not in rel.reltype:
            return None
        ext = rel.target_ref.rsplit('.', 1)[-1].lower() if '.' in rel.target_ref else 'png'
        if ext not in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'tif', 'tiff', 'emf', 'wmf'):
            ext = 'png'
        blob, ext = _crop_blob_by_src_rect(rel.target_part.blob, src_rect, ext)
        image_counter += 1
        fname = f'cover_img_{image_counter:03d}.{ext}'
        if assets_dir:
            fpath = os.path.join(assets_dir, fname)
            with open(fpath, 'wb') as f:
                f.write(blob)
            return fname
        return fname

    def _image_payload(elem):
        extent = None
        src_rect = {}
        rid = None
        drawing_nodes = list(elem.iter(f'{{{WP}}}inline')) + list(elem.iter(f'{{{WP}}}anchor'))
        for drawing in drawing_nodes:
            ext = drawing.find(f'{{{WP}}}extent')
            if ext is not None:
                extent = {'cx': ext.get('cx', '0'), 'cy': ext.get('cy', '0')}
            for sr in drawing.iter(f'{{{A}}}srcRect'):
                src_rect = {'l': sr.get('l', '0'), 't': sr.get('t', '0'), 'r': sr.get('r', '0'), 'b': sr.get('b', '0')}
            for blip in drawing.iter(f'{{{A}}}blip'):
                rid = blip.get(f'{{{R}}}embed')
                if rid:
                    break
            if rid:
                break
        asset = _save_image_by_rid(rid, src_rect) if rid else None
        if not asset:
            return None
        return {'asset': asset, 'extent': extent, 'srcRect': src_rect, 'rEmbed': rid}

    STOP_KW = ['摘 要', '摘要', '目  录', '目录', '目 录', 'ABSTRACT', '第1章', '1.1 ', '1.1.']
    SKIP_KW = ['页边距要求', '碳素笔', '完成后删除', '封面要求', '1.论文题目', '毕业论文（设计）题目为',
               '按答辩时间', '提交论文', '禁止使用']

    elements = []
    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sectPr':
            break

        if tag == 'tbl':
            tbl_props = _extract_tbl_props(child)
            rows_data = []
            for row in child.findall(f'{{{W}}}tr'):
                cells_data = []
                for tc in row.findall(f'{{{W}}}tc'):
                    tcPr = tc.find(f'{{{W}}}tcPr')
                    # Cell width
                    tcW = tcPr.find(f'{{{W}}}tcW') if tcPr is not None else None
                    cell_w = int(tcW.get(f'{{{W}}}w', '0')) if tcW is not None else 0
                    # Cell borders
                    cell_borders = {}
                    tcBorders = tcPr.find(f'{{{W}}}tcBorders') if tcPr is not None else None
                    if tcBorders is not None:
                        for b in tcBorders:
                            btag = b.tag.split('}')[-1]
                            bval = b.get(f'{{{W}}}val', 'nil')
                            if bval not in (None, 'nil', 'none'):
                                cell_borders[btag] = {
                                    'val': bval,
                                    'sz': b.get(f'{{{W}}}sz', '0'),
                                    'color': b.get(f'{{{W}}}color', '000000'),
                                }
                    # Cell paragraphs
                    cell_paras = []
                    for p in tc.findall(f'{{{W}}}p'):
                        pPr = p.find(f'{{{W}}}pPr')
                        jc = pPr.find(f'{{{W}}}jc') if pPr is not None else None
                        palign = jc.get(f'{{{W}}}val') if jc is not None else None
                        spacing = pPr.find(f'{{{W}}}spacing') if pPr is not None else None
                        p_line = spacing.get(f'{{{W}}}line') if spacing is not None else None
                        p_rule = spacing.get(f'{{{W}}}lineRule') if spacing is not None else None
                        p_before = spacing.get(f'{{{W}}}before') if spacing is not None else None
                        p_after = spacing.get(f'{{{W}}}after') if spacing is not None else None
                        p_ind = pPr.find(f'{{{W}}}ind') if pPr is not None else None
                        p_first = p_ind.get(f'{{{W}}}firstLine') if p_ind is not None else None
                        runs = []
                        for r in p.findall(f'{{{W}}}r'):
                            rPr = r.find(f'{{{W}}}rPr')
                            fn_ascii, fn_ea, fsz, fbold = '', '', 0, False
                            if rPr is not None:
                                rf = rPr.find(f'{{{W}}}rFonts')
                                if rf is not None:
                                    fn_ascii = rf.get(f'{{{W}}}ascii', '') or ''
                                    fn_ea = rf.get(f'{{{W}}}eastAsia', '') or ''
                                sz = rPr.find(f'{{{W}}}sz')
                                if sz is not None:
                                    fsz = int(sz.get(f'{{{W}}}val', '0')) // 2
                                _b_el = rPr.find(f'{{{W}}}b'); fbold = _b_el is not None and _b_el.get(f'{{{W}}}val', '1') not in ('0', 'false', 'False')
                            txt = ''.join(t.text or '' for t in r.findall(f'{{{W}}}t'))
                            payload = _image_payload(r)
                            if txt:
                                runs.append({'t': txt, 'fn': fn_ascii, 'fe': fn_ea, 'sz': fsz, 'b': fbold})
                            if payload:
                                runs.append({'t': '', 'fn': fn_ascii, 'fe': fn_ea, 'sz': fsz, 'b': fbold, **payload})
                            if not txt and not payload:
                                runs.append({'t': txt, 'fn': fn_ascii, 'fe': fn_ea, 'sz': fsz, 'b': fbold})
                        cell_paras.append({'al': palign, 'ls_val': p_line, 'ls_rule': p_rule,
                                           'sp_before': p_before, 'sp_after': p_after,
                                           'fl_indent': p_first, 'r': runs})
                    cells_data.append({'w': cell_w, 'tcPr': _extract_tc_props(tcPr), 'borders': cell_borders, 'p': cell_paras})
                rows_data.append(cells_data)
            elements.append({'type': 'table', 'role': _cover_table_role(rows_data), 'tblPr': tbl_props, 'rows': rows_data})
            continue

        if tag != 'p':
            continue

        # Extract paragraph data
        pPr = child.find(f'{{{W}}}pPr')
        has_sectPr = pPr.find(f'{{{W}}}sectPr') is not None if pPr is not None else False
        jc = pPr.find(f'{{{W}}}jc') if pPr is not None else None
        palign = jc.get(f'{{{W}}}val') if jc is not None else None

        spacing = pPr.find(f'{{{W}}}spacing') if pPr is not None else None
        line_val = spacing.get(f'{{{W}}}line') if spacing is not None else None
        lineRule = spacing.get(f'{{{W}}}lineRule') if spacing is not None else None
        before_val = spacing.get(f'{{{W}}}before') if spacing is not None else None
        after_val = spacing.get(f'{{{W}}}after') if spacing is not None else None

        indent = pPr.find(f'{{{W}}}ind') if pPr is not None else None
        first_line = indent.get(f'{{{W}}}firstLine') if indent is not None else None

        runs = []
        for r in child.findall(f'{{{W}}}r'):
            rPr = r.find(f'{{{W}}}rPr')
            fn_ascii, fn_ea, fsz, fbold = '', '', 0, False
            if rPr is not None:
                rf = rPr.find(f'{{{W}}}rFonts')
                if rf is not None:
                    fn_ascii = rf.get(f'{{{W}}}ascii', '') or ''
                    fn_ea = rf.get(f'{{{W}}}eastAsia', '') or ''
                sz = rPr.find(f'{{{W}}}sz')
                if sz is not None:
                    fsz = int(sz.get(f'{{{W}}}val', '0')) // 2
                _b_el = rPr.find(f'{{{W}}}b'); fbold = _b_el is not None and _b_el.get(f'{{{W}}}val', '1') not in ('0', 'false', 'False')
            txt = ''.join(t.text or '' for t in r.findall(f'{{{W}}}t'))
            runs.append({'t': txt, 'fn': fn_ascii, 'fe': fn_ea, 'sz': fsz, 'b': fbold})

        full_text = ''.join(r['t'] for r in runs).strip()

        # ── Heuristic: is this paragraph a format instruction? ──
        # Format notes specify fonts, sizes, alignments — not actual content.
        _fmt_font = any(f in full_text for f in ['黑体', '宋体', '楷体', '华文', '方正', 'Times New Roman'])
        _fmt_size = any(kw in full_text for kw in ['二号', '三号', '四号', '小四', '五号', '小五',
                                                    '号加粗', '号居中', 'pt', '号字'])
        _fmt_align = any(kw in full_text for kw in ['居中', '加粗', '缩进', '对齐', '行距', '段前', '段后',
                                                     '固定值', '倍行距', '1.5倍', '双倍'])
        _fmt_paren = '（' in full_text and '）' in full_text
        _is_fmt_note = (_fmt_paren and (_fmt_font or _fmt_size)) or (_fmt_font and _fmt_size and _fmt_align)
        _is_fmt_note = _is_fmt_note or (_fmt_paren and any(kw in full_text for kw in ['空一行', '空两行', '空行', '空  行']))
        _is_fmt_header = any(kw in full_text[:60] for kw in ['页眉页脚', '页眉', '页码', '字体要求',
                                                               '字号要求', '格式要求', '排版要求'])

        # Stop at abstract/TOC/body (real content, not format notes)
        if full_text and any(kw in full_text[:20] for kw in STOP_KW):
            if not _is_fmt_note:
                break

        # Skip individual format notes & section headers
        if full_text:
            if _is_fmt_note or _is_fmt_header:
                continue
            if any(kw in full_text[:30] for kw in SKIP_KW):
                continue
            if len(full_text) > 40 and '删除' in full_text[:80]:
                continue

        # Check for images and save binary assets so the generator can reinsert logos.
        img_payload = _image_payload(child)

        if img_payload:
            elements.append({
                'type': 'image',
                'al': palign, 'ls_val': line_val, 'ls_rule': lineRule,
                'sp_before': before_val, 'sp_after': after_val, 'fl_indent': first_line,
                **img_payload,
                'r': runs,
                'section_break_after': has_sectPr,
            })
        elif not full_text:
            elements.append({
                'type': 'empty',
                'al': palign, 'ls_val': line_val, 'ls_rule': lineRule,
                'sp_before': before_val, 'sp_after': after_val, 'fl_indent': first_line,
                'r': runs,  # run font sizes control paragraph height even when empty
                'section_break_after': has_sectPr,
            })
        else:
            elements.append({
                'type': 'para',
                'al': palign, 'ls_val': line_val, 'ls_rule': lineRule,
                'sp_before': before_val, 'sp_after': after_val, 'fl_indent': first_line,
                'r': runs,
                'section_break_after': has_sectPr,
            })

    return elements


if __name__ == '__main__':
    import sys
    path = sys.argv[1] if len(sys.argv) > 1 else 'Templates/模版.docx'
    fmt, md = extract(path)
    json_path = path.replace('.docx', '_format.json')
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(fmt, f, ensure_ascii=False, indent=2)
    md_path = path.replace('.docx', '_格式提取.md')
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md)
    print(f'Format JSON -> {json_path}')
    print(f'Format MD   -> {md_path}')
    print(f'Paragraphs: {len(fmt["paragraphs"])}  Tables: {len(fmt["tables"])}  Sections: {len(fmt["sections"])}')
