"""Script generator and generated-runtime regression cases."""
from __future__ import annotations

import json
import re
import zipfile
from typing import List
from xml.etree import ElementTree as ET

from docx import Document
from qa_conformance import check_conformance
from qa_conformance_modules.requirements import build_requirements
from regression_suite_modules.generated_docx import run_generated_case
from regression_suite_modules.harness import PNG_1X1, assert_true, base_content, base_format, case, new_workdir, write_json, write_sample_png
from script_generator import (
    RUNTIME_TEMPLATE,
    _extract_page_and_header,
    _front_matter_sections,
    _infer_style_profiles,
    _infer_template_rules,
    _normalize_numbered_section_order,
)

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _first_property_child_names(xml: str, prop_name: str) -> List[str]:
    root = ET.fromstring(xml.encode("utf-8"))
    prop = root.find(f".//{W_NS}{prop_name}")
    if prop is None:
        return []
    return [str(child.tag).rsplit("}", 1)[-1] for child in list(prop)]


@case
def conformance_style_check_ignores_static_toc_lines() -> None:
    content = base_content(
        [
            "Body paragraph after heading.",
        ]
    )
    content["sections"] = [
        {
            "heading": "2 Custom Chapter",
            "level": 1,
            "role": "body",
            "paragraphs": ["Body paragraph after heading."],
            "images": [],
        }
    ]
    result = run_generated_case("conformance_toc", content)
    conf = check_conformance(str(result["work"]), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in conf["issues"]]
    assert_true("STYLE_MISMATCH" not in codes, f"conformance matched TOC lines instead of body headings: {conf['issues']}")
    assert_true("docx_sections" in conf["counts"], "conformance report should name Word section count as docx_sections")
    assert_true("sections" not in conf["counts"], "ambiguous conformance count key 'sections' should not be emitted")


@case
def conformance_reference_labels_keep_template_cjk_font() -> None:
    content = base_content(["Body paragraph before references."])
    content["references"] = ["[1] 作者1. 中文参考文献与自动排版测试[J]. 综合研究评论, 2026."]
    result = run_generated_case("reference_cjk_font", content)
    conf = check_conformance(str(result["work"]), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in conf["issues"]]
    assert_true("STYLE_MISMATCH" not in codes, f"reference label triggered CJK font mismatch: {conf['issues']}")


@case
def conformance_checks_rich_text_and_duplicate_paragraphs() -> None:
    work = new_workdir("conformance_rich_duplicate_missing")
    content = base_content(
        [
            {
                "role": "rich_text",
                "text": "Inline formula should remain visible: x=1.",
                "runs": [
                    {"type": "text", "text": "Inline formula should remain visible: "},
                    {"type": "math", "text": "x=1", "math": [{"type": "inline", "text": "x=1", "latex": "x=1"}]},
                    {"type": "text", "text": "."},
                ],
                "math": [{"type": "inline", "text": "x=1", "latex": "x=1"}],
            },
            "Repeated body paragraph.",
            "Repeated body paragraph.",
        ]
    )
    write_json(work / "format.json", base_format())
    write_json(work / "content.json", content)
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Repeated body paragraph.")
    doc.save(work / "out.docx")
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {"content_formulas_rendered": 1}})
    conf = check_conformance(str(work), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in conf["issues"]]
    assert_true("CONTENT_PARAGRAPH_MISSING" in codes, f"conformance did not check rich_text/duplicate paragraphs: {conf['issues']}")


@case
def conformance_rejects_truncated_long_paragraphs() -> None:
    work = new_workdir("conformance_truncated_long_paragraph")
    full_text = (
        "This long paragraph starts with a stable prefix that previously could satisfy the loose matcher, "
        "but the final generated document must still contain the complete tail with conclusions and evidence."
    )
    content = base_content([full_text])
    write_json(work / "format.json", base_format())
    write_json(work / "content.json", content)
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(full_text[:95])
    doc.save(work / "out.docx")
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    conf = check_conformance(str(work), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in conf["issues"]]
    assert_true("CONTENT_PARAGRAPH_MISSING" in codes, f"truncated long paragraph passed conformance: {conf['issues']}")
    assert_true("content.json" in conf.get("next_action", ""), f"conformance next action should guide paragraph repair: {conf}")


@case
def conformance_requirements_count_mixed_inline_and_section_images() -> None:
    content = base_content([{"role": "image", "image": "inline.png"}])
    content["sections"][0]["images"] = ["inline.png", "", "section_only.png"]
    req = build_requirements(base_format(), content)
    assert_true(req["expected_counts"]["images"] == 2, f"conformance image requirements lost section-only image: {req['expected_counts']}")


@case
def script_generator_keeps_figure_reference_prose_as_body() -> None:
    fmt = base_format()
    fmt["style_profiles"] = {
        "body": {
            "font": "宋体",
            "size": 12.0,
            "align": "JUSTIFY",
            "line_spacing_fixed_pt": 28.0,
            "first_indent_cm": 0.74,
            "space_before_pt": 0.0,
            "space_after_pt": 0.0,
        },
        "figure_caption": {
            "font": "宋体",
            "size": 10.5,
            "align": "CENTER",
            "line_spacing_fixed_pt": 28.0,
            "first_indent_cm": 0.0,
            "space_before_pt": 6.0,
            "space_after_pt": 6.0,
        },
    }
    content = base_content(["图 1 展示了从数据到决策的机器学习研究流程。"])
    result = run_generated_case("figure_reference_prose_body", content, fmt=fmt)
    conf = check_conformance(str(result["work"]), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in conf["issues"]]
    assert_true("STYLE_MISMATCH" not in codes, f"figure reference prose was rendered as caption: {conf['issues']}")


@case
def script_generator_renders_native_footnotes_from_note_runs() -> None:
    note_text = "Generated native footnote text."
    content = base_content(
        [
            {
                "role": "rich_text",
                "text": "A sentence with a footnote.",
                "runs": [
                    {"type": "text", "text": "A sentence with a footnote"},
                    {"type": "note_ref", "note_type": "footnote", "source_id": "2", "text": note_text},
                    {"type": "text", "text": "."},
                ],
                "notes": [{"type": "footnote", "source_id": "2", "text": note_text}],
            }
        ]
    )
    content["_meta"]["footnote_references_extracted"] = 1
    content["_meta"]["footnote_definitions_extracted"] = 1
    result = run_generated_case("native_footnote_render", content)
    with zipfile.ZipFile(result["work"] / "out.docx") as zf:
        footnotes_xml = zf.read("word/footnotes.xml").decode("utf-8", errors="replace")
    manifest = json.loads((result["work"] / "build_manifest.json").read_text(encoding="utf-8"))

    assert_true("<w:footnoteReference" in result["xml"], f"document body has no footnoteReference: {result['xml'][:500]}")
    assert_true(note_text in footnotes_xml, f"footnote text missing from footnotes.xml: {footnotes_xml}")
    assert_true(manifest["counts"].get("footnote_references_rendered") == 1, f"footnote render count missing: {manifest}")
    assert_true(manifest["counts"].get("footnote_definitions_rendered") == 1, f"footnote definition count missing: {manifest}")


@case
def script_generator_renders_native_endnotes_from_note_runs() -> None:
    note_text = "Generated native endnote text."
    content = base_content(
        [
            {
                "role": "rich_text",
                "text": "A sentence with an endnote.",
                "runs": [
                    {"type": "text", "text": "A sentence with an endnote"},
                    {"type": "note_ref", "note_type": "endnote", "source_id": "5", "text": note_text},
                    {"type": "text", "text": "."},
                ],
                "notes": [{"type": "endnote", "source_id": "5", "text": note_text}],
            }
        ]
    )
    content["_meta"]["endnote_references_extracted"] = 1
    content["_meta"]["endnote_definitions_extracted"] = 1
    result = run_generated_case("native_endnote_render", content)
    with zipfile.ZipFile(result["work"] / "out.docx") as zf:
        endnotes_xml = zf.read("word/endnotes.xml").decode("utf-8", errors="replace")
    manifest = json.loads((result["work"] / "build_manifest.json").read_text(encoding="utf-8"))

    assert_true("<w:endnoteReference" in result["xml"], f"document body has no endnoteReference: {result['xml'][:500]}")
    assert_true(note_text in endnotes_xml, f"endnote text missing from endnotes.xml: {endnotes_xml}")
    assert_true(manifest["counts"].get("endnote_references_rendered") == 1, f"endnote render count missing: {manifest}")
    assert_true(manifest["counts"].get("endnote_definitions_rendered") == 1, f"endnote definition count missing: {manifest}")


@case
def script_generator_cover_removes_template_instructions_and_fills_fields() -> None:
    def run(text: str) -> Dict[str, Any]:
        return {"t": text, "fn": "Times New Roman", "fe": "宋体", "sz": 12, "b": False}

    def para(text: str) -> Dict[str, Any]:
        return {"al": "center", "r": [run(text)]}

    def cell(*paragraphs: str) -> Dict[str, Any]:
        return {"p": [para(text) for text in paragraphs]}

    fmt = base_format()
    fmt["cover"] = [
        {"type": "para", "al": "center", "r": [run("（完成时间按照答辩时间填写）")]},
        {"type": "para", "al": "center", "r": [run("  年  月  日")]},
        {
            "type": "table",
            "role": "cover_info_table",
            "rows": [
                [
                    cell("题    目："),
                    cell(
                        "Sample English Title",
                        "示例中文题目",
                        "英文题目(Times new Roman)/中文题目(宋体),三号加粗，1.5倍行距",
                    ),
                ],
                [cell("年级专业："), cell("表格行高0.9cm")],
                [cell("姓    名："), cell("楷体四号居中，1.5倍行距")],
                [cell("学    号："), cell("")],
                [cell("指导教师："), cell("")],
            ],
        },
        {"type": "para", "al": "left", "r": [run("2. An Overview of English Sports News Headlines   1")]},
    ]
    content = base_content(["Actual body paragraph."])
    content["cover_info"] = {
        "paper_title": "Real Thesis Title",
        "class_name": "English Major in Education (2020-2024)",
        "student_name": "Zhang San",
        "student_id": "2020123456",
        "advisor": "Prof. Li Si",
        "completion_date": "2026年 5月 6日",
    }

    result = run_generated_case("cover_template_instruction_cleanup", content, fmt=fmt)
    doc = Document(str(result["work"] / "out.docx"))
    text = "\n".join(
        [p.text for p in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    assert_true("Real Thesis Title" in text, f"real title was not rendered on cover: {text}")
    assert_true("English Major in Education" in text, f"class/major was not rendered on cover: {text}")
    assert_true("Zhang San" in text and "2020123456" in text and "Prof. Li Si" in text, f"cover fields missing: {text}")
    assert_true("2026年 5月 6日" in text, f"completion date was not rendered: {text}")
    for leaked in ["完成时间按照答辩时间填写", "英文题目(Times new Roman)", "表格行高0.9cm", "楷体四号居中", "An Overview of English Sports News Headlines"]:
        assert_true(leaked not in text, f"template instruction leaked into generated cover: {leaked} in {text}")
    assert_true("Sample English Title" not in text and "示例中文题目" not in text, f"sample template title leaked: {text}")
    field_runs = {}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip() in {"English Major in Education (2020-2024)", "Zhang San", "2020123456", "Prof. Li Si"}:
                            field_runs[run.text.strip()] = run
    for value in ["English Major in Education (2020-2024)", "2020123456", "Prof. Li Si"]:
        run = field_runs.get(value)
        assert_true(run is not None, f"cover field run missing: {value}")
        size = run.font.size.pt if run.font.size else 0
        min_size = 10.5 if value.startswith("English Major") else 13.5
        assert_true(size >= min_size, f"cover field inherited placeholder row-height style instead of value style: {value} size={size}")


@case
def script_generator_strips_template_instruction_from_header() -> None:
    fmt = base_format()
    fmt["sections"][0]["header"] = [
        {
            "text": "B. A. Thesis of English Major in Education (2022-2026)（新罗马字体，五号加粗居中）",
            "alignment": "CENTER",
            "runs": [
                {"text": "B. A. Thesis of English Major in Education (2022-2026)", "font": "Times New Roman", "size_pt": 10.5, "bold": True},
                {"text": "（新罗马字体，五号加粗居中）", "font": "Times New Roman", "size_pt": 9, "bold": True},
            ],
        }
    ]
    page = _extract_page_and_header(fmt)
    header = page.get("header") or {}
    assert_true(header.get("text") == "B. A. Thesis of English Major in Education (2022-2026)", f"header instruction was not stripped: {header}")


@case
def script_generator_infers_non_bold_english_body_sample() -> None:
    fmt = base_format()
    fmt["paragraphs"] = [
        {
            "text": "（正文从引言到参考文献不空行，不另起一页，1.5倍行距；段落首行缩进，两端对齐书写）",
            "runs": [{"font": "Times New Roman", "size_pt": 14, "bold": True}],
            "align": "LEFT",
        },
        {
            "text": "This template body sample is deliberately long enough to be selected as the body style source. It should stay non-bold, twelve point, and justified instead of inheriting a heading or instruction style.",
            "runs": [{"font": "Times New Roman", "size_pt": 12, "bold": False}],
            "align": "JUSTIFY",
        },
    ]
    profiles = _infer_style_profiles(fmt)
    body = profiles.get("body") or {}
    assert_true(body.get("font") == "Times New Roman", f"body font came from the wrong sample: {body}")
    assert_true(abs(float(body.get("size") or 0) - 12.0) < 0.01, f"body size came from the wrong sample: {body}")
    assert_true(body.get("bold") is False, f"body bold came from instruction/heading sample: {body}")


@case
def script_generator_ignores_toc_note_for_body_and_reference_bold() -> None:
    body_text = (
        "This is a real template body paragraph long enough to be selected as the body sample. "
        "It uses Times New Roman twelve point text and must remain non-bold even when nearby "
        "directory notes mention bold headings and references."
    )
    fmt = base_format()
    fmt["style_profiles"] = {
        "body": {"font": "Times New Roman", "size": 14.0, "bold": True, "align": "CENTER"},
        "reference": {"font": "宋体", "size": 15.0, "bold": True, "align": "CENTER"},
    }
    fmt["paragraphs"] = [
        {
            "text": "（备注：目录按1～2级标题编写，“正文的一级标题、参考文献、摘要、关键词”采用四号加粗书写，正文其他层次标题均采用Times New Roman体小四号书写。）",
            "runs": [{"font": "宋体", "size_pt": 12, "bold": False}],
            "align": "JUSTIFY",
        },
        {
            "text": "（正文从引言到参考文献不空行，不另起一页，1.5倍行距；段落首行缩进2字符，两端对齐书写。）",
            "runs": [{"font": "Times New Roman", "size_pt": 12, "bold": False}],
            "align": "LEFT",
            "indent": 0.85,
        },
        {
            "text": body_text,
            "runs": [{"font": "Times New Roman", "size_pt": 12, "bold": False}],
            "align": "JUSTIFY",
            "indent": 0.85,
        },
        {
            "text": "[1] Schiff, James A. John Updike Revisited [M]. New York: Twayne Publishers, 1998.",
            "runs": [{"font": "Times New Roman", "size_pt": 12, "bold": False}],
            "align": "LEFT",
        },
    ]
    profiles = _infer_style_profiles(fmt)
    body = profiles.get("body") or {}
    reference = profiles.get("reference") or {}
    assert_true(body.get("bold") is False, f"body inherited bold from TOC note: {body}")
    assert_true(body.get("align") == "JUSTIFY", f"body did not keep body alignment: {body}")
    assert_true(abs(float(body.get("size") or 0) - 12.0) < 0.01, f"body size came from TOC note: {body}")
    assert_true(reference.get("bold") is False, f"reference inherited bold from heading note: {reference}")
    assert_true(abs(float(reference.get("size") or 0) - 12.0) < 0.01, f"reference size came from heading note: {reference}")
    assert_true(reference.get("align") != "CENTER", f"reference entries should not inherit centered heading style: {reference}")


@case
def script_generator_keeps_chinese_front_matter_styles_from_template_rules() -> None:
    fmt = base_format()
    fmt["style_profiles"] = {
        "cn_title": {"font": "Arial", "size": 12.0, "bold": False, "align": "JUSTIFY"},
        "cn_abstract_body": {"font": "宋体", "size": 12.0, "bold": False, "align": "LEFT", "first_indent_cm": 0.85},
        "cn_keywords": {"font": "Times New Roman", "size": 14.0, "bold": True, "align": "LEFT"},
    }
    fmt["paragraphs"] = [
        {
            "text": "（备注：目录按1～2级标题编写，“正文的一级标题、参考文献、摘要、关键词”采用四号加粗书写，正文其他层次标题均采用Times New Roman体小四号书写。）",
            "runs": [{"font": "宋体", "size_pt": 12, "bold": False}],
            "align": "JUSTIFY",
        },
        {
            "text": "模因论视角下的英语体育新闻标题汉译研究",
            "runs": [{"font": "Times New Roman", "size_pt": 15, "bold": True}],
            "align": "CENTER",
        },
        {
            "text": "（宋体小三号，加粗，居中；单独成页）",
            "runs": [{"font": "宋体", "size_pt": 12, "bold": False}],
            "align": "CENTER",
        },
        {
            "text": "【摘要】（“摘要”二字宋体四号，加粗）摘要正文样本用于确认中文摘要内容不继承英文正文样式。",
            "runs": [
                {"text": "【", "font": "宋体", "size_pt": 14, "bold": False},
                {"text": "摘要", "font": "宋体", "size_pt": 14, "bold": True},
                {"text": "】", "font": "宋体", "size_pt": 14, "bold": False},
                {"text": "摘要正文样本用于确认中文摘要内容不继承英文正文样式。", "font": "宋体", "size_pt": 12, "bold": False},
            ],
            "align": "LEFT",
        },
        {
            "text": "【关键词】（“关键词三字”宋体四号加粗）模因论；体育新闻标题",
            "runs": [
                {"text": "【", "font": "宋体", "size_pt": 14, "bold": False},
                {"text": "关键词", "font": "宋体", "size_pt": 14, "bold": True},
                {"text": "】", "font": "宋体", "size_pt": 14, "bold": False},
                {"text": "模因论；体育新闻标题", "font": "宋体", "size_pt": 12, "bold": False},
            ],
            "align": "LEFT",
        },
        {
            "text": "（中文摘要内容必须与英文摘要完全对应。摘要和关键词内容小四号，1.5倍行距；关键词3-5个，关键词间用分号隔开）",
            "runs": [{"font": "Times New Roman", "size_pt": 12, "bold": False}],
            "align": "LEFT",
        },
        {
            "text": "This real English body sample is long enough to set the normal body style without taking over Chinese front matter roles.",
            "runs": [{"font": "Times New Roman", "size_pt": 12, "bold": False}],
            "align": "JUSTIFY",
        },
    ]
    profiles = _infer_style_profiles(fmt)
    assert_true(profiles["cn_title"]["font"] == "宋体" and profiles["cn_title"]["bold"] is True, f"cn title rule not applied: {profiles['cn_title']}")
    assert_true(abs(float(profiles["cn_title"]["size"]) - 15.0) < 0.01, f"cn title size wrong: {profiles['cn_title']}")
    assert_true(profiles["cn_abstract_body"]["font"] == "宋体" and profiles["cn_abstract_body"]["bold"] is False, f"cn abstract body inherited English style: {profiles['cn_abstract_body']}")
    assert_true(profiles["cn_keywords"]["font"] == "宋体" and profiles["cn_keywords"]["bold"] is False, f"cn keyword value style wrong: {profiles['cn_keywords']}")
    assert_true(profiles["cn_keywords_label"]["bold"] is True and profiles["cn_keywords_label"]["size"] >= 14, f"cn keyword label style wrong: {profiles['cn_keywords_label']}")


@case
def script_generator_renders_english_figure_caption_with_caption_style() -> None:
    img_src = new_workdir("english_caption_image_src")
    write_sample_png(img_src / "figure.png", width=480, height=260)
    fmt = base_format()
    fmt["style_profiles"] = {
        "body": {
            "font": "Times New Roman",
            "size": 12.0,
            "bold": False,
            "align": "JUSTIFY",
            "line_spacing_val": 1.5,
            "first_indent_cm": 0.0,
            "space_before_pt": 0.0,
            "space_after_pt": 0.0,
        },
        "figure_caption": {
            "font": "宋体",
            "size": 10.5,
            "bold": False,
            "align": "CENTER",
            "line_spacing_fixed_pt": 28.0,
            "first_indent_cm": 0.0,
            "space_before_pt": 6.0,
            "space_after_pt": 6.0,
        },
    }
    content = base_content([
        {"role": "image", "image": "figure.png"},
        "Figure 1. Distribution of meme types identified in the corpus.",
        "Body paragraph after the figure.",
    ])
    content["_meta"]["images_dir"] = str(img_src)
    content["_meta"]["images_extracted"] = 1
    content["sections"][0]["images"] = ["figure.png"]
    result = run_generated_case("english_figure_caption_style", content, fmt=fmt)
    conf = check_conformance(str(result["work"]), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in conf["issues"]]
    assert_true("STYLE_MISMATCH" not in codes, f"English Figure caption was rendered with body style: {conf['issues']}")


@case
def script_generator_honors_body_page_break_before() -> None:
    content = base_content([])
    content["sections"] = [
        {
            "heading": "1 First",
            "level": 1,
            "role": "body",
            "paragraphs": ["First body paragraph."],
            "images": [],
            "page_break_before": True,
        },
        {
            "heading": "2 Second",
            "level": 1,
            "role": "body",
            "paragraphs": ["Second body paragraph."],
            "images": [],
            "page_break_before": True,
        },
    ]
    result = run_generated_case("body_page_break_before", content)
    page_breaks = len(re.findall(r'<w:br[^>]+w:type="page"', result["xml"]))
    assert_true(page_breaks >= 1, "body page_break_before was not rendered as a page break")


@case
def script_generator_renders_chinese_backmatter_headings() -> None:
    content = base_content([])
    content["sections"] = [
        {
            "heading": "1 Body",
            "level": 1,
            "role": "body",
            "paragraphs": ["Body paragraph."],
            "images": [],
        },
        {
            "heading": "Acknowledgements",
            "level": 1,
            "role": "acknowledgement",
            "paragraphs": ["Thanks paragraph."],
            "images": [],
        },
        {
            "heading": "\u9644\u5f55",
            "level": 1,
            "role": "appendix",
            "paragraphs": ["Appendix paragraph."],
            "images": [],
        },
    ]
    result = run_generated_case("backmatter_unicode_headings", content)
    xml_compact = re.sub(r"\s+", "", result["xml"])
    assert_true("\u81f4\u8c22" in xml_compact, "acknowledgement heading was not rendered as Chinese text")
    assert_true("\u9644\u5f55" in xml_compact, "appendix heading was not rendered as Chinese text")
    codes = [item["code"] for item in result["report"]["issues"]]
    assert_true("CONTENT_HEADING_MISSING" not in codes, f"backmatter headings were still reported missing: {codes}")


@case
def script_generator_renders_backmatter_rich_items() -> None:
    work = new_workdir("backmatter_item_assets")
    img_dir = work / "figures"
    img_dir.mkdir()
    (img_dir / "dot.png").write_bytes(PNG_1X1)

    content = base_content([])
    content["_meta"]["images_dir"] = str(img_dir)
    content["_meta"]["images_extracted"] = 1
    content["_meta"]["tables_count"] = 1
    content["sections"] = [
        {
            "heading": "1 Body",
            "level": 1,
            "role": "body",
            "paragraphs": ["Body paragraph."],
            "images": [],
        },
        {
            "heading": "\u9644\u5f55",
            "level": 1,
            "role": "appendix",
            "paragraphs": [
                {"role": "table", "table_rows": [["A", "B"], ["1", "2"]]},
                {"role": "formula", "source": "latex", "latex": "E=mc^2", "text": "E=mc^2", "numbered": False},
                {"role": "image", "image": "dot.png"},
            ],
            "images": ["dot.png"],
        },
    ]
    result = run_generated_case("backmatter_rich_items", content)
    counts = result["manifest"]["counts"]
    assert_true(counts["content_tables_rendered"] == 1, f"appendix table was not rendered: {counts}")
    assert_true(counts["content_formulas_rendered"] == 1, f"appendix formula was not rendered: {counts}")
    assert_true(counts["content_images_rendered"] == 1, f"appendix image was not rendered: {counts}")


@case
def script_generator_renders_merged_table_cells() -> None:
    content = base_content(
        [
            {
                "role": "table",
                "table_rows": [
                    ["Merged header", "", "Score"],
                    ["Group A", "Alpha", "1"],
                    ["", "Beta", "2"],
                ],
                "table_merges": [
                    {"row": 0, "col": 0, "rowspan": 1, "colspan": 2},
                    {"row": 1, "col": 0, "rowspan": 2, "colspan": 1},
                ],
            }
        ],
        meta_tables=1,
    )
    result = run_generated_case("merged_table_cells", content)
    xml = result["xml"]
    assert_true(re.search(r"<w:gridSpan\b[^>]*/?w:val=\"2\"", xml), "horizontal merge gridSpan was not rendered")
    assert_true(re.search(r"<w:vMerge\b[^>]*/?w:val=\"restart\"", xml), "vertical merge restart was not rendered")
    assert_true(re.search(r"<w:vMerge\s*/>", xml), "vertical merge continuation was not rendered")
    counts = result["manifest"]["counts"]
    assert_true(counts.get("content_table_merges_rendered") == 2, f"table merge render count was not recorded: {counts}")


@case
def script_generator_renders_table_column_widths() -> None:
    content = base_content(
        [
            {
                "role": "table",
                "table_rows": [
                    ["Metric", "Description", "Value"],
                    ["A", "Longer explanatory text", "1"],
                ],
                "table_col_widths_twips": [1200, 2800, 1600],
            }
        ],
        meta_tables=1,
    )
    result = run_generated_case("table_column_widths", content)
    xml = result["xml"]
    assert_true(re.search(r"<w:tblLayout\b[^>]*/?w:type=\"fixed\"", xml), "fixed table layout was not rendered")
    for width in ("1200", "2800", "1600"):
        assert_true(re.search(rf"<w:gridCol\b[^>]*/?w:w=\"{width}\"", xml), f"gridCol width {width} was not rendered")
        assert_true(re.search(rf"<w:tcW\b[^>]*/?w:w=\"{width}\"", xml), f"cell tcW width {width} was not rendered")
    counts = result["manifest"]["counts"]
    assert_true(counts.get("content_table_widths_rendered") == 1, f"table width render count was not recorded: {counts}")


@case
def script_generator_repairs_partial_table_column_widths() -> None:
    content = base_content(
        [
            {
                "role": "table",
                "table_rows": [
                    ["Metric", "Description", "Unit", "Value"],
                    ["A", "Longer explanatory text", "kg", "1"],
                ],
                "table_col_widths_twips": [1600, 0],
            }
        ],
        meta_tables=1,
    )
    result = run_generated_case("table_partial_column_widths", content)
    grid_widths = [int(value) for value in re.findall(r"<w:gridCol\b[^>]*w:w=\"(\d+)\"", result["xml"])]
    assert_true(len(grid_widths) >= 4, f"expected four generated grid columns, got {grid_widths}")
    first_table_widths = grid_widths[:4]
    assert_true(all(width > 0 for width in first_table_widths), f"partial column widths produced hidden zero-width columns: {first_table_widths}")
    assert_true(first_table_widths[0] == 1600, f"known positive source width should be preserved before fallback repair: {first_table_widths}")

    tc_widths = [int(value) for value in re.findall(r"<w:tcW\b[^>]*w:w=\"(\d+)\"", result["xml"])]
    assert_true(len(tc_widths) >= 4, f"expected four generated cell widths, got {tc_widths}")
    assert_true(all(width > 0 for width in tc_widths[:4]), f"partial column widths produced zero-width cells: {tc_widths[:4]}")


@case
def script_generator_renders_table_layout_details() -> None:
    content = base_content(
        [
            {
                "role": "table",
                "table_rows": [
                    ["Header A", "Header B"],
                    ["Body A", "Body B"],
                ],
                "table_row_heights_twips": [{"val": 480, "rule": "exact"}, {"val": 360, "rule": "atLeast"}],
                "table_repeat_header_rows": 2,
                "table_cell_margins_twips": {"top": 80, "left": 120, "bottom": 90, "right": 140},
                "table_cell_overrides": [
                    {
                        "row": 0,
                        "col": 0,
                        "v_align": "top",
                        "margins_twips": {"top": 40, "left": 60, "bottom": 40, "right": 60},
                    }
                ],
            }
        ],
        meta_tables=1,
    )
    result = run_generated_case("table_layout_details", content)
    xml = result["xml"]
    assert_true(
        re.search(r"<w:trHeight\b(?=[^>]*w:val=\"480\")(?=[^>]*w:hRule=\"exact\")", xml),
        "exact row height was not rendered",
    )
    assert_true(
        re.search(r"<w:trHeight\b(?=[^>]*w:val=\"360\")(?=[^>]*w:hRule=\"atLeast\")", xml),
        "atLeast row height was not rendered",
    )
    assert_true(len(re.findall(r"<w:tblHeader\b", xml)) >= 2, "repeat header rows were not rendered")
    for side, width in {"top": "80", "left": "120", "bottom": "90", "right": "140"}.items():
        assert_true(
            re.search(rf"<w:{side}\b(?=[^>]*w:w=\"{width}\")(?=[^>]*w:type=\"dxa\")", xml),
            f"default table cell margin {side}={width} was not rendered",
        )
    assert_true(re.search(r"<w:vAlign\b[^>]*w:val=\"top\"", xml), "cell vertical alignment override was not rendered")
    for side, width in {"top": "40", "left": "60", "bottom": "40", "right": "60"}.items():
        assert_true(
            re.search(rf"<w:{side}\b(?=[^>]*w:w=\"{width}\")(?=[^>]*w:type=\"dxa\")", xml),
            f"cell-specific margin {side}={width} was not rendered",
        )
    counts = result["manifest"]["counts"]
    assert_true(counts.get("content_table_row_heights_rendered") == 1, f"row height count missing: {counts}")
    assert_true(counts.get("content_table_repeat_header_rows_rendered") == 2, f"repeat-header count missing: {counts}")
    assert_true(counts.get("content_table_cell_margins_rendered") == 1, f"default margin count missing: {counts}")
    assert_true(counts.get("content_table_cell_overrides_rendered") == 1, f"cell override count missing: {counts}")


@case
def script_generator_allows_tall_table_rows_to_split_across_pages() -> None:
    tall_cell = "\n".join(f"Long evidence line {idx}: generated tables must not force this row onto one page." for idx in range(28))
    content = base_content(
        [
            {
                "role": "table",
                "table_rows": [
                    ["Metric", "Evidence"],
                    ["Pagination", tall_cell],
                ],
                "table_repeat_header_rows": 1,
            }
        ],
        meta_tables=1,
    )
    result = run_generated_case("table_tall_row_can_split", content)
    xml = result["xml"]
    assert_true(len(re.findall(r"<w:tblHeader\b", xml)) == 1, "repeat header row was not preserved")
    assert_true(
        len(re.findall(r"<w:cantSplit\b", xml)) == 1,
        "tall table body row should be allowed to split across pages while the header stays together",
    )


@case
def script_generator_keeps_structured_caption_with_landscape_table_section() -> None:
    content = base_content(
        [
            {"role": "table_caption", "text": "表 1 Explicit landscape caption"},
            {
                "role": "table",
                "table_rows": [
                    [f"Wide header {idx}" for idx in range(1, 10)],
                    [f"Wide body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": {
                    "orientation": "landscape",
                    "page_width_twips": 15840,
                    "page_height_twips": 12240,
                    "margins_twips": {"left": 1440, "right": 1440, "top": 1440, "bottom": 1440},
                },
            },
            "Portrait body after landscape table.",
        ],
        meta_tables=1,
    )
    result = run_generated_case("table_caption_landscape_section", content, base_format())
    root = ET.fromstring(result["xml"].encode("utf-8"))
    body = root.find(f".//{W_NS}body")
    assert_true(body is not None, "generated document body missing")
    children = list(body)

    def child_text(child: ET.Element) -> str:
        return "".join(node.text or "" for node in child.iter(f"{W_NS}t"))

    def has_sectpr(child: ET.Element) -> bool:
        return child.tag == f"{W_NS}sectPr" or child.find(f".//{W_NS}sectPr") is not None

    def has_landscape_section(child: ET.Element) -> bool:
        for page_size in child.iter(f"{W_NS}pgSz"):
            if page_size.attrib.get(f"{W_NS}orient") == "landscape":
                return True
        return False

    table_idx = next((idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "Wide header 1" in child_text(child)), -1)
    caption_idx = max(
        (idx for idx, child in enumerate(children[:table_idx]) if "Explicit landscape caption" in child_text(child)),
        default=-1,
    )
    after_idx = next((idx for idx, child in enumerate(children) if "Portrait body after landscape table." in child_text(child)), -1)
    assert_true(caption_idx >= 0 and table_idx >= 0 and after_idx >= 0, "caption/table/following body markers missing")

    previous_section_break = max((idx for idx, child in enumerate(children[:table_idx]) if has_sectpr(child)), default=-1)
    landscape_section_break = next((idx for idx, child in enumerate(children[table_idx + 1 :], start=table_idx + 1) if has_landscape_section(child)), -1)
    assert_true(
        previous_section_break < caption_idx < table_idx < landscape_section_break < after_idx,
        "structured table caption should be inside the same landscape section as the wide table, with following body restored to portrait",
    )


@case
def script_generator_auto_landscapes_plain_wide_tables() -> None:
    content = base_content(
        [
            {"role": "table_caption", "text": "表 1 Auto landscape wide table"},
            {
                "role": "table",
                "table_rows": [
                    [f"Auto header {idx}" for idx in range(1, 10)],
                    [f"Auto body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
            },
            "Portrait body after auto landscape table.",
        ],
        meta_tables=1,
    )
    result = run_generated_case("auto_landscape_plain_wide_table", content, base_format())
    assert_true(
        result["xml"].count('w:orient="landscape"') == 1,
        "plain overwide tables should be rendered in one generated landscape section",
    )
    assert_true(
        result["manifest"]["counts"].get("content_auto_landscape_table_sections_rendered") == 1,
        "plain overwide table should be counted as an auto-landscape section",
    )

    root = ET.fromstring(result["xml"].encode("utf-8"))
    body = root.find(f".//{W_NS}body")
    assert_true(body is not None, "generated document body missing")
    children = list(body)

    def child_text(child: ET.Element) -> str:
        return "".join(node.text or "" for node in child.iter(f"{W_NS}t"))

    def has_sectpr(child: ET.Element) -> bool:
        return child.tag == f"{W_NS}sectPr" or child.find(f".//{W_NS}sectPr") is not None

    def has_landscape_section(child: ET.Element) -> bool:
        return any(page_size.attrib.get(f"{W_NS}orient") == "landscape" for page_size in child.iter(f"{W_NS}pgSz"))

    table_idx = next((idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "Auto header 1" in child_text(child)), -1)
    caption_idx = max((idx for idx, child in enumerate(children[:table_idx]) if "Auto landscape wide table" in child_text(child)), default=-1)
    after_idx = next((idx for idx, child in enumerate(children) if "Portrait body after auto landscape table." in child_text(child)), -1)
    assert_true(caption_idx >= 0 and table_idx >= 0 and after_idx >= 0, "auto landscape caption/table/following body markers missing")

    previous_section_break = max((idx for idx, child in enumerate(children[:table_idx]) if has_sectpr(child)), default=-1)
    landscape_section_break = next((idx for idx, child in enumerate(children[table_idx + 1 :], start=table_idx + 1) if has_landscape_section(child)), -1)
    assert_true(
        previous_section_break < caption_idx < table_idx < landscape_section_break < after_idx,
        "auto-landscaped table caption should stay with the table, and following body should return to portrait",
    )


@case
def script_generator_repeats_first_row_for_auto_landscape_long_wide_tables() -> None:
    rows = [[f"Long wide header {idx}" for idx in range(1, 10)]]
    rows.extend([[f"Long wide row {row}-{col}" for col in range(1, 10)] for row in range(1, 34)])
    content = base_content(
        [
            {"role": "table_caption", "text": "表 1 Long auto landscape table"},
            {
                "role": "table",
                "table_rows": rows,
                "table_col_widths_twips": [1200] * 9,
            },
            "Portrait body after long auto landscape table.",
        ],
        meta_tables=1,
    )
    result = run_generated_case("auto_landscape_long_wide_table_header", content, base_format())
    assert_true(
        result["xml"].count('w:orient="landscape"') == 1,
        "long overwide table should still be auto-landscaped",
    )
    assert_true(
        len(re.findall(r"<w:tblHeader\b", result["xml"])) == 1,
        "long auto-landscaped tables should repeat the first row when crossing pages",
    )
    counts = result["manifest"]["counts"]
    assert_true(
        counts.get("content_table_repeat_header_rows_rendered") == 1,
        f"default repeat-header count missing for long auto-landscaped table: {counts}",
    )


@case
def script_generator_does_not_default_repeat_header_for_short_tables() -> None:
    content = base_content(
        [
            {
                "role": "table",
                "table_rows": [["Header A", "Header B"], ["Body A", "Body B"]],
            },
            {
                "role": "table",
                "table_rows": [["Explicit off A", "Explicit off B"], ["Body A", "Body B"]],
                "table_repeat_header_rows": 0,
            },
        ],
        meta_tables=2,
    )
    result = run_generated_case("short_tables_no_default_repeat_header", content, base_format())
    assert_true(
        "<w:tblHeader" not in result["xml"],
        "short tables or explicit repeat-header=0 tables should not gain default repeated headers",
    )
    counts = result["manifest"]["counts"]
    assert_true(
        counts.get("content_table_repeat_header_rows_rendered", 0) == 0,
        f"short/default-off tables should not increment repeat-header count: {counts}",
    )


@case
def script_generator_groups_adjacent_landscape_tables_with_short_note() -> None:
    landscape_setup = {
        "orientation": "landscape",
        "page_width_twips": 15840,
        "page_height_twips": 12240,
        "margins_twips": {"left": 1440, "right": 1440, "top": 1440, "bottom": 1440},
    }
    content = base_content(
        [
            {"role": "table_caption", "text": "表 1 First landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"First header {idx}" for idx in range(1, 10)],
                    [f"First body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            "Notes between the two wide tables should stay with the landscape table group.",
            {"role": "table_caption", "text": "表 2 Second landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"Second header {idx}" for idx in range(1, 10)],
                    [f"Second body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            "Portrait body after adjacent landscape tables.",
        ],
        meta_tables=2,
    )
    result = run_generated_case("adjacent_landscape_tables_short_note", content, base_format())
    assert_true(
        result["xml"].count('w:orient="landscape"') == 1,
        "adjacent landscape tables separated by a short note should share one landscape section",
    )
    root = ET.fromstring(result["xml"].encode("utf-8"))
    body = root.find(f".//{W_NS}body")
    assert_true(body is not None, "generated document body missing")
    children = list(body)

    def child_text(child: ET.Element) -> str:
        return "".join(node.text or "" for node in child.iter(f"{W_NS}t"))

    def has_sectpr(child: ET.Element) -> bool:
        return child.tag == f"{W_NS}sectPr" or child.find(f".//{W_NS}sectPr") is not None

    def has_landscape_section(child: ET.Element) -> bool:
        return any(page_size.attrib.get(f"{W_NS}orient") == "landscape" for page_size in child.iter(f"{W_NS}pgSz"))

    table1_idx = next((idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "First header 1" in child_text(child)), -1)
    note_idx = next((idx for idx, child in enumerate(children) if "Notes between the two wide tables" in child_text(child)), -1)
    table2_idx = next((idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "Second header 1" in child_text(child)), -1)
    after_idx = next((idx for idx, child in enumerate(children) if "Portrait body after adjacent landscape tables." in child_text(child)), -1)
    caption1_idx = max((idx for idx, child in enumerate(children[:table1_idx]) if "First landscape table" in child_text(child)), default=-1)
    caption2_idx = max((idx for idx, child in enumerate(children[:table2_idx]) if "Second landscape table" in child_text(child)), default=-1)
    assert_true(min(caption1_idx, table1_idx, note_idx, caption2_idx, table2_idx, after_idx) >= 0, "landscape group markers missing")

    previous_section_break = max((idx for idx, child in enumerate(children[:table1_idx]) if has_sectpr(child)), default=-1)
    landscape_section_break = next((idx for idx, child in enumerate(children[table2_idx + 1 :], start=table2_idx + 1) if has_landscape_section(child)), -1)
    intermediate_section_breaks = [idx for idx, child in enumerate(children[table1_idx + 1 : table2_idx], start=table1_idx + 1) if has_sectpr(child)]
    assert_true(not intermediate_section_breaks, f"landscape group was split by intermediate section breaks: {intermediate_section_breaks}")
    assert_true(
        previous_section_break < caption1_idx < table1_idx < note_idx < caption2_idx < table2_idx < landscape_section_break < after_idx,
        "adjacent landscape tables, their captions, and the short note should be inside one landscape section before portrait body resumes",
    )


@case
def script_generator_groups_adjacent_landscape_tables_with_rich_bridge_notes() -> None:
    landscape_setup = {
        "orientation": "landscape",
        "page_width_twips": 15840,
        "page_height_twips": 12240,
        "margins_twips": {"left": 1440, "right": 1440, "top": 1440, "bottom": 1440},
    }
    content = base_content(
        [
            {"role": "table_caption", "text": "表 1 Rich first landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"Rich first header {idx}" for idx in range(1, 10)],
                    [f"Rich first body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            {
                "role": "rich_text",
                "text": "Note: beta remains significant between the two wide tables.",
                "runs": [
                    {"type": "text", "text": "Note: "},
                    {"type": "math", "text": r"\beta", "math": [{"type": "inline", "text": r"\beta", "latex": r"\beta"}]},
                    {"type": "text", "text": " remains significant between the two wide tables."},
                ],
                "math": [{"type": "inline", "text": r"\beta", "latex": r"\beta"}],
            },
            "A second short bridge note should not force portrait-page recovery.",
            {"role": "table_caption", "text": "表 2 Rich second landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"Rich second header {idx}" for idx in range(1, 10)],
                    [f"Rich second body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            "Portrait body after rich bridge landscape tables.",
        ],
        meta_tables=2,
    )
    result = run_generated_case("adjacent_landscape_tables_rich_bridge_notes", content, base_format())
    assert_true(
        result["xml"].count('w:orient="landscape"') == 1,
        "adjacent landscape tables separated by bounded rich bridge notes should share one landscape section",
    )
    root = ET.fromstring(result["xml"].encode("utf-8"))
    body = root.find(f".//{W_NS}body")
    assert_true(body is not None, "generated document body missing")
    children = list(body)

    def child_text(child: ET.Element) -> str:
        return "".join(node.text or "" for node in child.iter(f"{W_NS}t"))

    def has_sectpr(child: ET.Element) -> bool:
        return child.tag == f"{W_NS}sectPr" or child.find(f".//{W_NS}sectPr") is not None

    def has_landscape_section(child: ET.Element) -> bool:
        return any(page_size.attrib.get(f"{W_NS}orient") == "landscape" for page_size in child.iter(f"{W_NS}pgSz"))

    table1_idx = next((idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "Rich first header 1" in child_text(child)), -1)
    rich_note_idx = next((idx for idx, child in enumerate(children) if "remains significant between the two wide tables" in child_text(child)), -1)
    second_note_idx = next((idx for idx, child in enumerate(children) if "second short bridge note" in child_text(child)), -1)
    table2_idx = next((idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "Rich second header 1" in child_text(child)), -1)
    after_idx = next((idx for idx, child in enumerate(children) if "Portrait body after rich bridge landscape tables." in child_text(child)), -1)
    assert_true(min(table1_idx, rich_note_idx, second_note_idx, table2_idx, after_idx) >= 0, "rich landscape group markers missing")

    landscape_section_break = next((idx for idx, child in enumerate(children[table2_idx + 1 :], start=table2_idx + 1) if has_landscape_section(child)), -1)
    intermediate_section_breaks = [idx for idx, child in enumerate(children[table1_idx + 1 : table2_idx], start=table1_idx + 1) if has_sectpr(child)]
    assert_true(not intermediate_section_breaks, f"rich bridge landscape group was split by intermediate section breaks: {intermediate_section_breaks}")
    assert_true(
        table1_idx < rich_note_idx < second_note_idx < table2_idx < landscape_section_break < after_idx,
        "bounded rich bridge notes should stay between adjacent landscape tables before portrait body resumes",
    )


@case
def script_generator_splits_landscape_tables_around_display_math_bridge() -> None:
    landscape_setup = {
        "orientation": "landscape",
        "page_width_twips": 15840,
        "page_height_twips": 12240,
        "margins_twips": {"left": 1440, "right": 1440, "top": 1440, "bottom": 1440},
    }
    content = base_content(
        [
            {"role": "table_caption", "text": "表 1 Display bridge first landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"Display first header {idx}" for idx in range(1, 10)],
                    [f"Display first body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            {
                "role": "rich_text",
                "text": "Displayed equation bridge should resume portrait flow.",
                "runs": [
                    {"type": "text", "text": "Displayed equation bridge should resume portrait flow: "},
                    {"type": "math", "text": r"E=mc^2", "math": {"type": "display", "text": r"E=mc^2", "latex": r"E=mc^2"}},
                ],
                "math": {"type": "display", "text": r"E=mc^2", "latex": r"E=mc^2"},
            },
            {"role": "table_caption", "text": "表 2 Display bridge second landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"Display second header {idx}" for idx in range(1, 10)],
                    [f"Display second body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            "Portrait body after display bridge landscape tables.",
        ],
        meta_tables=2,
    )
    result = run_generated_case("landscape_tables_display_math_bridge_split", content, base_format())
    assert_true(
        result["xml"].count('w:orient="landscape"') == 2,
        "display-formula bridge content should split adjacent landscape tables into separate landscape sections",
    )

    root = ET.fromstring(result["xml"].encode("utf-8"))
    body = root.find(f".//{W_NS}body")
    assert_true(body is not None, "generated document body missing")
    children = list(body)

    def child_text(child: ET.Element) -> str:
        return "".join(node.text or "" for node in child.iter(f"{W_NS}t"))

    def has_landscape_section(child: ET.Element) -> bool:
        return any(page_size.attrib.get(f"{W_NS}orient") == "landscape" for page_size in child.iter(f"{W_NS}pgSz"))

    def has_portrait_section(child: ET.Element) -> bool:
        return any(page_size.attrib.get(f"{W_NS}orient") != "landscape" for page_size in child.iter(f"{W_NS}pgSz"))

    table1_idx = next(
        (idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "Display first header 1" in child_text(child)),
        -1,
    )
    formula_idx = next(
        (idx for idx, child in enumerate(children) if "Displayed equation bridge should resume portrait flow" in child_text(child)),
        -1,
    )
    table2_idx = next(
        (idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "Display second header 1" in child_text(child)),
        -1,
    )
    after_idx = next(
        (idx for idx, child in enumerate(children) if "Portrait body after display bridge landscape tables." in child_text(child)),
        -1,
    )
    assert_true(min(table1_idx, formula_idx, table2_idx, after_idx) >= 0, "display-math bridge landscape markers missing")

    first_landscape_section_break = next(
        (idx for idx, child in enumerate(children[table1_idx + 1 : formula_idx], start=table1_idx + 1) if has_landscape_section(child)),
        -1,
    )
    second_landscape_section_break = next(
        (idx for idx, child in enumerate(children[table2_idx + 1 : after_idx], start=table2_idx + 1) if has_landscape_section(child)),
        -1,
    )
    portrait_bridge_section_break = next(
        (idx for idx, child in enumerate(children[formula_idx + 1 : table2_idx], start=formula_idx + 1) if has_portrait_section(child)),
        -1,
    )
    assert_true(
        table1_idx < first_landscape_section_break < formula_idx < portrait_bridge_section_break < table2_idx < second_landscape_section_break < after_idx,
        "display-formula bridge should be closed as a portrait section between the two landscape table sections",
    )


@case
def script_generator_splits_landscape_tables_around_list_bridge() -> None:
    landscape_setup = {
        "orientation": "landscape",
        "page_width_twips": 15840,
        "page_height_twips": 12240,
        "margins_twips": {"left": 1440, "right": 1440, "top": 1440, "bottom": 1440},
    }
    bullet_bridge = "- Review item: this checklist paragraph is body flow between wide tables and must not be folded into the surrounding landscape section."
    content = base_content(
        [
            {"role": "table_caption", "text": "表 1 List bridge first landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"List first header {idx}" for idx in range(1, 10)],
                    [f"List first body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            bullet_bridge,
            {"role": "table_caption", "text": "表 2 List bridge second landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"List second header {idx}" for idx in range(1, 10)],
                    [f"List second body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            "Portrait body after list bridge landscape tables.",
        ],
        meta_tables=2,
    )
    result = run_generated_case("landscape_tables_list_bridge_split", content, base_format())
    assert_true(
        result["xml"].count('w:orient="landscape"') == 2,
        "list-like bridge content should split adjacent landscape tables into separate landscape sections",
    )

    root = ET.fromstring(result["xml"].encode("utf-8"))
    body = root.find(f".//{W_NS}body")
    assert_true(body is not None, "generated document body missing")
    children = list(body)

    def child_text(child: ET.Element) -> str:
        return "".join(node.text or "" for node in child.iter(f"{W_NS}t"))

    def has_landscape_section(child: ET.Element) -> bool:
        return any(page_size.attrib.get(f"{W_NS}orient") == "landscape" for page_size in child.iter(f"{W_NS}pgSz"))

    def has_portrait_section(child: ET.Element) -> bool:
        return any(page_size.attrib.get(f"{W_NS}orient") != "landscape" for page_size in child.iter(f"{W_NS}pgSz"))

    table1_idx = next(
        (idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "List first header 1" in child_text(child)),
        -1,
    )
    bullet_idx = next((idx for idx, child in enumerate(children) if "Review item: this checklist paragraph" in child_text(child)), -1)
    table2_idx = next(
        (idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "List second header 1" in child_text(child)),
        -1,
    )
    after_idx = next(
        (idx for idx, child in enumerate(children) if "Portrait body after list bridge landscape tables." in child_text(child)),
        -1,
    )
    assert_true(min(table1_idx, bullet_idx, table2_idx, after_idx) >= 0, "list bridge landscape markers missing")

    first_landscape_section_break = next(
        (idx for idx, child in enumerate(children[table1_idx + 1 : bullet_idx], start=table1_idx + 1) if has_landscape_section(child)),
        -1,
    )
    portrait_bridge_section_break = next(
        (idx for idx, child in enumerate(children[bullet_idx + 1 : table2_idx], start=bullet_idx + 1) if has_portrait_section(child)),
        -1,
    )
    second_landscape_section_break = next(
        (idx for idx, child in enumerate(children[table2_idx + 1 : after_idx], start=table2_idx + 1) if has_landscape_section(child)),
        -1,
    )
    assert_true(
        table1_idx < first_landscape_section_break < bullet_idx < portrait_bridge_section_break < table2_idx < second_landscape_section_break < after_idx,
        "list-like bridge content should be closed as a portrait section between the two landscape table sections",
    )


@case
def script_generator_splits_landscape_tables_around_chinese_numbered_bridge() -> None:
    landscape_setup = {
        "orientation": "landscape",
        "page_width_twips": 15840,
        "page_height_twips": 12240,
        "margins_twips": {"left": 1440, "right": 1440, "top": 1440, "bottom": 1440},
    }
    chinese_numbered_bridge = "（1）这是两个宽表之间的正文编号说明，应当回到纵向正文流。"
    content = base_content(
        [
            {"role": "table_caption", "text": "表 1 Chinese list bridge first landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"Chinese list first header {idx}" for idx in range(1, 10)],
                    [f"Chinese list first body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            chinese_numbered_bridge,
            {"role": "table_caption", "text": "表 2 Chinese list bridge second landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"Chinese list second header {idx}" for idx in range(1, 10)],
                    [f"Chinese list second body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            "Portrait body after Chinese numbered bridge landscape tables.",
        ],
        meta_tables=2,
    )
    result = run_generated_case("landscape_tables_chinese_numbered_bridge_split", content, base_format())
    assert_true(
        result["xml"].count('w:orient="landscape"') == 2,
        "Chinese-numbered bridge content should split adjacent landscape tables into separate landscape sections",
    )

    root = ET.fromstring(result["xml"].encode("utf-8"))
    body = root.find(f".//{W_NS}body")
    assert_true(body is not None, "generated document body missing")
    children = list(body)

    def child_text(child: ET.Element) -> str:
        return "".join(node.text or "" for node in child.iter(f"{W_NS}t"))

    def has_landscape_section(child: ET.Element) -> bool:
        return any(page_size.attrib.get(f"{W_NS}orient") == "landscape" for page_size in child.iter(f"{W_NS}pgSz"))

    def has_portrait_section(child: ET.Element) -> bool:
        return any(page_size.attrib.get(f"{W_NS}orient") != "landscape" for page_size in child.iter(f"{W_NS}pgSz"))

    table1_idx = next(
        (idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "Chinese list first header 1" in child_text(child)),
        -1,
    )
    bridge_idx = next((idx for idx, child in enumerate(children) if "两个宽表之间的正文编号说明" in child_text(child)), -1)
    table2_idx = next(
        (idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "Chinese list second header 1" in child_text(child)),
        -1,
    )
    after_idx = next(
        (idx for idx, child in enumerate(children) if "Portrait body after Chinese numbered bridge landscape tables." in child_text(child)),
        -1,
    )
    assert_true(min(table1_idx, bridge_idx, table2_idx, after_idx) >= 0, "Chinese numbered bridge landscape markers missing")

    first_landscape_section_break = next(
        (idx for idx, child in enumerate(children[table1_idx + 1 : bridge_idx], start=table1_idx + 1) if has_landscape_section(child)),
        -1,
    )
    portrait_bridge_section_break = next(
        (idx for idx, child in enumerate(children[bridge_idx + 1 : table2_idx], start=bridge_idx + 1) if has_portrait_section(child)),
        -1,
    )
    second_landscape_section_break = next(
        (idx for idx, child in enumerate(children[table2_idx + 1 : after_idx], start=table2_idx + 1) if has_landscape_section(child)),
        -1,
    )
    assert_true(
        table1_idx < first_landscape_section_break < bridge_idx < portrait_bridge_section_break < table2_idx < second_landscape_section_break < after_idx,
        "Chinese-numbered bridge content should be closed as a portrait section between the two landscape table sections",
    )


@case
def script_generator_splits_landscape_tables_around_bracketed_number_bridge() -> None:
    landscape_setup = {
        "orientation": "landscape",
        "page_width_twips": 15840,
        "page_height_twips": 12240,
        "margins_twips": {"left": 1440, "right": 1440, "top": 1440, "bottom": 1440},
    }
    bracketed_bridge = "[1] This numbered body item belongs to the portrait flow between two wide tables."
    content = base_content(
        [
            {"role": "table_caption", "text": "表 1 Bracketed list bridge first landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"Bracketed list first header {idx}" for idx in range(1, 10)],
                    [f"Bracketed list first body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            bracketed_bridge,
            {"role": "table_caption", "text": "表 2 Bracketed list bridge second landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"Bracketed list second header {idx}" for idx in range(1, 10)],
                    [f"Bracketed list second body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            "Portrait body after bracketed numbered bridge landscape tables.",
        ],
        meta_tables=2,
    )
    result = run_generated_case("landscape_tables_bracketed_number_bridge_split", content, base_format())
    assert_true(
        result["xml"].count('w:orient="landscape"') == 2,
        "bracketed-number bridge content should split adjacent landscape tables into separate landscape sections",
    )

    root = ET.fromstring(result["xml"].encode("utf-8"))
    body = root.find(f".//{W_NS}body")
    assert_true(body is not None, "generated document body missing")
    children = list(body)

    def child_text(child: ET.Element) -> str:
        return "".join(node.text or "" for node in child.iter(f"{W_NS}t"))

    def has_landscape_section(child: ET.Element) -> bool:
        return any(page_size.attrib.get(f"{W_NS}orient") == "landscape" for page_size in child.iter(f"{W_NS}pgSz"))

    def has_portrait_section(child: ET.Element) -> bool:
        return any(page_size.attrib.get(f"{W_NS}orient") != "landscape" for page_size in child.iter(f"{W_NS}pgSz"))

    table1_idx = next(
        (idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "Bracketed list first header 1" in child_text(child)),
        -1,
    )
    bridge_idx = next((idx for idx, child in enumerate(children) if "numbered body item belongs to the portrait flow" in child_text(child)), -1)
    table2_idx = next(
        (idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "Bracketed list second header 1" in child_text(child)),
        -1,
    )
    after_idx = next(
        (idx for idx, child in enumerate(children) if "Portrait body after bracketed numbered bridge landscape tables." in child_text(child)),
        -1,
    )
    assert_true(min(table1_idx, bridge_idx, table2_idx, after_idx) >= 0, "bracketed numbered bridge landscape markers missing")

    first_landscape_section_break = next(
        (idx for idx, child in enumerate(children[table1_idx + 1 : bridge_idx], start=table1_idx + 1) if has_landscape_section(child)),
        -1,
    )
    portrait_bridge_section_break = next(
        (idx for idx, child in enumerate(children[bridge_idx + 1 : table2_idx], start=bridge_idx + 1) if has_portrait_section(child)),
        -1,
    )
    second_landscape_section_break = next(
        (idx for idx, child in enumerate(children[table2_idx + 1 : after_idx], start=table2_idx + 1) if has_landscape_section(child)),
        -1,
    )
    assert_true(
        table1_idx < first_landscape_section_break < bridge_idx < portrait_bridge_section_break < table2_idx < second_landscape_section_break < after_idx,
        "bracketed-number bridge content should be closed as a portrait section between the two landscape table sections",
    )


@case
def script_generator_splits_landscape_tables_around_roman_number_bridge() -> None:
    landscape_setup = {
        "orientation": "landscape",
        "page_width_twips": 15840,
        "page_height_twips": 12240,
        "margins_twips": {"left": 1440, "right": 1440, "top": 1440, "bottom": 1440},
    }
    roman_bridge = "(iv) This roman-numbered body item belongs to the portrait flow between two wide tables."
    content = base_content(
        [
            {"role": "table_caption", "text": "表 1 Roman list bridge first landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"Roman list first header {idx}" for idx in range(1, 10)],
                    [f"Roman list first body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            roman_bridge,
            {"role": "table_caption", "text": "表 2 Roman list bridge second landscape table"},
            {
                "role": "table",
                "table_rows": [
                    [f"Roman list second header {idx}" for idx in range(1, 10)],
                    [f"Roman list second body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": landscape_setup,
            },
            "Portrait body after roman numbered bridge landscape tables.",
        ],
        meta_tables=2,
    )
    result = run_generated_case("landscape_tables_roman_number_bridge_split", content, base_format())
    assert_true(
        result["xml"].count('w:orient="landscape"') == 2,
        "roman-number bridge content should split adjacent landscape tables into separate landscape sections",
    )

    root = ET.fromstring(result["xml"].encode("utf-8"))
    body = root.find(f".//{W_NS}body")
    assert_true(body is not None, "generated document body missing")
    children = list(body)

    def child_text(child: ET.Element) -> str:
        return "".join(node.text or "" for node in child.iter(f"{W_NS}t"))

    def has_landscape_section(child: ET.Element) -> bool:
        return any(page_size.attrib.get(f"{W_NS}orient") == "landscape" for page_size in child.iter(f"{W_NS}pgSz"))

    def has_portrait_section(child: ET.Element) -> bool:
        return any(page_size.attrib.get(f"{W_NS}orient") != "landscape" for page_size in child.iter(f"{W_NS}pgSz"))

    table1_idx = next(
        (idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "Roman list first header 1" in child_text(child)),
        -1,
    )
    bridge_idx = next((idx for idx, child in enumerate(children) if "roman-numbered body item belongs to the portrait flow" in child_text(child)), -1)
    table2_idx = next(
        (idx for idx, child in enumerate(children) if child.tag == f"{W_NS}tbl" and "Roman list second header 1" in child_text(child)),
        -1,
    )
    after_idx = next(
        (idx for idx, child in enumerate(children) if "Portrait body after roman numbered bridge landscape tables." in child_text(child)),
        -1,
    )
    assert_true(min(table1_idx, bridge_idx, table2_idx, after_idx) >= 0, "roman numbered bridge landscape markers missing")

    first_landscape_section_break = next(
        (idx for idx, child in enumerate(children[table1_idx + 1 : bridge_idx], start=table1_idx + 1) if has_landscape_section(child)),
        -1,
    )
    portrait_bridge_section_break = next(
        (idx for idx, child in enumerate(children[bridge_idx + 1 : table2_idx], start=bridge_idx + 1) if has_portrait_section(child)),
        -1,
    )
    second_landscape_section_break = next(
        (idx for idx, child in enumerate(children[table2_idx + 1 : after_idx], start=table2_idx + 1) if has_landscape_section(child)),
        -1,
    )
    assert_true(
        table1_idx < first_landscape_section_break < bridge_idx < portrait_bridge_section_break < table2_idx < second_landscape_section_break < after_idx,
        "roman-number bridge content should be closed as a portrait section between the two landscape table sections",
    )


@case
def script_generator_does_not_group_landscape_tables_with_different_page_setups() -> None:
    first_setup = {
        "orientation": "landscape",
        "page_width_twips": 15840,
        "page_height_twips": 12240,
        "margins_twips": {"left": 1440, "right": 1440, "top": 1440, "bottom": 1440},
    }
    second_setup = {
        "orientation": "landscape",
        "page_width_twips": 16840,
        "page_height_twips": 11900,
        "margins_twips": {"left": 900, "right": 900, "top": 1200, "bottom": 1200},
    }
    content = base_content(
        [
            {"role": "table_caption", "text": "表 1 First source-section table"},
            {
                "role": "table",
                "table_rows": [
                    [f"First setup header {idx}" for idx in range(1, 10)],
                    [f"First setup body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1200] * 9,
                "source_section_page_setup": first_setup,
            },
            "Short note between source sections should not force incompatible page setups together.",
            {"role": "table_caption", "text": "表 2 Second source-section table"},
            {
                "role": "table",
                "table_rows": [
                    [f"Second setup header {idx}" for idx in range(1, 10)],
                    [f"Second setup body {idx}" for idx in range(1, 10)],
                ],
                "table_col_widths_twips": [1300] * 9,
                "source_section_page_setup": second_setup,
            },
            "Portrait body after different source-section tables.",
        ],
        meta_tables=2,
    )
    result = run_generated_case("different_landscape_page_setups", content, base_format())
    assert_true(
        result["xml"].count('w:orient="landscape"') == 2,
        "landscape tables with different source page setups should keep separate landscape sections",
    )


@case
def script_generator_renders_table_border_details() -> None:
    content = base_content(
        [
            {
                "role": "table",
                "table_rows": [
                    ["Header A", "Header B"],
                    ["Body A", "Body B"],
                ],
                "table_col_widths_twips": [1600, 2400],
                "table_borders": {
                    "top": {"val": "double", "sz": "12", "color": "4472C4", "space": "0"},
                    "insideH": {"val": "single", "sz": "4", "color": "808080", "space": "0"},
                },
                "table_cell_margins_twips": {"top": 80, "left": 120, "bottom": 90, "right": 140},
                "table_cell_overrides": [
                    {
                        "row": 0,
                        "col": 0,
                        "v_align": "top",
                        "margins_twips": {"top": 40, "left": 60, "bottom": 40, "right": 60},
                        "borders": {
                            "bottom": {"val": "dashed", "sz": "6", "color": "C00000", "space": "0"},
                            "right": {"val": "nil", "sz": "0", "color": "000000", "space": "0"},
                        },
                    }
                ],
            }
        ],
        meta_tables=1,
    )
    result = run_generated_case("table_border_details", content)
    xml = result["xml"]
    assert_true(
        re.search(r"<w:tblBorders>.*<w:top\b(?=[^>]*w:val=\"double\")(?=[^>]*w:sz=\"12\")(?=[^>]*w:color=\"4472C4\")", xml, re.S),
        "table top border was not rendered",
    )
    assert_true(
        re.search(r"<w:insideH\b(?=[^>]*w:val=\"single\")(?=[^>]*w:sz=\"4\")(?=[^>]*w:color=\"808080\")", xml),
        "table insideH border was not rendered",
    )
    assert_true(
        re.search(r"<w:bottom\b(?=[^>]*w:val=\"dashed\")(?=[^>]*w:sz=\"6\")(?=[^>]*w:color=\"C00000\")", xml),
        "cell bottom border was not rendered",
    )
    assert_true(
        re.search(r"<w:right\b(?=[^>]*w:val=\"nil\")(?=[^>]*w:sz=\"0\")(?=[^>]*w:color=\"000000\")", xml),
        "cell nil right border was not rendered",
    )
    assert_true(
        len(re.findall(r"<w:tcBorders>", xml)) == 1,
        "default three-line borders should not overwrite explicit source borders",
    )
    tbl_pr_children = _first_property_child_names(xml, "tblPr")
    for child_name in ("tblBorders", "tblLayout", "tblCellMar"):
        assert_true(child_name in tbl_pr_children, f"{child_name} was missing from tblPr: {tbl_pr_children}")
    assert_true(
        tbl_pr_children.index("tblBorders") < tbl_pr_children.index("tblLayout") < tbl_pr_children.index("tblCellMar"),
        f"table property order is invalid: {tbl_pr_children}",
    )
    tc_pr_children = _first_property_child_names(xml, "tcPr")
    for child_name in ("tcW", "tcBorders", "tcMar", "vAlign"):
        assert_true(child_name in tc_pr_children, f"{child_name} was missing from tcPr: {tc_pr_children}")
    assert_true(
        tc_pr_children.index("tcW") < tc_pr_children.index("tcBorders") < tc_pr_children.index("tcMar") < tc_pr_children.index("vAlign"),
        f"cell property order is invalid: {tc_pr_children}",
    )
    counts = result["manifest"]["counts"]
    assert_true(counts.get("content_table_borders_rendered") == 1, f"table border count missing: {counts}")
    assert_true(counts.get("content_table_cell_borders_rendered") == 1, f"cell border count missing: {counts}")


@case
def script_generator_keeps_borders_on_merged_table_cells() -> None:
    content = base_content(
        [
            {
                "role": "table",
                "table_rows": [
                    ["Merged header", "", "Score"],
                    ["Body A", "Body B", "1"],
                ],
                "table_merges": [
                    {"row": 0, "col": 0, "rowspan": 1, "colspan": 2},
                ],
                "table_cell_overrides": [
                    {
                        "row": 0,
                        "col": 0,
                        "borders": {
                            "top": {"val": "double", "sz": "12", "color": "4472C4"},
                            "bottom": {"val": "dashed", "sz": "6", "color": "C00000"},
                        },
                    }
                ],
            }
        ],
        meta_tables=1,
    )
    result = run_generated_case("merged_table_cell_borders", content)
    xml = result["xml"]
    assert_true(re.search(r"<w:gridSpan\b[^>]*/?w:val=\"2\"", xml), "horizontal merge gridSpan was not rendered")
    assert_true(
        re.search(r"<w:top\b(?=[^>]*w:val=\"double\")(?=[^>]*w:sz=\"12\")(?=[^>]*w:color=\"4472C4\")", xml),
        "merged cell top border was lost during merge",
    )
    assert_true(
        re.search(r"<w:bottom\b(?=[^>]*w:val=\"dashed\")(?=[^>]*w:sz=\"6\")(?=[^>]*w:color=\"C00000\")", xml),
        "merged cell bottom border was lost during merge",
    )


@case
def script_generator_cell_borders_do_not_trigger_default_table_borders() -> None:
    content = base_content(
        [
            {
                "role": "table",
                "table_rows": [
                    ["Header A", "Header B"],
                    ["Body A", "Body B"],
                ],
                "table_cell_overrides": [
                    {
                        "row": 0,
                        "col": 0,
                        "borders": {
                            "bottom": {"val": "dashed", "sz": "6", "color": "C00000"},
                        },
                    }
                ],
            }
        ],
        meta_tables=1,
    )
    result = run_generated_case("cell_border_overrides_default_table", content)
    xml = result["xml"]
    assert_true(
        len(re.findall(r"<w:tcBorders>", xml)) == 1,
        "cell-level source borders should prevent generated default three-line borders",
    )
    assert_true(
        re.search(r"<w:bottom\b(?=[^>]*w:val=\"dashed\")(?=[^>]*w:sz=\"6\")(?=[^>]*w:color=\"C00000\")", xml),
        "explicit cell bottom border was not rendered",
    )
    assert_true(
        not re.search(r"<w:top\b(?=[^>]*w:val=\"single\")(?=[^>]*w:sz=\"12\")(?=[^>]*w:color=\"000000\")", xml),
        "generated default table top border polluted a source cell-border-only table",
    )


@case
def script_generator_renders_two_level_nested_table_inside_parent_cell() -> None:
    img_src = new_workdir("nested_table_cell_image_src")
    write_sample_png(img_src / "top.png", width=180, height=120)
    write_sample_png(img_src / "nested.png", width=160, height=120)
    content = base_content(
        [
            {
                "role": "image",
                "image": "top.png",
                "caption": "Fig. 1 Top-level image",
            },
            {
                "role": "table",
                "table_rows": [
                    ["Outer A", "Outer B"],
                    ["Nested before\nNested after", "Outer D"],
                ],
                "table_cell_items": [
                    {
                        "row": 1,
                        "col": 0,
                        "items": [
                            {
                                "role": "table",
                                "location": "nested_table_cell",
                                "after_paragraph_index": 1,
                                "table_rows": [["Nested A", "Nested B"], ["Nested C", "Nested D"]],
                                "table_cell_items": [
                                    {
                                        "row": 1,
                                        "col": 1,
                                        "items": [
                                            {
                                                "role": "table",
                                                "location": "nested_table_cell",
                                                "table_rows": [["Deeper A", "Deeper B"]],
                                                "table_cell_items": [
                                                    {
                                                        "row": 0,
                                                        "col": 1,
                                                        "items": [
                                                            {
                                                                "role": "image",
                                                                "image": "nested.png",
                                                                "location": "table_cell",
                                                                "after_paragraph_index": 1,
                                                            }
                                                        ],
                                                    }
                                                ],
                                            }
                                        ],
                                    }
                                ],
                            }
                        ],
                    }
                ],
            }
        ],
        meta_tables=3,
    )
    content["_meta"]["images_extracted"] = 2
    content["_meta"]["images_dir"] = str(img_src)
    requirements = build_requirements(base_format(), content)
    assert_true(
        requirements.get("expected_counts", {}).get("images") == 2,
        f"nested table-cell image was not included in expected image counts: {requirements}",
    )
    result = run_generated_case("two_level_nested_table_inside_cell", content)
    xml = result["xml"]
    assert_true(len(re.findall(r"<w:tbl>", xml)) >= 3, "nested tables were not rendered as Word tables")
    assert_true("Nested A" in xml and "Nested D" in xml, "nested table text was missing from generated XML")
    assert_true("Deeper A" in xml and "Deeper B" in xml, "second-level nested table text was missing from generated XML")
    assert_true(xml.count("<w:drawing>") == 2, "top-level plus nested table-cell images should render exactly once each")
    assert_true(
        xml.find("Nested before") < xml.find("Nested A") < xml.find("Nested after"),
        "nested table was not rendered between the parent cell's before/after paragraphs",
    )
    assert_true(
        re.search(r"<w:tc\b[\s\S]*Nested before[\s\S]*<w:tbl>[\s\S]*Nested A[\s\S]*Nested D[\s\S]*<w:tbl>[\s\S]*Deeper A[\s\S]*Deeper B[\s\S]*</w:tbl>[\s\S]*Nested after", xml),
        "nested table was not rendered inside the parent cell in source order",
    )
    table_xmls = re.findall(r"<w:tbl\b.*?</w:tbl>", xml, flags=re.S)
    table_drawings = sum(table_xml.count("<w:drawing>") for table_xml in table_xmls)
    assert_true(table_drawings >= 1, "nested table-cell image rendered outside generated Word tables")
    assert_true(
        re.search(r"Deeper B[\s\S]*<w:drawing>", xml),
        "nested table-cell image was not rendered after its source cell text",
    )
    counts = result["manifest"]["counts"]
    assert_true(counts.get("content_nested_tables_rendered") == 2, f"nested table render count missing: {counts}")
    assert_true(counts.get("content_images_rendered") == 2, f"nested table-cell image render count missing: {counts}")


@case
def numbered_english_heading_is_not_front_title() -> None:
    cnt = {
        "sections": [
            {
                "heading": "1 Introduction",
                "level": 1,
                "role": "body",
                "paragraphs": ["Body"],
                "images": [],
            }
        ]
    }
    front = _front_matter_sections(cnt)
    assert_true(front.get("en_title") in ("", None), "numbered English body heading became English title")
    assert_true(0 not in front.get("front_indices", set()), "numbered English body heading was marked as front matter")


@case
def english_appendix_heading_is_not_front_title() -> None:
    cnt = {
        "sections": [
            {"heading": "1 Introduction", "level": 1, "role": "body", "paragraphs": ["Body"], "images": []},
            {"heading": "Appendix A Commands", "level": 1, "role": "appendix", "paragraphs": ["python run_pipeline.py"], "images": []},
        ]
    }
    front = _front_matter_sections(cnt)
    assert_true(front.get("en_title") in ("", None), "English appendix heading became English title")
    assert_true(1 not in front.get("front_indices", set()), "English appendix was marked as front matter")


@case
def script_generator_section_order_preserves_content_while_sorting_subsections() -> None:
    sections = [
        {"heading": "第1章 Intro", "level": 1, "paragraphs": ["chapter"], "images": []},
        {"heading": "1.2 Later", "level": 2, "paragraphs": ["later"], "images": []},
        {"heading": "1.2.2 Detail B", "level": 3, "paragraphs": ["b"], "images": []},
        {"heading": "1.2.1 Detail A", "level": 3, "paragraphs": ["a"], "images": []},
        {"heading": "1.1 Earlier", "level": 2, "paragraphs": ["earlier"], "images": []},
        {"heading": "第2章 Methods", "level": 1, "paragraphs": ["next"], "images": []},
    ]
    ordered = _normalize_numbered_section_order(sections)
    headings = [sec["heading"] for sec in ordered]
    assert_true(headings[:5] == ["第1章 Intro", "1.1 Earlier", "1.2 Later", "1.2.1 Detail A", "1.2.2 Detail B"], "numbered sections were not sorted safely")
    assert_true(ordered[1]["paragraphs"] == ["earlier"], "section content moved away from its heading")
    assert_true(headings[-1] == "第2章 Methods", "next numbered chapter was not preserved")


@case
def script_generator_template_rules_extract_page_header_and_flags() -> None:
    fmt = {
        "sections": [
            {
                "page_width_cm": 21.0,
                "page_height_cm": 29.7,
                "margin_top_cm": 2.5,
                "margin_bottom_cm": 2.4,
                "margin_left_cm": 2.8,
                "margin_right_cm": 2.2,
                "header": [
                    {
                        "text": "Thesis Header",
                        "alignment": "DEFAULT",
                        "runs": [{"text": "Thesis Header", "font": "Times New Roman", "size_pt": 9, "italic": True}],
                    }
                ],
            }
        ],
        "paragraphs": [
            {"text": "图1 系统结构"},
            {"text": "中文摘要不分自然段，英文题目要求大写字母，公式居中并编号，英文参考文献左对齐，参考文献悬挂缩进2字符。"},
        ],
    }
    page = _extract_page_and_header(fmt)
    rules = _infer_template_rules(fmt)
    assert_true(page["mt"] == 2.5 and page["mr"] == 2.2, "page margin extraction changed")
    assert_true(page["header"]["align"] == "CENTER", "DEFAULT header alignment was not normalized")
    assert_true(rules["cn_abstract_single_paragraph"] is True, "Chinese abstract rule was not detected")
    assert_true(rules["formula_center"] is True and rules["formula_numbered"] is True, "formula rules were not detected")
    assert_true(rules["reference_hanging_chars"] == 2.0, "reference hanging indent rule changed")


@case
def script_generator_style_profiles_apply_template_text_rules() -> None:
    body_text = "这是一个足够长的中文正文样本，用于模拟模板正文段落的自然排版特征，确保样式推断不会只依赖说明文字。"
    fmt = {
        "style_profiles": {},
        "paragraphs": [
            {
                "text": "论文正文使用宋体小四号，固定值28磅，首行缩进2字符，两端对齐。",
                "runs": [{"text": "论文正文使用宋体小四号，固定值28磅，首行缩进2字符，两端对齐。", "font": "宋体", "size_pt": 12}],
                "align": "LEFT",
            },
            {
                "text": "第1章 一级标题黑体三号居中加粗",
                "runs": [{"text": "第1章 一级标题黑体三号居中加粗", "font": "黑体", "size_pt": 16, "bold": True}],
                "align": "CENTER",
            },
            {
                "text": "图标题宋体五号居中",
                "runs": [{"text": "图标题宋体五号居中", "font": "宋体", "size_pt": 10.5}],
                "align": "CENTER",
            },
            {
                "text": "参考文献中文使用宋体小四号，悬挂缩进2字符。",
                "runs": [{"text": "参考文献中文使用宋体小四号，悬挂缩进2字符。", "font": "宋体", "size_pt": 12}],
                "align": "LEFT",
            },
            {
                "text": body_text,
                "runs": [{"text": body_text, "font": "宋体", "size_pt": 12}],
                "align": "JUSTIFY",
                "ls": 1.5,
                "indent": 0.74,
            },
        ],
    }
    profiles = _infer_style_profiles(fmt)
    assert_true(profiles["body"]["font"] == "宋体", "body font rule was not applied")
    assert_true(profiles["body"]["line_spacing_fixed_pt"] == 28.0, "fixed body line spacing rule was not applied")
    assert_true(profiles["body"]["align"] == "JUSTIFY", "body alignment rule was not normalized")
    assert_true(profiles["body"]["first_indent_cm"] > 0, "body first-line indent rule was not applied")
    assert_true(profiles["h1"]["font"] == "黑体" and profiles["h1"]["align"] == "CENTER", "h1 style inference changed")
    assert_true(profiles["figure_caption"]["size"] == 10.5 and profiles["figure_caption"]["align"] == "CENTER", "figure caption style rule changed")
    assert_true(profiles["reference"]["font"] == "宋体", "reference style rule changed")


@case
def script_generator_runtime_base_fragment_is_injected() -> None:
    assert_true("__BASE_RUNTIME__" not in RUNTIME_TEMPLATE, "base runtime placeholder leaked into generated template")
    assert_true(
        RUNTIME_TEMPLATE.lstrip().startswith("# -*- coding: utf-8 -*-"),
        "generated script header should come from base runtime",
    )
    assert_true("DATA = json.loads(__DATA_BLOB__)" in RUNTIME_TEMPLATE, "base runtime data bootstrap was not injected")
    assert_true("sys.dont_write_bytecode = True" in RUNTIME_TEMPLATE, "generated scripts should suppress __pycache__ clutter")
    assert_true(
        RUNTIME_TEMPLATE.index("sys.dont_write_bytecode = True") < RUNTIME_TEMPLATE.index("from latex_omath import latex_to_omath"),
        "bytecode suppression must run before local latex_omath imports",
    )
    assert_true("def apply_run_profile" in RUNTIME_TEMPLATE, "base run styling helper was not injected")
    assert_true("def add_text" in RUNTIME_TEMPLATE, "base text helper was not injected")
    assert_true("def setup_section" in RUNTIME_TEMPLATE, "base section setup helper was not injected")
    assert_true("def force_cover_headerless" in RUNTIME_TEMPLATE, "base cover cleanup helper was not injected")
    assert_true("def is_figure_caption_text" in RUNTIME_TEMPLATE, "caption prose classifier was not injected")
    assert_true("elif is_figure_caption_text(text):" in RUNTIME_TEMPLATE, "body renderer should use the safe caption classifier")
    assert_true(
        RUNTIME_TEMPLATE.index("def add_text") < RUNTIME_TEMPLATE.index("def render_cover_and_declarations"),
        "base text helpers should be defined before cover rendering",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def setup_section") < RUNTIME_TEMPLATE.index("def build_document"),
        "base page helpers should be defined before build orchestration",
    )


@case
def script_generator_runtime_content_helpers_fragment_is_injected() -> None:
    assert_true(
        "__CONTENT_HELPERS_RUNTIME__" not in RUNTIME_TEMPLATE,
        "content helper runtime placeholder leaked into generated template",
    )
    assert_true("def is_front_section_index" in RUNTIME_TEMPLATE, "front-section predicate was not injected")
    assert_true("def normalize_caption" in RUNTIME_TEMPLATE, "caption normalizer was not injected")
    assert_true("def clean_text_artifacts" in RUNTIME_TEMPLATE, "text cleanup helper was not injected")
    assert_true("def clean_formula_text" in RUNTIME_TEMPLATE, "formula cleanup helper was not injected")
    assert_true("def add_caption" in RUNTIME_TEMPLATE, "caption renderer helper was not injected")
    assert_true(RUNTIME_TEMPLATE.count("def clean_text_artifacts") == 1, "text cleanup helper should be injected exactly once")
    assert_true(
        RUNTIME_TEMPLATE.index("def clean_formula_text") < RUNTIME_TEMPLATE.index("def text_formula_to_latex"),
        "formula cleanup should be defined before plain-text formula conversion",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def add_caption") < RUNTIME_TEMPLATE.index("def render_table"),
        "caption helper should be defined before table/image rendering",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def is_front_section_index") < RUNTIME_TEMPLATE.index("def render_body"),
        "front-section predicate should be defined before body rendering",
    )


@case
def script_generator_runtime_formula_fragment_is_injected() -> None:
    assert_true("__FORMULA_RUNTIME__" not in RUNTIME_TEMPLATE, "formula runtime placeholder leaked into generated template")
    assert_true("def append_inline_formula" in RUNTIME_TEMPLATE, "inline formula runtime fragment was not injected")
    assert_true("def add_rich_text_runs" in RUNTIME_TEMPLATE, "rich_text runtime fragment was not injected")


@case
def script_generator_runtime_formula_text_fragment_is_injected() -> None:
    assert_true("__FORMULA_TEXT_RUNTIME__" not in RUNTIME_TEMPLATE, "formula text runtime placeholder leaked into generated template")
    assert_true("def chapter_number_from_heading" in RUNTIME_TEMPLATE, "chapter-number helper was not injected")
    assert_true("def text_formula_to_latex" in RUNTIME_TEMPLATE, "plain-text formula conversion helper was not injected")
    assert_true("def split_formula_number" in RUNTIME_TEMPLATE, "formula-number parsing helper was not injected")
    assert_true("def formula_has_number" in RUNTIME_TEMPLATE, "formula-number predicate was not injected")
    assert_true(RUNTIME_TEMPLATE.count("def text_formula_to_latex") == 1, "plain-text formula conversion helper should be injected exactly once")
    assert_true(
        RUNTIME_TEMPLATE.index("def clean_formula_text") < RUNTIME_TEMPLATE.index("def text_formula_to_latex"),
        "plain-text formula conversion should be defined after formula text cleanup",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def text_formula_to_latex") < RUNTIME_TEMPLATE.index("def render_formula"),
        "plain-text formula conversion should be defined before formula rendering",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def chapter_number_from_heading") < RUNTIME_TEMPLATE.index("def render_body"),
        "chapter-number helper should be defined before body rendering",
    )


@case
def script_generator_runtime_formula_render_fragment_is_injected() -> None:
    assert_true("__FORMULA_RENDER_RUNTIME__" not in RUNTIME_TEMPLATE, "formula render runtime placeholder leaked into generated template")
    assert_true("def next_formula_label" in RUNTIME_TEMPLATE, "formula label counter helper was not injected")
    assert_true("def render_plain_formula" in RUNTIME_TEMPLATE, "plain formula renderer was not injected")
    assert_true("def render_formula" in RUNTIME_TEMPLATE, "formula renderer was not injected")
    assert_true(RUNTIME_TEMPLATE.count("def render_formula") == 1, "formula renderer should be injected exactly once")
    assert_true(
        RUNTIME_TEMPLATE.index("def text_formula_to_latex") < RUNTIME_TEMPLATE.index("def render_formula"),
        "formula renderer should be defined after plain-text formula conversion",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def render_formula") < RUNTIME_TEMPLATE.index("def render_paragraph_item"),
        "formula renderer should be defined before body paragraph dispatch",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("FORMULA_COUNTERS = {}") < RUNTIME_TEMPLATE.index("if __name__ == '__main__'"),
        "formula counters should be initialized before the generated-script entrypoint runs",
    )


@case
def script_generator_runtime_media_table_fragment_is_injected() -> None:
    assert_true("__MEDIA_TABLE_RUNTIME__" not in RUNTIME_TEMPLATE, "media/table runtime placeholder leaked into generated template")
    assert_true("def render_table" in RUNTIME_TEMPLATE, "table runtime fragment was not injected")
    assert_true("def render_image" in RUNTIME_TEMPLATE, "image runtime fragment was not injected")
    assert_true("def add_code_block" in RUNTIME_TEMPLATE, "code-block runtime fragment was not injected")


@case
def script_generator_runtime_references_fragment_is_injected() -> None:
    assert_true("__REFERENCES_RUNTIME__" not in RUNTIME_TEMPLATE, "references runtime placeholder leaked into generated template")
    assert_true("def render_reference_entries" in RUNTIME_TEMPLATE, "reference runtime fragment was not injected")
    assert_true("def render_backmatter_section" in RUNTIME_TEMPLATE, "backmatter runtime fragment was not injected")
    assert_true("def collect_structural_backmatter" in RUNTIME_TEMPLATE, "structural backmatter runtime fragment was not injected")


@case
def script_generator_runtime_toc_fragment_is_injected() -> None:
    assert_true("__TOC_RUNTIME__" not in RUNTIME_TEMPLATE, "TOC runtime placeholder leaked into generated template")
    assert_true("def collect_toc_entries" in RUNTIME_TEMPLATE, "TOC entry collection runtime fragment was not injected")
    assert_true("def add_toc" in RUNTIME_TEMPLATE, "TOC rendering runtime fragment was not injected")
    assert_true("def _infer_heading_pages_from_word_com" in RUNTIME_TEMPLATE, "Word COM TOC page-resolution fragment was not injected")


@case
def script_generator_runtime_build_fragment_is_injected() -> None:
    assert_true("__BUILD_RUNTIME__" not in RUNTIME_TEMPLATE, "build runtime placeholder leaked into generated template")
    assert_true("def reset_build_stats" in RUNTIME_TEMPLATE, "build stats runtime fragment was not injected")
    assert_true("def write_build_manifest" in RUNTIME_TEMPLATE, "manifest writer runtime fragment was not injected")
    assert_true("def build_document" in RUNTIME_TEMPLATE, "document build runtime fragment was not injected")
    assert_true("def main" in RUNTIME_TEMPLATE, "generated-script main runtime fragment was not injected")


@case
def script_generator_runtime_cover_fragment_is_injected() -> None:
    assert_true("__COVER_RUNTIME__" not in RUNTIME_TEMPLATE, "cover runtime placeholder leaked into generated template")
    assert_true("def render_cover_and_declarations" in RUNTIME_TEMPLATE, "cover renderer runtime fragment was not injected")
    assert_true("def render_cover_table" in RUNTIME_TEMPLATE, "cover table runtime fragment was not injected")
    assert_true("def compute_cover_skip_indices" in RUNTIME_TEMPLATE, "cover spacer runtime fragment was not injected")
    assert_true("def set_cover_cell_margins" in RUNTIME_TEMPLATE, "cover cell-margin helper should not collide with body table helper")
    assert_true(
        RUNTIME_TEMPLATE.index("def render_cover_and_declarations") < RUNTIME_TEMPLATE.index("def render_front_matter"),
        "cover runtime should be defined before front matter rendering",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def set_cell_borders") < RUNTIME_TEMPLATE.index("def add_code_block"),
        "shared border helper should be defined before media/table runtime uses it",
    )


@case
def script_generator_runtime_front_matter_fragment_is_injected() -> None:
    assert_true("__FRONT_MATTER_RUNTIME__" not in RUNTIME_TEMPLATE, "front matter runtime placeholder leaked into generated template")
    assert_true("def section_text" in RUNTIME_TEMPLATE, "front matter section_text helper was not injected")
    assert_true("def add_keywords" in RUNTIME_TEMPLATE, "front matter keyword helper was not injected")
    assert_true("def render_front_matter" in RUNTIME_TEMPLATE, "front matter renderer runtime fragment was not injected")
    assert_true(RUNTIME_TEMPLATE.count("def render_front_matter") == 1, "front matter renderer should be injected exactly once")
    assert_true(
        RUNTIME_TEMPLATE.index("def add_rich_text_item") < RUNTIME_TEMPLATE.index("def render_front_matter"),
        "front matter renderer should be defined after rich-text formula helpers",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def add_toc") < RUNTIME_TEMPLATE.index("def render_front_matter"),
        "front matter renderer should be defined after TOC helpers",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def render_front_matter") < RUNTIME_TEMPLATE.index("def is_front_section_index"),
        "front matter renderer should stay before body/front-index helpers",
    )


@case
def script_generator_runtime_body_fragment_is_injected() -> None:
    assert_true("__BODY_RUNTIME__" not in RUNTIME_TEMPLATE, "body runtime placeholder leaked into generated template")
    assert_true("def render_paragraph_item" in RUNTIME_TEMPLATE, "body paragraph dispatcher runtime fragment was not injected")
    assert_true("def render_body" in RUNTIME_TEMPLATE, "body renderer runtime fragment was not injected")
    assert_true(RUNTIME_TEMPLATE.count("def render_body") == 1, "body renderer should be injected exactly once")
    assert_true(
        RUNTIME_TEMPLATE.index("def render_formula") < RUNTIME_TEMPLATE.index("def render_paragraph_item"),
        "body renderer should be defined after formula rendering helpers",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def render_table") < RUNTIME_TEMPLATE.index("def render_paragraph_item"),
        "body renderer should be defined after table/image helpers",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def render_reference_entries") < RUNTIME_TEMPLATE.index("def render_body"),
        "body renderer should be defined after reference/backmatter helpers",
    )
    assert_true(
        RUNTIME_TEMPLATE.index("def render_body") < RUNTIME_TEMPLATE.index("def build_document"),
        "body renderer should be defined before document build orchestration",
    )


@case
def lof_generated_when_figure_captions_present() -> None:
    img_dir = new_workdir("lof_fig_dir")
    write_sample_png(img_dir / "fig1.png")
    content = base_content([
        {"role": "figure", "image": "fig1.png", "caption": "Fig. 1 Test figure caption"},
        "Body paragraph after figure.",
    ])
    content["_meta"]["images_dir"] = str(img_dir)
    content["_meta"]["images_extracted"] = 1
    content["sections"][0]["images"] = ["fig1.png"]
    result = run_generated_case("lof_figures", content)
    xml = result["xml"]
    assert_true("Fig. 1 Test figure caption" in xml, "Figure caption should appear in LOF entry")
    # Verify LOF uses tab stops for page number alignment
    assert_true("<w:tab" in xml, "Tab stops should be present for page number alignment")


@case
def lot_generated_when_table_captions_present() -> None:
    content = base_content([
        {"role": "table_caption", "text": "Table 2 Test table caption"},
        {"role": "table", "table_rows": [["A", "B"], ["1", "2"]]},
    ])
    content["_meta"]["tables_count"] = 1
    result = run_generated_case("lot_tables", content)
    xml = result["xml"]
    # The heading uses 2-char-width spaces between characters like TOC heading.
    # Verify both characters appear in the same paragraph context.
    assert_true("表" in xml and "清" in xml and "单" in xml,
               "LOT heading characters should appear in output XML")
    assert_true("Table 2 Test table caption" in xml, "Table caption should appear in LOT entry")


@case
def no_lof_lot_when_no_captions() -> None:
    content = base_content(["Plain body paragraph with no figures or tables."])
    result = run_generated_case("no_lof_lot", content)
    xml = result["xml"]
    lof_heading = "图  清  单"
    lot_heading = "表  清  单"
    assert_true(lof_heading not in xml, f"LOF heading should not appear without figure captions")
    assert_true(lot_heading not in xml, f"LOT heading should not appear without table captions")


@case
def lot_not_generated_for_caption_like_heading_without_table() -> None:
    content = base_content(["Intro body."])
    content["sections"].append({
        "heading": "Table 1 Data sources and economic meaning",
        "level": 1,
        "role": "body",
        "paragraphs": ["Only prose under a caption-like heading, no table object."],
        "images": [],
    })
    content["_meta"]["tables_count"] = 0
    result = run_generated_case("lot_heading_without_table", content)
    doc = Document(result["work"] / "out.docx")
    compact_texts = {"".join((p.text or "").split()) for p in doc.paragraphs}
    assert_true("表清单" not in compact_texts, "LOT should not be generated without a real table")


@case
def lof_lot_both_generated_when_both_present() -> None:
    img_dir = new_workdir("lof_lot_both_dir")
    write_sample_png(img_dir / "fig1.png")
    content = base_content([
        {"role": "figure", "image": "fig1.png", "caption": "Fig. 1 Test figure caption"},
        {"role": "table_caption", "text": "Table 2 Test table caption"},
        {"role": "table", "table_rows": [["A", "B"]]},
    ])
    content["_meta"]["images_dir"] = str(img_dir)
    content["_meta"]["images_extracted"] = 1
    content["_meta"]["tables_count"] = 1
    content["sections"][0]["images"] = ["fig1.png"]
    result = run_generated_case("lof_lot_both", content)
    xml = result["xml"]
    assert_true("Fig. 1 Test figure caption" in xml, "Figure caption in LOF")
    assert_true("Table 2 Test table caption" in xml, "Table caption in LOT")
    # LOF should appear before LOT in document order
    fig_pos = xml.find("Fig. 1 Test figure caption")
    tbl_pos = xml.find("Table 2 Test table caption")
    assert_true(fig_pos < tbl_pos, f"LOF should appear before LOT (fig={fig_pos}, tbl={tbl_pos})")
    # Verify tab elements exist for LOF/LOT entries (right-aligned page numbers)
    assert_true("<w:tab" in xml, "Tab stops should be present in LOF/LOT entries")


@case
def lof_lot_profile_defaults_injected_into_runtime() -> None:
    assert_true("def add_figure_list" in RUNTIME_TEMPLATE, "add_figure_list runtime fragment was not injected")
    assert_true("def add_table_list" in RUNTIME_TEMPLATE, "add_table_list runtime fragment was not injected")
    assert_true("def collect_figure_entries" in RUNTIME_TEMPLATE, "collect_figure_entries runtime fragment was not injected")
    assert_true("def collect_table_entries" in RUNTIME_TEMPLATE, "collect_table_entries runtime fragment was not injected")
    assert_true("def add_list_entry" in RUNTIME_TEMPLATE, "add_list_entry runtime fragment was not injected")
    assert_true("def _infer_caption_pages_from_word_com" in RUNTIME_TEMPLATE,
                "_infer_caption_pages_from_word_com runtime fragment was not injected")
    assert_true("CAPTION_PAGE_MAP" in RUNTIME_TEMPLATE, "CAPTION_PAGE_MAP global was not injected")
    assert_true(
        RUNTIME_TEMPLATE.index("def add_figure_list") < RUNTIME_TEMPLATE.index("def render_front_matter"),
        "LOF/LOT helpers should be defined before front matter renderer uses them",
    )


@case
def lof_lot_paragraphs_excluded_from_conformance() -> None:
    img_dir = new_workdir("lof_conformance_dir")
    write_sample_png(img_dir / "fig_cap.png")
    content = base_content([
        {"role": "figure", "image": "fig_cap.png", "caption": "Fig. 9 Conformance caption test"},
        "Body paragraph.",
    ])
    content["_meta"]["images_dir"] = str(img_dir)
    content["_meta"]["images_extracted"] = 1
    content["sections"][0]["images"] = ["fig_cap.png"]
    result = run_generated_case("lof_conformance", content)
    conf = check_conformance(str(result["work"]), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in conf.get("issues", [])]
    assert_true("STYLE_MISMATCH" not in codes,
                f"conformance should not flag LOF lines as style mismatches: {conf.get('issues', [])}")


@case
def conformance_does_not_satisfy_caption_from_lof_lot_tab_line() -> None:
    work = new_workdir("lof_caption_missing_not_satisfied_by_list")
    content = base_content([
        {"role": "figure_caption", "text": "Fig. 1 Missing body caption"},
        "Body paragraph remains.",
    ])
    content["sections"][0]["heading"] = "正文"
    write_json(work / "format.json", base_format())
    write_json(work / "content.json", content)
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("目  录")
    list_line = doc.add_paragraph()
    list_line.add_run("Fig. 1 Missing body caption")
    list_line.add_run("\t")
    list_line.add_run("1")
    doc.add_paragraph("Body paragraph remains.")
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] Synthetic reference.")
    doc.save(work / "out.docx")
    conf = check_conformance(str(work), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in conf.get("issues", [])]
    assert_true(
        "CONTENT_PARAGRAPH_MISSING" in codes,
        f"conformance should not treat LOF/LOT list lines as body captions: {conf.get('issues', [])}",
    )
