"""Content parser regression cases."""
from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from typing import Any, Dict, List
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Pt
from lxml import etree

from content_parser import extract as extract_docx_content
from content_parser_modules.caption_flow import (
    is_figure_caption,
    is_table_caption,
    normalize_caption_spacing,
    pair_figure_blocks,
)
from content_parser_modules.body_dispatcher import append_text_or_code, parse_body_sections
from content_parser_modules.formula_extractor import _strip_trailing_formula_labels_from_xml
from content_parser_modules.image_extractor import ImageRegistry
from content_parser_modules.paragraph_stream import append_stream_run_group
from content_parser_modules.reference_collector import ReferenceCollector, is_reference_heading
from content_parser_modules.section_builder import (
    filter_content_sections,
    make_body_section,
    mark_first_body_page_break,
    postprocess_section_paragraphs,
)
from content_parser_modules.text_cleaner import clean_code_text, clean_text_artifacts
from formula_semantics import (
    CATEGORY_CONTAMINATED,
    CATEGORY_DISPLAY_MATH,
    CATEGORY_QUANTITY_TEXT,
    classify_formula_text,
    is_formula_problem_text,
    looks_like_formula_text as semantic_looks_like_formula_text,
    split_inline_math_spans,
)
from latex_omath import latex_to_omath
from qa_checker import check_output, write_reports
from content_parser_modules.source_audit import audit_docx_source

from regression_suite_modules.generated_docx import make_vml_picture_docx, omath_count, omath_para_count, run_generated_case
from regression_suite_modules.harness import (
    PNG_1X1,
    assert_true,
    base_content,
    base_format,
    case,
    fail,
    new_workdir,
    write_json,
    write_sample_png,
)


def _rewrite_docx_part(docx_path: Path, part_name: str, transform) -> None:
    original = docx_path.with_suffix(docx_path.suffix + ".src")
    docx_path.replace(original)
    with zipfile.ZipFile(original, "r") as zin, zipfile.ZipFile(docx_path, "w") as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename == part_name:
                text = data.decode("utf-8")
                data = transform(text).encode("utf-8")
            zout.writestr(info, data)
    original.unlink()


def _content_plain_text(content: Dict[str, Any]) -> str:
    pieces: List[str] = []
    for sec in content.get("sections") or []:
        if sec.get("heading"):
            pieces.append(str(sec.get("heading") or ""))
        for item in sec.get("paragraphs") or []:
            if isinstance(item, str):
                pieces.append(item)
            elif isinstance(item, dict):
                pieces.append(str(item.get("text") or item.get("code") or ""))
    return "\n".join(pieces)


@case
def caption_flow_pairs_images_with_nearby_captions() -> None:
    paragraphs = [
        {"role": "image", "image": "a.png"},
        "intervening prose",
        {"role": "figure_caption", "text": "Fig. 1 Demo"},
        {"role": "image", "image": "b.png"},
        {"role": "image", "image": "c.png"},
        {"role": "figure_caption", "text": "Figure 2 First"},
        {"role": "figure_caption", "text": "Figure 3 Second"},
    ]
    paired = pair_figure_blocks(paragraphs)
    assert_true(paired[0] == {"role": "figure", "image": "a.png", "caption": "Fig. 1 Demo"}, "image/body/caption was not paired")
    assert_true(paired[1] == "intervening prose", "intervening prose was not preserved after pairing")
    assert_true(paired[2] == {"role": "figure", "image": "b.png", "caption": "Figure 2 First"}, "first stacked image was not paired")
    assert_true(paired[3] == {"role": "figure", "image": "c.png", "caption": "Figure 3 Second"}, "second stacked image was not paired")
    assert_true(is_figure_caption("\u56fe1\u7ed3\u679c\u5bf9\u6bd4"), "Chinese figure caption was not detected")
    assert_true(is_figure_caption("图 1 机器学习研究流程示意图"), "Chinese noun-phrase figure caption was not detected")
    assert_true(not is_figure_caption("图 1 展示了从数据到决策的机器学习研究流程。"), "figure reference prose was misclassified as caption")
    assert_true(is_table_caption("表 1 不同模型在验证集上的表现"), "Chinese table caption was not detected")
    assert_true(not is_table_caption("表 1 显示 XGBoost 在多数指标上更高。"), "table reference prose was misclassified as caption")
    assert_true(normalize_caption_spacing("\u56fe1\u7ed3\u679c") == "\u56fe 1 \u7ed3\u679c", "Chinese figure caption spacing was not normalized")


@case
def text_cleaner_removes_editor_noise_and_preserves_code_lines() -> None:
    assert_true(clean_text_artifacts("[label](https://example.test)") == "label", "markdown link label was not preserved")
    assert_true(clean_text_artifacts("Plain Text") == "", "editor noise was not removed")
    assert_true(clean_text_artifacts("\u590d\u5236") == "", "Chinese editor noise was not removed")
    code = clean_code_text(" interface  Gi0/0/1 \nPlain Text\n ip   address  10.0.0.1 ")
    assert_true(code == "interface Gi0/0/1\nip address 10.0.0.1", f"code cleanup changed unexpectedly: {code!r}")


@case
def paragraph_stream_run_group_preserves_text_and_math_semantics() -> None:
    section = {"paragraphs": []}

    def append_text(section_obj: Dict[str, Any], text: str, in_appendix: bool = False) -> None:
        section_obj["paragraphs"].append({"role": "text", "text": text, "appendix": in_appendix})

    append_stream_run_group(
        section,
        [{"type": "text", "text": "plain text"}],
        append_text_or_code_func=append_text,
        in_appendix=True,
    )
    assert_true(section["paragraphs"][0] == {"role": "text", "text": "plain text", "appendix": True}, "plain run group bypassed text callback")

    append_stream_run_group(
        section,
        [{"type": "math", "text": "a=b", "math": [{"text": "a=b", "had_number_label": True}]}],
        append_text_or_code_func=append_text,
    )
    formula = section["paragraphs"][1]
    assert_true(formula["role"] == "formula", "math-only run group did not become a formula")
    assert_true(formula["numbered"] is True, "math had_number_label was not preserved")


@case
def reference_collector_exits_before_backmatter_and_keeps_tables() -> None:
    collector = ReferenceCollector(
        clean_text_func=clean_text_artifacts,
        is_backmatter_heading_func=lambda text: text == "Appendix A",
        normalize_heading_spacing_func=lambda text: text,
        classify_section_role_func=lambda heading, level: "appendix",
        table_rows_look_like_code_func=lambda rows: rows and rows[0] and rows[0][0].startswith("interface"),
        code_text_from_table_rows_func=lambda rows, clean_code_func=None: clean_code_func(rows[0][0]) if clean_code_func else rows[0][0],
        clean_code_func=clean_code_text,
    )
    assert_true(is_reference_heading("\u53c2\u8003\u6587\u732e"), "Chinese reference heading was not detected")
    assert_true(collector.start_if_heading("References"), "English reference heading was not accepted")
    assert_true(collector.consume_text("[1]  Example  Paper"), "active collector did not consume reference text")
    assert_true(collector.consume_table_rows([["interface  Gi0/0/1\n ip  address  10.0.0.1"]]), "active collector did not consume reference table")
    backmatter = collector.exit_to_backmatter_section("Appendix A", 1)
    assert_true(backmatter and backmatter["role"] == "appendix", "collector did not exit on backmatter")
    refs = collector.finish()
    assert_true(refs[0] == "[1] Example Paper", f"reference text was not cleaned: {refs[0]!r}")
    assert_true(refs[1]["role"] == "code", "reference code table was not preserved as code")
    assert_true(refs[1]["code"] == "interface Gi0/0/1\nip address 10.0.0.1", "reference code table was not cleaned")


@case
def section_builder_filters_placeholder_and_finalizes_blocks() -> None:
    sections = [
        make_body_section(),
        {"heading": "Abstract", "level": 1, "role": "en_abstract", "paragraphs": ["Summary"], "images": []},
        {"heading": "1 Introduction", "level": 1, "role": "body", "paragraphs": [], "images": []},
        {
            "heading": "2 Results",
            "level": 1,
            "role": "body",
            "paragraphs": [
                {"role": "image", "image": "fig1.png"},
                {"role": "figure_caption", "text": "Fig. 1. Result overview"},
            ],
            "images": ["fig1.png"],
        },
    ]
    filtered = filter_content_sections(sections)
    assert_true(filtered[0]["heading"] == "Abstract", "initial body placeholder was not filtered")
    assert_true(filtered[1]["heading"] == "1 Introduction", "empty structural heading was dropped")
    mark_first_body_page_break(filtered)
    assert_true(not filtered[0].get("page_break_before"), "front matter was marked as first body page")
    assert_true(filtered[1].get("page_break_before") is True, "first body section was not marked for page break")
    postprocess_section_paragraphs(filtered)
    assert_true(filtered[2]["paragraphs"][0]["role"] == "figure", "image and figure caption were not paired")
    assert_true(filtered[2]["paragraphs"][0]["caption"] == "Fig. 1. Result overview", "figure caption text changed")


@case
def body_dispatcher_routes_sections_references_tables_and_appendix_code() -> None:
    work = new_workdir("body_dispatcher")
    fig_dir = work / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("报名序号: [报名序号]")
    doc.add_paragraph("图 1 系统结构示意")
    doc.add_paragraph("参考文献")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "[1] Zhang S. Synthetic reference."
    doc.add_paragraph("附录 A")
    doc.add_paragraph("interface Gi0/0/1\n ip address 10.0.0.1")

    scratch = {"paragraphs": []}
    append_text_or_code(scratch, r"$$a=b$$")
    assert_true(scratch["paragraphs"][0]["role"] == "formula", "dispatcher text append did not preserve display math")

    result = parse_body_sections(doc, 0, ImageRegistry(str(fig_dir), "dispatch_img"))
    sections = filter_content_sections(result.sections)
    assert_true(sections[0]["heading"] == "1 Introduction", "body heading was not routed into a section")
    assert_true(result.placeholders_removed == 1, "unfilled placeholder paragraph was not filtered")
    assert_true(sections[0]["paragraphs"][0]["role"] == "figure_caption", "figure caption was not classified")
    assert_true(result.references and result.references[0]["role"] == "table", "reference table was not captured before appendix")
    assert_true(sections[1]["role"] == "appendix", "back matter did not exit reference collection")
    assert_true(sections[1]["paragraphs"][0]["role"] == "code", "appendix command text was not routed as code")


@case
def content_parser_preserves_body_custom_xml_wrapped_paragraph_and_table() -> None:
    work = new_workdir("parser_body_custom_xml_wrapper")
    docx = work / "body_custom_xml_wrapper.docx"
    doc = Document()
    doc.add_paragraph("1 Wrapped section")
    doc.add_paragraph("Wrapped paragraph from customXml")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "WrappedTableLeft"
    table.cell(0, 1).text = "WrappedTableRight"
    doc.add_paragraph("After wrapper paragraph")
    doc.save(docx)

    def wrap_body_content(xml: str) -> str:
        para_marker = "<w:t>Wrapped paragraph from customXml</w:t>"
        table_marker = "<w:t>WrappedTableLeft</w:t>"
        assert_true(para_marker in xml, "body customXml paragraph marker not found")
        assert_true(table_marker in xml, "body customXml table marker not found")
        para_start = xml.rfind("<w:p", 0, xml.find(para_marker))
        table_end = xml.find("</w:tbl>", xml.find(table_marker)) + len("</w:tbl>")
        assert_true(para_start >= 0 and table_end > para_start, "body customXml wrapper bounds not found")
        wrapped = xml[para_start:table_end]
        return (
            xml[:para_start]
            + '<w:customXml w:element="wrappedBody" w:uri="urn:synthetic">'
            + wrapped
            + "</w:customXml>"
            + xml[table_end:]
        )

    _rewrite_docx_part(docx, "word/document.xml", wrap_body_content)

    content = extract_docx_content(str(docx), output_dir=str(work))
    plain = _content_plain_text(content)
    assert_true("Wrapped paragraph from customXml" in plain, f"customXml-wrapped body paragraph was lost: {content}")
    assert_true("After wrapper paragraph" in plain, f"paragraph after customXml wrapper was lost or replaced: {content}")
    table_items = [
        item
        for section in content.get("sections") or []
        for item in section.get("paragraphs") or []
        if isinstance(item, dict) and item.get("role") == "table"
    ]
    assert_true(
        any(item.get("table_rows") == [["WrappedTableLeft", "WrappedTableRight"]] for item in table_items),
        f"customXml-wrapped body table was lost: {content}",
    )


@case
def content_parser_preserves_body_sdt_wrapped_paragraphs_and_table_in_place() -> None:
    work = new_workdir("parser_body_sdt_wrapper")
    docx = work / "body_sdt_wrapper.docx"
    doc = Document()
    doc.add_paragraph("1 Controlled section")
    doc.add_paragraph("Before content control wrapper")
    doc.add_paragraph("Controlled paragraph A")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "ControlledTableLeft"
    table.cell(0, 1).text = "ControlledTableRight"
    doc.add_paragraph("Controlled paragraph B")
    doc.add_paragraph("After content control wrapper")
    doc.save(docx)

    def wrap_body_content_control(xml: str) -> str:
        start_marker = "<w:t>Controlled paragraph A</w:t>"
        table_marker = "<w:t>ControlledTableLeft</w:t>"
        end_marker = "<w:t>Controlled paragraph B</w:t>"
        assert_true(start_marker in xml, "body sdt start paragraph marker not found")
        assert_true(table_marker in xml, "body sdt table marker not found")
        assert_true(end_marker in xml, "body sdt end paragraph marker not found")
        start = xml.rfind("<w:p", 0, xml.find(start_marker))
        end_para_start = xml.rfind("<w:p", 0, xml.find(end_marker))
        end = xml.find("</w:p>", end_para_start) + len("</w:p>")
        assert_true(start >= 0 and end > start, "body sdt wrapper bounds not found")
        wrapped = xml[start:end]
        return (
            xml[:start]
            + '<w:sdt><w:sdtPr><w:tag w:val="body-repeat-section"/></w:sdtPr><w:sdtContent>'
            + wrapped
            + "</w:sdtContent></w:sdt>"
            + xml[end:]
        )

    _rewrite_docx_part(docx, "word/document.xml", wrap_body_content_control)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [
        item
        for section in content.get("sections") or []
        for item in section.get("paragraphs") or []
    ]
    flow: List[str] = []
    for item in paragraphs:
        if isinstance(item, str):
            flow.append(item)
        elif isinstance(item, dict) and item.get("role") == "table":
            flow.append("TABLE:" + "|".join(item.get("table_rows", [[]])[0]))

    expected = [
        "Before content control wrapper",
        "Controlled paragraph A",
        "TABLE:ControlledTableLeft|ControlledTableRight",
        "Controlled paragraph B",
        "After content control wrapper",
    ]
    cursor = -1
    for marker in expected:
        try:
            cursor = flow.index(marker, cursor + 1)
        except ValueError:
            fail(f"body content-control wrapper did not preserve source-order marker {marker!r}: {flow}; content={content}")

    meta = content.get("_meta") or {}
    assert_true(
        meta.get("recovered_content_control_paragraphs") == 2,
        f"body content-control paragraph count should track in-place paragraphs only: {meta}",
    )
    assert_true(
        not any(item in {"ControlledTableLeft", "ControlledTableRight"} for item in flow),
        f"table-cell text from body content control leaked as loose paragraphs: {flow}",
    )


@case
def content_parser_preserves_body_sdt_inline_rich_media_formula_note_order() -> None:
    work = new_workdir("parser_body_sdt_inline_rich")
    img = work / "body_sdt_inline_image.png"
    write_sample_png(img, width=120, height=90)
    docx = work / "body_sdt_inline_rich.docx"
    doc = Document()
    doc.add_paragraph("1 Body controlled rich")
    para = doc.add_paragraph()
    para.add_run("Lead")
    para.add_run().add_picture(str(img))
    para.add_run("Formula line ")
    para._element.append(etree.fromstring(latex_to_omath(r"y=2", display=False).encode("utf-8")))
    para.add_run(" noted")
    doc.add_paragraph("After rich controlled paragraph")
    doc.save(docx)

    footnote_text = "Body controlled paragraph note must stay after the formula."

    def inject_reference(xml: str) -> str:
        replacements = [
            (
                '<w:t xml:space="preserve"> noted</w:t></w:r></w:p>',
                '<w:t xml:space="preserve"> noted</w:t></w:r><w:r><w:footnoteReference w:id="8"/></w:r></w:p>',
            ),
            (
                "<w:t> noted</w:t></w:r></w:p>",
                '<w:t> noted</w:t></w:r><w:r><w:footnoteReference w:id="8"/></w:r></w:p>',
            ),
        ]
        for old, new in replacements:
            if old in xml:
                return xml.replace(old, new, 1)
        return xml

    def wrap_body_sdt_and_inline_tail(xml: str) -> str:
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        root = etree.fromstring(xml.encode("utf-8"))
        ns = {"w": w_ns}
        paragraphs = root.xpath(".//w:body/w:p[.//w:t='Lead']", namespaces=ns)
        assert_true(paragraphs, "body rich paragraph not found")
        paragraph = paragraphs[0]
        children = list(paragraph)
        lead_idx = next(
            (
                idx
                for idx, child in enumerate(children)
                if child.xpath(".//w:t[text()='Lead']", namespaces=ns)
            ),
            -1,
        )
        assert_true(lead_idx >= 0 and lead_idx + 1 < len(children), "body rich inline tail injection point not found")
        tail = children[lead_idx + 1 :]
        for child in tail:
            paragraph.remove(child)

        inline_sdt = etree.Element(f"{{{w_ns}}}sdt")
        inline_pr = etree.SubElement(inline_sdt, f"{{{w_ns}}}sdtPr")
        inline_tag = etree.SubElement(inline_pr, f"{{{w_ns}}}tag")
        inline_tag.set(f"{{{w_ns}}}val", "body-inline-rich")
        inline_content = etree.SubElement(inline_sdt, f"{{{w_ns}}}sdtContent")
        hyperlink = etree.SubElement(inline_content, f"{{{w_ns}}}hyperlink")
        hyperlink.set(f"{{{w_ns}}}anchor", "synthetic-body-controlled-link")
        for child in tail:
            hyperlink.append(child)
        paragraph.append(inline_sdt)

        body = paragraph.getparent()
        paragraph_index = body.index(paragraph)
        body.remove(paragraph)
        block_sdt = etree.Element(f"{{{w_ns}}}sdt")
        block_pr = etree.SubElement(block_sdt, f"{{{w_ns}}}sdtPr")
        block_tag = etree.SubElement(block_pr, f"{{{w_ns}}}tag")
        block_tag.set(f"{{{w_ns}}}val", "body-rich-section")
        block_content = etree.SubElement(block_sdt, f"{{{w_ns}}}sdtContent")
        block_content.append(paragraph)
        body.insert(paragraph_index, block_sdt)
        return etree.tostring(root, encoding="unicode")

    _rewrite_docx_part(docx, "word/document.xml", inject_reference)
    _rewrite_docx_part(docx, "word/document.xml", wrap_body_sdt_and_inline_tail)
    with zipfile.ZipFile(docx, "a") as zf:
        zf.writestr(
            "word/footnotes.xml",
            (
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
                '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
                f'<w:footnote w:id="8"><w:p><w:r><w:t>{footnote_text}</w:t></w:r></w:p></w:footnote>'
                "</w:footnotes>"
            ),
        )

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [
        item
        for section in content.get("sections") or []
        for item in section.get("paragraphs") or []
    ]
    image_items = [item for item in paragraphs if isinstance(item, dict) and item.get("role") == "image"]
    rich_items = [item for item in paragraphs if isinstance(item, dict) and item.get("role") == "rich_text"]
    assert_true(image_items, f"body inline content-control image was lost: {paragraphs}")
    assert_true(rich_items, f"body inline content-control formula/note rich text was lost: {paragraphs}")
    rich_item = rich_items[0]
    assert_true(
        [run.get("type") for run in rich_item.get("runs") or []] == ["text", "math", "text", "note_ref"],
        f"body inline content-control rich run order changed: {rich_item}",
    )
    assert_true((rich_item.get("notes") or [{}])[0].get("text") == footnote_text, f"body inline content-control note text lost: {rich_item}")
    assert_true(
        not any(item == "LeadFormula line y=2 noted" for item in paragraphs if isinstance(item, str)),
        f"body inline content-control text was duplicated as flat fallback text: {paragraphs}",
    )
    meta = content.get("_meta") or {}
    assert_true(
        meta.get("recovered_content_control_paragraphs") == 1,
        f"body rich content-control paragraph count should count the wrapped body paragraph once: {meta}",
    )

    result = run_generated_case("body_sdt_inline_rich_render", content, base_format())
    xml = result["xml"]
    assert_true(xml.count("<w:drawing>") == 1, f"body inline content-control image did not render exactly once: {result['manifest']}")
    assert_true(omath_count(xml) >= 1, "body inline content-control formula did not render as native math")
    assert_true("<w:footnoteReference" in xml, "body inline content-control footnote did not render as native reference")
    assert_true(
        xml.find("Lead") < xml.find("<w:drawing>") < xml.find("Formula line") < xml.find("<w:footnoteReference"),
        "body inline content-control image/formula/footnote source order changed",
    )
    assert_true(result["manifest"]["counts"].get("footnote_references_rendered") == 1, f"footnote count missing: {result['manifest']}")
    assert_true(result["report"]["passed"] is True, f"body inline content-control rich render should pass QA: {result['report']}")


@case
def content_parser_uses_final_view_for_tracked_changes_and_ignores_comments() -> None:
    work = new_workdir("parser_tracked_changes_final_view")
    docx = work / "tracked_changes_final_view.docx"
    doc = Document()
    doc.add_paragraph("1 Revision boundary")
    doc.add_paragraph("Body before tail.")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell before tail"
    doc.save(docx)

    def inject_revisions_and_comment(xml: str) -> str:
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        root = etree.fromstring(xml.encode("utf-8"))
        ns = {"w": w_ns}
        body_para = root.xpath(".//w:body/w:p[.//w:t='Body before tail.']", namespaces=ns)[0]
        body_run = body_para.xpath("./w:r[.//w:t='Body before tail.']", namespaces=ns)[0]
        body_run.xpath(".//w:t", namespaces=ns)[0].text = "Body before "
        insert = etree.Element(f"{{{w_ns}}}ins")
        insert.set(f"{{{w_ns}}}id", "1")
        insert.set(f"{{{w_ns}}}author", "Synthetic")
        insert_run = etree.SubElement(insert, f"{{{w_ns}}}r")
        insert_text = etree.SubElement(insert_run, f"{{{w_ns}}}t")
        insert_text.text = "inserted "
        delete = etree.Element(f"{{{w_ns}}}del")
        delete.set(f"{{{w_ns}}}id", "2")
        delete.set(f"{{{w_ns}}}author", "Synthetic")
        delete_run = etree.SubElement(delete, f"{{{w_ns}}}r")
        delete_text = etree.SubElement(delete_run, f"{{{w_ns}}}delText")
        delete_text.text = "deleted "
        comment_start = etree.Element(f"{{{w_ns}}}commentRangeStart")
        comment_start.set(f"{{{w_ns}}}id", "3")
        comment_end = etree.Element(f"{{{w_ns}}}commentRangeEnd")
        comment_end.set(f"{{{w_ns}}}id", "3")
        comment_ref_run = etree.Element(f"{{{w_ns}}}r")
        etree.SubElement(comment_ref_run, f"{{{w_ns}}}commentReference").set(f"{{{w_ns}}}id", "3")
        tail_run = etree.Element(f"{{{w_ns}}}r")
        tail_text = etree.SubElement(tail_run, f"{{{w_ns}}}t")
        tail_text.text = "tail."
        body_para.extend([insert, delete, comment_start, tail_run, comment_end, comment_ref_run])

        cell_para = root.xpath(".//w:tbl//w:tc/w:p[.//w:t='Cell before tail']", namespaces=ns)[0]
        cell_run = cell_para.xpath("./w:r[.//w:t='Cell before tail']", namespaces=ns)[0]
        cell_run.xpath(".//w:t", namespaces=ns)[0].text = "Cell before "
        cell_insert = etree.Element(f"{{{w_ns}}}ins")
        cell_insert.set(f"{{{w_ns}}}id", "4")
        cell_insert.set(f"{{{w_ns}}}author", "Synthetic")
        cell_insert_run = etree.SubElement(cell_insert, f"{{{w_ns}}}r")
        cell_insert_text = etree.SubElement(cell_insert_run, f"{{{w_ns}}}t")
        cell_insert_text.text = "inserted "
        cell_delete = etree.Element(f"{{{w_ns}}}del")
        cell_delete.set(f"{{{w_ns}}}id", "5")
        cell_delete.set(f"{{{w_ns}}}author", "Synthetic")
        cell_delete_run = etree.SubElement(cell_delete, f"{{{w_ns}}}r")
        cell_delete_text = etree.SubElement(cell_delete_run, f"{{{w_ns}}}delText")
        cell_delete_text.text = "deleted "
        cell_tail_run = etree.Element(f"{{{w_ns}}}r")
        cell_tail_text = etree.SubElement(cell_tail_run, f"{{{w_ns}}}t")
        cell_tail_text.text = "tail"
        cell_para.extend([cell_insert, cell_delete, cell_tail_run])
        return etree.tostring(root, encoding="unicode")

    _rewrite_docx_part(docx, "word/document.xml", inject_revisions_and_comment)
    with zipfile.ZipFile(docx, "a") as zf:
        zf.writestr(
            "word/comments.xml",
            (
                '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:comment w:id="3" w:author="Synthetic"><w:p><w:r><w:t>Private reviewer comment</w:t></w:r></w:p></w:comment>'
                "</w:comments>"
            ),
        )

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    text = _content_plain_text(content)
    assert_true("Body before inserted tail." in text, f"body inserted revision text was not preserved in final-view content: {content}")
    assert_true("Body before deleted" not in text, f"body deleted revision text leaked into content: {content}")
    assert_true("Private reviewer comment" not in text, f"comment body leaked into extracted content: {content}")
    table_items = [
        item
        for sec in content.get("sections") or []
        for item in sec.get("paragraphs") or []
        if isinstance(item, dict) and item.get("role") == "table"
    ]
    assert_true(table_items, f"revision table was lost: {content}")
    assert_true(
        table_items[0].get("table_rows") == [["Cell before inserted tail"]],
        f"table inserted/deleted revision final view changed: {table_items[0]}",
    )
    audit = audit_docx_source(str(docx))
    codes = {issue.get("code") for issue in audit.get("issues") or []}
    assert_true("TRACKED_CHANGES_PRESENT" in codes, f"tracked change audit issue missing: {audit}")
    assert_true("COMMENTS_PRESENT" in codes, f"comment audit issue missing: {audit}")

    result = run_generated_case("tracked_changes_final_view_render", content, base_format())
    xml = result["xml"]
    assert_true("Body before inserted tail." in xml, "body inserted revision did not render in generated DOCX")
    assert_true("deleted" not in xml and "Private reviewer comment" not in xml, "deleted/comment text leaked into generated DOCX")
    assert_true("Cell before inserted tail" in xml, "table inserted revision did not render in generated DOCX")


@case
def content_parser_uses_final_view_text_for_revision_wrapped_headings() -> None:
    work = new_workdir("parser_revision_wrapped_headings")
    docx = work / "revision_wrapped_headings.docx"
    doc = Document()
    visible_heading = doc.add_paragraph("1 Visible moved heading")
    visible_heading.style = "Heading 1"
    doc.add_paragraph("Body under moved heading.")
    deleted_heading = doc.add_paragraph("2 Deleted old heading")
    deleted_heading.style = "Heading 1"
    doc.add_paragraph("Body after deleted heading.")
    doc.save(docx)

    def wrap_heading_revisions(xml: str) -> str:
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        root = etree.fromstring(xml.encode("utf-8"))
        ns = {"w": w_ns}

        moved_para = root.xpath(".//w:body/w:p[.//w:t='1 Visible moved heading']", namespaces=ns)[0]
        moved_run = moved_para.xpath("./w:r[.//w:t='1 Visible moved heading']", namespaces=ns)[0]
        moved_para.remove(moved_run)
        move_to = etree.Element(f"{{{w_ns}}}moveTo")
        move_to.set(f"{{{w_ns}}}id", "11")
        move_to.set(f"{{{w_ns}}}author", "Synthetic")
        move_to.append(moved_run)
        moved_para.append(move_to)

        deleted_para = root.xpath(".//w:body/w:p[.//w:t='2 Deleted old heading']", namespaces=ns)[0]
        deleted_run = deleted_para.xpath("./w:r[.//w:t='2 Deleted old heading']", namespaces=ns)[0]
        deleted_text = deleted_run.xpath(".//w:t", namespaces=ns)[0]
        deleted_text.tag = f"{{{w_ns}}}delText"
        deleted_para.remove(deleted_run)
        move_from = etree.Element(f"{{{w_ns}}}moveFrom")
        move_from.set(f"{{{w_ns}}}id", "12")
        move_from.set(f"{{{w_ns}}}author", "Synthetic")
        move_from.append(deleted_run)
        deleted_para.append(move_from)
        return etree.tostring(root, encoding="unicode")

    _rewrite_docx_part(docx, "word/document.xml", wrap_heading_revisions)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    headings = [section.get("heading") for section in content.get("sections") or []]
    assert_true("1 Visible moved heading" in headings, f"moveTo-wrapped heading was not routed as a section: {content}")
    assert_true("2 Deleted old heading" not in headings, f"moveFrom-wrapped deleted heading leaked as a section: {content}")
    visible_section = next(section for section in content.get("sections") or [] if section.get("heading") == "1 Visible moved heading")
    assert_true(
        "Body under moved heading." in visible_section.get("paragraphs", []),
        f"body paragraph was not kept under the moveTo heading: {content}",
    )
    text = _content_plain_text(content)
    assert_true("2 Deleted old heading" not in text, f"deleted revision heading text leaked into content: {content}")

    result = run_generated_case("revision_wrapped_heading_render", content, base_format())
    xml = result["xml"]
    assert_true("1 Visible moved heading" in xml, "moveTo-wrapped heading did not render in generated DOCX")
    assert_true("2 Deleted old heading" not in xml, "moveFrom-wrapped deleted heading rendered in generated DOCX")


@case
def content_parser_uses_final_view_for_revision_wrapped_table_rows() -> None:
    work = new_workdir("parser_revision_wrapped_table_rows")
    docx = work / "revision_wrapped_table_rows.docx"
    doc = Document()
    doc.add_paragraph("1 Table row revisions")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "VisibleMovedRow"
    table.cell(1, 1).text = "42"
    table.cell(2, 0).text = "DeletedMovedRow"
    table.cell(2, 1).text = "13"
    doc.save(docx)

    def wrap_table_rows(xml: str) -> str:
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        root = etree.fromstring(xml.encode("utf-8"))
        ns = {"w": w_ns}
        tbl = root.xpath(".//w:tbl[.//w:t='VisibleMovedRow']", namespaces=ns)[0]

        visible_row = tbl.xpath("./w:tr[.//w:t='VisibleMovedRow']", namespaces=ns)[0]
        visible_index = tbl.index(visible_row)
        tbl.remove(visible_row)
        move_to = etree.Element(f"{{{w_ns}}}moveTo")
        move_to.set(f"{{{w_ns}}}id", "31")
        move_to.set(f"{{{w_ns}}}author", "Synthetic")
        move_to.append(visible_row)
        tbl.insert(visible_index, move_to)

        deleted_row = tbl.xpath("./w:tr[.//w:t='DeletedMovedRow']", namespaces=ns)[0]
        deleted_index = tbl.index(deleted_row)
        tbl.remove(deleted_row)
        move_from = etree.Element(f"{{{w_ns}}}moveFrom")
        move_from.set(f"{{{w_ns}}}id", "32")
        move_from.set(f"{{{w_ns}}}author", "Synthetic")
        move_from.append(deleted_row)
        tbl.insert(deleted_index, move_from)
        return etree.tostring(root, encoding="unicode")

    _rewrite_docx_part(docx, "word/document.xml", wrap_table_rows)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    table_items = [
        item
        for section in content.get("sections") or []
        for item in section.get("paragraphs") or []
        if isinstance(item, dict) and item.get("role") == "table"
    ]
    assert_true(table_items, f"revision-wrapped table was lost: {content}")
    rows = table_items[0].get("table_rows") or []
    assert_true(["Metric", "Value"] in rows, f"header row changed: {rows}")
    assert_true(["VisibleMovedRow", "42"] in rows, f"moveTo-wrapped visible row was not preserved: {content}")
    assert_true(
        not any("DeletedMovedRow" in cell for row in rows for cell in row),
        f"moveFrom-wrapped deleted row leaked into table rows: {rows}",
    )

    result = run_generated_case("revision_wrapped_table_rows_render", content, base_format())
    xml = result["xml"]
    assert_true("VisibleMovedRow" in xml and "42" in xml, "moveTo-wrapped table row did not render")
    assert_true("DeletedMovedRow" not in xml and ">13<" not in xml, "moveFrom-wrapped table row rendered")


@case
def content_parser_uses_final_view_for_revision_wrapped_table_cells() -> None:
    work = new_workdir("parser_revision_wrapped_table_cells")
    docx = work / "revision_wrapped_table_cells.docx"
    doc = Document()
    doc.add_paragraph("1 Table cell revisions")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Current"
    table.cell(0, 2).text = "Previous"
    table.cell(1, 0).text = "RowLabel"
    table.cell(1, 1).text = "VisibleMovedCell"
    table.cell(1, 2).text = "DeletedMovedCell"
    doc.save(docx)

    def wrap_table_cells(xml: str) -> str:
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        root = etree.fromstring(xml.encode("utf-8"))
        ns = {"w": w_ns}
        tr = root.xpath(".//w:tr[.//w:t='VisibleMovedCell']", namespaces=ns)[0]

        visible_cell = tr.xpath("./w:tc[.//w:t='VisibleMovedCell']", namespaces=ns)[0]
        visible_index = tr.index(visible_cell)
        tr.remove(visible_cell)
        move_to = etree.Element(f"{{{w_ns}}}moveTo")
        move_to.set(f"{{{w_ns}}}id", "33")
        move_to.set(f"{{{w_ns}}}author", "Synthetic")
        move_to.append(visible_cell)
        tr.insert(visible_index, move_to)

        deleted_cell = tr.xpath("./w:tc[.//w:t='DeletedMovedCell']", namespaces=ns)[0]
        deleted_index = tr.index(deleted_cell)
        tr.remove(deleted_cell)
        move_from = etree.Element(f"{{{w_ns}}}moveFrom")
        move_from.set(f"{{{w_ns}}}id", "34")
        move_from.set(f"{{{w_ns}}}author", "Synthetic")
        move_from.append(deleted_cell)
        tr.insert(deleted_index, move_from)
        return etree.tostring(root, encoding="unicode")

    _rewrite_docx_part(docx, "word/document.xml", wrap_table_cells)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    table_items = [
        item
        for section in content.get("sections") or []
        for item in section.get("paragraphs") or []
        if isinstance(item, dict) and item.get("role") == "table"
    ]
    assert_true(table_items, f"revision-wrapped table-cell table was lost: {content}")
    rows = table_items[0].get("table_rows") or []
    flat = [cell for row in rows for cell in row]
    assert_true("VisibleMovedCell" in flat, f"moveTo-wrapped visible cell was not preserved: {rows}")
    assert_true("DeletedMovedCell" not in flat, f"moveFrom-wrapped deleted cell leaked into table rows: {rows}")

    result = run_generated_case("revision_wrapped_table_cells_render", content, base_format())
    xml = result["xml"]
    assert_true("VisibleMovedCell" in xml, "moveTo-wrapped table cell did not render")
    assert_true("DeletedMovedCell" not in xml, "moveFrom-wrapped table cell rendered")


@case
def content_parser_preserves_table_at_start_of_body() -> None:
    work = new_workdir("parser_table_first")

    docx = work / "table_first_then_text.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Accuracy"
    table.cell(1, 1).text = "98%"
    doc.add_paragraph("Body paragraph after the opening table.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    items = [item for sec in content.get("sections") or [] for item in sec.get("paragraphs") or []]
    table_items = [item for item in items if isinstance(item, dict) and item.get("table_rows")]
    assert_true(len(table_items) == 1, f"opening table was not preserved in content stream: {items}")
    assert_true(table_items[0]["table_rows"][1] == ["Accuracy", "98%"], "opening table rows changed")

    table_only_docx = work / "table_only.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Only"
    table.cell(0, 1).text = "Table"
    doc.save(table_only_docx)

    table_only = extract_docx_content(str(table_only_docx), output_dir=str(work / "only_out"))
    only_items = [item for sec in table_only.get("sections") or [] for item in sec.get("paragraphs") or []]
    assert_true(
        any(isinstance(item, dict) and item.get("table_rows") for item in only_items),
        f"table-only document became empty content: {table_only}",
    )


@case
def content_parser_preserves_merged_table_cells() -> None:
    work = new_workdir("parser_merged_table")
    docx = work / "merged_table.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "Merged header"
    table.cell(0, 2).text = "Score"
    table.cell(1, 0).merge(table.cell(2, 0)).text = "Group A"
    table.cell(1, 1).text = "Alpha"
    table.cell(1, 2).text = "1"
    table.cell(2, 1).text = "Beta"
    table.cell(2, 2).text = "2"
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    items = [item for sec in content.get("sections") or [] for item in sec.get("paragraphs") or []]
    table_items = [item for item in items if isinstance(item, dict) and item.get("role") == "table"]
    assert_true(len(table_items) == 1, f"merged table was not preserved as a table: {items}")
    table_item = table_items[0]
    assert_true(table_item["table_rows"][0] == ["Merged header", "", "Score"], f"horizontal span did not expand to a stable grid: {table_item}")
    assert_true(table_item["table_rows"][2][0] == "", f"vertical continuation cell should be an empty grid placeholder: {table_item}")
    merges = table_item.get("table_merges") or []
    assert_true(
        {"row": 0, "col": 0, "rowspan": 1, "colspan": 2} in merges,
        f"horizontal merge was not captured: {merges}",
    )
    assert_true(
        {"row": 1, "col": 0, "rowspan": 2, "colspan": 1} in merges,
        f"vertical merge was not captured: {merges}",
    )


@case
def content_parser_repairs_orphan_vmerge_without_losing_text() -> None:
    work = new_workdir("parser_orphan_vmerge_repair")
    docx = work / "orphan_vmerge.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Top left"
    table.cell(0, 1).text = "Top right"
    table.cell(1, 0).text = "Orphan continue should stay visible"
    table.cell(1, 1).text = "Bottom right"
    doc.save(docx)

    def inject_orphan_vmerge(xml: str) -> str:
        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        root = ET.fromstring(xml.encode("utf-8"))
        table_el = root.find(".//" + w_ns + "tbl")
        assert_true(table_el is not None, "test table missing")
        rows = table_el.findall(w_ns + "tr")
        assert_true(len(rows) == 2, "test rows missing")
        second_row_cells = rows[1].findall(w_ns + "tc")
        assert_true(len(second_row_cells) == 2, "test second row changed")
        tc_pr = second_row_cells[0].find(w_ns + "tcPr")
        if tc_pr is None:
            tc_pr = ET.Element(w_ns + "tcPr")
            second_row_cells[0].insert(0, tc_pr)
        ET.SubElement(tc_pr, w_ns + "vMerge")
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(docx, "word/document.xml", inject_orphan_vmerge)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    items = [item for sec in content.get("sections") or [] for item in sec.get("paragraphs") or []]
    table_items = [item for item in items if isinstance(item, dict) and item.get("role") == "table"]
    assert_true(len(table_items) == 1, f"orphan-vMerge table was not preserved as a table: {items}")
    table_item = table_items[0]
    assert_true(
        table_item.get("table_rows") == [
            ["Top left", "Top right"],
            ["Orphan continue should stay visible", "Bottom right"],
        ],
        f"orphan vMerge visible text was not repaired into a normal cell: {table_item}",
    )
    assert_true(not table_item.get("table_merges"), f"orphan vMerge should not create a fake merge: {table_item}")

    result = run_generated_case("parser_orphan_vmerge_repair_generated", content, base_format())
    assert_true(
        "Orphan continue should stay visible" in result["xml"],
        "generated DOCX lost visible text from repaired orphan vMerge cell",
    )
    assert_true(result["report"]["passed"] is True, f"orphan vMerge repair render should pass QA: {result['report']}")


@case
def content_parser_repairs_mismatched_vmerge_span_without_losing_text() -> None:
    work = new_workdir("parser_mismatched_vmerge_span_repair")
    docx = work / "mismatched_vmerge_span.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "Wide vertical start"
    table.cell(0, 2).text = "Top right"
    table.cell(1, 0).text = "Mismatched continuation should stay visible"
    table.cell(1, 1).text = "Bottom middle"
    table.cell(1, 2).text = "Bottom right"
    doc.save(docx)

    def inject_mismatched_vmerge(xml: str) -> str:
        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        root = ET.fromstring(xml.encode("utf-8"))
        table_el = root.find(".//" + w_ns + "tbl")
        assert_true(table_el is not None, "test table missing")
        rows = table_el.findall(w_ns + "tr")
        assert_true(len(rows) == 2, "test rows missing")
        first_row_cells = rows[0].findall(w_ns + "tc")
        second_row_cells = rows[1].findall(w_ns + "tc")
        assert_true(len(first_row_cells) == 2 and len(second_row_cells) == 3, "test merge grid changed")

        def ensure_tc_pr(cell):
            tc_pr = cell.find(w_ns + "tcPr")
            if tc_pr is None:
                tc_pr = ET.Element(w_ns + "tcPr")
                cell.insert(0, tc_pr)
            return tc_pr

        restart = ET.SubElement(ensure_tc_pr(first_row_cells[0]), w_ns + "vMerge")
        restart.set(w_ns + "val", "restart")
        ET.SubElement(ensure_tc_pr(second_row_cells[0]), w_ns + "vMerge")
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(docx, "word/document.xml", inject_mismatched_vmerge)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    items = [item for sec in content.get("sections") or [] for item in sec.get("paragraphs") or []]
    table_items = [item for item in items if isinstance(item, dict) and item.get("role") == "table"]
    assert_true(len(table_items) == 1, f"mismatched-vMerge table was not preserved as a table: {items}")
    table_item = table_items[0]
    assert_true(
        table_item.get("table_rows") == [
            ["Wide vertical start", "", "Top right"],
            ["Mismatched continuation should stay visible", "Bottom middle", "Bottom right"],
        ],
        f"mismatched vMerge visible text was not repaired into a normal cell: {table_item}",
    )
    merges = table_item.get("table_merges") or []
    assert_true(
        {"row": 0, "col": 0, "rowspan": 1, "colspan": 2} in merges,
        f"valid horizontal part of mismatched table should still be preserved: {table_item}",
    )
    assert_true(
        not any(int(merge.get("rowspan") or 1) > 1 for merge in merges),
        f"mismatched vMerge should not create a fake vertical merge: {table_item}",
    )

    result = run_generated_case("parser_mismatched_vmerge_span_repair_generated", content, base_format())
    assert_true(
        "Mismatched continuation should stay visible" in result["xml"],
        "generated DOCX lost visible text from repaired mismatched vMerge cell",
    )
    assert_true(result["report"]["passed"] is True, f"mismatched vMerge repair render should pass QA: {result['report']}")


@case
def content_parser_preserves_grid_before_vmerge_cells() -> None:
    work = new_workdir("parser_grid_before_vmerge")
    docx = work / "grid_before_vmerge.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Left"
    table.cell(0, 1).text = "Vertical start"
    table.cell(1, 1).text = "Vertical continue"
    doc.save(docx)

    def rewrite(xml: str) -> str:
        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        root = ET.fromstring(xml.encode("utf-8"))
        table_el = root.find(".//" + w_ns + "tbl")
        assert_true(table_el is not None, "test table missing")
        rows = table_el.findall(w_ns + "tr")
        assert_true(len(rows) == 2, "test rows missing")
        first_row_cells = rows[0].findall(w_ns + "tc")
        second_row_cells = rows[1].findall(w_ns + "tc")
        assert_true(len(first_row_cells) == 2 and len(second_row_cells) == 2, "test cell grid changed")

        def ensure_tc_pr(cell):
            tc_pr = cell.find(w_ns + "tcPr")
            if tc_pr is None:
                tc_pr = ET.Element(w_ns + "tcPr")
                cell.insert(0, tc_pr)
            return tc_pr

        restart = ET.SubElement(ensure_tc_pr(first_row_cells[1]), w_ns + "vMerge")
        restart.set(w_ns + "val", "restart")
        rows[1].remove(second_row_cells[0])
        tr_pr = ET.Element(w_ns + "trPr")
        grid_before = ET.SubElement(tr_pr, w_ns + "gridBefore")
        grid_before.set(w_ns + "val", "1")
        rows[1].insert(0, tr_pr)
        ET.SubElement(ensure_tc_pr(second_row_cells[1]), w_ns + "vMerge")
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(docx, "word/document.xml", rewrite)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    items = [item for sec in content.get("sections") or [] for item in sec.get("paragraphs") or []]
    table_items = [item for item in items if isinstance(item, dict) and item.get("role") == "table"]
    assert_true(len(table_items) == 1, f"gridBefore table was not preserved as a table: {items}")
    table_item = table_items[0]
    assert_true(
        table_item.get("table_rows") == [["Left", "Vertical start"], ["", ""]],
        f"gridBefore row offset was not preserved: {table_item}",
    )
    assert_true(
        {"row": 0, "col": 1, "rowspan": 2, "colspan": 1} in (table_item.get("table_merges") or []),
        f"gridBefore vertical merge was not captured: {table_item}",
    )


@case
def content_parser_preserves_table_column_widths() -> None:
    work = new_workdir("parser_table_widths")
    docx = work / "table_widths.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Description"
    table.cell(0, 2).text = "Value"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "Longer explanatory text"
    table.cell(1, 2).text = "1"
    doc.save(docx)

    def inject_widths(xml: str) -> str:
        grid = '<w:tblGrid><w:gridCol w:w="1200"/><w:gridCol w:w="2800"/><w:gridCol w:w="1600"/></w:tblGrid>'
        return re.sub(r"<w:tblGrid>.*?</w:tblGrid>", grid, xml, count=1, flags=re.S)

    _rewrite_docx_part(docx, "word/document.xml", inject_widths)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    items = [item for sec in content.get("sections") or [] for item in sec.get("paragraphs") or []]
    table_items = [item for item in items if isinstance(item, dict) and item.get("role") == "table"]
    assert_true(len(table_items) == 1, f"table with widths was not preserved as a table: {items}")
    assert_true(
        table_items[0].get("table_col_widths_twips") == [1200, 2800, 1600],
        f"table column widths were not extracted: {table_items[0]}",
    )


@case
def content_parser_repairs_gridspan_beyond_tblgrid_widths() -> None:
    work = new_workdir("parser_gridspan_beyond_tblgrid")
    docx = work / "gridspan_beyond_tblgrid.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "Overflow span should stay readable"
    table.cell(0, 1).text = "removed"
    table.cell(0, 2).text = "removed"
    doc.save(docx)

    def inject_overflow_span(xml: str) -> str:
        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        root = ET.fromstring(xml.encode("utf-8"))
        table_el = root.find(".//" + w_ns + "tbl")
        assert_true(table_el is not None, "test table missing")
        old_grid = table_el.find(w_ns + "tblGrid")
        if old_grid is not None:
            table_el.remove(old_grid)
        tbl_grid = ET.Element(w_ns + "tblGrid")
        for width in (2000, 2000):
            grid_col = ET.SubElement(tbl_grid, w_ns + "gridCol")
            grid_col.set(w_ns + "w", str(width))
        insert_at = 1 if len(list(table_el)) and table_el[0].tag == w_ns + "tblPr" else 0
        table_el.insert(insert_at, tbl_grid)

        row = table_el.find(w_ns + "tr")
        assert_true(row is not None, "test row missing")
        cells = row.findall(w_ns + "tc")
        assert_true(len(cells) == 3, "test cells missing")
        tc_pr = cells[0].find(w_ns + "tcPr")
        if tc_pr is None:
            tc_pr = ET.Element(w_ns + "tcPr")
            cells[0].insert(0, tc_pr)
        tc_w = tc_pr.find(w_ns + "tcW")
        if tc_w is None:
            tc_w = ET.SubElement(tc_pr, w_ns + "tcW")
        tc_w.set(w_ns + "type", "dxa")
        tc_w.set(w_ns + "w", "8400")
        grid_span = tc_pr.find(w_ns + "gridSpan")
        if grid_span is None:
            grid_span = ET.SubElement(tc_pr, w_ns + "gridSpan")
        grid_span.set(w_ns + "val", "3")
        row.remove(cells[2])
        row.remove(cells[1])
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(docx, "word/document.xml", inject_overflow_span)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    items = [item for sec in content.get("sections") or [] for item in sec.get("paragraphs") or []]
    table_items = [item for item in items if isinstance(item, dict) and item.get("role") == "table"]
    assert_true(len(table_items) == 1, f"gridSpan-overflow table was not preserved as a table: {items}")
    table_item = table_items[0]
    assert_true(
        table_item.get("table_rows") == [["Overflow span should stay readable", "", ""]],
        f"gridSpan overflow did not expand to a stable visible grid: {table_item}",
    )
    assert_true(
        {"row": 0, "col": 0, "rowspan": 1, "colspan": 3} in (table_item.get("table_merges") or []),
        f"overflow gridSpan merge was not preserved: {table_item}",
    )
    widths = table_item.get("table_col_widths_twips") or []
    assert_true(len(widths) == 3 and all(int(width or 0) > 0 for width in widths), f"overflow gridSpan produced zero-width columns: {table_item}")

    result = run_generated_case("parser_gridspan_beyond_tblgrid_generated", content, base_format())
    grid_widths = [int(value) for value in re.findall(r"<w:gridCol\b[^>]*w:w=\"(\d+)\"", result["xml"])]
    assert_true(
        len(grid_widths) >= 3 and all(width > 0 for width in grid_widths[:3]),
        f"generated DOCX kept a zero-width repaired overflow grid column: {grid_widths[:3]}",
    )
    assert_true(
        "Overflow span should stay readable" in result["xml"],
        "generated DOCX lost visible text from gridSpan-overflow cell",
    )
    assert_true(result["report"]["passed"] is True, f"gridSpan overflow repair render should pass QA: {result['report']}")


@case
def content_parser_preserves_landscape_section_for_wide_table() -> None:
    work = new_workdir("parser_landscape_section_wide_table")
    docx = work / "landscape_section_wide_table.docx"
    doc = Document()
    doc.add_paragraph("1 Landscape table")
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    table = doc.add_table(rows=2, cols=9)
    for ci in range(9):
        table.cell(0, ci).text = f"H{ci + 1}"
        table.cell(1, ci).text = f"Wide cell {ci + 1}"
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    items = [item for sec in content.get("sections") or [] for item in sec.get("paragraphs") or []]
    table_items = [item for item in items if isinstance(item, dict) and item.get("role") == "table"]
    assert_true(len(table_items) == 1, f"landscape-section table was not preserved as a table: {items}")
    table_item = table_items[0]
    source_setup = table_item.get("source_section_page_setup") or {}
    assert_true(
        source_setup.get("orientation") == "landscape",
        f"landscape section setup was not attached to the table: {table_item}",
    )
    assert_true(
        int(source_setup.get("page_width_twips") or 0) > int(source_setup.get("page_height_twips") or 0),
        f"landscape page dimensions were not preserved: {source_setup}",
    )

    result = run_generated_case("parser_landscape_section_wide_table_generated", content, base_format())
    assert_true('w:orient="landscape"' in result["xml"], "generated DOCX did not contain a landscape table section")
    assert_true(
        result["xml"].count('w:orient="landscape"') == 1,
        "generated DOCX should restore portrait template orientation after the landscape table",
    )
    grid_widths = [int(value) for value in re.findall(r"<w:gridCol\b[^>]*w:w=\"(\d+)\"", result["xml"])]
    assert_true(
        grid_widths[:9] == table_item.get("table_col_widths_twips"),
        f"landscape table widths should use source landscape text width, got {grid_widths[:9]}",
    )
    assert_true(
        result["manifest"]["counts"].get("content_landscape_table_sections_rendered", 0) == 1,
        f"landscape table section render count missing: {result['manifest']}",
    )
    assert_true(result["report"]["passed"] is True, f"landscape wide table render should pass QA: {result['report']}")


@case
def content_parser_preserves_table_layout_details() -> None:
    work = new_workdir("parser_table_layout_details")
    docx = work / "table_layout_details.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Body A"
    table.cell(1, 1).text = "Body B"
    doc.save(docx)

    def inject_layout(xml: str) -> str:
        xml = re.sub(
            r"(<w:tblPr>)",
            (
                r"\1<w:tblCellMar>"
                r'<w:top w:w="80" w:type="dxa"/>'
                r'<w:left w:w="120" w:type="dxa"/>'
                r'<w:bottom w:w="90" w:type="dxa"/>'
                r'<w:right w:w="140" w:type="dxa"/>'
                r"</w:tblCellMar>"
            ),
            xml,
            count=1,
        )
        row_index = 0

        def inject_row(match: re.Match[str]) -> str:
            nonlocal row_index
            row_index += 1
            if row_index == 1:
                props = '<w:trPr><w:trHeight w:val="480" w:hRule="exact"/><w:tblHeader w:val="true"/></w:trPr>'
            elif row_index == 2:
                props = '<w:trPr><w:trHeight w:val="360" w:hRule="atLeast"/></w:trPr>'
            else:
                props = ""
            return match.group(0) + props

        xml = re.sub(r"<w:tr(?:\s[^>]*)?>", inject_row, xml, count=2)
        xml = re.sub(
            r"(<w:tcPr>)",
            (
                r'\1<w:vAlign w:val="top"/>'
                r"<w:tcMar>"
                r'<w:top w:w="40" w:type="dxa"/>'
                r'<w:left w:w="60" w:type="dxa"/>'
                r'<w:bottom w:w="40" w:type="dxa"/>'
                r'<w:right w:w="60" w:type="dxa"/>'
                r"</w:tcMar>"
            ),
            xml,
            count=1,
        )
        return xml

    _rewrite_docx_part(docx, "word/document.xml", inject_layout)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    items = [item for sec in content.get("sections") or [] for item in sec.get("paragraphs") or []]
    table_items = [item for item in items if isinstance(item, dict) and item.get("role") == "table"]
    assert_true(len(table_items) == 1, f"layout-rich table was not preserved as a table: {items}")
    table_item = table_items[0]
    assert_true(
        table_item.get("table_row_heights_twips") == [{"val": 480, "rule": "exact"}, {"val": 360, "rule": "atLeast"}],
        f"table row heights were not extracted: {table_item}",
    )
    assert_true(table_item.get("table_repeat_header_rows") == 1, f"repeat header rows were not extracted: {table_item}")
    assert_true(
        table_item.get("table_cell_margins_twips") == {"top": 80, "left": 120, "bottom": 90, "right": 140},
        f"table default cell margins were not extracted: {table_item}",
    )
    overrides = table_item.get("table_cell_overrides") or []
    assert_true(
        {
            "row": 0,
            "col": 0,
            "v_align": "top",
            "margins_twips": {"top": 40, "left": 60, "bottom": 40, "right": 60},
        }
        in overrides,
        f"cell-specific layout override was not extracted: {overrides}",
    )


@case
def content_parser_preserves_table_border_details() -> None:
    work = new_workdir("parser_table_border_details")
    docx = work / "table_border_details.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Body A"
    table.cell(1, 1).text = "Body B"
    doc.save(docx)

    def inject_borders(xml: str) -> str:
        xml = re.sub(
            r"(<w:tblPr>)",
            (
                r"\1<w:tblBorders>"
                r'<w:top w:val="double" w:sz="12" w:color="4472C4" w:space="0"/>'
                r'<w:insideH w:val="single" w:sz="4" w:color="808080" w:space="0"/>'
                r"</w:tblBorders>"
            ),
            xml,
            count=1,
        )
        return re.sub(
            r"(<w:tcPr>)",
            (
                r"\1<w:tcBorders>"
                r'<w:bottom w:val="dashed" w:sz="6" w:color="C00000" w:space="0"/>'
                r'<w:right w:val="nil" w:sz="0" w:color="000000" w:space="0"/>'
                r"</w:tcBorders>"
            ),
            xml,
            count=1,
        )

    _rewrite_docx_part(docx, "word/document.xml", inject_borders)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    items = [item for sec in content.get("sections") or [] for item in sec.get("paragraphs") or []]
    table_items = [item for item in items if isinstance(item, dict) and item.get("role") == "table"]
    assert_true(len(table_items) == 1, f"border-rich table was not preserved as a table: {items}")
    table_item = table_items[0]
    table_borders = table_item.get("table_borders") or {}
    assert_true(
        table_borders.get("top") == {"val": "double", "sz": "12", "color": "4472C4", "space": "0"},
        f"table top border was not extracted: {table_item}",
    )
    assert_true(
        table_borders.get("insideH") == {"val": "single", "sz": "4", "color": "808080", "space": "0"},
        f"table insideH border was not extracted: {table_item}",
    )
    overrides = table_item.get("table_cell_overrides") or []
    first_cell = next((entry for entry in overrides if entry.get("row") == 0 and entry.get("col") == 0), {})
    assert_true(
        first_cell.get("borders", {}).get("bottom") == {"val": "dashed", "sz": "6", "color": "C00000", "space": "0"},
        f"cell bottom border was not extracted: {overrides}",
    )
    assert_true(
        first_cell.get("borders", {}).get("right") == {"val": "nil", "sz": "0", "color": "000000", "space": "0"},
        f"cell nil right border was not extracted: {overrides}",
    )


@case
def content_parser_preserves_two_level_nested_tables_in_cells() -> None:
    work = new_workdir("parser_two_level_nested_table_cell")
    img = work / "nested_image.png"
    img.write_bytes(PNG_1X1)
    docx = work / "nested_table_cell.docx"
    doc = Document()
    outer = doc.add_table(rows=2, cols=2)
    outer.cell(0, 0).text = "Outer A"
    outer.cell(0, 1).text = "Outer B"
    nested_host = outer.cell(1, 0)
    nested_host.text = "Nested before"
    outer.cell(1, 1).text = "Outer D"
    nested = nested_host.add_table(rows=2, cols=2)
    nested.cell(0, 0).text = "Nested A"
    nested.cell(0, 1).text = "Nested B"
    nested.cell(1, 0).text = "Nested C"
    nested.cell(1, 1).text = "Nested D"
    deeper = nested.cell(1, 1).add_table(rows=1, cols=2)
    deeper.cell(0, 0).text = "Deeper A"
    deeper.cell(0, 1).text = "Deeper B"
    deeper.cell(0, 1).paragraphs[0].add_run().add_picture(str(img))
    nested_host.add_paragraph("Nested after")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    assert_true(content.get("_meta", {}).get("tables_count") == 3, f"nested table count missing: {content.get('_meta')}")
    assert_true(content.get("_meta", {}).get("images_extracted") == 1, f"nested table image count missing: {content.get('_meta')}")
    items = [item for sec in content.get("sections") or [] for item in sec.get("paragraphs") or []]
    table_items = [item for item in items if isinstance(item, dict) and item.get("role") == "table"]
    assert_true(len(table_items) == 1, f"nested table should stay inside the outer table cell: {items}")
    outer_item = table_items[0]
    assert_true(
        outer_item.get("table_rows", [])[1][0] == "Nested before\nNested after",
        f"nested host text should preserve before/after paragraphs: {outer_item}",
    )
    cell_items = outer_item.get("table_cell_items") or []
    nested_cell = next((entry for entry in cell_items if entry.get("row") == 1 and entry.get("col") == 0), {})
    nested_items = [entry for entry in nested_cell.get("items") or [] if isinstance(entry, dict) and entry.get("role") == "table"]
    assert_true(nested_items, f"nested table was not attached to the parent cell: {outer_item}")
    nested_item = nested_items[0]
    assert_true(nested_item.get("location") == "nested_table_cell", f"nested table location missing: {nested_item}")
    assert_true(nested_item.get("after_paragraph_index") == 1, f"nested table insertion position missing: {nested_item}")
    assert_true(
        nested_item.get("table_rows") == [["Nested A", "Nested B"], ["Nested C", "Nested D"]],
        f"nested table rows changed: {nested_item}",
    )
    deeper_cell_items = nested_item.get("table_cell_items") or []
    deeper_cell = next((entry for entry in deeper_cell_items if entry.get("row") == 1 and entry.get("col") == 1), {})
    deeper_items = [entry for entry in deeper_cell.get("items") or [] if isinstance(entry, dict) and entry.get("role") == "table"]
    assert_true(deeper_items, f"second-level nested table was not attached to the nested cell: {nested_item}")
    assert_true(
        deeper_items[0].get("table_rows") == [["Deeper A", "Deeper B"]],
        f"second-level nested table rows changed: {deeper_items[0]}",
    )
    deeper_image_items = [
        item
        for entry in deeper_items[0].get("table_cell_items") or []
        for item in entry.get("items") or []
        if isinstance(item, dict) and item.get("role") == "image"
    ]
    assert_true(deeper_image_items, f"nested table-cell image was not attached to the nested cell: {deeper_items[0]}")
    assert_true(
        deeper_image_items[0].get("location") == "table_cell",
        f"nested table-cell image should keep table-cell origin: {deeper_image_items}",
    )
    top_level_table_images = [
        item
        for item in items
        if isinstance(item, dict) and item.get("role") == "image" and item.get("location") == "table_cell"
    ]
    assert_true(not top_level_table_images, f"nested table-cell image leaked out as a body image: {items}")

    result = run_generated_case("parser_two_level_nested_table_cell_generated", content, base_format())
    assert_true(result["xml"].count("<w:drawing>") == 1, "extracted nested table-cell image should render exactly once")
    table_xmls = re.findall(r"<w:tbl\b.*?</w:tbl>", result["xml"], flags=re.S)
    assert_true(
        sum(table_xml.count("<w:drawing>") for table_xml in table_xmls) >= 1,
        "extracted nested table-cell image rendered outside generated Word tables",
    )
    assert_true(
        result["manifest"]["counts"]["content_images_rendered"] == 1,
        f"extracted nested table-cell image render count changed: {result['manifest']}",
    )
    assert_true(result["report"]["passed"] is True, f"extracted nested table-cell image should pass QA: {result['report']}")


@case
def content_parser_preserves_three_level_nested_tables_in_cells() -> None:
    work = new_workdir("parser_three_level_nested_table_cell")
    docx = work / "three_level_nested_table_cell.docx"
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    outer_host = outer.cell(0, 0)
    outer_host.text = "Outer before"
    nested = outer_host.add_table(rows=1, cols=1)
    nested_host = nested.cell(0, 0)
    nested_host.text = "Nested before"
    deeper = nested_host.add_table(rows=1, cols=1)
    deeper_host = deeper.cell(0, 0)
    deeper_host.text = "Deeper before"
    deepest = deeper_host.add_table(rows=1, cols=1)
    deepest.cell(0, 0).text = "Deepest value"
    deeper_host.add_paragraph("Deeper after")
    nested_host.add_paragraph("Nested after")
    outer_host.add_paragraph("Outer after")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    assert_true(content.get("_meta", {}).get("tables_count") == 4, f"three-level nested table count missing: {content.get('_meta')}")
    items = [item for sec in content.get("sections") or [] for item in sec.get("paragraphs") or []]
    table_items = [item for item in items if isinstance(item, dict) and item.get("role") == "table"]
    assert_true(len(table_items) == 1, f"deeply nested tables should stay inside the outer table cell: {items}")
    outer_item = table_items[0]
    assert_true(
        outer_item.get("table_rows") == [["Outer before\nOuter after"]],
        f"outer cell direct text changed: {outer_item}",
    )
    nested_item = next(
        (
            item
            for entry in outer_item.get("table_cell_items") or []
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "table"
        ),
        None,
    )
    assert_true(nested_item, f"first nested table was not attached to the outer cell: {outer_item}")
    assert_true(nested_item.get("table_rows") == [["Nested before\nNested after"]], f"first nested table text changed: {nested_item}")
    deeper_item = next(
        (
            item
            for entry in nested_item.get("table_cell_items") or []
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "table"
        ),
        None,
    )
    assert_true(deeper_item, f"second nested table was not attached to the nested cell: {nested_item}")
    assert_true(deeper_item.get("table_rows") == [["Deeper before\nDeeper after"]], f"second nested table text changed: {deeper_item}")
    deepest_item = next(
        (
            item
            for entry in deeper_item.get("table_cell_items") or []
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "table"
        ),
        None,
    )
    assert_true(deepest_item, f"third nested table was not attached to the deeper cell: {deeper_item}")
    assert_true(deepest_item.get("table_rows") == [["Deepest value"]], f"third nested table rows changed: {deepest_item}")

    result = run_generated_case("parser_three_level_nested_table_cell_generated", content, base_format())
    assert_true("Deepest value" in result["xml"], "third-level nested table text did not render")
    assert_true(
        result["manifest"]["counts"].get("content_nested_tables_rendered", 0) >= 3,
        f"three nested table render count missing: {result['manifest']}",
    )
    assert_true(result["report"]["passed"] is True, f"three-level nested table render should pass QA: {result['report']}")


@case
def content_parser_preserves_nested_table_cell_inline_image_formula_and_footnote_order() -> None:
    work = new_workdir("parser_nested_table_cell_inline_image_formula_note")
    img = work / "nested_inline_image.png"
    write_sample_png(img, width=128, height=96)
    docx = work / "nested_table_cell_inline_image_formula_note.docx"
    doc = Document()
    doc.add_paragraph("1 Nested table inline content")
    outer = doc.add_table(rows=1, cols=1)
    host = outer.cell(0, 0)
    host.text = "Outer before"
    nested = host.add_table(rows=1, cols=1)
    para = nested.cell(0, 0).paragraphs[0]
    para.add_run("NestedLead")
    para.add_run().add_picture(str(img))
    para.add_run(r"Nested formula $a=1$ and ")
    para._element.append(etree.fromstring(latex_to_omath(r"b=2", display=False).encode("utf-8")))
    para.add_run(" noted")
    host.add_paragraph("Outer after")
    doc.save(docx)

    footnote_text = "Nested table-cell note must stay after the formula."

    def inject_reference(xml: str) -> str:
        replacements = [
            (
                '<w:t xml:space="preserve"> noted</w:t></w:r></w:p>',
                '<w:t xml:space="preserve"> noted</w:t></w:r><w:r><w:footnoteReference w:id="5"/></w:r></w:p>',
            ),
            (
                "<w:t> noted</w:t></w:r></w:p>",
                '<w:t> noted</w:t></w:r><w:r><w:footnoteReference w:id="5"/></w:r></w:p>',
            ),
        ]
        for old, new in replacements:
            if old in xml:
                return xml.replace(old, new, 1)
        return xml

    _rewrite_docx_part(docx, "word/document.xml", inject_reference)
    with zipfile.ZipFile(docx, "a") as zf:
        zf.writestr(
            "word/footnotes.xml",
            (
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
                '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
                f'<w:footnote w:id="5"><w:p><w:r><w:t>{footnote_text}</w:t></w:r></w:p></w:footnote>'
                "</w:footnotes>"
            ),
        )

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(len(table_items) == 1, f"nested table should stay inside the outer table cell: {paragraphs}")
    outer_item = table_items[0]
    assert_true(
        outer_item.get("table_rows") == [["Outer before\nOuter after"]],
        f"outer cell direct text should not absorb nested content: {outer_item}",
    )
    host_items = [
        item
        for entry in outer_item.get("table_cell_items") or []
        if entry.get("row") == 0 and entry.get("col") == 0
        for item in entry.get("items") or []
        if isinstance(item, dict)
    ]
    nested_items = [item for item in host_items if item.get("role") == "table"]
    assert_true(nested_items, f"nested table was not attached to the parent cell: {outer_item}")
    nested_item = nested_items[0]
    assert_true(
        nested_item.get("table_rows") == [["NestedLead\nNested formula $a=1$ and b=2 noted"]],
        f"nested table-cell text/image boundary changed: {nested_item}",
    )
    nested_cell_items = nested_item.get("table_cell_items") or []
    image_item = next(
        (
            item
            for entry in nested_cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "image"
        ),
        None,
    )
    rich_item = next(
        (
            item
            for entry in nested_cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "rich_text"
        ),
        None,
    )
    assert_true(image_item and image_item.get("after_paragraph_index") == 1, f"nested inline cell image position changed: {nested_cell_items}")
    assert_true(
        rich_item and rich_item.get("replace_paragraph_index") == 1,
        f"nested formula/note rich text should replace text after image: {nested_item}",
    )
    assert_true(
        [run.get("type") for run in rich_item.get("runs") or []] == ["text", "math", "text", "math", "text", "note_ref"],
        f"nested formula/note run order changed: {rich_item}",
    )

    result = run_generated_case("nested_table_cell_inline_image_formula_note_render", content, base_format())
    xml = result["xml"]
    assert_true("$a=1$" not in xml, "nested table-cell formula leaked LaTeX delimiters into generated XML")
    assert_true(omath_count(xml) >= 2, "nested table-cell LaTeX/OMML formulas did not both render as native math")
    assert_true(xml.count("<w:drawing>") == 1, f"nested inline cell image did not render exactly once: {result['manifest']}")
    assert_true("<w:footnoteReference" in xml, "nested table-cell footnote did not render as native reference")
    assert_true(
        xml.find("NestedLead") < xml.find("<w:drawing>") < xml.find("Nested formula") < xml.find("<w:footnoteReference"),
        "nested table-cell image/formula/footnote source order changed",
    )
    assert_true(
        result["manifest"]["counts"].get("content_nested_tables_rendered", 0) >= 1,
        f"nested table render count missing: {result['manifest']}",
    )
    assert_true(result["manifest"]["counts"].get("inline_formulas_rendered", 0) >= 2, f"inline formula count missing: {result['manifest']}")
    assert_true(result["manifest"]["counts"].get("footnote_references_rendered") == 1, f"footnote count missing: {result['manifest']}")
    assert_true(result["report"]["passed"] is True, f"nested inline image/formula/note table-cell render should pass QA: {result['report']}")


@case
def content_parser_does_not_count_cover_table_as_body_table() -> None:
    work = new_workdir("parser_cover_table_count")
    docx = work / "cover_table_count.docx"
    doc = Document()
    cover = doc.add_table(rows=1, cols=2)
    cover.cell(0, 0).text = "\u5b66\u6821\u7f16\u7801"
    cover.cell(0, 1).text = "10001"
    title = "\u9762\u5411\u7eff\u7535\u6d88\u7eb3\u7684\u591a\u6e90\u80fd\u6e90\u8c03\u5ea6\u7814\u7a76"
    p = doc.add_paragraph(title)
    p.runs[0].font.size = Pt(16)
    doc.add_paragraph("\u6458\u8981\uff1a\u672c\u6587\u6458\u8981\u7528\u4e8e\u9a8c\u8bc1\u5c01\u9762\u8868\u683c\u4e0d\u5e94\u8ba1\u5165\u6b63\u6587\u8868\u683c\u3002")
    doc.add_paragraph("\u5173\u952e\u8bcd\uff1a\u7eff\u7535\uff1b\u8c03\u5ea6")
    doc.add_heading("\u7b2c1\u7ae0 \u7eea\u8bba", 1)
    doc.add_paragraph("\u6b63\u6587\u6bb5\u843d\u3002")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    meta = content.get("_meta") or {}
    assert_true(meta.get("source_tables_count") == 1, f"source table count was not recorded: {meta}")
    assert_true(meta.get("tables_count") == 0, f"cover table was counted as a body table: {meta}")


@case
def content_parser_extracts_cover_paragraph_fields_and_skips_template_notes() -> None:
    work = new_workdir("parser_cover_paragraph_fields")
    docx = work / "cover_paragraph_fields.docx"
    doc = Document()
    doc.add_paragraph("本科毕业论文（设计）")
    doc.add_paragraph("年级专业：English Major in Education (2020-2024), Sanming University")
    doc.add_paragraph("姓    名：Zhang San")
    doc.add_paragraph("学    号：2020123456")
    doc.add_paragraph("指导教师：Prof. Li Si")
    doc.add_paragraph("（完成时间按照答辩时间填写）")
    doc.add_paragraph("2026年 5月 6日")
    title = "On Chinese Translation of English Sports News Headlines from the Perspective of Memetics"
    title_p = doc.add_paragraph(title)
    for run in title_p.runs:
        run.bold = True
        run.font.size = Pt(16)
    doc.add_paragraph("Zhang San")
    doc.add_paragraph("English Major in Education (2020-2024), Sanming University, Sanming, Fujian")
    doc.add_paragraph("Abstract: This is the real abstract paragraph and it should be the first front-matter content.")
    note = "摘要是论文内容的总结概括，主要概述论题的背景、目的、主要内容及结论。约200词，第三人称，不标注引用编号。"
    note_p = doc.add_paragraph(note)
    for run in note_p.runs:
        run.bold = True
        run.font.size = Pt(14)
    doc.add_paragraph("KEY WORDS: memetics; translation")
    doc.add_paragraph("1. Introduction")
    doc.add_paragraph("Actual body text should remain.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    cover_info = content.get("cover_info") or {}
    assert_true(cover_info.get("class_name", "").startswith("English Major"), f"class/major field missing: {cover_info}")
    assert_true(cover_info.get("student_name") == "Zhang San", f"name field missing: {cover_info}")
    assert_true(cover_info.get("student_id") == "2020123456", f"student id missing: {cover_info}")
    assert_true(cover_info.get("advisor") == "Prof. Li Si", f"advisor missing: {cover_info}")
    assert_true(cover_info.get("completion_date") == "2026年 5月 6日", f"date missing: {cover_info}")
    assert_true(cover_info.get("paper_title") == title, f"title missing from cover info: {cover_info}")

    all_text = "\n".join(
        [str(sec.get("heading") or "") for sec in content.get("sections") or []]
        + [
            str(item.get("text") if isinstance(item, dict) else item)
            for sec in content.get("sections") or []
            for item in sec.get("paragraphs") or []
        ]
    )
    assert_true("完成时间按照答辩时间填写" not in all_text, f"date instruction leaked into content: {all_text}")
    assert_true(note not in all_text, f"abstract format instruction leaked into content: {all_text}")
    assert_true("Actual body text should remain." in all_text, f"real body text was lost: {all_text}")


@case
def content_parser_recovers_late_chinese_abstract_from_references() -> None:
    work = new_workdir("parser_late_cn_front_matter")
    docx = work / "late_cn_front_matter.docx"
    doc = Document()
    doc.add_paragraph("English Thesis Title")
    doc.add_paragraph("Abstract: English abstract body.")
    doc.add_paragraph("key words: translation; memetics")
    doc.add_paragraph("1. Introduction")
    doc.add_paragraph("Actual body text should remain.")
    doc.add_paragraph("References")
    doc.add_paragraph("[1] Blackmore, S. The Meme Machine [M]. Oxford: Oxford University Press, 1999.")
    cn_title = "模因论视角下的英语体育新闻标题汉译研究"
    cn_abs = "随着中国的发展，体育新闻标题翻译具有重要意义。"
    cn_kw = "模因论；体育新闻标题；翻译策略"
    doc.add_paragraph(cn_title)
    doc.add_paragraph(f"【摘要】{cn_abs}")
    doc.add_paragraph(f"【关键词】{cn_kw}")
    doc.add_paragraph("Acknowledgements")
    doc.add_paragraph("Thanks.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    refs = content.get("references") or []
    assert_true(len(refs) == 1 and refs[0].startswith("[1]"), f"late Chinese front matter leaked into references: {refs}")
    assert_true(content.get("title_info", {}).get("title_cn") == cn_title, f"late Chinese title not recovered: {content.get('title_info')}")
    roles = [sec.get("role") for sec in content.get("sections") or []]
    assert_true("cn_abstract" in roles and "cn_keywords" in roles, f"late Chinese front matter sections missing: {roles}")
    cn_sections = {sec.get("role"): sec for sec in content.get("sections") or []}
    assert_true(cn_abs in "\n".join(cn_sections["cn_abstract"].get("paragraphs") or []), "Chinese abstract body missing")
    assert_true(cn_kw in "\n".join(cn_sections["cn_keywords"].get("paragraphs") or []), "Chinese keywords missing")


@case
def low_res_image_fragment_is_reported_and_not_upscaled() -> None:
    img_src = new_workdir("tiny_image_src")
    write_sample_png(img_src / "tiny.png", width=76, height=18)
    content = base_content([
        {"role": "figure", "image": "tiny.png", "caption": "Figure 1 broken label shard"}
    ])
    content["_meta"]["images_dir"] = str(img_src)
    content["_meta"]["images_extracted"] = 1
    content["sections"][0]["images"] = ["tiny.png"]
    result = run_generated_case("tiny_image_fragment", content)
    codes = [item["code"] for item in result["report"]["issues"]]
    assert_true("LOW_RES_IMAGE_FRAGMENT" in codes, "QA did not report a low-resolution image fragment")
    issue = next(item for item in result["report"]["issues"] if item["code"] == "LOW_RES_IMAGE_FRAGMENT")
    assert_true(issue["severity"] == "warning", f"contained image fragment should be a warning: {issue}")
    assert_true(result["manifest"]["counts"].get("content_image_fragments_contained") == 1, "contained image fragment was not counted")
    m = re.search(r"<wp:extent[^>]+cx=\"(\d+)\"[^>]+cy=\"(\d+)\"", result["xml"])
    assert_true(bool(m), "generated DOCX did not contain an image extent")
    width_inches = int(m.group(1)) / 914400
    assert_true(width_inches < 2.0, f"tiny image was upscaled to page width: {width_inches:.2f}in")


@case
def small_square_icon_is_not_reported_as_low_res_fragment() -> None:
    img_src = new_workdir("small_icon_src")
    write_sample_png(img_src / "icon.png", width=64, height=64)
    content = base_content([
        {"role": "figure", "image": "icon.png", "caption": "QR code icon"}
    ])
    content["_meta"]["images_dir"] = str(img_src)
    content["_meta"]["images_extracted"] = 1
    content["sections"][0]["images"] = ["icon.png"]
    result = run_generated_case("small_square_icon", content)
    codes = [item["code"] for item in result["report"]["issues"]]
    assert_true("LOW_RES_IMAGE_FRAGMENT" not in codes, f"legitimate small square image was reported as fragment: {result['report']['issues']}")


@case
def qa_manifest_detects_missing_image_render() -> None:
    img_src = new_workdir("image_missing_src")
    write_sample_png(img_src / "dot.png")
    content = base_content([
        {"role": "figure", "image": "dot.png", "caption": "Figure 1 sample"}
    ])
    content["_meta"]["images_dir"] = str(img_src)
    content["_meta"]["images_extracted"] = 1
    content["sections"][0]["images"] = ["dot.png"]
    result = run_generated_case("qa_missing_image", content)
    manifest_path = result["work"] / "build_manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["counts"]["content_images_rendered"] = 0
    write_json(manifest_path, manifest)
    report = check_output(str(result["work"]), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("IMAGE_COUNT_MISMATCH" in codes, "QA did not trust manifest for image mismatch")
    repair = report.get("repair_plan") or {}
    assert_true(repair.get("blocking_errors", 0) >= 1, "repair plan did not count blocking errors")
    steps = repair.get("steps") or []
    assert_true(steps and steps[0].get("severity") == "error", "repair plan did not prioritize errors")
    assert_true(any(step.get("code") == "IMAGE_COUNT_MISMATCH" for step in steps), "repair plan omitted image repair guidance")
    write_reports(report, str(result["work"]))
    assert_true((result["work"] / "qa_repair_plan.md").exists(), "repair plan markdown was not written")
    assert_true((result["work"] / "qa_fix_prompt.txt").exists(), "AI fix prompt was not written")


@case
def qa_manifest_detects_missing_formula_render() -> None:
    content = base_content([
        {"role": "formula", "latex": "a=b+c", "text": "a=b+c", "numbered": False}
    ])
    result = run_generated_case("qa_missing_formula", content)
    manifest_path = result["work"] / "build_manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["counts"]["content_formulas_rendered"] = 0
    write_json(manifest_path, manifest)
    report = check_output(str(result["work"]), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("FORMULA_COUNT_MISMATCH" in codes, "QA did not trust manifest for formula mismatch")


@case
def omml_display_item_builds_display_formula() -> None:
    xml = latex_to_omath("a=b+c", display=True)
    content = base_content([
        {
            "role": "formula",
            "source": "omml",
            "text": "a=b+c",
            "math": [{"type": "display", "xml": xml, "text": "a=b+c"}],
            "numbered": False,
        }
    ])
    result = run_generated_case("omml_display", content)
    assert_true(result["manifest"]["counts"]["display_formulas_rendered"] == 1, "OMML display formula was not counted")
    assert_true(omath_para_count(result["xml"]) == 1, "OMML display formula did not remain display")


@case
def content_parser_preserves_mixed_docx_tokens() -> None:
    work = new_workdir("parser_mixed")
    img = work / "dot.png"
    img.write_bytes(PNG_1X1)
    docx = work / "mixed.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("before ")
    p._element.append(etree.fromstring(latex_to_omath("x^2", display=False).encode("utf-8")))
    p.add_run(" after")
    p_img = doc.add_paragraph()
    p_img.add_run("image before ")
    p_img.add_run().add_picture(str(img))
    p_img._element.append(etree.fromstring(latex_to_omath("z^2", display=False).encode("utf-8")))
    p_img.add_run(" image after")
    p_disp = doc.add_paragraph()
    p_disp._element.append(etree.fromstring(latex_to_omath("a=b+c", display=True).encode("utf-8")))
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    rich_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "rich_text"]
    formula_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "formula"]
    image_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "image"]
    assert_true(any([r.get("type") for r in item.get("runs", [])] == ["text", "math", "text"] for item in rich_items), "mixed text/math run order was not preserved")
    assert_true(any((item.get("runs") or [])[0].get("text") == "before " and (item.get("runs") or [])[-1].get("text") == " after" for item in rich_items), "inline formula surrounding spaces were not preserved")
    assert_true(image_items, "image token disappeared from mixed paragraph")
    assert_true(any((item.get("math") or [{}])[0].get("type") == "display" for item in formula_items), "display OMML was not extracted as formula")


@case
def content_parser_extracts_vml_pictures() -> None:
    work = new_workdir("parser_vml")
    img = work / "dot.png"
    img.write_bytes(PNG_1X1)
    src_docx = work / "drawing.docx"
    vml_docx = work / "vml.docx"
    doc = Document()
    doc.add_paragraph().add_run().add_picture(str(img))
    doc.save(src_docx)
    make_vml_picture_docx(src_docx, vml_docx)
    content = extract_docx_content(str(vml_docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    assert_true(any(isinstance(p, dict) and p.get("role") == "image" for p in paragraphs), "VML picture was not extracted")


@case
def content_parser_recovers_textbox_and_content_control_text() -> None:
    work = new_workdir("parser_boxed_text")
    docx = work / "boxed_text.docx"
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Normal paragraph before boxed content.")
    doc.save(docx)

    textbox_text = "Recovered floating textbox paragraph."
    control_text = "Recovered content control paragraph."

    def inject_boxed_content(xml: str) -> str:
        block = (
            '<w:p><w:r><w:pict><v:shape xmlns:v="urn:schemas-microsoft-com:vml">'
            "<v:textbox><w:txbxContent>"
            f"<w:p><w:r><w:t>{textbox_text}</w:t></w:r></w:p>"
            "</w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>"
            "<w:sdt><w:sdtPr><w:tag w:val=\"control\"/></w:sdtPr><w:sdtContent>"
            f"<w:p><w:r><w:t>{control_text}</w:t></w:r></w:p>"
            "</w:sdtContent></w:sdt>"
        )
        return xml.replace("</w:body>", block + "</w:body>")

    _rewrite_docx_part(docx, "word/document.xml", inject_boxed_content)

    content = extract_docx_content(str(docx), output_dir=str(work))
    body_text = _content_plain_text(content)
    meta = content.get("_meta") or {}
    issues = (meta.get("source_audit") or {}).get("issues") or []
    textbox_issue = next((issue for issue in issues if issue.get("code") == "SOURCE_TEXTBOX_UNSUPPORTED"), {})

    assert_true(textbox_text in body_text, f"textbox body text was not recovered: {body_text}")
    assert_true(control_text in body_text, f"content-control body text was not recovered: {body_text}")
    assert_true(meta.get("recovered_textbox_paragraphs") == 1, f"textbox recovery count missing: {meta}")
    assert_true(meta.get("recovered_content_control_paragraphs") == 1, f"content-control recovery count missing: {meta}")
    assert_true(textbox_issue.get("severity") == "warning", f"recoverable textbox should require review, not block: {issues}")


@case
def content_parser_uses_final_view_for_revision_wrapped_textbox_text() -> None:
    work = new_workdir("parser_textbox_revision_final_view")
    docx = work / "textbox_revision_final_view.docx"
    doc = Document()
    doc.add_paragraph("1 Textbox revision")
    doc.save(docx)

    visible_text = "Visible textbox moved text."
    deleted_text = "Deleted textbox moved text."

    def inject_revised_textbox(xml: str) -> str:
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        block = (
            '<w:p><w:r><w:pict><v:shape xmlns:v="urn:schemas-microsoft-com:vml">'
            "<v:textbox><w:txbxContent><w:p>"
            f'<w:moveTo w:id="21" w:author="Synthetic"><w:r><w:t>{visible_text}</w:t></w:r></w:moveTo>'
            f'<w:moveFrom w:id="22" w:author="Synthetic"><w:r><w:t>{deleted_text}</w:t></w:r></w:moveFrom>'
            "</w:p></w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>"
        )
        assert_true(f'xmlns:w="{w_ns}"' in xml or "xmlns:w=" in xml, "document namespace should be present")
        return xml.replace("</w:body>", block + "</w:body>")

    _rewrite_docx_part(docx, "word/document.xml", inject_revised_textbox)

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    body_text = _content_plain_text(content)
    meta = content.get("_meta") or {}
    assert_true(visible_text in body_text, f"moveTo textbox text was not recovered: {content}")
    assert_true(deleted_text not in body_text, f"moveFrom textbox text leaked into final-view content: {content}")
    assert_true(meta.get("recovered_textbox_paragraphs") == 1, f"textbox recovery count changed unexpectedly: {meta}")

    result = run_generated_case("textbox_revision_final_view_render", content, base_format())
    xml = result["xml"]
    assert_true(visible_text in xml, "moveTo textbox text did not render in generated DOCX")
    assert_true(deleted_text not in xml, "moveFrom textbox text rendered in generated DOCX")


@case
def content_parser_keeps_body_content_control_that_partially_overlaps_table_text() -> None:
    work = new_workdir("parser_body_content_control_table_overlap")
    docx = work / "body_content_control_table_overlap.docx"
    doc = Document()
    doc.add_paragraph("1 Approval")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("Approved by committee")
    doc.add_paragraph("Plain paragraph after the table.")
    doc.save(docx)

    control_text = "Approved"

    def inject_body_content_control(xml: str) -> str:
        needle = "<w:t>Plain paragraph after the table.</w:t></w:r></w:p>"
        replacement = (
            '<w:sdt><w:sdtPr><w:tag w:val="body-control"/></w:sdtPr><w:sdtContent>'
            f"<w:p><w:r><w:t>{control_text}</w:t></w:r></w:p>"
            "</w:sdtContent></w:sdt>"
            + needle
        )
        assert_true(needle in xml, "body content-control injection point not found")
        return xml.replace(needle, replacement, 1)

    _rewrite_docx_part(docx, "word/document.xml", inject_body_content_control)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    loose_body_text = [item for item in paragraphs if isinstance(item, str)]
    meta = content.get("_meta") or {}

    assert_true(
        any(item == control_text for item in loose_body_text),
        f"body content-control text was mistaken for a table duplicate: {paragraphs}",
    )
    assert_true(
        meta.get("recovered_content_control_paragraphs") == 1,
        f"body content-control recovery count missing after table overlap: {meta}",
    )


@case
def content_parser_keeps_body_content_control_that_equals_table_cell_text() -> None:
    work = new_workdir("parser_body_content_control_table_exact")
    docx = work / "body_content_control_table_exact.docx"
    doc = Document()
    doc.add_paragraph("1 Approval")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("Approved")
    doc.add_paragraph("Plain paragraph after the table.")
    doc.save(docx)

    control_text = "Approved"

    def inject_body_content_control(xml: str) -> str:
        needle = "<w:t>Plain paragraph after the table.</w:t></w:r></w:p>"
        replacement = (
            '<w:sdt><w:sdtPr><w:tag w:val="body-control-exact"/></w:sdtPr><w:sdtContent>'
            f"<w:p><w:r><w:t>{control_text}</w:t></w:r></w:p>"
            "</w:sdtContent></w:sdt>"
            + needle
        )
        assert_true(needle in xml, "exact-match body content-control injection point not found")
        return xml.replace(needle, replacement, 1)

    _rewrite_docx_part(docx, "word/document.xml", inject_body_content_control)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    loose_body_text = [item for item in paragraphs if isinstance(item, str)]
    meta = content.get("_meta") or {}

    assert_true(
        any(item == control_text for item in loose_body_text),
        f"body content-control text equal to a table cell was dropped: {paragraphs}",
    )
    assert_true(
        meta.get("recovered_content_control_paragraphs") == 1,
        f"exact-match body content-control recovery count missing: {meta}",
    )


@case
def content_parser_preserves_table_cell_block_content_control_text() -> None:
    work = new_workdir("parser_table_cell_content_control")
    docx = work / "table_cell_content_control.docx"
    doc = Document()
    doc.add_paragraph("1 Controlled cell")
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].add_run("Before")
    cell.add_paragraph("After")
    doc.save(docx)

    control_text = "Controlled cell value must stay inside the table cell."

    def inject_cell_content_control(xml: str) -> str:
        needle = "<w:t>Before</w:t></w:r></w:p><w:p><w:r><w:t>After</w:t>"
        replacement = (
            "<w:t>Before</w:t></w:r></w:p>"
            '<w:sdt><w:sdtPr><w:tag w:val="cell-control"/></w:sdtPr><w:sdtContent>'
            f"<w:p><w:r><w:t>{control_text}</w:t></w:r></w:p>"
            "</w:sdtContent></w:sdt>"
            "<w:p><w:r><w:t>After</w:t>"
        )
        assert_true(needle in xml, "table-cell content-control injection point not found")
        return xml.replace(needle, replacement, 1)

    _rewrite_docx_part(docx, "word/document.xml", inject_cell_content_control)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with block content control was not preserved: {paragraphs}")
    table_item = table_items[0]
    expected_rows = [["Before\nControlled cell value must stay inside the table cell.\nAfter"]]
    assert_true(table_item.get("table_rows") == expected_rows, f"content-control text left the table cell: {table_item}")
    loose_body_text = [
        item
        for sec in content.get("sections") or []
        for item in sec.get("paragraphs") or []
        if isinstance(item, str)
    ]
    assert_true(
        not any(control_text in item for item in loose_body_text),
        f"content-control text was duplicated outside the source table cell: {paragraphs}",
    )

    result = run_generated_case("table_cell_content_control_render", content, base_format())
    xml = result["xml"]
    assert_true(xml.find("Before") < xml.find(control_text) < xml.find("After"), "content-control text rendered outside its source cell order")
    assert_true(result["report"]["passed"] is True, f"table-cell content-control render should pass QA: {result['report']}")


@case
def content_parser_preserves_table_cell_inline_content_control_text() -> None:
    work = new_workdir("parser_table_cell_inline_content_control")
    docx = work / "table_cell_inline_content_control.docx"
    doc = Document()
    doc.add_paragraph("1 Inline controlled cell")
    table = doc.add_table(rows=1, cols=1)
    para = table.cell(0, 0).paragraphs[0]
    para.add_run("Left")
    para.add_run("Right")
    doc.save(docx)

    control_text = "InlineCellValue"

    def inject_inline_content_control(xml: str) -> str:
        needle = "<w:t>Left</w:t></w:r><w:r><w:t>Right</w:t>"
        replacement = (
            "<w:t>Left</w:t></w:r>"
            '<w:sdt><w:sdtPr><w:tag w:val="inline-cell-control"/></w:sdtPr><w:sdtContent>'
            f"<w:r><w:t>{control_text}</w:t></w:r>"
            "</w:sdtContent></w:sdt>"
            "<w:r><w:t>Right</w:t>"
        )
        assert_true(needle in xml, "inline table-cell content-control injection point not found")
        return xml.replace(needle, replacement, 1)

    _rewrite_docx_part(docx, "word/document.xml", inject_inline_content_control)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with inline content control was not preserved: {paragraphs}")
    table_item = table_items[0]
    assert_true(
        table_item.get("table_rows") == [["LeftInlineCellValueRight"]],
        f"inline content-control text left the source table cell: {table_item}",
    )
    loose_body_text = [
        item
        for sec in content.get("sections") or []
        for item in sec.get("paragraphs") or []
        if isinstance(item, str)
    ]
    assert_true(
        not any(control_text in item for item in loose_body_text),
        f"inline content-control text was duplicated outside the source table cell: {paragraphs}",
    )

    result = run_generated_case("table_cell_inline_content_control_render", content, base_format())
    xml = result["xml"]
    assert_true(xml.find("Left") < xml.find(control_text) < xml.find("Right"), "inline content-control text rendered outside source order")
    assert_true(result["report"]["passed"] is True, f"inline table-cell content-control render should pass QA: {result['report']}")


@case
def content_parser_preserves_table_cell_nested_inline_content_control_text() -> None:
    work = new_workdir("parser_table_cell_nested_inline_content_control")
    docx = work / "table_cell_nested_inline_content_control.docx"
    doc = Document()
    doc.add_paragraph("1 Nested inline controlled cell")
    table = doc.add_table(rows=1, cols=1)
    para = table.cell(0, 0).paragraphs[0]
    para.add_run("OuterLeft")
    para.add_run("OuterRight")
    doc.save(docx)

    control_text = "NestedInlineCellValue"

    def inject_nested_inline_content_control(xml: str) -> str:
        needle = "<w:t>OuterLeft</w:t></w:r><w:r><w:t>OuterRight</w:t>"
        replacement = (
            "<w:t>OuterLeft</w:t></w:r>"
            '<w:sdt><w:sdtPr><w:tag w:val="outer-inline-cell-control"/></w:sdtPr><w:sdtContent>'
            '<w:sdt><w:sdtPr><w:tag w:val="inner-inline-cell-control"/></w:sdtPr><w:sdtContent>'
            f"<w:r><w:t>{control_text}</w:t></w:r>"
            "</w:sdtContent></w:sdt>"
            "</w:sdtContent></w:sdt>"
            "<w:r><w:t>OuterRight</w:t>"
        )
        assert_true(needle in xml, "nested inline table-cell content-control injection point not found")
        return xml.replace(needle, replacement, 1)

    _rewrite_docx_part(docx, "word/document.xml", inject_nested_inline_content_control)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with nested inline content control was not preserved: {paragraphs}")
    table_item = table_items[0]
    assert_true(
        table_item.get("table_rows") == [["OuterLeftNestedInlineCellValueOuterRight"]],
        f"nested inline content-control text left the source table cell: {table_item}",
    )
    loose_body_text = [
        item
        for sec in content.get("sections") or []
        for item in sec.get("paragraphs") or []
        if isinstance(item, str)
    ]
    assert_true(
        not any(control_text in item for item in loose_body_text),
        f"nested inline content-control text was duplicated outside the source table cell: {paragraphs}",
    )

    result = run_generated_case("table_cell_nested_inline_content_control_render", content, base_format())
    xml = result["xml"]
    assert_true(
        xml.find("OuterLeft") < xml.find(control_text) < xml.find("OuterRight"),
        "nested inline content-control text rendered outside source order",
    )
    assert_true(result["report"]["passed"] is True, f"nested inline table-cell content-control render should pass QA: {result['report']}")


@case
def content_parser_preserves_table_cell_content_control_simple_field_value() -> None:
    work = new_workdir("parser_table_cell_control_simple_field")
    docx = work / "table_cell_control_simple_field.docx"
    doc = Document()
    doc.add_paragraph("1 Field controlled cell")
    table = doc.add_table(rows=1, cols=1)
    para = table.cell(0, 0).paragraphs[0]
    para.add_run("Left")
    para.add_run("Right")
    doc.save(docx)

    field_text = "DisplayFieldValue"

    def inject_simple_field_content_control(xml: str) -> str:
        needle = "<w:t>Left</w:t></w:r><w:r><w:t>Right</w:t>"
        replacement = (
            "<w:t>Left</w:t></w:r>"
            '<w:sdt><w:sdtPr><w:tag w:val="field-cell-control"/></w:sdtPr><w:sdtContent>'
            '<w:fldSimple w:instr=" DOCPROPERTY SyntheticField ">'
            f"<w:r><w:t>{field_text}</w:t></w:r>"
            "</w:fldSimple>"
            "</w:sdtContent></w:sdt>"
            "<w:r><w:t>Right</w:t>"
        )
        assert_true(needle in xml, "simple-field table-cell content-control injection point not found")
        return xml.replace(needle, replacement, 1)

    _rewrite_docx_part(docx, "word/document.xml", inject_simple_field_content_control)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with simple-field content control was not preserved: {paragraphs}")
    table_item = table_items[0]
    assert_true(
        table_item.get("table_rows") == [["LeftDisplayFieldValueRight"]],
        f"simple-field visible result left the source table cell: {table_item}",
    )
    loose_body_text = [
        item
        for sec in content.get("sections") or []
        for item in sec.get("paragraphs") or []
        if isinstance(item, str)
    ]
    assert_true(
        not any(field_text in item for item in loose_body_text),
        f"simple-field visible result was duplicated outside the source table cell: {paragraphs}",
    )

    result = run_generated_case("table_cell_control_simple_field_render", content, base_format())
    xml = result["xml"]
    assert_true(xml.find("Left") < xml.find(field_text) < xml.find("Right"), "simple-field visible result rendered outside source order")
    assert_true(result["report"]["passed"] is True, f"simple-field content-control render should pass QA: {result['report']}")


@case
def content_parser_preserves_table_cell_content_control_custom_xml_value() -> None:
    work = new_workdir("parser_table_cell_control_custom_xml")
    docx = work / "table_cell_control_custom_xml.docx"
    doc = Document()
    doc.add_paragraph("1 Bound controlled cell")
    table = doc.add_table(rows=1, cols=1)
    para = table.cell(0, 0).paragraphs[0]
    para.add_run("Left")
    para.add_run("Right")
    doc.save(docx)

    bound_text = "BoundDropdownValue"

    def inject_custom_xml_content_control(xml: str) -> str:
        needle = "<w:t>Left</w:t></w:r><w:r><w:t>Right</w:t>"
        replacement = (
            "<w:t>Left</w:t></w:r>"
            '<w:sdt><w:sdtPr><w:tag w:val="bound-cell-control"/></w:sdtPr><w:sdtContent>'
            '<w:customXml w:element="boundValue" w:uri="urn:synthetic">'
            f"<w:r><w:t>{bound_text}</w:t></w:r>"
            "</w:customXml>"
            "</w:sdtContent></w:sdt>"
            "<w:r><w:t>Right</w:t>"
        )
        assert_true(needle in xml, "custom-xml table-cell content-control injection point not found")
        return xml.replace(needle, replacement, 1)

    _rewrite_docx_part(docx, "word/document.xml", inject_custom_xml_content_control)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with custom-xml content control was not preserved: {paragraphs}")
    table_item = table_items[0]
    assert_true(
        table_item.get("table_rows") == [["LeftBoundDropdownValueRight"]],
        f"custom-xml visible value left the source table cell: {table_item}",
    )
    loose_body_text = [
        item
        for sec in content.get("sections") or []
        for item in sec.get("paragraphs") or []
        if isinstance(item, str)
    ]
    assert_true(
        not any(bound_text in item for item in loose_body_text),
        f"custom-xml visible value was duplicated outside the source table cell: {paragraphs}",
    )

    result = run_generated_case("table_cell_control_custom_xml_render", content, base_format())
    xml = result["xml"]
    assert_true(xml.find("Left") < xml.find(bound_text) < xml.find("Right"), "custom-xml visible value rendered outside source order")
    assert_true(result["report"]["passed"] is True, f"custom-xml content-control render should pass QA: {result['report']}")


@case
def content_parser_preserves_table_cell_content_control_hyperlink_mixed_media_formula_note() -> None:
    work = new_workdir("parser_table_cell_control_hyperlink_mixed")
    img = work / "control_hyperlink_image.png"
    write_sample_png(img, width=128, height=96)
    docx = work / "table_cell_control_hyperlink_mixed.docx"
    doc = Document()
    doc.add_paragraph("1 Controlled hyperlink cell")
    table = doc.add_table(rows=1, cols=1)
    para = table.cell(0, 0).paragraphs[0]
    para.add_run("Lead")
    para.add_run().add_picture(str(img))
    para.add_run(r"Formula line $x=1$ and ")
    para._element.append(etree.fromstring(latex_to_omath(r"y=2", display=False).encode("utf-8")))
    para.add_run(" noted")
    doc.save(docx)

    footnote_text = "Controlled hyperlink cell note must stay after the formula."

    def inject_reference(xml: str) -> str:
        replacements = [
            (
                '<w:t xml:space="preserve"> noted</w:t></w:r></w:p>',
                '<w:t xml:space="preserve"> noted</w:t></w:r><w:r><w:footnoteReference w:id="7"/></w:r></w:p>',
            ),
            (
                "<w:t> noted</w:t></w:r></w:p>",
                '<w:t> noted</w:t></w:r><w:r><w:footnoteReference w:id="7"/></w:r></w:p>',
            ),
        ]
        for old, new in replacements:
            if old in xml:
                return xml.replace(old, new, 1)
        return xml

    def wrap_tail_in_content_control_hyperlink(xml: str) -> str:
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        root = etree.fromstring(xml.encode("utf-8"))
        ns = {"w": w_ns}
        paragraphs = root.xpath(".//w:tc//w:p[.//w:t='Lead']", namespaces=ns)
        assert_true(paragraphs, "controlled hyperlink table-cell paragraph not found")
        paragraph = paragraphs[0]
        children = list(paragraph)
        lead_idx = next(
            (
                idx
                for idx, child in enumerate(children)
                if child.xpath(".//w:t[text()='Lead']", namespaces=ns)
            ),
            -1,
        )
        assert_true(lead_idx >= 0 and lead_idx + 1 < len(children), "controlled hyperlink tail injection point not found")
        tail = children[lead_idx + 1 :]
        for child in tail:
            paragraph.remove(child)
        sdt = etree.Element(f"{{{w_ns}}}sdt")
        sdt_pr = etree.SubElement(sdt, f"{{{w_ns}}}sdtPr")
        tag = etree.SubElement(sdt_pr, f"{{{w_ns}}}tag")
        tag.set(f"{{{w_ns}}}val", "cell-control-hyperlink-mixed")
        sdt_content = etree.SubElement(sdt, f"{{{w_ns}}}sdtContent")
        hyperlink = etree.SubElement(sdt_content, f"{{{w_ns}}}hyperlink")
        hyperlink.set(f"{{{w_ns}}}anchor", "synthetic-controlled-link")
        for child in tail:
            hyperlink.append(child)
        paragraph.append(sdt)
        return etree.tostring(root, encoding="unicode")

    _rewrite_docx_part(docx, "word/document.xml", inject_reference)
    _rewrite_docx_part(docx, "word/document.xml", wrap_tail_in_content_control_hyperlink)
    with zipfile.ZipFile(docx, "a") as zf:
        zf.writestr(
            "word/footnotes.xml",
            (
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
                '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
                f'<w:footnote w:id="7"><w:p><w:r><w:t>{footnote_text}</w:t></w:r></w:p></w:footnote>'
                "</w:footnotes>"
            ),
        )

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with content-control hyperlink mixed cell was not preserved: {paragraphs}")
    table_item = table_items[0]
    assert_true(
        table_item.get("table_rows") == [["Lead\nFormula line $x=1$ and y=2 noted"]],
        f"content-control hyperlink mixed cell text/order changed: {table_item}",
    )
    cell_items = table_item.get("table_cell_items") or []
    image_item = next(
        (
            item
            for entry in cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "image"
        ),
        None,
    )
    rich_item = next(
        (
            item
            for entry in cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "rich_text"
        ),
        None,
    )
    assert_true(image_item and image_item.get("after_paragraph_index") == 1, f"content-control hyperlink image position changed: {cell_items}")
    assert_true(rich_item and rich_item.get("replace_paragraph_index") == 1, f"content-control hyperlink rich text missing: {table_item}")
    assert_true(
        [run.get("type") for run in rich_item.get("runs") or []] == ["text", "math", "text", "math", "text", "note_ref"],
        f"content-control hyperlink formula/note run order changed: {rich_item}",
    )

    result = run_generated_case("table_cell_control_hyperlink_mixed_render", content, base_format())
    xml = result["xml"]
    assert_true("$x=1$" not in xml, "content-control hyperlink LaTeX leaked delimiters into generated XML")
    assert_true(omath_count(xml) >= 2, "content-control hyperlink formulas did not both render as native math")
    assert_true(xml.count("<w:drawing>") == 1, f"content-control hyperlink image did not render exactly once: {result['manifest']}")
    assert_true("<w:footnoteReference" in xml, "content-control hyperlink footnote did not render as native reference")
    assert_true(
        xml.find("Lead") < xml.find("<w:drawing>") < xml.find("Formula line") < xml.find("<w:footnoteReference"),
        "content-control hyperlink image/formula/footnote source order changed",
    )
    assert_true(result["manifest"]["counts"].get("inline_formulas_rendered", 0) >= 2, f"inline formula count missing: {result['manifest']}")
    assert_true(result["manifest"]["counts"].get("footnote_references_rendered") == 1, f"footnote count missing: {result['manifest']}")
    assert_true(result["report"]["passed"] is True, f"content-control hyperlink mixed table-cell render should pass QA: {result['report']}")


@case
def content_parser_extracts_footnote_references_and_text() -> None:
    work = new_workdir("parser_footnotes")
    docx = work / "footnote_source.docx"
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("The method uses a calibrated dataset.")
    doc.save(docx)

    footnote_text = "Synthetic footnote text that must stay attached to the paragraph."

    def inject_reference(xml: str) -> str:
        return xml.replace(
            "<w:t>The method uses a calibrated dataset.</w:t></w:r></w:p>",
            '<w:t>The method uses a calibrated dataset.</w:t></w:r><w:r><w:footnoteReference w:id="2"/></w:r></w:p>',
        )

    _rewrite_docx_part(docx, "word/document.xml", inject_reference)
    with zipfile.ZipFile(docx, "a") as zf:
        zf.writestr(
            "word/footnotes.xml",
            (
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
                '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
                f'<w:footnote w:id="2"><w:p><w:r><w:t>{footnote_text}</w:t></w:r></w:p></w:footnote>'
                "</w:footnotes>"
            ),
        )

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    rich = next((p for p in paragraphs if isinstance(p, dict) and p.get("role") == "rich_text"), {})
    note_runs = [run for run in rich.get("runs") or [] if run.get("type") == "note_ref"]
    meta = content.get("_meta") or {}
    issues = (meta.get("source_audit") or {}).get("issues") or []
    footnote_issue = next((issue for issue in issues if issue.get("code") == "SOURCE_FOOTNOTE_UNSUPPORTED"), {})

    assert_true(note_runs, f"footnote reference did not stay in rich text runs: {paragraphs}")
    assert_true(note_runs[0].get("text") == footnote_text, f"footnote text not attached to reference: {note_runs}")
    assert_true(meta.get("footnote_references_extracted") == 1, f"footnote reference count missing: {meta}")
    assert_true(meta.get("footnote_definitions_extracted") == 1, f"footnote definition count missing: {meta}")
    assert_true(footnote_issue.get("severity") == "warning", f"extractable footnotes should require review, not block: {issues}")


@case
def content_parser_preserves_table_cell_footnote_reference() -> None:
    work = new_workdir("parser_table_cell_footnote")
    docx = work / "table_cell_footnote.docx"
    doc = Document()
    doc.add_paragraph("1 Table footnote")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("Metric value")
    doc.save(docx)

    footnote_text = "Table-cell footnote text must stay attached to the cell reference."

    def inject_reference(xml: str) -> str:
        return xml.replace(
            "<w:t>Metric value</w:t></w:r></w:p>",
            '<w:t>Metric value</w:t></w:r><w:r><w:footnoteReference w:id="3"/></w:r></w:p>',
        )

    _rewrite_docx_part(docx, "word/document.xml", inject_reference)
    with zipfile.ZipFile(docx, "a") as zf:
        zf.writestr(
            "word/footnotes.xml",
            (
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
                '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
                f'<w:footnote w:id="3"><w:p><w:r><w:t>{footnote_text}</w:t></w:r></w:p></w:footnote>'
                "</w:footnotes>"
            ),
        )

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with footnote cell was not preserved: {paragraphs}")
    table_item = table_items[0]
    cell_items = table_item.get("table_cell_items") or []
    rich_item = next(
        (
            item
            for entry in cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "rich_text"
        ),
        None,
    )
    note_runs = [run for run in (rich_item or {}).get("runs") or [] if run.get("type") == "note_ref"]
    assert_true(note_runs, f"table-cell footnote reference did not stay in rich text runs: {table_item}")
    assert_true(note_runs[0].get("text") == footnote_text, f"table-cell footnote text not attached: {note_runs}")

    result = run_generated_case("table_cell_footnote_render", content, base_format())
    assert_true("<w:footnoteReference" in result["xml"], "table-cell footnote did not render as native reference")
    assert_true(result["manifest"]["counts"].get("footnote_references_rendered") == 1, f"footnote count missing: {result['manifest']}")
    assert_true(result["report"]["passed"] is True, f"table-cell footnote render should pass QA: {result['report']}")


@case
def content_parser_preserves_table_cell_image_before_note_only_reference() -> None:
    work = new_workdir("parser_table_cell_image_note_only")
    img = work / "cell_image_note_only.png"
    write_sample_png(img, width=128, height=96)
    docx = work / "table_cell_image_note_only.docx"
    doc = Document()
    doc.add_paragraph("1 Table note-only reference")
    table = doc.add_table(rows=1, cols=1)
    para = table.cell(0, 0).paragraphs[0]
    para.add_run("Lead")
    para.add_run().add_picture(str(img))
    doc.save(docx)

    footnote_text = "Note-only anchor after an image must still render."

    def inject_reference(xml: str) -> str:
        return xml.replace(
            "</w:drawing></w:r></w:p>",
            '</w:drawing></w:r><w:r><w:footnoteReference w:id="6"/></w:r></w:p>',
            1,
        )

    _rewrite_docx_part(docx, "word/document.xml", inject_reference)
    with zipfile.ZipFile(docx, "a") as zf:
        zf.writestr(
            "word/footnotes.xml",
            (
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
                '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
                f'<w:footnote w:id="6"><w:p><w:r><w:t>{footnote_text}</w:t></w:r></w:p></w:footnote>'
                "</w:footnotes>"
            ),
        )

    content = extract_docx_content(str(docx), output_dir=str(work / "out"))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with image/note-only cell was not preserved: {paragraphs}")
    table_item = table_items[0]
    assert_true(table_item.get("table_rows") == [["Lead"]], f"note-only anchor should not add visible cell text: {table_item}")
    cell_items = table_item.get("table_cell_items") or []
    image_item = next(
        (
            item
            for entry in cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "image"
        ),
        None,
    )
    rich_item = next(
        (
            item
            for entry in cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "rich_text"
        ),
        None,
    )
    assert_true(image_item and image_item.get("after_paragraph_index") == 1, f"image position changed: {cell_items}")
    assert_true(rich_item and rich_item.get("replace_paragraph_index") == 1, f"note-only rich text position changed: {table_item}")
    assert_true([run.get("type") for run in rich_item.get("runs") or []] == ["note_ref"], f"note-only run changed: {rich_item}")
    assert_true((rich_item.get("notes") or [{}])[0].get("text") == footnote_text, f"note text not attached: {rich_item}")

    result = run_generated_case("table_cell_image_note_only_render", content, base_format())
    xml = result["xml"]
    assert_true(xml.count("<w:drawing>") == 1, f"image did not render exactly once: {result['manifest']}")
    assert_true("<w:footnoteReference" in xml, "note-only table-cell footnote did not render as native reference")
    assert_true(xml.find("Lead") < xml.find("<w:drawing>") < xml.find("<w:footnoteReference"), "image/note-only source order changed")
    assert_true(result["manifest"]["counts"].get("footnote_references_rendered") == 1, f"footnote count missing: {result['manifest']}")
    assert_true(result["report"]["passed"] is True, f"image/note-only table-cell render should pass QA: {result['report']}")


@case
def content_parser_detects_latex_delimited_formula_paragraphs() -> None:
    work = new_workdir("parser_latex_delimited")
    docx = work / "latex_delimited.docx"
    doc = Document()
    doc.add_paragraph("1 Formula Cases")
    doc.add_paragraph(r"$$L=\lim_{n\to\infty}\frac{1}{n}\sum_{i=1}^{n}x_i$$")
    doc.add_paragraph(r"$$M=\begin{matrix}a&b\\c&d\end{matrix}$$")
    doc.add_paragraph(r"$$I=\int_0^T f(t)\,dt$$")
    doc.add_paragraph(r"$$E_{total}=\sum_{t=1}^{24}P(t)\Delta t$$ (1.1)")
    doc.add_paragraph(r"$$R_{green}=\frac{E_{renew}}{E_{total}}\times100\%$$" + " \uff08A.1\uff09")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    formulas = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "formula"]
    assert_true(len(formulas) == 5, f"LaTeX-delimited paragraphs were not all formulas: {paragraphs}")
    assert_true(all(f.get("source") == "latex" and f.get("latex") for f in formulas), "LaTeX delimiters were not stripped into latex fields")
    assert_true(all("$$" not in (f.get("latex") or "") for f in formulas), "LaTeX delimiters leaked into formula latex")
    assert_true(any(f.get("text", "").endswith("\uff08A.1\uff09") and f.get("latex", "").startswith("R_{green}") for f in formulas), "appendix formula label was not handled")


@case
def formula_semantics_classifies_quantities_and_contamination() -> None:
    quantity = "全年发电量为 172.04 MWh"
    equation = "P_total(t)=P_AK+P_PM+P_NH3+P_L(t)=20.75+6*p_L(t)"
    contaminated = "当PRE(t)-PL(t)<0.1*41.5=4.15MW时,该时段无法开机;产量约束为="

    quantity_result = classify_formula_text(quantity)
    equation_result = classify_formula_text(equation)
    contaminated_result = classify_formula_text(contaminated)

    assert_true(quantity_result.category == CATEGORY_QUANTITY_TEXT, f"quantity became {quantity_result}")
    assert_true(not semantic_looks_like_formula_text(quantity), "quantity/unit text was mistaken for a formula")
    assert_true(equation_result.category == CATEGORY_DISPLAY_MATH, f"equation became {equation_result}")
    assert_true(semantic_looks_like_formula_text(equation), "standalone equation was not recognized")
    assert_true(contaminated_result.category == CATEGORY_CONTAMINATED, f"contaminated formula became {contaminated_result}")
    assert_true(not semantic_looks_like_formula_text(contaminated), "contaminated narrative was accepted as a formula")


@case
def content_parser_marks_formula_semantic_problems() -> None:
    work = new_workdir("parser_formula_semantics")
    docx = work / "formula_semantics.docx"
    quantity = "全年发电量为 172.04 MWh"
    equation = "P_total(t)=P_AK+P_PM+P_NH3+P_L(t)=20.75+6*p_L(t)"
    contaminated = "当PRE(t)-PL(t)<0.1*41.5=4.15MW时,该时段无法开机;产量约束为="
    doc = Document()
    doc.add_paragraph("1 Formula Semantics")
    doc.add_paragraph(quantity)
    doc.add_paragraph(equation)
    doc.add_paragraph(contaminated)
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    formulas = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "formula"]
    problems = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "formula_problem"]

    assert_true(quantity in paragraphs, "quantity text should remain ordinary text")
    assert_true(any(p.get("text") == equation for p in formulas), "standalone equation was not extracted as formula")
    assert_true(any(p.get("text") == contaminated for p in problems), "contaminated formula text was not marked")


@case
def formula_semantics_splits_inline_math_without_units() -> None:
    text = "模型中 P_total(t)=20.75+6*p_L(t)，全年电量为 172.04 MWh。"
    spans = split_inline_math_spans(text)
    assert_true(len(spans) == 1, f"expected one inline formula span, got {spans}")
    assert_true(spans[0]["text"] == "P_total(t)=20.75+6*p_L(t)", f"wrong inline span: {spans}")
    assert_true("172.04" not in spans[0]["text"], "quantity text was absorbed into inline formula")
    assert_true(not split_inline_math_spans("$100$ cost"), "plain dollar amount was mistaken for inline math")
    assert_true(not split_inline_math_spans("$$x=1$$"), "display math delimiter was mistaken for inline math")
    dollar_spans = split_inline_math_spans("变量 $x_i$ 表示第 i 个样本。")
    assert_true(dollar_spans and dollar_spans[0]["text"] == "x_i", f"valid dollar inline math was lost: {dollar_spans}")
    mixed = "由于约束仅为 u(t) = n 且 u(t) ∈ {0, 1}，这是标准选择问题。"
    mixed_spans = split_inline_math_spans(mixed)
    assert_true(not any("且" in str(s.get("text")) or "{" in str(s.get("text")) for s in mixed_spans), f"unsafe inline span accepted: {mixed_spans}")


@case
def prose_with_inline_formula_is_not_formula_problem() -> None:
    text = "假设三：每小时为一个调度时段。题目明确园区运行分析时段为 1 小时，即 ∆t = 1。"
    spans = split_inline_math_spans(text)
    assert_true(any(s.get("text") == "t = 1" for s in spans), f"inline formula was not found: {spans}")
    assert_true(not is_formula_problem_text(text), "ordinary prose with inline math should not block as formula_problem")


@case
def content_parser_builds_rich_text_for_plain_inline_formula() -> None:
    work = new_workdir("parser_plain_inline_formula")
    docx = work / "plain_inline_formula.docx"
    paragraph = "模型中 P_total(t)=20.75+6*p_L(t)，全年电量为 172.04 MWh。"
    doc = Document()
    doc.add_paragraph("1 Inline Formula")
    doc.add_paragraph(paragraph)
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    rich_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "rich_text"]
    assert_true(rich_items, f"inline formula paragraph was not converted to rich_text: {paragraphs}")
    runs = rich_items[0].get("runs") or []
    assert_true([r.get("type") for r in runs] == ["text", "math", "text"], f"inline formula runs are wrong: {runs}")
    result = run_generated_case("plain_inline_formula_render", content)
    assert_true(result["manifest"]["counts"]["inline_formulas_rendered"] == 1, "plain inline formula did not render as native inline math")
    assert_true(omath_count(result["xml"]) >= 1 and omath_para_count(result["xml"]) == 0, "plain inline formula rendered with wrong OMML shape")


@case
def content_parser_keeps_prose_inline_formula_as_rich_text() -> None:
    work = new_workdir("parser_prose_inline_formula")
    docx = work / "prose_inline_formula.docx"
    doc = Document()
    doc.add_paragraph("1 Formula Prose")
    doc.add_paragraph("假设三：每小时为一个调度时段。题目明确园区运行分析时段为 1 小时，即 ∆t = 1。")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    assert_true(not any(isinstance(p, dict) and p.get("role") == "formula_problem" for p in paragraphs), f"prose inline formula became formula_problem: {paragraphs}")
    rich = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "rich_text"]
    assert_true(rich and any(run.get("type") == "math" for run in rich[0].get("runs") or []), f"prose inline formula was not rich_text math: {paragraphs}")


@case
def numeric_formula_denominator_is_not_heading() -> None:
    work = new_workdir("parser_numeric_fragment_heading")
    docx = work / "numeric_fragment_heading.docx"
    doc = Document()
    doc.add_paragraph("论文题目: Numeric Fragment")
    doc.add_paragraph("8 离网运行分析")
    p = doc.add_paragraph("41 .5")
    run = p.runs[0]
    run.font.size = Pt(16)
    run.bold = True
    doc.add_paragraph("This text should stay under the previous heading.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    assert_true("41 .5" not in headings, f"numeric formula denominator became heading: {headings}")
    assert_true("41 .5" in body_text and "previous heading" in body_text, f"numeric fragment/body text was lost: {body_text}")


@case
def content_parser_repairs_split_ratio_formula_layout() -> None:
    work = new_workdir("parser_split_ratio_layout")
    docx = work / "split_ratio_layout.docx"
    doc = Document()
    doc.add_paragraph("1 Split Formula Layout")
    doc.add_paragraph("日总用电量和新能源发电量分别为：")
    doc.add_paragraph("24 24")
    doc.add_paragraph("Etotal = ∑ Ptotal(t) · ∆t, ERE = ∑ PRE(t) · ∆t")
    doc.add_paragraph("t=1")
    doc.add_paragraph("三项绿电直连指标定义为：")
    doc.add_paragraph("t=1")
    doc.add_paragraph("rself")
    p1 = doc.add_paragraph()
    p1._element.append(etree.fromstring(latex_to_omath(r"=Etotal-Esell-Ebuy\times100\%(6)(3)", display=True).encode("utf-8")))
    doc.add_paragraph("E")
    doc.add_paragraph("rgreen")
    doc.add_paragraph("RE")
    p2 = doc.add_paragraph()
    p2._element.append(etree.fromstring(latex_to_omath(r"=ERE-Esell\times100\%(7)(4)", display=True).encode("utf-8")))
    doc.add_paragraph("E")
    doc.add_paragraph("rup")
    doc.add_paragraph("total")
    p3 = doc.add_paragraph()
    p3._element.append(etree.fromstring(latex_to_omath(r"=Esell\times100\%(8)(5)", display=True).encode("utf-8")))
    doc.add_paragraph("E")
    doc.add_paragraph("RE")
    doc.add_paragraph("吨氨成本模型")
    doc.add_paragraph("吨氨成本由购电成本、风光度电成本和运维成本三部分构成，扣除售电收入：")
    doc.add_paragraph("∑24")
    doc.add_paragraph("[Pbuy(t) · λbuy(t) − Psell(t) · λsell] · 1000 + CRE + COM")
    doc.add_paragraph("其中风光度电成本 CRE = ∑24")
    doc.add_paragraph("[Pw(t) · 0.15 + Ppv(t) · 0.12] · 1000（元），运维成本")
    doc.add_paragraph("COM = ∑24")
    doc.add_paragraph("u(t)·(10×0.1+10×0.15+0.75×0.002)·1000(元),日产氨量QNH3=1.5×24=36")
    doc.add_paragraph("24")
    doc.add_paragraph("u(t)=n=Qtarget/3")
    doc.add_paragraph("t=1")
    doc.add_paragraph("Cton =")
    doc.add_paragraph("1")
    doc.add_paragraph("Qtarget")
    doc.add_paragraph("24")
    doc.add_paragraph("t=1")
    doc.add_paragraph("∆c(t) · u(t) + Cfixed]")
    doc.add_paragraph("(13)")
    doc.add_paragraph("max")
    doc.add_paragraph("α")
    doc.add_paragraph("24")
    doc.add_paragraph("QNH3 = 3 α(t) (20)")
    doc.add_paragraph("t=1")
    doc.add_paragraph("CRE=∑_{t=1}^{24} [Pw(t) · 0.15 + Ppv(t) · 0.12] · 1000")
    doc.add_paragraph("COM=∑24 (10×0.1+10×0.15+0.75×0.002)·1000")
    doc.add_paragraph("S = ∑10")
    doc.add_paragraph("x_i")
    doc.add_paragraph("K = ∑10")
    doc.add_paragraph("42")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    repaired = [p for p in paragraphs if isinstance(p, dict) and str(p.get("source") or "").startswith("repaired_")]
    repaired_text = "\n".join(p.get("text") or "" for p in repaired)
    residue = [p for p in paragraphs if p in ("rself", "rgreen", "rup", "RE") or (isinstance(p, str) and p.strip() == "E")]
    split_sum_problems = [p for p in paragraphs if isinstance(p, dict) and p.get("problem") == "split_sum_index_unknown"]

    assert_true(any(p.get("source") == "repaired_sum_bounds" for p in repaired), f"split sum bounds were not repaired: {paragraphs}")
    assert_true(sum(1 for p in repaired if p.get("source") == "repaired_ratio_cluster") == 3, f"ratio formulas were not repaired: {repaired}")
    assert_true(any(p.get("source") == "repaired_sum_prefix" for p in repaired), f"split leading sum prefix was not repaired: {paragraphs}")
    assert_true(sum(1 for p in repaired if p.get("source") == "repaired_labeled_sum_continuation") >= 2, f"labeled sum continuations were not repaired: {repaired}")
    assert_true(any(p.get("source") == "repaired_missing_sum_symbol" for p in repaired), f"missing sum symbol layout was not repaired: {repaired}")
    assert_true(any(p.get("source") == "repaired_fraction_sum_layout" for p in repaired), f"fraction plus sum layout was not repaired: {repaired}")
    assert_true(any(p.get("source") == "repaired_max_sum_layout" for p in repaired), f"max/min sum layout was not repaired: {repaired}")
    assert_true(any(p.get("source") == "repaired_inline_sum_missing_lower" for p in repaired), f"inline sum missing lower was not repaired from context: {repaired}")
    assert_true("rself=" in repaired_text and "rgreen=" in repaired_text and "rup=" in repaired_text, f"ratio variables missing: {repaired_text}")
    assert_true(all(r"\mathrm" in (p.get("latex") or "") for p in repaired if p.get("source") == "repaired_ratio_cluster"), "multi-letter repaired variables should use roman subscripts")
    latex_blob = "\n".join(p.get("latex") or "" for p in repaired)
    assert_true(r"\Deltat" not in latex_blob and r"\lambdabuy" not in latex_blob, f"greek-letter suffixes collapsed into invalid commands: {latex_blob}")
    assert_true(r"\Delta t" in latex_blob and r"\lambda_{\mathrm{buy}}" in latex_blob, f"greek-letter suffixes were not preserved: {latex_blob}")
    assert_true(r"\sum_{i=1}^{10}" in latex_blob, f"split sum index was not inferred from subscript variable: {latex_blob}")
    assert_true(r"\sum_{t=1}^{24}" in latex_blob and r"\frac{1}{Q_{\mathrm{target}}}" in latex_blob, f"time-indexed sum/fraction layout was not rendered: {latex_blob}")
    assert_true(r"\max_{\alpha}" in latex_blob, f"max sum layout did not preserve optimization variable: {latex_blob}")
    assert_true(not any((p.get("text") or "").startswith("K=") for p in repaired), f"index-less split sum was repaired by guessing: {repaired}")
    assert_true(split_sum_problems, "index-less split sum should be flagged instead of repaired by guessing")
    assert_true("(5)" not in latex_blob, f"stale source equation label leaked into repaired latex: {latex_blob}")
    assert_true(not residue, f"split ratio fragments leaked as body text: {residue}")
    result = run_generated_case("split_ratio_layout_render", content)
    assert_true(result["manifest"]["counts"]["display_formulas_rendered"] >= 7, "repaired split formulas did not render as display math")


@case
def content_parser_extracts_table_cell_images_and_flags_header_images() -> None:
    work = new_workdir("parser_table_header_images")
    img = work / "dot.png"
    img.write_bytes(PNG_1X1)
    docx = work / "table_header_images.docx"
    doc = Document()
    doc.sections[0].header.paragraphs[0].add_run().add_picture(str(img))
    doc.sections[0].footer.paragraphs[0].add_run().add_picture(str(img))
    doc.add_paragraph("1 Images")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).paragraphs[0].add_run().add_picture(str(img))
    table.cell(0, 1).text = "image in table cell"
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with image cell was not preserved: {paragraphs}")
    cell_items = table_items[0].get("table_cell_items") or []
    cell_images = [
        item
        for entry in cell_items
        if entry.get("row") == 0 and entry.get("col") == 0
        for item in entry.get("items") or []
        if isinstance(item, dict) and item.get("role") == "image"
    ]
    assert_true(cell_images, f"table-cell image was not attached to its source cell: {table_items[0]}")
    assert_true(cell_images[0].get("location") == "table_cell", f"table-cell image origin missing: {cell_images}")
    top_level_table_images = [
        p for p in paragraphs if isinstance(p, dict) and p.get("role") == "image" and p.get("location") == "table_cell"
    ]
    assert_true(not top_level_table_images, f"table-cell image leaked out as a body image: {paragraphs}")
    assert_true(content["_meta"]["images_extracted"] == 1, "header image should not be counted as a body image")
    non_body_locations = {item.get("location") for item in content["_meta"].get("non_body_images") or []}
    assert_true(any("header" in str(x) for x in non_body_locations), f"header image was not recorded as a non-body image: {non_body_locations}")
    assert_true(any("footer" in str(x) for x in non_body_locations), f"footer image was not recorded as a non-body image: {non_body_locations}")

    result = run_generated_case("table_cell_footer_image_render", content, base_format())
    codes = [item["code"] for item in result["report"]["issues"]]
    assert_true(result["manifest"]["counts"]["content_images_rendered"] == 1, f"table-cell body image was not rendered once: {result['manifest']}")
    table_xmls = re.findall(r"<w:tbl\b.*?</w:tbl>", result["xml"], flags=re.S)
    table_drawings = sum(xml.count("<w:drawing>") for xml in table_xmls)
    total_drawings = result["xml"].count("<w:drawing>")
    assert_true(total_drawings == 1, f"expected one rendered body drawing, saw {total_drawings}")
    assert_true(table_drawings == 1, "DOCX table-cell image rendered outside the generated Word table")
    assert_true("NON_BODY_IMAGE_UNSUPPORTED" in codes, f"non-body header/footer image did not block with a clear code: {codes}")


@case
def content_parser_preserves_table_cell_image_run_order() -> None:
    work = new_workdir("parser_table_cell_image_order")
    img_before = work / "before.png"
    img_after = work / "after.png"
    write_sample_png(img_before, width=120, height=90)
    write_sample_png(img_after, width=140, height=80)
    docx = work / "table_cell_image_order.docx"
    doc = Document()
    doc.add_paragraph("1 Image order")
    table = doc.add_table(rows=1, cols=2)
    left = table.cell(0, 0).paragraphs[0]
    left.add_run().add_picture(str(img_before))
    left.add_run("Image before text")
    right = table.cell(0, 1).paragraphs[0]
    right.add_run("Text before image")
    right.add_run().add_picture(str(img_after))
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with ordered image cells was not preserved: {paragraphs}")
    table_item = table_items[0]
    assert_true(
        table_item.get("table_rows") == [["Image before text", "Text before image"]],
        f"table text changed during image order extraction: {table_item}",
    )
    cell_items = table_item.get("table_cell_items") or []
    left_image = next(
        (
            item
            for entry in cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "image"
        ),
        None,
    )
    right_image = next(
        (
            item
            for entry in cell_items
            if entry.get("row") == 0 and entry.get("col") == 1
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "image"
        ),
        None,
    )
    assert_true(left_image and left_image.get("after_paragraph_index") == 0, f"image-before-text position lost: {cell_items}")
    assert_true(right_image and right_image.get("after_paragraph_index") == 1, f"text-before-image position lost: {cell_items}")

    result = run_generated_case("table_cell_image_order_render", content, base_format())
    xml = result["xml"]
    assert_true(xml.count("<w:drawing>") == 2, f"expected two rendered cell images: {result['manifest']}")
    assert_true(
        xml.find("<w:drawing>") < xml.find("Image before text"),
        "image-before-text cell rendered text before its image",
    )
    assert_true(
        xml.find("Text before image") < xml.rfind("<w:drawing>"),
        "text-before-image cell rendered image before its text",
    )
    assert_true(result["report"]["passed"] is True, f"table-cell image order render should pass QA: {result['report']}")


@case
def content_parser_preserves_table_cell_inline_omml_formula() -> None:
    work = new_workdir("parser_table_cell_inline_omml")
    docx = work / "table_cell_inline_omml.docx"
    doc = Document()
    doc.add_paragraph("1 Table formula")
    table = doc.add_table(rows=1, cols=1)
    cell_para = table.cell(0, 0).paragraphs[0]
    cell_para.add_run("Energy ")
    cell_para._element.append(etree.fromstring(latex_to_omath(r"E=mc^2", display=False).encode("utf-8")))
    cell_para.add_run(" model")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with inline OMML cell was not preserved: {paragraphs}")
    table_item = table_items[0]
    cell_items = table_item.get("table_cell_items") or []
    rich_item = next(
        (
            item
            for entry in cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "rich_text"
        ),
        None,
    )
    assert_true(rich_item, f"table-cell inline OMML was not preserved as structured rich text: {table_item}")
    assert_true(
        [run.get("type") for run in rich_item.get("runs") or []] == ["text", "math", "text"],
        f"table-cell inline OMML run order changed: {rich_item}",
    )

    result = run_generated_case("table_cell_inline_omml_render", content, base_format())
    assert_true(omath_count(result["xml"]) >= 1, "table-cell inline OMML rendered as plain text instead of native math")
    assert_true(result["manifest"]["counts"].get("inline_formulas_rendered", 0) >= 1, f"inline formula count missing: {result['manifest']}")
    assert_true(result["report"]["passed"] is True, f"table-cell inline OMML render should pass QA: {result['report']}")


@case
def content_parser_preserves_table_cell_latex_inline_formula() -> None:
    work = new_workdir("parser_table_cell_latex_inline")
    docx = work / "table_cell_latex_inline.docx"
    doc = Document()
    doc.add_paragraph("1 Table formula")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run(r"Energy $E=mc^2$ model")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with LaTeX inline cell was not preserved: {paragraphs}")
    table_item = table_items[0]
    cell_items = table_item.get("table_cell_items") or []
    rich_item = next(
        (
            item
            for entry in cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "rich_text"
        ),
        None,
    )
    assert_true(rich_item, f"table-cell LaTeX inline formula was not preserved as structured rich text: {table_item}")
    assert_true(
        [run.get("type") for run in rich_item.get("runs") or []] == ["text", "math", "text"],
        f"table-cell LaTeX inline formula run order changed: {rich_item}",
    )

    result = run_generated_case("table_cell_latex_inline_render", content, base_format())
    assert_true("$E=mc^2$" not in result["xml"], "table-cell LaTeX delimiter leaked into generated XML")
    assert_true(omath_count(result["xml"]) >= 1, "table-cell LaTeX inline formula rendered as plain text instead of native math")
    assert_true(result["manifest"]["counts"].get("inline_formulas_rendered", 0) >= 1, f"inline formula count missing: {result['manifest']}")
    assert_true(result["report"]["passed"] is True, f"table-cell LaTeX inline render should pass QA: {result['report']}")


@case
def content_parser_preserves_table_cell_mixed_image_latex_and_omml_order() -> None:
    work = new_workdir("parser_table_cell_mixed_media_formula")
    img = work / "cell_image.png"
    write_sample_png(img, width=128, height=96)
    docx = work / "table_cell_mixed_media_formula.docx"
    doc = Document()
    doc.add_paragraph("1 Mixed cell")
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].add_run("Lead paragraph")
    cell.add_paragraph().add_run().add_picture(str(img))
    formula_para = cell.add_paragraph()
    formula_para.add_run(r"Formula line $x=1$ and ")
    formula_para._element.append(etree.fromstring(latex_to_omath(r"y=2", display=False).encode("utf-8")))
    formula_para.add_run(" done")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with mixed media/formula cell was not preserved: {paragraphs}")
    table_item = table_items[0]
    assert_true(
        table_item.get("table_rows") == [["Lead paragraph\nFormula line $x=1$ and y=2 done"]],
        f"table text changed during mixed cell extraction: {table_item}",
    )
    cell_items = table_item.get("table_cell_items") or []
    image_item = next(
        (
            item
            for entry in cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "image"
        ),
        None,
    )
    rich_item = next(
        (
            item
            for entry in cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "rich_text"
        ),
        None,
    )
    assert_true(image_item and image_item.get("after_paragraph_index") == 1, f"cell image position changed: {cell_items}")
    assert_true(rich_item, f"mixed formula cell did not produce rich text: {table_item}")
    assert_true(
        [run.get("type") for run in rich_item.get("runs") or []] == ["text", "math", "text", "math", "text"],
        f"mixed LaTeX/OMML run order changed: {rich_item}",
    )

    result = run_generated_case("table_cell_mixed_media_formula_render", content, base_format())
    xml = result["xml"]
    assert_true("$x=1$" not in xml, "table-cell mixed formula leaked LaTeX delimiters into generated XML")
    assert_true(omath_count(xml) >= 2, "table-cell mixed LaTeX/OMML formulas did not both render as native math")
    assert_true(xml.count("<w:drawing>") == 1, f"cell image did not render exactly once: {result['manifest']}")
    assert_true(xml.find("<w:drawing>") < xml.find("Formula line"), "cell image moved after the following formula paragraph")
    assert_true(result["manifest"]["counts"].get("inline_formulas_rendered", 0) >= 2, f"inline formula count missing: {result['manifest']}")
    assert_true(result["report"]["passed"] is True, f"mixed table-cell media/formula render should pass QA: {result['report']}")


@case
def content_parser_preserves_table_cell_inline_image_before_formula_and_footnote() -> None:
    work = new_workdir("parser_table_cell_inline_image_formula_note")
    img = work / "cell_inline_image.png"
    write_sample_png(img, width=128, height=96)
    docx = work / "table_cell_inline_image_formula_note.docx"
    doc = Document()
    doc.add_paragraph("1 Inline cell")
    table = doc.add_table(rows=1, cols=1)
    para = table.cell(0, 0).paragraphs[0]
    para.add_run("Lead")
    para.add_run().add_picture(str(img))
    para.add_run(r"Formula line $x=1$ and ")
    para._element.append(etree.fromstring(latex_to_omath(r"y=2", display=False).encode("utf-8")))
    para.add_run(" noted")
    doc.save(docx)

    footnote_text = "Inline cell note must stay after the formula."

    def inject_reference(xml: str) -> str:
        replacements = [
            (
                '<w:t xml:space="preserve"> noted</w:t></w:r></w:p>',
                '<w:t xml:space="preserve"> noted</w:t></w:r><w:r><w:footnoteReference w:id="4"/></w:r></w:p>',
            ),
            (
                "<w:t> noted</w:t></w:r></w:p>",
                '<w:t> noted</w:t></w:r><w:r><w:footnoteReference w:id="4"/></w:r></w:p>',
            ),
        ]
        for old, new in replacements:
            if old in xml:
                return xml.replace(old, new, 1)
        return xml

    _rewrite_docx_part(docx, "word/document.xml", inject_reference)
    with zipfile.ZipFile(docx, "a") as zf:
        zf.writestr(
            "word/footnotes.xml",
            (
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
                '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
                f'<w:footnote w:id="4"><w:p><w:r><w:t>{footnote_text}</w:t></w:r></w:p></w:footnote>'
                "</w:footnotes>"
            ),
        )

    content = extract_docx_content(str(docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    table_items = [p for p in paragraphs if isinstance(p, dict) and p.get("role") == "table"]
    assert_true(table_items, f"table with inline image/formula/note cell was not preserved: {paragraphs}")
    table_item = table_items[0]
    assert_true(
        table_item.get("table_rows") == [["Lead\nFormula line $x=1$ and y=2 noted"]],
        f"inline image should split table-cell text at source image position: {table_item}",
    )
    cell_items = table_item.get("table_cell_items") or []
    image_item = next(
        (
            item
            for entry in cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "image"
        ),
        None,
    )
    rich_item = next(
        (
            item
            for entry in cell_items
            if entry.get("row") == 0 and entry.get("col") == 0
            for item in entry.get("items") or []
            if isinstance(item, dict) and item.get("role") == "rich_text"
        ),
        None,
    )
    assert_true(image_item and image_item.get("after_paragraph_index") == 1, f"inline cell image position changed: {cell_items}")
    assert_true(rich_item and rich_item.get("replace_paragraph_index") == 1, f"formula/note rich text should replace text after image: {table_item}")
    assert_true(
        [run.get("type") for run in rich_item.get("runs") or []] == ["text", "math", "text", "math", "text", "note_ref"],
        f"formula/note run order changed: {rich_item}",
    )

    result = run_generated_case("table_cell_inline_image_formula_note_render", content, base_format())
    xml = result["xml"]
    assert_true("$x=1$" not in xml, "inline table-cell formula leaked LaTeX delimiters into generated XML")
    assert_true(omath_count(xml) >= 2, "inline table-cell LaTeX/OMML formulas did not both render as native math")
    assert_true(xml.count("<w:drawing>") == 1, f"inline cell image did not render exactly once: {result['manifest']}")
    assert_true("<w:footnoteReference" in xml, "inline table-cell footnote did not render as native reference")
    assert_true(
        xml.find("Lead") < xml.find("<w:drawing>") < xml.find("Formula line") < xml.find("<w:footnoteReference"),
        "inline table-cell image/formula/footnote source order changed",
    )
    assert_true(result["manifest"]["counts"].get("inline_formulas_rendered", 0) >= 2, f"inline formula count missing: {result['manifest']}")
    assert_true(result["manifest"]["counts"].get("footnote_references_rendered") == 1, f"footnote count missing: {result['manifest']}")
    assert_true(result["report"]["passed"] is True, f"inline image/formula/note table-cell render should pass QA: {result['report']}")


@case
def content_parser_reports_unreadable_docx_image_relationships() -> None:
    work = new_workdir("parser_unreadable_docx_image")
    img = work / "dot.png"
    img.write_bytes(PNG_1X1)
    valid_docx = work / "valid_image.docx"
    bad_docx = work / "bad_image.docx"
    doc = Document()
    doc.add_paragraph("1 Images")
    doc.add_paragraph("Image before corrupted relationship.")
    doc.add_picture(str(img))
    doc.save(valid_docx)

    with zipfile.ZipFile(valid_docx, "r") as src, zipfile.ZipFile(bad_docx, "w", zipfile.ZIP_DEFLATED) as dst:
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename.startswith("word/media/") and info.filename.lower().endswith(".png"):
                data = b"not-a-readable-image"
            dst.writestr(info, data)

    content = extract_docx_content(str(bad_docx), output_dir=str(work))
    paragraphs = [p for sec in content["sections"] for p in sec.get("paragraphs", [])]
    assert_true(content["_meta"].get("images_extracted") == 0, f"unreadable DOCX image should not be counted as extracted: {content['_meta']}")
    failures = content["_meta"].get("image_extract_failures") or []
    assert_true(failures and "image" in str(failures[0].get("target") or ""), f"unreadable DOCX image relationship was not recorded: {failures}")
    assert_true(not any(isinstance(p, dict) and p.get("role") == "image" for p in paragraphs), f"unreadable DOCX image leaked into content stream: {paragraphs}")

    result = run_generated_case("unreadable_docx_image_generated", content, base_format())
    codes = [item["code"] for item in result["report"]["issues"]]
    assert_true("IMAGE_EXTRACT_FAILED" in codes, f"QA did not report unreadable DOCX image extraction failure: {codes}")
    action = str(result["report"].get("next_action") or "")
    assert_true("IMAGE_EXTRACT_FAILED" in action or "图片" in action, f"unreadable DOCX image next_action was too generic: {action}")


@case
def content_parser_keeps_references_before_english_appendix() -> None:
    work = new_workdir("parser_refs_appendix")
    docx = work / "refs_appendix.docx"
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Body paragraph before references.")
    doc.add_paragraph("References")
    doc.add_paragraph("[1] Synthetic reference one.")
    doc.add_paragraph("[2] Synthetic reference two.")
    doc.add_paragraph("Appendix A Reproducible Commands")
    doc.add_paragraph("python run_pipeline.py --mode developer")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    assert_true(len(content.get("references") or []) == 2, f"reference entries were not preserved: {content.get('references')}")
    assert_true(all("python run_pipeline" not in str(ref) for ref in content.get("references") or []), "appendix command leaked into references")
    appendix = [sec for sec in content.get("sections") or [] if sec.get("role") == "appendix"]
    assert_true(appendix and any("python run_pipeline" in str(p) for p in appendix[0].get("paragraphs") or []), "English appendix section was not preserved")


@case
def content_parser_extracts_english_title_before_abstract() -> None:
    work = new_workdir("parser_english_title")
    docx = work / "english_title.docx"
    doc = Document()
    p = doc.add_paragraph("中文论文标题示例用于测试")
    p.runs[0].font.size = Pt(16)
    doc.add_paragraph("Research on Template Driven Document Quality Assessment")
    doc.add_paragraph("Abstract: This is the abstract body.")
    doc.add_paragraph("Key words: document automation")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Body text.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    assert_true(content.get("title_info", {}).get("title_en") == "Research on Template Driven Document Quality Assessment", "English title before abstract was not extracted")
    assert_true(not any((sec.get("heading") or "").startswith("Research on") for sec in content.get("sections") or []), "English title leaked into body sections")


@case
def content_parser_extracts_docx_title_style_without_explicit_font_size() -> None:
    work = new_workdir("parser_title_style")
    docx = work / "title_style.docx"
    title = "\u9762\u5411\u4e2d\u6587\u6bd5\u4e1a\u8bba\u6587\u7684\u673a\u5668\u5b66\u4e60\u6392\u7248\u6d41\u6c34\u7ebf\u7814\u7a76"
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_heading("\u6458\u8981", 1)
    doc.add_paragraph("\u672c\u6587\u6458\u8981\u7528\u4e8e\u9a8c\u8bc1\u9898\u540d\u6837\u5f0f\u63d0\u53d6\u3002")
    doc.add_heading("1 \u7eea\u8bba", 1)
    doc.add_paragraph("Body text.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    assert_true(content.get("title_info", {}).get("title_cn") == title, f"title style was not extracted: {content.get('title_info')}")
    assert_true(title not in headings, f"title style leaked into body sections: {headings}")
    assert_true("\u6458\u8981" in headings, f"abstract heading was lost after title extraction: {headings}")


@case
def content_parser_extracts_plain_title_before_abstract() -> None:
    work = new_workdir("parser_plain_front_title")
    docx = work / "plain_front_title.docx"
    title = "\u9762\u5411\u7eff\u7535\u6d88\u7eb3\u7684\u591a\u6e90\u80fd\u6e90\u8c03\u5ea6\u4e0e\u7ecf\u6d4e\u6027\u8bc4\u4f30\u7814\u7a76"
    doc = Document()
    doc.add_paragraph(title)
    doc.add_paragraph("\u6458\u8981\uff1a\u672c\u6587\u6458\u8981\u7528\u4e8e\u9a8c\u8bc1\u65e0\u663e\u5f0f\u5b57\u53f7\u7684\u524d\u7f6e\u9898\u540d\u3002")
    doc.add_paragraph("\u5173\u952e\u8bcd\uff1a\u7eff\u7535\uff1b\u8c03\u5ea6")
    doc.add_heading("\u7b2c1\u7ae0 \u7eea\u8bba", 1)
    doc.add_paragraph("\u6b63\u6587\u6bb5\u843d\u3002")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    assert_true(content.get("title_info", {}).get("title_cn") == title, f"plain front title was not extracted: {content.get('title_info')}")
    assert_true(title not in headings, f"plain title leaked into body sections: {headings}")


@case
def content_parser_does_not_treat_chapter_heading_as_title() -> None:
    work = new_workdir("parser_chapter_not_title")
    docx = work / "chapter_not_title.docx"
    heading = "\u7b2c\u4e00\u7ae0 \u7eea\u8bba\u4e0e\u7814\u7a76\u80cc\u666f"
    doc = Document()
    doc.add_heading(heading, 1)
    doc.add_paragraph("\u8fd9\u662f\u6b63\u6587\u7b2c\u4e00\u6bb5\uff0c\u7528\u4e8e\u9a8c\u8bc1\u76f4\u63a5\u4ece\u7ae0\u8282\u5f00\u59cb\u7684\u6587\u6863\u3002")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    assert_true(content.get("title_info", {}).get("title_cn") != heading, f"chapter heading was misdetected as title: {content.get('title_info')}")
    assert_true(heading in headings, f"chapter heading was not preserved as a body section: {headings}")


@case
def content_parser_splits_chinese_enumerated_headings_after_keywords() -> None:
    work = new_workdir("parser_cn_enum")
    img = work / "dot.png"
    img.write_bytes(PNG_1X1)
    docx = work / "cn_enum.docx"
    doc = Document()
    doc.add_paragraph("摘要")
    doc.add_paragraph("本文用于测试。")
    doc.add_paragraph("关键词：")
    doc.add_paragraph("绿电直连；优化运行")
    doc.add_paragraph("一、问题重述")
    doc.add_paragraph("正文段落。")
    doc.add_paragraph().add_run().add_picture(str(img))
    doc.add_paragraph("图 1-1 示例图片")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    assert_true("一、问题重述" in headings, "Chinese enumerated heading was not split into a body section")
    body_sec = next(sec for sec in content.get("sections") or [] if sec.get("heading") == "一、问题重述")
    assert_true(body_sec.get("images"), "Image after Chinese enumerated heading was not assigned to body section")
    kw_sec = next((sec for sec in content.get("sections") or [] if sec.get("role") == "cn_keywords"), {})
    assert_true(not kw_sec.get("images"), "Image after body heading leaked into keyword front matter")



@case
def content_parser_skips_source_toc_and_records_placeholders() -> None:
    work = new_workdir("source_toc_placeholder")
    docx = work / "source_toc.docx"
    doc = Document()
    doc.add_paragraph("报名序号: [报名序号]")
    doc.add_paragraph("论文题目: Synthetic Energy Paper")
    doc.add_paragraph("目 录")
    doc.add_paragraph("一、 问题重述")
    doc.add_paragraph("172.04 MWh should not become a heading")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("一、 问题重述")
    doc.add_paragraph("This is the real body paragraph.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    assert_true(content.get("title_info", {}).get("title_cn") == "Synthetic Energy Paper", "labeled title was not extracted")
    assert_true(content.get("_meta", {}).get("source_toc_skipped_paragraphs", 0) >= 3, "source TOC block was not skipped")
    assert_true(content.get("_meta", {}).get("source_placeholders"), "source placeholders were not recorded")
    assert_true("172.04 MWh should not become a heading" not in headings, "TOC/body sentence became a heading")
    assert_true("一、 问题重述" in headings, "real body heading was lost after TOC skip")


@case
def content_parser_skips_source_toc_with_roman_page_numbers() -> None:
    work = new_workdir("source_toc_roman_pages")
    docx = work / "source_toc_roman_pages.docx"
    doc = Document()
    doc.add_paragraph("\u8bba\u6587\u9898\u76ee: Roman Page TOC")
    doc.add_paragraph("\u76ee\u5f55")
    doc.add_paragraph("\u6458\u8981........................................ I")
    doc.add_paragraph("Abstract.................................... II")
    doc.add_paragraph("1 \u7eea\u8bba..................................... 1")
    doc.add_paragraph("2 \u7cfb\u7edf\u8bbe\u8ba1................................. 5")
    doc.add_paragraph("\u53c2\u8003\u6587\u732e................................... 32")
    doc.add_paragraph("1 \u7eea\u8bba")
    doc.add_paragraph("This real body paragraph must not be collected as a reference.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    assert_true(content.get("_meta", {}).get("source_toc_skipped_paragraphs", 0) >= 5, "roman-page source TOC entries were not skipped")
    assert_true(not content.get("references"), f"source TOC reference line activated reference collection: {content.get('references')}")
    assert_true("1 \u7eea\u8bba" in headings, f"real body heading after roman TOC was lost: {headings}")
    assert_true("real body paragraph" in body_text, f"real body after roman TOC was lost: {body_text}")


@case
def content_parser_keeps_numbered_long_body_sentence_as_body() -> None:
    work = new_workdir("numbered_long_body_sentence")
    docx = work / "numbered_long_body_sentence.docx"
    long_sentence = (
        "1 绪论的第1个分析段落用于模拟完整论文正文，"
        "研究对象既包含宏观制度因素，也包含微观行为数据，"
        "用于检验文本抽取、分页和样式迁移的稳定性。"
    )
    doc = Document()
    doc.add_paragraph("论文题目: Numbered Sentence")
    doc.add_paragraph("1 绪论")
    doc.add_paragraph(long_sentence)
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    assert_true(long_sentence not in headings, f"numbered body sentence became heading: {headings}")
    assert_true(long_sentence in body_text, f"numbered body sentence was not preserved: {body_text}")


@case
def content_parser_exits_source_toc_without_page_break() -> None:
    work = new_workdir("source_toc_no_break")
    docx = work / "source_toc_no_break.docx"
    doc = Document()
    doc.add_paragraph("论文题目: No Break TOC")
    doc.add_paragraph("目录")
    doc.add_paragraph("1 绪论 1")
    doc.add_paragraph("1 正文第一章")
    doc.add_paragraph("This paragraph must survive even though no page break follows the source TOC.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    assert_true(content.get("_meta", {}).get("source_toc_skipped_paragraphs", 0) >= 2, "source TOC entries were not skipped")
    assert_true("1 正文第一章" in headings, f"real heading after source TOC was lost: {headings}")
    assert_true("must survive" in body_text, f"body text after source TOC was lost: {body_text}")


@case
def content_parser_skips_unpaged_source_toc_without_boundary() -> None:
    work = new_workdir("source_toc_unpaged_no_boundary")
    docx = work / "source_toc_unpaged_no_boundary.docx"
    doc = Document()
    doc.add_paragraph("论文题目: Unpaged TOC")
    doc.add_paragraph("目录")
    doc.add_paragraph("一、 问题重述")
    doc.add_paragraph("二、 问题分析")
    doc.add_paragraph("三、 模型假设")
    doc.add_paragraph("一、 问题重述")
    doc.add_paragraph("This is the real body paragraph after the repeated first heading.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    assert_true(content.get("_meta", {}).get("source_toc_skipped_paragraphs", 0) >= 4, "unpaged source TOC block was not skipped")
    assert_true("二、 问题分析" not in headings and "三、 模型假设" not in headings, f"unpaged TOC entries leaked as sections: {headings}")
    assert_true("一、 问题重述" in headings, f"real repeated heading was lost: {headings}")
    assert_true("real body paragraph" in body_text, f"body text after repeated heading was lost: {body_text}")


@case
def content_parser_skips_source_toc_with_duplicate_formula_fragments_before_boundary() -> None:
    work = new_workdir("source_toc_duplicate_fragments")
    docx = work / "source_toc_duplicate_fragments.docx"
    doc = Document()
    doc.add_paragraph("论文题目: Duplicate Fragment TOC")
    doc.add_paragraph("目 录")
    doc.add_paragraph("一、 问题重述")
    doc.add_paragraph("二、 问题分析")
    doc.add_paragraph("41 .5")
    doc.add_paragraph("41 .5")
    doc.add_paragraph("九、 问题五")
    doc.add_paragraph("十、 灵敏度分析")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("一、 问题重述")
    doc.add_paragraph("This is the real body paragraph.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    assert_true(content.get("_meta", {}).get("source_toc_skipped_paragraphs", 0) >= 8, "source TOC with duplicate fragments was not skipped to boundary")
    assert_true("九、 问题五" not in headings and "十、 灵敏度分析" not in headings, f"late TOC entries leaked: {headings}")
    assert_true("一、 问题重述" in headings and "real body paragraph" in body_text, f"real body after source TOC was lost: {headings} / {body_text}")


@case
def content_parser_preserves_real_directory_section_content() -> None:
    work = new_workdir("real_directory_section")
    docx = work / "real_directory_section.docx"
    doc = Document()
    doc.add_paragraph("论文题目: Directory Section")
    directory_heading = doc.add_paragraph("目录")
    directory_heading.style = doc.styles["Heading 1"]
    doc.add_paragraph("本节介绍产品目录与数据字典，不是源文档自动目录。")
    doc.add_paragraph("1 正文")
    doc.add_paragraph("This paragraph must remain in the document.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    assert_true(not content.get("_meta", {}).get("source_toc_skipped_paragraphs"), "real directory section was treated as a source TOC")
    assert_true("产品目录与数据字典" in body_text, f"real directory section content was lost: {body_text}")
    assert_true("must remain" in body_text, f"body text after real directory section was lost: {body_text}")


@case
def content_parser_preserves_real_directory_section_before_page_break() -> None:
    work = new_workdir("real_directory_section_page_break")
    docx = work / "real_directory_section_page_break.docx"
    doc = Document()
    doc.add_paragraph("论文题目: Directory Section With Page Break")
    directory_heading = doc.add_paragraph("目录")
    directory_heading.style = doc.styles["Heading 1"]
    subheading = doc.add_paragraph("一、 产品目录")
    subheading.style = doc.styles["Heading 2"]
    doc.add_paragraph("本节介绍产品目录与数据字典，不是源文档自动目录。")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("1 正文")
    doc.add_paragraph("This paragraph must remain after the page break.")
    doc.save(docx)

    content = extract_docx_content(str(docx), output_dir=str(work))
    body_text = "\n".join(
        p if isinstance(p, str) else str(p.get("text") or "")
        for sec in content.get("sections") or []
        for p in sec.get("paragraphs", [])
    )
    headings = [sec.get("heading") for sec in content.get("sections") or []]
    assert_true(not content.get("_meta", {}).get("source_toc_skipped_paragraphs"), "real directory section before a page break was treated as source TOC")
    assert_true("目录" in headings and "一、 产品目录" in headings, f"real directory headings were lost: {headings}")
    assert_true("产品目录与数据字典" in body_text and "must remain" in body_text, f"real directory/page-break content was lost: {body_text}")


@case
def qa_reports_toc_pollution_placeholders_and_formula_fragments() -> None:
    work = new_workdir("qa_semantic_guards")
    docx = work / "out.docx"
    doc = Document()
    doc.add_paragraph("报名序号: [报名序号]")
    doc.add_paragraph("172.04 MWh should not be a heading")
    doc.save(docx)
    content = base_content(["E", "rgreen", "RE"])
    content["_meta"]["source_placeholders"] = [{"paragraph": 1, "text": "报名序号: [报名序号]"}]
    content["sections"][0]["heading"] = "172.04 MWh should not be a heading"
    content["sections"][0]["paragraphs"] = [
        {"role": "formula", "text": "x=1(1)(1)", "numbered": True},
        {"role": "formula", "text": "吨/日对应的开机时段数n=Q/qrate∈24,21,18,15,12.产量约束为=", "numbered": True},
        "E",
        "rgreen",
        "RE",
    ]
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    (work / "build_generated.py").write_text("# synthetic\n", encoding="utf-8")
    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    for code in ["CONTENT_TOC_POLLUTION", "UNFILLED_PLACEHOLDER_TEXT", "FORMULA_NUMBER_CONFLICT", "FORMULA_TEXT_FRAGMENTED", "PLACEHOLDER_TEXT_LEFT"]:
        assert_true(code in codes, f"QA did not report {code}")
    formula_issue = next(item for item in report["issues"] if item["code"] == "FORMULA_TEXT_FRAGMENTED")
    assert_true("contaminated formula text" in formula_issue.get("detail", ""), "QA did not explain narrative text inside a formula")
    assert_true(formula_issue["severity"] == "warning", "fragmented formula text should be downgraded to warning")
    assert_true(report["passed"] is False, "semantic guard issues should fail QA")


