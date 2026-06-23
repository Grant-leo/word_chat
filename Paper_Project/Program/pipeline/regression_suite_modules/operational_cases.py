"""Privacy, visual QA, and CLI regression cases."""
from __future__ import annotations

import hashlib
import json
import os
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from privacy import sanitize_value
from qa_checker import check_output
from regression_suite_modules.harness import assert_true, base_content, base_format, case, new_workdir, write_json, write_sample_png

PIPELINE_DIR = Path(__file__).resolve().parents[1]
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
ET.register_namespace("w", "http://schemas.openxmlformats.org/wordprocessingml/2006/main")


def _rewrite_docx_part(docx_path: Path, part_name: str, transform) -> None:
    original = docx_path.with_suffix(docx_path.suffix + ".src")
    docx_path.replace(original)
    with zipfile.ZipFile(original, "r") as zin, zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename == part_name:
                text = data.decode("utf-8")
                data = transform(text).encode("utf-8")
            zout.writestr(info, data)
    original.unlink()


@case
def privacy_sanitizes_absolute_paths() -> None:
    data = {
        "path": r"X:\workspace\project\Outputs\private\file.docx",
        "tmp": str(Path(tempfile.gettempdir()) / "abc" / "file.pdf"),
        "embedded": r"Missing image: X:\workspace\project\Inputs\private\figure.png",
    }
    sanitized = sanitize_value(data, project_root=r"X:\workspace\project")
    text = json.dumps(sanitized, ensure_ascii=False)
    assert_true(r"X:\workspace" not in text and "project" not in text, "project path leaked")
    assert_true("<PROJECT>" in text and "<TEMP>" in text, "path labels missing")
    assert_true("Missing image: <PROJECT>/Inputs/private/figure.png" in text, "embedded absolute path was not sanitized")


@case
def private_corpus_inventory_classifies_realdata_without_content_leakage() -> None:
    from private_corpus_audit import audit_corpus

    work = new_workdir("private_corpus_inventory")
    corpus = work / "corpus"
    corpus.mkdir()
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Synthetic body text for inventory classification.")
    doc.save(corpus / "paper.docx")
    hmerge_doc = Document()
    hmerge_table = hmerge_doc.add_table(rows=1, cols=3)
    hmerge_table.cell(0, 0).text = "Legacy merged header"
    hmerge_table.cell(0, 1).text = ""
    hmerge_table.cell(0, 2).text = "Score"
    hmerge_path = corpus / "legacy_hmerge.docx"
    hmerge_doc.save(hmerge_path)

    def inject_legacy_hmerge(xml: str) -> str:
        root = ET.fromstring(xml.encode("utf-8"))
        table_el = root.find(".//" + W_NS + "tbl")
        assert_true(table_el is not None, "test table missing")
        row = table_el.find(W_NS + "tr")
        assert_true(row is not None, "test row missing")
        cells = row.findall(W_NS + "tc")
        assert_true(len(cells) == 3, "test cells missing")

        def ensure_tc_pr(cell):
            tc_pr = cell.find(W_NS + "tcPr")
            if tc_pr is None:
                tc_pr = ET.Element(W_NS + "tcPr")
                cell.insert(0, tc_pr)
            return tc_pr

        restart = ET.SubElement(ensure_tc_pr(cells[0]), W_NS + "hMerge")
        restart.set(W_NS + "val", "restart")
        ET.SubElement(ensure_tc_pr(cells[1]), W_NS + "hMerge")
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(hmerge_path, "word/document.xml", inject_legacy_hmerge)
    (corpus / "legacy.doc").write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1legacy")
    (corpus / "archive.rar").write_bytes(b"Rar!\x1a\x07\x00")

    result = audit_corpus(corpus, output_dir=work / "audit")
    categories = {item["relative_path"]: item["classification"] for item in result["items"]}
    assert_true(categories["paper.docx"] in {"content_candidate", "reference_candidate", "template_candidate"}, "valid DOCX was not classified as a document candidate")
    assert_true(categories["legacy.doc"] == "unsupported_or_conversion_needed", "legacy .doc should be isolated")
    assert_true(categories["archive.rar"] == "attachment_or_nonpaper", "archive should not enter document matrix")
    hmerge_item = next(item for item in result["items"] if item["relative_path"] == "legacy_hmerge.docx")
    assert_true(hmerge_item["features"].get("merged_cell_count") == 2, f"legacy hMerge was not counted in private inventory: {hmerge_item}")
    assert_true((work / "audit" / "inventory.json").exists(), "inventory.json was not written")
    assert_true((work / "audit" / "inventory.md").exists(), "inventory.md was not written")
    assert_true((work / "audit" / "review_queue.json").exists(), "review_queue.json was not written")
    report_text = (work / "audit" / "inventory.json").read_text(encoding="utf-8")
    assert_true("Synthetic body text" not in report_text, "private inventory leaked body text")


@case
def source_audit_detects_high_risk_docx_structures() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_structures")
    path = work / "risky.docx"
    doc = Document()
    doc.add_paragraph("Body text")
    doc.save(path)

    def inject_risky_structures(xml: str) -> str:
        return xml.replace(
            "</w:body>",
            (
                "<w:p><w:r><w:pict><v:shape xmlns:v=\"urn:schemas-microsoft-com:vml\">"
                "<v:textbox><w:txbxContent><w:p><w:r><w:t>Box</w:t></w:r></w:p></w:txbxContent>"
                "</v:textbox></v:shape></w:pict></w:r></w:p>"
                "<w:sdt><w:sdtContent><w:p><w:r><w:t>Control</w:t></w:r></w:p></w:sdtContent></w:sdt>"
                "<w:p><w:r><w:ins><w:t>Inserted</w:t></w:ins></w:r></w:p>"
                "<w:sectPr><w:pgSz w:w=\"16838\" w:h=\"11906\" w:orient=\"landscape\"/></w:sectPr>"
                "</w:body>"
            ),
        )

    _rewrite_docx_part(path, "word/document.xml", inject_risky_structures)
    with zipfile.ZipFile(path, "a") as zf:
        zf.writestr("word/footnotes.xml", "<w:footnotes xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"><w:footnote w:id=\"2\"><w:p><w:r><w:t>Footnote</w:t></w:r></w:p></w:footnote></w:footnotes>")
        zf.writestr("word/endnotes.xml", "<w:endnotes xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"><w:endnote w:id=\"2\"><w:p><w:r><w:t>Endnote</w:t></w:r></w:p></w:endnote></w:endnotes>")
        zf.writestr("word/comments.xml", "<w:comments xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"><w:comment w:id=\"0\"><w:p><w:r><w:t>Comment</w:t></w:r></w:p></w:comment></w:comments>")
        zf.writestr("word/embeddings/oleObject1.bin", b"ole")

    audit = audit_docx_source(path)
    codes = {issue["code"] for issue in audit["issues"]}
    expected = {
        "SOURCE_TEXTBOX_UNSUPPORTED",
        "SOURCE_FOOTNOTE_UNSUPPORTED",
        "SOURCE_ENDNOTE_UNSUPPORTED",
        "COMMENTS_PRESENT",
        "CONTENT_CONTROL_UNSUPPORTED",
        "TRACKED_CHANGES_PRESENT",
        "SOURCE_EMBEDDED_OBJECT_UNSUPPORTED",
        "SOURCE_LANDSCAPE_SECTION_UNSUPPORTED",
    }
    assert_true(expected <= codes, f"source audit missed high-risk structures: {codes}")


@case
def complex_table_and_image_format_boundaries_are_structural_qa_visible() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_table_image")
    path = work / "table_image.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))
    doc.save(path)
    with zipfile.ZipFile(path, "a") as zf:
        zf.writestr("word/media/image99.wmf", b"wmf")

    audit = audit_docx_source(path)
    codes = {issue["code"] for issue in audit["issues"]}
    assert_true("TABLE_MERGE_UNSUPPORTED" in codes or "COMPLEX_TABLE_UNSUPPORTED" in codes, f"merged table was not reported: {codes}")
    assert_true("CONTENT_IMAGE_FORMAT_UNSUPPORTED" in codes, f"unsupported image format was not reported: {codes}")


@case
def source_audit_counts_legacy_hmerge_as_merged_cells() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_legacy_hmerge")
    path = work / "legacy_hmerge.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "Legacy merged header"
    table.cell(0, 1).text = ""
    table.cell(0, 2).text = "Score"
    doc.save(path)

    def inject_legacy_hmerge(xml: str) -> str:
        root = ET.fromstring(xml.encode("utf-8"))
        table_el = root.find(".//" + W_NS + "tbl")
        assert_true(table_el is not None, "test table missing")
        row = table_el.find(W_NS + "tr")
        assert_true(row is not None, "test row missing")
        cells = row.findall(W_NS + "tc")
        assert_true(len(cells) == 3, "test cells missing")

        def ensure_tc_pr(cell):
            tc_pr = cell.find(W_NS + "tcPr")
            if tc_pr is None:
                tc_pr = ET.Element(W_NS + "tcPr")
                cell.insert(0, tc_pr)
            return tc_pr

        restart = ET.SubElement(ensure_tc_pr(cells[0]), W_NS + "hMerge")
        restart.set(W_NS + "val", "restart")
        ET.SubElement(ensure_tc_pr(cells[1]), W_NS + "hMerge")
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(path, "word/document.xml", inject_legacy_hmerge)

    audit = audit_docx_source(str(path))
    codes = {issue["code"] for issue in audit["issues"]}
    assert_true(audit["counts"].get("hmerge_count") == 2, f"legacy hMerge count missing: {audit}")
    assert_true(audit["counts"].get("merged_cell_count") == 2, f"legacy hMerge was not included in merged cell count: {audit}")
    assert_true("TABLE_MERGE_UNSUPPORTED" in codes, f"legacy hMerge should surface merged-table review guidance: {audit}")


@case
def source_audit_flags_mixed_gridspan_hmerge_as_irregular_table() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_mixed_gridspan_hmerge")
    path = work / "mixed_gridspan_hmerge.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    table.cell(0, 0).text = "Mixed encoded header"
    table.cell(0, 1).text = ""
    table.cell(0, 2).text = "Score"
    table.cell(0, 3).text = "Note"
    doc.save(path)

    def inject_mixed_gridspan_hmerge(xml: str) -> str:
        root = ET.fromstring(xml.encode("utf-8"))
        table_el = root.find(".//" + W_NS + "tbl")
        assert_true(table_el is not None, "test table missing")
        row = table_el.find(W_NS + "tr")
        assert_true(row is not None, "test row missing")
        cells = row.findall(W_NS + "tc")
        assert_true(len(cells) == 4, "test cells missing")

        def ensure_tc_pr(cell: ET.Element) -> ET.Element:
            tc_pr = cell.find(W_NS + "tcPr")
            if tc_pr is None:
                tc_pr = ET.Element(W_NS + "tcPr")
                cell.insert(0, tc_pr)
            return tc_pr

        first_pr = ensure_tc_pr(cells[0])
        grid_span = first_pr.find(W_NS + "gridSpan")
        if grid_span is None:
            grid_span = ET.SubElement(first_pr, W_NS + "gridSpan")
        grid_span.set(W_NS + "val", "2")
        hmerge = ET.SubElement(first_pr, W_NS + "hMerge")
        hmerge.set(W_NS + "val", "restart")
        ET.SubElement(ensure_tc_pr(cells[1]), W_NS + "hMerge")
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(path, "word/document.xml", inject_mixed_gridspan_hmerge)

    audit = audit_docx_source(str(path))
    codes = {issue["code"] for issue in audit["issues"]}
    assert_true("TABLE_MERGE_UNSUPPORTED" in codes, f"mixed gridSpan/hMerge should still report merged-table review: {audit}")
    assert_true("COMPLEX_TABLE_UNSUPPORTED" in codes, f"mixed gridSpan/hMerge ambiguity was not reported: {audit}")
    assert_true(audit["counts"].get("max_table_columns") == 4, f"mixed gridSpan/hMerge should not double-count table width: {audit}")
    assert_true(audit["counts"].get("irregular_table_count") == 1, f"irregular table count missing: {audit}")
    assert_true(audit["counts"].get("irregular_hmerge_count") == 1, f"irregular hMerge count missing: {audit}")
    assert_true(audit["counts"].get("irregular_grid_span_count") == 0, f"mixed gridSpan/hMerge should not be mislabeled as gridSpan overflow: {audit}")
    detail = " ".join(str(issue.get("detail") or "") for issue in audit["issues"] if issue.get("code") == "COMPLEX_TABLE_UNSUPPORTED")
    assert_true("irregular_hmerges=1" in detail, f"complex table detail did not name irregular hMerge count: {audit}")


@case
def source_audit_flags_mixed_gridspan_hmerge_vmerge_without_gridspan_noise() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_mixed_gridspan_hmerge_vmerge")
    path = work / "mixed_gridspan_hmerge_vmerge.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    table.cell(0, 0).text = "Mixed 2D block"
    table.cell(0, 1).text = ""
    table.cell(0, 2).text = "Score"
    table.cell(0, 3).text = "Note"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(1, 2).text = "1"
    table.cell(1, 3).text = "Keep"
    doc.save(path)

    def inject_mixed_gridspan_hmerge_vmerge(xml: str) -> str:
        root = ET.fromstring(xml.encode("utf-8"))
        table_el = root.find(".//" + W_NS + "tbl")
        assert_true(table_el is not None, "test table missing")
        rows = table_el.findall(W_NS + "tr")
        assert_true(len(rows) == 2, "test rows missing")

        def ensure_tc_pr(cell: ET.Element) -> ET.Element:
            tc_pr = cell.find(W_NS + "tcPr")
            if tc_pr is None:
                tc_pr = ET.Element(W_NS + "tcPr")
                cell.insert(0, tc_pr)
            return tc_pr

        for row_idx in (0, 1):
            cells = rows[row_idx].findall(W_NS + "tc")
            assert_true(len(cells) == 4, "test row cells missing")
            first_pr = ensure_tc_pr(cells[0])
            grid_span = first_pr.find(W_NS + "gridSpan")
            if grid_span is None:
                grid_span = ET.SubElement(first_pr, W_NS + "gridSpan")
            grid_span.set(W_NS + "val", "2")
            hmerge = ET.SubElement(first_pr, W_NS + "hMerge")
            hmerge.set(W_NS + "val", "restart")
            vmerge = ET.SubElement(first_pr, W_NS + "vMerge")
            if row_idx == 0:
                vmerge.set(W_NS + "val", "restart")

            second_pr = ensure_tc_pr(cells[1])
            ET.SubElement(second_pr, W_NS + "hMerge")
            second_vmerge = ET.SubElement(second_pr, W_NS + "vMerge")
            if row_idx == 0:
                second_vmerge.set(W_NS + "val", "restart")
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(path, "word/document.xml", inject_mixed_gridspan_hmerge_vmerge)

    audit = audit_docx_source(str(path))
    codes = {issue["code"] for issue in audit["issues"]}
    assert_true("TABLE_MERGE_UNSUPPORTED" in codes, f"mixed gridSpan/hMerge/vMerge should report merged-table review: {audit}")
    assert_true("COMPLEX_TABLE_UNSUPPORTED" in codes, f"mixed gridSpan/hMerge/vMerge ambiguity was not reported: {audit}")
    assert_true(audit["counts"].get("max_table_columns") == 4, f"mixed gridSpan/hMerge/vMerge should not double-count table width: {audit}")
    assert_true(audit["counts"].get("irregular_table_count") == 1, f"irregular table count missing: {audit}")
    assert_true(audit["counts"].get("irregular_hmerge_count") == 2, f"irregular hMerge count missing: {audit}")
    assert_true(audit["counts"].get("irregular_grid_span_count") == 0, f"mixed 2D merge should not be mislabeled as gridSpan overflow: {audit}")
    detail = " ".join(str(issue.get("detail") or "") for issue in audit["issues"] if issue.get("code") == "COMPLEX_TABLE_UNSUPPORTED")
    assert_true("irregular_hmerges=2" in detail, f"complex table detail did not name mixed 2D hMerge count: {audit}")


@case
def source_audit_counts_visible_mixed_hmerge_continuations_without_content_leakage() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_visible_mixed_hmerge_continuations")
    image = work / "continuation.png"
    write_sample_png(image, width=96, height=72)
    path = work / "visible_mixed_hmerge_continuations.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    table.cell(0, 0).text = "Mixed 2D block"
    table.cell(0, 1).text = "Synthetic duplicate text"
    table.cell(0, 2).text = "Score"
    table.cell(0, 3).text = "Note"
    table.cell(1, 0).text = ""
    table.cell(1, 1).paragraphs[0].add_run().add_picture(str(image))
    table.cell(1, 2).text = "1"
    table.cell(1, 3).text = "Keep"
    doc.save(path)

    def inject_mixed_gridspan_hmerge_vmerge(xml: str) -> str:
        root = ET.fromstring(xml.encode("utf-8"))
        table_el = root.find(".//" + W_NS + "tbl")
        assert_true(table_el is not None, "test table missing")
        rows = table_el.findall(W_NS + "tr")
        assert_true(len(rows) == 2, "test rows missing")

        def ensure_tc_pr(cell: ET.Element) -> ET.Element:
            tc_pr = cell.find(W_NS + "tcPr")
            if tc_pr is None:
                tc_pr = ET.Element(W_NS + "tcPr")
                cell.insert(0, tc_pr)
            return tc_pr

        for row_idx in (0, 1):
            cells = rows[row_idx].findall(W_NS + "tc")
            assert_true(len(cells) == 4, "test row cells missing")
            first_pr = ensure_tc_pr(cells[0])
            grid_span = first_pr.find(W_NS + "gridSpan")
            if grid_span is None:
                grid_span = ET.SubElement(first_pr, W_NS + "gridSpan")
            grid_span.set(W_NS + "val", "2")
            hmerge = ET.SubElement(first_pr, W_NS + "hMerge")
            hmerge.set(W_NS + "val", "restart")
            vmerge = ET.SubElement(first_pr, W_NS + "vMerge")
            if row_idx == 0:
                vmerge.set(W_NS + "val", "restart")

            second_pr = ensure_tc_pr(cells[1])
            ET.SubElement(second_pr, W_NS + "hMerge")
            second_vmerge = ET.SubElement(second_pr, W_NS + "vMerge")
            if row_idx == 0:
                second_vmerge.set(W_NS + "val", "restart")
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(path, "word/document.xml", inject_mixed_gridspan_hmerge_vmerge)

    audit = audit_docx_source(str(path))
    codes = {issue["code"] for issue in audit["issues"]}
    assert_true("COMPLEX_TABLE_UNSUPPORTED" in codes, f"visible mixed continuation ambiguity was not reported: {audit}")
    assert_true(audit["counts"].get("visible_hmerge_continuation_count") == 2, f"visible continuation count missing: {audit}")
    assert_true(
        "Synthetic duplicate text" not in str(audit),
        f"source audit leaked visible duplicate continuation text: {audit}",
    )
    detail = " ".join(str(issue.get("detail") or "") for issue in audit["issues"] if issue.get("code") == "COMPLEX_TABLE_UNSUPPORTED")
    assert_true("visible_hmerge_continuations=2" in detail, f"complex table detail did not name visible continuation count: {audit}")


@case
def source_audit_allows_four_level_nested_tables_and_flags_deeper_nesting() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_nested_tables")

    one_level = work / "one_level_nested.docx"
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "Nested"
    doc.save(one_level)
    one_audit = audit_docx_source(str(one_level))
    one_codes = {issue["code"] for issue in one_audit["issues"]}
    assert_true(
        "COMPLEX_TABLE_UNSUPPORTED" not in one_codes,
        f"one-level nested table should be handled by the engine, not blocked as complex: {one_audit}",
    )
    assert_true(one_audit["counts"].get("nested_table_max_depth") == 1, f"one-level nested depth missing: {one_audit}")

    two_level = work / "two_level_nested.docx"
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    deeper = nested.cell(0, 0).add_table(rows=1, cols=1)
    deeper.cell(0, 0).text = "Deeper"
    doc.save(two_level)
    two_audit = audit_docx_source(str(two_level))
    two_codes = {issue["code"] for issue in two_audit["issues"]}
    assert_true(
        "COMPLEX_TABLE_UNSUPPORTED" not in two_codes,
        f"two-level nested table should be handled by the engine, not blocked as complex: {two_audit}",
    )
    assert_true(two_audit["counts"].get("nested_table_max_depth") == 2, f"two-level nested depth missing: {two_audit}")

    three_level = work / "three_level_nested.docx"
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    deeper = nested.cell(0, 0).add_table(rows=1, cols=1)
    deepest = deeper.cell(0, 0).add_table(rows=1, cols=1)
    deepest.cell(0, 0).text = "Deepest"
    doc.save(three_level)
    three_audit = audit_docx_source(str(three_level))
    three_codes = {issue["code"] for issue in three_audit["issues"]}
    assert_true(
        "COMPLEX_TABLE_UNSUPPORTED" not in three_codes,
        f"three-level nested table should be handled by the engine, not blocked as complex: {three_audit}",
    )
    assert_true(three_audit["counts"].get("nested_table_max_depth") == 3, f"three-level nested depth missing: {three_audit}")

    four_level = work / "four_level_nested.docx"
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    deeper = nested.cell(0, 0).add_table(rows=1, cols=1)
    deepest = deeper.cell(0, 0).add_table(rows=1, cols=1)
    too_deep = deepest.cell(0, 0).add_table(rows=1, cols=1)
    too_deep.cell(0, 0).text = "Too deep"
    doc.save(four_level)
    four_audit = audit_docx_source(str(four_level))
    four_codes = {issue["code"] for issue in four_audit["issues"]}
    assert_true(
        "COMPLEX_TABLE_UNSUPPORTED" not in four_codes,
        f"four-level nested table should be handled by the engine, not blocked as complex: {four_audit}",
    )
    assert_true(four_audit["counts"].get("nested_table_max_depth") == 4, f"four-level nested depth missing: {four_audit}")

    five_level = work / "five_level_nested.docx"
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    deeper = nested.cell(0, 0).add_table(rows=1, cols=1)
    deepest = deeper.cell(0, 0).add_table(rows=1, cols=1)
    too_deep = deepest.cell(0, 0).add_table(rows=1, cols=1)
    beyond_limit = too_deep.cell(0, 0).add_table(rows=1, cols=1)
    beyond_limit.cell(0, 0).text = "Still too deep"
    doc.save(five_level)
    five_audit = audit_docx_source(str(five_level))
    five_codes = {issue["code"] for issue in five_audit["issues"]}
    assert_true("COMPLEX_TABLE_UNSUPPORTED" in five_codes, f"five-level nested table was not reported: {five_audit}")
    assert_true(five_audit["counts"].get("nested_table_max_depth") == 5, f"five-level nested depth missing: {five_audit}")


@case
def source_audit_flags_irregular_table_merge_grid() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_irregular_table")
    path = work / "irregular_table.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Top left"
    table.cell(0, 1).text = "Top right"
    table.cell(1, 0).text = "Orphan continue should not disappear silently"
    table.cell(1, 1).text = "Bottom right"
    doc.save(path)

    def inject_orphan_vmerge(xml: str) -> str:
        marker = '<w:t>Orphan continue should not disappear silently</w:t>'
        replacement = '<w:vMerge/></w:tcPr><w:p><w:r>' + marker
        return xml.replace(
            '<w:tcPr><w:tcW w:type="dxa" w:w="4320"/></w:tcPr><w:p><w:r>' + marker,
            '<w:tcPr><w:tcW w:type="dxa" w:w="4320"/>' + replacement,
            1,
        )

    _rewrite_docx_part(path, "word/document.xml", inject_orphan_vmerge)

    audit = audit_docx_source(str(path))
    codes = {issue["code"] for issue in audit["issues"]}
    assert_true("COMPLEX_TABLE_UNSUPPORTED" in codes, f"irregular merge grid was not reported: {audit}")
    assert_true(audit["counts"].get("irregular_table_count") == 1, f"irregular table count missing: {audit}")
    detail = " ".join(str(issue.get("detail") or "") for issue in audit["issues"] if issue.get("code") == "COMPLEX_TABLE_UNSUPPORTED")
    assert_true("irregular_tables=1" in detail, f"complex table detail did not name irregular table count: {audit}")


@case
def source_audit_flags_nonrectangular_legacy_hmerge_vmerge() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_nonrectangular_hmerge_vmerge")
    path = work / "nonrectangular_hmerge_vmerge.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Wide top"
    table.cell(0, 1).text = ""
    table.cell(0, 2).text = "Score"
    table.cell(1, 0).text = "Continuation should require review"
    table.cell(1, 1).text = "Beta"
    table.cell(1, 2).text = "1"
    doc.save(path)

    def inject_nonrectangular_hmerge_vmerge(xml: str) -> str:
        root = ET.fromstring(xml.encode("utf-8"))
        table_el = root.find(".//" + W_NS + "tbl")
        assert_true(table_el is not None, "test table missing")
        rows = table_el.findall(W_NS + "tr")
        assert_true(len(rows) == 2, "test rows missing")
        first_row_cells = rows[0].findall(W_NS + "tc")
        second_row_cells = rows[1].findall(W_NS + "tc")
        assert_true(len(first_row_cells) == 3 and len(second_row_cells) == 3, "test row cells missing")

        def ensure_tc_pr(cell: ET.Element) -> ET.Element:
            tc_pr = cell.find(W_NS + "tcPr")
            if tc_pr is None:
                tc_pr = ET.Element(W_NS + "tcPr")
                cell.insert(0, tc_pr)
            return tc_pr

        first_pr = ensure_tc_pr(first_row_cells[0])
        hmerge = ET.SubElement(first_pr, W_NS + "hMerge")
        hmerge.set(W_NS + "val", "restart")
        vmerge = ET.SubElement(first_pr, W_NS + "vMerge")
        vmerge.set(W_NS + "val", "restart")
        ET.SubElement(ensure_tc_pr(first_row_cells[1]), W_NS + "hMerge")
        ET.SubElement(ensure_tc_pr(second_row_cells[0]), W_NS + "vMerge")
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(path, "word/document.xml", inject_nonrectangular_hmerge_vmerge)

    audit = audit_docx_source(str(path))
    codes = {issue["code"] for issue in audit["issues"]}
    assert_true("COMPLEX_TABLE_UNSUPPORTED" in codes, f"nonrectangular hMerge/vMerge was not reported: {audit}")
    assert_true(audit["counts"].get("irregular_table_count") == 1, f"nonrectangular table count missing: {audit}")
    assert_true(audit["counts"].get("irregular_vmerge_count") == 1, f"nonrectangular vMerge count missing: {audit}")
    detail = " ".join(str(issue.get("detail") or "") for issue in audit["issues"] if issue.get("code") == "COMPLEX_TABLE_UNSUPPORTED")
    assert_true("irregular_vmerges=1" in detail, f"complex table detail did not name irregular vMerge count: {audit}")


@case
def source_audit_flags_landscape_wide_table_risk() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_landscape_wide_table")
    path = work / "landscape_wide_table.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=9)
    for idx, cell in enumerate(table.rows[0].cells):
        cell.text = f"C{idx + 1}"
    doc.save(path)

    def inject_landscape_section(xml: str) -> str:
        return xml.replace(
            "<w:pgSz w:w=\"12240\" w:h=\"15840\"/>",
            "<w:pgSz w:w=\"15840\" w:h=\"12240\" w:orient=\"landscape\"/>",
            1,
        )

    _rewrite_docx_part(path, "word/document.xml", inject_landscape_section)

    audit = audit_docx_source(str(path))
    codes = {issue["code"] for issue in audit["issues"]}
    assert_true("SOURCE_LANDSCAPE_SECTION_UNSUPPORTED" in codes, f"landscape section was not reported: {audit}")
    assert_true("COMPLEX_TABLE_UNSUPPORTED" in codes, f"wide table was not reported as complex: {audit}")
    assert_true(audit["counts"].get("wide_table_count") == 1, f"wide table count missing: {audit}")
    assert_true(audit["counts"].get("landscape_wide_table_risk_count") == 1, f"landscape wide table risk missing: {audit}")
    detail = " ".join(str(issue.get("detail") or "") for issue in audit["issues"] if issue.get("code") == "COMPLEX_TABLE_UNSUPPORTED")
    assert_true("wide_tables=1" in detail and "landscape_wide_tables=1" in detail, f"wide table detail incomplete: {audit}")


@case
def source_audit_counts_grid_after_row_omissions_as_columns() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_grid_after_wide_table")
    path = work / "grid_after_wide_table.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Visible first cell"
    doc.save(path)

    def inject_grid_after(xml: str) -> str:
        root = ET.fromstring(xml.encode("utf-8"))
        table_el = root.find(".//" + W_NS + "tbl")
        assert_true(table_el is not None, "table missing")
        old_grid = table_el.find(W_NS + "tblGrid")
        if old_grid is not None:
            table_el.remove(old_grid)
        tbl_grid = ET.Element(W_NS + "tblGrid")
        for _ in range(9):
            col = ET.SubElement(tbl_grid, W_NS + "gridCol")
            col.set(W_NS + "w", "900")
        table_el.insert(1, tbl_grid)
        row = table_el.find(W_NS + "tr")
        assert_true(row is not None, "row missing")
        tr_pr = row.find(W_NS + "trPr")
        if tr_pr is None:
            tr_pr = ET.Element(W_NS + "trPr")
            row.insert(0, tr_pr)
        grid_after = ET.SubElement(tr_pr, W_NS + "gridAfter")
        grid_after.set(W_NS + "val", "8")
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(path, "word/document.xml", inject_grid_after)

    audit = audit_docx_source(str(path))
    codes = {issue["code"] for issue in audit["issues"]}
    assert_true(audit["counts"].get("max_table_columns") == 9, f"gridAfter columns were not counted: {audit}")
    assert_true(audit["counts"].get("wide_table_count") == 1, f"gridAfter wide table was not counted: {audit}")
    assert_true("COMPLEX_TABLE_UNSUPPORTED" in codes, f"gridAfter wide table should require review: {audit}")
    detail = " ".join(str(issue.get("detail") or "") for issue in audit["issues"] if issue.get("code") == "COMPLEX_TABLE_UNSUPPORTED")
    assert_true("wide_tables=1" in detail, f"gridAfter wide table detail incomplete: {audit}")


@case
def source_audit_counts_revision_wrapped_table_rows_as_visible_width() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_revision_wrapped_wide_table")
    path = work / "revision_wrapped_wide_table.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=9)
    for idx, cell in enumerate(table.rows[0].cells):
        cell.text = f"Wide {idx + 1}"
    doc.save(path)

    def wrap_row_in_revision(xml: str) -> str:
        root = ET.fromstring(xml.encode("utf-8"))
        table_el = root.find(".//" + W_NS + "tbl")
        assert_true(table_el is not None, "table missing")
        row = table_el.find(W_NS + "tr")
        assert_true(row is not None, "row missing")
        row_index = list(table_el).index(row)
        table_el.remove(row)
        ins = ET.Element(W_NS + "ins")
        ins.set(W_NS + "id", "1")
        ins.set(W_NS + "author", "Regression")
        ins.append(row)
        table_el.insert(row_index, ins)
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(path, "word/document.xml", wrap_row_in_revision)

    audit = audit_docx_source(str(path))
    codes = {issue["code"] for issue in audit["issues"]}
    assert_true("TRACKED_CHANGES_PRESENT" in codes, f"revision wrapper should remain visible in audit: {audit}")
    assert_true(audit["counts"].get("max_table_columns") == 9, f"revision-wrapped row width was not counted: {audit}")
    assert_true(audit["counts"].get("wide_table_count") == 1, f"revision-wrapped wide row was not counted: {audit}")
    assert_true("COMPLEX_TABLE_UNSUPPORTED" in codes, f"revision-wrapped wide table should require review: {audit}")
    detail = " ".join(str(issue.get("detail") or "") for issue in audit["issues"] if issue.get("code") == "COMPLEX_TABLE_UNSUPPORTED")
    assert_true("wide_tables=1" in detail, f"revision-wrapped wide table detail incomplete: {audit}")


@case
def source_audit_ignores_deleted_revision_tables_for_table_risk() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_deleted_revision_table_risk")
    path = work / "deleted_revision_table_risk.docx"
    doc = Document()
    doc.add_paragraph("Visible body text.")
    doc.save(path)

    def inject_deleted_wide_table(xml: str) -> str:
        root = ET.fromstring(xml.encode("utf-8"))
        body = root.find(W_NS + "body")
        assert_true(body is not None, "document body missing")
        table_el = ET.Element(W_NS + "tbl")
        tbl_grid = ET.SubElement(table_el, W_NS + "tblGrid")
        for _ in range(9):
            col = ET.SubElement(tbl_grid, W_NS + "gridCol")
            col.set(W_NS + "w", "900")
        row = ET.SubElement(table_el, W_NS + "tr")
        for idx in range(9):
            cell = ET.SubElement(row, W_NS + "tc")
            tc_pr = ET.SubElement(cell, W_NS + "tcPr")
            tc_w = ET.SubElement(tc_pr, W_NS + "tcW")
            tc_w.set(W_NS + "type", "dxa")
            tc_w.set(W_NS + "w", "900")
            para = ET.SubElement(cell, W_NS + "p")
            run = ET.SubElement(para, W_NS + "r")
            text = ET.SubElement(run, W_NS + "t")
            text.text = f"Deleted {idx + 1}"
        move_from = ET.Element(W_NS + "moveFrom")
        move_from.set(W_NS + "id", "1")
        move_from.set(W_NS + "author", "Regression")
        move_from.append(table_el)
        sect_pr = body.find(W_NS + "sectPr")
        insert_at = list(body).index(sect_pr) if sect_pr is not None else len(list(body))
        body.insert(insert_at, move_from)
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(path, "word/document.xml", inject_deleted_wide_table)

    audit = audit_docx_source(str(path))
    codes = {issue["code"] for issue in audit["issues"]}
    assert_true("TRACKED_CHANGES_PRESENT" in codes, f"deleted revision wrapper should remain visible in audit: {audit}")
    assert_true(audit["counts"].get("table_count") == 0, f"deleted revision table should not count as a visible table: {audit}")
    assert_true(audit["counts"].get("max_table_columns") == 0, f"deleted revision table should not affect visible max columns: {audit}")
    assert_true(audit["counts"].get("wide_table_count") == 0, f"deleted revision table should not count as a visible wide table: {audit}")
    assert_true("COMPLEX_TABLE_UNSUPPORTED" not in codes, f"deleted revision table created a false complex-table warning: {audit}")


@case
def source_audit_does_not_double_count_nested_wide_tables() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_nested_wide_table_count")
    path = work / "nested_wide_table.docx"
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=9)
    for idx, cell in enumerate(nested.rows[0].cells):
        cell.text = f"N{idx + 1}"
    doc.save(path)

    audit = audit_docx_source(str(path))
    assert_true(audit["counts"].get("wide_table_count") == 1, f"nested wide table was double-counted: {audit}")


@case
def source_audit_does_not_mark_portrait_wide_table_as_landscape_risk() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_portrait_wide_then_landscape")
    path = work / "portrait_wide_then_landscape.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=9)
    for idx, cell in enumerate(table.rows[0].cells):
        cell.text = f"P{idx + 1}"
    doc.add_paragraph("Portrait section ends")
    doc.add_paragraph("Landscape section has no table")
    doc.save(path)

    def inject_final_landscape_section(xml: str) -> str:
        root = ET.fromstring(xml.encode("utf-8"))
        body = root.find(W_NS + "body")
        assert_true(body is not None, "document body missing")
        for para in body.findall(W_NS + "p"):
            text = "".join(node.text or "" for node in para.iter(W_NS + "t"))
            if text != "Portrait section ends":
                continue
            p_pr = para.find(W_NS + "pPr")
            if p_pr is None:
                p_pr = ET.Element(W_NS + "pPr")
                para.insert(0, p_pr)
            sect_pr = ET.Element(W_NS + "sectPr")
            pg_sz = ET.SubElement(sect_pr, W_NS + "pgSz")
            pg_sz.set(W_NS + "w", "12240")
            pg_sz.set(W_NS + "h", "15840")
            p_pr.append(sect_pr)
            break
        final_sect = body.find(W_NS + "sectPr")
        if final_sect is None:
            final_sect = ET.SubElement(body, W_NS + "sectPr")
        final_sect.clear()
        pg_sz = ET.SubElement(final_sect, W_NS + "pgSz")
        pg_sz.set(W_NS + "w", "15840")
        pg_sz.set(W_NS + "h", "12240")
        pg_sz.set(W_NS + "orient", "landscape")
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(path, "word/document.xml", inject_final_landscape_section)

    audit = audit_docx_source(str(path))
    assert_true(audit["counts"].get("wide_table_count") == 1, f"wide table count missing: {audit}")
    assert_true(audit["counts"].get("landscape_section_count") == 1, f"landscape section count missing: {audit}")
    assert_true(
        audit["counts"].get("landscape_wide_table_risk_count") == 0,
        f"portrait wide table was incorrectly marked as landscape risk: {audit}",
    )


@case
def source_audit_respects_grid_before_for_vmerge_columns() -> None:
    from content_parser_modules.source_audit import audit_docx_source

    work = new_workdir("source_audit_grid_before_vmerge")
    path = work / "grid_before_vmerge.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Left"
    table.cell(0, 1).text = "Vertical start"
    table.cell(1, 1).text = "Vertical continue"
    doc.save(path)

    def inject_grid_before_vmerge(xml: str) -> str:
        root = ET.fromstring(xml.encode("utf-8"))
        table_el = root.find(".//" + W_NS + "tbl")
        assert_true(table_el is not None, "table missing")
        rows = table_el.findall(W_NS + "tr")
        assert_true(len(rows) >= 2, "table rows missing")
        first_row_cells = rows[0].findall(W_NS + "tc")
        second_row_cells = rows[1].findall(W_NS + "tc")
        assert_true(len(first_row_cells) == 2 and len(second_row_cells) == 2, "test table shape changed")

        def ensure_tc_pr(cell):
            tc_pr = cell.find(W_NS + "tcPr")
            if tc_pr is None:
                tc_pr = ET.Element(W_NS + "tcPr")
                cell.insert(0, tc_pr)
            return tc_pr

        restart = ET.SubElement(ensure_tc_pr(first_row_cells[1]), W_NS + "vMerge")
        restart.set(W_NS + "val", "restart")
        rows[1].remove(second_row_cells[0])
        tr_pr = ET.Element(W_NS + "trPr")
        grid_before = ET.SubElement(tr_pr, W_NS + "gridBefore")
        grid_before.set(W_NS + "val", "1")
        rows[1].insert(0, tr_pr)
        ET.SubElement(ensure_tc_pr(second_row_cells[1]), W_NS + "vMerge")
        return ET.tostring(root, encoding="unicode")

    _rewrite_docx_part(path, "word/document.xml", inject_grid_before_vmerge)

    audit = audit_docx_source(str(path))
    codes = {issue["code"] for issue in audit["issues"]}
    assert_true("TABLE_MERGE_UNSUPPORTED" in codes, f"vMerge should still be reported as merged table: {audit}")
    assert_true("COMPLEX_TABLE_UNSUPPORTED" not in codes, f"valid gridBefore vMerge was misreported as complex: {audit}")
    assert_true(audit["counts"].get("irregular_vmerge_count") == 0, f"valid gridBefore vMerge was counted irregular: {audit}")


@case
def source_audit_issues_are_promoted_to_structural_qa() -> None:
    work = new_workdir("source_audit_structural_qa")
    doc = Document()
    doc.add_paragraph("Rendered body")
    doc.save(work / "最终论文.docx")
    issue_codes = [
        "SOURCE_FORMAT_UNSUPPORTED",
        "LEGACY_DOC_UNSUPPORTED",
        "SOURCE_TEXTBOX_UNSUPPORTED",
        "SOURCE_FOOTNOTE_UNSUPPORTED",
        "SOURCE_ENDNOTE_UNSUPPORTED",
        "TRACKED_CHANGES_PRESENT",
        "COMMENTS_PRESENT",
        "CONTENT_CONTROL_UNSUPPORTED",
        "SOURCE_EMBEDDED_OBJECT_UNSUPPORTED",
        "SOURCE_LANDSCAPE_SECTION_UNSUPPORTED",
        "CONTENT_IMAGE_FORMAT_UNSUPPORTED",
        "COMPLEX_TABLE_UNSUPPORTED",
        "TABLE_MERGE_UNSUPPORTED",
    ]
    write_json(
        work / "content.json",
        {
            "_meta": {
                "source_audit": {
                    "issues": [
                        {
                            "code": code,
                            "severity": "error" if code not in {"COMMENTS_PRESENT", "SOURCE_LANDSCAPE_SECTION_UNSUPPORTED", "COMPLEX_TABLE_UNSUPPORTED", "TABLE_MERGE_UNSUPPORTED"} else "warning",
                            "message": code,
                            "detail": "synthetic structural risk",
                        }
                        for code in issue_codes
                    ]
                }
            },
            "title_info": {"title_cn": "Synthetic"},
            "sections": [{"heading": "1 Introduction", "level": 1, "paragraphs": ["Rendered body"]}],
            "references": ["[1] Synthetic reference."],
        },
    )
    write_json(work / "format.json", {"_meta": {}, "paragraphs": [{"text": "style sample"}]})

    report = check_output(str(work), mode="developer")
    codes = {issue["code"] for issue in report["issues"]}
    assert_true(set(issue_codes) <= codes, f"structural QA did not promote source audit issues: {codes}")
    textbox = next(issue for issue in report["issues"] if issue["code"] == "SOURCE_TEXTBOX_UNSUPPORTED")
    assert_true("source_audit.py" in textbox["owner_developer"], f"source audit owner was not preserved: {textbox}")
    assert_true(report["passed"] is False, "source audit error codes should fail structural QA")


@case
def comparison_assessment_promotes_warning_runs_to_review() -> None:
    from comparison_assessment import assess_run

    work = new_workdir("comparison_assessment")
    structural = {
        "schema_version": 1,
        "passed": True,
        "status": "passed_with_warnings",
        "issues": [{"code": "REFERENCES_MISSING", "severity": "warning"}],
    }
    (work / "qa_report.json").write_text(json.dumps(structural, ensure_ascii=False), encoding="utf-8")
    assessment = assess_run(work)
    assert_true(assessment["decision"] == "PASSED_WITH_REVIEW", f"warning-only run should require review: {assessment}")
    assert_true(assessment["manual_review_required"] is True, "warning-only run should expose manual_review_required")
    assert_true("REFERENCES_MISSING" in assessment["blocking_issue_codes"], "leading warning code was not surfaced")


@case
def visual_sample_pages_pick_useful_pages() -> None:
    import qa_visual

    pages = ["cover", "contents", "blank", "1. Introduction", "middle"] + ["body"] * 7
    samples = qa_visual._sample_pages(12, pages)
    assert_true(1 in samples and 2 in samples and 4 in samples and 6 in samples, "sample page selection missed key pages")


@case
def visual_sample_pages_prioritize_late_risk_content_pages() -> None:
    import qa_visual

    pages = [
        "cover",
        "contents",
        "preface",
        "1. Introduction",
        "background prose",
        "method prose",
        "middle prose",
        "body prose",
        "图 2 模型结构",
        "Table 3 Ablation Results",
        "公式 (3.1) E = mc^2",
        "discussion",
        "references",
        "appendix",
    ]
    samples = qa_visual._sample_pages(14, pages)
    assert_true(len(samples) <= 6, f"sample page selection should stay bounded: {samples}")
    assert_true(
        1 in samples and 2 in samples and 4 in samples,
        f"sample page selection lost cover/TOC/body anchors: {samples}",
    )
    assert_true(
        9 in samples and 10 in samples and 11 in samples,
        f"sample page selection missed late figure/table/formula risk pages: {samples}",
    )


@case
def visual_qa_fails_closed_without_pdf_tools() -> None:
    import qa_visual

    work = new_workdir("visual_closed")
    (work / "final.docx").write_bytes(b"not a real docx; export is monkeypatched")
    fake_pdf = work / "fake.pdf"
    fake_pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")
    original_export = qa_visual._export_pdf
    original_pdfinfo = qa_visual._pdfinfo
    original_pages_text = qa_visual._pdf_pages_text
    original_render = qa_visual._render_samples
    try:
        qa_visual._export_pdf = lambda _docx, _visual_dir: str(fake_pdf)
        qa_visual._pdfinfo = lambda _pdf: {"available": False}
        qa_visual._pdf_pages_text = lambda _pdf, _visual_dir: []
        qa_visual._render_samples = lambda _pdf, _visual_dir, _pages: []
        report = qa_visual.check_visual(str(work), output_docx_name="final.docx")
    finally:
        qa_visual._export_pdf = original_export
        qa_visual._pdfinfo = original_pdfinfo
        qa_visual._pdf_pages_text = original_pages_text
        qa_visual._render_samples = original_render
    codes = [item["code"] for item in report["issues"]]
    assert_true(report["passed"] is False, "visual QA passed without pdfinfo/text validation")
    assert_true("PDFINFO_UNAVAILABLE" in codes, "missing pdfinfo was not reported")
    assert_true("Poppler" in report.get("next_action", ""), f"visual QA did not guide dependency repair: {report}")
    assert_true("下一步" in qa_visual.report_to_markdown(report), "visual report markdown should include next action")


@case
def visual_golden_compare_missing_does_not_auto_create_baseline() -> None:
    from qa_visual_modules.golden import _compare_or_update_golden

    work = new_workdir("visual_golden_missing")
    golden_dir = work / "golden"
    result = _compare_or_update_golden(
        str(work),
        counts={"pages": 1, "page_width_pt": 595.3, "page_height_pt": 841.9, "text_pages": 1, "all_page_images": 1},
        pages_text=["body"],
        image_stats={"page_hashes": ["0" * 16]},
        golden_dir=str(golden_dir),
        update_golden=False,
    )
    assert_true(result["status"] == "missing", f"compare-only missing golden should stay missing: {result}")
    assert_true(not list(golden_dir.glob("*.json")), "compare-only golden mode should not create baseline files")


@case
def visual_qa_fails_closed_when_required_wps_pdfinfo_fails() -> None:
    from qa_visual_modules import checks as visual_checks

    work = new_workdir("visual_wps_pdfinfo_failed")
    (work / "final.docx").write_bytes(b"not a real docx; export is monkeypatched")
    word_pdf = work / "word.pdf"
    wps_pdf = work / "wps.pdf"
    word_pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")
    wps_pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")

    original_export = visual_checks._export_pdf
    original_wps_export = visual_checks._export_wps_pdf
    original_pdfinfo = visual_checks._pdfinfo
    original_pages_text = visual_checks._pdf_pages_text
    original_render = visual_checks._render_samples
    try:
        visual_checks._export_pdf = lambda _docx, _visual_dir: str(word_pdf)
        visual_checks._export_wps_pdf = lambda _docx, _visual_dir: str(wps_pdf)

        def fake_pdfinfo(path):
            if Path(path).name == "wps.pdf":
                return {"available": True, "error": "xref table broken"}
            return {"available": True, "pages": 3, "page_width_pt": 595.3, "page_height_pt": 841.9}

        visual_checks._pdfinfo = fake_pdfinfo
        visual_checks._pdf_pages_text = lambda _pdf, _visual_dir: ["cover", "目录", "1 Introduction"]
        visual_checks._render_samples = lambda _pdf, _visual_dir, pages: [str(work / f"page_{page}.png") for page in pages]
        report = visual_checks.check_visual(str(work), output_docx_name="final.docx", require_wps=True, render_all_pages=False)
    finally:
        visual_checks._export_pdf = original_export
        visual_checks._export_wps_pdf = original_wps_export
        visual_checks._pdfinfo = original_pdfinfo
        visual_checks._pdf_pages_text = original_pages_text
        visual_checks._render_samples = original_render

    codes = [item["code"] for item in report["issues"]]
    assert_true(report["passed"] is False, f"required WPS visual QA should fail closed when WPS PDF metadata cannot be read: {report}")
    assert_true("WPS_PDFINFO_FAILED" in codes, f"WPS PDF metadata failure was not reported: {report}")
    assert_true("WPS" in report.get("next_action", "") and "重跑 visual QA" in report.get("next_action", ""), f"WPS PDF metadata failure lacked a concrete next action: {report}")


@case
def visual_qa_fails_closed_when_wps_page_size_differs() -> None:
    from qa_visual_modules import checks as visual_checks

    work = new_workdir("visual_wps_page_size_mismatch")
    (work / "final.docx").write_bytes(b"not a real docx; export is monkeypatched")
    word_pdf = work / "word.pdf"
    wps_pdf = work / "wps.pdf"
    word_pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")
    wps_pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")

    original_export = visual_checks._export_pdf
    original_wps_export = visual_checks._export_wps_pdf
    original_pdfinfo = visual_checks._pdfinfo
    original_pages_text = visual_checks._pdf_pages_text
    original_render = visual_checks._render_samples
    try:
        visual_checks._export_pdf = lambda _docx, _visual_dir: str(word_pdf)
        visual_checks._export_wps_pdf = lambda _docx, _visual_dir: str(wps_pdf)

        def fake_pdfinfo(path):
            if Path(path).name == "wps.pdf":
                return {"available": True, "pages": 3, "page_width_pt": 841.9, "page_height_pt": 595.3}
            return {"available": True, "pages": 3, "page_width_pt": 595.3, "page_height_pt": 841.9}

        visual_checks._pdfinfo = fake_pdfinfo
        visual_checks._pdf_pages_text = lambda _pdf, _visual_dir: ["cover", "目录", "1 Introduction"]
        visual_checks._render_samples = lambda _pdf, _visual_dir, pages: [str(work / f"page_{page}.png") for page in pages]
        report = visual_checks.check_visual(str(work), output_docx_name="final.docx", require_wps=False, render_all_pages=False)
    finally:
        visual_checks._export_pdf = original_export
        visual_checks._export_wps_pdf = original_wps_export
        visual_checks._pdfinfo = original_pdfinfo
        visual_checks._pdf_pages_text = original_pages_text
        visual_checks._render_samples = original_render

    codes = [item["code"] for item in report["issues"]]
    assert_true(report["passed"] is False, f"WPS page-size mismatch should fail visual QA even when WPS is optional: {report}")
    assert_true("WPS_PAGE_SIZE_MISMATCH" in codes, f"WPS page-size mismatch was not reported: {report}")
    action = report.get("next_action", "")
    assert_true("WPS" in action and ("纸张" in action or "页面尺寸" in action) and "重跑 visual QA" in action, f"WPS page-size mismatch lacked a concrete next action: {report}")


@case
def visual_qa_fails_closed_when_wps_text_pages_are_missing() -> None:
    from qa_visual_modules import checks as visual_checks

    work = new_workdir("visual_wps_text_page_mismatch")
    (work / "final.docx").write_bytes(b"not a real docx; export is monkeypatched")
    word_pdf = work / "word.pdf"
    wps_pdf = work / "wps.pdf"
    word_pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")
    wps_pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")

    original_export = visual_checks._export_pdf
    original_wps_export = visual_checks._export_wps_pdf
    original_pdfinfo = visual_checks._pdfinfo
    original_pages_text = visual_checks._pdf_pages_text
    original_render = visual_checks._render_samples
    try:
        visual_checks._export_pdf = lambda _docx, _visual_dir: str(word_pdf)
        visual_checks._export_wps_pdf = lambda _docx, _visual_dir: str(wps_pdf)

        def fake_pdfinfo(_path):
            return {"available": True, "pages": 3, "page_width_pt": 595.3, "page_height_pt": 841.9}

        def fake_pages_text(path, _visual_dir):
            if Path(path).name == "wps.pdf":
                return ["", "   ", ""]
            return ["cover", "目录", "1 Introduction"]

        visual_checks._pdfinfo = fake_pdfinfo
        visual_checks._pdf_pages_text = fake_pages_text
        visual_checks._render_samples = lambda _pdf, _visual_dir, pages: [str(work / f"page_{page}.png") for page in pages]
        report = visual_checks.check_visual(str(work), output_docx_name="final.docx", require_wps=False, render_all_pages=False)
    finally:
        visual_checks._export_pdf = original_export
        visual_checks._export_wps_pdf = original_wps_export
        visual_checks._pdfinfo = original_pdfinfo
        visual_checks._pdf_pages_text = original_pages_text
        visual_checks._render_samples = original_render

    codes = [item["code"] for item in report["issues"]]
    assert_true(report["passed"] is False, f"WPS text-page mismatch should fail visual QA even when WPS is optional: {report}")
    assert_true("WPS_TEXT_PAGE_MISMATCH" in codes, f"WPS text-page mismatch was not reported: {report}")
    assert_true(report["counts"].get("wps_text_pages") == 0, f"WPS text-page count was not recorded: {report}")
    action = report.get("next_action", "")
    assert_true("WPS" in action and ("文本" in action or "内容" in action) and "重跑 visual QA" in action, f"WPS text-page mismatch lacked a concrete next action: {report}")


@case
def visual_qa_preserves_word_and_wps_text_diagnostics() -> None:
    from qa_visual_modules import checks as visual_checks

    work = new_workdir("visual_wps_text_diagnostics")
    (work / "final.docx").write_bytes(b"not a real docx; export is monkeypatched")
    word_pdf = work / "word.pdf"
    wps_pdf = work / "wps.pdf"
    word_pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")
    wps_pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")

    original_export = visual_checks._export_pdf
    original_wps_export = visual_checks._export_wps_pdf
    original_pdfinfo = visual_checks._pdfinfo
    original_pages_text = visual_checks._pdf_pages_text
    original_render = visual_checks._render_samples
    original_which = visual_checks.shutil.which
    try:
        visual_checks._export_pdf = lambda _docx, _visual_dir: str(word_pdf)
        visual_checks._export_wps_pdf = lambda _docx, _visual_dir: str(wps_pdf)
        visual_checks._pdfinfo = lambda _path: {"available": True, "pages": 1, "page_width_pt": 595.3, "page_height_pt": 841.9}
        visual_checks._render_samples = lambda _pdf, _visual_dir, pages: [str(work / f"page_{page}.png") for page in pages]
        visual_checks.shutil.which = lambda name: "tool" if name == "pdftotext" else original_which(name)

        def fake_pages_text(path, visual_dir):
            text = "WPS diagnostic text\n" if Path(path).name == "wps.pdf" else "Word diagnostic text\n"
            Path(visual_dir, "rendered.txt").write_text(text, encoding="utf-8")
            return [text]

        visual_checks._pdf_pages_text = fake_pages_text
        report = visual_checks.check_visual(str(work), output_docx_name="final.docx", require_wps=False, render_all_pages=False)
    finally:
        visual_checks._export_pdf = original_export
        visual_checks._export_wps_pdf = original_wps_export
        visual_checks._pdfinfo = original_pdfinfo
        visual_checks._pdf_pages_text = original_pages_text
        visual_checks._render_samples = original_render
        visual_checks.shutil.which = original_which

    visual_dir = work / "visual_qa"
    assert_true(report["passed"] is True, f"diagnostic-preservation fixture should not introduce visual blockers: {report}")
    assert_true((visual_dir / "rendered_word.txt").read_text(encoding="utf-8") == "Word diagnostic text\n", "Word rendered text diagnostic was not preserved")
    assert_true((visual_dir / "rendered_wps.txt").read_text(encoding="utf-8") == "WPS diagnostic text\n", "WPS rendered text diagnostic was not preserved")
    assert_true((visual_dir / "rendered.txt").read_text(encoding="utf-8") == "Word diagnostic text\n", "legacy rendered.txt should remain the Word-rendered text artifact")


@case
def visual_qa_fails_closed_when_wps_sample_images_differ() -> None:
    from PIL import Image
    from qa_visual_modules import checks as visual_checks

    work = new_workdir("visual_wps_sample_image_mismatch")
    (work / "final.docx").write_bytes(b"not a real docx; export is monkeypatched")
    word_pdf = work / "word.pdf"
    wps_pdf = work / "wps.pdf"
    word_pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")
    wps_pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")

    original_export = visual_checks._export_pdf
    original_wps_export = visual_checks._export_wps_pdf
    original_pdfinfo = visual_checks._pdfinfo
    original_pages_text = visual_checks._pdf_pages_text
    original_render = visual_checks._render_samples
    try:
        visual_checks._export_pdf = lambda _docx, _visual_dir: str(word_pdf)
        visual_checks._export_wps_pdf = lambda _docx, _visual_dir: str(wps_pdf)
        visual_checks._pdfinfo = lambda _path: {"available": True, "pages": 1, "page_width_pt": 595.3, "page_height_pt": 841.9}
        visual_checks._pdf_pages_text = lambda _pdf, _visual_dir: ["same visible text"]

        def fake_render_samples(pdf, visual_dir, pages):
            sample_dir = Path(visual_dir) / "samples"
            sample_dir.mkdir(parents=True, exist_ok=True)
            rendered = []
            for page in pages:
                img = Image.new("L", (80, 80), color=255)
                pixels = img.load()
                if Path(pdf).name == "wps.pdf":
                    for x in range(40, 80):
                        for y in range(80):
                            pixels[x, y] = 0
                else:
                    for x in range(40):
                        for y in range(80):
                            pixels[x, y] = 0
                path = sample_dir / f"page_{page:03d}.png"
                img.save(path)
                rendered.append(str(path))
            return rendered

        visual_checks._render_samples = fake_render_samples
        report = visual_checks.check_visual(str(work), output_docx_name="final.docx", require_wps=False, render_all_pages=False)
    finally:
        visual_checks._export_pdf = original_export
        visual_checks._export_wps_pdf = original_wps_export
        visual_checks._pdfinfo = original_pdfinfo
        visual_checks._pdf_pages_text = original_pages_text
        visual_checks._render_samples = original_render

    codes = [item["code"] for item in report["issues"]]
    assert_true(report["passed"] is False, f"WPS sample-image mismatch should fail visual QA even when pages and text match: {report}")
    assert_true("WPS_SAMPLE_IMAGE_MISMATCH" in codes, f"WPS sample-image mismatch was not reported: {report}")
    assert_true(report["counts"].get("wps_sample_images") == 1, f"WPS sample image count was not recorded: {report}")
    assert_true(report["counts"].get("wps_sample_mismatches") == [1], f"WPS sample mismatch page was not recorded: {report}")
    action = report.get("next_action", "")
    assert_true("WPS" in action and ("样张" in action or "PNG" in action or "画面" in action) and "重跑 visual QA" in action, f"WPS sample-image mismatch lacked a concrete next action: {report}")


@case
def visual_qa_sanitizes_issue_details() -> None:
    import qa_visual

    work = new_workdir("visual_issue_privacy")
    report = qa_visual.check_visual(str(work), output_docx_name="missing.docx", project_root=str(work))
    text = json.dumps(report.get("issues") or [], ensure_ascii=False)
    assert_true(str(work) not in text and str(work).replace("\\", "/") not in text, "visual QA issue leaked output path")
    assert_true("<PROJECT>/missing.docx" in text, f"visual QA issue detail was not sanitized: {text}")


@case
def sample_pages_empty_when_page_count_unknown() -> None:
    import qa_visual

    assert_true(qa_visual._sample_pages(0, []) == [], "sample pages should be empty for unknown page count")


@case
def run_pipeline_missing_inputs_returns_nonzero() -> None:
    root = PIPELINE_DIR.parents[2]
    result = subprocess.run(
        [
            sys.executable,
            str(root / "run_pipeline.py"),
            "--template",
            "__missing_template__.docx",
            "--content",
            "__missing_content__.docx",
        ],
        cwd=str(root),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=60,
    )
    assert_true(result.returncode != 0, "run_pipeline returned success for missing inputs")


@case
def run_pipeline_help_localizes_agent_options() -> None:
    root = PIPELINE_DIR.parents[2]
    result = subprocess.run(
        [sys.executable, str(root / "run_pipeline.py"), "--help"],
        cwd=str(root),
        env={**os.environ, "PYTHONIOENCODING": "utf-8"},
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=60,
    )
    assert_true(result.returncode == 0, f"run_pipeline --help failed: {result.stderr}")
    assert_true("Agent 自动入口" in result.stdout and "自动修复闭环" in result.stdout, "help text lost novice-friendly Chinese option descriptions")
    assert_true("Agent-first mode" not in result.stdout and "Run a bounded" not in result.stdout, "help text still exposes old English option descriptions")


@case
def content_parser_cli_writes_outputs_outside_inputs() -> None:
    work = new_workdir("content_parser_cli_output")
    inputs = work / "Inputs"
    out_dir = work / "Outputs" / "content_cli"
    inputs.mkdir()
    docx = inputs / "paper.docx"
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Body paragraph.")
    doc.save(docx)

    result = subprocess.run(
        [
            sys.executable,
            str(PIPELINE_DIR / "content_parser.py"),
            str(docx),
            "--output-dir",
            str(out_dir),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=60,
    )
    assert_true(result.returncode == 0, f"content_parser CLI failed: {result.stdout}\n{result.stderr}")
    assert_true((out_dir / "paper_content.json").exists(), "content_parser CLI did not write JSON to output dir")
    assert_true((out_dir / "paper" / "figures").exists(), "content_parser CLI did not place figures under output dir")
    assert_true(not (inputs / "paper_content.json").exists(), "content_parser CLI wrote JSON beside input")
    assert_true(not (inputs / "paper" / "figures").exists(), "content_parser CLI wrote figures under Inputs")


@case
def content_parser_default_extract_writes_outputs_outside_inputs() -> None:
    from content_parser import extract as extract_docx_content

    work = new_workdir("content_parser_default_output")
    inputs = work / "Inputs"
    inputs.mkdir()
    docx = inputs / "paper.docx"
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Body paragraph.")
    doc.save(docx)

    old_cwd = os.getcwd()
    try:
        os.chdir(work)
        content = extract_docx_content(str(docx))
    finally:
        os.chdir(old_cwd)

    images_dir = Path(content["_meta"]["images_dir"])
    assert_true("Outputs" in images_dir.parts and "_content_parser_extract" in images_dir.parts, f"default images dir is unsafe: {images_dir}")
    assert_true(not (inputs / "paper" / "figures").exists(), "default content extraction wrote figures under Inputs")


@case
def extractors_reject_source_output_dirs() -> None:
    from content_parser import extract as extract_docx_content
    from format_extractor import extract as extract_format
    from md_parser import extract_content as extract_md_content

    work = new_workdir("extractor_output_dir_guard")
    inputs = work / "Inputs"
    templates = work / "Templates"
    inputs.mkdir()
    templates.mkdir()

    content_docx = inputs / "paper.docx"
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Body paragraph.")
    doc.save(content_docx)

    template_docx = templates / "template.docx"
    template = Document()
    template.add_paragraph("Template body.")
    template.save(template_docx)

    md = inputs / "paper.md"
    md.write_text("# Title\n\nBody paragraph.\n", encoding="utf-8")

    def expect_unsafe(call):
        try:
            call()
        except ValueError as exc:
            assert_true("Unsafe output_dir" in str(exc), f"unexpected safety error: {exc}")
            return
        raise AssertionError("extractor accepted a source output directory")

    expect_unsafe(lambda: extract_docx_content(str(content_docx), output_dir=str(inputs)))
    expect_unsafe(lambda: extract_docx_content(str(content_docx), output_dir=str(inputs / "Generated")))
    expect_unsafe(lambda: extract_format(str(template_docx), output_dir=str(templates)))
    expect_unsafe(lambda: extract_md_content(str(md), output_dir=str(inputs)))


@case
def extractor_cli_unsafe_output_shows_next_step() -> None:
    work = new_workdir("extractor_cli_unsafe_output")
    inputs = work / "Inputs"
    inputs.mkdir()
    docx = inputs / "paper.docx"
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Body paragraph.")
    doc.save(docx)

    result = subprocess.run(
        [
            sys.executable,
            str(PIPELINE_DIR / "content_parser.py"),
            str(docx),
            "--output-dir",
            str(inputs),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=60,
    )
    assert_true(result.returncode == 2, f"unsafe output dir should fail clearly: {result.stdout}\n{result.stderr}")
    assert_true("[NEXT]" in result.stdout and "Outputs/" in result.stdout, f"unsafe output dir lacked next step: {result.stdout}")

    md = inputs / "paper.md"
    md.write_text("# Title\n\nBody paragraph.\n", encoding="utf-8")
    md_result = subprocess.run(
        [
            sys.executable,
            str(PIPELINE_DIR / "md_parser.py"),
            str(md),
            "--output-dir",
            str(inputs),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=60,
    )
    assert_true(md_result.returncode == 2, f"unsafe MD output dir should fail clearly: {md_result.stdout}\n{md_result.stderr}")
    assert_true("[NEXT]" in md_result.stdout and "Outputs/" in md_result.stdout, f"unsafe MD output dir lacked next step: {md_result.stdout}")
    assert_true(not (inputs / "paper_format.json").exists(), "unsafe MD output dir wrote format JSON before failing")


@case
def format_extractor_cli_writes_outputs_outside_templates() -> None:
    work = new_workdir("format_extractor_cli_output")
    templates = work / "Templates"
    templates.mkdir()
    docx = templates / "template.docx"
    doc = Document()
    doc.add_paragraph("Template heading")
    doc.add_paragraph("Template body.")
    doc.save(docx)

    result = subprocess.run(
        [sys.executable, str(PIPELINE_DIR / "format_extractor.py"), str(docx)],
        cwd=str(work),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=60,
    )
    out_dir = work / "Outputs" / "_format_extractor_cli"
    assert_true(result.returncode == 0, f"format_extractor CLI failed: {result.stdout}\n{result.stderr}")
    assert_true((out_dir / "template_format.json").exists(), "format_extractor CLI did not write JSON to Outputs")
    assert_true((out_dir / "template_格式提取.md").exists(), "format_extractor CLI did not write MD report to Outputs")
    assert_true(not (templates / "template_format.json").exists(), "format_extractor CLI wrote JSON beside template")
    assert_true(not (templates / "template_assets").exists(), "format_extractor CLI wrote assets beside template")


@case
def format_extractor_default_assets_stay_outside_templates() -> None:
    from format_extractor import extract as extract_format

    work = new_workdir("format_extractor_default_assets")
    templates = work / "Templates"
    templates.mkdir()
    docx = templates / "template.docx"
    doc = Document()
    doc.add_paragraph("Template heading")
    doc.save(docx)

    old_cwd = os.getcwd()
    try:
        os.chdir(work)
        fmt, _ = extract_format(str(docx))
    finally:
        os.chdir(old_cwd)

    assets_dir = Path(fmt["_meta"]["assets_dir"])
    assert_true("Outputs" in assets_dir.parts and "_format_extractor_extract" in assets_dir.parts, f"default assets dir is unsafe: {assets_dir}")
    assert_true(not (templates / "template_assets").exists(), "default format extraction wrote assets beside template")


@case
def md_parser_cli_writes_outputs_outside_inputs() -> None:
    work = new_workdir("md_parser_cli_output")
    inputs = work / "Inputs"
    inputs.mkdir()
    md = inputs / "paper.md"
    md.write_text("# 格式说明\n\n正文：宋体，小四号。\n\n# 论文标题\n\n正文段落。\n", encoding="utf-8")

    result = subprocess.run(
        [sys.executable, str(PIPELINE_DIR / "md_parser.py"), str(md)],
        cwd=str(work),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=60,
    )
    out_dir = work / "Outputs" / "_md_parser_cli"
    assert_true(result.returncode == 0, f"md_parser CLI failed: {result.stdout}\n{result.stderr}")
    assert_true((out_dir / "paper_format.json").exists(), "md_parser CLI did not write format JSON to Outputs")
    assert_true((out_dir / "paper_content.json").exists(), "md_parser CLI did not write content JSON to Outputs")
    assert_true(not (inputs / "paper_format.json").exists(), "md_parser CLI wrote format JSON beside input")
    assert_true(not (inputs / "paper_content.json").exists(), "md_parser CLI wrote content JSON beside input")


@case
def public_template_download_requires_https_and_verifies_sha256() -> None:
    from public_template_suite_modules.storage import download_template, safe_download_url

    work = new_workdir("public_template_download_guard")
    local = work / "template.docx"
    local.write_bytes(b"PK" + b"x" * 3000)
    digest = hashlib.sha256(local.read_bytes()).hexdigest()

    assert_true(safe_download_url("https://example.com/template.docx"), "HTTPS template URL should be accepted")
    assert_true(not safe_download_url("http://example.com/template.docx"), "HTTP template URL should be rejected")
    assert_true(download_template({"id": "local", "file": str(local), "sha256": digest}) == local, "matching local sha256 should pass")

    try:
        download_template({"id": "local", "file": str(local), "sha256": "0" * 64})
    except RuntimeError as exc:
        assert_true("sha256 mismatch" in str(exc), f"unexpected checksum error: {exc}")
    else:
        raise AssertionError("template download accepted a sha256 mismatch")


@case
def public_template_visual_golden_baseline_is_opt_in() -> None:
    from public_template_suite import DEFAULT_GOLDEN_DIR, resolve_golden_dir

    assert_true(resolve_golden_dir(None, update_golden=False) is None, "--visual should not compare golden baselines by default")
    assert_true(resolve_golden_dir("", update_golden=False) is None, "empty golden-dir should disable golden comparison")
    assert_true(
        resolve_golden_dir(None, update_golden=True) == DEFAULT_GOLDEN_DIR,
        "--update-golden without explicit dir should use the default baseline directory",
    )
    custom = new_workdir("public_template_custom_golden") / "Golden"
    assert_true(resolve_golden_dir(str(custom), update_golden=False) == custom, "explicit golden-dir should be preserved")


@case
def qa_checker_cli_failure_returns_nonzero() -> None:
    work = new_workdir("qa_cli_nonzero")
    content = base_content([])
    content["_meta"]["missing_images"] = [{"source": "missing.png", "reason": "not_found"}]
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    result = subprocess.run(
        [
            sys.executable,
            str(PIPELINE_DIR / "qa_checker.py"),
            str(work),
            "--mode",
            "developer",
            "--docx",
            "out.docx",
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=60,
    )
    assert_true(result.returncode != 0, "qa_checker CLI returned success for a failed report")


@case
def qa_missing_image_detail_is_sanitized() -> None:
    work = new_workdir("qa_missing_image_privacy")
    private_path = str(work / "private" / "missing.png")
    content = base_content([])
    content["_meta"]["missing_images"] = [{"source": private_path, "reason": "not_found"}]
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    detail = "\n".join(str(item.get("detail") or "") for item in report["issues"])
    assert_true(private_path not in detail, "QA leaked an absolute missing-image path")
    assert_true("<TEMP>" in detail or "<ABS_PATH>" in detail or "<PROJECT>" in detail, "QA missing-image detail was not sanitized")

