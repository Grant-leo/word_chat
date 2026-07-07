"""Template profiling and format extraction regression cases."""
from __future__ import annotations

import json
import os

from docx import Document
from format_extractor import extract as extract_docx_format
from qa_checker_modules.format_phase import run_format_checks
from regression_suite_modules.generated_pdf import poppler_available, write_blank_pdf, write_text_pdf
from regression_suite_modules.harness import assert_true, base_format, case, new_workdir
from template_profiler import profile_format


@case
def template_profile_sanitizes_private_source() -> None:
    fmt = base_format(source="private_school_template.docx")
    fmt["_meta"]["assets_dir"] = r"E:\private\Templates\private_school_template_assets"
    profile = profile_format(fmt, project_root=r"E:\private")
    text = json.dumps(profile, ensure_ascii=False)
    assert_true("private_school_template" not in text, "template profile leaked source filename")
    assert_true(profile["source"]["source_ext"] == ".docx", "template profile lost source extension")
    assert_true(profile["source"]["has_assets_dir"] is True, "template profile lost assets flag")


@case
def template_profile_detects_chinese_formula_and_reference_rules() -> None:
    fmt = base_format()
    fmt["paragraphs"].append({"text": "公式应居中书写，参考文献按 GB/T 7714 排列。"})
    profile = profile_format(fmt)
    assert_true(profile["risk_flags"]["mentions_formula_rules"] is True, "Chinese formula rules were not detected")
    assert_true(profile["risk_flags"]["mentions_reference_rules"] is True, "Chinese reference rules were not detected")


@case
def format_extractor_stops_cover_before_spaced_abstract_heading() -> None:
    work = new_workdir("format_cover_stop_abstract")
    docx = work / "cover_stop.docx"
    doc = Document()
    doc.add_paragraph("Cover title")
    doc.add_paragraph("摘  要")
    doc.add_paragraph("Template abstract sample paragraph should not be replayed as cover.")
    doc.save(docx)

    fmt, _ = extract_docx_format(str(docx), output_dir=str(work))
    cover_text = "\n".join(
        "".join(run.get("t", "") for run in el.get("r", []))
        for el in fmt.get("cover") or []
        if isinstance(el, dict)
    )
    assert_true("Cover title" in cover_text, "cover text before abstract was not extracted")
    assert_true("摘" not in cover_text and "Template abstract sample" not in cover_text, "abstract page leaked into cover extraction")


@case
def pdf_instruction_template_extracts_profile() -> None:
    if not poppler_available():
        return
    work = new_workdir("pdf_instruction_template")
    pdf = work / "requirements.pdf"
    write_text_pdf(
        pdf,
        [
            ("Format requirements", 16, 72, 780),
            ("Page A4 margins top 2.5 cm bottom 2.5 cm left 3.0 cm right 2.5 cm", 11, 72, 742),
            ("Body font Times New Roman 12 pt justified line spacing 1.5", 11, 72, 718),
            ("Heading 1 font SimHei 16 pt bold centered", 11, 72, 694),
            ("Heading 2 font SimHei 14 pt bold left", 11, 72, 670),
            ("References font Times New Roman 12 pt", 11, 72, 646),
        ],
    )

    fmt, md = extract_docx_format(str(pdf), output_dir=str(work))
    pdf_meta = fmt.get("_meta", {}).get("pdf_template") or {}
    profile = profile_format(fmt)
    assert_true(pdf_meta.get("type") == "instruction_pdf", f"unexpected pdf type: {pdf_meta}")
    assert_true(pdf_meta.get("errors") == [], f"instruction PDF should not report errors: {pdf_meta}")
    assert_true(profile["source"]["source_ext"] == ".pdf", "template profile lost PDF source extension")
    assert_true(profile["capabilities"]["has_heading_styles"], "instruction PDF should provide heading style profiles")
    assert_true("PDF 模板格式提取" in md, "PDF extraction report title missing")


@case
def pdf_sparse_instruction_template_warns_limited_confidence() -> None:
    if not poppler_available():
        return
    work = new_workdir("pdf_sparse_instruction_template")
    pdf = work / "sparse_requirements.pdf"
    write_text_pdf(
        pdf,
        [
            ("Format requirements", 16, 72, 780),
            ("Page A4 margin left 3.0 cm right 2.5 cm", 11, 72, 742),
            ("Body font Times New Roman 12 pt justified line spacing 1.5", 11, 72, 718),
        ],
    )

    fmt, _ = extract_docx_format(str(pdf), output_dir=str(work))
    pdf_meta = fmt.get("_meta", {}).get("pdf_template") or {}
    warnings = pdf_meta.get("warnings") or []
    assert_true(
        any(str(item).startswith("PDF_TEMPLATE_INSTRUCTION_INCOMPLETE") for item in warnings),
        f"sparse instruction PDF should warn about incomplete rules: {warnings}",
    )


@case
def pdf_visual_template_extracts_lines_and_risk_profile() -> None:
    if not poppler_available():
        return
    work = new_workdir("pdf_visual_template")
    pdf = work / "visual_sample.pdf"
    write_text_pdf(
        pdf,
        [
            ("Synthetic Thesis Title", 20, 180, 780),
            ("Abstract", 16, 260, 735),
            ("This paper studies robust document generation from template samples.", 11, 72, 700),
            ("1 Introduction", 15, 72, 650),
            ("Template adaptation requires careful format extraction and QA.", 11, 72, 622),
            ("Figure 1 System architecture", 10, 210, 585),
            ("References", 15, 72, 540),
            ("[1] Doe J. Synthetic reference for regression testing.", 11, 72, 512),
        ],
    )

    fmt, _ = extract_docx_format(str(pdf), output_dir=str(work))
    pdf_meta = fmt.get("_meta", {}).get("pdf_template") or {}
    profile = profile_format(fmt)
    assert_true(pdf_meta.get("type") == "visual_sample_pdf", f"unexpected pdf type: {pdf_meta}")
    assert_true(len(fmt.get("paragraphs") or []) >= 6, "visual PDF lines were not converted into paragraphs")
    assert_true(profile["risk_flags"]["pdf_template_limited_confidence"] is True, "visual PDF risk flag missing")
    assert_true(profile["capabilities"]["has_page_geometry"] is True, "visual PDF page geometry missing")


@case
def pdf_scanned_template_surfaces_qa_error() -> None:
    if not poppler_available():
        return
    work = new_workdir("pdf_scanned_template")
    pdf = work / "blank_scan.pdf"
    write_blank_pdf(pdf)
    fmt, _ = extract_docx_format(str(pdf), output_dir=str(work))
    format_path = work / "format.json"
    format_path.write_text(json.dumps(fmt, ensure_ascii=False), encoding="utf-8")

    issues = []

    def add(code, level, message, detail=""):
        issues.append({"code": code, "level": level, "message": message, "detail": detail})

    run_format_checks({"format": str(format_path)}, {}, add)
    codes = {issue["code"] for issue in issues}
    assert_true("PDF_TEMPLATE_UNSUPPORTED" in codes, f"scanned PDF did not surface QA error: {issues}")


@case
def pdf_template_extractor_skips_broken_pdfinfo_path_shim() -> None:
    if os.name != "nt":
        return
    work = new_workdir("pdf_template_broken_pdfinfo_path")
    pdf = work / "blank_scan.pdf"
    write_blank_pdf(pdf)

    bad = work / "bad"
    good = work / "good"
    bad.mkdir()
    good.mkdir()
    (bad / "pdfinfo.cmd").write_text(
        "@echo off\r\necho The system cannot find the path specified. 1>&2\r\nexit /b 3\r\n",
        encoding="utf-8",
    )
    (good / "pdfinfo.cmd").write_text(
        "@echo off\r\necho Pages:          1\r\necho Page size:      595 x 842 pts\r\nexit /b 0\r\n",
        encoding="utf-8",
    )
    (good / "pdftotext.cmd").write_text(
        "@echo off\r\n"
        "echo %* | findstr /C:\"-bbox-layout\" >nul\r\n"
        "if %errorlevel%==0 (\r\n"
        "  set last=\r\n"
        "  for %%A in (%*) do set last=%%~A\r\n"
        "  > \"%last%\" echo ^<html^>^<body^>^<page width=\"595\" height=\"842\"^>^</page^>^</body^>^</html^>\r\n"
        ")\r\n"
        "exit /b 0\r\n",
        encoding="utf-8",
    )

    old_path = os.environ.get("PATH", "")
    os.environ["PATH"] = str(bad) + os.pathsep + str(good) + os.pathsep + old_path
    try:
        fmt, _ = extract_docx_format(str(pdf), output_dir=str(work))
    finally:
        os.environ["PATH"] = old_path

    format_path = work / "format.json"
    format_path.write_text(json.dumps(fmt, ensure_ascii=False), encoding="utf-8")

    issues = []

    def add(code, level, message, detail=""):
        issues.append({"code": code, "level": level, "message": message, "detail": detail})

    run_format_checks({"format": str(format_path)}, {}, add)
    codes = {issue["code"] for issue in issues}
    assert_true("PDF_TEMPLATE_UNSUPPORTED" in codes, f"blank PDF should remain unsupported after pdfinfo fallback: {issues}")
    assert_true("PDF_TEMPLATE_READ_FAILED" not in codes, f"broken first pdfinfo shim should not force read-failed: {issues}")

