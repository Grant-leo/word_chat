"""Structural QA and manifest regression cases."""
from __future__ import annotations

import json
from xml.etree import ElementTree as ET

from docx import Document

from qa_checker import check_output
from qa_checker_modules.content_samples import _content_toc_pollution_samples
from qa_checker_modules.content_metrics import (
    _content_text_chars,
    _count_content_formulas,
    _count_content_images,
    _count_content_tables,
)
from qa_checker_modules.registry import OWNER_BY_CODE
from qa_checker_modules.report_phase import build_report
from qa_checker_modules.repair import build_repair_plan
from qa_checker_modules.repair_guides import REPAIR_GUIDES
from qa_checker_modules.reports import repair_plan_to_markdown
from qa_conformance_modules.reports import build_report as build_conformance_report
from qa_conformance_modules.content_checks import _expected_paragraphs, _find_body_start_index, _find_para_by_text
from qa_conformance_modules.requirements import build_requirements
from qa_visual_modules.checks import check_visual

from regression_suite_modules.generated_docx import run_generated_case
from regression_suite_modules.harness import (
    assert_true,
    base_content,
    base_format,
    case,
    new_workdir,
    write_json,
    write_sample_png,
)

@case
def table_manifest_matches_structured_body_tables() -> None:
    content = base_content(
        [
            {"role": "table", "table_rows": [["A", "B"], ["1", "2"]]},
        ],
        meta_tables=99,
    )
    result = run_generated_case("table_manifest", content)
    assert_true(result["manifest"]["counts"]["content_tables_rendered"] == 1, "body table was not counted")
    assert_true(result["report"]["counts"]["content_tables"] == 1, "QA used raw doc.tables instead of structured tables")
    assert_true(not result["report"]["issues"], f"unexpected QA issues: {result['report']['issues']}")


@case
def qa_manifest_detects_missing_table_render() -> None:
    content = base_content([{"role": "table", "table_rows": [["A"], ["B"]]}])
    result = run_generated_case("qa_missing_table", content)
    manifest_path = result["work"] / "build_manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["counts"]["content_tables_rendered"] = 0
    write_json(manifest_path, manifest)
    report = check_output(str(result["work"]), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("TABLE_COUNT_MISMATCH" in codes, "QA did not trust manifest for table mismatch")


@case
def qa_manifest_detects_missing_footnote_render() -> None:
    work = new_workdir("qa_missing_footnote")
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("A sentence with a source footnote.")
    doc.save(work / "out.docx")
    content = base_content(
        [
            {
                "role": "rich_text",
                "text": "A sentence with a source footnote.",
                "runs": [
                    {"type": "text", "text": "A sentence with a source footnote"},
                    {"type": "note_ref", "note_type": "footnote", "source_id": "2", "text": "Missing rendered footnote."},
                    {"type": "text", "text": "."},
                ],
            }
        ]
    )
    content["_meta"]["footnote_references_extracted"] = 1
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {"footnote_references_rendered": 0}})
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    (work / "build_generated.py").write_text("# synthetic generated script\n", encoding="utf-8")

    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("FOOTNOTE_RENDER_COUNT_MISMATCH" in codes, f"QA did not report missing footnote render: {report['issues']}")


@case
def qa_flags_generated_script_unicode_escape_decoding() -> None:
    work = new_workdir("qa_generated_unicode_decode_guard")
    text = "中文字符保持原样：编码测试。"
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "import codecs\n"
        "text = codecs.decode('中文', 'unicode_escape')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" in codes, f"QA did not flag unsafe unicode decoding: {report}")
    assert_true(report["passed"] is False, f"unsafe unicode decoding should block structural QA: {report}")
    action = f"{report.get('next_action')}\n{json.dumps(report.get('repair_plan') or {}, ensure_ascii=False)}"
    assert_true("codecs.decode" in action and "build_generated.py" in action and "重跑" in action,
                f"unsafe unicode decode guidance was not beginner-facing enough: {action}")


@case
def qa_flags_generated_script_unicode_escape_decode_aliases() -> None:
    text = "中文字符保持原样：编码测试。"
    scripts = {
        "qa_generated_unicode_module_alias": (
            "import codecs as text_codecs\n"
            "text = text_codecs.decode('中文', 'unicode_escape')\n"
        ),
        "qa_generated_unicode_function_alias": (
            "from codecs import decode as decode_text\n"
            "text = decode_text('中文', 'unicode_escape')\n"
        ),
    }

    for name, script in scripts.items():
        work = new_workdir(name)
        doc = Document()
        doc.add_paragraph("Synthetic Thesis")
        doc.add_paragraph("1 Introduction")
        doc.add_paragraph(text)
        doc.save(work / "out.docx")
        write_json(work / "content.json", base_content([text]))
        write_json(work / "format.json", base_format())
        write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
        write_json(work / "workflow_mode.json", {"mode": "user"})
        (work / "build_generated.py").write_text(script, encoding="utf-8")

        report = check_output(str(work), mode="user", output_docx_name="out.docx")
        codes = [item["code"] for item in report["issues"]]
        assert_true("GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" in codes, f"QA did not flag unsafe unicode decode alias {name}: {report}")


@case
def qa_flags_generated_script_unicode_escape_normalized_codec_names() -> None:
    text = "中文字符保持原样：编码测试。"
    scripts = {
        "qa_generated_unicode_hyphen_codec": (
            "import codecs\n"
            "text = codecs.decode('中文', 'unicode-escape')\n"
        ),
        "qa_generated_unicode_space_case_codec": (
            "text = b'\\\\u4e2d\\\\u6587'.decode('Unicode Escape')\n"
        ),
        "qa_generated_raw_unicode_space_codec": (
            "import codecs\n"
            "text = codecs.decode('中文', encoding='raw unicode escape')\n"
        ),
    }

    for name, script in scripts.items():
        work = new_workdir(name)
        doc = Document()
        doc.add_paragraph("Synthetic Thesis")
        doc.add_paragraph("1 Introduction")
        doc.add_paragraph(text)
        doc.save(work / "out.docx")
        write_json(work / "content.json", base_content([text]))
        write_json(work / "format.json", base_format())
        write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
        write_json(work / "workflow_mode.json", {"mode": "user"})
        (work / "build_generated.py").write_text(script, encoding="utf-8")

        report = check_output(str(work), mode="user", output_docx_name="out.docx")
        codes = [item["code"] for item in report["issues"]]
        assert_true(
            "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" in codes,
            f"QA did not flag normalized unsafe unicode codec {name}: {report}",
        )


@case
def qa_flags_generated_script_unicode_escape_decoder_factories() -> None:
    text = "中文字符保持原样：编码测试。"
    scripts = {
        "qa_generated_unicode_getdecoder": (
            "import codecs\n"
            "decoder = codecs.getdecoder('unicode-escape')\n"
            "text = decoder(b'\\\\u4e2d\\\\u6587')[0]\n"
        ),
        "qa_generated_unicode_lookup_decode": (
            "import codecs\n"
            "text = codecs.lookup('Unicode Escape').decode(b'\\\\u4e2d\\\\u6587')[0]\n"
        ),
        "qa_generated_unicode_lookup_alias_decode": (
            "from codecs import lookup as lookup_codec\n"
            "text = lookup_codec('unicode-escape').decode(b'\\\\u4e2d\\\\u6587')[0]\n"
        ),
        "qa_generated_unicode_getdecoder_attribute_alias": (
            "import codecs\n"
            "decode_factory = codecs.getdecoder\n"
            "text = decode_factory('unicode_escape')(b'\\\\u4e2d\\\\u6587')[0]\n"
        ),
        "qa_generated_unicode_getdecoder_alias": (
            "from codecs import getdecoder as decode_factory\n"
            "text = decode_factory('raw unicode escape')(b'\\\\u4e2d\\\\u6587')[0]\n"
        ),
    }

    for name, script in scripts.items():
        work = new_workdir(name)
        doc = Document()
        doc.add_paragraph("Synthetic Thesis")
        doc.add_paragraph("1 Introduction")
        doc.add_paragraph(text)
        doc.save(work / "out.docx")
        write_json(work / "content.json", base_content([text]))
        write_json(work / "format.json", base_format())
        write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
        write_json(work / "workflow_mode.json", {"mode": "user"})
        (work / "build_generated.py").write_text(script, encoding="utf-8")

        report = check_output(str(work), mode="user", output_docx_name="out.docx")
        codes = [item["code"] for item in report["issues"]]
        assert_true(
            "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" in codes,
            f"QA did not flag unsafe unicode decoder factory {name}: {report}",
        )
        action = f"{report.get('next_action')}\n{json.dumps(report.get('repair_plan') or {}, ensure_ascii=False)}"
        assert_true(
            "getdecoder" in action or "lookup" in action,
            f"decoder-factory guidance did not name the unsafe factory route: {action}",
        )


@case
def qa_flags_generated_script_unicode_escape_static_codec_expressions() -> None:
    text = "中文字符保持原样：编码测试。"
    scripts = {
        "qa_generated_unicode_codec_variable": (
            "import codecs\n"
            "codec = 'unicode_escape'\n"
            "text = codecs.decode('中文', codec)\n"
        ),
        "qa_generated_unicode_codec_concat": (
            "import codecs\n"
            "codec = 'unicode' + '_escape'\n"
            "text = codecs.decode('中文', codec)\n"
        ),
        "qa_generated_unicode_codec_keyword_variable": (
            "codec = 'raw' + '_unicode_escape'\n"
            "text = b'abc'.decode(encoding=codec)\n"
        ),
        "qa_generated_unicode_factory_codec_variable": (
            "import codecs\n"
            "codec = 'Unicode Escape'\n"
            "decoder = codecs.getdecoder(codec)\n"
            "text = decoder(b'\\\\u4e2d\\\\u6587')[0]\n"
        ),
        "qa_generated_unicode_lookup_direct_concat": (
            "import codecs\n"
            "text = codecs.lookup('unicode' + '-escape').decode(b'\\\\u4e2d\\\\u6587')[0]\n"
        ),
    }

    for name, script in scripts.items():
        work = new_workdir(name)
        doc = Document()
        doc.add_paragraph("Synthetic Thesis")
        doc.add_paragraph("1 Introduction")
        doc.add_paragraph(text)
        doc.save(work / "out.docx")
        write_json(work / "content.json", base_content([text]))
        write_json(work / "format.json", base_format())
        write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
        write_json(work / "workflow_mode.json", {"mode": "user"})
        (work / "build_generated.py").write_text(script, encoding="utf-8")

        report = check_output(str(work), mode="user", output_docx_name="out.docx")
        codes = [item["code"] for item in report["issues"]]
        assert_true(
            "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" in codes,
            f"QA did not flag unsafe unicode static codec expression {name}: {report}",
        )


@case
def qa_flags_generated_script_unicode_escape_forwarded_to_wrapper() -> None:
    text = "中文字符保持原样：编码测试。"
    work = new_workdir("qa_generated_unicode_wrapper_forward")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "import codecs\n"
        "\n"
        "def decode_text(value, encoding):\n"
        "    return codecs.decode(value, encoding)\n"
        "\n"
        "text = decode_text('中文字符保持原样：编码测试。', 'unicode_escape')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" in codes,
        f"QA did not flag unsafe unicode codec forwarded through wrapper: {report}",
    )


@case
def qa_flags_generated_script_unicode_escape_decode_alias_assignment() -> None:
    text = "中文字符保持原样：编码测试。"
    work = new_workdir("qa_generated_unicode_decode_alias_assignment")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "import codecs\n"
        "decode_text = codecs.decode\n"
        "text = decode_text('中文字符保持原样：编码测试。', 'unicode_escape')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" in codes,
        f"QA did not flag unsafe unicode codec passed through decode function alias assignment: {report}",
    )


@case
def qa_flags_generated_script_unicode_escape_higher_order_decode() -> None:
    text = "中文字符保持原样：编码测试。"
    work = new_workdir("qa_generated_unicode_higher_order_decode")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "import codecs\n"
        "\n"
        "def apply_decoder(decoder, value, encoding):\n"
        "    return decoder(value, encoding)\n"
        "\n"
        "text = apply_decoder(codecs.decode, '中文字符保持原样：编码测试。', 'unicode_escape')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" in codes,
        f"QA did not flag unsafe unicode codec passed with codecs.decode function object: {report}",
    )


@case
def qa_flags_generated_script_general_codecs_decode_text_reencoding() -> None:
    text = "中文字符保持原样：编码测试。"
    scripts = {
        "qa_generated_codecs_decode_wrong_charset": (
            "import codecs\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = codecs.decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_codecs_decode_alias_wrong_charset": (
            "import codecs\n"
            "decode_text = codecs.decode\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = decode_text(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_codecs_decode_higher_order_wrong_charset": (
            "import codecs\n"
            "\n"
            "def apply_decoder(decoder, value, encoding):\n"
            "    return decoder(value, encoding, errors='ignore')\n"
            "\n"
            "text = apply_decoder(codecs.decode, '中文字符保持原样：编码测试。'.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_returned_higher_order_wrapper_codecs_decode_wrong_charset": (
            "import codecs\n"
            "\n"
            "def build_apply():\n"
            "    def apply_decoder(decoder, value, encoding):\n"
            "        return decoder(value, encoding, errors='ignore')\n"
            "    return apply_decoder\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_apply()(codecs.decode, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_dunder_import_codecs_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = __import__('codecs').decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_importlib_codecs_decode_wrong_charset": (
            "import importlib\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = importlib.import_module('codecs').decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_codecs_dict_decode_wrong_charset": (
            "import codecs\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = codecs.__dict__['decode'](text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_vars_codecs_decode_wrong_charset": (
            "import codecs\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = vars(codecs)['decode'](text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_globals_codecs_decode_wrong_charset": (
            "import codecs\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = globals()['codecs'].decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_literal_dict_decode_wrong_charset": (
            "import codecs\n"
            "funcs = {'decode': codecs.decode}\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = funcs['decode'](text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_literal_dict_module_decode_wrong_charset": (
            "import codecs\n"
            "modules = {'text_codecs': codecs}\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = modules['text_codecs'].decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_literal_dict_decode_alias_wrong_charset": (
            "import codecs\n"
            "decode_text = codecs.decode\n"
            "funcs = {'decode': decode_text}\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = funcs['decode'](text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_literal_dict_module_alias_wrong_charset": (
            "import codecs\n"
            "text_codecs = codecs\n"
            "modules = {'text_codecs': text_codecs}\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = modules['text_codecs'].decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_list_codecs_decode_wrong_charset": (
            "import codecs\n"
            "routes = [codecs.decode]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = routes[0](text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_tuple_getattr_codecs_decode_wrong_charset": (
            "import codecs\n"
            "routes = (getattr(codecs, 'decode'),)\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = routes[0](text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_function_returns_codecs_decode_wrong_charset": (
            "import codecs\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "def get_decoder():\n"
            "    return codecs.decode\n"
            "mojibake = get_decoder()(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_function_returns_getattr_codecs_decode_wrong_charset": (
            "import codecs\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "def get_decoder():\n"
            "    return getattr(codecs, 'decode')\n"
            "mojibake = get_decoder()(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_function_return_passed_to_higher_order_codecs_decode_wrong_charset": (
            "import codecs\n"
            "\n"
            "def get_decoder():\n"
            "    return codecs.decode\n"
            "\n"
            "def apply_decoder(decoder, value, encoding):\n"
            "    return decoder(value, encoding, errors='ignore')\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_decoder(get_decoder(), text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_function_return_passed_to_returned_higher_order_codecs_decode_wrong_charset": (
            "import codecs\n"
            "\n"
            "def get_decoder():\n"
            "    return codecs.decode\n"
            "\n"
            "def build_apply():\n"
            "    def apply_decoder(decoder, value, encoding):\n"
            "        return decoder(value, encoding, errors='ignore')\n"
            "    return apply_decoder\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_apply()(get_decoder(), text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_module_param_codecs_decode_wrong_charset": (
            "import codecs\n"
            "\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.decode(value, encoding, errors='ignore')\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(codecs, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_getattr_module_param_codecs_decode_wrong_charset": (
            "import codecs\n"
            "\n"
            "def apply_module(module, value, encoding):\n"
            "    return getattr(module, 'decode')(value, encoding, errors='ignore')\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(codecs, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_local_getattr_module_decode_alias_wrong_charset": (
            "import codecs\n"
            "\n"
            "def apply_module(module, value, encoding):\n"
            "    decode = getattr(module, 'decode')\n"
            "    return decode(value, encoding, errors='ignore')\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(codecs, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_required_arg_function_returns_getattr_module_decode_wrong_charset": (
            "import codecs\n"
            "\n"
            "def pick_decode(module):\n"
            "    return getattr(module, 'decode')\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = pick_decode(codecs)(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_required_arg_closure_returns_getattr_module_decode_wrong_charset": (
            "import codecs\n"
            "\n"
            "def build_decoder(module):\n"
            "    def pick_decode():\n"
            "        return getattr(module, 'decode')\n"
            "    return pick_decode()\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_decoder(codecs)(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_required_arg_nested_closure_returns_getattr_module_decode_wrong_charset": (
            "import codecs\n"
            "\n"
            "def build_decoder(module):\n"
            "    def outer():\n"
            "        def pick_decode():\n"
            "            return getattr(module, 'decode')\n"
            "        return pick_decode()\n"
            "    return outer()\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_decoder(codecs)(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_module_attribute_param_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.module = codecs\n"
            "\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.decode(value, encoding, errors='ignore')\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(box.module, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_module_return_passed_to_higher_order_codecs_decode_wrong_charset": (
            "import codecs\n"
            "\n"
            "def get_module():\n"
            "    return codecs\n"
            "\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.decode(value, encoding, errors='ignore')\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(get_module(), text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_module_container_return_passed_to_higher_order_codecs_decode_wrong_charset": (
            "import codecs\n"
            "\n"
            "def get_modules():\n"
            "    return [codecs]\n"
            "\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.decode(value, encoding, errors='ignore')\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(get_modules()[0], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_module_attribute_container_passed_to_higher_order_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.modules = [codecs]\n"
            "\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.decode(value, encoding, errors='ignore')\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(box.modules[0], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_attribute_container_return_passed_to_higher_order_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.modules = [codecs]\n"
            "\n"
            "def get_modules():\n"
            "    return box.modules\n"
            "\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.decode(value, encoding, errors='ignore')\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(get_modules()[0], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_nested_module_container_passed_to_higher_order_codecs_decode_wrong_charset": (
            "import codecs\n"
            "config = {'modules': [codecs]}\n"
            "\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.decode(value, encoding, errors='ignore')\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(config['modules'][0], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_instance_method_module_return_passed_to_higher_order_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def get_module(self):\n"
            "        return codecs\n"
            "holder = Holder()\n"
            "\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.decode(value, encoding, errors='ignore')\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(holder.get_module(), text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_bound_method_from_module_return_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def get_module(self):\n"
            "        return codecs\n"
            "holder = Holder()\n"
            "decode_text = holder.get_module().decode\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = decode_text(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_function_returns_bound_method_from_module_return_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def get_module(self):\n"
            "        return codecs\n"
            "holder = Holder()\n"
            "def get_decode():\n"
            "    return holder.get_module().decode\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = get_decode()(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_operator_attrgetter_codecs_decode_wrong_charset": (
            "import codecs\n"
            "import operator\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = operator.attrgetter('decode')(codecs)(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_operator_attrgetter_codecs_decode_alias_wrong_charset": (
            "import codecs\n"
            "from operator import attrgetter\n"
            "decode_text = attrgetter('decode')(codecs)\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = decode_text(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_function_returns_operator_attrgetter_codecs_decode_wrong_charset": (
            "import codecs\n"
            "from operator import attrgetter\n"
            "def get_decode():\n"
            "    return attrgetter('decode')(codecs)\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = get_decode()(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_operator_methodcaller_codecs_decode_wrong_charset": (
            "import codecs\n"
            "import operator\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = operator.methodcaller('decode', text.encode('utf-8'), 'gbk', errors='ignore')(codecs)\n"
        ),
        "qa_generated_operator_methodcaller_codecs_decode_alias_wrong_charset": (
            "import codecs\n"
            "from operator import methodcaller\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "decode_text = methodcaller('decode', text.encode('utf-8'), 'gbk', errors='ignore')\n"
            "mojibake = decode_text(codecs)\n"
        ),
        "qa_generated_function_returns_operator_methodcaller_codecs_decode_wrong_charset": (
            "import codecs\n"
            "from operator import methodcaller\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "def get_decode():\n"
            "    return methodcaller('decode', text.encode('utf-8'), 'gbk', errors='ignore')\n"
            "mojibake = get_decode()(codecs)\n"
        ),
        "qa_generated_lambda_returns_codecs_decode_wrong_charset": (
            "import codecs\n"
            "get_decode = lambda: codecs.decode\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = get_decode()(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_lambda_returns_codecs_module_wrong_charset": (
            "import codecs\n"
            "get_module = lambda: codecs\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = get_module().decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_returned_module_param_codecs_decode_wrong_charset": (
            "import codecs\n"
            "\n"
            "def build_apply():\n"
            "    def apply_module(module, value, encoding):\n"
            "        return module.decode(value, encoding, errors='ignore')\n"
            "    return apply_module\n"
            "\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_apply()(codecs, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_object_attribute_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.decode = codecs.decode\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = box.decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_setattr_object_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "setattr(box, 'decode', codecs.decode)\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = box.decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_object_alias_attribute_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.decode = codecs.decode\n"
            "alias = box\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = alias.decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_namespace_attribute_codecs_decode_wrong_charset": (
            "import codecs\n"
            "from types import SimpleNamespace\n"
            "ns = SimpleNamespace(decode=codecs.decode)\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = ns.decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_class_attribute_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    decode = codecs.decode\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = Holder.decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_attribute_function_return_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.decode = codecs.decode\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "def get_decoder():\n"
            "    return box.decode\n"
            "mojibake = get_decoder()(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_instance_attribute_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.decode = codecs.decode\n"
            "holder = Holder()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = holder.decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_setattr_instance_attribute_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        setattr(self, 'decode', codecs.decode)\n"
            "holder = Holder()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = holder.decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_instance_alias_attribute_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.decode = codecs.decode\n"
            "holder = Holder()\n"
            "alias = holder\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = alias.decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_instance_attribute_function_return_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.decode = codecs.decode\n"
            "holder = Holder()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "def get_decoder():\n"
            "    return holder.decode\n"
            "mojibake = get_decoder()(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_class_alias_instance_attribute_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.decode = codecs.decode\n"
            "Alias = Holder\n"
            "holder = Alias()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = holder.decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_factory_instance_attribute_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.decode = codecs.decode\n"
            "def make_holder():\n"
            "    return Holder()\n"
            "holder = make_holder()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = holder.decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_class_alias_temp_instance_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.decode = codecs.decode\n"
            "Alias = Holder\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = Alias().decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_staticmethod_factory_instance_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.decode = codecs.decode\n"
            "    @staticmethod\n"
            "    def make():\n"
            "        return Holder()\n"
            "holder = Holder.make()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = holder.decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_classmethod_temp_instance_codecs_decode_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.decode = codecs.decode\n"
            "    @classmethod\n"
            "    def make(cls):\n"
            "        return cls()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = Holder.make().decode(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_closure_wrapper_codecs_decode_wrong_charset": (
            "import codecs\n"
            "def build_decoder():\n"
            "    decoder = codecs.decode\n"
            "    def apply(value, encoding, errors=None):\n"
            "        return decoder(value, encoding, errors=errors)\n"
            "    return apply\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_decoder()(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_codecs_decode_local_function_return_wrong_charset": (
            "import codecs\n"
            "def build_decoder():\n"
            "    def local():\n"
            "        return codecs.decode\n"
            "    return local()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_decoder()(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
    }

    for name, script in scripts.items():
        work = new_workdir(name)
        doc = Document()
        doc.add_paragraph("Synthetic Thesis")
        doc.add_paragraph("1 Introduction")
        doc.add_paragraph(text)
        doc.save(work / "out.docx")
        write_json(work / "content.json", base_content([text]))
        write_json(work / "format.json", base_format())
        write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
        write_json(work / "workflow_mode.json", {"mode": "user"})
        (work / "build_generated.py").write_text(script, encoding="utf-8")

        report = check_output(str(work), mode="user", output_docx_name="out.docx")
        codes = [item["code"] for item in report["issues"]]
        assert_true(
            "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" in codes,
            f"QA did not flag general codecs.decode text re-decoding route {name}: {report}",
        )
        action = f"{report.get('next_action')}\n{json.dumps(report.get('repair_plan') or {}, ensure_ascii=False)}"
        assert_true(
            "codecs.decode" in action and "读取文件字节" in action,
            f"general codecs.decode guidance did not explain byte-boundary decoding: {action}",
        )


@case
def qa_flags_generated_script_getattr_codecs_decode_text_reencoding() -> None:
    text = "中文字符保持原样：编码测试。"
    scripts = {
        "qa_generated_codecs_decode_getattr_alias_wrong_charset": (
            "import codecs\n"
            "decode_text = getattr(codecs, 'decode')\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = decode_text(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_codecs_decode_getattr_direct_wrong_charset": (
            "import codecs as text_codecs\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = getattr(text_codecs, 'decode')(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
    }

    for name, script in scripts.items():
        work = new_workdir(name)
        doc = Document()
        doc.add_paragraph("Synthetic Thesis")
        doc.add_paragraph("1 Introduction")
        doc.add_paragraph(text)
        doc.save(work / "out.docx")
        write_json(work / "content.json", base_content([text]))
        write_json(work / "format.json", base_format())
        write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
        write_json(work / "workflow_mode.json", {"mode": "user"})
        (work / "build_generated.py").write_text(script, encoding="utf-8")

        report = check_output(str(work), mode="user", output_docx_name="out.docx")
        codes = [item["code"] for item in report["issues"]]
        assert_true(
            "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" in codes,
            f"QA did not flag getattr(codecs, 'decode') text re-decoding route {name}: {report}",
        )
        action = f"{report.get('next_action')}\n{json.dumps(report.get('repair_plan') or {}, ensure_ascii=False)}"
        assert_true(
            "getattr" in action and "读取文件字节" in action,
            f"getattr(codecs, 'decode') guidance did not name the dynamic route: {action}",
        )


@case
def qa_flags_generated_script_partial_codecs_decode_text_reencoding() -> None:
    text = "中文字符保持原样：编码测试。"
    scripts = {
        "qa_generated_codecs_decode_partial_alias_wrong_charset": (
            "import codecs\n"
            "from functools import partial\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "decode_text = partial(codecs.decode, encoding='gbk')\n"
            "mojibake = decode_text(text.encode('utf-8'), errors='ignore')\n"
        ),
        "qa_generated_codecs_decode_partial_direct_wrong_charset": (
            "import codecs as text_codecs\n"
            "import functools as ft\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = ft.partial(text_codecs.decode, encoding='gbk')(text.encode('utf-8'), errors='ignore')\n"
        ),
        "qa_generated_codecs_decode_partial_dynamic_charset": (
            "import codecs\n"
            "from functools import partial\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "decode_text = partial(codecs.decode)\n"
            "mojibake = decode_text(text.encode('utf-8'), errors='ignore')\n"
        ),
        "qa_generated_codecs_decode_partial_function_return_wrong_charset": (
            "import codecs\n"
            "from functools import partial\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "def get_decoder():\n"
            "    return partial(codecs.decode, encoding='gbk')\n"
            "mojibake = get_decoder()(text.encode('utf-8'), errors='ignore')\n"
        ),
    }

    for name, script in scripts.items():
        work = new_workdir(name)
        doc = Document()
        doc.add_paragraph("Synthetic Thesis")
        doc.add_paragraph("1 Introduction")
        doc.add_paragraph(text)
        doc.save(work / "out.docx")
        write_json(work / "content.json", base_content([text]))
        write_json(work / "format.json", base_format())
        write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
        write_json(work / "workflow_mode.json", {"mode": "user"})
        (work / "build_generated.py").write_text(script, encoding="utf-8")

        report = check_output(str(work), mode="user", output_docx_name="out.docx")
        codes = [item["code"] for item in report["issues"]]
        assert_true(
            "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" in codes,
            f"QA did not flag functools.partial(codecs.decode) text re-decoding route {name}: {report}",
        )
        action = f"{report.get('next_action')}\n{json.dumps(report.get('repair_plan') or {}, ensure_ascii=False)}"
        assert_true(
            "partial" in action and "读取文件字节" in action,
            f"functools.partial(codecs.decode) guidance did not name the partial route: {action}",
        )


@case
def qa_flags_generated_script_method_decode_text_reencoding() -> None:
    text = "中文字符保持原样：编码测试。"
    scripts = {
        "qa_generated_method_decode_wrong_charset_chain": (
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = text.encode('utf-8').decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_method_decode_wrong_charset_alias": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "mojibake = payload.decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_getattr_method_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "mojibake = getattr(payload, 'decode')('gbk', errors='ignore')\n"
        ),
        "qa_generated_getattr_method_decode_alias_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "decode_payload = getattr(payload, 'decode')\n"
            "mojibake = decode_payload('gbk', errors='ignore')\n"
        ),
        "qa_generated_bound_method_decode_alias_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "decode_payload = payload.decode\n"
            "mojibake = decode_payload('gbk', errors='ignore')\n"
        ),
        "qa_generated_dunder_getattribute_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "mojibake = payload.__getattribute__('decode')('gbk', errors='ignore')\n"
        ),
        "qa_generated_operator_methodcaller_decode_wrong_charset": (
            "import operator\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "mojibake = operator.methodcaller('decode', 'gbk', errors='ignore')(payload)\n"
        ),
        "qa_generated_operator_methodcaller_alias_wrong_charset": (
            "from operator import methodcaller\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "decode_payload = methodcaller('decode', 'gbk', errors='ignore')\n"
            "mojibake = decode_payload(payload)\n"
        ),
        "qa_generated_operator_attrgetter_decode_wrong_charset": (
            "import operator\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "mojibake = operator.attrgetter('decode')(payload)('gbk', errors='ignore')\n"
        ),
        "qa_generated_list_bound_method_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "decode_routes = [payload.decode]\n"
            "mojibake = decode_routes[0]('gbk', errors='ignore')\n"
        ),
        "qa_generated_tuple_getattr_method_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "decode_routes = (getattr(payload, 'decode'),)\n"
            "mojibake = decode_routes[0]('gbk', errors='ignore')\n"
        ),
        "qa_generated_dict_bound_method_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "decode_routes = {'decode': payload.decode}\n"
            "mojibake = decode_routes['decode']('gbk', errors='ignore')\n"
        ),
        "qa_generated_list_operator_methodcaller_decode_wrong_charset": (
            "import operator\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "decode_routes = [operator.methodcaller('decode', 'gbk', errors='ignore')]\n"
            "mojibake = decode_routes[0](payload)\n"
        ),
        "qa_generated_dict_operator_methodcaller_decode_wrong_charset": (
            "from operator import methodcaller\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "decode_routes = {'decode': methodcaller('decode', 'gbk', errors='ignore')}\n"
            "mojibake = decode_routes['decode'](payload)\n"
        ),
        "qa_generated_function_returns_bound_method_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "def get_decoder():\n"
            "    return payload.decode\n"
            "mojibake = get_decoder()('gbk', errors='ignore')\n"
        ),
        "qa_generated_function_returns_getattr_method_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "def get_decoder():\n"
            "    return getattr(payload, 'decode')\n"
            "mojibake = get_decoder()('gbk', errors='ignore')\n"
        ),
        "qa_generated_function_returns_operator_methodcaller_decode_wrong_charset": (
            "import operator\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "def get_decoder():\n"
            "    return operator.methodcaller('decode', 'gbk', errors='ignore')\n"
            "mojibake = get_decoder()(payload)\n"
        ),
        "qa_generated_str_constructor_wrong_charset_direct": (
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = str(text.encode('utf-8'), 'gbk', errors='ignore')\n"
        ),
        "qa_generated_str_constructor_wrong_charset_alias": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = text.encode('utf-8')\n"
            "mojibake = str(payload, encoding='gbk', errors='ignore')\n"
        ),
        "qa_generated_bytes_constructor_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = bytes(text, 'utf-8')\n"
            "mojibake = payload.decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_bytearray_constructor_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = bytearray(text, 'utf-8')\n"
            "mojibake = payload.decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_bytearray_str_constructor_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = bytearray(text, 'utf-8')\n"
            "mojibake = str(payload, 'gbk', errors='ignore')\n"
        ),
        "qa_generated_bytes_constructor_alias_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "make_bytes = bytes\n"
            "payload = make_bytes(text, 'utf-8')\n"
            "mojibake = payload.decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_bytearray_constructor_alias_str_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "make_array = bytearray\n"
            "payload = make_array(text, 'utf-8')\n"
            "mojibake = str(payload, 'gbk', errors='ignore')\n"
        ),
        "qa_generated_builtins_bytes_import_alias_decode_wrong_charset": (
            "from builtins import bytes as make_bytes\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = make_bytes(text, 'utf-8')\n"
            "mojibake = payload.decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_builtins_getattr_bytes_alias_decode_wrong_charset": (
            "import builtins\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "make_bytes = getattr(builtins, 'bytes')\n"
            "payload = make_bytes(text, 'utf-8')\n"
            "mojibake = payload.decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_builtins_getattr_bytearray_alias_str_wrong_charset": (
            "import builtins as bi\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "make_array = getattr(bi, 'bytearray')\n"
            "payload = make_array(text, 'utf-8')\n"
            "mojibake = str(payload, 'gbk', errors='ignore')\n"
        ),
        "qa_generated_bytes_new_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = bytes.__new__(bytes, text, 'utf-8')\n"
            "mojibake = payload.decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_constructor_alias_new_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "make_bytes = bytes\n"
            "payload = make_bytes.__new__(make_bytes, text, 'utf-8')\n"
            "mojibake = payload.decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_dunder_builtins_attr_bytes_str_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = __builtins__.bytes(text, 'utf-8')\n"
            "mojibake = str(payload, 'gbk', errors='ignore')\n"
        ),
        "qa_generated_dunder_builtins_getattr_bytes_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "make_bytes = getattr(__builtins__, 'bytes')\n"
            "payload = make_bytes(text, 'utf-8')\n"
            "mojibake = payload.decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_dunder_builtins_subscript_bytes_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "make_bytes = __builtins__['bytes']\n"
            "payload = make_bytes(text, 'utf-8')\n"
            "mojibake = payload.decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_type_bytes_alias_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "make_bytes = type(b'')\n"
            "payload = make_bytes(text, 'utf-8')\n"
            "mojibake = payload.decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_bytes_class_attr_str_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "make_bytes = b''.__class__\n"
            "payload = make_bytes(text, 'utf-8')\n"
            "mojibake = str(payload, 'gbk', errors='ignore')\n"
        ),
        "qa_generated_type_bytearray_alias_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "make_array = type(bytearray())\n"
            "payload = make_array(text, 'utf-8')\n"
            "mojibake = payload.decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_bytearray_class_attr_str_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "make_array = bytearray().__class__\n"
            "payload = make_array(text, 'utf-8')\n"
            "mojibake = str(payload, 'gbk', errors='ignore')\n"
        ),
        "qa_generated_memoryview_tobytes_decode_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "payload = memoryview(text.encode('utf-8')).tobytes()\n"
            "mojibake = payload.decode('gbk', errors='ignore')\n"
        ),
        "qa_generated_memoryview_alias_tobytes_str_wrong_charset": (
            "text = '中文字符保持原样：编码测试。'\n"
            "view = memoryview(text.encode('utf-8'))\n"
            "payload = view.tobytes()\n"
            "mojibake = str(payload, 'gbk', errors='ignore')\n"
        ),
    }

    for name, script in scripts.items():
        work = new_workdir(name)
        doc = Document()
        doc.add_paragraph("Synthetic Thesis")
        doc.add_paragraph("1 Introduction")
        doc.add_paragraph(text)
        doc.save(work / "out.docx")
        write_json(work / "content.json", base_content([text]))
        write_json(work / "format.json", base_format())
        write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
        write_json(work / "workflow_mode.json", {"mode": "user"})
        (work / "build_generated.py").write_text(script, encoding="utf-8")

        report = check_output(str(work), mode="user", output_docx_name="out.docx")
        codes = [item["code"] for item in report["issues"]]
        assert_true(
            "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" in codes,
            f"QA did not flag method .decode text re-decoding route {name}: {report}",
        )
        action = f"{report.get('next_action')}\n{json.dumps(report.get('repair_plan') or {}, ensure_ascii=False)}"
        assert_true(
            ".decode" in action and "读取文件字节" in action,
            f"method .decode guidance did not explain byte-boundary decoding: {action}",
        )


@case
def qa_flags_generated_script_general_codecs_decoder_factories_text_reencoding() -> None:
    text = "中文字符保持原样：编码测试。"
    scripts = {
        "qa_generated_codecs_getdecoder_wrong_charset": (
            "import codecs\n"
            "decoder = codecs.getdecoder('gbk')\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = decoder(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_decode_wrong_charset": (
            "import codecs\n"
            "codec = codecs.lookup('gbk')\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = codec.decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_list_wrong_charset": (
            "import codecs\n"
            "routes = [codecs.getdecoder('gbk')]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = routes[0](text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_alias_list_wrong_charset": (
            "import codecs\n"
            "decoder = codecs.getdecoder('gbk')\n"
            "routes = [decoder]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = routes[0](text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_function_return_wrong_charset": (
            "import codecs\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "def get_decoder():\n"
            "    return codecs.getdecoder('gbk')\n"
            "mojibake = get_decoder()(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_alias_function_return_wrong_charset": (
            "import codecs\n"
            "decoder = codecs.getdecoder('gbk')\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "def get_decoder():\n"
            "    return decoder\n"
            "mojibake = get_decoder()(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_list_wrong_charset": (
            "import codecs\n"
            "routes = [codecs.lookup('gbk')]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = routes[0].decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_function_return_wrong_charset": (
            "import codecs\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "def get_codec():\n"
            "    return codecs.lookup('gbk')\n"
            "mojibake = get_codec().decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_attribute_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.decoder = codecs.getdecoder('gbk')\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = box.decoder(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_attribute_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.codec = codecs.lookup('gbk')\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = box.codec.decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_namespace_attribute_wrong_charset": (
            "import codecs\n"
            "from types import SimpleNamespace\n"
            "ns = SimpleNamespace(decoder=codecs.getdecoder('gbk'))\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = ns.decoder(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_namespace_attribute_wrong_charset": (
            "import codecs\n"
            "from types import SimpleNamespace\n"
            "ns = SimpleNamespace(codec=codecs.lookup('gbk'))\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = ns.codec.decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_attribute_function_return_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.decoder = codecs.getdecoder('gbk')\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "def get_decoder():\n"
            "    return box.decoder\n"
            "mojibake = get_decoder()(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_attribute_function_return_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.codec = codecs.lookup('gbk')\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "def get_codec():\n"
            "    return box.codec\n"
            "mojibake = get_codec().decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_object_alias_attribute_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.decoder = codecs.getdecoder('gbk')\n"
            "alias = box\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = alias.decoder(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_object_alias_attribute_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.codec = codecs.lookup('gbk')\n"
            "alias = box\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = alias.codec.decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_class_attribute_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    decoder = codecs.getdecoder('gbk')\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = Holder.decoder(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_class_attribute_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    codec = codecs.lookup('gbk')\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = Holder.codec.decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_instance_attribute_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.decoder = codecs.getdecoder('gbk')\n"
            "holder = Holder()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = holder.decoder(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_instance_attribute_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.codec = codecs.lookup('gbk')\n"
            "holder = Holder()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = holder.codec.decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_instance_alias_attribute_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.codec = codecs.lookup('gbk')\n"
            "holder = Holder()\n"
            "alias = holder\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = alias.codec.decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_class_alias_instance_attribute_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.codec = codecs.lookup('gbk')\n"
            "Alias = Holder\n"
            "holder = Alias()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = holder.codec.decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_factory_instance_attribute_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.decoder = codecs.getdecoder('gbk')\n"
            "def make_holder():\n"
            "    return Holder()\n"
            "holder = make_holder()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = holder.decoder(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_class_alias_temp_instance_attribute_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.codec = codecs.lookup('gbk')\n"
            "Alias = Holder\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = Alias().codec.decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_staticmethod_factory_instance_attribute_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.decoder = codecs.getdecoder('gbk')\n"
            "    @staticmethod\n"
            "    def make():\n"
            "        return Holder()\n"
            "holder = Holder.make()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = holder.decoder(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_classmethod_temp_instance_attribute_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    def __init__(self):\n"
            "        self.codec = codecs.lookup('gbk')\n"
            "    @classmethod\n"
            "    def make(cls):\n"
            "        return cls()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = Holder.make().codec.decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_local_function_return_wrong_charset": (
            "import codecs\n"
            "def build_decoder():\n"
            "    def local():\n"
            "        return codecs.getdecoder('gbk')\n"
            "    return local()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_decoder()(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_lookup_local_function_return_wrong_charset": (
            "import codecs\n"
            "def build_codec():\n"
            "    def local():\n"
            "        return codecs.lookup('gbk')\n"
            "    return local()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_codec().decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_higher_order_factory_wrong_charset": (
            "import codecs\n"
            "def apply_factory(factory, value, encoding):\n"
            "    return factory(encoding)(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_factory(codecs.getdecoder, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_codecs_getdecoder_returned_higher_order_factory_wrong_charset": (
            "import codecs\n"
            "def build_apply():\n"
            "    def apply_factory(factory, value, encoding):\n"
            "        return factory(encoding)(value, errors='ignore')[0]\n"
            "    return apply_factory\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_apply()(codecs.getdecoder, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_codecs_lookup_higher_order_factory_wrong_charset": (
            "import codecs\n"
            "def apply_factory(factory, value, encoding):\n"
            "    return factory(encoding).decode(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_factory(codecs.lookup, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_codecs_lookup_returned_higher_order_factory_wrong_charset": (
            "import codecs\n"
            "def build_apply():\n"
            "    def apply_factory(factory, value, encoding):\n"
            "        return factory(encoding).decode(value, errors='ignore')[0]\n"
            "    return apply_factory\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_apply()(codecs.lookup, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_codecs_module_param_getdecoder_wrong_charset": (
            "import codecs\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.getdecoder(encoding)(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(codecs, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_getattr_codecs_module_param_getdecoder_wrong_charset": (
            "import codecs\n"
            "def apply_module(module, value, encoding):\n"
            "    return getattr(module, 'getdecoder')(encoding)(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(codecs, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_local_getattr_codecs_module_getdecoder_alias_wrong_charset": (
            "import codecs\n"
            "def apply_module(module, value, encoding):\n"
            "    factory = getattr(module, 'getdecoder')\n"
            "    return factory(encoding)(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(codecs, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_required_arg_function_returns_getattr_module_getdecoder_wrong_charset": (
            "import codecs\n"
            "def pick_getdecoder(module):\n"
            "    return getattr(module, 'getdecoder')\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = pick_getdecoder(codecs)('gbk')(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_required_arg_closure_returns_getattr_module_getdecoder_wrong_charset": (
            "import codecs\n"
            "def build_getdecoder(module):\n"
            "    def pick_getdecoder():\n"
            "        return getattr(module, 'getdecoder')\n"
            "    return pick_getdecoder()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_getdecoder(codecs)('gbk')(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_required_arg_nested_closure_returns_getattr_module_getdecoder_wrong_charset": (
            "import codecs\n"
            "def build_getdecoder(module):\n"
            "    def outer():\n"
            "        def pick_getdecoder():\n"
            "            return getattr(module, 'getdecoder')\n"
            "        return pick_getdecoder()\n"
            "    return outer()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_getdecoder(codecs)('gbk')(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_module_attribute_param_getdecoder_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.module = codecs\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.getdecoder(encoding)(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(box.module, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_module_container_return_param_getdecoder_wrong_charset": (
            "import codecs\n"
            "def get_modules():\n"
            "    return (codecs,)\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.getdecoder(encoding)(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(get_modules()[0], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_module_attribute_container_param_getdecoder_wrong_charset": (
            "import codecs\n"
            "from types import SimpleNamespace\n"
            "box = SimpleNamespace(modules=(codecs,))\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.getdecoder(encoding)(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(box.modules[0], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_attribute_container_return_param_getdecoder_wrong_charset": (
            "import codecs\n"
            "from types import SimpleNamespace\n"
            "box = SimpleNamespace(modules=(codecs,))\n"
            "def get_modules():\n"
            "    return box.modules\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.getdecoder(encoding)(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(get_modules()[0], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_nested_module_container_param_getdecoder_wrong_charset": (
            "import codecs\n"
            "config = {'modules': (codecs,)}\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.getdecoder(encoding)(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(config['modules'][0], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_staticmethod_module_container_param_getdecoder_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    @staticmethod\n"
            "    def get_modules():\n"
            "        return (codecs,)\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.getdecoder(encoding)(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(Holder.get_modules()[0], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_bound_getdecoder_from_module_return_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    @staticmethod\n"
            "    def get_module():\n"
            "        return codecs\n"
            "factory = Holder.get_module().getdecoder\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = factory('gbk')(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_function_returns_bound_getdecoder_from_module_return_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    @staticmethod\n"
            "    def get_module():\n"
            "        return codecs\n"
            "def get_factory():\n"
            "    return Holder.get_module().getdecoder\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = get_factory()('gbk')(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_module_param_lookup_wrong_charset": (
            "import codecs\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.lookup(encoding).decode(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(codecs, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_getattr_codecs_module_param_lookup_wrong_charset": (
            "import codecs\n"
            "def apply_module(module, value, encoding):\n"
            "    return getattr(module, 'lookup')(encoding).decode(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(codecs, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_local_getattr_codecs_module_lookup_alias_wrong_charset": (
            "import codecs\n"
            "def apply_module(module, value, encoding):\n"
            "    lookup = getattr(module, 'lookup')\n"
            "    return lookup(encoding).decode(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(codecs, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_required_arg_function_returns_getattr_module_lookup_wrong_charset": (
            "import codecs\n"
            "def pick_lookup(module):\n"
            "    return getattr(module, 'lookup')\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = pick_lookup(codecs)('gbk').decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_required_arg_closure_returns_getattr_module_lookup_wrong_charset": (
            "import codecs\n"
            "def build_lookup(module):\n"
            "    def pick_lookup():\n"
            "        return getattr(module, 'lookup')\n"
            "    return pick_lookup()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_lookup(codecs)('gbk').decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_required_arg_nested_closure_returns_getattr_module_lookup_wrong_charset": (
            "import codecs\n"
            "def build_lookup(module):\n"
            "    def outer():\n"
            "        def pick_lookup():\n"
            "            return getattr(module, 'lookup')\n"
            "        return pick_lookup()\n"
            "    return outer()\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_lookup(codecs)('gbk').decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_module_attribute_param_lookup_wrong_charset": (
            "import codecs\n"
            "class Box:\n"
            "    pass\n"
            "box = Box()\n"
            "box.module = codecs\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.lookup(encoding).decode(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(box.module, text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_module_container_return_param_lookup_wrong_charset": (
            "import codecs\n"
            "def get_modules():\n"
            "    return {'module': codecs}\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.lookup(encoding).decode(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(get_modules()['module'], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_class_attribute_module_container_param_lookup_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    modules = {'module': codecs}\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.lookup(encoding).decode(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(Holder.modules['module'], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_class_attribute_container_return_param_lookup_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    modules = {'module': codecs}\n"
            "def get_modules():\n"
            "    return Holder.modules\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.lookup(encoding).decode(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(get_modules()['module'], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_class_attribute_nested_module_container_param_lookup_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    config = {'modules': {'module': codecs}}\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.lookup(encoding).decode(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(Holder.config['modules']['module'], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_classmethod_nested_module_container_param_lookup_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    @classmethod\n"
            "    def get_config(cls):\n"
            "        return {'modules': {'module': codecs}}\n"
            "def apply_module(module, value, encoding):\n"
            "    return module.lookup(encoding).decode(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_module(Holder.get_config()['modules']['module'], text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_bound_lookup_from_nested_module_return_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    @classmethod\n"
            "    def get_config(cls):\n"
            "        return {'modules': {'module': codecs}}\n"
            "lookup = Holder.get_config()['modules']['module'].lookup\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = lookup('gbk').decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_function_returns_bound_lookup_from_nested_module_return_wrong_charset": (
            "import codecs\n"
            "class Holder:\n"
            "    @classmethod\n"
            "    def get_config(cls):\n"
            "        return {'modules': {'module': codecs}}\n"
            "def get_lookup():\n"
            "    return Holder.get_config()['modules']['module'].lookup\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = get_lookup()('gbk').decode(text.encode('utf-8'), errors='ignore')[0]\n"
        ),
        "qa_generated_codecs_getdecoder_factory_return_passed_to_higher_order_wrong_charset": (
            "import codecs\n"
            "def get_factory():\n"
            "    return codecs.getdecoder\n"
            "def apply_factory(factory, value, encoding):\n"
            "    return factory(encoding)(value, errors='ignore')[0]\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = apply_factory(get_factory(), text.encode('utf-8'), 'gbk')\n"
        ),
        "qa_generated_codecs_lookup_factory_return_passed_to_returned_higher_order_wrong_charset": (
            "import codecs\n"
            "def get_factory():\n"
            "    return codecs.lookup\n"
            "def build_apply():\n"
            "    def apply_factory(factory, value, encoding):\n"
            "        return factory(encoding).decode(value, errors='ignore')[0]\n"
            "    return apply_factory\n"
            "text = '中文字符保持原样：编码测试。'\n"
            "mojibake = build_apply()(get_factory(), text.encode('utf-8'), 'gbk')\n"
        ),
    }

    for name, script in scripts.items():
        work = new_workdir(name)
        doc = Document()
        doc.add_paragraph("Synthetic Thesis")
        doc.add_paragraph("1 Introduction")
        doc.add_paragraph(text)
        doc.save(work / "out.docx")
        write_json(work / "content.json", base_content([text]))
        write_json(work / "format.json", base_format())
        write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
        write_json(work / "workflow_mode.json", {"mode": "user"})
        (work / "build_generated.py").write_text(script, encoding="utf-8")

        report = check_output(str(work), mode="user", output_docx_name="out.docx")
        codes = [item["code"] for item in report["issues"]]
        assert_true(
            "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" in codes,
            f"QA did not flag general codecs decoder-factory route {name}: {report}",
        )
        action = f"{report.get('next_action')}\n{json.dumps(report.get('repair_plan') or {}, ensure_ascii=False)}"
        assert_true(
            ("getdecoder" in action or "lookup" in action) and "读取文件字节" in action,
            f"general decoder-factory guidance did not explain byte-boundary decoding: {action}",
        )


@case
def qa_does_not_flag_shadowed_method_factory_name_for_safe_decoder() -> None:
    text = "中文字符保持原样：编码测试。"
    work = new_workdir("qa_shadowed_method_factory_name_safe_decoder")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "import codecs\n"
        "class Safe:\n"
        "    def __init__(self):\n"
        "        self.decode = lambda value, encoding, errors=None: value.decode('utf-8')\n"
        "def make():\n"
        "    return Safe()\n"
        "class Holder:\n"
        "    def __init__(self):\n"
        "        self.decode = codecs.decode\n"
        "    @staticmethod\n"
        "    def make():\n"
        "        return Holder()\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = make().decode(text.encode('utf-8'), 'gbk', errors='ignore')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a class staticmethod as the same-scope global factory: {report}",
    )

    work = new_workdir("qa_shadowed_local_function_name_safe_decoder")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "import codecs\n"
        "def local():\n"
        "    return codecs.decode\n"
        "def build_decoder():\n"
        "    def local():\n"
        "        return lambda value, encoding, errors=None: value.decode('utf-8')\n"
        "    return local()\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = build_decoder()(text.encode('utf-8'), 'gbk', errors='ignore')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a global unsafe local() as a same-scope child local(): {report}",
    )

    work = new_workdir("qa_shadowed_returned_higher_order_wrapper_safe_decoder")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "import codecs\n"
        "def apply_decoder(decoder, value, encoding):\n"
        "    return decoder(value, encoding, errors='ignore')\n"
        "def build_apply():\n"
        "    def apply_decoder(decoder, value, encoding):\n"
        "        return value.decode('utf-8')\n"
        "    return apply_decoder\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = build_apply()(codecs.decode, text.encode('utf-8'), 'gbk')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a global higher-order wrapper as the returned safe local wrapper: {report}",
    )

    work = new_workdir("qa_shadowed_returned_factory_wrapper_safe_decoder")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "import codecs\n"
        "def apply_factory(factory, value, encoding):\n"
        "    return factory(encoding)(value, errors='ignore')[0]\n"
        "def build_apply():\n"
        "    def apply_factory(factory, value, encoding):\n"
        "        return value.decode('utf-8')\n"
        "    return apply_factory\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = build_apply()(codecs.getdecoder, text.encode('utf-8'), 'gbk')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a global decoder factory wrapper as the returned safe local wrapper: {report}",
    )

    work = new_workdir("qa_shadowed_returned_module_wrapper_safe_decoder")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "import codecs\n"
        "def apply_module(module, value, encoding):\n"
        "    return module.decode(value, encoding, errors='ignore')\n"
        "def build_apply():\n"
        "    def apply_module(module, value, encoding):\n"
        "        return value.decode('utf-8')\n"
        "    return apply_module\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = build_apply()(codecs, text.encode('utf-8'), 'gbk')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a global module wrapper as the returned safe local wrapper: {report}",
    )

    work = new_workdir("qa_safe_module_attribute_is_not_codecs_module")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "class Box:\n"
        "    pass\n"
        "box = Box()\n"
        "box.module = SafeModule()\n"
        "def apply_module(module, value, encoding):\n"
        "    return module.decode(value, encoding, errors='ignore')\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = apply_module(box.module, text.encode('utf-8'), 'gbk')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a safe object attribute as the codecs module: {report}",
    )

    work = new_workdir("qa_safe_module_container_return_is_not_codecs_module")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "def get_modules():\n"
        "    return [SafeModule()]\n"
        "def apply_module(module, value, encoding):\n"
        "    return module.decode(value, encoding, errors='ignore')\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = apply_module(get_modules()[0], text.encode('utf-8'), 'gbk')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a safe module returned from a container helper as codecs: {report}",
    )

    work = new_workdir("qa_safe_module_attribute_container_is_not_codecs_module")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "class Box:\n"
        "    pass\n"
        "box = Box()\n"
        "box.modules = [SafeModule()]\n"
        "def apply_module(module, value, encoding):\n"
        "    return module.decode(value, encoding, errors='ignore')\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = apply_module(box.modules[0], text.encode('utf-8'), 'gbk')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a safe module from an attribute container as codecs: {report}",
    )

    work = new_workdir("qa_safe_module_attribute_container_return_is_not_codecs_module")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "class Box:\n"
        "    pass\n"
        "box = Box()\n"
        "box.modules = [SafeModule()]\n"
        "def get_modules():\n"
        "    return box.modules\n"
        "def apply_module(module, value, encoding):\n"
        "    return module.decode(value, encoding, errors='ignore')\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = apply_module(get_modules()[0], text.encode('utf-8'), 'gbk')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a safe module from a returned attribute container as codecs: {report}",
    )

    work = new_workdir("qa_safe_nested_module_container_is_not_codecs_module")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "config = {'modules': [SafeModule()]}\n"
        "def apply_module(module, value, encoding):\n"
        "    return module.decode(value, encoding, errors='ignore')\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = apply_module(config['modules'][0], text.encode('utf-8'), 'gbk')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a safe nested container module as codecs: {report}",
    )

    work = new_workdir("qa_safe_instance_method_module_return_is_not_codecs_module")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "class Holder:\n"
        "    def get_module(self):\n"
        "        return SafeModule()\n"
        "holder = Holder()\n"
        "def apply_module(module, value, encoding):\n"
        "    return module.decode(value, encoding, errors='ignore')\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = apply_module(holder.get_module(), text.encode('utf-8'), 'gbk')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a safe instance method module return as codecs: {report}",
    )

    work = new_workdir("qa_safe_bound_method_from_instance_method_return_is_not_codecs")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "class Holder:\n"
        "    def get_module(self):\n"
        "        return SafeModule()\n"
        "holder = Holder()\n"
        "decode_text = holder.get_module().decode\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = decode_text(text.encode('utf-8'), 'gbk', errors='ignore')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a safe bound method from instance method return as codecs: {report}",
    )

    work = new_workdir("qa_safe_function_returned_bound_method_is_not_codecs")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "class Holder:\n"
        "    def get_module(self):\n"
        "        return SafeModule()\n"
        "holder = Holder()\n"
        "def get_decode():\n"
        "    return holder.get_module().decode\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = get_decode()(text.encode('utf-8'), 'gbk', errors='ignore')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a safe function-returned bound method as codecs: {report}",
    )

    work = new_workdir("qa_safe_operator_attrgetter_decode_is_not_codecs")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "from operator import attrgetter\n"
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "decode_text = attrgetter('decode')(SafeModule())\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = decode_text(text.encode('utf-8'), 'gbk', errors='ignore')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a safe attrgetter decode method as codecs: {report}",
    )

    work = new_workdir("qa_safe_operator_methodcaller_decode_is_not_codecs")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "from operator import methodcaller\n"
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "decode_via_module_shape = methodcaller('decode', text.encode('utf-8'), 'gbk', errors='ignore')(SafeModule())\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a safe methodcaller decode route as codecs: {report}",
    )

    work = new_workdir("qa_safe_local_getattr_module_decode_alias")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "def apply_module(module, value, encoding):\n"
        "    decode = getattr(module, 'decode')\n"
        "    return decode(value, encoding, errors='ignore')\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = apply_module(SafeModule(), text.encode('utf-8'), 'gbk')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a local getattr alias on a safe module as codecs: {report}",
    )

    work = new_workdir("qa_safe_required_arg_function_returns_getattr_module_decode")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "def pick_decode(module):\n"
        "    return getattr(module, 'decode')\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = pick_decode(SafeModule())(text.encode('utf-8'), 'gbk', errors='ignore')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a required-arg safe module getattr return as codecs: {report}",
    )

    work = new_workdir("qa_safe_required_arg_closure_returns_getattr_module_decode")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "def build_decoder(module):\n"
        "    def pick_decode():\n"
        "        return getattr(module, 'decode')\n"
        "    return pick_decode()\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = build_decoder(SafeModule())(text.encode('utf-8'), 'gbk', errors='ignore')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a required-arg closure safe module getattr return as codecs: {report}",
    )

    work = new_workdir("qa_safe_required_arg_nested_closure_returns_getattr_module_decode")
    doc = Document()
    doc.add_paragraph("Synthetic Thesis")
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph(text)
    doc.save(work / "out.docx")
    write_json(work / "content.json", base_content([text]))
    write_json(work / "format.json", base_format())
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text(
        "class SafeModule:\n"
        "    def decode(self, value, encoding, errors='strict'):\n"
        "        return value.decode('utf-8')\n"
        "def build_decoder(module):\n"
        "    def outer():\n"
        "        def pick_decode():\n"
        "            return getattr(module, 'decode')\n"
        "        return pick_decode()\n"
        "    return outer()\n"
        "text = '中文字符保持原样：编码测试。'\n"
        "roundtrip = build_decoder(SafeModule())(text.encode('utf-8'), 'gbk', errors='ignore')\n",
        encoding="utf-8",
    )

    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true(
        "GENERATED_SCRIPT_UNSAFE_UNICODE_DECODE" not in codes,
        f"QA falsely treated a required-arg nested-closure safe module getattr return as codecs: {report}",
    )


@case
def qa_counts_nested_note_refs_when_meta_is_missing() -> None:
    work = new_workdir("qa_nested_note_refs_without_meta")
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Nested notes should still be QA-visible.")
    doc.save(work / "out.docx")
    content = base_content(
        [
            {
                "role": "rich_text",
                "text": "Nested notes should still be QA-visible.",
                "items": [
                    {
                        "role": "rich_text",
                        "text": "Nested footnote",
                        "runs": [
                            {"type": "text", "text": "Nested footnote"},
                            {"type": "note_ref", "note_type": "footnote", "text": "Nested footnote text."},
                        ],
                    }
                ],
                "runs": [
                    {
                        "type": "text",
                        "text": "run cell note",
                        "table_cell_items": [
                            {
                                "row": 0,
                                "col": 0,
                                "items": [
                                    {
                                        "role": "rich_text",
                                        "text": "Nested endnote",
                                        "runs": [
                                            {"type": "text", "text": "Nested endnote"},
                                            {"type": "note_ref", "note_type": "endnote", "text": "Nested endnote text."},
                                        ],
                                    }
                                ],
                            }
                        ],
                    }
                ],
            }
        ]
    )
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(
        work / "build_manifest.json",
        {"schema_version": 1, "counts": {"footnote_references_rendered": 0, "endnote_references_rendered": 0}},
    )
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    (work / "build_generated.py").write_text("# synthetic generated script\n", encoding="utf-8")

    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("FOOTNOTE_RENDER_COUNT_MISMATCH" in codes, f"Nested footnote refs were not counted: {report['issues']}")
    assert_true("ENDNOTE_RENDER_COUNT_MISMATCH" in codes, f"Nested endnote refs were not counted: {report['issues']}")


@case
def code_table_is_not_body_table() -> None:
    content = base_content(
        [
            {
                "role": "code",
                "code": "interface GigabitEthernet0/0/1\nip address 10.0.0.1 255.255.255.0",
                "table_rows": [["interface GigabitEthernet0/0/1\nip address 10.0.0.1 255.255.255.0"]],
            }
        ],
        meta_tables=1,
    )
    result = run_generated_case("code_table", content)
    assert_true(result["manifest"]["counts"]["content_tables_rendered"] == 0, "code box was counted as body table")
    assert_true(result["report"]["counts"]["content_tables"] == 0, "QA counted code table_rows as body table")
    assert_true(not result["report"]["issues"], f"unexpected QA issues: {result['report']['issues']}")


@case
def qa_counts_xml_only_formula_items() -> None:
    content = base_content([
        {"role": "source_omml", "text": "E=mc^2", "xml": "<m:oMath/>"},
    ])
    assert_true(_count_content_formulas(content) == 1, "QA should count source OMML/XML formula items")


@case
def qa_counts_image_fields_even_when_role_varies() -> None:
    content = base_content([
        {"role": "media", "filename": "figure_a.png", "caption": "Figure A"},
    ])
    assert_true(_count_content_images(content) == 1, "QA should count image items by image fields, not only by role")


@case
def qa_counts_nested_rich_media_and_tables_recursively() -> None:
    content = base_content(
        [
            {
                "role": "rich_text",
                "text": "Rich parent with nested content.",
                "runs": [
                    {"type": "text", "text": "Before "},
                    {
                        "type": "text",
                        "text": "run-items",
                        "items": [
                            {"role": "image", "image": "run_item.png"},
                            {"role": "formula", "latex": "a+b"},
                            {"role": "table", "table_rows": [["run", "table"]]},
                        ],
                    },
                    {
                        "type": "text",
                        "text": "cell-items",
                        "table_cell_items": [
                            {
                                "row": 0,
                                "col": 0,
                                "items": [
                                    {"role": "image", "image": "run_cell.png"},
                                    {"role": "formula", "latex": "c+d"},
                                    {"role": "table", "table_rows": [["cell", "table"]]},
                                ],
                            }
                        ],
                    },
                ],
                "items": [
                    {"role": "image", "image": "rich_item.png"},
                    {"role": "formula", "latex": "e+f"},
                    {"role": "table", "table_rows": [["rich", "table"]]},
                ],
            }
        ]
    )

    assert_true(_count_content_images(content) == 3, "Structural QA should count nested rich_text images recursively")
    assert_true(_count_content_formulas(content) == 3, "Structural QA should count nested rich_text formulas recursively")
    assert_true(_count_content_tables(content) == 3, "Structural QA should count nested rich_text tables recursively")
    req = build_requirements(base_format(), content)
    assert_true(
        req["expected_counts"] == {"images": 3, "tables": 3, "formulas": 3},
        f"Strict QA expected counts should recurse into nested rich content: {req['expected_counts']}",
    )


@case
def qa_counts_deep_nested_rich_text_chars_recursively() -> None:
    deep_text = "深层中文正文" * 80
    content = base_content(
        [
            {
                "role": "rich_text",
                "text": "",
                "runs": [
                    {
                        "type": "text",
                        "text": "",
                        "items": [
                            {
                                "role": "rich_text",
                                "text": "",
                                "items": [
                                    {"role": "paragraph", "text": deep_text},
                                ],
                            }
                        ],
                    }
                ],
            }
        ]
    )

    assert_true(
        _content_text_chars(content) >= len(deep_text),
        "Structural QA should count deep nested rich text when judging content/docx text loss",
    )


@case
def conformance_expected_paragraphs_classifies_plain_caption_strings() -> None:
    content = base_content([
        "Table 1 Model results",
        "表 2 变量描述",
        "图 3 系统架构",
        "图 1 展示了系统架构。",
        "Figure 2 shows the workflow.",
        "This is a body paragraph.",
    ])
    expected = _expected_paragraphs(content)
    roles_by_text = {item["text"]: item["role"] for item in expected}
    assert_true(roles_by_text["Table 1 Model results"] == "table_caption", "plain English table caption string should use table_caption style")
    assert_true(roles_by_text["表 2 变量描述"] == "table_caption", "plain Chinese table caption string should use table_caption style")
    assert_true(roles_by_text["图 3 系统架构"] == "figure_caption", "plain Chinese figure caption string should use figure_caption style")
    assert_true(roles_by_text["图 1 展示了系统架构。"] == "body", "plain Chinese figure-reference prose should remain body")
    assert_true(roles_by_text["Figure 2 shows the workflow."] == "body", "plain English figure-reference prose should remain body")
    assert_true(roles_by_text["This is a body paragraph."] == "body", "ordinary strings should remain body paragraphs")


@case
def qa_toc_pollution_allows_numbered_multilevel_headings() -> None:
    content = {
        "sections": [
            {"heading": "1 Chapter 1", "level": 1, "paragraphs": ["Opening paragraph."]},
            {"heading": "1.1 Section 1.1", "level": 2, "paragraphs": ["Nested body paragraph."]},
            {"heading": "1.1.1 Detail", "level": 3, "paragraphs": ["Third-level content."]},
        ]
    }
    assert_true(not _content_toc_pollution_samples(content), "legitimate numbered multilevel headings were flagged as TOC pollution")


@case
def conformance_finds_inline_math_paragraphs_by_linearized_text() -> None:
    xml = """
    <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <w:r><w:t>Inline formula </w:t></w:r>
      <m:oMath>
        <m:r><m:t>E</m:t></m:r>
        <m:r><m:t>=</m:t></m:r>
        <m:r><m:t>m</m:t></m:r>
        <m:sSup><m:e><m:r><m:t>c</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>
      </m:oMath>
      <w:r><w:t> should remain editable.</w:t></w:r>
    </w:p>
    """
    para = ET.fromstring(xml)
    found = _find_para_by_text([para], "Inline formula E=mc^2 should remain editable.")
    assert_true(found is para, "conformance QA did not match paragraph text containing editable inline math")


@case
def conformance_body_start_keeps_default_body_before_first_heading() -> None:
    def para(text: str) -> ET.Element:
        return ET.fromstring(
            '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f"<w:r><w:t>{text}</w:t></w:r>"
            "</w:p>"
        )

    content = {
        "sections": [
            {
                "heading": "正文",
                "level": 1,
                "role": "body",
                "paragraphs": [
                    "This paragraph appears before the first explicit heading.",
                    "It must still be checked by strict QA.",
                ],
            },
            {
                "heading": "Methods",
                "level": 2,
                "role": "",
                "paragraphs": ["Method body text."],
            },
        ],
        "references": [],
    }
    expected = _expected_paragraphs(content)
    paragraphs = [
        para("Cover Title"),
        para("目 录"),
        para("Methods\t1"),
        para("This paragraph appears before the first explicit heading."),
        para("It must still be checked by strict QA."),
        para("Methods"),
        para("Method body text."),
    ]
    start = _find_body_start_index(paragraphs, expected)
    used: set[int] = set()
    assert_true(
        _find_para_by_text(paragraphs[start:], "This paragraph appears before the first explicit heading.", used) is not None,
        f"strict QA body start skipped default body paragraphs before the first explicit heading; start={start}",
    )


@case
def image_manifest_matches_rendered_body_images() -> None:
    img_src = new_workdir("image_src")
    write_sample_png(img_src / "dot.png")
    content = base_content([
        {"role": "figure", "image": "dot.png", "caption": "Figure 1 sample"}
    ])
    content["_meta"]["images_dir"] = str(img_src)
    content["_meta"]["images_extracted"] = 1
    content["sections"][0]["images"] = ["dot.png"]
    result = run_generated_case("image_manifest", content)
    assert_true(result["manifest"]["counts"]["content_images_rendered"] == 1, "body image was not counted")
    assert_true(result["report"]["counts"]["content_images"] == 1, "QA did not count body image occurrence")
    assert_true(not result["report"]["issues"], f"unexpected QA issues: {result['report']['issues']}")


@case
def qa_counts_mixed_inline_and_section_images() -> None:
    work = new_workdir("mixed_image_count")
    content = base_content([
        {"role": "image", "image": "inline.png"},
    ])
    content["sections"][0]["images"] = ["inline.png", "", "section_only.png"]
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    assert_true(report["counts"]["content_images"] == 2, f"mixed image count lost section-only images: {report['counts']}")


@case
def qa_reports_non_body_images_and_raw_latex_text() -> None:
    work = new_workdir("qa_non_body_latex")
    docx = work / "out.docx"
    doc = Document()
    doc.add_paragraph(r"$$x^2+y^2=z^2$$")
    doc.save(docx)
    content = base_content(["Body text"])
    content["_meta"]["non_body_images"] = [{"location": "section_1_header", "target": "media/image1.png"}]
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("NON_BODY_IMAGE_UNSUPPORTED" in codes, "QA did not flag unsupported header/footer images")
    assert_true("LATEX_DELIMITER_TEXT" in codes, "QA did not flag raw LaTeX delimiters left in final DOCX")
    assert_true(report["passed"] is False, "non-body images and raw LaTeX text should fail QA")


@case
def qa_reports_duplicate_front_matter_headings() -> None:
    work = new_workdir("qa_duplicate_front_heading")
    docx = work / "out.docx"
    doc = Document()
    doc.add_paragraph("摘  要")
    doc.add_paragraph("Synthetic title")
    doc.add_paragraph("摘要")
    doc.add_paragraph("Abstract body.")
    doc.save(docx)
    content = {
        "_meta": {"source": "synthetic.docx", "paragraphs": 1, "tables_count": 0, "images_extracted": 0},
        "title_info": {"title_cn": "Synthetic title"},
        "sections": [
            {"heading": "摘要", "level": 1, "role": "cn_abstract", "paragraphs": ["Abstract body."], "images": []}
        ],
        "references": ["[1] Synthetic reference."],
    }
    write_json(work / "content.json", content)
    write_json(work / "format.json", base_format())
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    (work / "build_generated.py").write_text("# synthetic\n", encoding="utf-8")
    report = check_output(str(work), mode="developer", output_docx_name="out.docx")
    codes = [item["code"] for item in report["issues"]]
    assert_true("DUPLICATE_FRONT_MATTER_HEADING" in codes, "QA did not report duplicate front matter heading")
    assert_true(report["passed"] is False, "duplicate front matter heading should fail QA")


@case
def qa_repair_plan_uses_relative_rebuild_command() -> None:
    work = new_workdir("qa_relative_rebuild_command")
    write_json(work / "workflow_mode.json", {"mode": "user"})
    report = {
        "mode": "user",
        "passed": True,
        "issues": [],
        "counts": {},
    }
    plan = build_repair_plan(report, str(work))
    command = plan["commands"]["rebuild_current_docx"]
    assert_true("build_generated.py" in command, "rebuild command should still point to build_generated.py")
    assert_true(str(work) not in command, f"rebuild command leaked absolute output path: {command}")


@case
def qa_repair_guides_cover_registered_issue_codes() -> None:
    missing = sorted(code for code in OWNER_BY_CODE if code not in REPAIR_GUIDES)
    assert_true(not missing, f"registered QA issue codes missing user-facing repair guides: {missing}")


@case
def qa_repair_plan_preserves_md_visual_workflow_command() -> None:
    work = new_workdir("qa_md_visual_workflow_command")
    write_json(
        work / "workflow_mode.json",
        {
            "mode": "user",
            "template": "demo.md",
            "content": "demo.md",
            "md": "demo.md",
            "qa_level": "visual",
            "golden_dir": "TestData/GoldenBaselines",
            "update_golden": True,
            "require_wps": True,
            "auto_repair": True,
            "repair_max_rounds": 4,
            "repair_stop_no_improve": 2,
        },
    )
    report = {
        "mode": "user",
        "passed": False,
        "issues": [{"code": "PLACEHOLDER_TEXT_LEFT", "severity": "error", "message": "placeholder"}],
        "counts": {},
    }
    plan = build_repair_plan(report, str(work))
    command = plan["commands"]["rerun_current_pipeline"]
    assert_true("--md demo.md" in command, f"MD workflow should rerun with --md: {command}")
    assert_true("--template" not in command and "--content" not in command, f"MD workflow should not be rewritten as template/content: {command}")
    assert_true("--qa-level visual" in command, f"QA level was not preserved: {command}")
    assert_true("--auto-repair" in command and "--repair-max-rounds 4" in command, f"auto repair options were not preserved: {command}")
    assert_true("--require-wps" in command and "--update-golden" in command, f"visual options were not preserved: {command}")
    assert_true("--golden-dir TestData/GoldenBaselines" in command, f"golden dir was not preserved: {command}")


@case
def qa_repair_plan_surfaces_next_action_and_resume_route() -> None:
    work = new_workdir("qa_repair_next_action_user_file")
    write_json(
        work / "workflow_mode.json",
        {
            "mode": "user",
            "template": "demo_template.docx",
            "content": "paper.md",
            "qa_level": "strict",
        },
    )
    user_file_report = {
        "mode": "user",
        "passed": False,
        "issues": [{"code": "CONTENT_IMAGE_MISSING", "severity": "error", "message": "missing image"}],
        "counts": {},
    }
    user_file_plan = build_repair_plan(user_file_report, str(work))
    assert_true("CONTENT_IMAGE_MISSING" in user_file_plan.get("next_action", ""), f"repair plan lost leading code: {user_file_plan}")
    assert_true("把缺失图片放回" in user_file_plan.get("next_action", ""), f"repair plan lost concrete user action: {user_file_plan}")
    assert_true(user_file_plan.get("resume_scope") == "input_files", f"user-file blocker should route to input files: {user_file_plan}")
    assert_true(
        user_file_plan.get("resume_command") == user_file_plan["commands"]["rerun_current_pipeline"],
        f"user-file blocker should resume with full pipeline rerun: {user_file_plan}",
    )
    assert_true(not user_file_plan["commands"].get("rebuild_current_docx"), f"user-file blocker should not suggest rebuild-only: {user_file_plan}")
    user_file_markdown = repair_plan_to_markdown(user_file_plan)
    assert_true("下一步" in user_file_markdown and "修复后运行" in user_file_markdown, f"repair markdown should show next action and resume command: {user_file_markdown}")

    work2 = new_workdir("qa_repair_next_action_generated_script")
    write_json(work2 / "workflow_mode.json", {"mode": "user"})
    generated_report = {
        "mode": "user",
        "passed": False,
        "issues": [{"code": "MISSING_DOCX", "severity": "error", "message": "missing docx"}],
        "counts": {},
    }
    generated_plan = build_repair_plan(generated_report, str(work2))
    assert_true("MISSING_DOCX" in generated_plan.get("next_action", ""), f"generated-script repair lost leading code: {generated_plan}")
    assert_true(generated_plan.get("resume_scope") == "current_docx", f"generated-script repair should route to current DOCX rebuild: {generated_plan}")
    assert_true(
        generated_plan.get("resume_command") == generated_plan["commands"]["rebuild_current_docx"],
        f"generated-script repair should resume with build_generated.py rebuild: {generated_plan}",
    )
    assert_true("修复后运行" in generated_plan.get("copy_to_ai_prompt", ""), f"AI prompt should include the resume command: {generated_plan}")


@case
def qa_repair_plan_warning_only_is_not_plain_pass() -> None:
    work = new_workdir("qa_repair_warning_only")
    write_json(
        work / "workflow_mode.json",
        {
            "mode": "user",
            "template": "demo_template.docx",
            "content": "paper.docx",
            "qa_level": "basic",
        },
    )
    report = {
        "mode": "user",
        "passed": True,
        "issues": [{"code": "REFERENCES_MISSING", "severity": "warning", "message": "references missing"}],
        "counts": {},
    }
    plan = build_repair_plan(report, str(work))
    markdown = repair_plan_to_markdown(plan)
    assert_true(plan["passed"] is True, f"warning-only QA should remain non-blocking: {plan}")
    assert_true(plan["warnings"] == 1, f"warning count should be visible: {plan}")
    assert_true("REFERENCES_MISSING" in plan["next_action"], f"repair plan lost warning code: {plan['next_action']}")
    assert_true("参考文献" in plan["next_action"], f"repair plan lost warning action: {plan['next_action']}")
    assert_true("警告" in plan["summary"], f"warning-only summary should not sound like a plain pass: {plan['summary']}")
    assert_true(plan["resume_scope"] == "warning_review", f"warning-only plan should route to warning review: {plan}")
    assert_true(
        plan["resume_command"] == plan["commands"]["rerun_current_pipeline"],
        f"warning-only plan should preserve a rerun command for fixes: {plan}",
    )
    assert_true("QA 已通过，仍建议" not in markdown, f"repair markdown should not hide warning-only issues: {markdown}")
    assert_true("REFERENCES_MISSING" in markdown and "警告" in markdown, f"repair markdown should surface warning details: {markdown}")


@case
def qa_report_markdown_lists_repair_plan_open_first_files() -> None:
    from qa_checker_modules.reports import report_to_markdown

    report = {
        "mode": "user",
        "passed": False,
        "output_dir_name": "demo",
        "next_action": "优先处理 `CONTENT_IMAGE_MISSING`。",
        "counts": {},
        "issues": [
            {
                "code": "CONTENT_IMAGE_MISSING",
                "severity": "error",
                "message": "missing image",
                "active_owner": "User input/template file",
            }
        ],
        "repair_plan": {
            "summary": "QA 发现 1 个阻断错误。",
            "output_dir": "Outputs/demo",
            "open_first": [
                "qa_repair_plan.md",
                "qa_report.md",
                "内容提取.md",
                "build_manifest.json",
                "最终论文.docx",
            ],
            "commands": {
                "rerun_current_pipeline": "python run_pipeline.py --mode user --template t.docx --content c.docx",
                "rebuild_current_docx": "python Outputs/demo/build_generated.py",
            },
            "steps": [],
        },
    }
    markdown = report_to_markdown(report)
    assert_true("## 先打开这些文件" in markdown, f"QA report should surface repair-plan review files: {markdown}")
    assert_true("Outputs/demo/qa_repair_plan.md" in markdown, f"QA report should point to repair plan: {markdown}")
    assert_true("Outputs/demo/内容提取.md" in markdown, f"QA report should point to content summary: {markdown}")
    assert_true("Outputs/demo/build_manifest.json" in markdown, f"QA report should point to build manifest: {markdown}")
    assert_true("Outputs/demo/最终论文.docx" in markdown, f"QA report should point to final DOCX: {markdown}")
    assert_true("## 可执行命令" in markdown, f"QA report should keep commands separate from open-first files: {markdown}")
    assert_true(markdown.index("## 先打开这些文件") < markdown.index("## 可执行命令"), f"open-first files should be listed before commands: {markdown}")


@case
def qa_report_next_action_names_first_repair_step() -> None:
    work = new_workdir("qa_next_action_first_step")
    write_json(work / "workflow_mode.json", {"mode": "user"})
    report = build_report(
        str(work),
        "user",
        {},
        [
            {
                "code": "CONTENT_IMAGE_MISSING",
                "severity": "error",
                "message": "missing image",
                "active_owner": "User input/template file",
            }
        ],
    )
    assert_true("CONTENT_IMAGE_MISSING" in report["next_action"], f"next_action lost the issue code: {report['next_action']}")
    assert_true("把缺失图片放回" in report["next_action"], f"next_action lost the beginner repair action: {report['next_action']}")
    assert_true("用户确认或补充输入文件" in report["next_action"], f"user-file routing disappeared: {report['next_action']}")


@case
def qa_report_next_action_names_warning_step() -> None:
    work = new_workdir("qa_next_action_warning_step")
    write_json(work / "workflow_mode.json", {"mode": "user"})
    report = build_report(
        str(work),
        "user",
        {},
        [
            {
                "code": "REFERENCES_MISSING",
                "severity": "warning",
                "message": "references missing",
            }
        ],
    )
    assert_true(report["passed"] is True, f"warning-only QA should remain non-blocking: {report}")
    assert_true("REFERENCES_MISSING" in report["next_action"], f"warning next_action lost the issue code: {report['next_action']}")
    assert_true("参考文献" in report["next_action"], f"warning next_action lost the beginner action: {report['next_action']}")
    assert_true("警告" in report["next_action"] or "warning" in report["next_action"], f"warning next_action should not sound like plain pass: {report['next_action']}")


@case
def qa_complex_table_warning_guides_visible_hmerge_continuations() -> None:
    work = new_workdir("qa_visible_hmerge_continuation_action")
    write_json(work / "workflow_mode.json", {"mode": "user"})
    report = build_report(
        str(work),
        "user",
        {"visible_hmerge_continuation_count": 2},
        [
            {
                "code": "COMPLEX_TABLE_UNSUPPORTED",
                "severity": "warning",
                "message": "complex table requires review",
                "detail": "irregular_hmerges=2 visible_hmerge_continuations=2",
            }
        ],
    )
    action = report["next_action"]
    plan_action = str((report.get("repair_plan") or {}).get("next_action") or "")
    step_action = str((((report.get("repair_plan") or {}).get("steps") or [{}])[0]).get("user_action") or "")
    combined = "\n".join([action, plan_action, step_action])
    assert_true("visible_hmerge_continuations" in combined, f"visible continuation detail disappeared from guidance: {combined}")
    assert_true("可见内容" in combined, f"guidance should tell users why the continuation needs review: {combined}")
    assert_true("最终 DOCX" in combined and "原文" in combined, f"guidance should tell users what to compare: {combined}")


@case
def qa_complex_table_warning_guides_visible_vmerge_continuations() -> None:
    work = new_workdir("qa_visible_vmerge_continuation_action")
    write_json(work / "workflow_mode.json", {"mode": "user"})
    report = build_report(
        str(work),
        "user",
        {"visible_vmerge_continuation_count": 1},
        [
            {
                "code": "COMPLEX_TABLE_UNSUPPORTED",
                "severity": "warning",
                "message": "complex table requires review",
                "detail": "visible_vmerge_continuations=1",
            }
        ],
    )
    action = report["next_action"]
    plan_action = str((report.get("repair_plan") or {}).get("next_action") or "")
    step_action = str((((report.get("repair_plan") or {}).get("steps") or [{}])[0]).get("user_action") or "")
    combined = "\n".join([action, plan_action, step_action])
    assert_true("visible_vmerge_continuations" in combined, f"visible vMerge continuation detail disappeared from guidance: {combined}")
    assert_true("可见内容" in combined, f"guidance should tell users why the vMerge continuation needs review: {combined}")
    assert_true("最终 DOCX" in combined and "原文" in combined, f"guidance should tell users what to compare: {combined}")


@case
def qa_json_reports_expose_explicit_status_labels() -> None:
    work = new_workdir("qa_json_status_labels")
    write_json(work / "workflow_mode.json", {"mode": "developer"})
    failed = build_report(
        str(work),
        "developer",
        {},
        [{"code": "CONTENT_IMAGE_MISSING", "severity": "error", "message": "missing image"}],
    )
    assert_true(failed["status"] == "failed", f"structural error status should be failed: {failed}")
    assert_true(failed["result_label"] == "未通过", f"structural error label should be explicit: {failed}")

    warning = build_report(
        str(work),
        "developer",
        {},
        [{"code": "REFERENCES_MISSING", "severity": "warning", "message": "missing refs"}],
    )
    assert_true(warning["status"] == "passed_with_warnings", f"structural warning status should be explicit: {warning}")
    assert_true(warning["result_label"] == "通过但有警告", f"structural warning label should be explicit: {warning}")

    clean = build_report(str(work), "developer", {}, [])
    assert_true(clean["status"] == "passed", f"structural clean status should be explicit: {clean}")
    assert_true(clean["result_label"] == "通过", f"structural clean label should be explicit: {clean}")

    conformance = build_conformance_report(
        str(work),
        "developer",
        {},
        [{"code": "STYLE_MISMATCH", "severity": "warning", "message": "style differs"}],
        project_root=str(work),
    )
    assert_true(conformance["status"] == "passed_with_warnings", f"strict QA status should expose warnings: {conformance}")
    assert_true(conformance["result_label"] == "通过但有警告", f"strict QA label should expose warnings: {conformance}")

    visual = check_visual(str(work), output_docx_name="missing.docx", project_root=str(work))
    assert_true(visual["status"] == "failed", f"visual missing-DOCX status should fail explicitly: {visual}")
    assert_true(visual["result_label"] == "未通过", f"visual missing-DOCX label should fail explicitly: {visual}")


@case
def qa_routes_user_file_errors_to_input_fix() -> None:
    work = new_workdir("qa_user_file_routing")
    write_json(
        work / "format.json",
        {
            "_meta": {
                "source": "blank_scan.pdf",
                "pdf_template": {
                    "type": "scanned_or_unsupported_pdf",
                    "errors": ["PDF_TEMPLATE_NO_TEXT"],
                    "confidence": 0.0,
                    "text_chars": 0,
                },
            },
            "paragraphs": [],
            "tables": [],
            "sections": [{"page_width_cm": 21.0, "page_height_cm": 29.7}],
            "cover": [],
            "style_profiles": {},
        },
    )
    write_json(work / "content.json", base_content(["Body text"]))
    write_json(work / "workflow_mode.json", {"mode": "user"})
    (work / "build_generated.py").write_text("# synthetic\n", encoding="utf-8")
    doc = Document()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Body text")
    doc.save(work / "out.docx")
    write_json(work / "build_manifest.json", {"schema_version": 1, "counts": {}})
    report = check_output(str(work), mode="user", output_docx_name="out.docx")
    issue = next(item for item in report["issues"] if item["code"] == "PDF_TEMPLATE_UNSUPPORTED")
    assert_true(issue["active_owner"] == "User input/template file", f"PDF template issue target was misleading: {issue}")
    assert_true("用户确认或补充输入文件" in report["next_action"], f"next action did not route to user file fix: {report['next_action']}")
    step = next(item for item in report["repair_plan"]["steps"] if item["code"] == "PDF_TEMPLATE_UNSUPPORTED")
    assert_true(step["target"] == "User input/template file", f"repair target was misleading: {step}")
    assert_true(not report["repair_plan"]["commands"].get("rebuild_current_docx"), "user-file-only error should not suggest rebuilding build_generated.py")
    assert_true("不要只修改 `build_generated.py`" in report["repair_plan"]["copy_to_ai_prompt"], "AI prompt should not route user-file-only errors to generated-script edits")

    work2 = new_workdir("qa_user_confirmation_routing")
    content = base_content(["Body text"])
    content["_meta"]["non_body_images"] = [{"location": "section_1_header", "target": "media/image1.png"}]
    write_json(work2 / "content.json", content)
    write_json(work2 / "format.json", base_format())
    write_json(work2 / "workflow_mode.json", {"mode": "user"})
    (work2 / "build_generated.py").write_text("# synthetic\n", encoding="utf-8")
    doc2 = Document()
    doc2.add_paragraph("1 Introduction")
    doc2.add_paragraph("Body text")
    doc2.save(work2 / "out.docx")
    write_json(work2 / "build_manifest.json", {"schema_version": 1, "counts": {}})
    report2 = check_output(str(work2), mode="user", output_docx_name="out.docx")
    issue2 = next(item for item in report2["issues"] if item["code"] == "NON_BODY_IMAGE_UNSUPPORTED")
    assert_true(issue2["active_owner"] == "User input/template file", f"user-confirmation issue target was misleading: {issue2}")
