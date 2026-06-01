"""Pipeline runner regression cases."""
from __future__ import annotations

import json
import io
import sys
import builtins
from contextlib import redirect_stdout
from pathlib import Path
from types import SimpleNamespace
from typing import Any, Dict, List

from docx import Document
from pipeline_runner.artifacts import (
    build_content_markdown,
    write_build_failure_report,
    write_content_artifacts,
    write_extraction_failure_report,
    write_format_artifacts,
)
from pipeline_runner.build_phase import generate_and_build_docx_phase
from pipeline_runner.cli import DEFAULT_GOLDEN_DIR, build_arg_parser, dispatch_cli
from pipeline_runner.contracts import (
    has_contract_errors,
    validate_build_manifest,
    validate_content_data,
    validate_format_data,
    validate_qa_report,
)
from pipeline_runner.context import (
    create_unique_output_dir,
    normalize_qa_level,
    resolve_inputs,
    write_workflow_mode,
)
from pipeline_runner.dependencies import load_optional_dependencies
from pipeline_runner.execution import ScriptExecutionResult, run_generated_script
from pipeline_runner.io import choose_file, choose_mode, normalize_mode, scan_inputs
from pipeline_runner.qa import QADependencies, run_qa_phases
from pipeline_runner.repair_loop import run_repair_loop
from pipeline_runner.reports import qa_status_fields
from pipeline_runner.summary import build_completion_summary, write_agent_preflight_report, write_agent_summary
from pipeline_runner.template_phase import write_template_profile_phase, write_template_requirements_phase
from pipeline_runner.verification import VerificationError, _merge_content_results, double_verify

from regression_suite_modules.generated_pdf import poppler_available, write_blank_pdf, write_text_pdf
from regression_suite_modules.harness import (
    PNG_1X1,
    assert_true,
    base_content,
    base_format,
    case,
    fail,
    new_workdir,
    write_json,
)

PIPELINE_DIR = Path(__file__).resolve().parents[1]
REPO_ROOT = Path(__file__).resolve().parents[4]


def _write_repair_loop_fixture(work: Path, *, final_docx: str = "final.docx") -> Path:
    work.mkdir(exist_ok=True)
    (work / "format.json").write_text(json.dumps(base_format(), ensure_ascii=False, indent=2), encoding="utf-8")
    (work / "content.json").write_text(
        json.dumps(base_content(["This body paragraph should remain in the final document."]), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    write_workflow_mode(
        str(work),
        mode="user",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="basic",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
        auto_repair=True,
        repair_max_rounds=5,
        repair_stop_no_improve=2,
    )
    build_script = f"""from docx import Document
import json
import os

BASE = os.path.dirname(__file__)
OUT = os.path.join(BASE, {final_docx!r})

def main():
    doc = Document()
    doc.add_paragraph('1 Introduction')
    doc.add_paragraph('This body paragraph should remain in the final document.')
    doc.add_paragraph('TODO')
    doc.save(OUT)
    with open(os.path.join(BASE, 'build_manifest.json'), 'w', encoding='utf-8') as handle:
        json.dump({{'schema_version': 1, 'counts': {{'content_images_rendered': 0, 'content_tables_rendered': 0, 'content_formulas_rendered': 0}}}}, handle)

if __name__ == '__main__':
    main()
"""
    path = work / "build_generated.py"
    path.write_text(build_script, encoding="utf-8")
    return path


def _fake_repair_report(out_dir: str, *, code: str, severity: str = "error", message: str = "synthetic") -> Dict[str, Any]:
    from qa_checker_modules.repair import build_repair_plan
    from qa_checker_modules.reports import write_reports

    report = {
        "schema_version": 1,
        "mode": "user",
        "passed": severity != "error",
        "counts": {},
        "issues": [{"code": code, "severity": severity, "message": message, "detail": ""}],
        "next_action": "repair",
        "output_dir_name": Path(out_dir).name,
    }
    report.update(qa_status_fields(report["passed"], report["issues"]))
    report["repair_plan"] = build_repair_plan(report, out_dir)
    write_reports(report, out_dir)
    return report


def _write_conformance_report(out_dir: str, *, passed: bool) -> Dict[str, Any]:
    issues = [] if passed else [
        {
            "code": "STYLE_MISMATCH",
            "severity": "error",
            "message": "Some final DOCX paragraphs do not match template role styles.",
            "detail": "reference `[1] 作者1.`: eastAsia font Times New Roman != 宋体",
        }
    ]
    report = {
        "schema_version": 1,
        "mode": "user",
        "passed": passed,
        "counts": {},
        "issues": issues,
        "next_action": "Fix Outputs/<run>/build_generated.py and rerun it.",
    }
    report.update(qa_status_fields(report["passed"], report["issues"]))
    Path(out_dir, "conformance_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


@case
def pipeline_contracts_accept_current_handoffs() -> None:
    manifest = {
        "schema_version": 1,
        "counts": {
            "content_images_rendered": 0,
            "content_tables_rendered": 1,
            "content_formulas_rendered": 2,
        },
    }
    qa_report = {
        "schema_version": 1,
        "mode": "developer",
        "passed": True,
        "status": "passed",
        "result_label": "通过",
        "counts": {},
        "issues": [],
        "next_action": "ok",
    }
    all_issues = (
        validate_format_data(base_format())
        + validate_content_data(base_content(["Body text"]))
        + validate_build_manifest(manifest)
        + validate_qa_report(qa_report)
    )
    assert_true(not has_contract_errors(all_issues), "valid current handoff structures should satisfy contracts")


@case
def pipeline_contracts_report_structural_errors() -> None:
    issues = (
        validate_format_data({"paragraphs": {}})
        + validate_content_data({"sections": {}})
        + validate_build_manifest({"counts": {"content_images_rendered": -1}})
        + validate_qa_report({"passed": "yes", "status": "maybe", "result_label": "", "counts": [], "issues": [{"code": ""}]})
        + validate_qa_report({"passed": False, "counts": {}, "issues": []})
        + validate_qa_report({
            "passed": True,
            "status": "passed",
            "result_label": "通过",
            "counts": {},
            "issues": [{"code": "WARN", "severity": "warning", "message": "review"}],
        })
    )
    codes = {issue.code for issue in issues}
    assert_true("FORMAT_SECTIONS_MISSING" in codes, "missing format sections was not reported")
    assert_true("FORMAT_PARAGRAPHS_NOT_LIST" in codes, "invalid format paragraphs was not reported")
    assert_true("CONTENT_SECTIONS_NOT_LIST" in codes, "invalid content sections was not reported")
    assert_true("MANIFEST_COUNT_INVALID" in codes, "invalid manifest count was not reported")
    assert_true("QA_REPORT_PASSED_NOT_BOOL" in codes, "invalid QA passed flag was not reported")
    assert_true("QA_REPORT_STATUS_MISSING" in codes, "missing QA status was not reported")
    assert_true("QA_REPORT_STATUS_INVALID" in codes, "invalid QA status was not reported")
    assert_true("QA_REPORT_RESULT_LABEL_MISSING" in codes, "missing QA result label was not reported")
    assert_true("QA_REPORT_RESULT_LABEL_EMPTY" in codes, "invalid QA result label was not reported")
    assert_true("QA_REPORT_STATUS_MISMATCH" in codes, "mismatched QA status was not reported")
    assert_true("QA_REPORT_RESULT_LABEL_MISMATCH" in codes, "mismatched QA result label was not reported")
    assert_true("QA_REPORT_ISSUE_CODE_MISSING" in codes, "invalid QA issue code was not reported")
    assert_true(has_contract_errors(issues), "structural contract errors should be marked as errors")


@case
def pipeline_io_helpers_scan_inputs_and_normalize_modes() -> None:
    work = new_workdir("pipeline_io")
    (work / "paper.docx").write_text("x", encoding="utf-8")
    (work / "notes.md").write_text("x", encoding="utf-8")
    (work / "template.PDF").write_text("x", encoding="utf-8")
    (work / "~$paper.docx").write_text("x", encoding="utf-8")
    (work / "ignore.txt").write_text("x", encoding="utf-8")
    assert_true(scan_inputs(str(work)) == ["notes.md", "paper.docx"], "scan_inputs should skip temp and unsupported files")
    assert_true(scan_inputs(str(work), exts=(".docx", ".pdf")) == ["paper.docx", "template.PDF"], "scan_inputs should respect extension filters case-insensitively")
    assert_true(normalize_mode("developer") == "developer", "developer mode should be preserved")
    assert_true(normalize_mode("bad-mode") == "user", "unknown mode should fall back to user")


@case
def pipeline_interactive_interrupts_show_next_steps() -> None:
    original_input = builtins.input

    def assert_interrupt_guidance(func, input_exc):
        def raising_input(_prompt=""):
            raise input_exc

        builtins.input = raising_input
        buffer = io.StringIO()
        try:
            with redirect_stdout(buffer):
                try:
                    func()
                except SystemExit as exc:
                    assert_true(exc.code == 1, f"interactive cancel exit code should be 1, got {exc.code}")
                else:
                    fail("interactive cancel did not exit")
        finally:
            builtins.input = original_input
        text = buffer.getvalue()
        assert_true("下一步" in text and "--agent-auto" in text, f"cancel guidance missing next step: {text}")

    assert_interrupt_guidance(lambda: choose_file(["a.docx", "b.docx"], "模板"), KeyboardInterrupt())
    assert_interrupt_guidance(lambda: choose_mode(default="user"), EOFError())


@case
def pipeline_cli_dispatches_md_parameter_and_single_file_modes() -> None:
    parser = build_arg_parser()
    calls: List[Dict[str, Any]] = []

    def fake_run(template_file, content_file, **kwargs):
        calls.append({"template": template_file, "content": content_file, **kwargs})
        return "ok"

    def expect_exit_zero(args, template_dir="", inputs_dir=""):
        try:
            dispatch_cli(args, run_pipeline=fake_run, template_dir=template_dir, inputs_dir=inputs_dir)
        except SystemExit as exc:
            assert_true(exc.code == 0, f"dispatch returned nonzero exit: {exc.code}")
            return
        fail("dispatch should exit through exit_from_result")

    expect_exit_zero(parser.parse_args(["--md", "paper.md", "--mode", "developer", "--qa-level", "basic", "--no-qa"]))
    assert_true(calls[-1]["template"] is None and calls[-1]["content"] is None, "MD mode should not require template/content")
    assert_true(calls[-1]["md_file"] == "paper.md", "MD mode did not pass md_file")
    assert_true(calls[-1]["mode"] == "developer" and calls[-1]["run_qa"] is False, "MD mode options were not preserved")

    expect_exit_zero(parser.parse_args(["--template", "t.docx", "--content", "c.docx"]))
    assert_true(calls[-1]["template"] == "t.docx" and calls[-1]["content"] == "c.docx", "parameter mode files changed")
    assert_true(calls[-1]["mode"] == "user", "non-interactive auto mode should default to user")

    expect_exit_zero(parser.parse_args(["--template", "t.docx", "--content", "c.docx", "--auto-repair", "--no-qa", "--repair-max-rounds", "4"]))
    assert_true(calls[-1]["run_qa"] is True and calls[-1]["auto_repair"] is True, "auto repair should force QA on")
    assert_true(calls[-1]["repair_max_rounds"] == 4, "auto repair max rounds option was not passed")
    assert_true(calls[-1]["golden_dir"] is None, "visual golden baseline should be opt-in by default")

    work = new_workdir("pipeline_cli")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    (template_dir / "only.pdf").write_bytes(b"")
    (inputs_dir / "only.md").write_text("# Paper", encoding="utf-8")
    expect_exit_zero(parser.parse_args(["--mode", "user"]), template_dir=str(template_dir), inputs_dir=str(inputs_dir))
    assert_true(calls[-1]["template"] == "only.pdf" and calls[-1]["content"] == "only.md", "single-file interactive dispatch changed")

    expect_exit_zero(parser.parse_args(["--template", "given.docx", "--mode", "user"]), template_dir=str(template_dir), inputs_dir=str(inputs_dir))
    assert_true(calls[-1]["template"] == "given.docx" and calls[-1]["content"] == "only.md", "partial template argument should still auto-select content")
    expect_exit_zero(parser.parse_args(["--content", "given.md", "--mode", "user"]), template_dir=str(template_dir), inputs_dir=str(inputs_dir))
    assert_true(calls[-1]["template"] == "only.pdf" and calls[-1]["content"] == "given.md", "partial content argument should still auto-select template")

    expect_exit_zero(parser.parse_args(["--template", "t.docx", "--content", "c.docx", "--qa-level", "visual", "--update-golden"]))
    assert_true(calls[-1]["golden_dir"] == DEFAULT_GOLDEN_DIR, "update-golden without explicit dir should use default baseline directory")

    expect_exit_zero(parser.parse_args(["--agent-auto", "--mode", "auto", "--no-qa"]), template_dir=str(template_dir), inputs_dir=str(inputs_dir))
    assert_true(calls[-1]["template"] == "only.pdf" and calls[-1]["content"] == "only.md", "agent-auto should non-interactively select single template/content")
    assert_true(calls[-1]["mode"] == "user", "agent-auto auto mode should default to user")
    assert_true(calls[-1]["run_qa"] is True and calls[-1]["auto_repair"] is True, "agent-auto should force QA and user auto-repair")
    assert_true(calls[-1]["agent_auto"] is True, "agent-auto flag was not passed to the pipeline")

    md_only = new_workdir("pipeline_cli_agent_md_only")
    md_templates = md_only / "Templates"
    md_inputs = md_only / "Inputs"
    md_templates.mkdir()
    md_inputs.mkdir()
    (md_inputs / "paper.md").write_text("# 格式说明\n\n# 正文", encoding="utf-8")
    expect_exit_zero(parser.parse_args(["--agent-auto"]), template_dir=str(md_templates), inputs_dir=str(md_inputs))
    assert_true(calls[-1]["template"] is None and calls[-1]["content"] is None, "agent-auto pure MD should not pass template/content")
    assert_true(calls[-1]["md_file"] == "paper.md", "agent-auto should use the single MD file as pure Markdown mode when no template exists")

    (md_inputs / "second.md").write_text("# Another", encoding="utf-8")
    try:
        dispatch_cli(parser.parse_args(["--agent-auto"]), run_pipeline=fake_run, template_dir=str(md_templates), inputs_dir=str(md_inputs))
    except SystemExit as exc:
        assert_true(exc.code == 2, f"agent-auto multiple MD candidates should ask for selection, got exit {exc.code}")
        report_path = md_only / "Outputs" / "_agent_preflight_latest" / "agent_preflight_report.json"
        report = json.loads(report_path.read_text(encoding="utf-8"))
        md_steps = "\n".join(report.get("next_steps") or [])
        assert_true("使用 Inputs/paper.md 作为纯 Markdown 输入" in md_steps, f"pure MD preflight should offer direct agent replies: {report}")
        assert_true("使用 Inputs/second.md 作为纯 Markdown 输入" in md_steps, f"pure MD preflight should include every candidate: {report}")
    else:
        fail("agent-auto should not blind-pick among multiple pure MD candidates")


@case
def pipeline_agent_auto_guides_missing_and_ambiguous_inputs() -> None:
    parser = build_arg_parser()
    calls: List[Dict[str, Any]] = []

    def fake_run(template_file, content_file, **kwargs):
        calls.append({"template": template_file, "content": content_file, **kwargs})
        return "ok"

    def expect_exit(args, template_dir, inputs_dir, expected_code, expected_text):
        buf = io.StringIO()
        with redirect_stdout(buf):
            try:
                dispatch_cli(args, run_pipeline=fake_run, template_dir=str(template_dir), inputs_dir=str(inputs_dir))
            except SystemExit as exc:
                assert_true(exc.code == expected_code, f"unexpected exit {exc.code}, output={buf.getvalue()}")
                assert_true(expected_text in buf.getvalue(), f"agent-auto output missing guidance {expected_text!r}: {buf.getvalue()}")
                report_path = Path(template_dir).parent / "Outputs" / "_agent_preflight_latest" / "agent_preflight_report.json"
                assert_true(report_path.exists(), f"agent-auto preflight report missing: {report_path}")
                report = json.loads(report_path.read_text(encoding="utf-8"))
                assert_true(report.get("next_steps"), f"preflight report missing next steps: {report}")
                return report
        fail("agent-auto should exit instead of silently continuing")

    empty = new_workdir("pipeline_agent_auto_empty")
    empty_templates = empty / "Templates"
    empty_inputs = empty / "Inputs"
    empty_templates.mkdir()
    empty_inputs.mkdir()
    expect_exit(parser.parse_args(["--agent-auto"]), empty_templates, empty_inputs, 1, "没有找到可运行的模板和内容")

    missing_content = new_workdir("pipeline_agent_auto_missing_content")
    templates = missing_content / "Templates"
    inputs = missing_content / "Inputs"
    templates.mkdir()
    inputs.mkdir()
    (templates / "one.docx").write_bytes(b"")
    expect_exit(parser.parse_args(["--agent-auto"]), templates, inputs, 1, "请把模板 DOCX/PDF 放入 Templates/")

    ambiguous_template = new_workdir("pipeline_agent_auto_multi_template")
    templates = ambiguous_template / "Templates"
    inputs = ambiguous_template / "Inputs"
    templates.mkdir()
    inputs.mkdir()
    (templates / "a.docx").write_bytes(b"")
    (templates / "b.pdf").write_bytes(b"")
    (inputs / "paper.docx").write_bytes(b"")
    template_report = expect_exit(parser.parse_args(["--agent-auto"]), templates, inputs, 2, "模板存在多个候选")
    template_steps = "\n".join(template_report["next_steps"])
    assert_true("作为模板" in template_steps, f"template ambiguity next step is unclear: {template_report}")
    assert_true("使用 Templates/a.docx 作为模板" in template_steps, f"template ambiguity should offer a copyable Agent reply: {template_report}")
    assert_true("使用 Templates/b.pdf 作为模板" in template_steps, f"template ambiguity should include every candidate: {template_report}")

    ambiguous_content = new_workdir("pipeline_agent_auto_multi_content")
    templates = ambiguous_content / "Templates"
    inputs = ambiguous_content / "Inputs"
    templates.mkdir()
    inputs.mkdir()
    (templates / "template.docx").write_bytes(b"")
    (inputs / "a.docx").write_bytes(b"")
    (inputs / "b.md").write_text("# Paper", encoding="utf-8")
    content_report = expect_exit(parser.parse_args(["--agent-auto"]), templates, inputs, 2, "内容存在多个候选")
    content_steps = "\n".join(content_report["next_steps"])
    assert_true("作为内容" in content_steps, f"content ambiguity next step is unclear: {content_report}")
    assert_true("作为模板" not in content_steps, f"content ambiguity should not mention template selection: {content_report}")
    assert_true("使用 Inputs/a.docx 作为内容" in content_steps, f"content ambiguity should offer a copyable Agent reply: {content_report}")
    assert_true("使用 Inputs/b.md 作为内容" in content_steps, f"content ambiguity should include every candidate: {content_report}")
    assert_true(not calls, "missing/ambiguous agent-auto inputs should not call the pipeline")


@case
def pipeline_agent_preflight_report_lists_source_folders_and_formats() -> None:
    work = new_workdir("pipeline_agent_preflight_source_folders")
    json_path, md_path = write_agent_preflight_report(
        str(work / "Outputs"),
        status="blocked_missing_input",
        message="没有找到可运行的模板和内容。",
        next_steps=[
            "把模板 DOCX/PDF 放入 Templates/，把内容 DOCX 或 Markdown 放入 Inputs/。",
            "放好后让 Agent 重新运行自动入口。",
        ],
        candidates={"Templates": ["a.docx"], "Inputs": ["paper.docx", "paper.md"]},
    )
    report = json.loads(Path(json_path).read_text(encoding="utf-8"))
    folders = {item.get("folder"): item for item in report.get("source_folders") or []}
    assert_true("Templates" in folders, f"preflight JSON should name where templates go: {report}")
    assert_true("Inputs" in folders, f"preflight JSON should name where content goes: {report}")
    assert_true({".docx", ".pdf"}.issubset(set(folders["Templates"].get("accepted_extensions") or [])), f"template formats missing: {report}")
    assert_true({".docx", ".md"}.issubset(set(folders["Inputs"].get("accepted_extensions") or [])), f"content formats missing: {report}")
    text = Path(md_path).read_text(encoding="utf-8")
    assert_true("## 文件应该放哪里" in text, f"preflight markdown should have a stable source-folder section: {text}")
    assert_true("Templates/" in text and ".docx" in text and ".pdf" in text, f"template source-folder guidance missing: {text}")
    assert_true("Inputs/" in text and ".docx" in text and ".md" in text, f"input source-folder guidance missing: {text}")
    assert_true("python run_pipeline.py --agent-auto" in text, f"preflight markdown should give a rerun route: {text}")


@case
def pipeline_context_resolves_inputs_outputs_and_workflow() -> None:
    work = new_workdir("pipeline_context")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    outputs_dir = work / "Outputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    outputs_dir.mkdir()
    (template_dir / "template.PDF").write_text("x", encoding="utf-8")
    (inputs_dir / "paper.md").write_text("# Paper", encoding="utf-8")
    nested_dir = inputs_dir / "nested"
    nested_dir.mkdir()
    (nested_dir / "nested_paper.md").write_text("# Nested Paper", encoding="utf-8")

    resolution = resolve_inputs("template.PDF", "paper.md", None, str(template_dir), str(inputs_dir))
    assert_true(resolution.ok, f"expected inputs to resolve: {resolution.error}")
    assert_true(resolution.inputs.use_md_content is True, "markdown content flag was not set")
    assert_true(resolution.inputs.use_md_format is False, "PDF template should use format_extractor, not md_parser")
    assert_true(resolution.inputs.content_name == "paper", "content stem was not preserved")
    missing = resolve_inputs("missing.docx", "paper.md", None, str(template_dir), str(inputs_dir))
    assert_true(str(template_dir) not in str(missing.error), f"missing template error leaked absolute path: {missing.error}")
    assert_true("Templates/missing.docx" in str(missing.error), f"missing template error lost next-step location: {missing.error}")
    missing_abs = resolve_inputs(str(template_dir / "missing_abs.docx"), "paper.md", None, str(template_dir), str(inputs_dir))
    assert_true(str(template_dir) not in str(missing_abs.error), f"absolute missing template error leaked path: {missing_abs.error}")
    assert_true("missing_abs.docx" in str(missing_abs.error), f"absolute missing template error lost filename: {missing_abs.error}")
    assert_true(normalize_qa_level("bad-level") == "strict", "invalid QA level should fall back to strict")

    first_dir, first_name = create_unique_output_dir(str(outputs_dir), "paper", today="2026-05-26")
    second_dir, second_name = create_unique_output_dir(str(outputs_dir), "paper", today="2026-05-26")
    assert_true(first_name == "2026-05-26_paper", "first output folder name changed")
    assert_true(second_name == "2026-05-26_paper_2", "duplicate output folder suffix changed")
    workflow_path = write_workflow_mode(
        first_dir,
        mode="developer",
        template_path=resolution.inputs.template_path,
        content_path=resolution.inputs.content_path,
        run_qa=True,
        qa_level="visual",
        golden_dir="Golden",
        update_golden=False,
        require_wps=True,
    )
    workflow = json.loads(Path(workflow_path).read_text(encoding="utf-8"))
    assert_true(workflow["mode"] == "developer", "workflow mode was not written")
    assert_true(workflow["qa_level"] == "visual", "workflow QA level was not written")
    assert_true(workflow["require_wps"] is True, "workflow require_wps flag was not written")
    assert_true(workflow["auto_repair"] is False, "workflow auto_repair default changed")

    nested = resolve_inputs("template.PDF", "nested/nested_paper.md", None, str(template_dir), str(inputs_dir))
    assert_true(nested.ok, f"expected nested content to resolve: {nested.error}")
    nested_workflow_path = write_workflow_mode(
        second_dir,
        mode="developer",
        template_path=nested.inputs.template_path,
        content_path=nested.inputs.content_path,
        run_qa=True,
        qa_level="strict",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
    )
    nested_workflow = json.loads(Path(nested_workflow_path).read_text(encoding="utf-8"))
    assert_true(nested_workflow["content"] == "nested/nested_paper.md", f"workflow lost nested input path: {nested_workflow}")
    assert_true(str(inputs_dir) not in nested_workflow["content"], f"workflow leaked absolute input path: {nested_workflow}")


@case
def pipeline_workflow_external_input_avoids_fake_rerun_command() -> None:
    from qa_checker_modules.repair import build_repair_plan
    from qa_checker_modules.report_phase import build_report

    work = new_workdir("pipeline_workflow_external_input")
    outputs_dir = work / "Outputs"
    external_dir = work / "External Files"
    outputs_dir.mkdir()
    external_dir.mkdir()
    external_md = external_dir / "external paper.md"
    external_md.write_text("# External Paper\n\n![missing](missing.png)", encoding="utf-8")

    workflow_path = write_workflow_mode(
        str(outputs_dir),
        mode="user",
        template_path=str(external_md),
        content_path=str(external_md),
        run_qa=True,
        qa_level="strict",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
    )
    workflow = json.loads(Path(workflow_path).read_text(encoding="utf-8"))
    workflow_text = json.dumps(workflow, ensure_ascii=False)
    assert_true(workflow.get("md") == "", f"external MD should not become a fake basename rerun arg: {workflow}")
    assert_true(workflow.get("input_location_warnings"), f"external input workflow should record beginner guidance: {workflow}")
    assert_true(str(external_dir) not in workflow_text, f"workflow leaked an absolute external path: {workflow}")

    external_inputs_dir = work / "Other Project" / "Inputs"
    external_inputs_dir.mkdir(parents=True)
    lookalike_md = external_inputs_dir / "lookalike.md"
    lookalike_md.write_text("# Lookalike Paper\n\n![missing](missing.png)", encoding="utf-8")
    lookalike_workflow_path = write_workflow_mode(
        str(outputs_dir),
        mode="user",
        template_path=str(lookalike_md),
        content_path=str(lookalike_md),
        run_qa=True,
        qa_level="strict",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
    )
    lookalike_workflow = json.loads(Path(lookalike_workflow_path).read_text(encoding="utf-8"))
    lookalike_text = json.dumps(lookalike_workflow, ensure_ascii=False)
    assert_true(lookalike_workflow.get("md") == "", f"external lookalike Inputs folder should not become a rerun arg: {lookalike_workflow}")
    assert_true(lookalike_workflow.get("input_location_warnings"), f"lookalike external Inputs folder should still guide users: {lookalike_workflow}")
    assert_true(str(external_inputs_dir.parent) not in lookalike_text, f"lookalike workflow leaked an absolute external path: {lookalike_workflow}")

    report = {
        "mode": "user",
        "passed": False,
        "issues": [{"code": "CONTENT_IMAGE_MISSING", "severity": "error", "message": "missing image"}],
        "counts": {},
    }
    plan = build_repair_plan(report, str(outputs_dir))
    assert_true(not plan["commands"].get("rerun_current_pipeline"), f"repair plan should not offer a non-runnable basename command: {plan}")
    assert_true(not plan.get("resume_command"), f"input-location blocker should omit fake resume command: {plan}")
    assert_true("Inputs/" in plan.get("next_action", "") and "放入" in plan.get("next_action", ""), f"repair plan should tell users where to put the file: {plan}")

    qa_report = build_report(
        str(outputs_dir),
        "user",
        {},
        [{"code": "CONTENT_IMAGE_MISSING", "severity": "error", "message": "missing image"}],
    )
    assert_true("Inputs/" in qa_report.get("next_action", "") and "放入" in qa_report.get("next_action", ""), f"qa_report next_action should carry input-location guidance: {qa_report}")

    write_json(outputs_dir / "qa_report.json", qa_report)
    summary_json, _summary_md = write_agent_summary(
        str(outputs_dir),
        outputs_dir.name,
        "最终论文.docx",
        "user",
        pipeline_status="failed",
    )
    summary = json.loads(Path(summary_json).read_text(encoding="utf-8"))
    action_text = "\n".join(summary.get("next_actions") or summary.get("manual_check_required") or [])
    assert_true("Inputs/" in action_text and "放入" in action_text, f"agent summary should surface input-location guidance: {summary}")


@case
def pipeline_auto_repair_patches_build_script_and_reruns_qa() -> None:
    from qa_checker import check_and_write as qa_check_and_write

    work = new_workdir("pipeline_auto_repair_placeholder")
    private_dir = work / "Inputs"
    private_dir.mkdir()
    private_file = private_dir / "private.docx"
    private_file.write_bytes(b"private-user-content")
    private_before = private_file.read_bytes()
    core_file = PIPELINE_DIR / "pipeline_runner" / "qa.py"
    core_before = core_file.read_bytes()

    build_path = _write_repair_loop_fixture(work)
    first_build = run_generated_script(str(build_path), str(work), python_executable=sys.executable)
    assert_true(first_build.returncode == 0, f"synthetic build failed: {first_build.stderr}")

    deps = QADependencies(
        qa_check_and_write=qa_check_and_write,
        conformance_check_and_write=None,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="basic",
        project_root=str(REPO_ROOT),
        max_rounds=3,
        stop_no_improve=2,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )

    assert_true(result.ok, f"auto repair did not converge: {result.status}")
    assert_true((work / "repair_loop_report.json").exists(), "repair loop JSON report missing")
    assert_true((work / "repair_loop_report.md").exists(), "repair loop markdown report missing")
    assert_true("AUTO_REPAIR_PLACEHOLDER_CLEANUP_V1" in build_path.read_text(encoding="utf-8"), "build script was not patched")
    import zipfile

    with zipfile.ZipFile(work / "final.docx") as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    assert_true("TODO" not in xml, "placeholder text remained after auto repair")
    qa = json.loads((work / "qa_report.json").read_text(encoding="utf-8"))
    assert_true(qa.get("passed") is True, f"QA was not rerun to passing state: {qa.get('issues')}")
    loop_report = json.loads((work / "repair_loop_report.json").read_text(encoding="utf-8"))
    assert_true(loop_report["rounds_run"] >= 1, "repair loop did not record a repair round")
    assert_true(
        "查看 qa_report.md 和 repair_loop_report.md 中的剩余 warning，确认不会影响交付。" not in loop_report.get("manual_check_required", []),
        "zero-warning repair report should not ask users to review remaining warnings",
    )
    loop_text = (work / "repair_loop_report.md").read_text(encoding="utf-8")
    assert_true(
        "交付前查看 qa_report.md" not in loop_text,
        "zero-warning repair markdown should not tell users to inspect remaining warnings",
    )
    assert_true(private_file.read_bytes() == private_before, "auto repair modified a private input file")
    assert_true(core_file.read_bytes() == core_before, "auto repair modified a core engine file")


@case
def pipeline_fallback_qa_reports_expose_explicit_status_labels() -> None:
    work = new_workdir("pipeline_fallback_status_structural")
    write_workflow_mode(
        str(work),
        mode="user",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="basic",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
    )
    result = run_qa_phases(
        str(work),
        mode="user",
        output_docx_name="missing.docx",
        qa_level="basic",
        project_root=str(work),
        deps=QADependencies(
            qa_check_and_write=None,
            conformance_check_and_write=None,
            visual_check_and_write=None,
            optional_import_detail=lambda name: "synthetic missing dependency",
        ),
    )
    assert_true(result is False, "missing structural QA dependency should stop the pipeline")
    structural = json.loads((work / "qa_report.json").read_text(encoding="utf-8"))
    assert_true(structural["status"] == "failed", f"structural dependency report should expose failed status: {structural}")
    assert_true(structural["result_label"] == "未通过", f"structural dependency report should expose result label: {structural}")

    work2 = new_workdir("pipeline_fallback_status_conformance")
    write_workflow_mode(
        str(work2),
        mode="user",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="strict",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
    )
    result2 = run_qa_phases(
        str(work2),
        mode="user",
        output_docx_name="missing.docx",
        qa_level="strict",
        project_root=str(work2),
        deps=QADependencies(
            qa_check_and_write=lambda out_dir, mode="user", output_docx_name="missing.docx": _fake_repair_report(out_dir, code="NO_ISSUE", severity="info"),
            conformance_check_and_write=None,
            visual_check_and_write=None,
            optional_import_detail=lambda name: "synthetic conformance dependency",
        ),
    )
    assert_true(result2 is False, "missing conformance QA dependency should stop the pipeline")
    conformance = json.loads((work2 / "conformance_report.json").read_text(encoding="utf-8"))
    assert_true(conformance["status"] == "failed", f"conformance dependency report should expose failed status: {conformance}")
    assert_true(conformance["result_label"] == "未通过", f"conformance dependency report should expose result label: {conformance}")

    work3 = new_workdir("pipeline_fallback_status_visual")
    write_workflow_mode(
        str(work3),
        mode="user",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="visual",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
    )
    result3 = run_qa_phases(
        str(work3),
        mode="user",
        output_docx_name="missing.docx",
        qa_level="visual",
        project_root=str(work3),
        deps=QADependencies(
            qa_check_and_write=lambda out_dir, mode="user", output_docx_name="missing.docx": _fake_repair_report(out_dir, code="NO_ISSUE", severity="info"),
            conformance_check_and_write=lambda out_dir, mode="user", output_docx_name="missing.docx", project_root="": _write_conformance_report(out_dir, passed=True),
            visual_check_and_write=None,
            optional_import_detail=lambda name: "synthetic visual dependency",
        ),
    )
    assert_true(result3 is False, "missing visual QA dependency should stop the pipeline")
    visual = json.loads((work3 / "visual_report.json").read_text(encoding="utf-8"))
    assert_true(visual["status"] == "failed", f"visual dependency report should expose failed status: {visual}")
    assert_true(visual["result_label"] == "未通过", f"visual dependency report should expose result label: {visual}")

    build = write_build_failure_report(
        str(new_workdir("pipeline_fallback_status_build")),
        mode="user",
        stderr="synthetic build failure",
        stdout="",
    )
    assert_true(build["status"] == "failed", f"build failure report should expose failed status: {build}")
    assert_true(build["result_label"] == "未通过", f"build failure report should expose result label: {build}")

    extraction = write_extraction_failure_report(
        str(new_workdir("pipeline_fallback_status_extract")),
        mode="developer",
        label="Content",
        error="synthetic verification mismatch",
        target="content_parser.py",
    )
    assert_true(extraction["status"] == "failed", f"extraction failure report should expose failed status: {extraction}")
    assert_true(extraction["result_label"] == "未通过", f"extraction failure report should expose result label: {extraction}")


@case
def pipeline_auto_repair_patches_reference_east_asia_mismatch() -> None:
    work = new_workdir("pipeline_auto_repair_reference_font")
    build_path = _write_repair_loop_fixture(work)
    build_path.write_text(
        build_path.read_text(encoding="utf-8")
        + """
def add_reference_mixed_runs(p, text, prof):
    # Chinese parts use the role's CJK font; Latin/numeric punctuation uses Times New Roman.
    for seg in re.findall(r'[\\u4e00-\\u9fff]+|[^\\u4e00-\\u9fff]+', text):
        r = p.add_run(seg)
        if has_cjk(seg):
            apply_run_profile(r, prof, seg, force_latin='Times New Roman')
        else:
            p_latin = dict(prof); p_latin['font'] = 'Times New Roman'
            apply_run_profile(r, p_latin, seg, force_latin='Times New Roman')
""",
        encoding="utf-8",
    )
    first_build = run_generated_script(str(build_path), str(work), python_executable=sys.executable)
    assert_true(first_build.returncode == 0, f"synthetic build failed: {first_build.stderr}")

    def fake_qa(out_dir, mode="user", output_docx_name="final.docx"):
        return _fake_repair_report(out_dir, code="NO_ISSUE", severity="info")

    def fake_conformance(out_dir, mode="user", output_docx_name="final.docx", project_root=""):
        patched = "AUTO_REPAIR_REFERENCE_EAST_ASIA_FONT_V1" in Path(out_dir, "build_generated.py").read_text(encoding="utf-8")
        return _write_conformance_report(out_dir, passed=patched)

    deps = QADependencies(
        qa_check_and_write=fake_qa,
        conformance_check_and_write=fake_conformance,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="strict",
        project_root=str(REPO_ROOT),
        max_rounds=3,
        stop_no_improve=2,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )
    assert_true(result.ok, f"reference font auto repair did not converge: {result.status}")
    assert_true("AUTO_REPAIR_REFERENCE_EAST_ASIA_FONT_V1" in build_path.read_text(encoding="utf-8"), "reference font patch was not applied")


@case
def pipeline_auto_repair_stops_after_no_improvement() -> None:
    work = new_workdir("pipeline_auto_repair_no_improve")
    build_path = _write_repair_loop_fixture(work)
    calls = {"qa": 0}

    def fake_qa(out_dir, mode="user", output_docx_name="final.docx"):
        calls["qa"] += 1
        return _fake_repair_report(out_dir, code="MISSING_DOCX", message="synthetic persistent missing docx")

    deps = QADependencies(
        qa_check_and_write=fake_qa,
        conformance_check_and_write=None,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="basic",
        project_root=str(REPO_ROOT),
        max_rounds=5,
        stop_no_improve=2,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )
    report = json.loads((work / "repair_loop_report.json").read_text(encoding="utf-8"))
    assert_true(not result.ok and result.status == "stopped_no_improvement", f"unexpected stop status: {result.status}")
    assert_true(calls["qa"] >= 3, "QA was not rerun during no-improvement loop")
    assert_true((work / "final.docx").exists(), "rebuild action did not run build_generated.py")
    assert_true(report["stop_detail"], "no-improvement stop detail was not recorded")
    assert_true(report.get("next_action"), f"repair loop report should expose a beginner-facing next action: {report}")
    assert_true(report.get("resume_scope") == "current_docx", f"no-improvement repair loop should resume from the current DOCX script: {report}")
    assert_true("build_generated.py" in report.get("resume_command", ""), f"repair loop report should provide a copyable resume command: {report}")
    loop_text = (work / "repair_loop_report.md").read_text(encoding="utf-8")
    assert_true("下一步" in loop_text and "修复范围" in loop_text, "repair loop markdown should show the next action and scope")
    assert_true(build_path.read_text(encoding="utf-8").count("AUTO_REPAIR") == 0, "no-improvement rebuild should not patch script")


@case
def pipeline_auto_repair_stops_for_needs_user_file() -> None:
    work = new_workdir("pipeline_auto_repair_needs_file")
    _write_repair_loop_fixture(work)
    calls = {"qa": 0}

    def fake_qa(out_dir, mode="user", output_docx_name="final.docx"):
        calls["qa"] += 1
        return _fake_repair_report(out_dir, code="CONTENT_IMAGE_MISSING", message="synthetic missing image")

    deps = QADependencies(
        qa_check_and_write=fake_qa,
        conformance_check_and_write=None,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="basic",
        project_root=str(REPO_ROOT),
        max_rounds=5,
        stop_no_improve=2,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )
    report = json.loads((work / "repair_loop_report.json").read_text(encoding="utf-8"))
    assert_true(not result.ok and result.status == "stopped_needs_user_file", f"unexpected stop status: {result.status}")
    assert_true(calls["qa"] == 1, "needs-user-file blocker should stop before rebuild loops")
    assert_true(report["blockers"] and report["blockers"][0]["code"] == "CONTENT_IMAGE_MISSING", "blocker was not recorded")
    assert_true(report.get("resume_scope") == "input_files", f"needs-user-file stop should route to input files: {report}")
    assert_true("CONTENT_IMAGE_MISSING" in report.get("next_action", ""), f"needs-user-file stop should name the blocking code: {report}")
    assert_true("完整流水线" in report.get("next_action", ""), f"needs-user-file stop should tell users to rerun the full pipeline after fixing input: {report}")


@case
def pipeline_auto_repair_build_failure_has_resume_handoff() -> None:
    work = new_workdir("pipeline_auto_repair_build_failure")
    _write_repair_loop_fixture(work)
    calls = {"build": 0}

    def fake_qa(out_dir, mode="user", output_docx_name="final.docx"):
        return _fake_repair_report(out_dir, code="PLACEHOLDER_TEXT_LEFT", message="synthetic placeholder")

    def failing_build(_gen_py_path, _out_dir, python_executable):
        calls["build"] += 1
        return ScriptExecutionResult(returncode=1, stdout="", stderr="synthetic build failure")

    deps = QADependencies(
        qa_check_and_write=fake_qa,
        conformance_check_and_write=None,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="basic",
        project_root=str(REPO_ROOT),
        max_rounds=2,
        stop_no_improve=1,
        deps=deps,
        run_generated_script=failing_build,
        python_executable=sys.executable,
    )
    report = json.loads((work / "repair_loop_report.json").read_text(encoding="utf-8"))
    assert_true(not result.ok and result.status == "stopped_build_failed", f"unexpected build-failure stop status: {result.status}")
    assert_true(calls["build"] == 1, "auto repair should attempt one rebuild before stopping")
    assert_true(report.get("resume_scope") == "current_docx", f"build failure should route to current docx repair: {report}")
    assert_true("build_generated.py" in report.get("resume_command", ""), f"build failure should provide rebuild command: {report}")
    assert_true("构建错误" in report.get("next_action", "") and "build_generated.py" in report.get("next_action", ""), f"build failure next action should be concrete: {report}")


@case
def pipeline_auto_repair_qa_crash_has_resume_handoff() -> None:
    work = new_workdir("pipeline_auto_repair_qa_crash")
    _write_repair_loop_fixture(work)
    private_path = work / "private" / "qa_checker.py"

    def crashing_qa(out_dir, mode="user", output_docx_name="final.docx"):
        raise RuntimeError(f"synthetic QA crash at {private_path}")

    deps = QADependencies(
        qa_check_and_write=crashing_qa,
        conformance_check_and_write=None,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="basic",
        project_root=str(work),
        max_rounds=2,
        stop_no_improve=1,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )
    report_text = (work / "repair_loop_report.json").read_text(encoding="utf-8")
    report = json.loads(report_text)
    qa_report = json.loads((work / "qa_report.json").read_text(encoding="utf-8"))
    plan = json.loads((work / "qa_repair_plan.json").read_text(encoding="utf-8"))
    assert_true(not result.ok, "auto repair should stop when structural QA crashes")
    assert_true(report["status"] == "stopped_needs_user_input", f"unexpected structural crash stop status: {report}")
    assert_true("STRUCTURAL_QA_FAILED" in report["final_error_codes"], f"structural crash code missing from repair report: {report}")
    assert_true(report.get("resume_scope") == "manual_or_dependency", f"structural crash should route to dependency/manual repair: {report}")
    assert_true("run_pipeline.py" in report.get("resume_command", ""), f"structural crash should preserve a rerun command: {report}")
    assert_true("--auto-repair" in report.get("resume_command", ""), f"structural crash rerun command should preserve auto repair: {report}")
    assert_true("qa_checker.py" in report.get("next_action", ""), f"structural crash next action should name qa_checker.py: {report}")
    assert_true(qa_report["issues"][0]["code"] == "STRUCTURAL_QA_FAILED", f"structural crash QA report used wrong code: {qa_report}")
    assert_true(plan.get("resume_scope") == "full_pipeline", f"structural crash repair plan should route to full pipeline: {plan}")
    assert_true(str(work) not in report_text, "repair-loop structural crash report leaked an absolute path")
    assert_true("<PROJECT>" in report_text, f"structural crash detail was not sanitized: {report}")
    assert_true((work / "repair_loop_report.md").exists(), "structural crash should write repair_loop_report.md")


@case
def pipeline_auto_repair_conformance_crash_has_resume_handoff() -> None:
    work = new_workdir("pipeline_auto_repair_conformance_crash")
    _write_repair_loop_fixture(work)
    private_path = work / "private" / "qa_conformance.py"

    def passing_qa(out_dir, mode="user", output_docx_name="final.docx"):
        return _fake_repair_report(out_dir, code="NO_ISSUE", severity="info")

    def crashing_conformance(out_dir, mode="user", output_docx_name="final.docx", project_root=""):
        raise RuntimeError(f"synthetic conformance crash at {private_path}")

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=crashing_conformance,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="strict",
        project_root=str(work),
        max_rounds=2,
        stop_no_improve=1,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )
    report_text = (work / "repair_loop_report.json").read_text(encoding="utf-8")
    report = json.loads(report_text)
    conformance_report = json.loads((work / "conformance_report.json").read_text(encoding="utf-8"))
    assert_true(not result.ok, "auto repair should stop when conformance QA crashes")
    assert_true("CONFORMANCE_QA_FAILED" in report["final_error_codes"], f"conformance crash code missing from repair report: {report}")
    assert_true(report["blockers"] and report["blockers"][0]["code"] == "CONFORMANCE_QA_FAILED", f"conformance crash should be a blocker: {report}")
    assert_true("qa_conformance.py" in report.get("next_action", ""), f"conformance crash next action should name qa_conformance.py: {report}")
    assert_true("run_pipeline.py" in report.get("resume_command", ""), f"conformance crash should preserve a rerun command: {report}")
    assert_true("--auto-repair" in report.get("resume_command", ""), f"conformance crash rerun command should preserve auto repair: {report}")
    assert_true(conformance_report["issues"][0]["code"] == "CONFORMANCE_QA_FAILED", f"wrong conformance crash code: {conformance_report}")
    assert_true(str(work) not in report_text, "repair-loop conformance crash report leaked an absolute path")
    assert_true("<PROJECT>" in report_text, f"conformance crash detail was not sanitized: {report}")


@case
def pipeline_auto_repair_visual_crash_has_resume_handoff() -> None:
    work = new_workdir("pipeline_auto_repair_visual_crash")
    _write_repair_loop_fixture(work)
    private_path = work / "private" / "qa_visual.py"

    def passing_qa(out_dir, mode="user", output_docx_name="final.docx"):
        return _fake_repair_report(out_dir, code="NO_ISSUE", severity="info")

    def passing_conformance(out_dir, mode="user", output_docx_name="final.docx", project_root=""):
        return _write_conformance_report(out_dir, passed=True)

    def crashing_visual(out_dir, output_docx_name="final.docx", project_root="", render_all_pages=True, require_wps=False, golden_dir=None, update_golden=False):
        raise RuntimeError(f"synthetic visual crash at {private_path}")

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=passing_conformance,
        visual_check_and_write=crashing_visual,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="visual",
        project_root=str(work),
        max_rounds=2,
        stop_no_improve=1,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )
    report_text = (work / "repair_loop_report.json").read_text(encoding="utf-8")
    report = json.loads(report_text)
    visual_report = json.loads((work / "visual_report.json").read_text(encoding="utf-8"))
    assert_true(not result.ok, "auto repair should stop when visual QA crashes")
    assert_true("VISUAL_QA_FAILED" in report["final_error_codes"], f"visual crash code missing from repair report: {report}")
    assert_true(report["blockers"] and report["blockers"][0]["code"] == "VISUAL_QA_FAILED", f"visual crash should be a blocker: {report}")
    assert_true("qa_visual.py" in report.get("next_action", ""), f"visual crash next action should name qa_visual.py: {report}")
    assert_true("run_pipeline.py" in report.get("resume_command", ""), f"visual crash should preserve a rerun command: {report}")
    assert_true("--auto-repair" in report.get("resume_command", ""), f"visual crash rerun command should preserve auto repair: {report}")
    assert_true(visual_report["issues"][0]["code"] == "VISUAL_QA_FAILED", f"wrong visual crash code: {visual_report}")
    assert_true(str(work) not in report_text, "repair-loop visual crash report leaked an absolute path")
    assert_true("<PROJECT>" in report_text, f"visual crash detail was not sanitized: {report}")


@case
def pipeline_auto_repair_strict_requires_conformance_dependency() -> None:
    work = new_workdir("pipeline_auto_repair_missing_conformance")
    _write_repair_loop_fixture(work)
    _fake_repair_report(str(work), code="NO_ISSUE", severity="info")
    raw_detail = f"missing dependency detail at {work / 'private' / 'qa_conformance.py'}"

    deps = QADependencies(
        qa_check_and_write=lambda out_dir, mode="user", output_docx_name="final.docx": _fake_repair_report(out_dir, code="NO_ISSUE", severity="info"),
        conformance_check_and_write=None,
        visual_check_and_write=None,
        optional_import_detail=lambda name: raw_detail,
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="strict",
        project_root=str(work),
        max_rounds=1,
        stop_no_improve=1,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )
    report = json.loads((work / "repair_loop_report.json").read_text(encoding="utf-8"))
    conformance_report = json.loads((work / "conformance_report.json").read_text(encoding="utf-8"))
    assert_true(not result.ok, "strict auto repair should not converge without conformance QA")
    assert_true("CONFORMANCE_QA_UNAVAILABLE" in report["final_error_codes"], f"missing conformance error was not recorded: {report}")
    assert_true(report["blockers"] and report["blockers"][0]["code"] == "CONFORMANCE_QA_UNAVAILABLE", "missing conformance dependency should be a blocker")
    assert_true("conformance_report.md" in conformance_report.get("next_action", ""), f"dependency report did not explain next action: {conformance_report}")
    conformance_text = json.dumps(conformance_report, ensure_ascii=False)
    assert_true(str(work) not in conformance_text, "repair dependency report leaked an absolute path")
    assert_true("<PROJECT>" in conformance_text, f"repair dependency report did not sanitize detail: {conformance_report}")


@case
def pipeline_auto_repair_visual_preserves_visual_options() -> None:
    work = new_workdir("pipeline_auto_repair_visual_options")
    build_path = _write_repair_loop_fixture(work)
    visual_calls: List[Dict[str, Any]] = []
    qa_calls = {"count": 0}

    def fake_qa(out_dir, mode="user", output_docx_name="final.docx"):
        qa_calls["count"] += 1
        if qa_calls["count"] == 1:
            return _fake_repair_report(out_dir, code="PLACEHOLDER_TEXT_LEFT", message="synthetic placeholder")
        return _fake_repair_report(out_dir, code="NO_ISSUE", severity="info")

    def fake_conformance(out_dir, mode="user", output_docx_name="final.docx", project_root=""):
        return _write_conformance_report(out_dir, passed=True)

    def fake_visual(out_dir, output_docx_name="final.docx", project_root="", render_all_pages=True, require_wps=False, golden_dir=None, update_golden=False):
        visual_calls.append(
            {
                "require_wps": require_wps,
                "golden_dir": golden_dir,
                "update_golden": update_golden,
                "render_all_pages": render_all_pages,
            }
        )
        report = {
            "schema_version": 1,
            "mode": "user",
            "passed": True,
            "counts": {},
            "issues": [],
            "next_action": "ok",
        }
        Path(out_dir, "visual_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        return report

    deps = QADependencies(
        qa_check_and_write=fake_qa,
        conformance_check_and_write=fake_conformance,
        visual_check_and_write=fake_visual,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="visual",
        project_root=str(REPO_ROOT),
        max_rounds=3,
        stop_no_improve=2,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
        golden_dir=str(work / "Golden"),
        update_golden=True,
        require_wps=True,
    )
    assert_true(result.ok, f"visual auto repair did not converge: {result.status}")
    assert_true("AUTO_REPAIR_PLACEHOLDER_CLEANUP_V1" in build_path.read_text(encoding="utf-8"), "placeholder repair was not applied")
    assert_true(len(visual_calls) >= 2, "visual QA was not rerun after repair")
    assert_true(all(call["require_wps"] is True for call in visual_calls), f"require_wps was not preserved: {visual_calls}")
    assert_true(all(call["update_golden"] is True for call in visual_calls), f"update_golden was not preserved: {visual_calls}")
    assert_true(all(str(call["golden_dir"]).endswith("Golden") for call in visual_calls), f"golden_dir was not preserved: {visual_calls}")


@case
def pipeline_auto_repair_blocks_unrepairable_visual_errors() -> None:
    work = new_workdir("pipeline_auto_repair_visual_blocker")
    _write_repair_loop_fixture(work)

    def passing_qa(out_dir, mode="user", output_docx_name="final.docx"):
        return _fake_repair_report(out_dir, code="NO_ISSUE", severity="info")

    def passing_conformance(out_dir, mode="user", output_docx_name="final.docx", project_root=""):
        return _write_conformance_report(out_dir, passed=True)

    def failing_visual(out_dir, output_docx_name="final.docx", project_root="", render_all_pages=True, require_wps=False, golden_dir=None, update_golden=False):
        report = {
            "schema_version": 1,
            "mode": "user",
            "passed": False,
            "counts": {},
            "issues": [{"severity": "error", "code": "PAGE_IMAGE_UNREADABLE", "message": "bad png", "detail": "page=2"}],
            "next_action": "Fix rendering tools.",
        }
        Path(out_dir, "visual_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        return report

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=passing_conformance,
        visual_check_and_write=failing_visual,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="visual",
        project_root=str(REPO_ROOT),
        max_rounds=2,
        stop_no_improve=1,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )
    report = json.loads((work / "repair_loop_report.json").read_text(encoding="utf-8"))
    assert_true(not result.ok, "unrepairable visual QA error should stop auto repair")
    assert_true(result.status == "stopped_needs_user_input", f"visual blocker should not be reported as missing user file: {result.status}")
    assert_true(report["blockers"] and report["blockers"][0]["code"] == "PAGE_IMAGE_UNREADABLE", f"visual error was not surfaced as blocker: {report}")


@case
def pipeline_auto_repair_wps_page_mismatch_has_visual_rerun_handoff() -> None:
    work = new_workdir("pipeline_auto_repair_wps_mismatch")
    _write_repair_loop_fixture(work)

    def passing_qa(out_dir, mode="user", output_docx_name="final.docx"):
        return _fake_repair_report(out_dir, code="NO_ISSUE", severity="info")

    def passing_conformance(out_dir, mode="user", output_docx_name="final.docx", project_root=""):
        return _write_conformance_report(out_dir, passed=True)

    def failing_visual(out_dir, output_docx_name="final.docx", project_root="", render_all_pages=True, require_wps=False, golden_dir=None, update_golden=False):
        report = {
            "schema_version": 1,
            "mode": "user",
            "passed": False,
            "counts": {},
            "issues": [{"severity": "error", "code": "WPS_PAGE_COUNT_MISMATCH", "message": "page mismatch", "detail": "word=10 wps=11"}],
            "next_action": "Compare Word and WPS PDFs.",
        }
        Path(out_dir, "visual_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        return report

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=passing_conformance,
        visual_check_and_write=failing_visual,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="visual",
        project_root=str(REPO_ROOT),
        max_rounds=2,
        stop_no_improve=1,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )
    report = json.loads((work / "repair_loop_report.json").read_text(encoding="utf-8"))
    assert_true(not result.ok, "WPS page mismatch should stop auto repair")
    assert_true(report["blockers"] and report["blockers"][0]["code"] == "WPS_PAGE_COUNT_MISMATCH", f"WPS mismatch should be a blocker: {report}")
    assert_true("重跑 visual QA" in report.get("next_action", ""), f"WPS mismatch handoff should tell users to rerun visual QA: {report}")


@case
def pipeline_auto_repair_wps_page_size_mismatch_has_visual_rerun_handoff() -> None:
    work = new_workdir("pipeline_auto_repair_wps_page_size_mismatch")
    _write_repair_loop_fixture(work)

    def passing_qa(out_dir, mode="user", output_docx_name="final.docx"):
        return _fake_repair_report(out_dir, code="NO_ISSUE", severity="info")

    def passing_conformance(out_dir, mode="user", output_docx_name="final.docx", project_root=""):
        return _write_conformance_report(out_dir, passed=True)

    def failing_visual(out_dir, output_docx_name="final.docx", project_root="", render_all_pages=True, require_wps=False, golden_dir=None, update_golden=False):
        report = {
            "schema_version": 1,
            "mode": "user",
            "passed": False,
            "counts": {},
            "issues": [{"severity": "error", "code": "WPS_PAGE_SIZE_MISMATCH", "message": "page size mismatch", "detail": "word=595.3x841.9 wps=841.9x595.3"}],
            "next_action": "Compare Word and WPS paper sizes.",
        }
        Path(out_dir, "visual_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        return report

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=passing_conformance,
        visual_check_and_write=failing_visual,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="visual",
        project_root=str(REPO_ROOT),
        max_rounds=2,
        stop_no_improve=1,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )
    report = json.loads((work / "repair_loop_report.json").read_text(encoding="utf-8"))
    assert_true(not result.ok, "WPS page-size mismatch should stop auto repair")
    assert_true(result.status == "stopped_needs_user_input", f"WPS page-size mismatch should require manual confirmation: {result.status}")
    assert_true(report["blockers"] and report["blockers"][0]["code"] == "WPS_PAGE_SIZE_MISMATCH", f"WPS page-size mismatch should be a blocker: {report}")
    assert_true("WPS" in report.get("next_action", "") and "重跑 visual QA" in report.get("next_action", ""), f"WPS page-size mismatch handoff should tell users to rerun visual QA: {report}")


@case
def pipeline_auto_repair_wps_text_mismatch_has_visual_rerun_handoff() -> None:
    work = new_workdir("pipeline_auto_repair_wps_text_mismatch")
    _write_repair_loop_fixture(work)

    def passing_qa(out_dir, mode="user", output_docx_name="final.docx"):
        return _fake_repair_report(out_dir, code="NO_ISSUE", severity="info")

    def passing_conformance(out_dir, mode="user", output_docx_name="final.docx", project_root=""):
        return _write_conformance_report(out_dir, passed=True)

    def failing_visual(out_dir, output_docx_name="final.docx", project_root="", render_all_pages=True, require_wps=False, golden_dir=None, update_golden=False):
        report = {
            "schema_version": 1,
            "mode": "user",
            "passed": False,
            "counts": {"text_pages": 12, "wps_text_pages": 0},
            "issues": [{"severity": "error", "code": "WPS_TEXT_PAGE_MISMATCH", "message": "text page mismatch", "detail": "word_text_pages=12 wps_text_pages=0"}],
            "next_action": "Compare Word and WPS text pages.",
        }
        Path(out_dir, "visual_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        return report

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=passing_conformance,
        visual_check_and_write=failing_visual,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="visual",
        project_root=str(REPO_ROOT),
        max_rounds=2,
        stop_no_improve=1,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )
    report = json.loads((work / "repair_loop_report.json").read_text(encoding="utf-8"))
    assert_true(not result.ok, "WPS text-page mismatch should stop auto repair")
    assert_true(result.status == "stopped_needs_user_input", f"WPS text-page mismatch should require manual confirmation: {result.status}")
    assert_true(report["blockers"] and report["blockers"][0]["code"] == "WPS_TEXT_PAGE_MISMATCH", f"WPS text-page mismatch should be a blocker: {report}")
    assert_true("WPS" in report.get("next_action", "") and "重跑 visual QA" in report.get("next_action", ""), f"WPS text-page mismatch handoff should tell users to rerun visual QA: {report}")


@case
def pipeline_auto_repair_wps_sample_mismatch_has_visual_rerun_handoff() -> None:
    work = new_workdir("pipeline_auto_repair_wps_sample_mismatch")
    _write_repair_loop_fixture(work)

    def passing_qa(out_dir, mode="user", output_docx_name="final.docx"):
        return _fake_repair_report(out_dir, code="NO_ISSUE", severity="info")

    def passing_conformance(out_dir, mode="user", output_docx_name="final.docx", project_root=""):
        return _write_conformance_report(out_dir, passed=True)

    def failing_visual(out_dir, output_docx_name="final.docx", project_root="", render_all_pages=True, require_wps=False, golden_dir=None, update_golden=False):
        report = {
            "schema_version": 1,
            "mode": "user",
            "passed": False,
            "counts": {"sample_images": 2, "wps_sample_images": 2, "wps_sample_mismatches": [1]},
            "issues": [{"severity": "error", "code": "WPS_SAMPLE_IMAGE_MISMATCH", "message": "sample image mismatch", "detail": "pages=1"}],
            "next_action": "Compare Word and WPS sample PNGs.",
        }
        Path(out_dir, "visual_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        return report

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=passing_conformance,
        visual_check_and_write=failing_visual,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="visual",
        project_root=str(REPO_ROOT),
        max_rounds=2,
        stop_no_improve=1,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )
    report = json.loads((work / "repair_loop_report.json").read_text(encoding="utf-8"))
    assert_true(not result.ok, "WPS sample-image mismatch should stop auto repair")
    assert_true(result.status == "stopped_needs_user_input", f"WPS sample-image mismatch should require manual confirmation: {result.status}")
    assert_true(report["blockers"] and report["blockers"][0]["code"] == "WPS_SAMPLE_IMAGE_MISMATCH", f"WPS sample-image mismatch should be a blocker: {report}")
    assert_true("WPS" in report.get("next_action", "") and ("样张" in report.get("next_action", "") or "PNG" in report.get("next_action", "")) and "重跑 visual QA" in report.get("next_action", ""), f"WPS sample-image mismatch handoff should tell users to rerun visual QA: {report}")


@case
def pipeline_auto_repair_report_paths_are_sanitized() -> None:
    work = new_workdir("pipeline_auto_repair_sanitized_paths")
    build_path = _write_repair_loop_fixture(work)
    first_build = run_generated_script(str(build_path), str(work), python_executable=sys.executable)
    assert_true(first_build.returncode == 0, f"synthetic build failed: {first_build.stderr}")

    from qa_checker import check_and_write as qa_check_and_write

    deps = QADependencies(
        qa_check_and_write=qa_check_and_write,
        conformance_check_and_write=None,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    result = run_repair_loop(
        str(work),
        mode="user",
        output_docx_name="final.docx",
        qa_level="basic",
        project_root=str(REPO_ROOT),
        max_rounds=3,
        stop_no_improve=2,
        deps=deps,
        run_generated_script=run_generated_script,
        python_executable=sys.executable,
    )
    report_text = (work / "repair_loop_report.json").read_text(encoding="utf-8")
    report = json.loads(report_text)
    assert_true(result.ok, f"auto repair did not converge: {result.status}")
    assert_true(not Path(report["output_dir"]).is_absolute(), f"output_dir leaked absolute path: {report['output_dir']}")
    assert_true(not Path(report["final_docx"]).is_absolute(), f"final_docx leaked absolute path: {report['final_docx']}")
    assert_true(str(work) not in report_text, "repair loop report leaked the absolute output path")


@case
def pipeline_runs_pdf_template_end_to_end_strict() -> None:
    if not poppler_available():
        return
    if str(REPO_ROOT) not in sys.path:
        sys.path.insert(0, str(REPO_ROOT))
    import run_pipeline as root_runner

    work = new_workdir("pipeline_pdf_e2e")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    outputs_dir = work / "Outputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    outputs_dir.mkdir()

    write_text_pdf(
        template_dir / "pdf_template.pdf",
        [
            ("Format requirements", 16, 72, 780),
            ("Page A4 margins top 2.5 cm bottom 2.5 cm left 3.0 cm right 2.5 cm", 11, 72, 742),
            ("Body font Times New Roman 12 pt justified line spacing 1.5", 11, 72, 718),
            ("Heading 1 font SimHei 16 pt bold centered", 11, 72, 694),
            ("Heading 2 font SimHei 14 pt bold left", 11, 72, 670),
            ("References font Times New Roman 12 pt hanging indent", 11, 72, 646),
        ],
    )

    content_doc = Document()
    content_doc.add_heading("PDF Template End-to-End Demo", level=0)
    content_doc.add_paragraph("Abstract")
    content_doc.add_paragraph("This synthetic document verifies that a PDF template can drive the reusable Word pipeline.")
    content_doc.add_heading("1 Introduction", level=1)
    content_doc.add_paragraph("The body includes enough text for extraction, generation, and strict QA to exercise the normal handoff.")
    content_doc.add_heading("1.1 Method", level=2)
    content_doc.add_paragraph("The method section checks heading hierarchy and paragraph rendering.")
    content_doc.add_paragraph("References")
    content_doc.add_paragraph("[1] Synthetic Author. PDF template regression case. 2026.")
    content_doc.save(inputs_dir / "paper.docx")

    old_dirs = (root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR)
    try:
        root_runner.TEMPLATE_DIR = str(template_dir)
        root_runner.INPUTS_DIR = str(inputs_dir)
        root_runner.OUTPUTS_DIR = str(outputs_dir)
        result = root_runner.run("pdf_template.pdf", "paper.docx", mode="developer", qa_level="strict")
    finally:
        root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR = old_dirs

    assert_true(result, "PDF template end-to-end pipeline returned no output directory")
    result_dir = Path(result)
    assert_true((result_dir / "最终论文.docx").exists(), "PDF template E2E did not build final DOCX")
    qa = json.loads((result_dir / "qa_report.json").read_text(encoding="utf-8"))
    assert_true(qa.get("passed") is True, f"PDF template E2E strict QA failed: {qa.get('issues')}")


@case
def pipeline_blocks_scanned_pdf_template_before_generation() -> None:
    if not poppler_available():
        return
    if str(REPO_ROOT) not in sys.path:
        sys.path.insert(0, str(REPO_ROOT))
    import run_pipeline as root_runner

    work = new_workdir("pipeline_scanned_pdf_template_blocker")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    outputs_dir = work / "Outputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    outputs_dir.mkdir()

    write_blank_pdf(template_dir / "scanned_template.pdf")

    content_doc = Document()
    content_doc.add_heading("Scanned PDF Template Demo", level=0)
    content_doc.add_paragraph("This body should not be processed because the PDF template is unsupported.")
    content_doc.save(inputs_dir / "paper.docx")

    old_dirs = (root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR)
    try:
        root_runner.TEMPLATE_DIR = str(template_dir)
        root_runner.INPUTS_DIR = str(inputs_dir)
        root_runner.OUTPUTS_DIR = str(outputs_dir)
        result = root_runner.run("scanned_template.pdf", "paper.docx", mode="developer", qa_level="strict")
    finally:
        root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR = old_dirs

    assert_true(result is None, "scanned PDF template should stop the pipeline before generation")
    out_dirs = sorted(outputs_dir.iterdir())
    assert_true(out_dirs, "scanned PDF template should still create a report handoff directory")
    out_dir = out_dirs[-1]
    assert_true((out_dir / "qa_report.json").exists(), "scanned PDF blocker should write qa_report.json")
    assert_true((out_dir / "qa_repair_plan.json").exists(), "scanned PDF blocker should write qa_repair_plan.json")
    assert_true((out_dir / "agent_summary.json").exists(), "scanned PDF blocker should write agent_summary.json")
    assert_true(not (out_dir / "build_generated.py").exists(), "unsupported PDF template should not generate build_generated.py")
    assert_true(not (out_dir / "最终论文.docx").exists(), "unsupported PDF template should not build a misleading final DOCX")

    report = json.loads((out_dir / "qa_report.json").read_text(encoding="utf-8"))
    codes = {item.get("code") for item in report.get("issues") or []}
    assert_true("PDF_TEMPLATE_UNSUPPORTED" in codes, f"scanned PDF blocker should report PDF_TEMPLATE_UNSUPPORTED: {report}")
    assert_true("PDF_TEMPLATE_UNSUPPORTED" in report.get("next_action", ""), f"qa_report should name the PDF issue in next_action: {report}")
    profile = json.loads((out_dir / "template_profile.json").read_text(encoding="utf-8"))
    risks = profile.get("risk_flags") or {}
    assert_true(risks.get("pdf_template_unsupported") is True, f"scanned PDF should keep the unsupported risk flag: {profile}")
    assert_true(not risks.get("pdf_template_read_failed"), f"scanned PDF should not be mislabeled as read failed: {profile}")
    assert_true(not risks.get("pdf_template_dependency_missing"), f"scanned PDF should not be mislabeled as dependency missing: {profile}")
    plan = json.loads((out_dir / "qa_repair_plan.json").read_text(encoding="utf-8"))
    assert_true(plan.get("resume_scope") == "input_files", f"scanned PDF should route users to replace/OCR the template: {plan}")
    assert_true("DOCX" in plan.get("next_action", "") or "OCR" in plan.get("next_action", ""), f"scanned PDF next action should name DOCX/OCR route: {plan}")
    summary = json.loads((out_dir / "agent_summary.json").read_text(encoding="utf-8"))
    action_text = "\n".join(summary.get("next_actions") or summary.get("manual_check_required") or [])
    assert_true("PDF_TEMPLATE_UNSUPPORTED" in action_text and ("OCR" in action_text or "DOCX" in action_text), f"agent summary lost scanned-PDF next step: {summary}")


@case
def pipeline_blocks_pdf_template_when_poppler_missing_with_dependency_guidance() -> None:
    if str(REPO_ROOT) not in sys.path:
        sys.path.insert(0, str(REPO_ROOT))
    import run_pipeline as root_runner
    import format_extractor_modules.pdf_template as pdf_template

    work = new_workdir("pipeline_pdf_template_poppler_missing")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    outputs_dir = work / "Outputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    outputs_dir.mkdir()

    write_text_pdf(
        template_dir / "requirements.pdf",
        [
            ("Format requirements", 16, 72, 780),
            ("Page A4 margins top 2.5 cm bottom 2.5 cm left 3.0 cm right 2.5 cm", 11, 72, 742),
            ("Body font Times New Roman 12 pt justified line spacing 1.5", 11, 72, 718),
            ("Heading 1 font SimHei 16 pt bold centered", 11, 72, 694),
        ],
    )

    content_doc = Document()
    content_doc.add_heading("Missing Poppler PDF Template Demo", level=0)
    content_doc.add_paragraph("This body should not be processed while PDF template dependencies are missing.")
    content_doc.save(inputs_dir / "paper.docx")

    old_dirs = (root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR)
    old_which = pdf_template.shutil.which
    try:
        root_runner.TEMPLATE_DIR = str(template_dir)
        root_runner.INPUTS_DIR = str(inputs_dir)
        root_runner.OUTPUTS_DIR = str(outputs_dir)
        pdf_template.shutil.which = lambda _name: None
        result = root_runner.run("requirements.pdf", "paper.docx", mode="developer", qa_level="strict")
    finally:
        root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR = old_dirs
        pdf_template.shutil.which = old_which

    assert_true(result is None, "missing Poppler should stop the pipeline before generation")
    out_dirs = sorted(outputs_dir.iterdir())
    assert_true(out_dirs, "missing Poppler should still create a report handoff directory")
    out_dir = out_dirs[-1]
    assert_true(not (out_dir / "build_generated.py").exists(), "missing Poppler should not generate build_generated.py")
    assert_true(not (out_dir / "最终论文.docx").exists(), "missing Poppler should not build a misleading final DOCX")

    report = json.loads((out_dir / "qa_report.json").read_text(encoding="utf-8"))
    codes = {item.get("code") for item in report.get("issues") or []}
    assert_true("PDF_TEMPLATE_DEPENDENCY_MISSING" in codes, f"missing Poppler should have a dependency-specific code: {report}")
    assert_true("Poppler" in report.get("next_action", "") and "重跑" in report.get("next_action", ""), f"qa_report should tell users to fix Poppler and rerun: {report}")
    profile = json.loads((out_dir / "template_profile.json").read_text(encoding="utf-8"))
    risks = profile.get("risk_flags") or {}
    assert_true(risks.get("pdf_template_dependency_missing") is True, f"missing Poppler should have a dependency risk flag: {profile}")
    assert_true(not risks.get("pdf_template_read_failed"), f"missing Poppler should not be mislabeled as read failed: {profile}")
    assert_true(not risks.get("pdf_template_unsupported"), f"missing Poppler should not be mislabeled as unsupported: {profile}")
    plan = json.loads((out_dir / "qa_repair_plan.json").read_text(encoding="utf-8"))
    assert_true(plan.get("resume_scope") == "environment", f"missing Poppler should route to environment repair: {plan}")
    assert_true("Poppler" in plan.get("next_action", "") and plan.get("resume_command"), f"repair plan should keep a rerun command after Poppler repair: {plan}")
    summary = json.loads((out_dir / "agent_summary.json").read_text(encoding="utf-8"))
    action_text = "\n".join(summary.get("next_actions") or summary.get("manual_check_required") or [])
    assert_true("PDF_TEMPLATE_DEPENDENCY_MISSING" in action_text and "Poppler" in action_text, f"agent summary lost Poppler repair guidance: {summary}")


@case
def pipeline_blocks_corrupt_pdf_template_with_read_failure_guidance() -> None:
    if str(REPO_ROOT) not in sys.path:
        sys.path.insert(0, str(REPO_ROOT))
    import run_pipeline as root_runner

    work = new_workdir("pipeline_corrupt_pdf_template_blocker")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    outputs_dir = work / "Outputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    outputs_dir.mkdir()

    (template_dir / "damaged_template.pdf").write_bytes(b"%PDF-1.4\nnot a complete pdf\n")

    content_doc = Document()
    content_doc.add_heading("Corrupt PDF Template Demo", level=0)
    content_doc.add_paragraph("This body should not be processed while the PDF template is unreadable.")
    content_doc.save(inputs_dir / "paper.docx")

    old_dirs = (root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR)
    try:
        root_runner.TEMPLATE_DIR = str(template_dir)
        root_runner.INPUTS_DIR = str(inputs_dir)
        root_runner.OUTPUTS_DIR = str(outputs_dir)
        result = root_runner.run("damaged_template.pdf", "paper.docx", mode="developer", qa_level="strict")
    finally:
        root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR = old_dirs

    assert_true(result is None, "corrupt PDF template should stop the pipeline before generation")
    out_dirs = sorted(outputs_dir.iterdir())
    assert_true(out_dirs, "corrupt PDF template should still create a report handoff directory")
    out_dir = out_dirs[-1]
    assert_true(not (out_dir / "build_generated.py").exists(), "corrupt PDF template should not generate build_generated.py")
    assert_true(not (out_dir / "最终论文.docx").exists(), "corrupt PDF template should not build a misleading final DOCX")

    report = json.loads((out_dir / "qa_report.json").read_text(encoding="utf-8"))
    codes = {item.get("code") for item in report.get("issues") or []}
    assert_true("PDF_TEMPLATE_READ_FAILED" in codes, f"corrupt PDF should have a read-failure code: {report}")
    assert_true("PDF_TEMPLATE_UNSUPPORTED" not in codes, f"corrupt PDF should not be mislabeled as scanned/textless: {report}")
    assert_true("PDF_TEMPLATE_READ_FAILED" in report.get("next_action", ""), f"qa_report should name the PDF read issue: {report}")
    profile = json.loads((out_dir / "template_profile.json").read_text(encoding="utf-8"))
    risks = profile.get("risk_flags") or {}
    assert_true(risks.get("pdf_template_read_failed") is True, f"corrupt PDF should have a read-failure risk flag: {profile}")
    assert_true(not risks.get("pdf_template_dependency_missing"), f"corrupt PDF should not be mislabeled as dependency missing: {profile}")
    assert_true(not risks.get("pdf_template_unsupported"), f"corrupt PDF should not be mislabeled as scanned/textless: {profile}")
    plan = json.loads((out_dir / "qa_repair_plan.json").read_text(encoding="utf-8"))
    assert_true(plan.get("resume_scope") == "input_files", f"corrupt PDF should route users to replace/re-export the template: {plan}")
    assert_true("重新导出" in plan.get("next_action", "") or "可正常打开" in plan.get("next_action", ""), f"corrupt PDF next action should name re-export/openable-PDF route: {plan}")
    assert_true("OCR" not in plan.get("next_action", ""), f"corrupt PDF should not tell users to OCR as the primary route: {plan}")
    summary = json.loads((out_dir / "agent_summary.json").read_text(encoding="utf-8"))
    action_text = "\n".join(summary.get("next_actions") or summary.get("manual_check_required") or [])
    assert_true("PDF_TEMPLATE_READ_FAILED" in action_text and ("重新导出" in action_text or "可正常打开" in action_text), f"agent summary lost corrupt-PDF next step: {summary}")


@case
def pipeline_blocks_protected_pdf_template_with_unlock_guidance() -> None:
    if str(REPO_ROOT) not in sys.path:
        sys.path.insert(0, str(REPO_ROOT))
    import run_pipeline as root_runner
    import format_extractor_modules.pdf_template as pdf_template

    work = new_workdir("pipeline_protected_pdf_template_blocker")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    outputs_dir = work / "Outputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    outputs_dir.mkdir()

    (template_dir / "protected_template.pdf").write_bytes(b"%PDF-1.4\n% synthetic protected pdf placeholder\n")

    content_doc = Document()
    content_doc.add_heading("Protected PDF Template Demo", level=0)
    content_doc.add_paragraph("This body should not be processed while the PDF template is password protected.")
    content_doc.save(inputs_dir / "paper.docx")

    old_dirs = (root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR)
    old_which = pdf_template.shutil.which
    old_run = pdf_template.subprocess.run

    def fake_which(name: str) -> str:
        if name in {"pdfinfo", "pdftotext"}:
            return f"C:/fake-poppler/{name}.exe"
        return old_which(name)

    def fake_run(*_args, **_kwargs):
        return SimpleNamespace(
            returncode=1,
            stdout=b"",
            stderr=b"Command Line Error: Incorrect password\nThis file is encrypted and requires a password.\n",
        )

    try:
        root_runner.TEMPLATE_DIR = str(template_dir)
        root_runner.INPUTS_DIR = str(inputs_dir)
        root_runner.OUTPUTS_DIR = str(outputs_dir)
        pdf_template.shutil.which = fake_which
        pdf_template.subprocess.run = fake_run
        result = root_runner.run("protected_template.pdf", "paper.docx", mode="developer", qa_level="strict")
    finally:
        root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR = old_dirs
        pdf_template.shutil.which = old_which
        pdf_template.subprocess.run = old_run

    assert_true(result is None, "protected PDF template should stop the pipeline before generation")
    out_dirs = sorted(outputs_dir.iterdir())
    assert_true(out_dirs, "protected PDF template should still create a report handoff directory")
    out_dir = out_dirs[-1]
    assert_true(not (out_dir / "build_generated.py").exists(), "protected PDF template should not generate build_generated.py")
    assert_true(not (out_dir / "最终论文.docx").exists(), "protected PDF template should not build a misleading final DOCX")

    report = json.loads((out_dir / "qa_report.json").read_text(encoding="utf-8"))
    codes = {item.get("code") for item in report.get("issues") or []}
    assert_true("PDF_TEMPLATE_PROTECTED" in codes, f"protected PDF should have a password/permission-specific code: {report}")
    assert_true("PDF_TEMPLATE_READ_FAILED" not in codes, f"protected PDF should not be mislabeled as corrupt/unreadable: {report}")
    assert_true("PDF_TEMPLATE_UNSUPPORTED" not in codes, f"protected PDF should not be mislabeled as scanned/textless: {report}")
    next_action = report.get("next_action", "")
    assert_true(
        "PDF_TEMPLATE_PROTECTED" in next_action and "密码" in next_action and ("权限" in next_action or "无保护" in next_action),
        f"qa_report should tell users to unlock/export an unprotected PDF: {report}",
    )
    profile = json.loads((out_dir / "template_profile.json").read_text(encoding="utf-8"))
    risks = profile.get("risk_flags") or {}
    assert_true(risks.get("pdf_template_protected") is True, f"protected PDF should have a protected risk flag: {profile}")
    assert_true(not risks.get("pdf_template_read_failed"), f"protected PDF should not be mislabeled as read failed: {profile}")
    assert_true(not risks.get("pdf_template_unsupported"), f"protected PDF should not be mislabeled as unsupported: {profile}")
    plan = json.loads((out_dir / "qa_repair_plan.json").read_text(encoding="utf-8"))
    assert_true(plan.get("resume_scope") == "input_files", f"protected PDF should route users to replace/unlock the template: {plan}")
    assert_true("密码" in plan.get("next_action", "") and ("权限" in plan.get("next_action", "") or "无保护" in plan.get("next_action", "")), f"protected PDF next action should name password/permission repair: {plan}")
    summary = json.loads((out_dir / "agent_summary.json").read_text(encoding="utf-8"))
    action_text = "\n".join(summary.get("next_actions") or summary.get("manual_check_required") or [])
    assert_true("PDF_TEMPLATE_PROTECTED" in action_text and "密码" in action_text, f"agent summary lost protected-PDF next step: {summary}")


@case
def pipeline_sparse_pdf_instruction_warning_names_missing_rules() -> None:
    if not poppler_available():
        return
    if str(REPO_ROOT) not in sys.path:
        sys.path.insert(0, str(REPO_ROOT))
    import run_pipeline as root_runner

    work = new_workdir("pipeline_sparse_pdf_instruction_warning")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    outputs_dir = work / "Outputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    outputs_dir.mkdir()

    write_text_pdf(
        template_dir / "sparse_requirements.pdf",
        [
            ("Format requirements", 16, 72, 780),
            ("Page A4 margin left 3.0 cm right 2.5 cm", 11, 72, 742),
            ("Body font Times New Roman 12 pt justified line spacing 1.5", 11, 72, 718),
        ],
    )

    content_doc = Document()
    content_doc.add_heading("Sparse PDF Instruction Demo", level=0)
    content_doc.add_heading("Introduction", level=1)
    content_doc.add_paragraph("This body can be formatted, but the PDF template does not describe headings, captions, or references.")
    content_doc.save(inputs_dir / "paper.docx")

    old_dirs = (root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR)
    try:
        root_runner.TEMPLATE_DIR = str(template_dir)
        root_runner.INPUTS_DIR = str(inputs_dir)
        root_runner.OUTPUTS_DIR = str(outputs_dir)
        result = root_runner.run("sparse_requirements.pdf", "paper.docx", mode="developer", qa_level="basic")
    finally:
        root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR = old_dirs

    assert_true(result is not None, "sparse instruction PDF should continue with warning-only QA")
    out_dirs = sorted(outputs_dir.iterdir())
    assert_true(out_dirs, "sparse instruction PDF should create an output directory")
    out_dir = out_dirs[-1]
    assert_true((out_dir / "最终论文.docx").exists(), "sparse instruction PDF should still build the DOCX")

    report = json.loads((out_dir / "qa_report.json").read_text(encoding="utf-8"))
    codes = [item.get("code") for item in report.get("issues") or []]
    assert_true("PDF_TEMPLATE_INSTRUCTION_INCOMPLETE" in codes, f"sparse PDF should name incomplete instruction rules: {report}")
    next_action = report.get("next_action", "")
    assert_true(
        "PDF_TEMPLATE_INSTRUCTION_INCOMPLETE" in next_action and "标题" in next_action and "题注" in next_action and "参考文献" in next_action,
        f"sparse PDF next_action should name the missing rule families: {report}",
    )

    plan = json.loads((out_dir / "qa_repair_plan.json").read_text(encoding="utf-8"))
    assert_true(plan.get("resume_scope") == "warning_review", f"sparse PDF warning should route to warning review: {plan}")
    assert_true("标题" in plan.get("next_action", "") and "题注" in plan.get("next_action", "") and "参考文献" in plan.get("next_action", ""), f"repair plan should name missing sparse-PDF rules: {plan}")
    profile = json.loads((out_dir / "template_profile.json").read_text(encoding="utf-8"))
    risks = profile.get("risk_flags") or {}
    assert_true(risks.get("pdf_template_instruction_incomplete") is True, f"profile should expose sparse instruction risk: {profile}")
    summary = json.loads((out_dir / "agent_summary.json").read_text(encoding="utf-8"))
    action_text = "\n".join(summary.get("next_actions") or summary.get("manual_check_required") or [])
    assert_true("PDF_TEMPLATE_INSTRUCTION_INCOMPLETE" in action_text and "参考文献" in action_text, f"agent summary lost sparse-PDF next step: {summary}")


@case
def pipeline_landscape_pdf_template_warning_names_orientation_review() -> None:
    if not poppler_available():
        return
    if str(REPO_ROOT) not in sys.path:
        sys.path.insert(0, str(REPO_ROOT))
    import run_pipeline as root_runner

    work = new_workdir("pipeline_landscape_pdf_template_warning")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    outputs_dir = work / "Outputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    outputs_dir.mkdir()

    write_text_pdf(
        template_dir / "landscape_requirements.pdf",
        [
            ("Format requirements", 16, 72, 520),
            ("Page landscape A4 margins top 2.0 cm bottom 2.0 cm left 2.0 cm right 2.0 cm", 11, 72, 492),
            ("Body font Times New Roman 12 pt justified line spacing 1.5", 11, 72, 468),
            ("Heading 1 font SimHei 16 pt bold centered", 11, 72, 444),
            ("Heading 2 font SimHei 14 pt bold left", 11, 72, 420),
            ("References font Times New Roman 12 pt", 11, 72, 396),
        ],
        page_width=842.0,
        page_height=595.0,
    )

    content_doc = Document()
    content_doc.add_heading("Landscape PDF Template Demo", level=0)
    content_doc.add_heading("Introduction", level=1)
    content_doc.add_paragraph("This body should build, but users need an explicit orientation review step.")
    content_doc.add_heading("References", level=1)
    content_doc.add_paragraph("[1] Doe J. Landscape template warning regression.")
    content_doc.save(inputs_dir / "paper.docx")

    old_dirs = (root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR)
    try:
        root_runner.TEMPLATE_DIR = str(template_dir)
        root_runner.INPUTS_DIR = str(inputs_dir)
        root_runner.OUTPUTS_DIR = str(outputs_dir)
        result = root_runner.run("landscape_requirements.pdf", "paper.docx", mode="developer", qa_level="basic")
    finally:
        root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR = old_dirs

    assert_true(result is not None, "landscape PDF template should continue with warning-only QA")
    out_dirs = sorted(outputs_dir.iterdir())
    assert_true(out_dirs, "landscape PDF template should create an output directory")
    out_dir = out_dirs[-1]
    assert_true((out_dir / "最终论文.docx").exists(), "landscape PDF template should still build the DOCX")

    report = json.loads((out_dir / "qa_report.json").read_text(encoding="utf-8"))
    codes = [item.get("code") for item in report.get("issues") or []]
    assert_true("PDF_TEMPLATE_LANDSCAPE_PAGE" in codes, f"landscape PDF should name orientation review: {report}")
    next_action = report.get("next_action", "")
    assert_true(
        "PDF_TEMPLATE_LANDSCAPE_PAGE" in next_action and "横向" in next_action and ("页面方向" in next_action or "纸张方向" in next_action),
        f"landscape PDF next_action should tell users to review orientation: {report}",
    )
    plan = json.loads((out_dir / "qa_repair_plan.json").read_text(encoding="utf-8"))
    assert_true(plan.get("resume_scope") == "warning_review", f"landscape PDF warning should route to warning review: {plan}")
    assert_true("横向" in plan.get("next_action", "") and "Word/WPS" in plan.get("next_action", ""), f"repair plan should name Word/WPS orientation review: {plan}")
    profile = json.loads((out_dir / "template_profile.json").read_text(encoding="utf-8"))
    risks = profile.get("risk_flags") or {}
    assert_true(risks.get("pdf_template_landscape_page") is True, f"profile should expose landscape PDF risk: {profile}")
    summary = json.loads((out_dir / "agent_summary.json").read_text(encoding="utf-8"))
    action_text = "\n".join(summary.get("next_actions") or summary.get("manual_check_required") or [])
    assert_true("PDF_TEMPLATE_LANDSCAPE_PAGE" in action_text and "横向" in action_text, f"agent summary lost landscape-PDF next step: {summary}")


@case
def pipeline_visual_pdf_template_warning_names_visual_review() -> None:
    if not poppler_available():
        return
    if str(REPO_ROOT) not in sys.path:
        sys.path.insert(0, str(REPO_ROOT))
    import run_pipeline as root_runner

    work = new_workdir("pipeline_visual_pdf_template_warning")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    outputs_dir = work / "Outputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    outputs_dir.mkdir()

    write_text_pdf(
        template_dir / "visual_sample.pdf",
        [
            ("Synthetic Thesis Title", 20, 180, 780),
            ("Abstract", 16, 260, 735),
            ("This paper studies robust document generation from template samples.", 11, 72, 700),
            ("1 Introduction", 15, 72, 650),
            ("Template adaptation requires careful format extraction and QA.", 11, 72, 622),
            ("Figure 1 System architecture", 10, 210, 585),
            ("References", 15, 72, 540),
            ("[1] Doe J. Synthetic reference for regression testing.", 11, 72, 512),
        ],
    )

    content_doc = Document()
    content_doc.add_heading("Visual PDF Template Demo", level=0)
    content_doc.add_paragraph("Abstract")
    content_doc.add_paragraph("This body should build, but the template format comes from a visual PDF sample.")
    content_doc.add_heading("1 Introduction", level=1)
    content_doc.add_paragraph("Users need explicit Word/WPS review guidance because PDF samples only approximate Word styles.")
    content_doc.add_heading("References", level=1)
    content_doc.add_paragraph("[1] Doe J. Visual sample template warning regression.")
    content_doc.save(inputs_dir / "paper.docx")

    old_dirs = (root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR)
    try:
        root_runner.TEMPLATE_DIR = str(template_dir)
        root_runner.INPUTS_DIR = str(inputs_dir)
        root_runner.OUTPUTS_DIR = str(outputs_dir)
        result = root_runner.run("visual_sample.pdf", "paper.docx", mode="developer", qa_level="basic")
    finally:
        root_runner.TEMPLATE_DIR, root_runner.INPUTS_DIR, root_runner.OUTPUTS_DIR = old_dirs

    assert_true(result is not None, "visual sample PDF template should continue with warning-only QA")
    out_dirs = sorted(outputs_dir.iterdir())
    assert_true(out_dirs, "visual sample PDF template should create an output directory")
    out_dir = out_dirs[-1]
    assert_true((out_dir / "最终论文.docx").exists(), "visual sample PDF template should still build the DOCX")

    report = json.loads((out_dir / "qa_report.json").read_text(encoding="utf-8"))
    codes = [item.get("code") for item in report.get("issues") or []]
    assert_true("PDF_TEMPLATE_VISUAL_APPROXIMATION" in codes, f"visual sample PDF should name visual-review warning: {report}")
    next_action = report.get("next_action", "")
    assert_true(
        "PDF_TEMPLATE_VISUAL_APPROXIMATION" in next_action and "Word/WPS" in next_action and ("样张" in next_action or "视觉" in next_action),
        f"visual sample PDF next_action should tell users to review approximated layout in Word/WPS: {report}",
    )
    plan = json.loads((out_dir / "qa_repair_plan.json").read_text(encoding="utf-8"))
    assert_true(plan.get("resume_scope") == "warning_review", f"visual sample PDF warning should route to warning review: {plan}")
    assert_true("Word/WPS" in plan.get("next_action", "") and ("样张" in plan.get("next_action", "") or "视觉" in plan.get("next_action", "")), f"repair plan should name Word/WPS visual review: {plan}")
    profile = json.loads((out_dir / "template_profile.json").read_text(encoding="utf-8"))
    risks = profile.get("risk_flags") or {}
    assert_true(risks.get("pdf_template_visual_approximation") is True, f"profile should expose visual sample PDF risk: {profile}")
    summary = json.loads((out_dir / "agent_summary.json").read_text(encoding="utf-8"))
    action_text = "\n".join(summary.get("next_actions") or summary.get("manual_check_required") or [])
    assert_true("PDF_TEMPLATE_VISUAL_APPROXIMATION" in action_text and "Word/WPS" in action_text, f"agent summary lost visual-sample-PDF next step: {summary}")


@case
def pipeline_dependencies_loads_optional_modules_and_reports_missing() -> None:
    def marker(name):
        return lambda *args, **kwargs: name

    modules = {
        "qa_checker": SimpleNamespace(check_and_write=marker("qa")),
        "qa_conformance": SimpleNamespace(check_and_write=marker("conformance"), write_requirements=marker("requirements")),
        "template_profiler": SimpleNamespace(write_profile=marker("profile")),
        "md_parser": SimpleNamespace(extract_format=marker("md_format"), extract_content=marker("md_content")),
    }

    def fake_import_module(name):
        if name == "qa_visual":
            raise ImportError("visual unavailable")
        return modules[name]

    deps = load_optional_dependencies(import_module=fake_import_module)
    assert_true(deps.qa_check_and_write() == "qa", "qa checker dependency was not loaded")
    assert_true(deps.conformance_check_and_write() == "conformance", "conformance dependency was not loaded")
    assert_true(deps.write_template_requirements() == "requirements", "template requirements dependency was not loaded")
    assert_true(deps.write_template_profile() == "profile", "template profiler dependency was not loaded")
    assert_true(deps.extract_md_format() == "md_format" and deps.extract_md_content() == "md_content", "md parser dependencies were not loaded")
    assert_true(deps.visual_check_and_write is None, "missing visual dependency should be None")
    assert_true("visual unavailable" in deps.optional_import_detail("qa_visual"), "missing dependency detail was not preserved")
    assert_true(deps.optional_import_detail("qa_checker") == "", "available dependency should not report an error detail")


@case
def pipeline_artifacts_write_format_and_content_handoffs() -> None:
    work = new_workdir("pipeline_artifacts")
    fmt_json_path, fmt_md_path = write_format_artifacts(base_format(), "# Format", str(work))
    assert_true(Path(fmt_json_path).exists(), "format.json was not written")
    assert_true(Path(fmt_md_path).read_text(encoding="utf-8") == "# Format", "format markdown changed")

    content = base_content(
        [
            {"text": "A paragraph with math", "math": [{"text": "E=mc^2"}]},
            {"role": "figure", "image": "fig2.png", "caption": "Figure 2 Demo"},
            {"role": "image", "image": "cell.png", "location": "table_cell"},
            {"role": "table_caption", "text": "表 1 指标对比"},
            {"role": "table", "table_rows": [["Metric", "Value"], ["Accuracy", "98%"]]},
            "Plain paragraph",
        ]
    )
    content["sections"][0]["images"] = ["fig1.png", "fig2.png"]
    cnt_json_path, cnt_md_path = write_content_artifacts(content, str(work), str(work / "paper.docx"))
    summary = Path(cnt_md_path).read_text(encoding="utf-8")
    assert_true(Path(cnt_json_path).exists(), "content.json was not written")
    assert_true("# 内容提取" in summary, "content report title missing")
    assert_true("- [图片] fig1.png" in summary, "content report image line missing")
    assert_true("[图片] fig2.png" in summary, "structured figure was not summarized as an image")
    assert_true("[图片] cell.png" in summary, "structured image item was not summarized as an image")
    assert_true("[结构化内容]" not in summary, "image/table/formula summary leaked an opaque structured-content label")
    assert_true(summary.count("[图片] fig2.png") == 1, "structured figure image was duplicated in content report")
    assert_true("[表格] 2行 x 2列" in summary, "structured table was not summarized as a table")
    assert_true("表 1 指标对比" in summary, "table caption was not preserved in content markdown")
    assert_true(summary.count("[公式]") == 0, "non-formula structured content was mislabeled as formula")
    assert_true("(+1公式)" in summary, "content report math count missing")
    assert_true("## 参考文献" in build_content_markdown(content, str(work / "paper.docx")), "references section missing")


@case
def pipeline_execution_runs_generated_script_with_utf8_output() -> None:
    work = new_workdir("pipeline_execution")
    script = work / "build_generated.py"
    script.write_text("print('生成完成')\n", encoding="utf-8")
    result = run_generated_script(str(script), str(work), python_executable=sys.executable)
    assert_true(result.returncode == 0, f"generated script returned {result.returncode}: {result.stderr}")
    assert_true("生成完成" in result.stdout, "UTF-8 stdout was not decoded")


@case
def pipeline_build_phase_generates_and_blocks_on_failure() -> None:
    work = new_workdir("pipeline_build_phase")
    steps: List[str] = []
    calls: List[str] = []

    def fake_step(label):
        steps.append(label)

    def fake_generate(fmt_json_path, cnt_json_path, out_dir, output_docx_name):
        calls.append(f"generate:{Path(out_dir).name}:{output_docx_name}")
        Path(out_dir, "build_generated.py").write_text("print('ok')\n", encoding="utf-8")
        return 11

    def fake_run(gen_py_path, out_dir, python_executable):
        calls.append(f"run:{Path(gen_py_path).name}:{python_executable}")
        return ScriptExecutionResult(returncode=0, stdout="built\n", stderr="")

    ok = generate_and_build_docx_phase(
        "format.json",
        "content.json",
        str(work),
        "folder",
        output_docx_name="out.docx",
        generate_script=fake_generate,
        run_generated_script=fake_run,
        python_executable="python",
        step=fake_step,
    )
    assert_true(ok is True, "build phase should pass when generated script succeeds")
    assert_true(
        steps[:2]
        == [
            "Phase 4/6: 生成构建脚本",
            "Phase 5/6: 构建最终 docx（生成静态目录；可用 Word COM 时写入页码）",
        ],
        f"unexpected build steps: {steps}",
    )
    assert_true(calls == [f"generate:{work.name}:out.docx", "run:build_generated.py:python"], f"unexpected build calls: {calls}")

    def failing_run(gen_py_path, out_dir, python_executable):
        return ScriptExecutionResult(returncode=1, stdout="", stderr="boom")

    failed = generate_and_build_docx_phase(
        "format.json",
        "content.json",
        str(work),
        "folder",
        output_docx_name="out.docx",
        generate_script=fake_generate,
        run_generated_script=failing_run,
        python_executable="python",
        step=fake_step,
    )
    assert_true(failed is False, "build phase should block when generated script fails")


@case
def pipeline_qa_runs_strict_and_blocks_failures() -> None:
    work = new_workdir("pipeline_qa")
    calls: List[str] = []

    def passing_qa(out_dir, mode, output_docx_name):
        calls.append("qa")
        report = {"passed": True, "issues": [], "counts": {}, "mode": mode}
        report.update(qa_status_fields(report["passed"], report["issues"]))
        return report

    def passing_conformance(out_dir, mode, output_docx_name, project_root):
        calls.append("conformance")
        report = {"passed": True, "issues": [], "counts": {}, "mode": mode}
        report.update(qa_status_fields(report["passed"], report["issues"]))
        return report

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=passing_conformance,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    assert_true(
        run_qa_phases(str(work), mode="developer", output_docx_name="最终论文.docx", qa_level="strict", project_root=str(work), deps=deps),
        "strict QA should pass when structural and conformance QA pass",
    )
    assert_true(calls == ["qa", "conformance"], f"unexpected QA call order: {calls}")

    def failing_qa(out_dir, mode, output_docx_name):
        report = {
            "passed": False,
            "counts": {},
            "issues": [{"severity": "error", "code": "TEST_ERROR", "message": "Synthetic failure", "active_owner": "developer"}],
            "repair_plan": {"steps": []},
        }
        report.update(qa_status_fields(report["passed"], report["issues"]))
        return report

    failing_deps = QADependencies(
        qa_check_and_write=failing_qa,
        conformance_check_and_write=passing_conformance,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    assert_true(
        not run_qa_phases(str(work), mode="developer", output_docx_name="最终论文.docx", qa_level="basic", project_root=str(work), deps=failing_deps),
        "QA should block the pipeline when the required QA report fails",
    )


@case
def pipeline_qa_contract_checks_strict_and_visual_reports() -> None:
    def passing_qa(out_dir, mode, output_docx_name):
        report = {"passed": True, "issues": [], "counts": {}, "mode": mode}
        report.update(qa_status_fields(report["passed"], report["issues"]))
        return report

    def invalid_conformance(out_dir, mode, output_docx_name, project_root):
        return {"passed": True, "issues": [], "counts": {}, "mode": mode}

    strict_work = new_workdir("pipeline_qa_contract_strict_report")
    strict_deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=invalid_conformance,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    strict_buffer = io.StringIO()
    with redirect_stdout(strict_buffer):
        assert_true(
            run_qa_phases(
                str(strict_work),
                mode="developer",
                output_docx_name="最终论文.docx",
                qa_level="strict",
                project_root=str(strict_work),
                deps=strict_deps,
            ),
            "contract warnings should not block an otherwise passing strict QA run",
        )
    strict_output = strict_buffer.getvalue()
    assert_true("[CONTRACT] conformance_report.json" in strict_output, f"strict report contract warning was not printed: {strict_output}")
    assert_true("QA_REPORT_STATUS_MISSING" in strict_output, f"strict report missing-status contract issue was not printed: {strict_output}")

    def passing_conformance(out_dir, mode, output_docx_name, project_root):
        report = {"passed": True, "issues": [], "counts": {}, "mode": mode}
        report.update(qa_status_fields(report["passed"], report["issues"]))
        return report

    def invalid_visual(out_dir, output_docx_name, project_root, render_all_pages, require_wps, golden_dir, update_golden):
        return {"passed": True, "issues": [], "counts": {}, "mode": "developer"}

    visual_work = new_workdir("pipeline_qa_contract_visual_report")
    visual_deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=passing_conformance,
        visual_check_and_write=invalid_visual,
        optional_import_detail=lambda name: "",
    )
    visual_buffer = io.StringIO()
    with redirect_stdout(visual_buffer):
        assert_true(
            run_qa_phases(
                str(visual_work),
                mode="developer",
                output_docx_name="最终论文.docx",
                qa_level="visual",
                project_root=str(visual_work),
                deps=visual_deps,
            ),
            "contract warnings should not block an otherwise passing visual QA run",
        )
    visual_output = visual_buffer.getvalue()
    assert_true("[CONTRACT] visual_report.json" in visual_output, f"visual report contract warning was not printed: {visual_output}")
    assert_true("QA_REPORT_STATUS_MISSING" in visual_output, f"visual report missing-status contract issue was not printed: {visual_output}")


@case
def pipeline_qa_points_to_conformance_report_when_strict_fails() -> None:
    work = new_workdir("pipeline_qa_conformance_hint")

    def passing_qa(out_dir, mode, output_docx_name):
        return {"passed": True, "issues": [], "counts": {}, "mode": mode, "repair_plan": {"steps": [], "passed": True}}

    def failing_conformance(out_dir, mode, output_docx_name, project_root):
        return {
            "passed": False,
            "issues": [
                {"severity": "error", "code": "STYLE_MISMATCH", "message": f"bad style {idx}", "detail": "body"}
                for idx in range(10)
            ],
            "counts": {},
            "next_action": "Fix conformance mismatch.",
        }

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=failing_conformance,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    buf = io.StringIO()
    with redirect_stdout(buf):
        ok = run_qa_phases(str(work), mode="developer", output_docx_name="最终论文.docx", qa_level="strict", project_root=str(work), deps=deps)
    output = buf.getvalue()
    assert_true(not ok, "strict QA should fail when conformance fails")
    assert_true("conformance_report.md" in output, f"conformance report was not routed in terminal hint: {output}")
    assert_true("请看 conformance_report.md" in output, f"overflow summary pointed at the wrong report: {output}")
    assert_true("结构 QA 已通过" in output, f"terminal hint should avoid implying structural QA repair plan is enough: {output}")


@case
def pipeline_qa_terminal_labels_warning_only_reports() -> None:
    work = new_workdir("pipeline_qa_warning_terminal")

    def warning_qa(out_dir, mode, output_docx_name):
        return {
            "passed": True,
            "issues": [
                {
                    "severity": "warning",
                    "code": "REFERENCES_MISSING",
                    "message": "references missing",
                    "active_owner": "content_parser.py",
                }
            ],
            "counts": {},
            "mode": mode,
            "next_action": "QA 已通过但有警告需要人工确认。优先处理 `REFERENCES_MISSING`：请补参考文献或确认无需参考文献。",
            "repair_plan": {"steps": [], "passed": True},
        }

    deps = QADependencies(
        qa_check_and_write=warning_qa,
        conformance_check_and_write=None,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    buf = io.StringIO()
    with redirect_stdout(buf):
        ok = run_qa_phases(str(work), mode="developer", output_docx_name="最终论文.docx", qa_level="basic", project_root=str(work), deps=deps)
    output = buf.getvalue()
    assert_true(ok, "warning-only structural QA should not block the pipeline")
    assert_true("通过但有警告" in output, f"terminal summary hid warning-only status: {output}")
    assert_true("下一步" in output and "REFERENCES_MISSING" in output, f"terminal warning summary should show next action and code: {output}")
    assert_true("请补参考文献" in output, f"terminal warning summary lost beginner action: {output}")


@case
def pipeline_qa_writes_report_when_structural_dependency_missing() -> None:
    work = new_workdir("pipeline_qa_missing_structural_report")
    raw_detail = f"missing module at {work / 'private' / 'qa_checker.py'}"
    write_workflow_mode(
        str(work),
        mode="developer",
        template_path=str(work / "Templates" / "template.docx"),
        content_path=str(work / "Inputs" / "content.docx"),
        run_qa=True,
        qa_level="basic",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
    )

    deps = QADependencies(
        qa_check_and_write=None,
        conformance_check_and_write=None,
        visual_check_and_write=None,
        optional_import_detail=lambda name: raw_detail,
    )
    buf = io.StringIO()
    with redirect_stdout(buf):
        ok = run_qa_phases(str(work), mode="developer", output_docx_name="最终论文.docx", qa_level="basic", project_root=str(work), deps=deps)
    output = buf.getvalue()
    report = json.loads((work / "qa_report.json").read_text(encoding="utf-8"))
    plan = json.loads((work / "qa_repair_plan.json").read_text(encoding="utf-8"))
    report_text = json.dumps(report, ensure_ascii=False)
    plan_text = json.dumps(plan, ensure_ascii=False)
    assert_true(not ok, "structural QA should fail closed when qa_checker dependency is missing")
    assert_true((work / "qa_report.md").exists(), "missing structural dependency should write qa_report.md")
    assert_true((work / "qa_repair_plan.md").exists(), "missing structural dependency should write qa_repair_plan.md")
    assert_true((work / "qa_fix_prompt.txt").exists(), "missing structural dependency should write qa_fix_prompt.txt")
    assert_true(report["issues"][0]["code"] == "STRUCTURAL_QA_UNAVAILABLE", f"wrong structural dependency issue: {report}")
    assert_true(plan.get("resume_scope") == "full_pipeline", f"structural dependency repair plan should route to full pipeline: {plan}")
    assert_true("run_pipeline.py" in plan.get("resume_command", ""), f"structural dependency plan should keep a rerun command: {plan}")
    assert_true("--template template.docx --content content.docx" in plan.get("resume_command", ""), f"structural dependency rerun command lost file arguments: {plan}")
    assert_true("qa_checker.py" in plan.get("next_action", ""), f"structural dependency plan lost concrete next action: {plan}")
    assert_true(str(work) not in report_text and str(work) not in plan_text, "structural dependency handoff leaked an absolute path")
    assert_true("<PROJECT>" in report_text and "<PROJECT>" in plan_text, f"structural dependency detail was not sanitized: {report} {plan}")
    assert_true("qa_report.md" in output and "qa_repair_plan.md" in output, f"terminal output did not route to structural reports: {output}")


@case
def pipeline_qa_writes_report_when_structural_checker_crashes() -> None:
    work = new_workdir("pipeline_qa_structural_crash_report")
    private_path = work / "private" / "qa_checker.py"
    write_workflow_mode(
        str(work),
        mode="developer",
        template_path=str(work / "Templates" / "template.docx"),
        content_path=str(work / "Inputs" / "content.docx"),
        run_qa=True,
        qa_level="basic",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
    )

    def crashing_qa(out_dir, mode, output_docx_name):
        raise RuntimeError(f"synthetic QA crash at {private_path}")

    deps = QADependencies(
        qa_check_and_write=crashing_qa,
        conformance_check_and_write=None,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    buf = io.StringIO()
    with redirect_stdout(buf):
        ok = run_qa_phases(str(work), mode="developer", output_docx_name="最终论文.docx", qa_level="basic", project_root=str(work), deps=deps)
    output = buf.getvalue()
    report = json.loads((work / "qa_report.json").read_text(encoding="utf-8"))
    plan = json.loads((work / "qa_repair_plan.json").read_text(encoding="utf-8"))
    report_text = json.dumps(report, ensure_ascii=False)
    plan_text = json.dumps(plan, ensure_ascii=False)
    assert_true(not ok, "structural QA crash should fail closed")
    assert_true(report["issues"][0]["code"] == "STRUCTURAL_QA_FAILED", f"wrong structural crash issue: {report}")
    assert_true(plan.get("resume_scope") == "full_pipeline", f"structural crash should route to full pipeline: {plan}")
    assert_true("run_pipeline.py" in plan.get("resume_command", ""), f"structural crash should preserve a rerun command: {plan}")
    assert_true("qa_checker.py" in plan.get("next_action", ""), f"structural crash plan lost concrete next action: {plan}")
    assert_true(str(work) not in report_text and str(work) not in plan_text, "structural crash handoff leaked an absolute path")
    assert_true("<PROJECT>" in report_text and "<PROJECT>" in plan_text, f"structural crash detail was not sanitized: {report} {plan}")
    assert_true("qa_report.md" in output and "qa_repair_plan.md" in output, f"terminal output did not route to structural crash reports: {output}")


@case
def pipeline_qa_writes_report_when_conformance_dependency_missing() -> None:
    work = new_workdir("pipeline_qa_missing_conformance_report")
    raw_detail = f"missing module at {work / 'private' / 'qa_conformance.py'}"

    def passing_qa(out_dir, mode, output_docx_name):
        return {"passed": True, "issues": [], "counts": {}, "mode": mode, "repair_plan": {"steps": [], "passed": True}}

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=None,
        visual_check_and_write=None,
        optional_import_detail=lambda name: raw_detail,
    )
    buf = io.StringIO()
    with redirect_stdout(buf):
        ok = run_qa_phases(str(work), mode="developer", output_docx_name="最终论文.docx", qa_level="strict", project_root=str(work), deps=deps)
    output = buf.getvalue()
    report = json.loads((work / "conformance_report.json").read_text(encoding="utf-8"))
    report_text = json.dumps(report, ensure_ascii=False)
    assert_true(not ok, "strict QA should fail closed when conformance dependency is missing")
    assert_true((work / "conformance_report.md").exists(), "missing conformance dependency should write markdown report")
    assert_true(report["issues"][0]["code"] == "CONFORMANCE_QA_UNAVAILABLE", f"wrong dependency issue: {report}")
    assert_true(str(work) not in report_text, "QA dependency report leaked an absolute path")
    assert_true("<PROJECT>" in report_text, f"QA dependency report did not sanitize detail: {report}")
    assert_true("conformance_report.md" in output, f"terminal output did not route to conformance report: {output}")


@case
def pipeline_qa_writes_report_when_conformance_checker_crashes() -> None:
    work = new_workdir("pipeline_qa_conformance_crash_report")
    private_path = work / "private" / "qa_conformance.py"

    def passing_qa(out_dir, mode, output_docx_name):
        return {"passed": True, "issues": [], "counts": {}, "mode": mode, "repair_plan": {"steps": [], "passed": True}}

    def crashing_conformance(out_dir, mode, output_docx_name, project_root):
        raise RuntimeError(f"synthetic conformance crash at {private_path}")

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=crashing_conformance,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "",
    )
    buf = io.StringIO()
    with redirect_stdout(buf):
        ok = run_qa_phases(str(work), mode="developer", output_docx_name="最终论文.docx", qa_level="strict", project_root=str(work), deps=deps)
    output = buf.getvalue()
    report = json.loads((work / "conformance_report.json").read_text(encoding="utf-8"))
    report_text = json.dumps(report, ensure_ascii=False)
    assert_true(not ok, "strict QA should fail closed when conformance checker crashes")
    assert_true((work / "conformance_report.md").exists(), "conformance crash should write markdown report")
    assert_true(report["issues"][0]["code"] == "CONFORMANCE_QA_FAILED", f"wrong conformance crash issue: {report}")
    assert_true(str(work) not in report_text, "conformance crash report leaked an absolute path")
    assert_true("<PROJECT>" in report_text, f"conformance crash report did not sanitize detail: {report}")
    assert_true("conformance_report.md" in output, f"terminal output did not route to conformance crash report: {output}")


@case
def pipeline_qa_writes_report_when_visual_dependency_missing() -> None:
    work = new_workdir("pipeline_qa_missing_visual_report")

    def passing_qa(out_dir, mode, output_docx_name):
        return {"passed": True, "issues": [], "counts": {}, "mode": mode, "repair_plan": {"steps": [], "passed": True}}

    def passing_conformance(out_dir, mode, output_docx_name, project_root):
        return {"passed": True, "issues": [], "counts": {}, "next_action": "ok"}

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=passing_conformance,
        visual_check_and_write=None,
        optional_import_detail=lambda name: "missing visual",
    )
    buf = io.StringIO()
    with redirect_stdout(buf):
        ok = run_qa_phases(str(work), mode="developer", output_docx_name="最终论文.docx", qa_level="visual", project_root=str(work), deps=deps)
    output = buf.getvalue()
    report = json.loads((work / "visual_report.json").read_text(encoding="utf-8"))
    assert_true(not ok, "visual QA should fail closed when visual dependency is missing")
    assert_true((work / "visual_report.md").exists(), "missing visual dependency should write markdown report")
    assert_true(report["issues"][0]["code"] == "VISUAL_QA_UNAVAILABLE", f"wrong dependency issue: {report}")
    assert_true("visual_report.md" in output, f"terminal output did not route to visual report: {output}")


@case
def pipeline_qa_writes_report_when_visual_checker_crashes() -> None:
    work = new_workdir("pipeline_qa_visual_crash_report")
    private_path = work / "private" / "qa_visual.py"

    def passing_qa(out_dir, mode, output_docx_name):
        return {"passed": True, "issues": [], "counts": {}, "mode": mode, "repair_plan": {"steps": [], "passed": True}}

    def passing_conformance(out_dir, mode, output_docx_name, project_root):
        return {"passed": True, "issues": [], "counts": {}, "next_action": "ok"}

    def crashing_visual(out_dir, output_docx_name, project_root, render_all_pages, require_wps, golden_dir, update_golden):
        raise RuntimeError(f"synthetic visual crash at {private_path}")

    deps = QADependencies(
        qa_check_and_write=passing_qa,
        conformance_check_and_write=passing_conformance,
        visual_check_and_write=crashing_visual,
        optional_import_detail=lambda name: "",
    )
    buf = io.StringIO()
    with redirect_stdout(buf):
        ok = run_qa_phases(str(work), mode="developer", output_docx_name="最终论文.docx", qa_level="visual", project_root=str(work), deps=deps)
    output = buf.getvalue()
    report = json.loads((work / "visual_report.json").read_text(encoding="utf-8"))
    report_text = json.dumps(report, ensure_ascii=False)
    assert_true(not ok, "visual QA should fail closed when visual checker crashes")
    assert_true((work / "visual_report.md").exists(), "visual crash should write markdown report")
    assert_true(report["issues"][0]["code"] == "VISUAL_QA_FAILED", f"wrong visual crash issue: {report}")
    assert_true(str(work) not in report_text, "visual crash report leaked an absolute path")
    assert_true("<PROJECT>" in report_text, f"visual crash report did not sanitize detail: {report}")
    assert_true("visual_report.md" in output, f"terminal output did not route to visual crash report: {output}")


@case
def pipeline_strict_and_visual_reports_surface_specific_next_actions() -> None:
    from qa_conformance_modules.reports import build_report as build_conformance_report
    from qa_visual_modules.checks import _next_action as visual_next_action

    work = new_workdir("pipeline_report_specific_next_actions")
    placeholder_report = build_conformance_report(
        str(work),
        "user",
        {},
        [{"code": "PLACEHOLDER_TEXT_LEFT", "severity": "error", "message": "placeholder remains", "detail": ""}],
        project_root=str(work),
    )
    assert_true("PLACEHOLDER_TEXT_LEFT" in placeholder_report["next_action"], f"placeholder strict action lost the issue code: {placeholder_report}")
    assert_true("占位符" in placeholder_report["next_action"], f"placeholder strict report used a generic action: {placeholder_report}")
    assert_true("strict QA" in placeholder_report["next_action"], f"strict report should tell users how to verify after fixing: {placeholder_report}")

    word_field_report = build_conformance_report(
        str(work),
        "user",
        {},
        [{"code": "WORD_FIELD_ERROR", "severity": "error", "message": "field error remains", "detail": ""}],
        project_root=str(work),
    )
    assert_true("WORD_FIELD_ERROR" in word_field_report["next_action"], f"Word field strict action lost the issue code: {word_field_report}")
    assert_true("Word 域" in word_field_report["next_action"] and "重跑 strict QA" in word_field_report["next_action"], f"Word field strict report used a generic action: {word_field_report}")

    strict_warning_report = build_conformance_report(
        str(work),
        "user",
        {},
        [{"code": "STYLE_MISMATCH", "severity": "warning", "message": "style needs review", "detail": "body font differs"}],
        project_root=str(work),
    )
    assert_true(strict_warning_report["passed"] is True, f"warning-only strict report should remain non-blocking: {strict_warning_report}")
    assert_true("STYLE_MISMATCH" in strict_warning_report["next_action"], f"strict warning action lost the issue code: {strict_warning_report}")
    assert_true("警告" in strict_warning_report["next_action"], f"strict warning action should not sound like a plain pass: {strict_warning_report}")
    assert_true("样式" in strict_warning_report["next_action"] and "strict QA" in strict_warning_report["next_action"], f"strict warning action should tell users what to check and how to rerun: {strict_warning_report}")
    assert_true("机器检查已通过" not in strict_warning_report["next_action"], f"warning-only strict action should not hide the warning: {strict_warning_report}")

    invalid_pages = visual_next_action([{"code": "PDF_PAGE_COUNT_INVALID", "severity": "error", "message": "no pages"}])
    assert_true("PDF_PAGE_COUNT_INVALID" in invalid_pages, f"invalid-page visual action lost the issue code: {invalid_pages}")
    assert_true("没有有效页面" in invalid_pages and "重跑 visual QA" in invalid_pages, f"invalid-page visual action is too generic: {invalid_pages}")

    unreadable_pages = visual_next_action([{"code": "PAGE_IMAGE_UNREADABLE", "severity": "error", "message": "bad png"}])
    assert_true("PAGE_IMAGE_UNREADABLE" in unreadable_pages, f"unreadable-page visual action lost the issue code: {unreadable_pages}")
    assert_true("不可读页面" in unreadable_pages and "重跑 visual QA" in unreadable_pages, f"unreadable-page visual action is too generic: {unreadable_pages}")

    wps_mismatch = visual_next_action([{"code": "WPS_PAGE_COUNT_MISMATCH", "severity": "error", "message": "wps pages differ"}])
    assert_true("WPS_PAGE_COUNT_MISMATCH" in wps_mismatch, f"WPS mismatch visual action lost the issue code: {wps_mismatch}")
    assert_true("WPS" in wps_mismatch and "分页差异" in wps_mismatch and "重跑 visual QA" in wps_mismatch, f"WPS mismatch action should tell users how to resume: {wps_mismatch}")

    wps_page_size = visual_next_action([{"code": "WPS_PAGE_SIZE_MISMATCH", "severity": "error", "message": "wps page size differs"}])
    assert_true("WPS" in wps_page_size and ("纸张" in wps_page_size or "页面尺寸" in wps_page_size) and "重跑 visual QA" in wps_page_size, f"WPS page-size mismatch action should tell users how to resume: {wps_page_size}")

    wps_text_mismatch = visual_next_action([{"code": "WPS_TEXT_PAGE_MISMATCH", "severity": "error", "message": "wps text pages missing"}])
    assert_true("WPS" in wps_text_mismatch and ("文本" in wps_text_mismatch or "内容" in wps_text_mismatch) and "重跑 visual QA" in wps_text_mismatch, f"WPS text mismatch action should tell users how to resume: {wps_text_mismatch}")
    assert_true("rendered_word.txt" in wps_text_mismatch and "rendered_wps.txt" in wps_text_mismatch, f"WPS text mismatch action should name preserved diagnostics: {wps_text_mismatch}")

    wps_sample_render_failed = visual_next_action([{"code": "WPS_SAMPLE_RENDER_FAILED", "severity": "error", "message": "wps sample render failed"}])
    assert_true("visual_qa/samples/" in wps_sample_render_failed and "visual_qa/wps/samples/" in wps_sample_render_failed, f"WPS sample render action should name both sample directories: {wps_sample_render_failed}")
    assert_true("重跑 visual QA" in wps_sample_render_failed, f"WPS sample render action should tell users how to resume: {wps_sample_render_failed}")

    wps_sample_mismatch = visual_next_action([{"code": "WPS_SAMPLE_IMAGE_MISMATCH", "severity": "error", "message": "wps sample images differ"}])
    assert_true("WPS" in wps_sample_mismatch and ("样张" in wps_sample_mismatch or "PNG" in wps_sample_mismatch or "画面" in wps_sample_mismatch) and "重跑 visual QA" in wps_sample_mismatch, f"WPS sample-image mismatch action should tell users how to resume: {wps_sample_mismatch}")

    wps_pdfinfo = visual_next_action([{"code": "WPS_PDFINFO_FAILED", "severity": "error", "message": "wps pdfinfo failed"}])
    assert_true("WPS" in wps_pdfinfo and "PDF" in wps_pdfinfo and "重跑 visual QA" in wps_pdfinfo, f"WPS PDF metadata action should tell users how to resume: {wps_pdfinfo}")

    wps_invalid_pages = visual_next_action([{"code": "WPS_PAGE_COUNT_INVALID", "severity": "error", "message": "wps pages missing"}])
    assert_true("WPS" in wps_invalid_pages and "有效页面" in wps_invalid_pages and "重跑 visual QA" in wps_invalid_pages, f"WPS invalid page-count action should tell users how to resume: {wps_invalid_pages}")

    missing_golden = visual_next_action([{"code": "GOLDEN_BASELINE_MISSING", "severity": "warning", "message": "no baseline"}])
    assert_true("GOLDEN_BASELINE_MISSING" in missing_golden, f"warning-only visual action lost the issue code: {missing_golden}")
    assert_true("黄金基线" in missing_golden and "--update-golden" in missing_golden, f"warning-only visual action should still guide users: {missing_golden}")
    assert_true("机器检查已通过" not in missing_golden, f"warning-only visual action should not sound fully done: {missing_golden}")


@case
def pipeline_warning_only_markdown_result_labels_are_explicit() -> None:
    from qa_checker_modules.reports import (
        repair_plan_to_markdown as structural_repair_markdown,
        report_to_markdown as structural_report_markdown,
    )
    from qa_conformance_modules.reports import report_to_markdown as conformance_report_markdown
    from qa_visual_modules.reports import report_to_markdown as visual_report_markdown

    structural_report = {
        "mode": "user",
        "passed": True,
        "output_dir_name": "demo",
        "next_action": "优先处理 `REFERENCES_MISSING`。",
        "issues": [{"code": "REFERENCES_MISSING", "severity": "warning", "message": "references missing"}],
        "counts": {},
    }
    structural_plan = {
        "passed": True,
        "warnings": 1,
        "summary": "QA 没有阻断错误，但发现 1 个警告。",
        "output_dir": "Outputs/demo",
        "next_action": "优先处理 `REFERENCES_MISSING`。",
        "steps": [{"code": "REFERENCES_MISSING", "severity": "warning", "title": "没有识别到参考文献"}],
    }
    conformance_report = {
        "passed": True,
        "mode": "user",
        "output_dir_name": "demo",
        "next_action": "strict 合规 QA 没有阻断错误，但有警告 `STYLE_MISMATCH` 需要人工确认。",
        "issues": [{"code": "STYLE_MISMATCH", "severity": "warning", "message": "style needs review"}],
        "counts": {},
    }
    visual_report = {
        "passed": True,
        "output_dir_name": "demo",
        "next_action": "黄金基线缺失；首次建立视觉基线时可用 --update-golden 生成。",
        "issues": [{"code": "GOLDEN_BASELINE_MISSING", "severity": "warning", "message": "no baseline"}],
        "counts": {},
    }

    assert_true("- 结果：通过但有警告" in structural_report_markdown(structural_report), "structural warning-only report should label pass-with-warning")
    assert_true("- 结果：已通过但有警告" in structural_repair_markdown(structural_plan), "structural warning-only repair plan should label pass-with-warning")
    assert_true("- 结果：通过但有警告" in conformance_report_markdown(conformance_report), "strict warning-only report should label pass-with-warning")
    assert_true("- 结果：通过但有警告" in visual_report_markdown(visual_report), "visual warning-only report should label pass-with-warning")


@case
def pipeline_conformance_report_markdown_lists_review_artifacts() -> None:
    from qa_conformance_modules.reports import report_to_markdown as conformance_report_markdown

    report = {
        "passed": False,
        "mode": "developer",
        "output_dir_name": "demo",
        "next_action": "对照内容和最终 DOCX 修复遗漏段落后重跑 strict QA。",
        "counts": {"content_paragraphs": 10, "docx_paragraphs": 9},
        "issues": [
            {
                "code": "CONTENT_PARAGRAPH_MISSING",
                "severity": "error",
                "message": "content paragraph missing",
                "detail": "paragraph=9",
            }
        ],
    }
    markdown = conformance_report_markdown(report)
    assert_true("## 核对入口" in markdown, f"strict report should list review artifacts: {markdown}")
    assert_true("Outputs/demo/最终论文.docx" in markdown, f"strict report should point to final DOCX: {markdown}")
    assert_true("Outputs/demo/内容提取.md" in markdown, f"strict report should point to content summary: {markdown}")
    assert_true("Outputs/demo/content.json" in markdown, f"strict report should point to structured content: {markdown}")
    assert_true("Outputs/demo/build_manifest.json" in markdown, f"strict report should point to build manifest: {markdown}")
    assert_true("Outputs/demo/template_requirements.json" in markdown, f"strict report should point to template requirements: {markdown}")


@case
def pipeline_visual_report_markdown_lists_diagnostic_artifacts() -> None:
    from qa_visual_modules.reports import report_to_markdown as visual_report_markdown

    report = {
        "passed": False,
        "output_dir_name": "demo",
        "next_action": "比较 Word/WPS 样张后重跑 visual QA。",
        "counts": {"sample_images": 2, "wps_sample_images": 2},
        "issues": [
            {
                "code": "WPS_SAMPLE_IMAGE_MISMATCH",
                "severity": "error",
                "message": "WPS sample images differ.",
                "detail": "pages=1",
            }
        ],
        "artifacts": {
            "pdf": "<PROJECT>/Outputs/demo/visual_qa/rendered.pdf",
            "word_text": "<PROJECT>/Outputs/demo/visual_qa/rendered_word.txt",
            "wps_pdf": "<PROJECT>/Outputs/demo/visual_qa/rendered_wps.pdf",
            "wps_text": "<PROJECT>/Outputs/demo/visual_qa/rendered_wps.txt",
            "samples": [
                "<PROJECT>/Outputs/demo/visual_qa/samples/page_001-01.png",
                "<PROJECT>/Outputs/demo/visual_qa/samples/page_002-02.png",
            ],
            "wps_samples": [
                "<PROJECT>/Outputs/demo/visual_qa/wps/samples/page_001-01.png",
                "<PROJECT>/Outputs/demo/visual_qa/wps/samples/page_002-02.png",
            ],
        },
    }
    markdown = visual_report_markdown(report)
    assert_true("## 诊断产物" in markdown, f"visual report should list diagnostic artifacts: {markdown}")
    assert_true("rendered.pdf" in markdown and "rendered_wps.pdf" in markdown, f"visual report should name PDF artifacts: {markdown}")
    assert_true("rendered_word.txt" in markdown and "rendered_wps.txt" in markdown, f"visual report should name text diagnostics: {markdown}")
    assert_true("visual_qa/samples/" in markdown and "visual_qa/wps/samples/" in markdown, f"visual report should name both sample directories: {markdown}")


@case
def pipeline_summary_mentions_outputs_and_mode() -> None:
    summary = build_completion_summary("2026-05-27_demo", "最终论文.docx", "developer")
    assert_true("Outputs/2026-05-27_demo/" in summary, "output directory missing from completion summary")
    assert_true("最终论文.docx" in summary, "output docx missing from completion summary")
    assert_true("当前模式: 开发者" in summary, "developer mode missing from completion summary")
    assert_true("build_generated.py" in summary, "user fine-tuning target missing from completion summary")
    assert_true("agent_summary.md" in summary, "agent handoff summary missing from completion summary")


@case
def pipeline_agent_summary_writes_user_handoff() -> None:
    work = new_workdir("pipeline_agent_summary")
    write_workflow_mode(
        str(work),
        mode="user",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="strict",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
        auto_repair=True,
        agent_auto=True,
        repair_max_rounds=3,
        repair_stop_no_improve=2,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    write_json(
        work / "qa_report.json",
        {
            "passed": True,
            "issues": [],
            "counts": {},
            "next_action": "ok",
        },
    )
    write_json(
        work / "conformance_report.json",
        {
            "passed": True,
            "issues": [],
            "counts": {},
            "next_action": "ok",
        },
    )
    write_json(
        work / "repair_loop_report.json",
        {
            "status": "converged",
            "rounds_run": 1,
            "final_errors": 0,
            "final_warnings": 1,
            "manual_check_required": ["打开最终 DOCX 核对图片。"],
        },
    )
    json_path, md_path = write_agent_summary(str(work), "2026-05-29_demo", "最终论文.docx", "user")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    assert_true(summary["agent_auto"] is True and summary["auto_repair"] is True, "agent summary should preserve workflow flags")
    assert_true(summary["output_docx"] == "Outputs/2026-05-29_demo/最终论文.docx", "agent summary should use run-relative output paths")
    assert_true(summary["repair_loop"]["status"] == "converged", "repair loop status missing from agent summary")
    assert_true(summary["manual_check_required"] == ["打开最终 DOCX 核对图片。"], "manual checks were not preserved")
    assert_true("Agent 排版摘要" in text and "最终论文" in text, "agent summary markdown missing user-facing handoff")


@case
def pipeline_agent_summary_does_not_claim_qa_passed_when_qa_disabled() -> None:
    work = new_workdir("pipeline_agent_summary_no_qa")
    write_workflow_mode(
        str(work),
        mode="user",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=False,
        qa_level="basic",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
        auto_repair=False,
        agent_auto=False,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    json_path, _ = write_agent_summary(str(work), "2026-05-29_noqa", "最终论文.docx", "user")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    assert_true(summary["status_label"] == "已生成 DOCX，未运行自动 QA", f"summary overstated QA status: {summary['status_label']}")


@case
def pipeline_agent_summary_requires_expected_strict_reports() -> None:
    work = new_workdir("pipeline_agent_summary_missing_strict_report")
    write_workflow_mode(
        str(work),
        mode="user",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="strict",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
        auto_repair=True,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    write_json(work / "qa_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    json_path, md_path = write_agent_summary(str(work), "2026-05-29_missing_strict", "最终论文.docx", "user")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    assert_true(summary["status_label"] == "已生成 DOCX，但自动 QA 报告不完整", f"summary overstated QA status: {summary['status_label']}")
    assert_true(summary["missing_required_reports"] == ["conformance"], f"missing strict report was not recorded: {summary}")
    assert_true("DOCX/XML 合规 QA 未生成" in text, "agent summary should tell users which required QA report is missing")


@case
def pipeline_agent_summary_failed_before_docx_does_not_blame_missing_later_qa() -> None:
    work = new_workdir("pipeline_agent_summary_failed_before_docx")
    write_workflow_mode(
        str(work),
        mode="user",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="strict",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
        auto_repair=False,
        agent_auto=True,
    )
    write_json(
        work / "qa_report.json",
        {
            "passed": False,
            "issues": [{"severity": "error", "code": "EXTRACTION_VERIFICATION_FAILED"}],
            "counts": {},
            "next_action": "查看 qa_repair_plan.md，先处理提取验证失败。",
        },
    )
    json_path, md_path = write_agent_summary(
        str(work),
        "2026-05-29_failed_before_docx",
        "最终论文.docx",
        "user",
        pipeline_status="failed",
        note="模板格式提取多次验证无法收敛。",
    )
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    assert_true(summary["missing_required_reports"] == [], f"early failure should not require later QA reports: {summary}")
    assert_true("合规 QA 未生成" not in text, "early failure should not send users toward conformance QA")
    assert_true("提取验证失败" in text or "qa_repair_plan.md" in text, "early failure should preserve the real next action")


@case
def pipeline_agent_summary_surfaces_structural_repair_steps() -> None:
    work = new_workdir("pipeline_agent_summary_repair_steps")
    write_workflow_mode(
        str(work),
        mode="user",
        template_path="template.docx",
        content_path="content.md",
        run_qa=True,
        qa_level="basic",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
        auto_repair=True,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    _fake_repair_report(str(work), code="CONTENT_IMAGE_MISSING", message="missing markdown image")

    json_path, md_path = write_agent_summary(str(work), "2026-05-31_missing_image", "最终论文.docx", "user")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    action_text = "\n".join(summary.get("next_actions") or [])

    assert_true("CONTENT_IMAGE_MISSING" in action_text, f"summary lost the structural issue code: {summary}")
    assert_true("把缺失图片放回" in action_text, f"summary did not surface the beginner action: {summary}")
    assert_true("CONTENT_IMAGE_MISSING" in text and "把缺失图片放回" in text, "agent summary markdown should show the concrete repair step")


@case
def pipeline_agent_summary_surfaces_structural_warning_steps() -> None:
    from qa_checker_modules.report_phase import build_report as build_structural_report
    from qa_checker_modules.reports import write_reports as write_structural_reports

    work = new_workdir("pipeline_agent_summary_structural_warning_steps")
    write_workflow_mode(
        str(work),
        mode="user",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="basic",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
        auto_repair=False,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    report = build_structural_report(
        str(work),
        "user",
        {},
        [{"code": "REFERENCES_MISSING", "severity": "warning", "message": "references missing"}],
    )
    write_structural_reports(report, str(work))

    json_path, md_path = write_agent_summary(str(work), "2026-05-31_structural_warning", "最终论文.docx", "user")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    action_text = "\n".join(summary.get("next_actions") or summary.get("manual_check_required") or [])

    assert_true("警告" in summary["status_label"], f"summary should label warning-only structural QA: {summary}")
    assert_true("REFERENCES_MISSING" in action_text, f"summary lost the structural warning code: {summary}")
    assert_true("参考文献" in action_text, f"summary lost the structural warning action: {summary}")
    assert_true("REFERENCES_MISSING" in text, "agent summary markdown should show the structural warning code")


@case
def pipeline_agent_summary_surfaces_conformance_issue_steps() -> None:
    work = new_workdir("pipeline_agent_summary_conformance_steps")
    write_workflow_mode(
        str(work),
        mode="user",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="strict",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
        auto_repair=True,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    write_json(work / "qa_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(
        work / "conformance_report.json",
        {
            "passed": False,
            "issues": [
                {
                    "code": "STYLE_MISMATCH",
                    "severity": "error",
                    "message": "paragraph style mismatch",
                    "detail": "body: eastAsia font Times New Roman != 宋体",
                }
            ],
            "counts": {},
            "next_action": "Fix generic conformance mismatch.",
        },
    )

    json_path, md_path = write_agent_summary(str(work), "2026-05-31_conformance", "最终论文.docx", "user")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    action_text = "\n".join(summary.get("next_actions") or [])

    assert_true("STYLE_MISMATCH" in action_text, f"summary lost the conformance issue code: {summary}")
    assert_true("conformance_report.md" in action_text and "重跑 strict QA" in action_text, f"summary did not surface a strict-QA next action: {summary}")
    assert_true("STYLE_MISMATCH" in text, "agent summary markdown should show the conformance issue code")


@case
def pipeline_agent_summary_surfaces_visual_issue_steps() -> None:
    work = new_workdir("pipeline_agent_summary_visual_steps")
    write_workflow_mode(
        str(work),
        mode="developer",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="visual",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
        auto_repair=False,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    write_json(work / "qa_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(work / "conformance_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(
        work / "visual_report.json",
        {
            "passed": False,
            "issues": [
                {
                    "code": "PDFINFO_UNAVAILABLE",
                    "severity": "error",
                    "message": "pdfinfo is not available",
                    "detail": "",
                }
            ],
            "counts": {},
            "next_action": "Install render tools.",
        },
    )

    json_path, md_path = write_agent_summary(str(work), "2026-05-31_visual", "最终论文.docx", "developer")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    action_text = "\n".join(summary.get("next_actions") or [])

    assert_true("PDFINFO_UNAVAILABLE" in action_text, f"summary lost the visual issue code: {summary}")
    assert_true("Poppler" in action_text and "visual QA" in action_text, f"summary did not surface a visual-QA next action: {summary}")
    assert_true("PDFINFO_UNAVAILABLE" in text, "agent summary markdown should show the visual issue code")


@case
def pipeline_agent_summary_surfaces_visual_warning_steps() -> None:
    work = new_workdir("pipeline_agent_summary_visual_warning_steps")
    write_workflow_mode(
        str(work),
        mode="developer",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="visual",
        golden_dir="TestData/GoldenBaselines",
        update_golden=False,
        require_wps=False,
        auto_repair=False,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    write_json(work / "qa_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(work / "conformance_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(
        work / "visual_report.json",
        {
            "passed": True,
            "issues": [
                {
                    "code": "GOLDEN_BASELINE_MISSING",
                    "severity": "warning",
                    "message": "Golden baseline was requested but no baseline exists.",
                    "detail": "",
                }
            ],
            "counts": {},
            "next_action": "黄金基线缺失；首次建立视觉基线时可用 --update-golden 生成。",
        },
    )

    json_path, md_path = write_agent_summary(str(work), "2026-05-31_visual_warning", "最终论文.docx", "developer")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    action_text = "\n".join(summary.get("next_actions") or summary.get("manual_check_required") or [])

    assert_true("警告" in summary["status_label"], f"summary should not hide warning-only QA status: {summary}")
    assert_true("GOLDEN_BASELINE_MISSING" in action_text, f"summary lost the visual warning code: {summary}")
    assert_true("--update-golden" in action_text and "visual QA" in action_text, f"summary did not surface a visual warning next action: {summary}")
    assert_true("GOLDEN_BASELINE_MISSING" in text, "agent summary markdown should show the visual warning code")


@case
def pipeline_agent_summary_surfaces_wps_mismatch_rerun_step() -> None:
    work = new_workdir("pipeline_agent_summary_wps_mismatch")
    write_workflow_mode(
        str(work),
        mode="developer",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="visual",
        golden_dir=None,
        update_golden=False,
        require_wps=True,
        auto_repair=False,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    write_json(work / "qa_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(work / "conformance_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(
        work / "visual_report.json",
        {
            "passed": False,
            "issues": [
                {
                    "code": "WPS_PAGE_COUNT_MISMATCH",
                    "severity": "error",
                    "message": "WPS page count differs",
                    "detail": "word=10 wps=11",
                }
            ],
            "counts": {},
            "next_action": "Compare Word and WPS PDFs.",
        },
    )

    json_path, md_path = write_agent_summary(str(work), "2026-06-01_wps_mismatch", "最终论文.docx", "developer")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    action_text = "\n".join(summary.get("next_actions") or [])

    assert_true("WPS_PAGE_COUNT_MISMATCH" in action_text, f"summary lost the WPS mismatch issue code: {summary}")
    assert_true("重跑 visual QA" in action_text, f"summary should tell users to rerun visual QA after WPS mismatch repair: {summary}")
    assert_true("WPS_PAGE_COUNT_MISMATCH" in text and "重跑 visual QA" in text, "agent summary markdown should show the WPS rerun step")


@case
def pipeline_agent_summary_surfaces_wps_page_size_rerun_step() -> None:
    work = new_workdir("pipeline_agent_summary_wps_page_size")
    write_workflow_mode(
        str(work),
        mode="developer",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="visual",
        golden_dir=None,
        update_golden=False,
        require_wps=True,
        auto_repair=False,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    write_json(work / "qa_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(work / "conformance_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(
        work / "visual_report.json",
        {
            "passed": False,
            "issues": [
                {
                    "code": "WPS_PAGE_SIZE_MISMATCH",
                    "severity": "error",
                    "message": "WPS page size differs",
                    "detail": "word=595.3x841.9 wps=841.9x595.3",
                }
            ],
            "counts": {},
            "next_action": "",
        },
    )

    json_path, md_path = write_agent_summary(str(work), "2026-06-01_wps_page_size", "最终论文.docx", "developer")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    action_text = "\n".join(summary.get("next_actions") or [])

    assert_true("WPS_PAGE_SIZE_MISMATCH" in action_text, f"summary lost the WPS page-size issue code: {summary}")
    assert_true(("纸张" in action_text or "页面尺寸" in action_text) and "重跑 visual QA" in action_text, f"summary should tell users to rerun visual QA after WPS page-size repair: {summary}")
    assert_true("WPS_PAGE_SIZE_MISMATCH" in text and "重跑 visual QA" in text, "agent summary markdown should show the WPS page-size rerun step")


@case
def pipeline_agent_summary_surfaces_wps_text_rerun_step() -> None:
    work = new_workdir("pipeline_agent_summary_wps_text")
    write_workflow_mode(
        str(work),
        mode="developer",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="visual",
        golden_dir=None,
        update_golden=False,
        require_wps=True,
        auto_repair=False,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    write_json(work / "qa_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(work / "conformance_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(
        work / "visual_report.json",
        {
            "passed": False,
            "issues": [
                {
                    "code": "WPS_TEXT_PAGE_MISMATCH",
                    "severity": "error",
                    "message": "WPS text pages are missing",
                    "detail": "word_text_pages=12 wps_text_pages=0",
                }
            ],
            "counts": {"text_pages": 12, "wps_text_pages": 0},
            "next_action": "",
        },
    )

    json_path, md_path = write_agent_summary(str(work), "2026-06-01_wps_text", "最终论文.docx", "developer")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    action_text = "\n".join(summary.get("next_actions") or [])

    assert_true("WPS_TEXT_PAGE_MISMATCH" in action_text, f"summary lost the WPS text issue code: {summary}")
    assert_true(("文本" in action_text or "内容" in action_text) and "重跑 visual QA" in action_text, f"summary should tell users to rerun visual QA after WPS text repair: {summary}")
    assert_true("rendered_word.txt" in action_text and "rendered_wps.txt" in action_text, f"summary should name WPS text diagnostic files: {summary}")
    assert_true("WPS_TEXT_PAGE_MISMATCH" in text and "重跑 visual QA" in text, "agent summary markdown should show the WPS text rerun step")


@case
def pipeline_agent_summary_surfaces_wps_sample_rerun_step() -> None:
    work = new_workdir("pipeline_agent_summary_wps_sample")
    write_workflow_mode(
        str(work),
        mode="developer",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="visual",
        golden_dir=None,
        update_golden=False,
        require_wps=True,
        auto_repair=False,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    write_json(work / "qa_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(work / "conformance_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(
        work / "visual_report.json",
        {
            "passed": False,
            "issues": [
                {
                    "code": "WPS_SAMPLE_IMAGE_MISMATCH",
                    "severity": "error",
                    "message": "WPS sample images differ",
                    "detail": "pages=1",
                }
            ],
            "counts": {"sample_images": 2, "wps_sample_images": 2, "wps_sample_mismatches": [1]},
            "next_action": "",
        },
    )

    json_path, md_path = write_agent_summary(str(work), "2026-06-01_wps_sample", "最终论文.docx", "developer")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    action_text = "\n".join(summary.get("next_actions") or [])

    assert_true("WPS_SAMPLE_IMAGE_MISMATCH" in action_text, f"summary lost the WPS sample issue code: {summary}")
    assert_true(("样张" in action_text or "PNG" in action_text or "画面" in action_text) and "重跑 visual QA" in action_text, f"summary should tell users to rerun visual QA after WPS sample repair: {summary}")
    assert_true("visual_qa/samples/" in action_text and "visual_qa/wps/samples/" in action_text, f"summary should name both WPS/Word sample directories: {summary}")
    assert_true("WPS_SAMPLE_IMAGE_MISMATCH" in text and "重跑 visual QA" in text, "agent summary markdown should show the WPS sample rerun step")


@case
def pipeline_agent_summary_surfaces_wps_sample_render_paths() -> None:
    work = new_workdir("pipeline_agent_summary_wps_sample_render")
    write_workflow_mode(
        str(work),
        mode="developer",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="visual",
        golden_dir=None,
        update_golden=False,
        require_wps=True,
        auto_repair=False,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    write_json(work / "qa_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(work / "conformance_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(
        work / "visual_report.json",
        {
            "passed": False,
            "issues": [
                {
                    "code": "WPS_SAMPLE_RENDER_FAILED",
                    "severity": "error",
                    "message": "WPS sample images failed to render",
                    "detail": "pages=4 rendered=2",
                }
            ],
            "counts": {"sample_images": 4, "wps_sample_images": 2},
            "next_action": "",
        },
    )

    json_path, md_path = write_agent_summary(str(work), "2026-06-01_wps_sample_render", "最终论文.docx", "developer")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    action_text = "\n".join(summary.get("next_actions") or [])

    assert_true("WPS_SAMPLE_RENDER_FAILED" in action_text, f"summary lost the WPS sample render issue code: {summary}")
    assert_true("visual_qa/samples/" in action_text and "visual_qa/wps/samples/" in action_text, f"summary should name both sample directories: {summary}")
    assert_true("重跑 visual QA" in action_text, f"summary should tell users to rerun visual QA after WPS sample render repair: {summary}")
    assert_true("WPS_SAMPLE_RENDER_FAILED" in text and "visual_qa/wps/samples/" in text, "agent summary markdown should show the WPS sample render paths")


@case
def pipeline_agent_summary_surfaces_wps_pdfinfo_rerun_step() -> None:
    work = new_workdir("pipeline_agent_summary_wps_pdfinfo")
    write_workflow_mode(
        str(work),
        mode="developer",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="visual",
        golden_dir=None,
        update_golden=False,
        require_wps=True,
        auto_repair=False,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    write_json(work / "qa_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(work / "conformance_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(
        work / "visual_report.json",
        {
            "passed": False,
            "issues": [
                {
                    "code": "WPS_PDFINFO_FAILED",
                    "severity": "error",
                    "message": "WPS PDF metadata failed",
                    "detail": "xref table broken",
                }
            ],
            "counts": {},
            "next_action": "",
        },
    )

    json_path, md_path = write_agent_summary(str(work), "2026-06-01_wps_pdfinfo", "最终论文.docx", "developer")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    action_text = "\n".join(summary.get("next_actions") or [])

    assert_true("WPS_PDFINFO_FAILED" in action_text, f"summary lost the WPS PDF metadata issue code: {summary}")
    assert_true("WPS" in action_text and "重跑 visual QA" in action_text, f"summary should tell users to rerun visual QA after WPS PDF metadata repair: {summary}")
    assert_true("WPS_PDFINFO_FAILED" in text and "重跑 visual QA" in text, "agent summary markdown should show the WPS PDF metadata rerun step")


@case
def pipeline_agent_summary_markdown_labels_warning_reports_explicitly() -> None:
    work = new_workdir("pipeline_agent_summary_warning_report_labels")
    write_workflow_mode(
        str(work),
        mode="developer",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="visual",
        golden_dir="TestData/GoldenBaselines",
        update_golden=False,
        require_wps=False,
        auto_repair=False,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    write_json(
        work / "qa_report.json",
        {
            "passed": True,
            "issues": [{"code": "REFERENCES_MISSING", "severity": "warning", "message": "references missing"}],
            "counts": {},
            "next_action": "结构 QA 没有阻断错误，但有警告 `REFERENCES_MISSING` 需要人工确认。",
        },
    )
    write_json(
        work / "conformance_report.json",
        {
            "passed": True,
            "issues": [{"code": "STYLE_MISMATCH", "severity": "warning", "message": "style needs review"}],
            "counts": {},
            "next_action": "strict 合规 QA 没有阻断错误，但有警告 `STYLE_MISMATCH` 需要人工确认。",
        },
    )
    write_json(
        work / "visual_report.json",
        {
            "passed": True,
            "issues": [{"code": "GOLDEN_BASELINE_MISSING", "severity": "warning", "message": "no baseline"}],
            "counts": {},
            "next_action": "visual QA 通过但缺少黄金基线，需要人工确认。",
        },
    )

    json_path, md_path = write_agent_summary(str(work), "2026-05-31_warning_labels", "最终论文.docx", "developer")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")

    assert_true(summary["reports"]["structural"]["result_label"] == "通过但有警告", f"structural JSON label should flag warnings: {summary}")
    assert_true(summary["reports"]["conformance"]["result_label"] == "通过但有警告", f"strict JSON label should flag warnings: {summary}")
    assert_true(summary["reports"]["visual"]["result_label"] == "通过但有警告", f"visual JSON label should flag warnings: {summary}")
    assert_true(summary["reports"]["structural"]["status"] == "passed_with_warnings", f"structural JSON status should flag warnings: {summary}")
    assert_true(summary["reports"]["conformance"]["status"] == "passed_with_warnings", f"strict JSON status should flag warnings: {summary}")
    assert_true(summary["reports"]["visual"]["status"] == "passed_with_warnings", f"visual JSON status should flag warnings: {summary}")
    assert_true("结构 QA：通过但有警告" in text, f"structural summary line should not say plain pass: {text}")
    assert_true("DOCX/XML 合规 QA：通过但有警告" in text, f"strict summary line should not say plain pass: {text}")
    assert_true("视觉 QA：通过但有警告" in text, f"visual summary line should not say plain pass: {text}")


@case
def pipeline_agent_summary_localizes_manual_checks_and_skips_empty_warning_prompt() -> None:
    work = new_workdir("pipeline_agent_summary_manual_checks")
    write_workflow_mode(
        str(work),
        mode="user",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="strict",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
        auto_repair=True,
        agent_auto=True,
    )
    (work / "最终论文.docx").write_bytes(b"synthetic")
    write_json(work / "qa_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(work / "conformance_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(
        work / "repair_loop_report.json",
        {
            "status": "converged",
            "rounds_run": 0,
            "final_errors": 0,
            "final_warnings": 0,
            "manual_check_required": [
                "用 Word/WPS 打开最终 DOCX，核对分页、图片、公式、表格和目录。",
                "查看 qa_report.md 和 repair_loop_report.md 中的剩余 warning，确认不会影响交付。",
            ],
            "remaining_manual_note": "当前启用的 QA 未报告剩余 warning。",
        },
    )
    json_path, md_path = write_agent_summary(str(work), "2026-05-29_manual", "最终论文.docx", "user")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    assert_true(summary["manual_check_required"] == ["用 Word/WPS 打开最终 DOCX，核对分页、图片、公式、表格和目录。"], f"manual checks were noisy: {summary['manual_check_required']}")
    assert_true("Review remaining warnings" not in text and "剩余 warning" not in text, "agent summary should not expose warning boilerplate when no warnings remain")


@case
def pipeline_agent_summary_surfaces_repair_loop_next_action() -> None:
    work = new_workdir("pipeline_agent_summary_repair_loop_next")
    write_workflow_mode(
        str(work),
        mode="user",
        template_path="template.docx",
        content_path="content.docx",
        run_qa=True,
        qa_level="basic",
        golden_dir=None,
        update_golden=False,
        require_wps=False,
        auto_repair=True,
        repair_max_rounds=2,
        repair_stop_no_improve=1,
    )
    write_json(work / "qa_report.json", {"passed": True, "issues": [], "counts": {}, "next_action": "ok"})
    write_json(
        work / "repair_loop_report.json",
        {
            "status": "stopped_build_failed",
            "ok": False,
            "rounds_run": 1,
            "final_errors": 1,
            "final_warnings": 0,
            "next_action": "自动修复后重建失败；先打开 repair_loop_report.md，再检查 build_generated.py。",
            "resume_scope": "current_docx",
            "resume_command": "python Outputs/demo/build_generated.py",
            "manual_check_required": ["用 Word/WPS 打开最终 DOCX，核对分页、图片、公式、表格和目录。"],
        },
    )
    json_path, md_path = write_agent_summary(str(work), "2026-05-31_repair_loop_next", "最终论文.docx", "user", pipeline_status="failed")
    summary = json.loads(Path(json_path).read_text(encoding="utf-8"))
    text = Path(md_path).read_text(encoding="utf-8")
    action_text = "\n".join(summary.get("next_actions") or summary.get("manual_check_required") or [])
    assert_true("repair_loop_report.md" in action_text and "build_generated.py" in action_text, f"agent summary lost repair-loop next action: {summary}")
    assert_true("repair_loop_report.md" in text and "build_generated.py" in text, "agent summary markdown should surface repair-loop next action")


@case
def pipeline_verification_double_verify_arbitrates_format_mismatch() -> None:
    calls: List[int] = []

    def flaky_format_extractor(path):
        calls.append(len(calls) + 1)
        paragraph_count = 1 if len(calls) != 2 else 2
        fmt = {
            "_meta": {"paragraphs": paragraph_count},
            "paragraphs": [{"runs": []} for _ in range(paragraph_count)],
            "tables": [],
            "sections": [],
        }
        return fmt, "# Format"

    result, report = double_verify(flaky_format_extractor, "synthetic.docx", "Format")
    assert_true(len(calls) == 3, "format mismatch should trigger a third arbitration run")
    assert_true(len(result["paragraphs"]) == 1, "third run should arbitrate back to the majority shape")
    assert_true(report == "# Format", "format markdown payload should be preserved")


@case
def pipeline_verification_arbitrates_content_mismatch() -> None:
    work = new_workdir("pipeline_verify_content_majority")
    calls: List[int] = []

    def flaky_content_extractor(path, output_dir=None):
        calls.append(len(calls) + 1)
        references = [{"text": "Synthetic reference"}] if len(calls) == 2 else []
        return {
            "_meta": {"images_extracted": 0, "image_extract_failures": [], "non_body_images": []},
            "sections": [{"heading": "Body", "paragraphs": ["Text"]}],
            "references": references,
        }

    result = double_verify(flaky_content_extractor, "synthetic.docx", "Content", output_dir=str(work))
    assert_true(len(calls) == 3, "content mismatch should trigger a third arbitration run")
    assert_true(result["references"] == [], "third run should arbitrate content back to the majority shape")


@case
def pipeline_verification_converges_incremental_content_and_materializes_images() -> None:
    work = new_workdir("pipeline_verify_content_converge")
    calls: List[int] = []

    def incremental_content_extractor(path, output_dir=None):
        calls.append(len(calls) + 1)
        run_no = len(calls)
        base = Path(path).stem
        fig_dir = Path(output_dir) / base / "figures"
        fig_dir.mkdir(parents=True, exist_ok=True)

        image_names = [f"img{i}.png" for i in range(1, run_no + 1)]
        for image_name in image_names:
            (fig_dir / image_name).write_bytes(PNG_1X1)

        references = [{"text": f"[{i}] Ref {i}"} for i in range(1, run_no + 1)]
        return {
            "_meta": {
                "images_extracted": len(image_names),
                "images_dir": str(fig_dir),
                "image_extract_failures": [{"target": "transient", "error": "miss"}] if run_no == 1 else [],
                "non_body_images": [],
            },
            "sections": [
                {
                    "heading": "Body",
                    "level": 1,
                    "paragraphs": [{"role": "image", "image": image_names[-1]}],
                    "images": [image_names[-1]],
                }
            ],
            "references": references,
        }

    result = double_verify(incremental_content_extractor, "synthetic.docx", "Content", output_dir=str(work))
    assert_true(len(calls) == 5, "content convergence should keep collecting unstable recoverable content up to the cap")
    assert_true(result["_meta"].get("converged_extraction"), "convergence metadata was not recorded")
    assert_true(len(result["references"]) == 5, f"references were not merged: {result['references']}")
    assert_true(result["sections"][0]["images"] == ["img1.png", "img2.png", "img3.png", "img4.png", "img5.png"], "section images were not unioned")

    paragraph_images = sorted(
        item.get("image")
        for item in result["sections"][0]["paragraphs"]
        if isinstance(item, dict) and item.get("image")
    )
    assert_true(paragraph_images == ["img1.png", "img2.png", "img3.png", "img4.png", "img5.png"], f"paragraph images were not merged: {paragraph_images}")
    final_fig_dir = work / "synthetic" / "figures"
    assert_true((final_fig_dir / "img1.png").exists() and (final_fig_dir / "img5.png").exists(), "converged images were not materialized")
    assert_true(not (work / "_extract_verify_runs").exists(), "isolated verification directories were not cleaned")
    assert_true(result["_meta"].get("image_extract_failures") == [], "transient image failures should not survive convergence")


@case
def pipeline_verification_inserts_recovered_items_near_anchors() -> None:
    def content_with(paragraphs, images=None):
        return {
            "_meta": {
                "images_extracted": len(images or []),
                "image_extract_failures": [],
                "non_body_images": [],
            },
            "sections": [
                {
                    "heading": "Body",
                    "level": 1,
                    "role": "body",
                    "paragraphs": paragraphs,
                    "images": images or [],
                }
            ],
            "references": [],
        }

    merged, _reason = _merge_content_results(
        [
            content_with(["before", {"role": "image", "image": "img.png"}, "middle", "after"], ["img.png"]),
            content_with(["before", {"role": "formula", "source": "latex", "latex": "E=mc^2", "text": "E=mc^2"}, "middle", "after"]),
            content_with(["before", {"role": "table", "table_rows": [["A"], ["1"]]}, "middle", "after"]),
        ]
    )
    paragraphs = merged["sections"][0]["paragraphs"]
    middle_index = paragraphs.index("middle")
    formula_index = next(i for i, item in enumerate(paragraphs) if isinstance(item, dict) and item.get("role") == "formula")
    table_index = next(i for i, item in enumerate(paragraphs) if isinstance(item, dict) and item.get("table_rows"))
    assert_true(formula_index < middle_index, f"recovered formula drifted to section tail: {paragraphs}")
    assert_true(table_index < middle_index, f"recovered table drifted to section tail: {paragraphs}")


@case
def pipeline_verification_fails_when_no_majority_signature() -> None:
    calls: List[int] = []

    def unstable_content_extractor(path):
        calls.append(len(calls) + 1)
        return {
            "_meta": {"images_extracted": len(calls), "image_extract_failures": [], "non_body_images": []},
            "sections": [{"heading": f"Body {idx}", "paragraphs": ["Text"]} for idx in range(len(calls))],
            "references": [],
        }

    try:
        double_verify(unstable_content_extractor, "synthetic.docx", "Content")
    except VerificationError as exc:
        assert_true("verification failed" in str(exc), "failure should explain that verification failed")
        assert_true(len(calls) == 3, "unresolved mismatch should stop after the third run")
        return
    fail("unresolved extraction mismatch should raise VerificationError")


@case
def run_pipeline_writes_repair_plan_for_verification_failure() -> None:
    if str(REPO_ROOT) not in sys.path:
        sys.path.insert(0, str(REPO_ROOT))
    import run_pipeline as runner

    work = new_workdir("run_pipeline_verification_failure")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    outputs_dir = work / "Outputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    outputs_dir.mkdir()
    (template_dir / "template.docx").write_text("template", encoding="utf-8")
    (inputs_dir / "content.docx").write_text("content", encoding="utf-8")

    original = {
        "TEMPLATE_DIR": runner.TEMPLATE_DIR,
        "INPUTS_DIR": runner.INPUTS_DIR,
        "OUTPUTS_DIR": runner.OUTPUTS_DIR,
        "double_verify": runner.double_verify,
    }

    def failing_verify(*args, **kwargs):
        raise VerificationError("synthetic verification mismatch")

    try:
        runner.TEMPLATE_DIR = str(template_dir)
        runner.INPUTS_DIR = str(inputs_dir)
        runner.OUTPUTS_DIR = str(outputs_dir)
        runner.double_verify = failing_verify
        result = runner.run("template.docx", "content.docx", mode="user", run_qa=True)
    finally:
        runner.TEMPLATE_DIR = original["TEMPLATE_DIR"]
        runner.INPUTS_DIR = original["INPUTS_DIR"]
        runner.OUTPUTS_DIR = original["OUTPUTS_DIR"]
        runner.double_verify = original["double_verify"]

    assert_true(result is None, "pipeline should stop cleanly on verification failure")
    out_dirs = sorted(outputs_dir.iterdir())
    assert_true(out_dirs, "pipeline did not create an output directory")
    report = json.loads((out_dirs[-1] / "qa_report.json").read_text(encoding="utf-8"))
    assert_true(report["issues"][0]["code"] == "EXTRACTION_VERIFICATION_FAILED", f"wrong verification failure issue: {report}")
    assert_true((out_dirs[-1] / "qa_repair_plan.md").exists(), "verification failure did not write repair plan")
    plan = json.loads((out_dirs[-1] / "qa_repair_plan.json").read_text(encoding="utf-8"))
    assert_true("build_manifest.json" not in plan["open_first"], f"verification failure plan should not point at later-stage artifacts: {plan}")
    assert_true(not plan["commands"].get("rebuild_current_docx"), f"verification failure should not suggest rebuilding a missing generated script: {plan}")


@case
def run_pipeline_writes_repair_plan_for_build_failure() -> None:
    if str(REPO_ROOT) not in sys.path:
        sys.path.insert(0, str(REPO_ROOT))
    import run_pipeline as runner

    work = new_workdir("run_pipeline_build_failure")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    outputs_dir = work / "Outputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    outputs_dir.mkdir()
    (template_dir / "template.docx").write_text("template", encoding="utf-8")
    (inputs_dir / "content.docx").write_text("content", encoding="utf-8")

    original = {
        "TEMPLATE_DIR": runner.TEMPLATE_DIR,
        "INPUTS_DIR": runner.INPUTS_DIR,
        "OUTPUTS_DIR": runner.OUTPUTS_DIR,
        "double_verify": runner.double_verify,
        "generate_script": runner.generate_script,
        "run_generated_script": runner.run_generated_script,
    }

    def fake_verify(_extractor, _path, label, **_kwargs):
        if label == "Format":
            return base_format(), "# Format"
        return base_content(["Body text"])

    def fake_generate(_fmt_json_path, _cnt_json_path, out_dir, _output_docx_name):
        Path(out_dir, "build_generated.py").write_text("raise RuntimeError('synthetic build failure')\n", encoding="utf-8")
        return 45

    def failing_build(_gen_py_path, _out_dir, python_executable):
        return ScriptExecutionResult(
            returncode=1,
            stdout="",
            stderr=f"Traceback synthetic failure at {work}\\private\\source.docx",
        )

    try:
        runner.TEMPLATE_DIR = str(template_dir)
        runner.INPUTS_DIR = str(inputs_dir)
        runner.OUTPUTS_DIR = str(outputs_dir)
        runner.double_verify = fake_verify
        runner.generate_script = fake_generate
        runner.run_generated_script = failing_build
        result = runner.run("template.docx", "content.docx", mode="user", run_qa=True)
    finally:
        runner.TEMPLATE_DIR = original["TEMPLATE_DIR"]
        runner.INPUTS_DIR = original["INPUTS_DIR"]
        runner.OUTPUTS_DIR = original["OUTPUTS_DIR"]
        runner.double_verify = original["double_verify"]
        runner.generate_script = original["generate_script"]
        runner.run_generated_script = original["run_generated_script"]

    assert_true(result is None, "pipeline should stop cleanly on generated-script build failure")
    out_dirs = sorted(outputs_dir.iterdir())
    assert_true(out_dirs, "pipeline did not create an output directory for the build failure")
    out_dir = out_dirs[-1]
    report = json.loads((out_dir / "qa_report.json").read_text(encoding="utf-8"))
    assert_true(report["issues"][0]["code"] == "MISSING_DOCX", f"build failure should be routed as missing final DOCX: {report}")
    assert_true(str(work) not in report["issues"][0].get("detail", ""), f"build failure detail leaked an absolute private path: {report}")
    plan = json.loads((out_dir / "qa_repair_plan.json").read_text(encoding="utf-8"))
    assert_true(plan.get("resume_scope") == "current_docx", f"build failure should route to current DOCX rebuild: {plan}")
    assert_true("build_generated.py" in plan.get("resume_command", ""), f"build failure plan should provide rebuild command: {plan}")
    assert_true("build_generated.py" in plan.get("open_first", []), f"build failure plan should ask users to open build_generated.py first: {plan}")
    assert_true("qa_repair_plan" in plan.get("next_action", "") or "qa_report" in plan.get("next_action", ""), f"build failure next action should point users to the standard report handoff: {plan}")
    summary = json.loads((out_dir / "agent_summary.json").read_text(encoding="utf-8"))
    action_text = "\n".join(summary.get("next_actions") or summary.get("manual_check_required") or [])
    assert_true("MISSING_DOCX" in action_text and "build_generated.py" in action_text and "qa_repair_plan" in action_text, f"agent summary lost the build-failure next step: {summary}")


@case
def run_pipeline_agent_auto_writes_preflight_report_for_missing_file() -> None:
    if str(REPO_ROOT) not in sys.path:
        sys.path.insert(0, str(REPO_ROOT))
    import run_pipeline as runner

    work = new_workdir("run_pipeline_agent_preflight_missing_file")
    template_dir = work / "Templates"
    inputs_dir = work / "Inputs"
    outputs_dir = work / "Outputs"
    template_dir.mkdir()
    inputs_dir.mkdir()
    outputs_dir.mkdir()

    original = {
        "TEMPLATE_DIR": runner.TEMPLATE_DIR,
        "INPUTS_DIR": runner.INPUTS_DIR,
        "OUTPUTS_DIR": runner.OUTPUTS_DIR,
    }
    try:
        runner.TEMPLATE_DIR = str(template_dir)
        runner.INPUTS_DIR = str(inputs_dir)
        runner.OUTPUTS_DIR = str(outputs_dir)
        result = runner.run(None, None, md_file="missing.md", mode="user", agent_auto=True)
    finally:
        runner.TEMPLATE_DIR = original["TEMPLATE_DIR"]
        runner.INPUTS_DIR = original["INPUTS_DIR"]
        runner.OUTPUTS_DIR = original["OUTPUTS_DIR"]

    assert_true(result is None, "missing explicit agent input should stop before pipeline run")
    report_path = outputs_dir / "_agent_preflight_latest" / "agent_preflight_report.json"
    assert_true(report_path.exists(), "missing explicit agent input should write a preflight report")
    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert_true(report["status"] == "blocked_input_resolution", f"wrong preflight status: {report}")
    assert_true(report.get("next_steps"), f"preflight report should include next steps: {report}")


@case
def run_pipeline_completion_step_is_named() -> None:
    root = PIPELINE_DIR.parents[2]
    text = (root / "run_pipeline.py").read_text(encoding="utf-8")
    assert_true("step('完成')" in text or 'step("完成")' in text, "completion step should have a user-visible label")
    assert_true("step('??')" not in text and 'step("??")' not in text, "placeholder completion step leaked into run_pipeline.py")


@case
def pipeline_template_phase_writes_profile_and_requirements() -> None:
    work = new_workdir("pipeline_template_phase")
    calls: List[str] = []

    def fake_profile(fmt, out_dir, project_root):
        calls.append(f"profile:{project_root}")
        return {
            "capabilities": {
                "has_cover": True,
                "has_heading_styles": False,
                "has_caption_styles": True,
            },
            "risk_flags": {"mixed_headers": True, "low_sample": False},
        }

    profile = write_template_profile_phase(
        base_format(),
        str(work),
        project_root="ROOT",
        write_template_profile=fake_profile,
    )
    assert_true(profile["capabilities"]["has_cover"] is True, "profile return value was not preserved")

    def fake_requirements(fmt, content, out_dir):
        calls.append("requirements")
        Path(out_dir, "template_requirements.json").write_text("{}", encoding="utf-8")

    ok = write_template_requirements_phase(
        base_format(),
        base_content(["Body"]),
        str(work),
        write_template_requirements=fake_requirements,
        optional_import_detail=lambda name: " detail",
    )
    assert_true(ok is True, "requirements phase should report success")
    assert_true(calls == ["profile:ROOT", "requirements"], f"unexpected template phase calls: {calls}")
    assert_true((work / "template_requirements.json").exists(), "requirements writer was not called")

    skipped = write_template_requirements_phase(
        base_format(),
        base_content(["Body"]),
        str(work),
        write_template_requirements=None,
        optional_import_detail=lambda name: " detail",
    )
    assert_true(skipped is False, "missing requirements writer should report skipped")


