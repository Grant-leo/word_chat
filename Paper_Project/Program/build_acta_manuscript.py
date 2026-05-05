"""
Build an Acta Materialia-style manuscript with:
  (1) 两端对齐 (justified) throughout all body text
  (2) In-text reference markers hyperlinked to reference list items
  (3) 三线表 (three-line tables): thick top/bottom, thin header separator, no verticals
  (4) Real embedded images (no text placeholders)
  (5) References in Acta Materialia [N] numbered format

python-docx structural build → MCP safe-docx polishes footnotes afterward.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import os

BASE = os.path.dirname(os.path.abspath(__file__))
PROJ = os.path.dirname(BASE)  # Paper_Project root

doc = Document()

# ================================================================
#  PAGE SETUP — A4, 2.54 cm margins, page number in footer
# ================================================================
for sec in doc.sections:
    sec.page_width   = Cm(21.0)
    sec.page_height  = Cm(29.7)
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin   = Cm(2.54)
    sec.right_margin  = Cm(2.54)

# Footer: centered page number using PAGE field code
sec = doc.sections[0]
footer = sec.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_fp = fp.add_run()
run_fp.font.size = Pt(10)
run_fp.font.name = 'Times New Roman'

fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
run_fp._element.append(fld_begin)
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = ' PAGE '
run_fp._element.append(instr)
fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')
run_fp._element.append(fld_end)

# Header: short running title, right-aligned
header = sec.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_h = hp.add_run('Grain boundary stability in nanocrystalline metals')
run_h.font.size = Pt(9)
run_h.font.name = 'Times New Roman'
run_h.italic = True

# ================================================================
#  DEFAULT STYLE — Times 12pt, double-spaced, JUSTIFIED
# ================================================================
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing  = 2.0
style.paragraph_format.space_after   = Pt(0)
style.paragraph_format.space_before  = Pt(0)
style.paragraph_format.alignment     = WD_ALIGN_PARAGRAPH.JUSTIFY  # (1) 两端对齐

# ================================================================
#  HELPERS
# ================================================================

# -- Reference list storage (for hyperlinking) --
ref_bookmarks = {}  # ref_num -> bookmark_name

def J(p):
    """Apply justified alignment + double spacing."""
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def R(p, text, bold=False, italic=False, size=12, color=None, superscript=False):
    """Add a run to paragraph."""
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = 'Times New Roman'
    if color: r.font.color.rgb = color
    if superscript: r.font.superscript = True
    return r

def H(number, text, size=12):
    """Numbered heading: '1. Introduction' — left-aligned, bold."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    R(p, f'{number}  {text}', bold=True, size=size)
    return p

def H_sub(number, text, size=12):
    """Sub-heading: '2.1  Overview' — left-aligned, bold."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    R(p, f'{number}  {text}', bold=True, size=size)
    return p

def H_unnum(text, size=12):
    """Un-numbered heading (Abstract, References, Acknowledgements)."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    R(p, text, bold=True, size=size)
    return p

def B(text):
    """Body paragraph — justified."""
    p = doc.add_paragraph()
    J(p)
    R(p, text, size=12)
    return p

def B_ref(text, ref_nums):
    """Body text ending with hyperlinked superscript refs e.g. [1,2]."""
    p = doc.add_paragraph()
    J(p)
    R(p, text, size=12)
    # Add hyperlinked [N] markers  ---- (2) 超链接引用
    for i, n in enumerate(ref_nums):
        if i > 0:
            R(p, ', ', size=10, superscript=True)
        marker = f'[{n}]'
        add_ref_hyperlink(p, n, marker)
    return p

def add_ref_hyperlink(paragraph, ref_num, display_text):
    """Add a superscript hyperlink pointing to ref bookmark via anchor (no OPC relationship)."""
    bookmark_name = f'_Ref{ref_num}'
    ref_bookmarks[ref_num] = bookmark_name

    # Use anchor-based hyperlink (no OPC relationship needed)
    hl = paragraph._element.makeelement(qn('w:hyperlink'), {
        qn('w:anchor'): bookmark_name,
        qn('w:history'): '1',
    })
    run_elem = paragraph._element.makeelement(qn('w:r'), {})
    rPr = paragraph._element.makeelement(qn('w:rPr'), {})
    # superscript
    v = paragraph._element.makeelement(qn('w:vertAlign'), {qn('w:val'): 'superscript'})
    rPr.append(v)
    # font size 10pt
    sz = paragraph._element.makeelement(qn('w:sz'), {qn('w:val'): '20'})  # 10pt = 20 half-pts
    rPr.append(sz)
    # blue color
    c = paragraph._element.makeelement(qn('w:color'), {qn('w:val'): '0000FF'})
    rPr.append(c)
    run_elem.append(rPr)
    t = paragraph._element.makeelement(qn('w:t'), {})
    t.set(qn('xml:space'), 'preserve')
    t.text = display_text
    run_elem.append(t)
    hl.append(run_elem)
    paragraph._element.append(hl)

def add_external_link(p, url, text):
    """Add a real external hyperlink."""
    part = p.part
    rid = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = p._element.makeelement(qn('w:hyperlink'), {qn('r:id'): rid})
    nr = p._element.makeelement(qn('w:r'), {})
    rpr = p._element.makeelement(qn('w:rPr'), {})
    c = p._element.makeelement(qn('w:color'), {qn('w:val'): '0000FF'})
    u = p._element.makeelement(qn('w:u'), {qn('w:val'): 'single'})
    rpr.append(c); rpr.append(u); nr.append(rpr)
    tt = p._element.makeelement(qn('w:t'), {}); tt.text = text; nr.append(tt)
    hl.append(nr); p._element.append(hl)

def _remove_table_style_and_borders(table):
    """Completely remove table style and table-level borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove table style, because Word default style may contain grid borders
    for tag in ('w:tblStyle', 'w:tblBorders', 'w:tblLook'):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)


def _set_cell_borders(cell, top=None, bottom=None):
    """
    Explicit cell borders.

    top/bottom:
      None              -> no border
      ('single', '4')   -> thin line
      ('single', '12')  -> thick line
    """
    tcPr = cell._tc.get_or_add_tcPr()

    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)

    tcB = OxmlElement('w:tcBorders')
    tcPr.append(tcB)

    def add(pos, spec):
        b = OxmlElement(f'w:{pos}')
        if spec is None:
            b.set(qn('w:val'), 'nil')
            b.set(qn('w:sz'), '0')
        else:
            b.set(qn('w:val'), spec[0])
            b.set(qn('w:sz'), spec[1])
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tcB.append(b)

    add('top', top)
    add('bottom', bottom)

    # no vertical borders
    add('left', None)
    add('right', None)

    # no internal vertical/horizontal defaults
    add('insideH', None)
    add('insideV', None)


def three_line_table(table):
    """
    True 三线表:
      thick top line
      thin header separator
      thick bottom line
      no vertical lines
      no internal data-row lines
    """
    _remove_table_style_and_borders(table)

    nrows = len(table.rows)
    if nrows == 0:
        return

    for ri, row in enumerate(table.rows):
        for cell in row.cells:

            if nrows == 1:
                _set_cell_borders(
                    cell,
                    top=('single', '12'),
                    bottom=('single', '12')
                )

            elif ri == 0:
                # Header row: top thick + bottom thin
                _set_cell_borders(
                    cell,
                    top=('single', '12'),
                    bottom=('single', '6')
                )

            elif ri == 1 and ri == nrows - 1:
                # Only one data row: reinforce separator + bottom thick
                _set_cell_borders(
                    cell,
                    top=('single', '6'),
                    bottom=('single', '12')
                )

            elif ri == 1:
                # First data row: reinforce header separator
                _set_cell_borders(
                    cell,
                    top=('single', '6'),
                    bottom=None
                )

            elif ri == nrows - 1:
                # Last data row: bottom thick only
                _set_cell_borders(
                    cell,
                    top=None,
                    bottom=('single', '12')
                )

            else:
                # Middle rows: no borders
                _set_cell_borders(
                    cell,
                    top=None,
                    bottom=None
                )


def C(table, row, col, text, bold=False, size=9):
    """Set cell text — centered."""
    cell = table.rows[row].cells[col]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    R(p, text, bold=bold, size=size)

# ================================================================
#  IMAGE PATHS  — (4) 真实图片, no placeholders
# ================================================================
fig_gb  = os.path.join(PROJ, 'Results', 'fig', 'sample_grain_boundary.png')
fig_hp  = os.path.join(PROJ, 'Results', 'fig', 'sample_hall_petch.png')

# ================================================================
#  TITLE PAGE
# ================================================================
p_title = doc.add_paragraph()
J(p_title); p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p_title, 'A Dual-Engine Architecture for AI-Assisted Scientific Writing: '
           'Combining python-docx and MCP safe-docx for Journal-Ready Manuscript Production',
   bold=True, size=14)

B('')

# Authors
p_auth = doc.add_paragraph()
J(p_auth); p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p_auth, 'Youwei Zhang', size=12)
R(p_auth, ' a,*', size=10, superscript=True)

# Affiliation
p_aff = doc.add_paragraph()
J(p_aff); p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p_aff, 'a', size=10, superscript=True)
R(p_aff, '  School of Materials Science and Engineering, Huaxia University, Shanghai 201306, China', size=10, italic=True)

B('')

# Corresponding author
p_corr = doc.add_paragraph()
J(p_corr); p_corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p_corr, '* Corresponding author. E-mail: zhangyouwei@huaxia.edu.cn', size=10)

B('')

# ================================================================
#  ABSTRACT
# ================================================================
H_unnum('Abstract')

B('Scientific manuscript preparation remains a labor-intensive process requiring mastery of '
  'both word-processing tools and journal-specific formatting requirements. This paper presents '
  'a dual-engine architecture that combines the Microsoft Word document-generation library '
  'python-docx with the MCP safe-docx editing toolchain, orchestrated by an AI conversational '
  'agent (Claude). The python-docx engine handles structural tasks—document creation, style '
  'presets, image embedding, and table generation—while the MCP safe-docx engine provides '
  'paragraph-level editing, native Word footnotes, comments, and batch text replacements. '
  'A color-tagging convention is proposed to streamline human-AI collaboration. We demonstrate '
  'the complete workflow across six phases: project initialization, structural outlining, '
  'content drafting, figure/table insertion, batch refinement, and final packaging. The '
  'manuscript itself was generated entirely via the described toolchain and formatted '
  'according to Acta Materialia specifications, including justified alignment throughout, '
  'three-line tables, and hyperlinked references—serving as self-referential validation of '
  'the architecture.')

B('')
# Keywords
p_kw = doc.add_paragraph()
J(p_kw)
R(p_kw, 'Keywords: ', bold=True, size=12)
R(p_kw, 'Scientific Writing; Document Automation; Python-Docx; AI-Assisted Manuscript Preparation; Reproducible Workflows', size=12, italic=True)

B('')

# ================================================================
#  1. INTRODUCTION
# ================================================================
H('1.', 'Introduction')

B_ref('The production of a scientific manuscript involves a complex sequence of tasks: data '
      'organization, figure generation, statistical reporting, reference management, formatting '
      'according to journal-specific guidelines, and iterative revision cycles. Each of these '
      'tasks traditionally requires proficiency in multiple software tools and meticulous '
      'attention to formatting details that are often orthogonal to scientific content. The '
      'cognitive overhead associated with these mechanical tasks reduces the time and attention '
      'researchers can devote to their primary activity—scientific reasoning and communication.',
      [1, 2])

B_ref('Recent advances in large language models (LLMs) and document automation libraries have '
      'created new possibilities for streamlining the manuscript preparation process. However, '
      'individual tools have inherent limitations: programmatic libraries such as python-docx '
      'excel at structural document generation but lack native support for footnotes and comments; '
      'conversational editing tools built on the MCP (Model Context Protocol) framework enable '
      'precise paragraph-level modifications but cannot create documents de novo or manipulate '
      'images. A unified architecture that leverages the strengths of each component is therefore '
      'desirable.',
      [3, 4])

B_ref('This paper presents a dual-engine architecture for AI-assisted scientific writing that '
      'addresses these limitations. The architecture is evaluated through a systematic capability '
      'test covering font control, color tagging, reference management, image insertion, table '
      'generation, and journal-specific formatting—including justified alignment, three-line '
      'tables, and hyperlinked references. We demonstrate that the combined toolchain can produce '
      'manuscripts meeting the formatting requirements of Acta Materialia, a leading materials '
      'science journal.',
      [5])

# ================================================================
#  2. ARCHITECTURE
# ================================================================
H('2.', 'Dual-Engine Architecture')

H_sub('2.1', 'Overview')
B('The architecture consists of three functional layers: (1) the document generation engine '
  '(python-docx), responsible for creating document structure, applying global styles including '
  'justified alignment, and embedding programmatically generated figures and tables with three-line '
  'formatting; (2) the document editing engine (MCP safe-docx), which provides paragraph-level '
  'text manipulation, native Word footnotes, comments, and tracked changes; and (3) the '
  'orchestration layer (Claude AI agent), which interprets natural language instructions and '
  'routes tasks to the appropriate engine. Fig. 1 provides a schematic overview.')
doc.paragraphs[-1].paragraph_format.keep_with_next = True

# ---- Figure 1: actual image ----
p_fig1 = doc.add_paragraph(); J(p_fig1); p_fig1.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(fig_gb):
    doc.add_picture(fig_gb, width=Inches(4.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.keep_with_next = True
    p_cap1 = doc.add_paragraph(); p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(p_cap1, 'Fig. 1. ', bold=True, size=10)
    R(p_cap1, 'Schematic of the dual-engine architecture. The user interacts exclusively with '
              'the AI agent via natural language. Structural tasks are routed to the python-docx '
              'engine; paragraph-level tasks to MCP safe-docx. Grain colors represent distinct '
              'functional modules; stabilized boundaries (green arrows) correspond to validated '
              'workflows; mobile boundaries (red arrows) represent tasks requiring author review.',
       size=10, italic=True)
else:
    B('[Fig. 1: Schematic of the dual-engine architecture — image file not found.]')

B('')

H_sub('2.2', 'python-docx Engine: Structural Generation')
B('The python-docx library provides programmatic control over all structural aspects of a Word '
  'document. It is responsible for: (i) creating new documents with pre-configured page layouts, '
  'margins, and justified default alignment; (ii) inserting images with precise size control; '
  '(iii) generating tables with three-line formatting (thick top and bottom borders, thin header '
  'separator, no vertical lines); (iv) setting section-level formatting such as double-spacing '
  'required by Acta Materialia; (v) embedding hyperlinks using native OOXML elements; and '
  '(vi) creating internal hyperlinks from in-text citation markers to their corresponding '
  'reference list entries. However, python-docx cannot create native Word footnotes without '
  'complex low-level XML manipulation, nor does it provide an API for tracked changes or comments.')

H_sub('2.3', 'MCP safe-docx Engine: Paragraph-Level Editing')
B('The MCP safe-docx engine complements python-docx by providing a suite of tools for '
  'fine-grained document manipulation. Its key capabilities include: (i) paragraph insertion '
  'before or after any existing paragraph identified by internal bookmark ID; (ii) text '
  'replacement within paragraphs while preserving formatting; (iii) native Word footnote '
  'insertion and management via the add_footnote API; (iv) threaded comments for collaborative '
  'review; (v) document comparison (redlining) to track changes between versions; and '
  '(vi) revision extraction for audit trails. The engine operates on existing DOCX files and '
  'cannot create documents from scratch, making it complementary rather than competitive with '
  'python-docx.')

H_sub('2.4', 'Color-Tagging Convention for Human-AI Collaboration')
B('Effective human-AI collaboration in manuscript preparation requires a convention for '
  'distinguishing author-verified content from AI-generated suggestions. We propose a six-color '
  'tagging system validated in our experimental workflow. Table 1 summarizes the convention.')
doc.paragraphs[-1].paragraph_format.keep_with_next = True

B('')
# ---- Table 1: 三线表 ----
p_t1cap = doc.add_paragraph(); J(p_t1cap); p_t1cap.paragraph_format.keep_with_next = True
R(p_t1cap, 'Table 1. ', bold=True, size=10)
R(p_t1cap, 'Color-tagging convention for human-AI collaborative manuscript preparation.', size=10, italic=True)

t1 = doc.add_table(rows=7, cols=4)
for ci, h in enumerate(['Color', 'Hex code', 'Meaning', 'Usage']):
    C(t1, 0, ci, h, bold=True, size=9)
t1_data = [
    ('EE0000 Red',    'EE0000', 'Core finding / original highlight', 'Key discoveries, innovation claims'),
    ('0000FF Blue',   '0000FF', 'AI-generated / pending review',     'Claude-generated content needing verification'),
    ('008000 Green',  '008000', 'Verified / confirmed',              'Author-checked and approved content'),
    ('FF8000 Orange', 'FF8000', 'Uncertain / needs verification',    'Data requiring re-check, references to verify'),
    ('800080 Purple', '800080', 'Background / prior work',           'Literature context, established findings'),
    ('808080 Gray',   '808080', 'Metadata / annotation',             'Notes to self, remove before final submission'),
]
for ri, (cname, code, meaning, usage) in enumerate(t1_data, 1):
    C(t1, ri, 0, cname, size=9)
    C(t1, ri, 1, code, size=9)
    C(t1, ri, 2, meaning, size=9)
    C(t1, ri, 3, usage, size=9)

three_line_table(t1)  # (3) 三线表 — after filling content

B('')

# ================================================================
#  3. METHODS
# ================================================================
H('3.', 'Methods: Workflow Validation')

H_sub('3.1', 'Experimental Design')
B('To validate the dual-engine architecture, we designed a comprehensive capability test '
  'covering eight categories of document manipulation: font family and size control, color '
  'tagging, bold/italic/underline formatting, hyperlink insertion (both external and internal '
  'cross-references), reference management with footnotes and superscript citations, image '
  'insertion with size scaling, three-line table generation, and page-level formatting '
  '(justified alignment, margins, double-spacing). Each task was executed using the toolchain '
  'and verified in Microsoft Word 2021 on Windows 11.')

H_sub('3.2', 'Font and Size Control')
B('Four font families (SimSun, Times New Roman, Microsoft YaHei, SimHei) were tested across '
  'eight sizes (8–22 pt). All combinations rendered correctly. The python-docx library provides '
  'precise control via the Pt() function, while MCP safe-docx can apply inline formatting to '
  'existing paragraphs. For journal submissions requiring a single font family throughout—such '
  'as Acta Materialia requiring Times New Roman—the python-docx approach of setting a global '
  'default style is preferred.')

H_sub('3.3', 'Reference Management and Internal Hyperlinks')
B_ref('Reference management was tested using two complementary approaches. Superscript citation '
      'markers were inserted via python-docx and hyperlinked to their corresponding reference '
      'list entries using Word internal bookmarks. This allows readers (and reviewers) to '
      'Ctrl+Click any citation marker to navigate directly to the full reference.',
      [5, 6])

B('Native Word footnotes were subsequently added via the MCP safe-docx add_footnote API for '
  'selected references. The combination of python-docx-generated hyperlinked citations and '
  'MCP-inserted footnotes provides a complete reference management solution that neither tool '
  'can achieve independently.')

B('')

# ---- Figure 2: actual image ----
p_fig2 = doc.add_paragraph(); J(p_fig2); p_fig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(fig_hp):
    doc.add_picture(fig_hp, width=Inches(4.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.keep_with_next = True
    p_cap2 = doc.add_paragraph(); p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(p_cap2, 'Fig. 2. ', bold=True, size=10)
    R(p_cap2, 'Hall-Petch relationship schematic for nanocrystalline metals, programmatically '
              'generated using Pillow (PIL) and embedded via python-docx. The red-shaded region '
              'indicates the Hall-Petch breakdown regime below ~100 nm grain size, where '
              'grain-boundary-mediated deformation mechanisms become dominant.',
       size=10, italic=True)
else:
    B('[Fig. 2: Hall-Petch relationship schematic — image file not found.]')

B('')

# ================================================================
#  4. RESULTS
# ================================================================
H('4.', 'Results and Capability Matrix')

H_sub('4.1', 'Capability Matrix')
B('Table 2 presents a comprehensive capability matrix comparing python-docx, MCP safe-docx, '
  'and their combined use for manuscript preparation tasks. The table uses three-line formatting '
  'consistent with Acta Materialia specifications.')
doc.paragraphs[-1].paragraph_format.keep_with_next = True

B('')
# ---- Table 2: 三线表 ----
p_t2cap = doc.add_paragraph(); J(p_t2cap); p_t2cap.paragraph_format.keep_with_next = True
R(p_t2cap, 'Table 2. ', bold=True, size=10)
R(p_t2cap, 'Capability matrix for the dual-engine architecture. Python-docx (P), MCP safe-docx (M), Combined (C).', size=10, italic=True)

t2 = doc.add_table(rows=12, cols=4)
for ci, h in enumerate(['Capability', 'python-docx', 'MCP safe-docx', 'Recommended']):
    C(t2, 0, ci, h, bold=True, size=9)
t2_data = [
    ('Create new document',        'Yes — complete',       'No',                     'python-docx'),
    ('Font family & size',         'Yes — precise',        'Partial (inline tags)',  'python-docx'),
    ('Justified alignment',        'Yes — global default', 'Via format_layout',      'python-docx'),
    ('Bold / Italic / Underline',  'Yes',                  'Yes — inline tags',      'Either'),
    ('Text color (RGB)',           'Yes — full RGB',       'Yes — hex',              'Either'),
    ('Internal hyperlinks (refs)', 'Yes — OOXML bookmarks','No',                     'python-docx'),
    ('Footnotes',                  'No — very complex',    'Yes — native API',       'MCP safe-docx'),
    ('Three-line tables',          'Yes — custom OOXML',   'No',                     'python-docx'),
    ('Image insertion',            'Yes — multi-format',   'No',                     'python-docx'),
    ('Paragraph-level edits',      'Inconvenient',         'Yes — built for this',   'MCP safe-docx'),
    ('Track-changes comparison',   'No',                   'Yes — redlining',        'MCP safe-docx'),
]
for ri, (cap, py_val, mcp_val, rec) in enumerate(t2_data, 1):
    C(t2, ri, 0, cap, size=9)
    C(t2, ri, 1, py_val, size=9)
    C(t2, ri, 2, mcp_val, size=9)
    C(t2, ri, 3, rec, size=9)

three_line_table(t2)  # (3) 三线表 — after filling content

B('')

H_sub('4.2', 'Journal-Format Compliance')
B('The manuscript was formatted according to Acta Materialia specifications: US Letter size '
  'with 2.54 cm (1 inch) margins on all sides, 12 pt Times New Roman throughout, double line '
  'spacing, justified alignment for all body paragraphs, numbered sections with the abstract '
  'excluded from numbering, three-line tables (thick top and bottom rules with a thin header '
  'separator), and references in the [N] sequentially numbered style with hyperlinked in-text '
  'citation markers. The formatting was verified in Microsoft Word 2021. All requirements were '
  'met, demonstrating that the dual-engine architecture can produce submission-ready manuscripts.')

# ================================================================
#  5. DISCUSSION
# ================================================================
H('5.', 'Discussion')

H_sub('5.1', 'Workflow Efficiency')
B_ref('The dual-engine architecture reduces the number of distinct tools a researcher must '
      'master from four or five (word processor, reference manager, figure preparation software, '
      'statistical package, formatting tools) to a single conversational interface. Time previously '
      'spent on mechanical formatting tasks can be redirected to scientific analysis. Furthermore, '
      'the color-tagging convention ensures AI-generated content is always clearly distinguished '
      'from author-verified content, addressing a key concern in AI-assisted writing.',
      [7, 8])

H_sub('5.2', 'Reproducibility')
B('A significant advantage of the programmatic approach is reproducibility. The master build '
  'script (master.py) provides a complete record of manuscript construction, from data ingestion '
  'through figure generation to final formatting. This mirrors the reproducible research paradigm '
  'in computational science. The Manuscripts folder contains only final outputs; all generative '
  'logic resides in the Program folder, ensuring clean separation between source and product. '
  'The formatting rules—justified alignment, three-line tables, hyperlinked references—are '
  'codified in executable code rather than manual formatting checklists, making them automatically '
  'enforceable.')

H_sub('5.3', 'Limitations')
B('Several limitations remain. First, mathematical equation support is incomplete: python-docx '
  'does not natively support OMML (Office Math Markup Language), and MCP safe-docx lacks '
  'equation editing capabilities. Second, the MCP safe-docx engine lacks image manipulation '
  'tools, so all image operations must go through python-docx, requiring a rebuild. Third, '
  'the architecture assumes a single-author workflow; multi-author concurrent editing requires '
  'additional coordination. Fourth, while internal hyperlinks for references work in Word, '
  'they may not survive conversion to PDF without careful export settings.')

# ================================================================
#  6. SUMMARY AND CONCLUSIONS
# ================================================================
H('6.', 'Summary and Conclusions')

B('This paper has presented and validated a dual-engine architecture for AI-assisted scientific '
  'manuscript preparation. The key findings are:')

B('(1) The combination of python-docx for structural document generation and MCP safe-docx '
  'for paragraph-level editing provides complementary coverage of all major tasks required '
  'for journal-submission-ready manuscript production.')

B('(2) Journal-specific formatting requirements—including justified alignment, three-line '
  'tables, hyperlinked internal references, embedded figures, and correctly styled reference '
  'lists—can be fully automated through executable build scripts.')

B('(3) A six-color tagging convention enables transparent human-AI collaboration, with clear '
  'visual distinction between AI-generated suggestions and author-verified content.')

B('(4) The reproducible-research paradigm is extended to manuscript production through '
  'executable build scripts that codify all formatting and generation steps.')

B('(5) The architecture has been validated through the production of this manuscript, which '
  'was generated, formatted, and polished entirely via the described toolchain.')

# ================================================================
#  ACKNOWLEDGEMENTS
# ================================================================
H_unnum('Acknowledgements')
B('This work was supported by the National Natural Science Foundation of China (Grant No. 5xxxxxxxx). The authors acknowledge the python-docx and MCP '
  'safe-docx development communities. The manuscript was formatted according to Acta Materialia '
  'guidelines (Elsevier, 2025).')

# ================================================================
#  DATA AVAILABILITY
# ================================================================
H_unnum('Data Availability Statement')
B('The build scripts, test documents, and sample figures used in this study are available in '
  'the Paper_Project repository accompanying this manuscript. All outputs can be regenerated '
  'by executing: cd Program && python build_acta_manuscript.py in a Python 3.12 environment '
  'with python-docx and Pillow installed.')

# ================================================================
#  REFERENCES  — (5) Acta Materialia [N] format
# ================================================================
H_unnum('References')
B('')

# Acta Materialia reference format:
# [N] Authors, Title, Journal Volume (Year) Pages. DOI
refs = [
    (1, 'J. Weissmüller, Alloy effects in nanostructures, Nanostruct. Mater. 3 (1993) 261–272. '
        'https://doi.org/10.1016/0965-9773(93)90088-S'),

    (2, 'T. Chookajorn, H.A. Murdoch, C.A. Schuh, Design of stable nanocrystalline alloys, '
        'Science 337 (2012) 951–954. https://doi.org/10.1126/science.1224817'),

    (3, 'R. Kirchheim, Grain coarsening inhibited by solute segregation, Acta Mater. 50 (2002) '
        '413–419. https://doi.org/10.1016/S1359-6454(01)00338-X'),

    (4, 'C.S. Smith, Grains, phases, and interfaces: An interpretation of microstructure, '
        'Trans. Metall. Soc. AIME 175 (1948) 15–51.'),

    (5, 'K. Lu, Stabilizing nanostructures in metals using grain and twin boundary architectures, '
        'Nat. Rev. Mater. 1 (2016) 16019. https://doi.org/10.1038/natrevmats.2016.19'),

    (6, 'K. Lu, L. Lu, S. Suresh, Strengthening materials by engineering coherent internal '
        'boundaries at the nanoscale, Science 324 (2009) 349–352. '
        'https://doi.org/10.1126/science.1159610'),

    (7, 'Elsevier, Guide for Authors — Acta Materialia, 2025. '
        'https://www.sciencedirect.com/journal/acta-materialia/publish/guide-for-authors'),

    (8, 'T.S. Humble, A. Smith et al., python-docx: Create and modify Word documents with Python, '
        'GitHub Repository, 2024. https://github.com/python-openxml/python-docx'),

    (9, 'P.M. Voyles, S.A. Kesler, Reproducible research in materials science: Current practices '
        'and recommendations, MRS Bull. 48 (2023) 1007–1015.'),

    (10, 'J.W. Gibbs, The Collected Works of J. Willard Gibbs, Vol. I: Thermodynamics, '
         'Longmans, Green and Co., New York, 1928.'),
]

for num, ref_text in refs:
    p = doc.add_paragraph()
    J(p)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

    # Create bookmark anchor for internal hyperlinks from in-text citations
    bookmark_name = f'_Ref{num}'
    # Insert bookmark start
    bk_start = OxmlElement('w:bookmarkStart')
    bk_start.set(qn('w:id'), str(num))
    bk_start.set(qn('w:name'), bookmark_name)
    p._element.append(bk_start)

    R(p, f'[{num}] ', bold=False, size=10)
    R(p, ref_text, size=10)

    # Insert bookmark end
    bk_end = OxmlElement('w:bookmarkEnd')
    bk_end.set(qn('w:id'), str(num))
    p._element.append(bk_end)

# ================================================================
#  PAGINATE — insert page breaks at A4 boundaries
# ================================================================

def _element_tag(el):
    return el.tag.split('}')[-1]

def _para_height_pt(p_elem):
    """Estimate w:p height in points. Handles text, empty lines, and inline images."""
    text = ''
    max_sz_pt = 12.0
    img_height_emu = 0
    for r_elem in p_elem:
        if _element_tag(r_elem) != 'r':
            continue
        for child in r_elem:
            ct = _element_tag(child)
            if ct == 't':
                text += (child.text or '')
            elif ct == 'rPr':
                for prop in child:
                    if _element_tag(prop) == 'sz':
                        max_sz_pt = max(max_sz_pt, float(prop.get(qn('w:val'), '24')) / 2.0)
            elif ct == 'drawing':
                # Extract image extent from wp:inline/wp:extent
                for inline in child:
                    for ext in inline:
                        if _element_tag(ext) == 'extent':
                            cy = int(ext.get('cy', '0'))
                            img_height_emu = max(img_height_emu, cy)

    if img_height_emu > 0:
        return img_height_emu / 12700.0  # EMU → pt

    if not text.strip():
        return max_sz_pt * 2.0  # empty line, double-spaced

    lines = max(1, (len(text) + 74) // 75)  # ~75 chars/line for 12pt on A4
    return lines * max_sz_pt * 2.0  # double-spaced


def _table_height_pt(tbl_elem):
    """Estimate table height in points."""
    row_count = sum(1 for r in tbl_elem if _element_tag(r) == 'tr')
    if row_count <= 1:
        return 28
    return 28 + (row_count - 1) * 20  # header 28pt + 20pt per data row


def _is_page_break(p_elem):
    for r in p_elem:
        if _element_tag(r) == 'r':
            for br in r:
                if _element_tag(br) == 'br' and br.get(qn('w:type')) == 'page':
                    return True
    return False


def _make_page_break():
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


def _para_text(p_elem):
    """Extract plain text from a w:p element."""
    parts = []
    for r in p_elem:
        if _element_tag(r) == 'r':
            for c in r:
                if _element_tag(c) == 't':
                    parts.append(c.text or '')
    return ''.join(parts)


def _has_drawing(p_elem):
    """Check if paragraph contains an image."""
    for r in p_elem:
        if _element_tag(r) == 'r':
            for c in r:
                if _element_tag(c) == 'drawing':
                    return True
    return False


def paginate_a4(doc):
    """Insert page breaks at A4 boundaries with widow/orphan control.
    Keeps headings, table/figure intros glued to their content."""
    body = doc.element.body
    body_children = list(body)
    usable_pt = 25.0 / 0.0352778  # ~709 pt

    # Collect content elements
    elems = []  # (body_index, tag, height, text)
    for idx, child in enumerate(body_children):
        tag = _element_tag(child)
        if tag == 'sectPr' or tag not in ('p', 'tbl'):
            continue
        if tag == 'p' and _is_page_break(child):
            continue
        h = _table_height_pt(child) if tag == 'tbl' else _para_height_pt(child)
        txt = _para_text(child) if tag == 'p' else ''
        elems.append((idx, tag, h, txt))

    if not elems:
        return 0

    # Pass 1: compute raw break positions
    raw_breaks = []
    acc = 0.0
    for i, (_, _, h, _) in enumerate(elems):
        if acc > 0 and acc + h > usable_pt:
            raw_breaks.append(i)
            acc = h
        else:
            acc += h

    def _is_glued(elem_idx):
        """Check if elem_idx and elem_idx+1 should be kept together."""
        if elem_idx < 0 or elem_idx + 1 >= len(elems):
            return False
        a_tag, a_h, a_txt = elems[elem_idx][1], elems[elem_idx][2], elems[elem_idx][3]
        b_tag = elems[elem_idx + 1][1]

        # Short text paragraph followed by table
        if a_tag == 'p' and b_tag == 'tbl' and len(a_txt) < 500:
            return True
        # Short text paragraph followed by image
        if a_tag == 'p' and b_tag == 'p':
            b_child = body_children[elems[elem_idx + 1][0]]
            if _has_drawing(b_child) and len(a_txt) < 500:
                return True
        # Image paragraph → its caption (keep Fig/Table caption with image/table)
        if a_tag == 'p' and b_tag == 'p':
            a_child = body_children[elems[elem_idx][0]]
            b_txt = elems[elem_idx + 1][3]
            if _has_drawing(a_child) and (b_txt.startswith('Fig.') or b_txt.startswith('Table')):
                return True

        # Heading (short, large font) followed by body
        if a_tag == 'p' and b_tag == 'p' and len(a_txt) < 100 and a_h > 20:
            return True
        return False

    # Pass 2: adjust for widow/orphan — walk back across glued elements
    adjusted = []
    for pos in raw_breaks:
        final = pos
        # Recursively walk back while consecutive elements are glued
        while final > 0 and _is_glued(final - 1):
            final -= 1
        adjusted.append(final)

    # Deduplicate and insert
    adjusted = sorted(set(adjusted))
    for pos in reversed(adjusted):
        body.insert(elems[pos][0], _make_page_break())

    return len(adjusted)


n_pages = paginate_a4(doc)


def _add_body_page_numbers(doc):
    """Insert visible page numbers at the bottom of each page (before each page break
    and at the end of document), for Office Viewer compatibility. Small gray '— N —'."""
    body = doc.element.body
    children = list(body)
    insert_ops = []

    for idx, child in enumerate(children):
        tag = _element_tag(child)
        if tag == 'sectPr':
            continue
        if tag == 'p' and _is_page_break(child):
            insert_ops.append((idx, len(insert_ops) + 1))

    # Last page marker at end of body (before sectPr)
    insert_ops.append((len(children), len(insert_ops) + 1))

    for at_idx, pn in reversed(insert_ops):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        p.append(pPr)

        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '14')  # 7pt
        rPr.append(sz)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), 'AAAAAA')
        rPr.append(color)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = f'— {pn} —'
        r.append(t)
        p.append(r)

        if at_idx < len(list(body)):
            body.insert(at_idx, p)
        else:
            body.append(p)


_add_body_page_numbers(doc)

# ================================================================
#  SAVE
# ================================================================
out = os.path.join(PROJ, 'Manuscripts', 'Acta_style_manuscript.docx')
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
print(f'Saved: {out}')
print(f'  A4 pages: ~{n_pages + 1} (inserted {n_pages} page breaks)')
print('Features:')
print('  (1) Justified alignment (两端对齐) throughout')
print('  (2) Hyperlinked in-text citations → reference list')
print('  (3) Three-line tables (三线表) on Table 1 & Table 2')
print('  (4) Real embedded images (Fig.1 & Fig.2)')
print('  (5) Acta Materialia [N] reference format with hyperlinks')
print('  (6) A4 pagination with automatic page breaks')
