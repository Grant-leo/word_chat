"""
build_comprehensive_doc.py — 综合功能演示脚本
演示 python-docx + Pillow 的全部排版能力：
  字体/字号/颜色/格式/超链接/参考文献/代码块/项目符号/图片插入/程序化绘图/三线表
输出: ../Manuscripts/示例_综合功能演示.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image, ImageDraw, ImageFont
import os, random, math

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.normpath(os.path.join(BASE, '..', 'Manuscripts', '示例_综合功能演示.docx'))
FIG_DIR = os.path.normpath(os.path.join(BASE, '..', 'Results', 'fig'))
os.makedirs(os.path.dirname(OUT), exist_ok=True)
os.makedirs(FIG_DIR, exist_ok=True)

doc = Document()

# A4 page setup
for sec in doc.sections:
    sec.page_width   = Cm(21.0)
    sec.page_height  = Cm(29.7)
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin   = Cm(2.54)
    sec.right_margin  = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ============================================================
# 工具函数
# ============================================================

def add_hyperlink(paragraph, url, text):
    """插入可点击的外部超链接（蓝色下划线）"""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = paragraph._element.makeelement(qn('w:hyperlink'), {qn('r:id'): r_id})
    nr = paragraph._element.makeelement(qn('w:r'), {})
    rpr = paragraph._element.makeelement(qn('w:rPr'), {})
    c = paragraph._element.makeelement(qn('w:color'), {qn('w:val'): '0000FF'})
    u = paragraph._element.makeelement(qn('w:u'), {qn('w:val'): 'single'})
    rpr.append(c); rpr.append(u); nr.append(rpr)
    t = paragraph._element.makeelement(qn('w:t'), {})
    t.text = text; nr.append(t)
    hl.append(nr); paragraph._element.append(hl)

def H(text, size=16, font_name='黑体', color=RGBColor(0,51,153)):
    """一级标题：加粗居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.name = font_name; r.font.color.rgb = color
    return p

def H2(text, size=13, font_name='微软雅黑', color=None):
    """二级标题：加粗左对齐"""
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.name = font_name
    if color: r.font.color.rgb = color
    return p

def body(text, font_name='宋体', size=11, color=None):
    """正文段落"""
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = font_name
    if color: r.font.color.rgb = color
    return p

def bullet(text, font_name='宋体', size=11):
    """项目符号段落"""
    p = doc.add_paragraph()
    r = p.add_run('  * ' + text); r.font.size = Pt(size); r.font.name = font_name
    return p

def code_block(text):
    """代码块样式：Consolas 9pt 灰色"""
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(9); r.font.name = 'Consolas'; r.font.color.rgb = RGBColor(80,80,80)
    return p

def caption(text):
    """图注：居中、斜体、9pt 灰色"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(9); r.font.name = '宋体'; r.italic = True; r.font.color.rgb = RGBColor(100,100,100)
    return p

def sep():
    """分隔线"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('—' * 40); r.font.size = Pt(6); r.font.color.rgb = RGBColor(180,180,180)

def sup_ref(text, ref_id):
    """正文 + 蓝色上标引用标记"""
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(11); r.font.name = '宋体'
    r2 = p.add_run(ref_id); r2.font.superscript = True; r2.font.color.rgb = RGBColor(0,0,255); r2.font.size = Pt(8)
    return p

# ============================================================
# 三线表（来自 Acta Materialia 标准实现）
# ============================================================

def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)

def _set_tc_border(cell, top=None, bottom=None, left='nil', right='nil'):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcB = OxmlElement('w:tcBorders')
    tcPr.append(tcB)
    def add_border(pos, val='nil', sz='0'):
        b = OxmlElement(f'w:{pos}')
        b.set(qn('w:val'), val)
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tcB.append(b)
    if top is not None:
        add_border('top', top[0], top[1])
    else:
        add_border('top', 'nil', '0')
    if bottom is not None:
        add_border('bottom', bottom[0], bottom[1])
    else:
        add_border('bottom', 'nil', '0')
    add_border('left', left, '0')
    add_border('right', right, '0')

def three_line_table(table):
    """将表格转为三线表：顶粗线 / 表头下细线 / 底粗线 / 无竖线"""
    _remove_table_borders(table)
    nrows = len(table.rows)
    if nrows == 0:
        return
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            if ri == 0 and nrows == 1:
                _set_tc_border(cell, top=('single','12'), bottom=('single','12'))
            elif ri == 0:
                _set_tc_border(cell, top=('single','12'), bottom=('single','4'))
            elif ri == 1 and ri == nrows - 1:
                _set_tc_border(cell, top=('single','4'), bottom=('single','12'))
            elif ri == 1:
                _set_tc_border(cell, top=('single','4'))
            elif ri == nrows - 1:
                _set_tc_border(cell, bottom=('single','12'))
            else:
                _set_tc_border(cell)

def set_cell(cell, text, bold=False, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格文本"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = '宋体'

# ============================================================
# Pillow 程序化生成图片
# ============================================================

def generate_gb_schematic():
    """生成纳米晶晶界稳定化示意图"""
    path = os.path.join(FIG_DIR, 'demo_grain_boundary.png')
    if os.path.exists(path):
        return path
    random.seed(42)
    img = Image.new('RGB', (800, 500), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.rectangle([(0, 0), (800, 40)], fill=(41, 65, 122))
    draw.text((250, 6), 'Grain Boundary Stabilization Schematic', fill=(255, 255, 255))
    draw.text((320, 26), '晶界稳定化示意图', fill=(255, 255, 220))
    for _ in range(20):
        cx, cy = random.randint(80, 720), random.randint(120, 460)
        color = random.choice([(255,230,200), (200,230,255), (230,255,200), (255,220,220), (240,240,200)])
        n_vertices = random.randint(5, 8)
        radius = random.randint(30, 70)
        points = []
        for v in range(n_vertices):
            angle = 2 * math.pi * v / n_vertices + random.uniform(-0.2, 0.2)
            r = radius * random.uniform(0.7, 1.3)
            points.append((cx + int(r * math.cos(angle)), cy + int(r * math.sin(angle))))
        draw.polygon(points, fill=color, outline=(80, 80, 80), width=2)
    img.save(path, 'PNG')
    return path

def generate_hall_petch_chart():
    """生成霍尔-佩奇关系示意图"""
    path = os.path.join(FIG_DIR, 'demo_hall_petch.png')
    if os.path.exists(path):
        return path
    chart = Image.new('RGB', (600, 350), color=(255, 255, 255))
    d = ImageDraw.Draw(chart)
    d.line([(80, 300), (580, 300)], fill=(0, 0, 0), width=2)
    d.line([(80, 300), (80, 30)], fill=(0, 0, 0), width=2)
    d.text((260, 320), 'Grain Size D (nm)', fill=(0, 0, 0))
    d.rectangle([(80, 20), (160, 280)], fill=(255, 200, 200), outline=(200, 0, 0))
    d.text((85, 170), 'HP Breakdown', fill=(200, 0, 0))
    bars = [
        (120, 200, '10', (0, 100, 200)),
        (220, 170, '50', (50, 120, 200)),
        (320, 140, '100', (100, 140, 200)),
        (420, 110, '500', (150, 160, 200)),
        (520, 80, 'CG', (200, 180, 200)),
    ]
    for x, h, lab, col in bars:
        d.rectangle([(x-28, 300-h), (x+28, 300)], fill=col, outline=(0, 0, 0))
        d.text((x-15, 310), lab, fill=(0, 0, 0))
    chart.save(path, 'PNG')
    return path

# ============================================================
# 封面
# ============================================================

cov = doc.add_paragraph(); cov.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cov.add_run('Word 文档处理能力测试'); r.bold = True
r.font.size = Pt(22); r.font.name = '宋体'; r.font.color.rgb = RGBColor(0,51,153)

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('python-docx + Pillow + MCP safe-docx  全功能演示'); r.font.size = Pt(13)
r.font.name = '微软雅黑'; r.font.color.rgb = RGBColor(102,102,102)

ver = doc.add_paragraph(); ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ver.add_run('本文件由 build_comprehensive_doc.py 程序化生成'); r.font.size = Pt(10)
r.font.name = '宋体'; r.font.color.rgb = RGBColor(150,150,150)

doc.add_paragraph()

# ============================================================
# 目录
# ============================================================

toc_items = [
    '一、字体与字号控制',
    '二、文字颜色与标记体系',
    '三、粗体 / 斜体 / 下划线格式组合',
    '四、代码块样式',
    '五、项目符号列表',
    '六、超链接插入',
    '七、参考文献：上标引用',
    '八、图片插入与尺寸控制',
    '九、程序化生成数据图',
    '十、综合能力矩阵（三线表）',
]
for item in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(item); r.font.size = Pt(12); r.font.name = '宋体'

doc.add_paragraph()
sep()

# ============================================================
# 一、字体与字号
# ============================================================

H('一、字体与字号控制', 16)
body('以下演示 python-docx 对不同字体族和字号的精确控制：', '宋体', 11)

font_samples = [
    (8,   '宋体',             False, False, False, '8pt  宋体 -- 极小字（脚注 / 页眉用）'),
    (9,   'Times New Roman',  False, True,  False, '9pt  Times New Roman Italic -- 图注 / 表格注文'),
    (10.5,'宋体',             False, False, False, '10.5pt 宋体 -- 中文公文正文标准'),
    (12,  'Times New Roman',  False, False, False, '12pt Times New Roman -- 英文正文标准'),
    (14,  '微软雅黑',         True,  False, False, '14pt 微软雅黑 Bold -- 段落标题候选'),
    (16,  '黑体',             True,  False, False, '16pt 黑体 Bold -- 二级标题常用'),
    (18,  '宋体',             True,  False, False, '18pt 宋体 Bold -- 一级标题'),
    (22,  'Times New Roman',  True,  False, False, '22pt Times New Roman Bold -- 文档大标题'),
]
for size, font_name, bold, italic, underline, text in font_samples:
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = font_name
    r.bold = bold; r.italic = italic; r.underline = underline

sep()

# ============================================================
# 二、颜色标记体系
# ============================================================

H('二、文字颜色与标记体系', 16)

color_samples = [
    (RGBColor(238,   0,   0), 'EE0000 红色 -- 重要结论 / 原文核心标记'),
    (RGBColor(  0,   0, 255), '0000FF 蓝色 -- AI 补充建议 / 待审核内容'),
    (RGBColor(  0, 128,   0), '008000 绿色 -- 已验证 / 已确认结论'),
    (RGBColor(255, 128,   0), 'FF8000 橙色 -- 待确认 / 存疑 / 需进一步验证'),
    (RGBColor(128,   0, 128), '800080 紫色 -- 背景补充 / 前人工作总结'),
    (RGBColor(128, 128, 128), '808080 灰色 -- 注释 / 辅助信息'),
]
for color, text in color_samples:
    p = doc.add_paragraph()
    # 色块
    block = p.add_run('  ')
    block.font.color.rgb = color; block.font.size = Pt(14)
    # 说明文字
    r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(12); r.font.name = '宋体'

body('颜色标记策略：红 = 核心新知 / 蓝 = 补充待审 / 绿 = 已确认 / 橙 = 存疑 / 紫 = 背景 / 灰 = 注释。', '宋体', 10)

sep()

# ============================================================
# 三、格式组合
# ============================================================

H('三、粗体 / 斜体 / 下划线格式组合', 16)

style_samples = [
    (True, False, False, '粗体 Bold -- 强调关键词、标题、重要结论'),
    (False, True, False, '斜体 Italic -- 外文术语、变量符号、书名'),
    (False, False, True, '下划线 Underline -- 次要标记、链接指示'),
    (True, True, False, '粗体+斜体 Bold+Italic -- 强强调，常用于核心概念'),
    (True, False, True, '粗体+下划线 Bold+Underline -- 警告/特别注意'),
    (True, True, True, '粗体+斜体+下划线 三者组合 -- 极强调（慎用）'),
]
for bold, italic, underline, text in style_samples:
    p = doc.add_paragraph()
    r = p.add_run(f'  {text}'); r.bold = bold; r.italic = italic; r.underline = underline
    r.font.size = Pt(12); r.font.name = 'Times New Roman'

sep()

# ============================================================
# 四、代码块样式 (from tutorial)
# ============================================================

H('四、代码块样式', 16)
body('以下演示 Consolas 等宽字体的代码块渲染效果：', '宋体', 11)

code_block('  def three_line_table(table):')
code_block('      """将表格转为三线表：顶粗线 / 表头下细线 / 底粗线 / 无竖线"""')
code_block('      _remove_table_borders(table)')
code_block('      nrows = len(table.rows)')
code_block('      for ri, row in enumerate(table.rows):')
code_block('          for cell in row.cells:')
code_block('              if ri == 0:')
code_block('                  _set_tc_border(cell, top=("single","12"), bottom=("single","4"))')
code_block('              elif ri == nrows - 1:')
code_block('                  _set_tc_border(cell, bottom=("single","12"))')

body('代码块使用 Consolas 9pt 灰色，适合在文档中嵌入技术说明。', '宋体', 10)

sep()

# ============================================================
# 五、项目符号列表 (from tutorial)
# ============================================================

H('五、项目符号列表', 16)
body('以下演示项目符号列表的各种使用场景：', '宋体', 11)
doc.add_paragraph()

body('python-docx 的核心能力：', '宋体', 11)
bullet('创建新文档、设置页面布局、页边距控制')
bullet('插入多格式图片（PNG / JPEG / TIFF / BMP / GIF）')
bullet('创建表格并精确控制单元格格式（含三线表）')
bullet('字体控制：粗体、斜体、下划线、字号、颜色')
bullet('插入外部超链接（Ctrl+Click 可跳转）')

doc.add_paragraph()
body('MCP safe-docx 的核心能力：', '宋体', 11)
bullet('段落级编辑：插入、删除、替换文本')
bullet('添加 Word 原生脚注（python-docx 做不到）')
bullet('添加批注 / 注释（python-docx 做不到）')
bullet('文档对比（红色修订跟踪）')
bullet('批量查找替换')

doc.add_paragraph()
body('双引擎铁律：python-docx 建框架 → MCP 精修。不要反过来，也别在 MCP 修改后重跑 python-docx。', '宋体', 10)

sep()

# ============================================================
# 六、超链接
# ============================================================

H('六、超链接插入', 16)
body('python-docx 可通过 OOXML 直写插入真实超链接（Ctrl+Click 可跳转）：', '宋体', 11)

links = [
    ('https://doi.org/10.1016/0965-9773(93)90088-S', 'Weissmueller J. Nanostructured Materials, 1993'),
    ('https://doi.org/10.1126/science.1224817', 'Chookajorn T, Schuh C A. Science, 2012'),
    ('https://doi.org/10.1016/j.actamat.2010.10.027', 'Kirchheim R. Acta Materialia, 2002'),
]
for url, text in links:
    p = doc.add_paragraph()
    r = p.add_run('  - '); r.font.size = Pt(11); r.font.name = '宋体'
    add_hyperlink(p, url, text)

body('注：MCP safe-docx 同样支持通过 <a href="..."> 标签插入超链接。', '宋体', 10)

sep()

# ============================================================
# 七、参考文献：上标引用
# ============================================================

H('七、参考文献：上标引用', 16)
body('以下段落使用蓝色上标引用标记，模仿学术论文中的文献引用格式：', '宋体', 11)

p1 = doc.add_paragraph()
r = p1.add_run('Weissmueller 模型'); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
r = p1.add_run(' 建立了晶界能随溶质偏聚变化的热力学框架'); r.font.name = '宋体'; r.font.size = Pt(12)
r = p1.add_run('[1]'); r.font.superscript = True; r.font.color.rgb = RGBColor(0,0,255); r.font.size = Pt(9)

p2 = doc.add_paragraph()
r = p2.add_run('Zener 钉扎'); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
r = p2.add_run(' 是动力学稳定晶粒尺寸的经典机制'); r.font.name = '宋体'; r.font.size = Pt(12)
r = p2.add_run('[2]'); r.font.superscript = True; r.font.color.rgb = RGBColor(0,0,255); r.font.size = Pt(9)

p3 = doc.add_paragraph()
r = p3.add_run('Gibbs 吸附等温式'); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
r = p3.add_run(' d-gamma = -Gamma x d-mu 给出了晶界能变化与溶质过剩量的定量关系'); r.font.name = '宋体'; r.font.size = Pt(12)
r = p3.add_run('[3]'); r.font.superscript = True; r.font.color.rgb = RGBColor(0,0,255); r.font.size = Pt(9)

doc.add_paragraph()
ref_p = doc.add_paragraph()
r = ref_p.add_run('【参考文献列表】'); r.bold = True; r.font.size = Pt(11); r.font.name = '宋体'

refs = [
    '[1] Weissmueller J. Alloy effects in nanostructures. Nanostructured Materials, 1993, 3(1-6): 261-272.',
    '[2] Smith C S. Grains, phases, and interfaces. Trans. Metall. Soc. AIME, 1948, 175: 15-51.',
    '[3] Gibbs J W. The Collected Works of J. Willard Gibbs. Vol. I. Longmans, Green and Co., 1928.',
    '[4] Chookajorn T, Murdoch H A, Schuh C A. Design of stable nanocrystalline alloys. Science, 2012, 337: 951-954.',
    '[5] Kirchheim R. Grain coarsening inhibited by solute segregation. Acta Materialia, 2002, 50(2): 413-419.',
]
for ref in refs:
    p = doc.add_paragraph()
    r = p.add_run(ref); r.font.size = Pt(9); r.font.name = 'Times New Roman'

body('提示：MCP safe-docx 的 add_footnote 工具可添加 Word 原生脚注，更适合正式投稿。python-docx 难以原生实现脚注。', '宋体', 10)

sep()

# ============================================================
# 八、图片插入与尺寸控制
# ============================================================

H('八、图片插入与尺寸控制', 16)
body('以下演示 4 种不同尺寸的图片嵌入——从缩略图到全页宽图：', '宋体', 11)

img_path = generate_gb_schematic()

size_demos = [
    (1.2, '1.2 英寸 -- 缩略图'),
    (2.5, '2.5 英寸 -- 正文插图'),
    (4.5, '4.5 英寸 -- 跨栏大图'),
    (6.0, '6.0 英寸 -- 全页宽图'),
]
for size_inch, label in size_demos:
    p = doc.add_paragraph()
    r = p.add_run(f'{label}:'); r.font.size = Pt(10); r.font.name = '宋体'; r.bold = True
    doc.add_picture(img_path, width=Inches(size_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

caption('图1: 纳米晶晶界稳定化示意图（同图多尺寸缩放演示）')

sep()

# ============================================================
# 九、程序化生成数据图
# ============================================================

H('九、程序化生成数据图', 16)
body('以下图表由 Pillow 程序化生成后嵌入 Word——适合数据驱动的自动报告生成：', '宋体', 11)

chart_path = generate_hall_petch_chart()
doc.add_picture(chart_path, width=Inches(5.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption('图2: 霍尔-佩奇关系示意图（Pillow 程序化绘图 + python-docx 嵌入 Word）')

body('程序化绘图的优势：可通过数据驱动自动生成/更新图表，无需手动截图，适合实验数据自动报告生成流程。', '宋体', 10)

sep()

# ============================================================
# 十、综合能力矩阵（三线表）
# ============================================================

H('十、综合能力矩阵（三线表）', 16)
body('下表使用三线表（顶粗线 / 表头下细线 / 底粗线，无竖线），总结 Win Word 处理完整工具链：', '宋体', 11)

data = [
    ['创建新文档',     'Yes — 完善',       'No',                'python-docx'],
    ['字体族设置',     'Yes — 精准',       '部分（内联标签）',  'python-docx 预设'],
    ['字号控制',       'Yes — 精确到 0.5pt','有限',             'python-docx'],
    ['粗/斜/下划线',   'Yes',              'Yes — 内联标签',   '两者皆可'],
    ['文字颜色',       'Yes — RGB',        'Yes — hex color',   '两者皆可'],
    ['超链接',         'Yes — 原生 OOXML', 'Yes — <a> 标签',   '两者皆可'],
    ['脚注',           'No — 极复杂',      'Yes — 原生 API',   'MCP safe-docx'],
    ['评论/批注',      'No',               'Yes — 原生 API',   'MCP safe-docx'],
    ['图片插入',       'Yes — 多格式',     'No',                'python-docx'],
    ['段落级编辑',     '不方便',           'Yes — 便捷',        'MCP safe-docx'],
    ['表格',           'Yes — 完善',       'Warning — 有限',    'python-docx 建 + MCP 改'],
]

nrows = len(data) + 1
table = doc.add_table(rows=nrows, cols=4)

headers = ['能力项', 'python-docx', 'MCP safe-docx', '推荐方案']
for i, h in enumerate(headers):
    set_cell(table.rows[0].cells[i], h, bold=True, size=10)

for ri, row_data in enumerate(data, 1):
    for ci, cell_text in enumerate(row_data):
        set_cell(table.rows[ri].cells[ci], cell_text, size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)

three_line_table(table)

doc.add_paragraph()
body('三线表实现原理：完全移除 tblBorders → 每个单元格独立设置 tcBorders → 顶行 top=12(粗) bottom=4(细) / 中间行无横线 / 底行 bottom=12(粗) / 全部竖线=nil。', '宋体', 9)

sep()

# ============================================================
# 页脚
# ============================================================

body('本演示文档涵盖的全部功能清单：', '宋体', 11)
bullet('8 种字号 × 4 种字体族')
bullet('6 色标记体系（红色/蓝色/绿色/橙色/紫色/灰色）')
bullet('6 种格式组合（粗体/斜体/下划线的排列组合）')
bullet('代码块样式（Consolas 等宽字体，灰色）')
bullet('项目符号列表')
bullet('3 个可点击 DOI 超链接')
bullet('上标引用标记 [1][2][3] + 5 条参考文献')
bullet('Pillow 程序化生成 2 张图（晶界示意图 + Hall-Petch 图）')
bullet('4 种图片尺寸（1.2 / 2.5 / 4.5 / 6.0 英寸）')
bullet('1 张三线表（顶粗线 / 表头下细线 / 底粗线）')
doc.add_paragraph()

sep()
footer = doc.add_paragraph(); footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run('本测试文档由 build_comprehensive_doc.py 程序化生成，全面覆盖 Word 文档处理能力。')
r.font.size = Pt(9); r.font.name = '宋体'; r.italic = True; r.font.color.rgb = RGBColor(128,128,128)
r = footer.add_run('\n生成方式：python-docx 创建框架/样式/图片/表格 + MCP safe-docx 补充脚注/评论/段落编辑')
r.font.size = Pt(8); r.font.name = '宋体'; r.italic = True; r.font.color.rgb = RGBColor(150,150,150)

# ============================================================
# 保存
# ============================================================

doc.save(OUT)
print(f'Saved: {OUT}')
print('Contains:')
print('  1. Font/Size: 8pt~22pt, 4 font families')
print('  2. Colors: 6-color tagging system with color blocks')
print('  3. Formatting: 6 style combinations')
print('  4. Code block: Consolas 9pt gray')
print('  5. Bullet lists: 2 groups')
print('  6. Hyperlinks: 3 clickable DOI links')
print('  7. References: superscript [1][2][3] + 5-item bibliography')
print('  8. Images: 4 sizes (1.2" to 6.0")')
print('  9. Chart: 1 programmatic Hall-Petch diagram')
print('  10. Table: 12-row x 4-col three-line table (capability matrix)')
